#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""TS-CAP V1.0 - Tính sức chịu tải dọc trục cọc theo TCVN 11823-10:2017 / AASHTO LRFD.

Lịch sử sửa đổi chi tiết: xem CHANGELOG_TS-CAP.md kèm theo bộ mã nguồn (không phát hành cùng bản build).

CẢNH BÁO KỸ THUẬT: trước khi dùng cho thiết kế chính thức cần benchmark với hồ sơ tính tay/Excel gốc
và rà lại lựa chọn hệ số sức kháng theo điều kiện dự án, phương pháp thí nghiệm/xác minh SCT hiện trường.

QA-OCR AUTO 2026-07-06: tích hợp song song pipeline OCR ảnh cũ và pipeline hybrid bảo toàn lớp;
QA-OCR SPT-SAFE 2026-07-06: không dùng đồ thị SPT để tự sinh/nâng SPT mặc định; chỉ lấy số đọc trực tiếp từ bảng SPT.
QA-OCR SPT-NO-GUESS 2026-07-06: không tự điền 100 sau refusal; giữ SPT=0 nếu OCR đọc trực tiếp trong bảng; không nội suy/chèn SPT.
QA-GEO RQD 2026-07-06: đá phong hóa có RQD<20 tự quy về IGM (loại 5); RQD>=20 giữ đá phong hóa (loại 4).
chương trình tự chấm điểm/gộp kết quả để giảm rủi ro mất lớp địa chất thật.
"""

from __future__ import annotations

import csv
import json
import copy
import math
import os
import sys
import subprocess
import platform
import hashlib
import shutil
import traceback
import re
import unicodedata
import tempfile
import uuid
import zlib
import threading
import queue
import contextlib
import io
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

APP_NAME = "TS-CAP V1.0"
APP_TITLE = "Tính sức chịu tải cọc theo đất nền"
APP_SUBTITLE = "TCVN 11823-10:2017 / AASHTO LRFD - Chương Nền móng"
APP_AUTHOR = "Tác Giả: Nguyễn Ngọc Dũng\nPhòng QLTK - Khối XD PPP\nTập đoàn SunGroup"


def _ts_app_base_dir():
    """Base folder for source mode, PyInstaller, and Nuitka standalone."""
    candidates = []
    try:
        if getattr(sys, "frozen", False) or "__compiled__" in globals():
            candidates.append(os.path.dirname(os.path.abspath(sys.executable)))
    except Exception:
        pass
    try:
        candidates.append(os.path.dirname(os.path.abspath(sys.executable)))
    except Exception:
        pass
    try:
        candidates.append(os.path.dirname(os.path.abspath(__file__)))
    except Exception:
        pass
    for path in candidates:
        if path and os.path.isdir(path):
            return path
    return os.getcwd()


def _ts_icon_path(filename):
    base_dir = _ts_app_base_dir()
    candidates = []
    try:
        meipass = getattr(sys, "_MEIPASS", None)
        if meipass:
            candidates.extend([
                os.path.join(meipass, "assets", "icons", filename),
                os.path.join(meipass, filename),
            ])
    except Exception:
        pass
    candidates.extend([
        os.path.join(base_dir, "assets", "icons", filename),
        os.path.join(base_dir, filename),
        os.path.join(os.getcwd(), "assets", "icons", filename),
    ])
    try:
        candidates.append(os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "icons", filename))
    except Exception:
        pass
    for path in candidates:
        if path and os.path.exists(path):
            return path
    return None


def apply_app_icon(window, icon_stem):
    """Apply window icon in source, PyInstaller, and Nuitka builds."""
    try:
        ico_path = _ts_icon_path(f"{icon_stem}.ico")
        if ico_path and os.name == "nt":
            window.iconbitmap(default=ico_path)
    except Exception:
        pass
    try:
        png_path = _ts_icon_path(f"{icon_stem}.png")
        if png_path:
            import tkinter as _tk
            _icon_image = _tk.PhotoImage(file=png_path)
            refs = getattr(window, "_ts_icon_refs", [])
            refs.append(_icon_image)
            window._ts_icon_refs = refs
            window.iconphoto(True, _icon_image)
    except Exception:
        pass



def safe_lift_window(win: Any) -> None:
    """Đưa cửa sổ lên trước trong thời gian ngắn, không giữ topmost vĩnh viễn."""
    try:
        if win is None:
            return
        try:
            win.deiconify()
        except Exception:
            pass
        win.lift()
        try:
            win.focus_force()
        except Exception:
            pass
        try:
            win.attributes("-topmost", True)
            win.after(300, lambda w=win: w.attributes("-topmost", False))
        except Exception:
            pass
    except Exception:
        pass


def safe_release_grabs(root: Any) -> None:
    """Cứu hộ GUI: nhả grab nếu một dialog/progress bị ẩn hoặc exception giữa chừng."""
    try:
        if root is None:
            return
        grabbed = root.grab_current()
        if grabbed is not None:
            grabbed.grab_release()
    except Exception:
        pass


def _no_window_kwargs() -> dict:
    """QA fix O2: ẩn cửa sổ console khi gọi subprocess trên bản build windowed (Windows)."""
    if os.name == "nt":
        return {"creationflags": getattr(subprocess, "CREATE_NO_WINDOW", 0x08000000)}
    return {}


PA_MPA = 0.101
GAMMA_W_KN_M3 = 9.81
# QA fix P4: trọng lượng bản thân là tải CÓ LỢI khi kiểm nhổ -> nhân γDC = 0.9 theo LRFD.
GAMMA_DC_UPLIFT = 0.90

# Hệ số sức kháng mặc định, lấy theo các giá trị đang dùng trong bảng Excel mẫu và Bảng hệ số TCVN/AASHTO.
PHI_BORE_SHAFT_CLAY = 0.45
PHI_BORE_SHAFT_SAND = 0.55
PHI_BORE_SHAFT_IGM = 0.60
PHI_BORE_SHAFT_IGM_SPT_LT100 = 0.70
PHI_BORE_SHAFT_IGM_SPT_GE100 = 0.60
PHI_BORE_TIP_CLAY = 0.40
PHI_BORE_TIP_SAND = 0.50
PHI_BORE_TIP_IGM = PHI_BORE_SHAFT_IGM_SPT_GE100  # legacy/default; mũi IGM thực tế lấy theo N60 như thân: N<100=0.70, N>=100=0.60
PHI_ROCK_SIDE = 0.55
PHI_ROCK_TIP = 0.50
PHI_DRIVEN_SHAFT_CLAY = 0.35
PHI_DRIVEN_SHAFT_SAND = 0.30
PHI_DRIVEN_TIP_CLAY = 0.35
PHI_DRIVEN_TIP_SAND = 0.30

# Hệ số sức kháng nhổ tách riêng, không dùng nhầm hệ số nén.
# Các giá trị mặc định đặt theo hướng bảo thủ để đối chiếu TVQT/TCVN/AASHTO theo từng dự án.
PHI_UPLIFT_BORE_CLAY = 0.35
PHI_UPLIFT_BORE_SAND = 0.45
PHI_UPLIFT_BORE_IGM = 0.45
PHI_UPLIFT_ROCK_SIDE = 0.40
PHI_UPLIFT_DRIVEN_CLAY = 0.25
PHI_UPLIFT_DRIVEN_SAND = 0.25
PHI_EXTREME_UPLIFT_DEFAULT = 0.80

# Cọc đóng đất rời - Meyerhof SPT theo TCVN 11823-10:2017.
# qp = 0.038*N160*(Db/D) MPa, qλ = 0.4*N160 MPa đối với cát, 0.3*N160 MPa đối với cát bột không pha sét.
DRIVEN_SAND_SIDE_DISPLACEMENT_KPA_PER_N = 1.9
DRIVEN_SAND_SIDE_NONDISPLACEMENT_KPA_PER_N = 0.96

# Hệ số tải trọng ma sát âm DD theo TCVN 11823-3, Bảng 4.
# Dùng giá trị lớn nhất khi DD là tác dụng bất lợi trong kiểm toán nén cọc.
GAMMA_DD_TOMLINSON = 1.40
GAMMA_DD_LAMBDA = 1.05
GAMMA_DD_ONEILL_REESE = 1.25

LIMIT_STATE_STRENGTH = "CĐ"
LIMIT_STATE_EXTREME = "ĐB"
LIMIT_STATE_LABELS = {
    LIMIT_STATE_STRENGTH: "TTGHCĐ - Cường độ",
    LIMIT_STATE_EXTREME: "TTGHĐB - Đặc biệt",
}

# Hệ số TTGHĐB trong file TVQT đang lấy 1.0 cho sức kháng bên/mũi đất sét, đất cát và đá.
# Với IGM/cọc đóng, V0.1.4 cũng tách trạng thái riêng và mặc định TTGHĐB = 1.0 để tránh dùng nhầm hệ số TTGHCĐ.
PHI_EXTREME_DEFAULT = 1.00

SOIL_TYPE_LABELS = {
    0: "Không khí/Hang karst",
    1: "Cát/đất rời",
    2: "Sét/đất dính",
    3: "Đá nguyên khối",
    4: "Đá nứt vỡ/phong hóa",
    5: "IGM/đá mềm",
    6: "Cuội sỏi/đất rời",
}

PILE_MODE_CHOICES = [
    "Cọc khoan trong đất",
    "Cọc đóng",
    "Cọc khoan trong đá",
]

# Theme fallback, đồng bộ key với n2d_theme_library.py nếu file đó đặt cùng thư mục.
DEFAULT_THEME = {
    "bg": "#F6FBFF", "panel": "#FFFFFF", "sidebar": "#1D4ED8", "sidebar2": "#93C5FD",
    "accent": "#2563EB", "accent_dark": "#1D4ED8", "text": "#0F172A", "muted": "#475569",
    "button": "#E0F2FE", "button_active": "#BAE6FD", "big_button": "#DBEAFE",
    "tree_head": "#E0F2FE", "tree_row": "#FFFFFF", "tree_alt": "#F8FAFC", "tree_fg": "#0F172A",
    "progress": "#22C55E", "trough": "#DCFCE7", "border": "#93C5FD", "entry_bg": "white", "entry_fg": "black", "note": "#EFF6FF",
}

try:
    from n2d_theme_library import THEME_PRESETS as _N2D_THEME_PRESETS, UI_THEME_LABELS as _N2D_THEME_LABELS, normalize_ui_theme as _normalize_ui_theme
except Exception:
    _N2D_THEME_PRESETS = {"CLEAN_LIGHT": DEFAULT_THEME}
    _N2D_THEME_LABELS = {"CLEAN_LIGHT": "Clean Light / Material - sáng tối giản"}
    def _normalize_ui_theme(value: Any = "") -> str:
        return "CLEAN_LIGHT"


BOREHOLE_OCR_ENGINE_LABELS = {
    "TESSERACT": "Tesseract/OpenCV (mặc định)",
    # QA-OCR v4: RapidOCR đọc chữ nhỏ tốt hơn Tesseract nhưng CHẬM trên CPU khi gọi đại trà
    # theo từng crop. Chỉ nên bật thử nghiệm/cứu ảnh khó; pipeline chính mặc định là Tesseract.
    "RAPID": "RapidOCR (thử nghiệm - chậm, chỉ dùng cứu ảnh khó)",
    "PADDLE": "PaddleOCR/OpenCV (thử nghiệm)",
}

def _normalize_borehole_ocr_engine(value: Any = "") -> str:
    txt = str(value or "").strip().upper()
    if "RAPID" in txt:
        return "RAPID"
    if "PADDLE" in txt:
        return "PADDLE"
    return "TESSERACT"


def _normalize_item_name(value: Any) -> str:
    """Chuẩn hóa tên hạng mục để khớp giữa bảng thông tin riêng và bảng địa chất.

    Mục tiêu là tránh lỗi do dấu tiếng Việt, khoảng trắng kép, dấu gạch, ký tự NBSP...
    Ví dụ: "Trụ T2", "Tru  T2", "TRỤ-T2" đều có cùng key "trut2".
    Không tự map T2 sang Trụ T2 nếu không trùng key để tránh dùng nhầm địa chất.
    """
    text = str(value or "").strip().replace("\u00a0", " ")
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = text.replace("đ", "d").replace("Đ", "D")
    text = text.lower()
    return re.sub(r"[^a-z0-9]+", "", text)


def safe_filename(name: Any, default: str = "HANG_MUC") -> str:
    """Làm sạch tên file Windows/Linux để dùng cho báo cáo, tránh lỗi ký tự cấm."""
    s = str(name or default).strip()
    bad = '<>:"/\\|?*\n\r\t'
    for ch in bad:
        s = s.replace(ch, "_")
    s = "_".join(s.split())
    return (s[:120] or default)


PILE_TYPE_CODE_LABELS = {
    "1": "Cọc khoan nhồi",
    "2": "Cọc đóng",
    "3": "Cọc ép",
    # QA fix P1: thêm lựa chọn tiết diện cọc đóng/ép. Mã 2/3 giữ nguyên nghĩa cũ = cọc VUÔNG đặc.
    "2T": "Cọc đóng tròn",
    "3T": "Cọc ép tròn",
    "2O": "Cọc đóng ống",
    "3O": "Cọc ép ống",
}


def normalize_pile_type_choice(value: Any) -> str:
    """Đọc loại cọc từ mã ngắn hoặc chuỗi cũ.

    Bảng nhập liệu dùng mã để gõ nhanh:
    1 = Cọc khoan nhồi; 2 = Cọc đóng vuông; 3 = Cọc ép vuông;
    2T/3T = Cọc đóng/ép tròn đặc; 2O/3O = Cọc đóng/ép ống (PHC), cột Ds = đường kính trong.
    Vẫn hỗ trợ template cũ ghi nguyên cụm chữ để không làm hỏng dữ liệu cũ.
    """
    raw = str(value or "").strip()
    if not raw:
        return ""
    key = raw.replace(",", ".")
    try:
        f = float(key)
        if abs(f - round(f)) < 1e-9:
            key = str(int(round(f)))
    except Exception:
        pass
    key = key.strip().upper()
    if key in PILE_TYPE_CODE_LABELS:
        return PILE_TYPE_CODE_LABELS[key]
    low = _strip_accents(raw).lower()
    is_ong = (" ong" in f" {low}") or ("phc" in low)
    is_tron = "tron" in low
    if "ep" in low:
        if is_ong:
            return "Cọc ép ống"
        if is_tron:
            return "Cọc ép tròn"
        return "Cọc ép"
    if "dong" in low:
        if is_ong:
            return "Cọc đóng ống"
        if is_tron:
            return "Cọc đóng tròn"
        return "Cọc đóng"
    if "khoan" in low or "nhồi" in raw.lower() or "nhoi" in low:
        return "Cọc khoan nhồi"
    return raw


def pile_type_to_input_code(value: Any) -> str:
    """Đổi loại cọc về mã 1/2/3 để ghi vào bảng nhập liệu khi biết chắc loại cọc."""
    label = normalize_pile_type_choice(value)
    for code, name in PILE_TYPE_CODE_LABELS.items():
        if _normalize_item_name(label) == _normalize_item_name(name):
            return code
    return str(value or "").strip()



# ==========================================
# LICENSE DÙNG CHUNG TS-PILE / TS-COL / TS-CAP
# ==========================================
SECRET_SALT = "Dung_Dev_Security_Key_2026!@#"
API_URL = "https://script.google.com/macros/s/AKfycbyzL6t1TROfT5LqTkZ_D0LkMZ4K50aKNBzfHg3-_XIYddToVyjGkBc_WfU7gaa0wuzMbA/exec"

LICENSE_RUNTIME_INFO = {
    "machine_id": "",
    "is_active": False,
    "days_left": None,
    "checked_at": "",
    "message": "",
    "reason": "",
}

def _format_days_left(days_left: Any) -> str:
    try:
        d = int(float(days_left))
        if d < 0:
            return "0 ngày"
        return f"{d} ngày"
    except Exception:
        return "Không xác định"

# ==========================================
# QA fix R1/R3: cache kích hoạt để chạy offline có thời hạn + ổn định machine-id
# ==========================================
LICENSE_OFFLINE_GRACE_DAYS = 14

def _n2d_license_dir() -> str:
    base = os.path.join(os.path.expanduser("~"), ".n2d_license")
    try:
        os.makedirs(base, exist_ok=True)
    except Exception:
        pass
    return base

def _license_cache_path() -> str:
    return os.path.join(_n2d_license_dir(), "ts_suite_license.json")

def _machine_id_store_path() -> str:
    return os.path.join(_n2d_license_dir(), "machine_id.txt")

def _offline_trial_path():
    return os.path.join(_n2d_license_dir(), "ts_suite_offline_trial.json")

def _offline_trial_sig(machine_id, day):
    return hashlib.sha256(f"{machine_id}|{day}|OFFLINE_TRIAL|{SECRET_SALT}".encode("utf-8")).hexdigest()

def _offline_trial_save(machine_id, day=None):
    try:
        if day is None:
            day = datetime.now().strftime("%Y-%m-%d")
        data = {"machine_id": str(machine_id), "start_day": str(day),
                "sig": _offline_trial_sig(machine_id, day)}
        with open(_offline_trial_path(), "w", encoding="utf-8") as fh:
            json.dump(data, fh)
        return True
    except Exception:
        return False

def _offline_trial_load(machine_id):
    try:
        with open(_offline_trial_path(), "r", encoding="utf-8") as fh:
            data = json.load(fh)
        if str(data.get("machine_id", "")) != str(machine_id):
            return None
        day = str(data.get("start_day", ""))
        if str(data.get("sig", "")) != _offline_trial_sig(machine_id, day):
            return None
        return day
    except Exception:
        return None

def _offline_trial_check(machine_id):
    """Trả (được chạy?, ngày offline còn lại). Nếu chưa có file trial thì tạo mới."""
    day = _offline_trial_load(machine_id)
    if not day:
        today = datetime.now().strftime("%Y-%m-%d")
        if _offline_trial_save(machine_id, today):
            return True, LICENSE_OFFLINE_GRACE_DAYS
        return False, 0
    try:
        start = datetime.strptime(day, "%Y-%m-%d")
        delta = (datetime.now() - start).days
        if delta < 0:
            return False, 0
        remain = max(LICENSE_OFFLINE_GRACE_DAYS - delta, 0)
        return remain > 0, remain
    except Exception:
        return False, 0

def _license_cache_sig(machine_id: str, day: str, days_left: Any) -> str:
    return hashlib.sha256(f"{machine_id}|{day}|{days_left}|{SECRET_SALT}".encode("utf-8")).hexdigest()

def _license_cache_save(machine_id: str, days_left: Any) -> None:
    try:
        day = datetime.now().strftime("%Y-%m-%d")
        data = {"machine_id": str(machine_id), "last_ok": day, "days_left": days_left,
                "sig": _license_cache_sig(machine_id, day, days_left)}
        with open(_license_cache_path(), "w", encoding="utf-8") as fh:
            json.dump(data, fh)
    except Exception:
        pass

def _license_cache_clear() -> None:
    try:
        os.remove(_license_cache_path())
    except Exception:
        pass

def _license_offline_check(machine_id):
    """Trả (được chạy offline?, ngày bản quyền còn lại, ngày offline còn lại).

    14 ngày chỉ là giới hạn chạy offline. Giá trị trả về cho GUI/app vẫn là
    số ngày bản quyền ước tính còn lại từ lần server xác nhận gần nhất.
    """
    try:
        with open(_license_cache_path(), "r", encoding="utf-8") as fh:
            data = json.load(fh)
        if str(data.get("machine_id", "")) != str(machine_id):
            return False, 0, 0
        day = str(data.get("last_ok", ""))
        if str(data.get("sig", "")) != _license_cache_sig(machine_id, day, data.get("days_left")):
            return False, 0, 0
        last = datetime.strptime(day, "%Y-%m-%d")
        delta = (datetime.now() - last).days
        if delta < 0:
            return False, 0, 0
        cached_days = int(float(data.get("days_left")))
        license_remaining = max(cached_days - delta, 0)
        offline_remaining = max(LICENSE_OFFLINE_GRACE_DAYS - delta, 0)
        ok = (cached_days - delta) > 0 and offline_remaining > 0
        return ok, license_remaining, offline_remaining
    except Exception:
        return False, 0, 0

def get_machine_id() -> str:
    'Lấy mã máy theo đúng cách TS-PILE/TS-COL để dùng chung license.'
    hwid = ""
    errors = []
    if platform.system() == "Windows":
        for cmd in (
            'powershell "(Get-CimInstance -Class Win32_ComputerSystemProduct).UUID"',
            'wmic csproduct get uuid',
        ):
            try:
                out = subprocess.check_output(cmd, shell=True, stderr=subprocess.DEVNULL, **_no_window_kwargs()).decode(errors="ignore").strip()
                if "wmic" in cmd.lower():
                    parts = [line.strip() for line in out.splitlines() if line.strip() and "uuid" not in line.lower()]
                    out = parts[0] if parts else ""
                if out:
                    hwid = out
                    break
            except Exception as exc:
                errors.append(str(exc))
    if hwid:
        mid = hashlib.sha256((str(hwid) + SECRET_SALT).encode()).hexdigest()
        try:
            with open(_machine_id_store_path(), "w", encoding="utf-8") as fh:
                fh.write(mid)
        except Exception:
            pass
        return mid
    # QA fix R3: dùng machine-id đã lưu trước khi rơi về fallback MAC, tránh lệch máy đã kích hoạt.
    try:
        with open(_machine_id_store_path(), "r", encoding="utf-8") as fh:
            stored = fh.read().strip()
        if stored:
            return stored
    except Exception:
        pass
    node = platform.node() or "unknown-node"
    mac = uuid.getnode()
    hwid = f"{platform.system()}|{node}|{mac}|{platform.machine()}"
    return hashlib.sha256((str(hwid) + SECRET_SALT).encode()).hexdigest()

def check_server_trial() -> Tuple[bool, Any]:
    machine_id = get_machine_id()
    LICENSE_RUNTIME_INFO["machine_id"] = machine_id
    LICENSE_RUNTIME_INFO["checked_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        import requests
        response = requests.post(API_URL, json={"machine_id": machine_id}, timeout=8)
        data = response.json()
        status = str(data.get("status", "")).strip().lower()
        if status == "active":
            days_left = data.get("days_left")
            _license_cache_save(machine_id, days_left)
            LICENSE_RUNTIME_INFO.update({
                "is_active": True,
                "days_left": days_left,
                "reason": "active",
                "message": str(data.get("message", "License active") or "License active"),
            })
            return True, days_left
        inactive_statuses = {"expired", "inactive", "revoked", "not_found", "blocked", "disabled", "not_active", "not-active", "unregistered"}
        if status in inactive_statuses:
            _license_cache_clear()
            LICENSE_RUNTIME_INFO.update({
                "is_active": False,
                "days_left": 0,
                "reason": "expired",
                "message": str(data.get("message", "Hết hạn hoặc chưa kích hoạt") or "Hết hạn hoặc chưa kích hoạt"),
            })
            return False, 0
        raise RuntimeError(f"Phản hồi license không hợp lệ: status={status or 'missing'}")
    except Exception:
        ok_offline, license_days, offline_days = _license_offline_check(machine_id)
        if ok_offline:
            LICENSE_RUNTIME_INFO.update({
                "is_active": True,
                "days_left": license_days,
                "reason": "offline",
                "message": f"Không kết nối được máy chủ bản quyền; đang chạy chế độ offline (còn {offline_days} ngày offline).",
            })
            return True, license_days
        ok_trial, trial_days = _offline_trial_check(machine_id)
        if ok_trial:
            LICENSE_RUNTIME_INFO.update({
                "is_active": True,
                "days_left": trial_days,
                "reason": "offline_trial",
                "message": f"Không kiểm tra được online; đang chạy offline lần đầu (còn {trial_days} ngày).",
            })
            return True, trial_days
        LICENSE_RUNTIME_INFO.update({
            "is_active": False,
            "days_left": 0,
            "reason": "offline_trial_expired",
            "message": "Không kiểm tra được bản quyền online và đã quá 14 ngày dùng offline kể từ lần mở đầu tiên. Vui lòng kết nối mạng rồi mở lại phần mềm.",
        })
        return False, 0

def _notify_license_blocked() -> None:
    """Hiện thông báo chặn theo đúng lý do (hết hạn vs lỗi mạng) trong LICENSE_RUNTIME_INFO."""
    if str(LICENSE_RUNTIME_INFO.get("reason", "")) in ("network", "offline_trial_expired"):
        title = "Không kiểm tra được bản quyền"
        body = (LICENSE_RUNTIME_INFO.get("message") or
                "Không kết nối được máy chủ bản quyền. Vui lòng kết nối mạng rồi mở lại phần mềm.")
    else:
        title = "Hết hạn"
        body = "Hết thời gian sử dụng hoặc chưa kích hoạt bản quyền."
    try:
        root_hidden = tk.Tk()
        root_hidden.withdraw()
        messagebox.showerror(title, body)
        root_hidden.destroy()
    except Exception:
        pass


def security_check_or_exit() -> Any:
    'Chặn TS-CAP nếu license chung không còn hợp lệ.'
    ok, days_left = check_server_trial()
    if not ok:
        _notify_license_blocked()
        sys.exit()
    return days_left

def _strip_accents(value: Any) -> str:
    text = str(value or "")
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return text.replace("đ", "d").replace("Đ", "D")

def _display_item_name(value: Any) -> str:
    return str(value or "").strip().replace("\u00a0", " ")


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None:
            return default
        s = str(value).strip().replace("\u00a0", "").replace(",", ".")
        if not s:
            return default
        return float(s)
    except Exception:
        return default


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(round(_safe_float(value, default)))
    except Exception:
        return default


def _fmt(x: Any, nd: int = 3) -> str:
    try:
        if x is None:
            return ""
        v = float(x)
        if abs(v) < 1e-12:
            v = 0.0
        if int(nd) <= 0:
            return f"{v:.0f}"
        return f"{v:.{nd}f}".rstrip("0").rstrip(".")
    except Exception:
        return str(x or "")


def linear_interp(x: float, xs: List[float], ys: List[float], clamp: bool = True) -> float:
    if not xs or not ys or len(xs) != len(ys):
        return 0.0
    pairs = sorted((float(a), float(b)) for a, b in zip(xs, ys))
    xs2 = [p[0] for p in pairs]
    ys2 = [p[1] for p in pairs]
    x = float(x)
    if x <= xs2[0]:
        if clamp or len(xs2) == 1:
            return ys2[0]
    if x >= xs2[-1]:
        if clamp or len(xs2) == 1:
            return ys2[-1]
    for i in range(len(xs2) - 1):
        if xs2[i] <= x <= xs2[i + 1]:
            den = xs2[i + 1] - xs2[i]
            if abs(den) < 1e-12:
                return ys2[i]
            t = (x - xs2[i]) / den
            return ys2[i] + t * (ys2[i + 1] - ys2[i])
    # extrapolate if requested
    if x < xs2[0]:
        den = xs2[1] - xs2[0]
        return ys2[0] + (x - xs2[0]) * (ys2[1] - ys2[0]) / den if abs(den) > 1e-12 else ys2[0]
    den = xs2[-1] - xs2[-2]
    return ys2[-1] + (x - xs2[-1]) * (ys2[-1] - ys2[-2]) / den if abs(den) > 1e-12 else ys2[-1]


def bilinear_interp(x: float, y: float, xs: List[float], ys: List[float], grid: List[List[float]]) -> float:
    """Nội suy song tuyến tính cho các bảng dạng Excel NoisuyMatran."""
    if not xs or not ys or not grid:
        return 0.0
    xs = [float(v) for v in xs]
    ys = [float(v) for v in ys]
    x = max(min(float(x), max(xs)), min(xs))
    y = max(min(float(y), max(ys)), min(ys))

    def bracket(v, arr):
        arr2 = sorted(arr)
        if v <= arr2[0]:
            return arr.index(arr2[0]), arr.index(arr2[0])
        if v >= arr2[-1]:
            return arr.index(arr2[-1]), arr.index(arr2[-1])
        for i in range(len(arr2) - 1):
            if arr2[i] <= v <= arr2[i + 1]:
                return arr.index(arr2[i]), arr.index(arr2[i + 1])
        return 0, 0

    ix0, ix1 = bracket(x, xs)
    iy0, iy1 = bracket(y, ys)
    x0, x1 = xs[ix0], xs[ix1]
    y0, y1 = ys[iy0], ys[iy1]
    q11 = float(grid[iy0][ix0])
    q21 = float(grid[iy0][ix1])
    q12 = float(grid[iy1][ix0])
    q22 = float(grid[iy1][ix1])
    if ix0 == ix1 and iy0 == iy1:
        return q11
    if ix0 == ix1:
        return linear_interp(y, [y0, y1], [q11, q12])
    if iy0 == iy1:
        return linear_interp(x, [x0, x1], [q11, q21])
    tx = (x - x0) / max(x1 - x0, 1e-12)
    ty = (y - y0) / max(y1 - y0, 1e-12)
    return (q11 * (1 - tx) * (1 - ty) + q21 * tx * (1 - ty) + q12 * (1 - tx) * ty + q22 * tx * ty)


def alpha_clay_bored(su_mpa: float) -> float:
    """Hệ số alpha cho cọc khoan trong đất dính theo Excel TVQT/TCVN.

    TVQT dùng:
    - alpha = 0.55 khi Su/pa <= 1.5
    - alpha = 0.55 - 0.1*(Su/pa - 1.5) khi 1.5 < Su/pa <= 2.5
    - ngoài miền này bảng Excel trả "unknown"; trong tool vẫn kẹp alpha=0.45
      và sinh cảnh báo ở bước tính lớp để không làm gián đoạn batch.
    """
    su = max(float(su_mpa), 0.0)
    ratio = su / PA_MPA if PA_MPA > 0 else 0.0
    if ratio <= 1.5:
        return 0.55
    return max(0.45, 0.55 - 0.1 * (ratio - 1.5))


def clay_alpha_domain_warning(layer_name: str, su_mpa: float) -> str:
    """Trả cảnh báo nếu Su/pa vượt miền bảng TVQT."""
    ratio = max(float(su_mpa), 0.0) / PA_MPA if PA_MPA > 0 else 0.0
    if ratio > 2.5:
        return f"Lớp {layer_name}: Su/pa={ratio:.2f} > 2.5; ngoài miền bảng α, lấy α=0.45."
    return ""


# Bảng α Tomlinson cho cọc đóng/ép trong đất dính theo quy trình TCVN/AASHTO.
# Hàng là Su (MPa), cột là Db/D. Dữ liệu nhập từ bảng quy trình người dùng cung cấp.
_TOMLINSON_DB_OVER_D = [0.0, 10.0, 20.0, 40.0, 100.0]
_TOMLINSON_SU_MPA = [0.0, 0.025, 0.05, 0.075, 0.10, 0.125, 0.15, 0.175, 0.20, 1.0]
_TOMLINSON_ALPHA_TABLES = {
    "sand_gravel": [
        [1.00, 1.00, 1.00, 1.00, 1.00],
        [1.00, 1.00, 1.00, 1.00, 1.00],
        [1.00, 1.00, 1.00, 0.95, 0.95],
        [1.00, 1.00, 1.00, 0.84, 0.84],
        [1.00, 1.00, 0.94, 0.67, 0.67],
        [1.00, 1.00, 0.82, 0.50, 0.50],
        [1.00, 1.00, 0.74, 0.40, 0.40],
        [1.00, 1.00, 0.74, 0.40, 0.40],
        [1.00, 1.00, 0.74, 0.40, 0.40],
        [1.00, 1.00, 0.74, 0.40, 0.40],
    ],
    "soft_clay": [
        [0.55, 0.55, 0.85, 0.85, 0.85],
        [0.55, 0.55, 0.85, 0.85, 0.85],
        [0.40, 0.40, 0.76, 0.76, 0.76],
        [0.34, 0.34, 0.71, 0.71, 0.71],
        [0.29, 0.29, 0.69, 0.69, 0.69],
        [0.26, 0.26, 0.66, 0.66, 0.66],
        [0.25, 0.25, 0.64, 0.64, 0.64],
        [0.23, 0.23, 0.61, 0.61, 0.61],
        [0.20, 0.20, 0.57, 0.57, 0.57],
        [0.15, 0.15, 0.50, 0.50, 0.50],
    ],
    "stiff_clay": [
        [1.00, 1.00, 1.00, 1.00, 1.00],
        [1.00, 1.00, 1.00, 1.00, 1.00],
        [0.94, 0.94, 0.97, 1.00, 1.00],
        [0.85, 0.85, 0.925, 1.00, 1.00],
        [0.66, 0.66, 0.77, 0.88, 0.88],
        [0.42, 0.42, 0.555, 0.69, 0.69],
        [0.29, 0.29, 0.375, 0.46, 0.46],
        [0.24, 0.24, 0.295, 0.35, 0.35],
        [0.20, 0.20, 0.25, 0.30, 0.30],
        [0.20, 0.20, 0.24, 0.28, 0.28],
    ],
}
_TOMLINSON_CASE_LABELS = {
    "sand_gravel": "lớp trên là cát/cát cuội",
    "soft_clay": "lớp trên là sét mềm",
    "stiff_clay": "lớp trên là sét nửa cứng đến cứng",
    "clay_worst": "lớp trên là sét; chọn bảng bất lợi theo SPT/Su",
}


def alpha_clay_driven_tomlinson(su_mpa: float, db_over_d: float = 0.0, cover_case: str = "stiff_clay") -> float:
    """Hệ số α Tomlinson cho cọc đóng/ép trong đất dính."""
    case = str(cover_case or "stiff_clay").strip().lower()
    su = max(float(su_mpa or 0.0), 0.0)
    x = max(float(db_over_d or 0.0), 0.0)
    if case == "clay_worst":
        a_soft = bilinear_interp(x, su, _TOMLINSON_DB_OVER_D, _TOMLINSON_SU_MPA, _TOMLINSON_ALPHA_TABLES["soft_clay"])
        a_stiff = bilinear_interp(x, su, _TOMLINSON_DB_OVER_D, _TOMLINSON_SU_MPA, _TOMLINSON_ALPHA_TABLES["stiff_clay"])
        return min(a_soft, a_stiff)
    if case not in _TOMLINSON_ALPHA_TABLES:
        case = "stiff_clay"
    return bilinear_interp(x, su, _TOMLINSON_DB_OVER_D, _TOMLINSON_SU_MPA, _TOMLINSON_ALPHA_TABLES[case])


def _is_nondisplacement_driven_pile(pile_type: str = "") -> bool:
    """Cọc không chiếm chỗ: cọc H, cọc ống mũi hở/thành mỏng không tạo nút đất."""
    txt = _strip_accents(pile_type).lower()
    keys = ["khong chiem cho", "khong tao nut", "mui ho", "open end", "open-end", "non-displacement", "nondisplacement"]
    return any(k in txt for k in keys) or re.search(r"\bcoc\s*h\b", txt) is not None


def driven_sand_side_nominal_kpa(n160: float, pile_type: str = "") -> Tuple[float, str]:
    """Ma sát bên danh định cọc đóng/ép trong đất rời theo Meyerhof SPT."""
    factor = DRIVEN_SAND_SIDE_NONDISPLACEMENT_KPA_PER_N if _is_nondisplacement_driven_pile(pile_type) else DRIVEN_SAND_SIDE_DISPLACEMENT_KPA_PER_N
    return max(factor * max(float(n160), 0.0), 0.0), f"Cọc đóng/ép đất rời Meyerhof: qs={factor:.2f}·N160 kPa"


def _driven_sand_tip_limit_case(text: Any = "") -> str:
    plain = _strip_accents(text).lower()
    if any(k in plain for k in ("cat bot", "cat bui", "bui", "bot", "silt", "silty", "ml")):
        return "silty_sand"
    return "sand"


def driven_sand_tip_nominal_kpa(n160: float, db_over_d: float, soil_text: Any = "") -> Tuple[float, Dict[str, float], str]:
    """Sức kháng mũi danh định cọc đóng/ép trong đất rời theo Meyerhof SPT."""
    n = max(float(n160), 0.0)
    ratio = max(float(db_over_d or 0.0), 0.0)
    limit_case = _driven_sand_tip_limit_case(soil_text)
    qlambda_mpa = (0.30 if limit_case == "silty_sand" else 0.40) * n
    qp_mpa_raw = 0.038 * n * ratio
    qp_mpa = min(qp_mpa_raw, qlambda_mpa)
    note = f"Cọc đóng/ép đất rời Meyerhof: qp=0.038·N160·Db/D={qp_mpa_raw:.3f} MPa; qλ={qlambda_mpa:.3f} MPa"
    return qp_mpa * 1000.0, {"N160": n, "Db_over_D": ratio, "qp_mpa_raw": qp_mpa_raw, "qp_mpa_limit": qlambda_mpa, "sand_tip_limit_case": limit_case}, note


def _interval_overlap_len(top1: float, bottom1: float, top2: float, bottom2: float) -> float:
    """Chiều dài giao nhau của hai đoạn cao độ [bottom, top]."""
    a_top, a_bot = max(float(top1), float(bottom1)), min(float(top1), float(bottom1))
    b_top, b_bot = max(float(top2), float(bottom2)), min(float(top2), float(bottom2))
    return max(min(a_top, b_top) - max(a_bot, b_bot), 0.0)


def clay_c_phi_side_nominal_kpa(c_mpa: float, phi_deg: float, sigma_v_eff_mpa: float) -> Tuple[float, float, str]:
    """Sức kháng bên danh định đất dính theo lựa chọn C,φ.

    Đây là nhánh tùy chọn thực dụng để dùng bộ chỉ tiêu c,φ trong bảng địa chất:
    qs = c + σ'v*tanφ, đổi MPa -> kPa. Mặc định module vẫn dùng SPT -> Su.
    """
    c = max(float(c_mpa), 0.0)
    phi = max(float(phi_deg), 0.0)
    sig = max(float(sigma_v_eff_mpa), 0.0)
    qs_mpa = c + sig * math.tan(math.radians(phi))
    return max(qs_mpa * 1000.0, 0.0), phi, "Đất dính C,φ: qs=c+σ'v·tanφ"


def clay_c_phi_tip_nominal_kpa(c_mpa: float, phi_deg: float, sigma_v_eff_mpa: float) -> Tuple[float, Dict[str, float], str]:
    """Sức kháng mũi danh định đất dính theo C,φ, dạng sức chịu tải tổng quát.

    qp = c*Nc + σ'v*Nq. Khi φ≈0, lấy Nc=9, Nq=1.
    """
    c = max(float(c_mpa), 0.0)
    phi = max(float(phi_deg), 0.0)
    sig = max(float(sigma_v_eff_mpa), 0.0)
    if phi <= 1e-6:
        Nq, Nc = 1.0, 9.0
    else:
        tanphi = math.tan(math.radians(phi))
        Nq = math.exp(math.pi * tanphi) * (math.tan(math.radians(45.0 + phi / 2.0)) ** 2.0)
        Nc = (Nq - 1.0) / max(tanphi, 1e-12)
    qp_mpa = c * Nc + sig * Nq
    return max(qp_mpa * 1000.0, 0.0), {"Nc": Nc, "Nq": Nq, "c_mpa": c, "phi_deg": phi}, "Đất dính C,φ: qp=cNc+σ'vNq"


def beta_sand_bored(n60: float, n1_60: float, sigma_v_eff_mpa: float, mode: str = "sand", m_sand: float = 0.6) -> Tuple[float, float, float]:
    """Trả về beta, phi_deg, sigma_p_eff cho cọc khoan trong đất rời.

    Theo bảng Excel/TCVN 11823-10 đang dùng:
    - góc ma sát tính toán: phi_f = 27.5 + 9.2*log10((N1)60)
    - ứng suất tiền cố kết hiệu quả: sigma'p = 0.47*N60^m*pa cho cát
      hoặc sigma'p = 0.15*N60*pa cho sỏi.

    Vì vậy KHÔNG được dùng (N1)60 cho sigma'p. V0.1.1 đã truyền nhầm
    một biến vào cả hai vị trí; V0.1.2 tách rõ N60 và (N1)60.
    """
    n60_val = max(float(n60), 1e-6)
    n160_val = max(float(n1_60), 1e-6)
    sv = max(float(sigma_v_eff_mpa), 1e-5)
    phi_deg = 27.5 + 9.2 * math.log10(n160_val)
    sin_phi = math.sin(math.radians(phi_deg))
    tan_phi = math.tan(math.radians(phi_deg))
    if str(mode).strip().lower().startswith("gravel") or "sỏi" in str(mode).lower():
        sigma_p = 0.15 * n60_val * PA_MPA
    else:
        sigma_p = 0.47 * (n60_val ** float(m_sand)) * PA_MPA
    sigma_p = max(sigma_p, 1e-9)
    beta = (1.0 - sin_phi) * tan_phi * ((sigma_p / sv) ** sin_phi)
    return max(beta, 0.0), phi_deg, sigma_p


def cn_overburden_tcvn(sigma_v_eff_mpa: float) -> float:
    """Hệ số hiệu chỉnh áp lực phủ CN theo TVQT/TCVN.

    CN = 0.77*log10(1.92/sigma'v), giới hạn CN <= 2.0.
    Không ép (N1)60 <= N60, vì TVQT tính N160 = CN*N60.
    """
    sv = max(float(sigma_v_eff_mpa), 1e-5)
    cn = 0.77 * math.log10(1.92 / sv)
    return max(min(cn, 2.0), 0.0)


def n1_60_corrected(n_spt: float, sigma_v_eff_mpa: float, cap_to_n: bool = False) -> float:
    n = max(float(n_spt), 0.0)
    corr = n * cn_overburden_tcvn(sigma_v_eff_mpa)
    if cap_to_n:
        corr = min(corr, n)
    return max(corr, 0.0)


def rock_alpha_e_from_rqd(rqd: float, open_joint: bool = True) -> float:
    """Nội suy hệ số alpha_E theo RQD và tình trạng khe nứt đá.

    Bảng TCVN 11823-10:2017/O'Neill & Reese:
    RQD = 20,30,50,70,100%;
    khe nứt khép kín: 0.45,0.50,0.60,0.85,1.00;
    khe nứt hở hoặc có trét mùn: 0.45,0.50,0.55,0.55,0.85.
    open_joint=True nghĩa là khe nứt hở/có mùn.
    """
    xs = [20, 30, 50, 70, 100]
    closed = [0.45, 0.50, 0.60, 0.85, 1.00]
    openj = [0.45, 0.50, 0.55, 0.55, 0.85]
    return linear_interp(float(rqd), xs, openj if open_joint else closed, clamp=True)


def ksp_factor(spacing_crack_mm: float, width_crack_mm: float, diameter_m: float) -> float:
    """Hệ số Ksp cho mũi IGM/đá mềm theo khe nứt.

    spacing_crack_mm = sd: khoảng cách trung bình giữa các khe nứt, mm.
    width_crack_mm = td: chiều rộng/độ mở khe nứt, mm.
    Giá trị mặc định GUI: sd=50 mm, td=5 mm để thiên về bất lợi khi thiếu số liệu chi tiết.
    """
    sd = max(float(spacing_crack_mm), 1e-6)
    td = max(float(width_crack_mm), 0.0)
    d_mm = max(float(diameter_m) * 1000.0, 1e-6)
    return (3.0 + sd / d_mm) / (10.0 * math.sqrt(1.0 + 300.0 * td / sd))


def depth_factor_socket(socket_depth_m: float, diameter_m: float) -> float:
    return min(1.0 + 0.4 * max(socket_depth_m, 0.0) / max(diameter_m, 1e-9), 3.4)


def rock_side_nominal_kpa(qu_mpa: float, fc_mpa: float, rqd: float, supported: bool = True, open_joint: bool = True) -> Tuple[float, float, str]:
    """Sức kháng thành bên danh định trong đá theo TVQT/TCVN, trả kPa.

    - Thi công có chống đỡ thành hố / điều kiện bảo thủ theo Eq.97:
      qs = 0.65*αE*pa*sqrt(min(qu,f'c)/pa). Không bỏ αE khi giới hạn theo f'c.
    - Thi công không chống đỡ / thành hố tự ổn định theo Eq.96:
      qs = C*pa*sqrt(min(qu,f'c)/pa), C=1.0 cho thiết bị khoan thông thường.

    V0.2.33: rà lại theo TCVN 11823-10; cả Eq.96 và Eq.97 đều dùng qu_eff=min(qu,f'c) khi qu>f'c.
    Điều kiện khe nứt đá dùng để xác định αE theo RQD: khe nứt khép kín hoặc khe nứt hở/có mùn.
    qu và f'c nhập MPa; kết quả đổi sang kPa.
    """
    qu = max(float(qu_mpa), 0.0)
    fc = max(float(fc_mpa), 0.0)
    if qu <= 0:
        return 0.0, 0.0, "Thiếu qu đá"
    ae = rock_alpha_e_from_rqd(rqd, open_joint)
    qu_eff = min(qu, fc) if fc > 0 else qu
    lim_note = "; dùng min(qu,f'c)" if fc > 0 and qu > fc else ""
    if supported:
        qs_mpa = 0.65 * ae * PA_MPA * math.sqrt(max(qu_eff, 0.0) / PA_MPA)
        return qs_mpa * 1000.0, ae, "Đá - thi công có chống đỡ: 0.65*αE*pa*sqrt(min(qu,f'c)/pa)" + lim_note
    qs_mpa = PA_MPA * math.sqrt(max(qu_eff, 0.0) / PA_MPA)
    return qs_mpa * 1000.0, 1.0, "Đá - thi công không chống đỡ: C*pa*sqrt(min(qu,f'c)/pa), C=1.0" + lim_note


def igm_side_nominal_kpa(n60: float, qu_mpa: float, alpha: float, joint_factor: float, use_spt_when_missing_qu: bool = False) -> Tuple[float, str]:
    """Sức kháng thành bên danh định IGM.

    Quy trình mặc định:
    - Loại 1: SPT < 100 -> fsn = 0.0036*N (MPa).
    - Loại 2: SPT >= 100 và qu < 4.7 MPa -> fsn = 0.098*sqrt(qu) (MPa).
    - Loại 3: 4.7 <= qu < 23.9 MPa -> fsn = 0.81*qu^0.51 (MPa).
    - Nếu SPT >= 100 nhưng thiếu qu: chỉ dùng công thức SPT khi người dùng chọn rõ trong GUI.

    alpha và joint_factor được giữ trong chữ ký hàm để tương thích GUI/các bản cũ, nhưng không dùng
    trong quy trình IGM này.
    """
    n = max(float(n60), 0.0)
    qu = max(float(qu_mpa), 0.0)
    if n < 100.0:
        return 0.0036 * n * 1000.0, "IGM loại 1: SPT<100 -> fsn=0.0036N MPa"
    if qu <= 0.0:
        if use_spt_when_missing_qu:
            return 0.0036 * n * 1000.0, "IGM thiếu qu: người dùng chọn tính theo SPT -> fsn=0.0036N MPa"
        return 0.0, "IGM SPT>=100 nhưng thiếu qu -> chưa tính fsn"
    if qu < 4.7:
        return 0.098 * math.sqrt(qu) * 1000.0, "IGM loại 2: SPT>=100, qu<4.7MPa -> fsn=0.098√qu MPa"
    if qu < 23.9:
        return 0.81 * (qu ** 0.51) * 1000.0, "IGM loại 3: 4.7<=qu<23.9MPa -> fsn=0.81qu^0.51 MPa"
    return 0.81 * (qu ** 0.51) * 1000.0, "IGM qu>=23.9MPa ngoài miền khuyến nghị; fsn=0.81qu^0.51 MPa"


def igm_tip_nominal_kpa(n60: float, qu_mpa: float, crack_spacing_mm: float, crack_width_mm: float, diameter_m: float, socket_depth_m: float, use_spt_when_missing_qu: bool = False) -> Tuple[float, Dict[str, float], str]:
    """Sức kháng mũi danh định IGM.

    - Loại 1: SPT < 100 -> qBN = 0.0439*N (MPa).
    - Loại 2: SPT >= 100 và qu < 4.7 MPa -> qBN = 0.057*(1+Ld/D)*sqrt(qu) (MPa).
    - Nếu SPT >= 100 nhưng thiếu qu: chỉ dùng công thức SPT khi người dùng chọn rõ trong GUI.
    - Với qu >= 4.7 MPa, giữ nhánh mũi theo qu/Ksp của bản trước do bảng ảnh không nêu công thức mũi riêng.
    """
    n = max(float(n60), 0.0)
    qu = max(float(qu_mpa), 0.0)
    D = max(float(diameter_m), 1e-9)
    Ld = max(float(socket_depth_m), 0.0)
    if n < 100.0:
        return 0.0439 * n * 1000.0, {"N60": n}, "IGM loại 1: SPT<100 -> qBN=0.0439N MPa"
    if qu <= 0.0:
        if use_spt_when_missing_qu:
            return 0.0439 * n * 1000.0, {"N60": n, "igm_missing_qu_policy": "use_spt"}, "IGM thiếu qu: người dùng chọn tính theo SPT -> qBN=0.0439N MPa"
        return 0.0, {"N60": n}, "IGM SPT>=100 nhưng thiếu qu -> chưa tính qBN"
    if qu < 4.7:
        qp_mpa = 0.057 * (1.0 + Ld / D) * math.sqrt(qu)
        return qp_mpa * 1000.0, {"N60": n, "qu_mpa": qu, "socket_depth_m": Ld, "Ld_over_D": Ld / D}, "IGM loại 2: SPT>=100, qu<4.7MPa -> qBN=0.057(1+Ld/D)√qu MPa"
    ksp = ksp_factor(crack_spacing_mm, crack_width_mm, D)
    df = depth_factor_socket(Ld, D)
    qp_kpa = 3.0 * qu * ksp * df * 1000.0
    return qp_kpa, {"ksp": ksp, "depth_factor": df, "N60": n, "qu_mpa": qu, "socket_depth_m": Ld}, "IGM qu>=4.7MPa -> qBN=3quKsp*d"


def hoek_brown_qp(sigma_vb_mpa: float, qu_mpa: float, mi: float, gsi: float, disturbance: float) -> Tuple[float, Dict[str, float]]:
    """Sức kháng mũi đá nứt vỡ theo Hoek-Brown/GSI, trả MPa.

    Theo TCVN 11823-10 Eq.99-Eq.100:
        A  = σ'vb + qu * [mb*(σ'vb/qu)]^a
        qp = A    + qu * [mb*(A/qu) + s]^a
    Lưu ý: theo Eq.100 của TCVN, bước tính A không có +s.
    Trả thêm qp_uncapped_mpa; giới hạn 2.5qu được áp ở _tip_resistance theo lựa chọn người dùng.
    """
    qu = max(float(qu_mpa), 1e-6)
    sig = max(float(sigma_vb_mpa), 0.0)
    D = max(min(float(disturbance), 1.0), 0.0)
    gsi = float(gsi)
    mi = float(mi)
    s = math.exp((gsi - 100.0) / max(9.0 - 3.0 * D, 1e-9))
    a = 0.5 + (1.0 / 6.0) * (math.exp(-gsi / 15.0) - math.exp(-20.0 / 3.0))
    mb = mi * math.exp((gsi - 100.0) / max(28.0 - 14.0 * D, 1e-9))
    A = sig + qu * max(mb * (sig / qu), 0.0) ** a
    qp = A + qu * max(mb * (A / qu) + s, 0.0) ** a
    return qp, {"s": s, "a": a, "mb": mb, "A_mpa": A, "qb_mpa": A, "qp_uncapped_mpa": qp}


def _is_one_row_group_layout(value: Any) -> bool:
    """Xác định bệ 1 hàng cọc.

    V0.2.6: cột nhập liệu đổi thành Số hàng cọc, người dùng nhập số cho ngắn gọn.
    Vẫn giữ tương thích với chuỗi cũ "1 hàng cọc" / "Nhiều hàng cọc".
    """
    s = str(value or "").strip().lower()
    s_plain = s.replace("đ", "d")
    try:
        # Chỉ xem là 1 hàng khi giá trị số thực sự <= 1.5; tránh lỗi "10" bị nhận là 1 hàng.
        return float(s_plain.replace(",", ".")) <= 1.5
    except Exception:
        pass
    return ("1" in s and "hàng" in s) or ("1" in s_plain and "hang" in s_plain) or "mot hang" in s_plain or "một hàng" in s


def group_factor_one_row(spacing_over_d: float) -> float:
    """Hệ số nhóm cho bệ 1 hàng cọc.

    Nội suy theo S/D = khoảng cách cọc / đường kính cọc.
    Không dùng 2.72D làm mốc bảng; 2.72D chỉ là một giá trị S/D thực tế do người dùng nhập.
    Theo yêu cầu bảng tính hiện hành: 2.5D -> 0.90, 4.0D -> 1.00.
    """
    return linear_interp(spacing_over_d, [2.5, 4.0], [0.90, 1.00], clamp=True)


def group_factor_clay(spacing_over_d: float, layout: str = "Nhiều hàng cọc") -> float:
    """Hệ số nhóm đất dính theo S/D.

    Bệ 1 hàng dùng bảng riêng. Bệ nhiều hàng: 2.5D -> 0.65, 6.0D -> 1.00.
    """
    if _is_one_row_group_layout(layout):
        return group_factor_one_row(spacing_over_d)
    return linear_interp(spacing_over_d, [2.5, 6.0], [0.65, 1.00], clamp=True)


def group_factor_sand(spacing_over_d: float, layout: str = "Nhiều hàng cọc") -> float:
    """Hệ số nhóm đất rời/IGM/đá mềm theo S/D.

    Bệ 1 hàng dùng bảng riêng. Bệ nhiều hàng: 2.5D -> 0.67, 4.0D -> 1.00.
    """
    if _is_one_row_group_layout(layout):
        return group_factor_one_row(spacing_over_d)
    return linear_interp(spacing_over_d, [2.5, 4.0], [0.67, 1.00], clamp=True)


def group_factor_for_soil(soil_type: int, spacing_over_d: float, row_count: Any = "2", ignore_igm_rock: bool = False) -> Tuple[float, str]:
    """Hệ số nhóm theo từng lớp đất.

    - Đất dính áp bảng đất dính.
    - Đất rời/cuội sỏi áp bảng đất rời.
    - IGM/đá dùng bảng đất rời nếu không bỏ qua theo option.
    - Hệ số được nội suy theo S/D = khoảng cách cọc / D cọc.
    """
    st = int(soil_type or 0)
    if st == 2:
        return group_factor_clay(spacing_over_d, row_count), "fg đất dính"
    if st in (1, 6):
        return group_factor_sand(spacing_over_d, row_count), "fg đất rời"
    if st in (3, 4, 5):
        if ignore_igm_rock:
            return 1.0, "fg IGM/đá bỏ qua"
        return group_factor_sand(spacing_over_d, row_count), "fg IGM/đá"
    return 1.0, "fg=1.0"


@dataclass
class SoilLayer:
    name: str = ""
    bottom_elev_m: float = 0.0
    soil_type: int = 1  # 1=sand, 2=clay, 3=IGM/rock
    n_spt: float = 0.0      # Nₕₜ/N đo hiện trường nhập trong bảng địa chất; N60 = Nₕₜ*ER/60.
    gamma_kN_m3: float = 18.0
    c_mpa: float = 0.0      # C hoặc lực dính, MPa. Giữ tên c_mpa để dùng cho phương pháp C,φ.
    su_mpa: float = 0.0     # Su tương thích file cũ; V0.2 UI không nhập trực tiếp.
    phi_deg: float = 0.0
    qu_mpa: float = 0.0
    rqd: float = 50.0
    gsi: float = 30.0
    mi: float = 9.0
    disturbance: float = 0.5
    comment: str = ""


@dataclass
class PileInput:
    project: str = ""
    item: str = ""
    mode: str = "Cọc khoan trong đất"
    pile_type: str = "Cọc khoan nhồi"  # Cọc khoan nhồi / Cọc đóng (ép)
    analysis_type_auto: str = ""
    n_piles: int = 1
    diameter_mm: float = 1200.0
    # QA fix P7: 0 = mũi bằng D. Tránh trường hợp gọi API chỉ set diameter_mm mà mũi âm thầm dùng 1200.
    tip_diameter_mm: float = 0.0
    # QA fix P1: đường kính trong cho cọc ống (2O/3O); 0 = cọc đặc.
    driven_inner_dia_mm: float = 0.0
    spacing_m: float = 3.0
    pile_count_in_group: int = 1
    group_layout: str = "2"  # V0.2.6: Số hàng cọc; 1 = bệ 1 hàng, >=2 = nhiều hàng
    cap_width_m: float = 0.0
    cap_length_m: float = 0.0
    cap_thickness_m: float = 0.0
    ground_elev_m: float = 0.0
    cap_bottom_elev_m: float = 0.0
    pile_tip_elev_m: float = -30.0
    water_elev_m: float = 0.0
    concrete_gamma_kN_m3: float = 24.5
    fc_mpa: float = 30.0
    fy_mpa: float = 400.0
    n_rebars: int = 0
    rebar_dia_mm: float = 0.0
    structural_phi: float = 0.75
    stirrup_type: int = 1
    exclude_top_bored_m: float = 1.5
    sand_preconsolidation_mode: str = "sand"
    sand_m: float = 0.6
    igm_alpha: float = 0.25
    igm_joint_factor: float = 0.45
    igm_missing_qu_policy: str = "require_qu"  # require_qu / use_spt: xử lý IGM SPT>=100 nhưng thiếu qu
    rock_socket_depth_m: Optional[float] = None
    rock_side_method: str = "fractured"  # legacy: normal/fractured/ignore, chỉ giữ để đọc file cũ
    rock_construction_condition: str = "Có chống đỡ"  # Có chống đỡ / Không chống đỡ
    rock_joint_condition: str = "Khe nứt hở hoặc có mùn"  # Khe nứt khép kín / Khe nứt hở hoặc có mùn
    rock_tip_condition: str = "fractured"  # intact/fractured
    include_rock_tip: bool = False
    allow_rock_tip_exceed_25qu: bool = False  # Chỉ bật khi có thử tải/kinh nghiệm cho phép Qp đá vượt 2.5qu
    spt_er_percent: float = 60.0   # N60 = Nₕₜ * ER/60
    spt_input_mode: str = "Nₕₜ"  # Nₕₜ: số búa hiện trường; N60: dữ liệu nhập đã hiệu chỉnh năng lượng.
    clay_use_c_phi: bool = False
    ignore_group_igm_rock: bool = False
    allow_geology_extrapolation: bool = False  # Cho phép kéo dài lớp cuối xuống mũi cọc khi thiếu chiều sâu khảo sát
    driven_sand_method: str = "Meyerhof SPT"  # Nhánh cọc đóng đất rời: Meyerhof SPT; Nordlund cần thêm tham số riêng
    driven_clay_alpha_method: str = "Tomlinson"
    include_downdrag: bool = False
    downdrag_top_elev_m: float = 0.0
    downdrag_bottom_elev_m: float = 0.0
    downdrag_factor: float = 0.0
    force_cd_kn: float = 0.0       # Pu nén lớn nhất theo cọc đơn
    force_db_kn: float = 0.0       # Pu nén lớn nhất theo cọc đơn
    uplift_cd_kn: float = 0.0      # Lực nhổ lớn nhất/cọc TTGHCĐ, lấy trị tuyệt đối của N âm
    uplift_db_kn: float = 0.0      # Lực nhổ lớn nhất/cọc TTGHĐB, lấy trị tuyệt đối của N âm
    cap_force_cd_kn: float = 0.0   # Nội lực đáy bệ TTGHCĐ để kiểm nhóm
    cap_force_db_kn: float = 0.0   # Nội lực đáy bệ TTGHĐB để kiểm nhóm
    crack_spacing_mm: float = 50.0
    crack_width_mm: float = 5.0
    rock_open_joint: bool = True
    layers: List[SoilLayer] = field(default_factory=list)

    @property
    def diameter_m(self) -> float:
        return max(self.diameter_mm / 1000.0, 1e-9)

    @property
    def tip_diameter_m(self) -> float:
        return max((self.tip_diameter_mm or self.diameter_mm) / 1000.0, 1e-9)

    @property
    def _is_driven_mode(self) -> bool:
        mode_text = (self.mode + " " + self.pile_type).lower()
        return "đóng" in mode_text or "dong" in mode_text or "ép" in mode_text or "ep" in mode_text

    @property
    def driven_shape(self) -> str:
        """QA fix P1: tiết diện cọc đóng/ép: VUONG (mặc định, mã 2/3), TRON (2T/3T), ONG (2O/3O)."""
        if not self._is_driven_mode:
            return "TRON"
        text = _strip_accents(self.pile_type).lower()
        if " ong" in f" {text}" or "phc" in text:
            return "ONG"
        if "tron" in text:
            return "TRON"
        return "VUONG"

    @property
    def inner_diameter_m(self) -> float:
        return max(float(self.driven_inner_dia_mm or 0.0), 0.0) / 1000.0

    @property
    def perimeter_m(self) -> float:
        # QA fix P1: chu vi theo tiết diện thật: vuông = 4D; tròn/ống = πD.
        # Trước đây mọi cọc đóng/ép bị coi là vuông (4D) làm ma sát thành bên
        # của cọc tròn/ống bị phóng đại ~27%.
        if self._is_driven_mode and self.driven_shape == "VUONG":
            return 4.0 * self.diameter_m
        return math.pi * self.diameter_m

    @property
    def area_m2(self) -> float:
        """Diện tích mũi cọc dùng cho sức kháng mũi.

        Cọc ống coi như bịt mũi/nút đất kín -> dùng tiết diện đặc πD²/4 (có cảnh báo trong validate).
        """
        if self._is_driven_mode:
            if self.driven_shape == "VUONG":
                return self.diameter_m ** 2
            return math.pi * self.diameter_m ** 2 / 4.0
        return math.pi * self.tip_diameter_m ** 2 / 4.0

    @property
    def shaft_area_gross_m2(self) -> float:
        """Diện tích tiết diện thực: dùng cho Pr vật liệu và trọng lượng cọc (cọc ống trừ lõi rỗng)."""
        if self._is_driven_mode:
            if self.driven_shape == "VUONG":
                return self.diameter_m ** 2
            if self.driven_shape == "ONG":
                din = min(self.inner_diameter_m, self.diameter_m)
                return math.pi * max(self.diameter_m ** 2 - din ** 2, 0.0) / 4.0
            return math.pi * self.diameter_m ** 2 / 4.0
        return math.pi * self.diameter_m ** 2 / 4.0

    @property
    def pile_length_m(self) -> float:
        return max(self.cap_bottom_elev_m - self.pile_tip_elev_m, 0.0)


@dataclass
class LayerCalc:
    name: str
    top_elev_m: float
    bottom_elev_m: float
    thickness_m: float
    skin_length_m: float
    soil_type: int
    soil_label: str
    n_spt: float      # Nₕₜ/N nhập
    n60: float         # N60 đã hiệu chỉnh năng lượng ER
    n1_60: float       # (N1)60/N160 hiệu chỉnh áp lực phủ
    gamma_eff_kN_m3: float
    sigma_v_eff_mpa: float
    c_mpa: float
    su_mpa: float
    alpha_or_beta: float
    phi_deg: float
    qs_factored_kpa: float          # TTGHCĐ - giữ tên cũ để tương thích
    qs_nominal_kpa: float
    qs_factored_kn: float           # TTGHCĐ - giữ tên cũ để tương thích
    qs_nominal_kn: float
    qs_extreme_kpa: float = 0.0     # TTGHĐB
    qs_extreme_kn: float = 0.0      # TTGHĐB
    phi_strength: float = 0.0
    phi_extreme: float = 1.0
    group_factor: float = 1.0
    qu_mpa: float = 0.0
    rqd: float = 0.0
    note: str = ""
    qs_uplift_factored_kpa: float = 0.0
    qs_uplift_factored_kn: float = 0.0
    qs_uplift_extreme_kpa: float = 0.0
    qs_uplift_extreme_kn: float = 0.0
    downdrag_length_m: float = 0.0
    downdrag_kn: float = 0.0


@dataclass
class LimitStateCapacity:
    code: str
    label: str
    qshaft_kn: float
    qtip_kn: float
    compression_single_gross_kn: float
    pile_weight_effective_kn: float
    pile_weight_dry_kn: float
    buoyancy_kn: float
    compression_single_net_kn: float
    uplift_single_magnitude_kn: float
    uplift_single_signed_kn: float
    group_factor: float
    compression_group_single_gross_kn: float
    compression_group_single_net_kn: float
    compression_group_total_kn: float
    material_pr_kn: float
    governing_kn: float
    governing_note: str
    qshaft_uplift_kn: float = 0.0
    uplift_group_total_kn: float = 0.0
    downdrag_kn: float = 0.0
    compression_single_net_before_downdrag_kn: float = 0.0
    compression_group_total_before_downdrag_kn: float = 0.0


@dataclass
class CapacityResult:
    pile_input: PileInput
    layers: List[LayerCalc]
    qshaft_nominal_kn: float
    qtip_nominal_kn: float
    strength: LimitStateCapacity
    extreme: LimitStateCapacity
    warnings: List[str] = field(default_factory=list)
    toe_info: Dict[str, Any] = field(default_factory=dict)

    # Các property sau giữ tương thích với bản cũ, lấy theo TTGHCĐ.
    @property
    def qshaft_factored_kn(self) -> float:
        return self.strength.qshaft_kn

    @property
    def qtip_factored_kn(self) -> float:
        return self.strength.qtip_kn

    @property
    def compression_single_factored_kn(self) -> float:
        return self.strength.compression_single_gross_kn

    @property
    def compression_single_net_kn(self) -> float:
        return self.strength.compression_single_net_kn

    @property
    def uplift_single_factored_kn(self) -> float:
        return self.strength.uplift_single_magnitude_kn

    @property
    def group_factor(self) -> float:
        return self.strength.group_factor

    @property
    def compression_group_single_factored_kn(self) -> float:
        return self.strength.compression_group_single_gross_kn

    @property
    def compression_group_single_net_kn(self) -> float:
        return self.strength.compression_group_single_net_kn

    @property
    def compression_group_total_kn(self) -> float:
        return self.strength.compression_group_total_kn

    @property
    def material_pr_kn(self) -> float:
        return self.strength.material_pr_kn

    @property
    def governing_geotech_kn(self) -> float:
        return self.strength.governing_kn

    @property
    def governing_note(self) -> str:
        return self.strength.governing_note


class SCTCalculator:
    @staticmethod
    def structural_capacity_nominal_kn(inp: PileInput) -> float:
        # Pn = k1*(0.85fc(Ag-Ast)+fyAst), đơn vị kN.
        Ag = inp.shaft_area_gross_m2 * 1e6
        Ast = 0.0
        if inp.n_rebars > 0 and inp.rebar_dia_mm > 0:
            Ast = inp.n_rebars * math.pi * inp.rebar_dia_mm ** 2 / 4.0
        k1 = 0.85 if int(inp.stirrup_type) == 1 else 0.80
        pn_n = k1 * (0.85 * inp.fc_mpa * max(Ag - Ast, 0.0) + inp.fy_mpa * Ast)
        return pn_n / 1000.0

    @staticmethod
    def structural_phi_for_limit_state(limit_state: str) -> float:
        # TVQT dùng 0.75 cho TTGHCĐ và 1.0 cho TTGHĐB.
        return 1.0 if str(limit_state).upper() in (LIMIT_STATE_EXTREME, "DB", "ĐB", "EXTREME") else 0.75

    @staticmethod
    def structural_capacity_kn(inp: PileInput, limit_state: str = LIMIT_STATE_STRENGTH) -> float:
        return SCTCalculator.structural_capacity_nominal_kn(inp) * SCTCalculator.structural_phi_for_limit_state(limit_state)

    @staticmethod
    def _is_driven(inp: PileInput) -> bool:
        txt = _strip_accents((inp.mode + " " + inp.pile_type)).lower()
        return ("dong" in txt) or ("ep" in txt)

    @staticmethod
    def _uplift_phi(inp: PileInput, soil_type: int, limit_state: str = LIMIT_STATE_STRENGTH) -> float:
        """Hệ số sức kháng nhổ tách riêng với hệ số nén."""
        st = int(soil_type or 0)
        ls = str(limit_state).upper()
        if ls in (LIMIT_STATE_EXTREME, "DB", "ĐB", "EXTREME"):
            return PHI_EXTREME_UPLIFT_DEFAULT
        if SCTCalculator._is_driven(inp):
            if st in (1, 6):
                return PHI_UPLIFT_DRIVEN_SAND
            if st == 2:
                return PHI_UPLIFT_DRIVEN_CLAY
            if st == 5:
                return PHI_UPLIFT_BORE_IGM
            if st in (3, 4):
                return PHI_UPLIFT_ROCK_SIDE
            return 0.0
        if st in (1, 6):
            return PHI_UPLIFT_BORE_SAND
        if st == 2:
            return PHI_UPLIFT_BORE_CLAY
        if st == 5:
            return PHI_UPLIFT_BORE_IGM
        if st in (3, 4):
            return PHI_UPLIFT_ROCK_SIDE
        return 0.0

    @staticmethod
    def default_downdrag_load_factor(inp: PileInput) -> float:
        """Hệ số tải trọng ma sát âm γDD theo TCVN 11823-3, Bảng 4.

        Dùng giá trị lớn nhất khi DD là tác dụng bất lợi trong kiểm toán nén:
        - Cọc đóng/ép theo α Tomlinson: 1.40
        - Cọc đóng/ép theo λ: 1.05
        - Cọc khoan theo O'Neill & Reese: 1.25
        """
        if SCTCalculator._is_driven(inp):
            method = _strip_accents(str(getattr(inp, "driven_clay_alpha_method", "Tomlinson") or "Tomlinson")).lower()
            if "lambda" in method or method.strip() in ("lamda", "λ"):
                return GAMMA_DD_LAMBDA
            return GAMMA_DD_TOMLINSON
        return GAMMA_DD_ONEILL_REESE

    @staticmethod
    def _downdrag_zone(inp: PileInput) -> Optional[Tuple[float, float]]:
        if not bool(getattr(inp, "include_downdrag", False)):
            return None
        top = float(getattr(inp, "downdrag_top_elev_m", inp.ground_elev_m) or inp.ground_elev_m)
        bot = float(getattr(inp, "downdrag_bottom_elev_m", inp.pile_tip_elev_m) or inp.pile_tip_elev_m)
        if abs(top - bot) < 1e-9:
            return None
        return max(top, bot), min(top, bot)

    @staticmethod
    def _preprocess_layers(inp: PileInput) -> List[str]:
        """Chuẩn hóa dữ liệu lớp trước khi tính, đặc biệt lớp 0 không khí/hang karst."""
        warnings: List[str] = []
        for ly in inp.layers:
            try:
                if int(ly.soil_type or 0) == 0:
                    changed = []
                    if abs(float(ly.gamma_kN_m3 or 0.0)) > 1e-9:
                        changed.append(f"γ {ly.gamma_kN_m3:g}->0")
                    ly.gamma_kN_m3 = 0.0
                    ly.n_spt = 0.0
                    ly.c_mpa = 0.0
                    ly.su_mpa = 0.0
                    ly.phi_deg = 0.0
                    ly.qu_mpa = 0.0
                    ly.rqd = 0.0
                    if changed:
                        warnings.append(f"Lớp {ly.name or '?'} type 0 Không khí/Hang karst: tự ép " + ", ".join(changed) + ".")
            except Exception:
                continue
        return warnings

    @staticmethod
    def _validate_input(inp: PileInput) -> List[str]:
        warnings: List[str] = []
        if _safe_float(getattr(inp, "diameter_mm", 0.0), 0.0) <= 0.0:
            warnings.append("Chưa nhập đường kính/cạnh cọc D; cần bổ sung trước khi dùng kết quả tính toán.")
        if not str(getattr(inp, "pile_type", "") or "").strip():
            warnings.append("Chưa chọn loại cọc; nhập 1 = khoan nhồi, 2/3 = đóng/ép vuông, 2T/3T = đóng/ép tròn, 2O/3O = cọc ống (Ds = D trong).")
        if SCTCalculator._is_driven(inp):
            shape = getattr(inp, "driven_shape", "VUONG")
            if shape == "ONG":
                din = _safe_float(getattr(inp, "driven_inner_dia_mm", 0.0), 0.0)
                dout = _safe_float(getattr(inp, "diameter_mm", 0.0), 0.0)
                if din <= 0.0:
                    warnings.append("Cọc ống nhưng chưa nhập đường kính trong (cột Ds); Pr vật liệu/trọng lượng tạm tính như cọc tròn đặc.")
                elif din >= dout > 0.0:
                    warnings.append("Đường kính trong cọc ống >= đường kính ngoài; kiểm tra lại cột Ds.")
                warnings.append("Cọc ống: sức kháng mũi tính với tiết diện đặc (coi như bịt mũi/nút đất kín); "
                                "Pr vật liệu tính như BTCT thường - cọc PHC ứng suất trước cần đối chiếu catalogue nhà sản xuất.")
            elif shape == "VUONG":
                warnings.append("Cọc đóng/ép mã 2/3 được tính là cọc VUÔNG đặc (chu vi 4D, diện tích D×D); "
                                "nếu thực tế là cọc tròn/ống hãy dùng mã 2T/3T hoặc 2O/3O.")
        layers = SCTCalculator._sorted_layers(inp)
        if not layers:
            return warnings
        last_bottom = min(float(ly.bottom_elev_m) for ly in layers)
        if inp.pile_tip_elev_m < last_bottom - 1e-9 and not getattr(inp, "allow_geology_extrapolation", False):
            warnings.append(f"Mũi cọc ({inp.pile_tip_elev_m:g}) nằm dưới đáy lớp khảo sát cuối ({last_bottom:g}); không ngoại suy địa tầng. Cần bổ sung lớp địa chất hoặc bật option cho phép ngoại suy.")
        if inp.pile_tip_elev_m < last_bottom - 1e-9 and getattr(inp, "allow_geology_extrapolation", False):
            warnings.append(f"Đã cho phép ngoại suy lớp cuối từ {last_bottom:g} xuống mũi cọc {inp.pile_tip_elev_m:g}; cần kiểm chứng địa chất.")
        top = inp.ground_elev_m
        prev_bottom = top
        for ly in layers:
            if ly.bottom_elev_m > prev_bottom + 1e-9:
                warnings.append(f"Cao độ đáy lớp {ly.name or '?'} không giảm dần so với lớp phía trên; kiểm tra bảng địa chất.")
            prev_bottom = ly.bottom_elev_m

        # Cảnh báo thiếu số liệu địa chất nhập từ bảng. Các cờ _missing_* được gắn ở lớp GUI
        # khi đọc raw cell để không nhầm giá trị mặc định nội bộ là dữ liệu người dùng đã nhập.
        missing_gamma_layers: List[str] = []
        missing_spt_layers: List[str] = []
        missing_rock_layers: List[str] = []
        for idx, ly in enumerate(layers):
            st = int(getattr(ly, "soil_type", 0) or 0)
            lname = str(getattr(ly, "name", "") or f"L{idx+1}").strip() or f"L{idx+1}"
            if st != 0 and bool(getattr(ly, "_missing_gamma", False)):
                missing_gamma_layers.append(lname)
            if st in (3, 4) and bool(getattr(ly, "_missing_rock_data", False)):
                missing_rock_layers.append(lname)
            # Không cảnh báo thiếu SPT cho 1-2 lớp đầu, lớp 0 và lớp đá dùng bộ chỉ tiêu đá riêng.
            if idx >= 2 and st not in (0, 3, 4) and bool(getattr(ly, "_missing_spt", False)):
                missing_spt_layers.append(lname)
        if missing_rock_layers:
            warnings.append("Thiếu số liệu lớp đá: " + ", ".join(missing_rock_layers) + ".")
        if missing_gamma_layers:
            warnings.append("Thiếu trọng lượng riêng các lớp đất: " + ", ".join(missing_gamma_layers) + ".")
        if missing_spt_layers:
            warnings.append("Thiếu SPT của lớp " + ", ".join(missing_spt_layers) + ".")

        zone = SCTCalculator._downdrag_zone(inp)
        if getattr(inp, "include_downdrag", False):
            if zone is None:
                warnings.append("Đã bật ma sát âm nhưng chưa khai báo vùng hợp lệ; không tính downdrag.")
            else:
                ztop, zbot = zone
                if ztop <= zbot:
                    warnings.append("Vùng ma sát âm không hợp lệ.")
                if ztop > inp.cap_bottom_elev_m + 1e-9:
                    warnings.append("Đỉnh vùng ma sát âm nằm phía trên đáy bệ; chỉ xét phần giao với đoạn cọc.")
                if zbot < inp.pile_tip_elev_m - 1e-9:
                    warnings.append("Đáy vùng ma sát âm nằm dưới mũi cọc; chỉ xét phần giao với đoạn cọc.")
        return warnings

    @staticmethod
    def _shaft_phi(inp: PileInput, soil_type: int, limit_state: str = LIMIT_STATE_STRENGTH, n60: Optional[float] = None) -> float:
        st = int(soil_type)
        ls = str(limit_state).upper()
        if ls in (LIMIT_STATE_EXTREME, "DB", "ĐB", "EXTREME"):
            return PHI_EXTREME_DEFAULT
        if SCTCalculator._is_driven(inp):
            if st in (1, 6):
                return PHI_DRIVEN_SHAFT_SAND
            if st == 2:
                return PHI_DRIVEN_SHAFT_CLAY
            if st == 5:
                nv = float(n60) if n60 is not None else 100.0
                return PHI_BORE_SHAFT_IGM_SPT_LT100 if nv < 100.0 else PHI_BORE_SHAFT_IGM_SPT_GE100
            return PHI_ROCK_SIDE
        if st in (1, 6):
            return PHI_BORE_SHAFT_SAND
        if st == 2:
            return PHI_BORE_SHAFT_CLAY
        if st == 5:
            nv = float(n60) if n60 is not None else 100.0
            return PHI_BORE_SHAFT_IGM_SPT_LT100 if nv < 100.0 else PHI_BORE_SHAFT_IGM_SPT_GE100
        if st in (3, 4):
            return PHI_ROCK_SIDE
        return 0.0

    @staticmethod
    def _tip_phi(inp: PileInput, soil_type: int, limit_state: str = LIMIT_STATE_STRENGTH, n60: Optional[float] = None) -> float:
        st = int(soil_type)
        ls = str(limit_state).upper()
        if ls in (LIMIT_STATE_EXTREME, "DB", "ĐB", "EXTREME"):
            return PHI_EXTREME_DEFAULT
        mode_l = (inp.mode + " " + inp.pile_type).lower()
        if SCTCalculator._is_driven(inp):
            if st in (1, 6):
                return PHI_DRIVEN_TIP_SAND
            if st == 2:
                return PHI_DRIVEN_TIP_CLAY
            if st == 5:
                nv = float(n60) if n60 is not None else 100.0
                return PHI_BORE_SHAFT_IGM_SPT_LT100 if nv < 100.0 else PHI_BORE_SHAFT_IGM_SPT_GE100
            return PHI_ROCK_TIP
        if st in (1, 6):
            return PHI_BORE_TIP_SAND
        if st == 2:
            return PHI_BORE_TIP_CLAY
        if st == 5:
            nv = float(n60) if n60 is not None else 100.0
            return PHI_BORE_SHAFT_IGM_SPT_LT100 if nv < 100.0 else PHI_BORE_SHAFT_IGM_SPT_GE100
        if st in (3, 4):
            return PHI_ROCK_TIP
        return 0.0

    @staticmethod
    def pile_weight_effective_kn(inp: PileInput) -> Tuple[float, float, float]:
        """Trả về (W khô, W hữu hiệu, lực đẩy nổi) của bản thân cọc, kN.

        W hữu hiệu tính theo mực nước ngầm:
        - đoạn trên nước: gamma_c
        - đoạn dưới nước: gamma_c - gamma_w
        """
        A = inp.shaft_area_gross_m2
        top = inp.cap_bottom_elev_m
        bot = inp.pile_tip_elev_m
        L = max(top - bot, 0.0)
        gamma_c = max(inp.concrete_gamma_kN_m3, 0.0)
        w_dry = A * L * gamma_c
        if L <= 0:
            return 0.0, 0.0, 0.0
        water = inp.water_elev_m
        if water >= top:
            L_sub = L
        elif water <= bot:
            L_sub = 0.0
        else:
            L_sub = water - bot
        L_sub = max(min(L_sub, L), 0.0)
        L_above = L - L_sub
        gamma_sub = max(gamma_c - GAMMA_W_KN_M3, 0.0)
        w_eff = A * (L_above * gamma_c + L_sub * gamma_sub)
        return w_dry, w_eff, max(w_dry - w_eff, 0.0)

    @staticmethod
    def _sorted_layers(inp: PileInput) -> List[SoilLayer]:
        layers = [ly for ly in inp.layers if str(ly.name).strip() or abs(float(ly.bottom_elev_m)) > 1e-12]
        # Cao độ giảm theo chiều sâu: top lớn, bottom nhỏ.
        layers.sort(key=lambda ly: ly.bottom_elev_m, reverse=True)
        return layers

    @staticmethod
    def _unit_weight_effective(gamma_kN_m3: float, z_mid_elev: float, water_elev: float) -> float:
        # Nếu dưới mực nước, lấy gamma' = gamma - gamma_w. Nếu trên mực nước, lấy gamma.
        if z_mid_elev <= water_elev:
            return max(gamma_kN_m3 - GAMMA_W_KN_M3, 0.1)
        return gamma_kN_m3

    @staticmethod
    def _sigma_v_eff_at(inp: PileInput, elev_m: float, layer_tops: List[Tuple[float, SoilLayer]]) -> float:
        """Tích phân ứng suất hữu hiệu từ mặt đất/cao độ bắt đầu đến elev_m, trả MPa."""
        if elev_m >= inp.ground_elev_m:
            return 0.0
        stress_kpa = 0.0
        current_top = inp.ground_elev_m
        for bottom, ly in layer_tops:
            seg_top = current_top
            seg_bot = bottom
            if seg_bot >= seg_top:
                current_top = seg_bot
                continue
            # lấy phần từ seg_top xuống max(seg_bot, elev_m)
            a = seg_top
            b = max(seg_bot, elev_m)
            if a > b:
                mid = (a + b) / 2.0
                gamma_layer = 0.0 if int(getattr(ly, "soil_type", 0) or 0) == 0 else max(float(ly.gamma_kN_m3 or 0.0), 0.0)
                gamma_eff = SCTCalculator._unit_weight_effective(gamma_layer, mid, inp.water_elev_m) if gamma_layer > 0 else 0.0
                # nếu mực nước cắt giữa đoạn, chia đôi để chính xác hơn
                if b < inp.water_elev_m < a:
                    stress_kpa += max(a - inp.water_elev_m, 0.0) * gamma_layer
                    stress_kpa += max(inp.water_elev_m - b, 0.0) * max(gamma_layer - GAMMA_W_KN_M3, 0.0)
                else:
                    stress_kpa += (a - b) * gamma_eff
            if elev_m >= seg_bot:
                break
            current_top = seg_bot
        return max(stress_kpa / 1000.0, 0.0)

    @staticmethod
    def _build_layer_segments(inp: PileInput) -> List[Tuple[float, float, SoilLayer]]:
        layers = SCTCalculator._sorted_layers(inp)
        segments: List[Tuple[float, float, SoilLayer]] = []
        if not layers:
            return segments
        top = inp.ground_elev_m
        pile_top = inp.cap_bottom_elev_m
        pile_tip = inp.pile_tip_elev_m
        # Nếu đáy bệ/cọc nằm trên mặt đất, tự tạo đoạn không khí từ đáy bệ đến mặt đất.
        if pile_top > inp.ground_elev_m + 1e-9:
            air = SoilLayer(name="Không khí", bottom_elev_m=inp.ground_elev_m, soil_type=0, n_spt=0.0, gamma_kN_m3=0.0)
            segments.append((pile_top, inp.ground_elev_m, air))
        for ly in layers:
            bottom = ly.bottom_elev_m
            seg_top = min(top, pile_top)
            seg_bottom = max(bottom, pile_tip)
            # đoạn giao giữa lớp đất và đoạn cọc
            inter_top = min(top, pile_top)
            inter_bottom = max(bottom, pile_tip)
            if inter_top > inter_bottom:
                segments.append((inter_top, inter_bottom, ly))
            top = bottom
            if bottom <= pile_tip:
                break
        # Không tự kéo lớp cuối xuống mũi nếu chưa cho phép ngoại suy.
        if segments and segments[-1][1] > pile_tip and getattr(inp, "allow_geology_extrapolation", False):
            last_top, last_bottom, last_ly = segments[-1]
            segments[-1] = (last_top, pile_tip, last_ly)
        return segments

    @staticmethod
    def _layer_tops(inp: PileInput) -> List[Tuple[float, SoilLayer]]:
        layers = SCTCalculator._sorted_layers(inp)
        return [(ly.bottom_elev_m, ly) for ly in layers]

    @staticmethod
    def _n60_from_input(inp: PileInput, layer: SoilLayer) -> float:
        """Đổi SPT nhập trong bảng địa chất sang N60.

        Mặc định bảng địa chất nhập Nₕₜ/N đo hiện trường và đổi theo ER:
            N60 = Nₕₜ * ER / 60

        Để đối chiếu với các bản cũ hoặc dữ liệu đã hiệu chỉnh, có thể chọn chế độ
        "N60" trong tab Thông số chung; khi đó số nhập ở cột SPT được dùng trực tiếp
        là N60 và không nhân lại hệ số ER.
        """
        n_in = max(float(getattr(layer, "n_spt", 0.0) or 0.0), 0.0)
        mode = str(getattr(inp, "spt_input_mode", "Nₕₜ") or "Nₕₜ").upper().replace(" ", "")
        if "N60" in mode:
            return n_in
        er = max(float(getattr(inp, "spt_er_percent", 60.0) or 60.0), 0.0)
        if er <= 0.0:
            er = 60.0
        return n_in * er / 60.0

    @staticmethod
    def _spt_note(inp: PileInput, n_input: float, n60: float) -> str:
        mode = str(getattr(inp, "spt_input_mode", "Nₕₜ") or "Nₕₜ").upper().replace(" ", "")
        if "N60" in mode:
            return f"SPT nhập đã là N60={n60:.2f}"
        return f"Nₕₜ={n_input:.2f}, ER={getattr(inp, 'spt_er_percent', 60.0):.0f}%, N60={n60:.2f}"

    @staticmethod
    def _rock_supported_construction(inp: PileInput) -> bool:
        txt = _strip_accents(str(getattr(inp, "rock_construction_condition", "") or "")).lower()
        if not txt:
            legacy = str(getattr(inp, "rock_side_method", "fractured") or "fractured").lower()
            return not legacy.startswith("normal")
        return "khong" not in txt

    @staticmethod
    def _rock_open_joint(inp: PileInput) -> bool:
        txt = _strip_accents(str(getattr(inp, "rock_joint_condition", "") or "")).lower()
        if txt:
            return not ("khep" in txt and "kin" in txt)
        return bool(getattr(inp, "rock_open_joint", True))

    @staticmethod
    def _clay_su_for_layer(inp: PileInput, ly: SoilLayer) -> float:
        """Su của lớp sét dùng cho các nhánh alpha/mũi cọc, MPa."""
        c_mpa = ly.c_mpa if ly.c_mpa > 0 else ly.su_mpa
        n60 = SCTCalculator._n60_from_input(inp, ly)
        return max(c_mpa, 0.0) if bool(getattr(inp, "clay_use_c_phi", False)) else 0.006 * max(n60, 0.0)

    @staticmethod
    def _tomlinson_cover_case(inp: PileInput, ly: SoilLayer) -> str:
        """Chọn bảng α Tomlinson theo lớp nằm ngay phía trên lớp sét đang xét."""
        layers = SCTCalculator._sorted_layers(inp)
        idx = None
        for i, item in enumerate(layers):
            if item is ly:
                idx = i
                break
        # Trường hợp đối tượng bị copy/tái tạo, dùng cao độ + tên để tìm gần đúng.
        if idx is None:
            for i, item in enumerate(layers):
                if str(item.name).strip() == str(ly.name).strip() and abs(float(item.bottom_elev_m) - float(ly.bottom_elev_m)) < 1e-7:
                    idx = i
                    break
        above = layers[idx - 1] if idx is not None and idx > 0 else None
        def clay_case_by_su_n60(layer: SoilLayer) -> str:
            su_val = SCTCalculator._clay_su_for_layer(inp, layer)
            n60_val = SCTCalculator._n60_from_input(inp, layer)
            case_su = "soft_clay" if su_val < 0.05 else "stiff_clay"
            case_n = "soft_clay" if n60_val < 8.0 else "stiff_clay"
            if case_su != case_n:
                return "clay_worst"
            return case_su

        if above is not None:
            st_above = int(getattr(above, "soil_type", 0) or 0)
            if st_above in (1, 6):
                return "sand_gravel"
            if st_above == 2:
                return clay_case_by_su_n60(above)
        # Nếu không xác định được lớp phía trên, dùng chính trạng thái lớp sét đang xét.
        return clay_case_by_su_n60(ly)

    @staticmethod
    def _tomlinson_db_over_d(inp: PileInput, elev_m: float) -> float:
        """Db/D theo độ sâu điểm xét dưới cao độ thiên nhiên/mặt đất."""
        return max(float(inp.ground_elev_m) - float(elev_m), 0.0) / max(float(inp.diameter_m), 1e-9)

    @staticmethod
    def _skin_resistance(inp: PileInput) -> Tuple[List[LayerCalc], float, float, float, List[str]]:
        warnings: List[str] = []
        driven_igm_rock_warned: set = set()
        segments = SCTCalculator._build_layer_segments(inp)
        layer_tops = SCTCalculator._layer_tops(inp)
        result_layers: List[LayerCalc] = []
        qsf = 0.0
        qsf_ex = 0.0
        qsn = 0.0
        mode_l = (inp.mode + " " + inp.pile_type).lower()
        is_driven = SCTCalculator._is_driven(inp)
        exclude_top = inp.exclude_top_bored_m if (("khoan" in mode_l) and not is_driven) else 0.0
        skin_start_elev = inp.cap_bottom_elev_m - max(exclude_top, 0.0)
        dd_zone = SCTCalculator._downdrag_zone(inp)
        dd_factor = max(float(getattr(inp, "downdrag_factor", 0.0) or SCTCalculator.default_downdrag_load_factor(inp)), 0.0)

        for top, bottom, ly in segments:
            th = max(top - bottom, 0.0)
            # Chiều dài ma sát hình học sau khi bỏ đoạn đầu cọc khoan.
            skin_top = min(top, skin_start_elev)
            skin_bottom = bottom
            skin_full_len = max(skin_top - skin_bottom, 0.0)
            dd_len = 0.0
            if dd_zone and skin_full_len > 1e-9:
                dd_top, dd_bot = dd_zone
                dd_len = _interval_overlap_len(skin_top, skin_bottom, dd_top, dd_bot)
            # Vùng ma sát âm không được tính là ma sát dương chịu nén; đồng thời không lấy vào nhổ để bảo thủ.
            skin_len = max(skin_full_len - dd_len, 0.0)
            mid_elev = (skin_top + skin_bottom) / 2.0 if skin_full_len > 0 else (top + bottom) / 2.0
            sigma_v = SCTCalculator._sigma_v_eff_at(inp, mid_elev, layer_tops)
            gamma_eff = 0.0 if int(getattr(ly, "soil_type", 0) or 0) == 0 else SCTCalculator._unit_weight_effective(ly.gamma_kN_m3, mid_elev, inp.water_elev_m)
            n1_input = max(ly.n_spt, 0.0)
            n60 = SCTCalculator._n60_from_input(inp, ly)
            st = int(ly.soil_type)
            if st == 0:
                skin_len = 0.0
                dd_len = 0.0
            is_sand = st in (1, 6)
            is_clay = st == 2
            is_intact_rock = st == 3
            is_fractured_rock = st == 4
            is_igm = st == 5
            n1 = n1_60_corrected(n60, sigma_v, cap_to_n=False) if is_sand else n60
            c_mpa = ly.c_mpa if ly.c_mpa > 0 else ly.su_mpa
            su = 0.006 * n60 if not inp.clay_use_c_phi else max(c_mpa, 0.0)
            alpha_beta = 0.0
            phi_deg = ly.phi_deg
            qs_nominal_kpa = 0.0
            note = ""

            if skin_full_len <= 1e-9 or st == 0:
                note = "Không tính ma sát thành bên"
            elif is_driven:
                if is_sand:
                    qs_nominal_kpa, note = driven_sand_side_nominal_kpa(n1, inp.pile_type)
                    alpha_beta = 0.0
                    note += f"; {SCTCalculator._spt_note(inp, n1_input, n60)}; N160={n1:.2f}"
                elif is_clay:
                    if inp.clay_use_c_phi:
                        qs_nominal_kpa, alpha_beta, note = clay_c_phi_side_nominal_kpa(c_mpa, ly.phi_deg, sigma_v)
                        note = "Cọc đóng/ép - " + note
                    else:
                        db_over_d = SCTCalculator._tomlinson_db_over_d(inp, mid_elev)
                        cover_case = SCTCalculator._tomlinson_cover_case(inp, ly)
                        alpha_beta = alpha_clay_driven_tomlinson(su, db_over_d, cover_case)
                        qs_nominal_kpa = alpha_beta * su * 1000.0
                        case_label = _TOMLINSON_CASE_LABELS.get(cover_case, cover_case)
                        note = f"Cọc đóng/ép - đất dính: qs=αSu; α Tomlinson nội suy theo {case_label}, Su={su:.3f}MPa, Db/D={db_over_d:.2f}; " + SCTCalculator._spt_note(inp, n1_input, n60)
                elif is_igm:
                    qs_nominal_kpa, note = igm_side_nominal_kpa(n60, ly.qu_mpa, inp.igm_alpha, inp.igm_joint_factor, str(getattr(inp, "igm_missing_qu_policy", "require_qu")).lower() == "use_spt")
                    alpha_beta = inp.igm_alpha
                    note = "Cọc đóng/ép - " + note
                    # QA fix P3: công thức IGM xây dựng cho cọc khoan nhồi; cảnh báo khi áp cho cọc đóng/ép.
                    if ly.name not in driven_igm_rock_warned:
                        driven_igm_rock_warned.add(ly.name)
                        warnings.append(f"Lớp {ly.name or '?'}: cọc đóng/ép trong IGM dùng công thức của cọc khoan (ngoài phạm vi phương pháp); nên kiểm soát bằng độ chối/thử động khi thi công.")
                else:
                    qs_nominal_kpa, alpha_beta, note = rock_side_nominal_kpa(ly.qu_mpa, inp.fc_mpa, ly.rqd, supported=SCTCalculator._rock_supported_construction(inp), open_joint=SCTCalculator._rock_open_joint(inp))
                    note = "Cọc đóng/ép trong đá - " + note
                    if ly.name not in driven_igm_rock_warned:
                        driven_igm_rock_warned.add(ly.name)
                        warnings.append(f"Lớp {ly.name or '?'}: cọc đóng/ép trong đá dùng công thức thành bên của cọc khoan (ngoài phạm vi phương pháp); nên kiểm soát bằng độ chối/thử động khi thi công.")
            else:
                if is_sand:
                    mode = "gravel" if st == 6 else inp.sand_preconsolidation_mode
                    alpha_beta, phi_calc, sigma_p = beta_sand_bored(n60, n1, sigma_v, mode, inp.sand_m)
                    phi_deg = phi_calc
                    qs_nominal_kpa = alpha_beta * sigma_v * 1000.0
                    note = f"Cọc khoan - đất rời: {SCTCalculator._spt_note(inp, n1_input, n60)}; σ'p dùng N60, φ dùng N160={n1:.2f}; σ'p={sigma_p:.3f} MPa"
                elif is_clay:
                    if inp.clay_use_c_phi:
                        qs_nominal_kpa, alpha_beta, note = clay_c_phi_side_nominal_kpa(c_mpa, ly.phi_deg, sigma_v)
                        note = "Cọc khoan - " + note
                    else:
                        warn = clay_alpha_domain_warning(ly.name, su)
                        if warn:
                            warnings.append(warn)
                        alpha_beta = alpha_clay_bored(su)
                        qs_nominal_kpa = alpha_beta * su * 1000.0
                        note = "Cọc khoan - đất dính: αSu, " + SCTCalculator._spt_note(inp, n1_input, n60)
                elif is_igm:
                    qs_nominal_kpa, note = igm_side_nominal_kpa(n60, ly.qu_mpa, inp.igm_alpha, inp.igm_joint_factor, str(getattr(inp, "igm_missing_qu_policy", "require_qu")).lower() == "use_spt")
                    alpha_beta = inp.igm_alpha
                    note = "Cọc khoan - " + note
                elif is_intact_rock or is_fractured_rock:
                    if ly.qu_mpa <= 0:
                        warnings.append(f"Lớp {ly.name}: chưa nhập qu đá, không tính đúng ma sát thành bên đá.")
                    qs_nominal_kpa, alpha_beta, note = rock_side_nominal_kpa(
                        ly.qu_mpa, inp.fc_mpa, ly.rqd,
                        supported=SCTCalculator._rock_supported_construction(inp),
                        open_joint=SCTCalculator._rock_open_joint(inp),
                    )
                    note = "Cọc khoan trong đá: " + note
                else:
                    note = "Loại đất chưa hỗ trợ rõ; không tính ma sát thành bên"

            phi_cd = SCTCalculator._shaft_phi(inp, st, LIMIT_STATE_STRENGTH, n60=n60)
            phi_db = SCTCalculator._shaft_phi(inp, st, LIMIT_STATE_EXTREME, n60=n60)
            phi_up_cd = SCTCalculator._uplift_phi(inp, st, LIMIT_STATE_STRENGTH)
            phi_up_db = SCTCalculator._uplift_phi(inp, st, LIMIT_STATE_EXTREME)
            qs_factored_kpa = qs_nominal_kpa * phi_cd
            qs_extreme_kpa = qs_nominal_kpa * phi_db
            # Sức kháng nhổ dùng Rs danh định và hệ số φup riêng; không thêm hệ số giảm Rs để tránh double reduction.
            qs_uplift_cd_kpa = qs_nominal_kpa * phi_up_cd
            qs_uplift_db_kpa = qs_nominal_kpa * phi_up_db
            qf = inp.perimeter_m * qs_factored_kpa * skin_len
            qex = inp.perimeter_m * qs_extreme_kpa * skin_len
            qn = inp.perimeter_m * qs_nominal_kpa * skin_len
            qup_cd = inp.perimeter_m * qs_uplift_cd_kpa * skin_len
            qup_db = inp.perimeter_m * qs_uplift_db_kpa * skin_len
            dd_kn = inp.perimeter_m * qs_nominal_kpa * dd_len * dd_factor
            qsf += qf
            qsf_ex += qex
            qsn += qn
            if dd_len > 1e-9:
                note += f"; ma sát âm L={dd_len:.2f}m, γDD={dd_factor:.2f}, DD={dd_kn:.2f}kN; không tính ma sát dương trong vùng này"
            result_layers.append(LayerCalc(
                name=ly.name, top_elev_m=top, bottom_elev_m=bottom, thickness_m=th, skin_length_m=skin_len,
                soil_type=ly.soil_type, soil_label=SOIL_TYPE_LABELS.get(ly.soil_type, "?"), n_spt=ly.n_spt,
                n60=n60, n1_60=n1, gamma_eff_kN_m3=gamma_eff, sigma_v_eff_mpa=sigma_v, c_mpa=c_mpa, su_mpa=su,
                alpha_or_beta=alpha_beta, phi_deg=phi_deg, qs_factored_kpa=qs_factored_kpa,
                qs_nominal_kpa=qs_nominal_kpa, qs_factored_kn=qf, qs_nominal_kn=qn,
                qs_extreme_kpa=qs_extreme_kpa, qs_extreme_kn=qex, phi_strength=phi_cd, phi_extreme=phi_db, qu_mpa=ly.qu_mpa, rqd=ly.rqd, note=note,
                qs_uplift_factored_kpa=qs_uplift_cd_kpa, qs_uplift_factored_kn=qup_cd,
                qs_uplift_extreme_kpa=qs_uplift_db_kpa, qs_uplift_extreme_kn=qup_db,
                downdrag_length_m=dd_len, downdrag_kn=dd_kn,
            ))
        return result_layers, qsf, qsf_ex, qsn, warnings

    @staticmethod
    def _toe_layer(inp: PileInput) -> Optional[SoilLayer]:
        layers = SCTCalculator._sorted_layers(inp)
        if not layers:
            return None
        top = inp.ground_elev_m
        for ly in layers:
            if top >= inp.pile_tip_elev_m >= ly.bottom_elev_m:
                return ly
            top = ly.bottom_elev_m
        if getattr(inp, "allow_geology_extrapolation", False):
            return layers[-1]
        return None


    @staticmethod
    def _toe_layer_top_and_embedment(inp: PileInput, toe: SoilLayer) -> Tuple[float, float]:
        """Trả về cao độ đỉnh lớp mũi và chiều dài mũi cọc ngập trong lớp đó."""
        top = float(getattr(inp, "ground_elev_m", 0.0) or 0.0)
        for ly in SCTCalculator._sorted_layers(inp):
            bottom = float(getattr(ly, "bottom_elev_m", top) or top)
            same_obj = ly is toe
            same_val = (str(getattr(ly, "name", "")).strip() == str(getattr(toe, "name", "")).strip() and abs(bottom - float(getattr(toe, "bottom_elev_m", bottom) or bottom)) < 1e-7)
            if same_obj or same_val:
                return top, max(top - float(inp.pile_tip_elev_m), 0.0)
            top = bottom
        return float(getattr(inp, "ground_elev_m", 0.0) or 0.0), 0.0


    @staticmethod
    def _rock_tip_geology_conditions(inp: PileInput) -> Tuple[Dict[str, Any], List[str]]:
        """Kiểm tra điều kiện tự động để dùng qp=2.5qu cho mũi đá tốt.

        Điều kiện dùng 2.5qu theo TCVN được quy đổi như sau:
        - mũi nằm trong lớp loại 3: Đá nguyên khối / đá tốt;
        - điều kiện khe nứt global là khe nứt khép kín;
        - điều kiện thi công global là không chống đỡ / thành hố ổn định;
        - trong phạm vi 2B dưới mũi, địa tầng đều là đá tốt loại 3;
        - chiều sâu ngàm liên tục vào lớp đá tốt tại mũi > 1.5B.
        Nếu một trong các điều kiện này không thỏa, mũi đá dùng Hoek-Brown/GSI.
        """
        warnings: List[str] = []
        D = max(float(inp.diameter_m), 1e-9)
        tip = float(inp.pile_tip_elev_m)
        check_bot = tip - 2.0 * D
        layers = SCTCalculator._sorted_layers(inp)
        ctx: Dict[str, Any] = {
            "joint_ok": not SCTCalculator._rock_open_joint(inp),
            "construction_ok": not SCTCalculator._rock_supported_construction(inp),
            "good_rock_2b": False,
            "socket_depth_good_m": 0.0,
            "socket_ok": False,
            "toe_good_rock": False,
            "coverage_2b_m": 0.0,
            "required_2b_m": 2.0 * D,
            "required_socket_m": 1.5 * D,
            "use_25qu": False,
        }
        if not layers:
            warnings.append("Không có địa tầng để kiểm tra điều kiện Qp=2.5qu cho mũi đá.")
            return ctx, warnings

        top = inp.ground_elev_m
        toe_top = None
        toe_layer = None
        for ly in layers:
            bottom = ly.bottom_elev_m
            if top >= tip >= bottom:
                toe_top = top
                toe_layer = ly
                break
            top = bottom
        if toe_layer is None and getattr(inp, "allow_geology_extrapolation", False):
            # Chỉ dùng cho cảnh báo; bản chất vẫn không đủ để khẳng định 2B dưới mũi.
            toe_layer = layers[-1]
            toe_top = layers[-2].bottom_elev_m if len(layers) >= 2 else inp.ground_elev_m
        if toe_layer is None:
            warnings.append("Không xác định được lớp đá tại mũi để kiểm tra điều kiện Qp=2.5qu.")
            return ctx, warnings

        ctx["toe_good_rock"] = int(getattr(toe_layer, "soil_type", 0) or 0) == 3
        if ctx["toe_good_rock"] and toe_top is not None:
            ctx["socket_depth_good_m"] = max(float(toe_top) - tip, 0.0)
            ctx["socket_ok"] = ctx["socket_depth_good_m"] > 1.5 * D + 1e-9

        # Kiểm tra đoạn 2B dưới mũi có phủ đủ địa tầng và toàn bộ là đá tốt loại 3 hay không.
        coverage = 0.0
        has_bad = False
        top = inp.ground_elev_m
        for ly in layers:
            bottom = ly.bottom_elev_m
            ov = _interval_overlap_len(top, bottom, tip, check_bot)
            if ov > 1e-9:
                coverage += ov
                if int(getattr(ly, "soil_type", 0) or 0) != 3:
                    has_bad = True
            top = bottom
            if bottom <= check_bot:
                break
        ctx["coverage_2b_m"] = coverage
        ctx["good_rock_2b"] = (coverage >= 2.0 * D - 1e-8) and not has_bad

        if not ctx["joint_ok"]:
            warnings.append("Không dùng Qp=2.5qu: điều kiện khe nứt đá không phải khe nứt khép kín/không chèn lấp.")
        if not ctx["construction_ok"]:
            warnings.append("Không dùng Qp=2.5qu: điều kiện thi công đang là có chống đỡ/không tương ứng thành hố đá ổn định.")
        if not ctx["toe_good_rock"]:
            warnings.append("Không dùng Qp=2.5qu: mũi không nằm trong lớp đá tốt loại 3.")
        if not ctx["good_rock_2b"]:
            warnings.append(f"Không dùng Qp=2.5qu: chưa chứng minh đá tốt liên tục trong 2B dưới mũi; cần {2.0*D:.2f} m, phủ {coverage:.2f} m.")
        if not ctx["socket_ok"]:
            warnings.append(f"Không dùng Qp=2.5qu: chiều sâu ngàm vào đá tốt {ctx['socket_depth_good_m']:.2f} m ≤ 1.5B={1.5*D:.2f} m.")

        ctx["use_25qu"] = bool(ctx["joint_ok"] and ctx["construction_ok"] and ctx["toe_good_rock"] and ctx["good_rock_2b"] and ctx["socket_ok"])
        return ctx, warnings

    @staticmethod
    def _rock_hb_tip_qp_mpa(inp: PileInput, toe: SoilLayer, sigma_tip_mpa: float) -> Tuple[float, Dict[str, Any], str, List[str]]:
        """Tính Qp mũi đá theo Hoek-Brown và giới hạn qp <= 2.5qu nếu chưa có thử tải."""
        warnings: List[str] = []
        qu = max(float(getattr(toe, "qu_mpa", 0.0) or 0.0), 0.0)
        if qu <= 0.0:
            return 0.0, {"qp_mpa": 0.0, "qp_uncapped_mpa": 0.0, "qp_limit_25qu_mpa": 0.0, "hb_capped": False}, "Thiếu qu đá; không tính Qp Hoek-Brown", ["Mũi đá thiếu qu, không tính được Qp theo Hoek-Brown."]
        qp_raw, hb = hoek_brown_qp(sigma_tip_mpa, qu, toe.mi, toe.gsi, toe.disturbance)
        qp_limit = 2.5 * qu
        allow_exceed = bool(getattr(inp, "allow_rock_tip_exceed_25qu", False))
        qp = qp_raw if allow_exceed else min(qp_raw, qp_limit)
        capped = (not allow_exceed) and qp_raw > qp_limit + 1e-9
        if capped:
            warnings.append(f"Qp Hoek-Brown={qp_raw:.3f} MPa > 2.5qu={qp_limit:.3f} MPa; đã giới hạn theo 2.5qu.")
        if allow_exceed and qp_raw > qp_limit + 1e-9:
            warnings.append("Đã bật cho phép Qp đá vượt 2.5qu; cần có thử tải/kinh nghiệm địa phương để bảo vệ giá trị này.")
        hb.update({"qp_mpa": qp, "qp_uncapped_mpa": qp_raw, "qp_limit_25qu_mpa": qp_limit, "hb_capped": capped, "allow_exceed_25qu": allow_exceed})
        note = "Mũi đá dùng Hoek-Brown/GSI"
        if capped:
            note += ", giới hạn Qp≤2.5qu"
        elif allow_exceed:
            note += ", cho phép vượt 2.5qu theo lựa chọn"
        return qp, hb, note, warnings

    @staticmethod
    def _weak_clay_below_tip_2d(inp: PileInput) -> Tuple[bool, List[str]]:
        """Kiểm tra Su<0.024 MPa trong phạm vi 2D dưới mũi cọc đất dính."""
        warnings: List[str] = []
        layers = SCTCalculator._sorted_layers(inp)
        if not layers:
            return False, warnings
        tip = inp.pile_tip_elev_m
        check_bot = tip - 2.0 * inp.diameter_m
        if min(ly.bottom_elev_m for ly in layers) > check_bot + 1e-9:
            warnings.append(f"Địa chất chưa phủ đủ 2D dưới mũi cọc đất dính: cần tới cao độ {check_bot:g} để kiểm tra Su<0.024 MPa.")
        top = inp.ground_elev_m
        weak = False
        for ly in layers:
            bottom = ly.bottom_elev_m
            if _interval_overlap_len(top, bottom, tip, check_bot) > 1e-9 and int(ly.soil_type or 0) == 2:
                c_mpa = ly.c_mpa if ly.c_mpa > 0 else ly.su_mpa
                n60 = SCTCalculator._n60_from_input(inp, ly)
                su = 0.006 * n60 if not inp.clay_use_c_phi else max(c_mpa, 0.0)
                if su < 0.024:
                    weak = True
                    warnings.append(f"Lớp {ly.name or '?'} trong phạm vi 2D dưới mũi có Su={su:.4f} MPa < 0.024 MPa; giảm Nc còn 0.67Nc.")
            top = bottom
            if bottom <= check_bot:
                break
        return weak, warnings

    @staticmethod
    def _tip_resistance(inp: PileInput, layer_results: List[LayerCalc]) -> Tuple[float, float, float, Dict[str, Any], List[str]]:
        warnings: List[str] = []
        toe = SCTCalculator._toe_layer(inp)
        if toe is None:
            return 0.0, 0.0, 0.0, {}, ["Chưa có lớp đất/đá tại mũi cọc."]
        layer_tops = SCTCalculator._layer_tops(inp)
        sigma_tip = SCTCalculator._sigma_v_eff_at(inp, inp.pile_tip_elev_m, layer_tops)
        L = inp.pile_length_m
        D = inp.diameter_m
        Ap = inp.area_m2
        c_mpa = toe.c_mpa if toe.c_mpa > 0 else toe.su_mpa
        n_toe60 = SCTCalculator._n60_from_input(inp, toe)
        n_toe160 = n1_60_corrected(n_toe60, sigma_tip, cap_to_n=False)
        su = 0.006 * n_toe60 if not inp.clay_use_c_phi else max(c_mpa, 0.0)
        st = toe.soil_type
        mode_l = (inp.mode + " " + inp.pile_type).lower()
        tip_info: Dict[str, Any] = {"toe_layer": toe.name, "soil_type": st, "sigma_tip_mpa": sigma_tip}
        qpf_kpa = 0.0
        qpn_kpa = 0.0
        note = ""

        is_sand = st in (1, 6)
        is_clay = st == 2
        is_intact_rock = st == 3
        is_fractured_rock = st == 4
        is_igm = st == 5
        if st == 0:
            warnings.append("Mũi cọc nằm trong lớp Không khí/Hang karst; không tính Qp và cần kiểm tra lại cao độ mũi/địa chất.")

        if is_intact_rock or is_fractured_rock:
            if not inp.include_rock_tip:
                qpn_kpa = 0.0
                qpf_kpa = 0.0
                note = "Mũi đặt trong đá: Bỏ qua sức kháng mũi, chỉ lấy sức kháng thân"
                tip_info.update({"rock_tip_case": "Bỏ qua sức kháng mũi đá", "qp_mpa": 0.0, "phi_qp": 0.0})
            else:
                # V0.2.32: chỉ dùng qp=2.5qu khi tự kiểm được điều kiện đá tốt/khe nứt/thi công/2B/ngàm 1.5B.
                # Nếu không thỏa, kể cả mũi đang nằm trong loại 3, chuyển sang Hoek-Brown/GSI và khống chế 2.5qu.
                if is_intact_rock:
                    rock_ctx, wctx = SCTCalculator._rock_tip_geology_conditions(inp)
                    tip_info.update({"rock_tip_conditions": rock_ctx})
                    if rock_ctx.get("use_25qu"):
                        qp_mpa = 2.5 * max(toe.qu_mpa, 0.0)
                        qpn_kpa = qp_mpa * 1000.0
                        qpf_kpa = qpn_kpa * PHI_ROCK_TIP
                        note = "Đá tốt: thỏa khe nứt/thi công/2B dưới mũi/ngàm >1.5B -> Qp=2.5qu"
                        tip_info.update({"rock_tip_case": "Đá tốt thỏa điều kiện 2.5qu", "qp_mpa": qp_mpa, "qp_limit_25qu_mpa": qp_mpa, "phi_qp": PHI_ROCK_TIP})
                    else:
                        warnings.extend(wctx)
                        qp_mpa, hb, hb_note, whb = SCTCalculator._rock_hb_tip_qp_mpa(inp, toe, sigma_tip)
                        warnings.extend(whb)
                        qpn_kpa = qp_mpa * 1000.0
                        qpf_kpa = qpn_kpa * PHI_ROCK_TIP
                        note = "Đá loại 3 nhưng không đủ điều kiện 2.5qu: " + hb_note
                        tip_info.update(hb)
                        tip_info.update({"rock_tip_case": "Hoek-Brown do không đủ điều kiện 2.5qu", "phi_qp": PHI_ROCK_TIP})
                else:
                    qp_mpa, hb, hb_note, whb = SCTCalculator._rock_hb_tip_qp_mpa(inp, toe, sigma_tip)
                    warnings.extend(whb)
                    qpn_kpa = qp_mpa * 1000.0
                    qpf_kpa = qpn_kpa * PHI_ROCK_TIP
                    note = "Đá nứt vỡ/phong hóa: " + hb_note
                    tip_info.update(hb)
                    tip_info.update({"rock_tip_case": "Đá nứt vỡ/phong hóa Hoek-Brown", "phi_qp": PHI_ROCK_TIP})
        elif is_igm:
            qu = toe.qu_mpa
            socket_depth = inp.rock_socket_depth_m if inp.rock_socket_depth_m is not None else max(0.0, inp.pile_tip_elev_m - toe.bottom_elev_m)
            qpn_kpa, extra, note = igm_tip_nominal_kpa(n_toe60, qu, inp.crack_spacing_mm, inp.crack_width_mm, inp.diameter_m, socket_depth, str(getattr(inp, "igm_missing_qu_policy", "require_qu")).lower() == "use_spt")
            phi_qp = SCTCalculator._tip_phi(inp, 5, LIMIT_STATE_STRENGTH, n_toe60)
            qpf_kpa = qpn_kpa * phi_qp
            tip_info.update(extra)
            tip_info.update({"qp_mpa": qpn_kpa / 1000.0, "phi_qp": phi_qp, "N60": n_toe60})
        elif is_sand:
            if SCTCalculator._is_driven(inp):
                toe_top, db_embed_m = SCTCalculator._toe_layer_top_and_embedment(inp, toe)
                db_over_d = db_embed_m / max(inp.diameter_m, 1e-9)
                soil_text = f"{getattr(toe, 'name', '')} {getattr(toe, 'comment', '')}"
                qpn_kpa, extra, note = driven_sand_tip_nominal_kpa(n_toe160, db_over_d, soil_text)
                phi_qp = PHI_DRIVEN_TIP_SAND
                tip_info.update(extra)
                tip_info.update({"N60": n_toe60, "N160": n_toe160, "toe_layer_top_m": toe_top, "toe_embedment_m": db_embed_m, "qp_mpa": qpn_kpa / 1000.0, "phi_qp": phi_qp})
            else:
                n_toe = n_toe60
                # Cọc khoan trong đất rời: qp=0.057N60, giới hạn 3.0MPa nếu không có thử tải.
                # QA fix P2: công thức chỉ có hiệu lực với N60 <= 50 theo TCVN 11823-10.
                if n_toe > 50.0:
                    warnings.append(f"Mũi cọc khoan trong đất rời có N60={n_toe:.0f} > 50: công thức qp=0.057N60 chỉ hiệu lực với N60<=50; xem xét khai lớp mũi là IGM (loại 5).")
                qp_mpa = min(0.057 * n_toe, 3.0)
                phi_qp = PHI_BORE_TIP_SAND
                qpn_kpa = qp_mpa * 1000.0
                note = "Cọc khoan đất rời/cuội sỏi: qp=min(0.057N60,3.0MPa)"
                tip_info.update({"N60": n_toe, "qp_mpa": qp_mpa, "phi_qp": phi_qp})
            qpf_kpa = qpn_kpa * phi_qp
        elif is_clay:
            driven_tip = SCTCalculator._is_driven(inp)
            phi_qp = PHI_DRIVEN_TIP_CLAY if driven_tip else PHI_BORE_TIP_CLAY
            if driven_tip and not inp.clay_use_c_phi:
                qp_mpa = 9.0 * su
                qpn_kpa = qp_mpa * 1000.0
                note = "Cọc đóng/ép - đất dính: qp=9Su"
                tip_info.update({"Nc": 9.0, "Su_mpa": su, "qp_uncapped_mpa": qp_mpa, "qp_mpa": qp_mpa, "phi_qp": phi_qp})
            elif inp.clay_use_c_phi:
                qpn_kpa, extra, note = clay_c_phi_tip_nominal_kpa(c_mpa, toe.phi_deg, sigma_tip)
                tip_info.update(extra)
                qp_mpa = qpn_kpa / 1000.0
            else:
                Nc0 = min(9.0, 6.0 * (1.0 + 0.2 * L / D))
                weak_2d, wweak = SCTCalculator._weak_clay_below_tip_2d(inp)
                warnings.extend(wweak)
                Nc = Nc0 * (2.0 / 3.0 if weak_2d else 1.0)
                qp_uncapped_mpa = Nc * su
                qp_mpa = min(qp_uncapped_mpa, 4.0)
                qpn_kpa = qp_mpa * 1000.0
                note = "Cọc khoan - đất dính: qp=Nc·Su ≤ 4.0MPa; kiểm Su<0.024MPa trong phạm vi 2D dưới mũi"
                tip_info.update({"Nc": Nc, "Nc0": Nc0, "Su_mpa": su, "weak_clay_2D": weak_2d, "qp_uncapped_mpa": qp_uncapped_mpa, "qp_cap_mpa": 4.0, "qp_capped": qp_uncapped_mpa > 4.0 + 1e-9})
                if qp_uncapped_mpa > 4.0 + 1e-9:
                    warnings.append(f"Qp đất dính Nc·Su={qp_uncapped_mpa:.3f} MPa > 4.0 MPa; đã giới hạn theo TCVN 11823-10.")
            qpf_kpa = qpn_kpa * phi_qp
            tip_info.update({"qp_mpa": qpn_kpa / 1000.0, "phi_qp": phi_qp})
        else:
            note = "Không có sức kháng mũi"

        phi_cd = SCTCalculator._tip_phi(inp, st, LIMIT_STATE_STRENGTH, n_toe60)
        phi_db = SCTCalculator._tip_phi(inp, st, LIMIT_STATE_EXTREME, n_toe60)
        if (is_intact_rock or is_fractured_rock) and not inp.include_rock_tip:
            phi_cd = 0.0
            phi_db = 0.0
        qpf_kpa = qpn_kpa * phi_cd
        qpf_ex_kpa = qpn_kpa * phi_db
        qpf_kn = qpf_kpa * Ap
        qpf_ex_kn = qpf_ex_kpa * Ap
        qpn_kn = qpn_kpa * Ap
        tip_info.update({
            "note": note, "qpf_kpa": qpf_kpa, "qpf_extreme_kpa": qpf_ex_kpa, "qpn_kpa": qpn_kpa,
            "Qpf_kn": qpf_kn, "Qpf_extreme_kn": qpf_ex_kn, "Qpn_kn": qpn_kn,
            "phi_qp_strength": phi_cd, "phi_qp_extreme": phi_db,
        })
        return qpf_kn, qpf_ex_kn, qpn_kn, tip_info, warnings

    @staticmethod
    def calculate(inp: PileInput) -> CapacityResult:
        warnings: List[str] = []
        if inp.pile_length_m <= 0:
            warnings.append("Chiều dài cọc <= 0. Kiểm tra cao độ đáy bệ và cao độ mũi cọc.")
        if not inp.layers:
            warnings.append("Chưa có lớp đất/đá.")
        if getattr(inp, "include_downdrag", False) and float(getattr(inp, "downdrag_factor", 0.0) or 0.0) <= 0.0:
            inp.downdrag_factor = SCTCalculator.default_downdrag_load_factor(inp)
        warnings += SCTCalculator._preprocess_layers(inp)
        warnings += SCTCalculator._validate_input(inp)
        layers, qsf_cd, qsf_db, qsn, w1 = SCTCalculator._skin_resistance(inp)
        warnings += w1
        qpf_cd, qpf_db, qpn, toe_info, w2 = SCTCalculator._tip_resistance(inp, layers)
        warnings += w2

        w_dry, w_eff, buoyancy = SCTCalculator.pile_weight_effective_kn(inp)

        # Group factor: không áp nhầm bảng cọc khoan đất rời cho cọc đóng/ép.
        # Hệ số nhóm được dùng cho sức kháng địa kỹ thuật quy đổi/cọc; W' không bị nhân fg.
        s_over_d = inp.spacing_m / max(inp.diameter_m, 1e-9)
        is_driven = SCTCalculator._is_driven(inp)
        for lr in layers:
            if inp.pile_count_in_group > 1 and lr.skin_length_m > 0:
                if is_driven and int(lr.soil_type or 0) in (1, 6):
                    fg_i, fg_note = 1.0, "fg cọc đóng/ép đất rời không dùng bảng cọc khoan"
                else:
                    fg_i, fg_note = group_factor_for_soil(lr.soil_type, s_over_d, inp.group_layout, inp.ignore_group_igm_rock)
            else:
                fg_i, fg_note = 1.0, "fg=1.0"
            lr.group_factor = fg_i
            if lr.skin_length_m > 0:
                lr.note = (lr.note + f"; {fg_note}={fg_i:.3f} theo S/D={s_over_d:.3f}").strip('; ')

        qsf_cd_group_side = sum(lr.qs_factored_kn * lr.group_factor for lr in layers)
        qsf_db_group_side = sum(lr.qs_extreme_kn * lr.group_factor for lr in layers)
        fg_cd_weighted = qsf_cd_group_side / qsf_cd if qsf_cd > 1e-9 else 1.0
        fg_db_weighted = qsf_db_group_side / qsf_db if qsf_db > 1e-9 else fg_cd_weighted
        qsup_cd = sum(lr.qs_uplift_factored_kn for lr in layers)
        qsup_db = sum(lr.qs_uplift_extreme_kn for lr in layers)
        downdrag_kn = sum(lr.downdrag_kn for lr in layers) if getattr(inp, "include_downdrag", False) else 0.0
        if downdrag_kn > 0:
            warnings.append(f"Đã tính ma sát âm DD={downdrag_kn:.2f} kN/cọc với γDD={inp.downdrag_factor:.2f}; DD được cộng vào tải nén và vùng DD bị loại khỏi ma sát dương.")

        def build_limit_state(code: str, qshaft: float, qtip: float, qshaft_uplift: float, fg_weighted: float) -> LimitStateCapacity:
            gross = qshaft + qtip
            net_before_dd = gross - w_eff
            # Áp fg cho sức kháng địa kỹ thuật của cọc đơn trong nhóm; trọng lượng cọc không nhân fg.
            group_gross = gross * fg_weighted
            group_net = group_gross - w_eff
            group_total_before_dd = group_net * max(inp.pile_count_in_group, 1)
            net = net_before_dd  # sức kháng không đổi; DD được cộng vào tải khi kiểm toán nén/FOS.
            group_total = group_total_before_dd
            # Quy ước nhổ: dùng qshaft_uplift đã nhân φup riêng; Qp không tham gia nhổ.
            # QA fix P4: trọng lượng bản thân là tải có lợi khi chống nhổ -> nhân γDC = 0.9.
            w_uplift = GAMMA_DC_UPLIFT * w_eff
            uplift_mag = qshaft_uplift + w_uplift
            uplift_signed = -uplift_mag
            uplift_group_total = (qshaft_uplift * fg_weighted + w_uplift) * max(inp.pile_count_in_group, 1)
            pr = SCTCalculator.structural_capacity_kn(inp, code)
            governing = min(net, group_net, pr) if pr > 0 else min(net, group_net)
            note = "Min([Q] cọc đơn theo đất nền, [Q] nhóm quy đổi/cọc, Pr vật liệu); tải kiểm toán nén được cộng thêm ma sát âm nếu bật option"
            return LimitStateCapacity(
                code=code, label=LIMIT_STATE_LABELS.get(code, code), qshaft_kn=qshaft, qtip_kn=qtip,
                compression_single_gross_kn=gross, pile_weight_effective_kn=w_eff, pile_weight_dry_kn=w_dry,
                buoyancy_kn=buoyancy, compression_single_net_kn=net, uplift_single_magnitude_kn=uplift_mag,
                uplift_single_signed_kn=uplift_signed, group_factor=fg_weighted, compression_group_single_gross_kn=group_gross,
                compression_group_single_net_kn=group_net, compression_group_total_kn=group_total, material_pr_kn=pr,
                governing_kn=governing, governing_note=note, qshaft_uplift_kn=qshaft_uplift,
                uplift_group_total_kn=uplift_group_total, downdrag_kn=downdrag_kn,
                compression_single_net_before_downdrag_kn=net_before_dd,
                compression_group_total_before_downdrag_kn=group_total_before_dd,
            )

        strength = build_limit_state(LIMIT_STATE_STRENGTH, qsf_cd, qpf_cd, qsup_cd, fg_cd_weighted)
        extreme = build_limit_state(LIMIT_STATE_EXTREME, qsf_db, qpf_db, qsup_db, fg_db_weighted)
        return CapacityResult(
            pile_input=inp, layers=layers, qshaft_nominal_kn=qsn, qtip_nominal_kn=qpn,
            strength=strength, extreme=extreme, warnings=warnings, toe_info=toe_info
        )



class EditableTree:
    GRID_SEPARATOR_TAG = "__grid_separator__"

    def __init__(self, parent, columns: List[Tuple[str, str, int]], height: int = 12, combo_columns: Optional[Dict[str, List[str]]] = None, copy_suffix_col: Optional[int] = None, visual_grid: bool = False):
        self.columns = columns
        self.combo_columns = combo_columns or {}
        self.copy_suffix_col = copy_suffix_col
        self.visual_grid = bool(visual_grid)
        self._last_right_click_event = None
        self._context_rowid: Optional[str] = None
        self._context_col_index: Optional[int] = None
        # Undo/Redo dạng snapshot theo bảng, tương tự thao tác trong Excel/Word.
        # Mỗi lần người dùng sửa ô, paste, thêm/xóa/copy dòng hoặc thao tác menu sẽ lưu trạng thái trước đó.
        self._undo_stack: List[List[List[str]]] = []
        self._redo_stack: List[List[List[str]]] = []
        self._history_limit = 80
        self._history_restoring = False
        # Cho phép từng bảng thêm lệnh riêng vào menu chuột phải mà không phá các lệnh mặc định.
        self.extra_context_menu_commands: List[Any] = []
        show_cols = [c[0] for c in columns]
        self.tree = ttk.Treeview(parent, columns=show_cols, show="headings", height=height, selectmode="extended")
        for key, label, width in columns:
            # Treeview không hỗ trợ gridline thật theo từng cell; với bảng địa chất bật visual_grid,
            # thêm ký tự phân cách vào heading/cell để nhìn giống bảng Excel hơn.
            head_text = f"│ {label}" if self.visual_grid else label
            self.tree.heading(key, text=head_text)
            self.tree.column(key, width=width, anchor=tk.CENTER)
        yscroll = ttk.Scrollbar(parent, orient=tk.VERTICAL, command=self.tree.yview)
        xscroll = ttk.Scrollbar(parent, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)
        self.tree.grid(row=0, column=0, sticky="nsew")
        yscroll.grid(row=0, column=1, sticky="ns")
        xscroll.grid(row=1, column=0, sticky="ew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)
        self.tree.bind("<Double-1>", self._edit_cell)
        self.tree.bind("<Button-3>", self._show_context_menu)
        self.tree.bind("<Control-d>", self.fill_selected_empty_cells_from_above)
        self.tree.bind("<Control-D>", self.fill_selected_empty_cells_from_above)
        self.tree.bind("<Control-z>", self.undo)
        self.tree.bind("<Control-Z>", self.undo)
        self.tree.bind("<Control-y>", self.redo)
        self.tree.bind("<Control-Y>", self.redo)
        if self.visual_grid:
            try:
                self.tree.tag_configure(self.GRID_SEPARATOR_TAG, foreground="#94A3B8", background="#E2E8F0")
            except Exception:
                pass

    def _strip_visual_cell(self, value: Any) -> str:
        s = str(value or "")
        if self.visual_grid:
            s = s.strip()
            # Dòng kẻ ngang mô phỏng.
            if s and all(ch in "─━-— " for ch in s):
                return ""
            # Bỏ ký tự kẻ đứng chỉ dùng để hiển thị.
            s = s.strip("│ ").strip()
        return s

    def _display_values(self, values: List[Any]) -> List[str]:
        vals = ["" if x is None else str(x) for x in values]
        ncol = len(self.columns)
        vals = (vals + [""] * ncol)[:ncol]
        if not self.visual_grid:
            return vals
        return [f"│ {v}" for v in vals]

    def _separator_values(self) -> List[str]:
        vals = []
        for _, _, width in self.columns:
            n = max(4, min(18, int(width / 9)))
            vals.append("─" * n)
        return vals

    def _is_separator_row(self, rowid: str) -> bool:
        try:
            return self.GRID_SEPARATOR_TAG in set(self.tree.item(rowid, "tags") or ())
        except Exception:
            return False

    def data_children(self) -> List[str]:
        return [i for i in self.tree.get_children() if not self._is_separator_row(i)]

    def _row_values_raw(self, rowid: str) -> List[str]:
        vals = list(self.tree.item(rowid, "values"))
        vals = [self._strip_visual_cell(v) for v in vals]
        ncol = len(self.columns)
        return (vals + [""] * ncol)[:ncol]

    def _insert_raw_row(self, index, values: List[Any], tags: Tuple[str, ...] = ()) -> str:
        rowid = self.tree.insert("", index, values=self._display_values(list(values)), tags=tuple(tags))
        if self.visual_grid:
            # Dòng kẻ ngang mô phỏng ngay dưới mỗi lớp. Dòng này không được trả về trong get_rows().
            self.tree.insert("", index if index == tk.END else self.tree.index(rowid) + 1,
                             values=self._separator_values(), tags=(self.GRID_SEPARATOR_TAG,))
        return rowid

    def _nearest_data_row_from_y(self, y: int) -> str:
        rowid = self.tree.identify_row(y)
        if rowid and not self._is_separator_row(rowid):
            return rowid
        # Nếu bấm đúng vào dòng kẻ ngang, lấy dòng dữ liệu phía trên để lệnh vẫn hoạt động tự nhiên.
        children = list(self.tree.get_children())
        if rowid in children:
            idx = children.index(rowid)
            for j in range(idx - 1, -1, -1):
                if not self._is_separator_row(children[j]):
                    return children[j]
        return ""

    def _show_context_menu(self, event):
        rowid = self._nearest_data_row_from_y(event.y)
        colid = self.tree.identify_column(event.x)
        col_index = None
        try:
            col_index = int(str(colid).replace("#", "")) - 1
        except Exception:
            col_index = None
        if rowid:
            if rowid not in self.tree.selection():
                self.tree.selection_set(rowid)
            self.tree.focus(rowid)
        self._last_right_click_event = event
        self._context_rowid = rowid
        self._context_col_index = col_index if col_index is not None and 0 <= col_index < len(self.columns) else None
        menu = tk.Menu(self.tree, tearoff=0)
        menu.add_command(label="Undo (Ctrl+Z)", command=self.undo)
        menu.add_command(label="Redo (Ctrl+Y)", command=self.redo)
        menu.add_separator()
        menu.add_command(label="Edit ô", command=lambda: self._edit_cell(self._last_right_click_event))
        menu.add_command(label="Paste từ Clipboard", command=lambda: self.paste_from_clipboard(self.tree.winfo_toplevel()))
        menu.add_command(label="Copy dòng đã chọn", command=self.copy_selected_row)
        menu.add_command(label="Copy giá trị ô đang chọn", command=self.fill_clicked_cell_down_to_selected)
        menu.add_command(label="Copy ô trên", command=self.fill_selected_empty_cells_from_above)
        menu.add_command(label="Copy ô dưới", command=self.fill_selected_empty_cells_from_below)
        extra_cmds = list(getattr(self, "extra_context_menu_commands", []) or [])
        if extra_cmds:
            menu.add_separator()
            for spec in extra_cmds:
                if spec is None or spec == "separator":
                    menu.add_separator()
                    continue
                try:
                    label, cmd = spec[0], spec[1]
                    menu.add_command(label=str(label), command=cmd)
                except Exception:
                    continue
        menu.add_separator()
        menu.add_command(label="Xóa dòng", command=self.delete_selected_row)
        menu.add_command(label="Xóa tất cả", command=self.clear)
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def _normalize_rows_for_history(self, rows: List[List[Any]]) -> List[List[str]]:
        """Chuẩn hóa snapshot để so sánh/lưu undo-redo ổn định."""
        ncol = len(self.columns)
        out: List[List[str]] = []
        for r in rows or []:
            vals = list(r) if isinstance(r, (list, tuple)) else [r]
            out.append(["" if v is None else str(v) for v in (vals + [""] * ncol)[:ncol]])
        return out

    def _current_history_state(self) -> List[List[str]]:
        return self._normalize_rows_for_history(self.get_rows())

    def push_undo_state(self):
        """Lưu trạng thái hiện tại trước khi sửa bảng."""
        if getattr(self, "_history_restoring", False):
            return
        state = self._current_history_state()
        if self._undo_stack and self._undo_stack[-1] == state:
            return
        self._undo_stack.append(state)
        limit = max(int(getattr(self, "_history_limit", 80) or 80), 10)
        if len(self._undo_stack) > limit:
            self._undo_stack = self._undo_stack[-limit:]
        self._redo_stack.clear()

    def clear_history(self):
        self._undo_stack.clear()
        self._redo_stack.clear()

    def _restore_history_state(self, rows: List[List[Any]]):
        self._history_restoring = True
        try:
            self.set_rows(rows, record_undo=False)
        finally:
            self._history_restoring = False

    def undo(self, event=None):
        """Ctrl+Z: hoàn tác thao tác gần nhất trên bảng đang focus."""
        if not self._undo_stack:
            return "break" if event is not None else None
        current = self._current_history_state()
        prev = self._undo_stack.pop()
        if current != prev:
            self._redo_stack.append(current)
        self._restore_history_state(prev)
        return "break" if event is not None else None

    def redo(self, event=None):
        """Ctrl+Y: làm lại thao tác vừa Undo."""
        if not self._redo_stack:
            return "break" if event is not None else None
        current = self._current_history_state()
        nxt = self._redo_stack.pop()
        if current != nxt:
            self._undo_stack.append(current)
        self._restore_history_state(nxt)
        return "break" if event is not None else None

    def _append_copy_suffix(self, text: Any) -> str:
        base = str(text or "").strip()
        return (base if base else "Hạng mục") + " (1)"

    def copy_selected_row(self):
        data_children = self.data_children()
        sel = set(self.tree.selection())
        sel_sorted = [rowid for rowid in data_children if rowid in sel]
        if not sel_sorted:
            messagebox.showwarning("Copy dòng", "Chưa chọn dòng cần copy.")
            return
        rows = self.get_rows()
        indices = [data_children.index(r) for r in sel_sorted]
        insert_idx = max(indices) + 1
        copied = []
        for idx in indices:
            vals = list(rows[idx])
            while len(vals) < len(self.columns):
                vals.append("")
            if self.copy_suffix_col is not None and 0 <= int(self.copy_suffix_col) < len(vals):
                vals[int(self.copy_suffix_col)] = self._append_copy_suffix(vals[int(self.copy_suffix_col)])
            copied.append(vals)
        self.push_undo_state()
        self.set_rows(rows[:insert_idx] + copied + rows[insert_idx:])

    def _ordered_selection(self) -> List[str]:
        data_children = self.data_children()
        sel = set(self.tree.selection())
        return [rowid for rowid in data_children if rowid in sel]

    def _set_tree_cell(self, rowid: str, col_index: int, value: Any):
        vals = self._row_values_raw(rowid)
        if 0 <= col_index < len(vals):
            vals[col_index] = value
            self.tree.item(rowid, values=self._display_values(vals))

    def fill_clicked_cell_down_to_selected(self):
        """Copy giá trị của ô đang bấm chuột phải xuống cùng cột của các dòng đã chọn.

        Dùng được cho mọi cột: Hạng mục, Lớp, Loại, SPT, γ, C, φ, qu, RQD...
        """
        rowid = self._context_rowid
        col_index = self._context_col_index
        if not rowid or col_index is None:
            messagebox.showwarning("Copy giá trị ô đang chọn", "Hãy bấm chuột phải đúng vào ô có giá trị nguồn.")
            return
        vals = self._row_values_raw(rowid)
        value = vals[col_index] if col_index < len(vals) else ""
        children = self.data_children()
        selected = self._ordered_selection()
        if rowid not in selected:
            selected = [rowid]
        row_pos = children.index(rowid) if rowid in children else -1
        targets = [r for r in selected if r != rowid and (row_pos < 0 or children.index(r) > row_pos)]
        # Nếu dòng nguồn nằm dưới cùng vùng chọn, vẫn copy sang các dòng còn lại.
        if not targets:
            targets = [r for r in selected if r != rowid]
        if not targets:
            messagebox.showinfo("Copy giá trị ô đang chọn", "Hãy chọn thêm các dòng cần copy giá trị.")
            return
        self.push_undo_state()
        for r in targets:
            self._set_tree_cell(r, col_index, value)

    def fill_selected_empty_cells_from_above(self, event=None):
        """Điền các ô trống của dòng đang chọn bằng giá trị cùng cột ở dòng ngay trên.

        Lệnh này chỉ điền ô trống, không ghi đè dữ liệu đã nhập. Áp dụng cho toàn bộ cột.
        """
        children = self.data_children()
        selected = self._ordered_selection()
        if not selected:
            return "break" if event is not None else None
        history_saved = False
        for rowid in selected:
            try:
                idx = children.index(rowid)
            except ValueError:
                continue
            if idx <= 0:
                continue
            prev = children[idx - 1]
            vals = self._row_values_raw(rowid)
            prev_vals = self._row_values_raw(prev)
            changed = False
            for c in range(len(self.columns)):
                if str(vals[c]).strip() == "" and str(prev_vals[c]).strip() != "":
                    vals[c] = prev_vals[c]
                    changed = True
            if changed:
                if not history_saved:
                    self.push_undo_state()
                    history_saved = True
                self.tree.item(rowid, values=self._display_values(vals))
        return "break" if event is not None else None

    def fill_selected_empty_cells_from_below(self, event=None):
        """Điền các ô trống của dòng đang chọn bằng giá trị cùng cột ở dòng ngay dưới.

        Lệnh này tương tự Copy ô trên: chỉ điền ô trống, không ghi đè dữ liệu đã nhập.
        Khi chọn nhiều dòng liên tiếp, xử lý từ dưới lên để giá trị có thể được kéo ngược lên tự nhiên.
        """
        children = self.data_children()
        selected = self._ordered_selection()
        if not selected:
            return "break" if event is not None else None
        history_saved = False
        for rowid in reversed(selected):
            try:
                idx = children.index(rowid)
            except ValueError:
                continue
            if idx >= len(children) - 1:
                continue
            next_row = children[idx + 1]
            vals = self._row_values_raw(rowid)
            next_vals = self._row_values_raw(next_row)
            changed = False
            for c in range(len(self.columns)):
                if str(vals[c]).strip() == "" and str(next_vals[c]).strip() != "":
                    vals[c] = next_vals[c]
                    changed = True
            if changed:
                if not history_saved:
                    self.push_undo_state()
                    history_saved = True
                self.tree.item(rowid, values=self._display_values(vals))
        return "break" if event is not None else None

    def delete_selected_row(self):
        sel = set(self.tree.selection())
        if not sel:
            return
        data_children = self.data_children()
        delete_set = set()
        all_children = list(self.tree.get_children())
        for rowid in data_children:
            if rowid in sel:
                delete_set.add(rowid)
                # Xóa cả dòng kẻ ngang liền sau dòng dữ liệu.
                try:
                    idx = all_children.index(rowid)
                    if idx + 1 < len(all_children) and self._is_separator_row(all_children[idx + 1]):
                        delete_set.add(all_children[idx + 1])
                except Exception:
                    pass
        if delete_set:
            self.push_undo_state()
        for rowid in list(delete_set):
            try:
                self.tree.delete(rowid)
            except Exception:
                pass

    def add_blank_rows(self, n: int = 1):
        rows = self.get_rows()
        rows.extend([["" for _ in self.columns] for _ in range(max(int(n), 1))])
        self.push_undo_state()
        self.set_rows(rows)

    def clear(self, record_undo: bool = True):
        if record_undo and self.tree.get_children():
            self.push_undo_state()
        for item in self.tree.get_children():
            self.tree.delete(item)

    def paste_from_clipboard(self, root):
        try:
            text = root.clipboard_get()
        except Exception:
            messagebox.showwarning("Clipboard", "Không đọc được clipboard.")
            return
        rows = []
        for line in text.splitlines():
            if not line.strip():
                continue
            parts = [p.strip() for p in line.split("\t")]
            if len(parts) == 1:
                parts = [p.strip() for p in line.replace(",", "\t").split("\t")]
            rows.append(parts)
        if not rows:
            return
        ncol = len(self.columns)
        data_rows = []
        for idx, r in enumerate(rows):
            # Chỉ bỏ dòng đầu nếu rõ ràng là header. Không bỏ các dòng dữ liệu có tên lớp/hạng mục dạng chữ.
            if idx == 0 and len(rows) > 1:
                joined = "|".join(str(x).strip().lower() for x in r)
                header_keys = ["hạng mục", "hang muc", "item", "lớp", "lop", "cao độ", "cao do", "loại", "loai", "n1", "n60", "gamma", "loại cọc", "loai coc"]
                if any(k in joined for k in header_keys):
                    continue
            vals = (r + [""] * ncol)[:ncol]
            data_rows.append(vals)
        self.push_undo_state()
        self.set_rows(data_rows)

    def get_rows(self) -> List[List[str]]:
        return [self._row_values_raw(i) for i in self.data_children()]

    def set_rows(self, rows: List[List[Any]], record_undo: bool = False):
        if record_undo:
            self.push_undo_state()
        self.clear(record_undo=False)
        ncol = len(self.columns)
        for r in rows:
            vals = ["" if x is None else x for x in (list(r) + [""] * ncol)[:ncol]]
            self._insert_raw_row(tk.END, vals)

    def _edit_cell(self, event):
        region = self.tree.identify("region", event.x, event.y)
        if region != "cell":
            return
        rowid = self._nearest_data_row_from_y(event.y)
        colid = self.tree.identify_column(event.x)
        if not rowid or not colid:
            return
        try:
            col_index = int(colid.replace("#", "")) - 1
        except Exception:
            return
        if not (0 <= col_index < len(self.columns)):
            return
        bbox = self.tree.bbox(rowid, colid)
        if not bbox:
            return
        x, y, w, h = bbox
        vals = self._row_values_raw(rowid)
        old = vals[col_index] if col_index < len(vals) else ""
        col_key = self.columns[col_index][0]
        choices = self.combo_columns.get(col_key)
        if choices:
            entry = ttk.Combobox(self.tree, values=choices, state="readonly")
            entry.set(old if old in choices else (choices[0] if choices else ""))
        else:
            entry = ttk.Entry(self.tree)
            entry.insert(0, old)
        entry.place(x=x, y=y, width=w, height=h)
        entry.focus_set()
        committed = [False]
        def commit(_=None):
            if committed[0]:
                return
            committed[0] = True
            try:
                val = entry.get()
            except Exception:
                return
            vals2 = self._row_values_raw(rowid)
            old2 = vals2[col_index] if col_index < len(vals2) else ""
            if str(old2) != str(val):
                self.push_undo_state()
                vals2[col_index] = val
                self.tree.item(rowid, values=self._display_values(vals2))
            try:
                entry.destroy()
            except Exception:
                pass
        entry.bind("<Return>", commit)
        entry.bind("<FocusOut>", commit)
        if choices:
            entry.bind("<<ComboboxSelected>>", commit)


class SCTApp:
    def __init__(self, root: Optional[tk.Tk] = None):
        self.root = root if root is not None else tk.Tk()
        self._main_thread_ident = threading.get_ident()
        self._destroyed = False
        self._calc_busy = False
        self._export_busy = False
        self.root.title(f"{APP_NAME} - {APP_TITLE}")
        apply_app_icon(self.root, "ts_cap")
        self.root.geometry("1180x760")
        self.theme_key = tk.StringVar(value="AQUA_FRESH")
        self.borehole_ocr_engine = tk.StringVar(value="TESSERACT")
        self.pal = self._get_palette(self.theme_key.get())
        self.vars: Dict[str, tk.Variable] = {}
        self.last_result: Optional[CapacityResult] = None
        self.last_warning_rows: List[Tuple[str, str]] = []
        self.warning_button = None
        self.item_geo_map: Dict[str, str] = {}  # key hạng mục tính toán -> key lỗ khoan địa chất được chọn
        self.item_geo_map_display: Dict[str, str] = {}
        # Dữ liệu riêng từng hạng mục nhưng không đưa vào bảng chính để bảng cọc tính toán không quá dài.
        self.item_downdrag_data: Dict[str, Dict[str, str]] = {}  # key -> {item, top, bottom, gamma}
        self.item_uplift_data: Dict[str, Dict[str, str]] = {}     # key -> {item, uplift_cd, uplift_db}
        self.project_file_path: str = ""  # file .tscap đang mở/lưu
        self._setup_style()
        self._build_ui()
        self._bind_global_undo_redo_shortcuts()
        self._bind_ux_rescue_shortcuts()
        try:
            self.root.bind("<Destroy>", self._on_root_destroy, add="+")
        except Exception:
            pass
        # Không tự nạp dữ liệu demo khi khởi động. Người dùng nhập/paste/import dữ liệu thực tế hoặc mở Template dữ liệu.

    def _on_root_destroy(self, event=None) -> None:
        try:
            if event is None or getattr(event, "widget", None) is self.root:
                self._destroyed = True
                safe_release_grabs(self.root)
        except Exception:
            pass

    def _root_alive(self) -> bool:
        try:
            return (not getattr(self, "_destroyed", False)) and bool(self.root.winfo_exists())
        except Exception:
            return False

    def _bind_ux_rescue_shortcuts(self) -> None:
        """Phím cứu hộ khi dialog/modal/progress bị ẩn sau cửa sổ khác."""
        try:
            self.root.bind_all("<F12>", lambda e: safe_lift_window(self.root), add="+")
            self.root.bind_all("<Control-Shift-Escape>", lambda e: safe_release_grabs(self.root), add="+")
            self.root.bind_all("<Control-Shift-BackSpace>", lambda e: safe_release_grabs(self.root), add="+")
        except Exception:
            pass

    def _is_main_thread(self) -> bool:
        try:
            return threading.get_ident() == getattr(self, "_main_thread_ident", None)
        except Exception:
            return True

    def _set_status(self, text: Any) -> None:
        """Cập nhật status an toàn khi tác vụ OCR chạy nền."""
        msg = str(text or "")
        try:
            if self._is_main_thread():
                if hasattr(self, "status"):
                    self.status.set(msg)
            else:
                root = getattr(self, "root", None)
                if root is not None:
                    root.after(0, lambda m=msg: hasattr(self, "status") and self.status.set(m))
        except Exception:
            pass

    def _update_ui_idle(self) -> None:
        """Chỉ update giao diện trong main thread, tránh Tkinter treo/lỗi khi chạy worker."""
        try:
            if self._is_main_thread():
                self.root.update_idletasks()
        except Exception:
            pass

    def _bind_global_undo_redo_shortcuts(self):
        """Bắt Ctrl+Z/Ctrl+Y khi focus nằm trong một bảng EditableTree."""
        try:
            self.root.bind_all("<Control-z>", self._handle_global_undo, add="+")
            self.root.bind_all("<Control-Z>", self._handle_global_undo, add="+")
            self.root.bind_all("<Control-y>", self._handle_global_redo, add="+")
            self.root.bind_all("<Control-Y>", self._handle_global_redo, add="+")
            self.root.bind_all("<Control-s>", self._handle_save_project_shortcut, add="+")
            self.root.bind_all("<Control-S>", self._handle_save_project_shortcut, add="+")
            self.root.bind_all("<Control-o>", self._handle_open_project_shortcut, add="+")
            self.root.bind_all("<Control-O>", self._handle_open_project_shortcut, add="+")
        except Exception:
            pass

    def _focused_editable_tree(self) -> Optional[EditableTree]:
        """Trả bảng đang focus; chỉ Undo/Redo bảng, không chặn Ctrl+Z của Entry/Text khác."""
        try:
            fw = self.root.focus_get()
        except Exception:
            fw = None
        if fw is None:
            return None
        try:
            # Khi đang sửa trực tiếp trong ô bằng Entry/Combobox hoặc đang gõ ở Text,
            # không bắt Ctrl+Z/Ctrl+Y để tránh phá thao tác nhập liệu hiện hành.
            if fw.winfo_class() in ("Entry", "TEntry", "Text", "TCombobox", "Combobox"):
                return None
        except Exception:
            pass
        for tbl in (getattr(self, "item_table", None), getattr(self, "layer_table", None), getattr(self, "result_table", None)):
            if not tbl:
                continue
            try:
                if fw == tbl.tree or str(fw).startswith(str(tbl.tree)):
                    return tbl
            except Exception:
                continue
        return None

    def _handle_global_undo(self, event=None):
        tbl = self._focused_editable_tree()
        if tbl is not None:
            return tbl.undo(event)
        return None

    def _handle_global_redo(self, event=None):
        tbl = self._focused_editable_tree()
        if tbl is not None:
            return tbl.redo(event)
        return None

    def _handle_save_project_shortcut(self, event=None):
        # Không chặn Ctrl+S khi đang gõ trong ô nhập liệu hệ thống nếu người dùng đang edit Entry/Text.
        try:
            fw = self.root.focus_get()
            if fw is not None and fw.winfo_class() in ("Entry", "TEntry", "Text", "TCombobox", "Combobox"):
                return None
        except Exception:
            pass
        self.save_project_file()
        return "break"

    def _handle_open_project_shortcut(self, event=None):
        try:
            fw = self.root.focus_get()
            if fw is not None and fw.winfo_class() in ("Entry", "TEntry", "Text", "TCombobox", "Combobox"):
                return None
        except Exception:
            pass
        self.open_project_file()
        return "break"

    # ------------------------------------------------------------------
    # Lưu / mở toàn bộ dự án TS-CAP
    # ------------------------------------------------------------------
    def _project_snapshot(self) -> Dict[str, Any]:
        """Tạo snapshot JSON cho toàn bộ dữ liệu người dùng đang nhập.

        File dự án .tscap chỉ lưu dữ liệu đầu vào, mapping, dữ liệu phụ và bảng kết quả hiển thị.
        Đối tượng tính toán nội bộ last_results không được serialize; sau khi mở dự án nên bấm Tính toán lại
        trước khi xuất báo cáo để kết quả/báo cáo luôn đồng bộ với code hiện hành.
        """
        vars_data: Dict[str, Any] = {}
        for key, var in getattr(self, "vars", {}).items():
            try:
                vars_data[key] = var.get()
            except Exception:
                vars_data[key] = ""
        try:
            result_summary = self.summary_text.get("1.0", "end-1c") if hasattr(self, "summary_text") else ""
        except Exception:
            result_summary = ""
        result_filter = {}
        for name in ("result_filter_metric", "result_filter_mode", "result_filter_a", "result_filter_b"):
            try:
                obj = getattr(self, name, None)
                result_filter[name] = obj.get() if obj is not None else ""
            except Exception:
                result_filter[name] = ""
        return {
            "file_type": 'TS-CAP project',
            "format_version": 1,
            "app_name": APP_NAME,
            "saved_at": datetime.now().isoformat(timespec="seconds"),
            "theme_key": str(self.theme_key.get() if hasattr(self, "theme_key") else ""),
            "borehole_ocr_engine": _normalize_borehole_ocr_engine(self.borehole_ocr_engine.get() if hasattr(self, "borehole_ocr_engine") else "TESSERACT"),
            "vars": vars_data,
            "item_rows": self.item_table.get_rows() if hasattr(self, "item_table") else [],
            "layer_rows": self.layer_table.get_rows() if hasattr(self, "layer_table") else [],
            "item_geo_map": dict(getattr(self, "item_geo_map", {}) or {}),
            "item_geo_map_display": dict(getattr(self, "item_geo_map_display", {}) or {}),
            "item_downdrag_data": dict(getattr(self, "item_downdrag_data", {}) or {}),
            "item_uplift_data": dict(getattr(self, "item_uplift_data", {}) or {}),
            "result_rows": self.result_table.get_rows() if hasattr(self, "result_table") else [],
            "result_summary": result_summary,
            "result_filter": result_filter,
        }

    def save_project_file(self):
        """Lưu toàn bộ dữ liệu dự án đang nhập vào file .tscap dạng JSON UTF-8."""
        try:
            default_name = safe_filename(self.vars.get("project").get() if self.vars.get("project") else 'TS_CAP_Project', 'TS_CAP_Project')
        except Exception:
            default_name = 'TS_CAP_Project'
        path = getattr(self, "project_file_path", "") or ""
        if not path:
            path = filedialog.asksaveasfilename(
                title='Lưu dự án TS-CAP',
                defaultextension=".tscap",
                initialfile=f"{default_name}.tscap",
                filetypes=[('TS-CAP Project', "*.tscap"), ("JSON", "*.json"), ("All files", "*.*")],
            )
        if not path:
            return
        try:
            data = self._project_snapshot()
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            self.project_file_path = path
            self._set_status(f"Đã lưu dự án: {os.path.basename(path)}")
            messagebox.showinfo("Lưu dự án", f"Đã lưu dự án:\n{path}")
        except Exception as exc:
            traceback.print_exc()
            messagebox.showerror("Lưu dự án", f"Không lưu được dự án:\n{exc}")

    def save_project_file_as(self):
        old_path = getattr(self, "project_file_path", "")
        self.project_file_path = ""
        try:
            self.save_project_file()
        finally:
            if not getattr(self, "project_file_path", "") and old_path:
                self.project_file_path = old_path

    def open_project_file(self):
        """Mở file .tscap và thay toàn bộ dữ liệu đang nhập bằng dữ liệu trong file."""
        path = filedialog.askopenfilename(
            title='Mở dự án TS-CAP',
            filetypes=[('TS-CAP Project', "*.tscap"), ("JSON", "*.json"), ("All files", "*.*")],
        )
        if not path:
            return
        # Hỏi lại để tránh người dùng mất dữ liệu chưa lưu.
        try:
            has_current = bool((self.item_table.get_rows() if hasattr(self, "item_table") else []) or (self.layer_table.get_rows() if hasattr(self, "layer_table") else []))
        except Exception:
            has_current = False
        if has_current:
            ok = messagebox.askyesno("Mở dự án", "Mở dự án sẽ thay dữ liệu đang nhập trên màn hình. Tiếp tục?")
            if not ok:
                return
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self._apply_project_snapshot(data)
            self.project_file_path = path
            self._set_status(f"Đã mở dự án: {os.path.basename(path)}")
            messagebox.showinfo("Open", "Đã mở dự án.")
        except Exception as exc:
            traceback.print_exc()
            messagebox.showerror("Mở dự án", f"Không mở được dự án:\n{exc}")

    def _apply_project_snapshot(self, data: Dict[str, Any]):
        if not isinstance(data, dict):
            raise ValueError("File dự án không đúng định dạng JSON.")
        if str(data.get("file_type", "")).strip() and 'TS-CAP' not in str(data.get("file_type", "")):
            # Vẫn cho mở JSON cũ nếu thiếu file_type, nhưng cảnh báo khi file_type rõ ràng không đúng.
            raise ValueError('File này không phải dự án TS-CAP.')

        # Theme
        try:
            theme_key = str(data.get("theme_key", "") or "")
            if theme_key:
                self.theme_key.set(theme_key)
                self._setup_style()
        except Exception:
            pass

        # Tùy chọn OCR lỗ khoan scan.
        try:
            if hasattr(self, "borehole_ocr_engine"):
                self.borehole_ocr_engine.set(_normalize_borehole_ocr_engine(data.get("borehole_ocr_engine", "TESSERACT")))
        except Exception:
            pass

        # Biến thông số chung.
        vars_data = data.get("vars", {}) or {}
        if isinstance(vars_data, dict):
            for key, val in vars_data.items():
                try:
                    if key in self.vars:
                        self.vars[key].set(val)
                except Exception:
                    continue

        # Bảng chính.
        item_rows = data.get("item_rows", []) or []
        layer_rows = data.get("layer_rows", []) or []
        if hasattr(self, "item_table"):
            self.item_table.set_rows(item_rows, record_undo=False)
            self.item_table.clear_history()
        if hasattr(self, "layer_table"):
            self.layer_table.set_rows(layer_rows, record_undo=False)
            try:
                self._merge_first_col_visual(self.layer_table)
            except Exception:
                pass
            self.layer_table.clear_history()

        # Dữ liệu phụ theo hạng mục/lỗ khoan.
        self.item_geo_map = dict(data.get("item_geo_map", {}) or {})
        self.item_geo_map_display = dict(data.get("item_geo_map_display", {}) or {})
        self.item_downdrag_data = dict(data.get("item_downdrag_data", {}) or {})
        self.item_uplift_data = dict(data.get("item_uplift_data", {}) or {})

        # Khôi phục bảng kết quả hiển thị để người dùng xem lại; kết quả tính nội bộ sẽ được tính lại khi bấm Tính toán.
        if hasattr(self, "result_table"):
            self.result_table.set_rows(data.get("result_rows", []) or [], record_undo=False)
            self.result_table.clear_history()
        try:
            self.summary_text.config(state=tk.NORMAL)
            self.summary_text.delete("1.0", tk.END)
            self.summary_text.insert("1.0", str(data.get("result_summary", "") or ""))
        except Exception:
            pass
        rf = data.get("result_filter", {}) or {}
        if isinstance(rf, dict):
            for name in ("result_filter_metric", "result_filter_mode", "result_filter_a", "result_filter_b"):
                try:
                    obj = getattr(self, name, None)
                    if obj is not None and name in rf:
                        obj.set(rf.get(name, ""))
                except Exception:
                    pass

        self.last_result = None
        self.last_results = []

    def _get_palette(self, key: Any) -> Dict[str, str]:
        k = _normalize_ui_theme(key)
        return dict(_N2D_THEME_PRESETS.get(k, DEFAULT_THEME))

    def _setup_style(self):
        'Đồng bộ style với TS-PILE/TS-COL.\n\n        Nguyên tắc V0.2.50:\n        - Vùng nhập liệu dùng nền phẳng theo theme, không dùng nền chữ riêng cho label.\n        - Nút, notebook, bảng và entry dùng chung palette trong n2d_theme_library.py.\n        - Sidebar/toolbar dùng tk widget nên cần refresh màu riêng sau khi đổi theme.\n        '
        self.pal = self._get_palette(self.theme_key.get())
        pal = self.pal
        try:
            self.root.configure(bg=pal["bg"])
        except Exception:
            pass
        try:
            self.root.option_add("*Font", "Arial 10")
        except Exception:
            pass

        style = ttk.Style(self.root)
        self.style = style
        try:
            style.theme_use("clam")
        except Exception:
            pass

        # Nền chữ thống nhất theo nền form để tránh các ô chữ bị đóng khung màu.
        label_bg = pal["bg"]
        style.configure("TFrame", background=pal["bg"])
        style.configure("Panel.TFrame", background=pal["bg"])
        style.configure("Card.TFrame", background=pal["bg"], relief="flat")
        style.configure("TLabel", background=label_bg, foreground=pal["text"], font=("Arial", 10))
        style.configure("Card.TLabel", background=label_bg, foreground=pal["text"], font=("Arial", 10))
        style.configure("Muted.TLabel", background=label_bg, foreground=pal["muted"], font=("Arial", 9))
        style.configure("Title.TLabel", background=pal["bg"], foreground=pal["accent_dark"], font=("Arial", 18, "bold"))
        style.configure("Subtitle.TLabel", background=pal["bg"], foreground=pal["muted"], font=("Arial", 9))
        style.configure("Header.TLabel", background=pal["bg"], foreground=pal["accent_dark"], font=("Arial", 17, "bold"))
        style.configure("TLabelframe", background=pal["bg"], bordercolor=pal["border"], lightcolor=pal["border"], darkcolor=pal["border"])
        style.configure("TLabelframe.Label", background=pal["bg"], foreground=pal["accent_dark"], font=("Arial", 10, "bold"))
        style.configure("TNotebook", background=pal["bg"], bordercolor=pal["border"])
        style.configure("TNotebook.Tab", background=pal["button"], foreground=pal["text"], padding=(9, 5), font=("Arial", 9, "bold"))
        style.map("TNotebook.Tab", background=[("selected", pal["panel"]), ("active", pal["button_active"])], foreground=[("selected", pal["accent_dark"])])
        style.configure("TButton", font=("Arial", 9), padding=(6, 4), background=pal["button"], foreground=pal["text"], bordercolor=pal["border"])
        style.map("TButton", background=[("active", pal["button_active"]), ("pressed", pal["big_button"])], foreground=[("disabled", "#888888")])
        style.configure("Small.TButton", font=("Arial", 9), padding=(4, 3), background=pal["button"], foreground=pal["text"])
        style.configure("About.TButton", font=("Arial", 8, "bold"), padding=(4, 2), background=pal["button"], foreground=pal["accent_dark"], bordercolor=pal["border"])
        style.map("About.TButton", background=[("active", pal["button_active"]), ("pressed", pal["big_button"])])
        style.configure("Accent.TButton", font=("Arial", 10, "bold"), padding=(9, 6), background=pal["big_button"], foreground=pal["text"])
        style.configure("Primary.TButton", font=("Arial", 10, "bold"), padding=(12, 8), background=pal["accent"], foreground="white")
        style.map("Accent.TButton", background=[("active", pal["button_active"]), ("pressed", pal["border"])])
        style.map("Primary.TButton", background=[("active", pal["accent_dark"]), ("pressed", pal["accent_dark"])], foreground=[("active", "white")])
        style.configure("TCheckbutton", background=pal["bg"], foreground=pal["text"], font=("Arial", 10))
        style.configure("TRadiobutton", background=pal["bg"], foreground=pal["text"], font=("Arial", 10))
        style.configure("TEntry", font=("Arial", 10), fieldbackground=pal["entry_bg"], background=pal["entry_bg"], foreground=pal["entry_fg"], insertcolor=pal["entry_fg"], bordercolor=pal["border"])
        style.configure("TCombobox", font=("Arial", 10), fieldbackground=pal["entry_bg"], background=pal["entry_bg"], foreground=pal["entry_fg"], arrowcolor=pal["accent_dark"], bordercolor=pal["border"])
        style.map("TCombobox", fieldbackground=[("readonly", pal["entry_bg"]), ("!disabled", pal["entry_bg"])], background=[("readonly", pal["entry_bg"]), ("!disabled", pal["entry_bg"])], foreground=[("readonly", pal["entry_fg"]), ("!disabled", pal["entry_fg"])])
        style.configure("Treeview", background=pal.get("tree_row", "white"), foreground=pal.get("tree_fg", pal["text"]), rowheight=26, fieldbackground=pal.get("tree_row", "white"), font=("Arial", 10), bordercolor=pal["border"])
        style.configure("Treeview.Heading", background=pal.get("tree_head", pal["button"]), foreground=pal["text"], font=("Arial", 10, "bold"), bordercolor=pal["border"])
        style.map("Treeview", background=[("selected", pal["accent"])], foreground=[("selected", "white")])
        try:
            self._refresh_n2d_shell_colors()
        except Exception:
            pass

    def _paint_tk_children(self, widget):
        """Cập nhật màu cho tk.Frame/tk.Label/tk.Button trong shell giao diện."""
        pal = self.pal
        try:
            cls = widget.winfo_class()
            if cls in ("Frame", "Toplevel"):
                # Sidebar có hàm riêng; các frame còn lại theo nền form.
                if widget is getattr(self, "sidebar_frame", None):
                    widget.configure(bg=pal["sidebar"])
                elif widget is getattr(self, "sidebar_separator", None):
                    widget.configure(bg=pal["sidebar2"])
                else:
                    widget.configure(bg=pal["bg"])
            elif cls == "Label":
                widget.configure(bg=pal["bg"], fg=pal["text"])
            elif cls == "Text":
                widget.configure(bg=pal.get("entry_bg", "white"), fg=pal.get("entry_fg", "black"), insertbackground=pal.get("entry_fg", "black"))
        except Exception:
            pass
        for child in widget.winfo_children():
            self._paint_tk_children(child)

    def _refresh_n2d_shell_colors(self):
        pal = self.pal
        try:
            self._paint_tk_children(self.root)
        except Exception:
            pass
        try:
            self.sidebar_frame.configure(bg=pal["sidebar"])
            for child in self.sidebar_frame.winfo_children():
                if child.winfo_class() == "Label":
                    child.configure(bg=pal["sidebar"], fg="white")
                elif child.winfo_class() == "Frame":
                    child.configure(bg=pal["sidebar2"] if child is getattr(self, "sidebar_separator", None) else pal["sidebar"])
        except Exception:
            pass
        try:
            self.sidebar_title.configure(bg=pal["sidebar"], fg="white")
            self.sidebar_caption.configure(bg=pal["sidebar"], fg="#F8FAFC")
            if hasattr(self, "sidebar_files_label"):
                self.sidebar_files_label.configure(bg=pal["sidebar"], fg="white")
            self.sidebar_group_label.configure(bg=pal["sidebar"], fg="white")
            if hasattr(self, "author_frame"):
                self.author_frame.configure(bg=pal["sidebar"], bd=0, highlightthickness=0)
            for lbl in getattr(self, "author_labels", []):
                lbl.configure(bg=pal["sidebar"], fg="#FFF7EA", bd=0, highlightthickness=0)
        except Exception:
            pass
        try:
            self.title_label.configure(bg=pal["bg"], fg=pal["accent_dark"])
            self.subtitle_label.configure(bg=pal["bg"], fg=pal["muted"])
            self.header_frame.configure(bg=pal["bg"])
            self.title_frame.configure(bg=pal["bg"])
            if hasattr(self, "project_buttons_frame"):
                self.project_buttons_frame.configure(bg=pal["bg"])
            if hasattr(self, "header_actions_frame"):
                self.header_actions_frame.configure(bg=pal["bg"])
            self.theme_frame.configure(bg=pal["bg"])
        except Exception:
            pass
        try:
            self.summary_text.configure(bg=pal.get("entry_bg", "white"), fg=pal.get("entry_fg", "black"), insertbackground=pal.get("entry_fg", "black"))
        except Exception:
            pass
        try:
            self._update_file_buttons()
        except Exception:
            pass
        try:
            self._update_workflow_buttons()
        except Exception:
            pass

    def _update_file_buttons(self):
        """Cập nhật màu/nền cho nhóm Files trong sidebar."""
        pal = self.pal
        btn_bg = pal.get("button", "#E0F2FE")
        btn_active = pal.get("button_active", "#BAE6FD")
        btn_fg = pal.get("accent_dark", pal.get("accent", "#1D4ED8"))
        border = pal.get("sidebar2", pal.get("border", "#93C5FD"))
        for btn in getattr(self, "file_buttons", []):
            try:
                btn.configure(
                    bg=btn_bg,
                    fg=btn_fg,
                    activebackground=btn_active,
                    activeforeground=btn_fg,
                    font=("Arial", 9, "bold"),
                    relief=tk.RIDGE,
                    bd=1,
                    cursor="hand2",
                    anchor="w",
                    justify=tk.LEFT,
                    padx=10,
                    pady=7,
                    highlightthickness=1,
                    highlightbackground=border,
                    highlightcolor=border,
                )
            except Exception:
                pass

    def _go_to_workflow_step(self, idx: int):
        try:
            self.nb.select(idx)
            self._update_workflow_buttons()
        except Exception:
            pass

    def _update_workflow_buttons(self):
        pal = self.pal
        try:
            active_idx = self.nb.index(self.nb.select())
        except Exception:
            active_idx = 0
        btn_bg = pal.get("button", "#E0F2FE")
        btn_active = pal.get("button_active", "#BAE6FD")
        btn_fg = pal.get("accent_dark", pal.get("accent", "#1D4ED8"))
        border = pal.get("sidebar2", pal.get("border", "#93C5FD"))
        for i, btn in enumerate(getattr(self, "workflow_buttons", [])):
            label = self.workflow_steps[i][0] if i < len(getattr(self, "workflow_steps", [])) else btn.cget("text")
            is_active = (i == active_idx)
            try:
                btn.configure(
                    text=label,
                    bg=(pal.get("panel", "white") if is_active else btn_bg),
                    fg=btn_fg,
                    activebackground=btn_active,
                    activeforeground=btn_fg,
                    font=("Arial", 9, "bold" if is_active else "normal"),
                    relief=(tk.SUNKEN if is_active else tk.RIDGE),
                    bd=1,
                    cursor="hand2",
                    anchor="w",
                    justify=tk.LEFT,
                    padx=10,
                    pady=7,
                    highlightthickness=1,
                    highlightbackground=border,
                    highlightcolor=border,
                )
            except Exception:
                pass

    def _var(self, name: str, default: Any = "") -> tk.StringVar:
        v = tk.StringVar(value=str(default))
        self.vars[name] = v
        return v

    def _build_ui(self):
        'Shell giao diện kiểu TS-PILE/TS-COL: sidebar trái + vùng làm việc phải.'
        pal = self.pal
        self.root.minsize(1120, 720)

        self.sidebar_frame = tk.Frame(self.root, width=230, bg=pal["sidebar"])
        self.sidebar_frame.pack(side=tk.LEFT, fill=tk.Y)
        self.sidebar_frame.pack_propagate(False)

        self.sidebar_title = tk.Label(
            self.sidebar_frame,
            text='TS-CAP V1.0',
            font=("Arial", 18, "bold"),
            fg="white",
            bg=pal["sidebar"],
            justify=tk.LEFT,
        )
        self.sidebar_title.pack(anchor=tk.W, padx=20, pady=(24, 4))
        self.sidebar_caption = tk.Label(
            self.sidebar_frame,
            text="Tính sức chịu tải cọc\ntheo đất nền",
            font=("Arial", 9),
            fg="#F8FAFC",
            bg=pal["sidebar"],
            justify=tk.LEFT,
        )
        self.sidebar_caption.pack(anchor=tk.W, padx=20, pady=(0, 16))
        self.sidebar_separator = tk.Frame(self.sidebar_frame, height=1, bg=pal["sidebar2"])
        self.sidebar_separator.pack(fill=tk.X, padx=20, pady=(0, 18))

        # Sidebar chia 2 nhóm chính: Files và Menu.
        # Các nhãn nhóm cùng thẳng hàng; các nút con cùng thẳng hàng.
        self.sidebar_files_label = tk.Label(self.sidebar_frame, text="Files", font=("Arial", 10, "bold"), fg="white", bg=pal["sidebar"])
        self.sidebar_files_label.pack(anchor=tk.W, padx=20, pady=(8, 8))
        self.file_buttons = []
        file_actions = [
            ("Open", self.open_project_file),
            ("Save", self.save_project_file),
            ("Save As", self.save_project_file_as),
        ]
        for label, cmd in file_actions:
            btn = tk.Button(
                self.sidebar_frame,
                text=label,
                command=cmd,
                wraplength=190,
            )
            btn.pack(fill=tk.X, padx=20, pady=4)
            self.file_buttons.append(btn)

        self.sidebar_group_label = tk.Label(self.sidebar_frame, text="Menu", font=("Arial", 10, "bold"), fg="white", bg=pal["sidebar"])
        self.sidebar_group_label.pack(anchor=tk.W, padx=20, pady=(18, 8))

        self.workflow_steps = [
            ("Thông số chung", 0),
            ("Thông tin cọc tính toán", 1),
            ("Thông số địa chất", 2),
            ("Tính toán", 3),
        ]
        self.workflow_buttons = []
        for label, tab_idx in self.workflow_steps:
            btn = tk.Button(
                self.sidebar_frame,
                text=label,
                command=lambda idx=tab_idx: self._go_to_workflow_step(idx),
                wraplength=190,
            )
            btn.pack(fill=tk.X, padx=20, pady=4)
            self.workflow_buttons.append(btn)

        self.author_frame = tk.Frame(self.sidebar_frame, bg=pal["sidebar"], bd=0, highlightthickness=0)
        self.author_frame.pack(side=tk.BOTTOM, anchor=tk.W, fill=tk.X, padx=20, pady=22)
        self.author_labels = []
        author_lines = [line.strip() for line in APP_AUTHOR.split("\n") if line.strip()]
        if author_lines:
            lbl = tk.Label(self.author_frame, text=author_lines[0], font=("Arial", 10, "bold"), fg="#FFF7EA", bg=pal["sidebar"], bd=0, highlightthickness=0, wraplength=200, justify=tk.LEFT)
            lbl.pack(anchor=tk.W, pady=(0, 8))
            self.author_labels.append(lbl)
            for line in author_lines[1:]:
                lbl = tk.Label(self.author_frame, text=line, font=("Arial", 10), fg="#FFF7EA", bg=pal["sidebar"], bd=0, highlightthickness=0, wraplength=200, justify=tk.LEFT)
                lbl.pack(anchor=tk.W, pady=(2, 2))
                self.author_labels.append(lbl)

        right = tk.Frame(self.root, bg=pal["bg"], padx=16, pady=12)
        right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        self.header_frame = tk.Frame(right, bg=pal["bg"])
        self.header_frame.pack(fill=tk.X, pady=(0, 10))

        self.header_actions_frame = tk.Frame(self.header_frame, bg=pal["bg"])
        self.header_actions_frame.pack(side=tk.RIGHT, anchor=tk.NE)
        ttk.Button(self.header_actions_frame, text="Settings", style="About.TButton", command=self.show_settings).pack(side=tk.LEFT, padx=(0, 4), pady=(2, 0))
        ttk.Button(self.header_actions_frame, text="License", style="About.TButton", command=self.show_license).pack(side=tk.LEFT, padx=(0, 4), pady=(2, 0))
        ttk.Button(self.header_actions_frame, text="About", style="About.TButton", command=self.show_about).pack(side=tk.LEFT, padx=(0, 0), pady=(2, 0))
        # Giữ biến theme_frame để các hàm đổi theme cũ không bị ảnh hưởng.
        self.theme_frame = self.header_actions_frame

        self.title_frame = tk.Frame(self.header_frame, bg=pal["bg"])
        self.title_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.title_label = tk.Label(self.title_frame, text=APP_TITLE, font=("Arial", 18, "bold"), fg=pal["accent_dark"], bg=pal["bg"])
        self.title_label.pack(anchor=tk.W)
        self.subtitle_label = tk.Label(self.title_frame, text=APP_NAME, font=("Arial", 9), fg=pal["muted"], bg=pal["bg"])
        self.subtitle_label.pack(anchor=tk.W, pady=(2, 0))

        self.nb = ttk.Notebook(right)
        self.nb.pack(fill=tk.BOTH, expand=True)
        self.tab_input = ttk.Frame(self.nb, padding=10)
        self.tab_items = ttk.Frame(self.nb, padding=10)
        self.tab_layers = ttk.Frame(self.nb, padding=10)
        self.tab_result = ttk.Frame(self.nb, padding=10)
        self.nb.add(self.tab_input, text=" Thông số chung ")
        self.nb.add(self.tab_items, text=" Thông tin cọc tính toán ")
        self.nb.add(self.tab_layers, text=" Thông số địa chất ")
        self.nb.add(self.tab_result, text=" Tính toán ")
        self.nb.bind("<<NotebookTabChanged>>", lambda _e: self._update_workflow_buttons())

        self._build_input_tab()
        self._build_items_tab()
        self._build_layers_tab()
        self._build_result_tab()

        action_bar = ttk.Frame(right, padding=(0, 8, 0, 0))
        action_bar.pack(fill=tk.X)
        ttk.Button(action_bar, text="Tính toán", style="Primary.TButton", width=18, command=self.calculate).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(action_bar, text="Xuất báo cáo", style="Accent.TButton", width=18, command=self.export_report).pack(side=tk.LEFT, padx=6)
        self.warning_button = ttk.Button(action_bar, text="Warning", width=14, command=self.show_warning_panel, state=tk.DISABLED)
        self.warning_button.pack(side=tk.LEFT, padx=6)
        self.status = tk.StringVar(value="Sẵn sàng - chưa nạp dữ liệu")
        ttk.Label(action_bar, textvariable=self.status, style="Muted.TLabel").pack(side=tk.RIGHT, padx=6)

        self._refresh_n2d_shell_colors()

    def _set_theme_by_label(self, label: str):
        for k, v in _N2D_THEME_LABELS.items():
            if v == label:
                self.theme_key.set(k)
                break
        self._setup_style()
        try:
            self._set_status(f"Đã đổi theme: {label}")
        except Exception:
            pass

    def _change_theme(self, _=None):
        label = ""
        try:
            label = self.cbo_theme.get()
        except Exception:
            label = _N2D_THEME_LABELS.get(self.theme_key.get(), "")
        if label:
            self._set_theme_by_label(label)

    def _center_window(self, win):
        try:
            win.update_idletasks()
            w = win.winfo_width()
            h = win.winfo_height()
            x = self.root.winfo_rootx() + max((self.root.winfo_width() - w) // 2, 0)
            y = self.root.winfo_rooty() + max((self.root.winfo_height() - h) // 2, 0)
            win.geometry(f"{w}x{h}+{x}+{y}")
        except Exception:
            pass

    def show_about(self):
        win = tk.Toplevel(self.root)
        win.title(f"About - {APP_NAME}")
        win.transient(self.root)
        win.grab_set()
        win.resizable(False, False)
        win.configure(bg=self.pal["panel"])
        frm = tk.Frame(win, bg=self.pal["panel"], padx=26, pady=22)
        frm.pack(fill=tk.BOTH, expand=True)
        tk.Label(frm, text=APP_NAME, font=("Arial", 24, "bold"), fg=self.pal["accent_dark"], bg=self.pal["panel"]).pack(anchor=tk.W)
        tk.Label(frm, text=APP_TITLE, font=("Arial", 12, "italic"), fg=self.pal["muted"], bg=self.pal["panel"]).pack(anchor=tk.W, pady=(2, 14))
        tk.Frame(frm, height=1, bg=self.pal["border"]).pack(fill=tk.X, pady=(0, 14))
        tk.Label(frm, text="Tác giả", font=("Arial", 11, "bold"), fg=self.pal["text"], bg=self.pal["panel"]).pack(anchor=tk.W)
        tk.Label(frm, text="Nguyễn Ngọc Dũng", font=("Arial", 14, "bold"), fg=self.pal["accent_dark"], bg=self.pal["panel"]).pack(anchor=tk.W, pady=(2, 8))
        tk.Label(frm, text="Phòng QLTK - Khối XD PPP\nTập đoàn SunGroup", justify=tk.LEFT, font=("Arial", 10), fg=self.pal["text"], bg=self.pal["panel"]).pack(anchor=tk.W, pady=(0, 16))
        ttk.Button(frm, text="OK", style="Accent.TButton", command=win.destroy).pack(fill=tk.X)
        self._center_window(win)

    def show_license(self):
        win = tk.Toplevel(self.root)
        win.title(f"License - {APP_NAME}")
        win.geometry("560x250")
        win.resizable(False, False)
        win.configure(bg=self.pal["bg"])
        try:
            win.transient(self.root)
            win.grab_set()
        except Exception:
            pass
        frm = ttk.Frame(win, padding=18)
        frm.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frm, text="License", font=("Arial", 20, "bold"), foreground=self.pal["accent_dark"]).grid(row=0, column=0, columnspan=3, sticky=tk.W, pady=(0, 14))
        ttk.Label(frm, text="Phần mềm:", font=("Arial", 10, "bold")).grid(row=1, column=0, sticky=tk.W, padx=(0, 10), pady=6)
        ttk.Label(frm, text=APP_NAME, font=("Arial", 11, "bold"), foreground=self.pal["accent_dark"]).grid(row=1, column=1, sticky=tk.W, pady=6)
        full_machine_id = get_machine_id()
        ttk.Label(frm, text="HWID / Mã thiết bị:", font=("Arial", 10, "bold")).grid(row=2, column=0, sticky=tk.W, padx=(0, 10), pady=6)
        hwid_var = tk.StringVar(value=full_machine_id)
        hwid_entry = ttk.Entry(frm, textvariable=hwid_var, width=48, state="readonly")
        hwid_entry.grid(row=2, column=1, sticky=tk.EW, pady=6)
        def copy_hwid():
            self.root.clipboard_clear()
            self.root.clipboard_append(full_machine_id)
            self._set_status("Đã copy HWID vào clipboard")
        ttk.Button(frm, text="Copy", command=copy_hwid).grid(row=2, column=2, sticky=tk.W, padx=(8, 0), pady=6)
        ttk.Label(frm, text="Gửi HWID này cho người quản lý license khi cần kích hoạt/gia hạn.", style="Muted.TLabel", wraplength=500).grid(row=3, column=0, columnspan=3, sticky=tk.W, pady=(10, 4))
        ttk.Button(frm, text="Đóng", style="Accent.TButton", command=win.destroy).grid(row=4, column=1, sticky=tk.E, pady=(16, 0))
        frm.columnconfigure(1, weight=1)
        self._center_window(win)

    def show_settings(self):
        win = tk.Toplevel(self.root)
        win.title(f"Settings - {APP_NAME}")
        win.transient(self.root)
        win.grab_set()
        win.resizable(False, False)
        frm = ttk.Frame(win, padding=18)
        frm.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frm, text="Settings", font=("Arial", 20, "bold"), foreground=self.pal["accent_dark"]).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 12))
        ttk.Label(frm, text="Theme giao diện", font=("Arial", 11, "bold")).grid(row=1, column=0, sticky=tk.W, pady=6, padx=(0, 12))
        theme_values = [v for k, v in _N2D_THEME_LABELS.items()]
        theme_var = tk.StringVar(value=_N2D_THEME_LABELS.get(self.theme_key.get(), theme_values[0] if theme_values else ""))
        cb_theme = ttk.Combobox(frm, textvariable=theme_var, values=theme_values, state="readonly", width=48)
        cb_theme.grid(row=1, column=1, sticky=tk.EW, pady=6)
        ttk.Label(frm, text="OCR ảnh/PDF scan", font=("Arial", 11, "bold")).grid(row=2, column=0, sticky=tk.W, pady=6, padx=(0, 12))
        engine_values = list(BOREHOLE_OCR_ENGINE_LABELS.values())
        engine_key = _normalize_borehole_ocr_engine(self.borehole_ocr_engine.get() if hasattr(self, "borehole_ocr_engine") else "TESSERACT")
        engine_var = tk.StringVar(value=BOREHOLE_OCR_ENGINE_LABELS.get(engine_key, BOREHOLE_OCR_ENGINE_LABELS["TESSERACT"]))
        cb_engine = ttk.Combobox(frm, textvariable=engine_var, values=engine_values, state="readonly", width=48)
        cb_engine.grid(row=2, column=1, sticky=tk.EW, pady=6)
        ttk.Label(
            frm,
            text="PDF CAD vẫn ưu tiên đọc text/vector - hãy dùng PDF gốc thay vì ảnh chụp khi có thể. Engine mặc định Tesseract; RapidOCR là lựa chọn thử nghiệm cho ảnh khó (chậm hơn đáng kể trên CPU, cần pip install rapidocr-onnxruntime).",
            style="Muted.TLabel",
            wraplength=520,
        ).grid(row=3, column=0, columnspan=2, sticky=tk.W, pady=(0, 4))
        frm.columnconfigure(1, weight=1)
        def apply_settings(close_after=False):
            self._set_theme_by_label(theme_var.get())
            eng_txt = str(engine_var.get())
            if "Rapid" in eng_txt:
                selected_engine = "RAPID"
            elif "Paddle" in eng_txt:
                selected_engine = "PADDLE"
            else:
                selected_engine = "TESSERACT"
            self.borehole_ocr_engine.set(selected_engine)
            # Dọn cache OCR crop khi đổi engine để tránh dùng lại kết quả của engine cũ.
            try:
                self._borehole_ocr_crop_cache = {}
                self._borehole_ocr_crop_data_cache = {}
            except Exception:
                pass
            if selected_engine == "PADDLE":
                try:
                    import paddleocr  # type: ignore  # noqa: F401
                    import paddle  # type: ignore  # noqa: F401
                except Exception as exc:
                    messagebox.showwarning(
                        "PaddleOCR",
                        "Chưa dùng được PaddleOCR. Chương trình sẽ tự dùng Tesseract.\n"
                        f"Chi tiết: {exc}"
                    )
            elif selected_engine == "RAPID":
                try:
                    import rapidocr_onnxruntime  # type: ignore  # noqa: F401
                except Exception as exc:
                    messagebox.showwarning(
                        "RapidOCR",
                        "Chưa dùng được RapidOCR. Cài bằng lệnh: pip install rapidocr-onnxruntime\n"
                        "Chương trình sẽ tự dùng Tesseract cho đến khi cài xong.\n"
                        f"Chi tiết: {exc}"
                    )
            if close_after:
                win.destroy()
        btns = ttk.Frame(frm)
        btns.grid(row=4, column=0, columnspan=2, sticky=tk.EW, pady=(12, 0))
        ttk.Button(btns, text="OK", style="Accent.TButton", command=lambda: apply_settings(True)).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=3)
        ttk.Button(btns, text="Apply", command=lambda: apply_settings(False)).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=3)
        ttk.Button(btns, text="Cancel", command=win.destroy).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=3)
        self._center_window(win)

    def _add_row(self, parent, row, label, var, width=15, combo_values=None):
        ttk.Label(parent, text=label, style="Card.TLabel").grid(row=row, column=0, sticky="w", padx=6, pady=4)
        if combo_values:
            w = ttk.Combobox(parent, textvariable=var, values=combo_values, width=width, state="readonly")
        else:
            w = ttk.Entry(parent, textvariable=var, width=width)
        w.grid(row=row, column=1, sticky="ew", padx=6, pady=4)
        return w

    def _help_assets(self) -> Dict[str, Dict[str, Any]]:
        # V0.2.27: bỏ hoàn toàn cơ chế hiện ảnh ghi chú trong giao diện.
        return {}

    def _asset_path(self, filename: str) -> Optional[str]:
        return None

    def _hide_help_popup(self, event=None):
        popup = getattr(self, "_help_popup", None)
        if popup is not None:
            try:
                popup.destroy()
            except Exception:
                pass
        self._help_popup = None
        self._help_popup_key = None
        self._help_popup_refs = []

    def _show_help_popup(self, key: str, x_root: Optional[int] = None, y_root: Optional[int] = None):
        # V0.2.27: không hiện ảnh/popup nữa.
        return

    def _attach_help_popup(self, widget, key: str):
        # V0.2.27: bỏ trigger click/hover để tránh ảnh tự hiện.
        return

    def _on_layer_tree_click(self, event):
        # V0.2.27: click vào GSI/mi không mở ảnh nữa.
        return

    def _build_input_tab(self):
        outer = ttk.Frame(self.tab_input)
        outer.pack(fill=tk.BOTH, expand=True)
        left = ttk.LabelFrame(outer, text="Thông số chung", padding=10)
        right = ttk.LabelFrame(outer, text="Thông số mô hình tính", padding=10)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 6))
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(6, 0))
        for f in (left, right):
            f.grid_columnconfigure(1, weight=1)
        r = 0
        self._add_row(left, r, "Dự án", self._var("project", ""), 30); r += 1
        self._add_row(left, r, "γ bê tông (kN/m³)", self._var("concrete_gamma", "24.5"), 12); r += 1
        self._add_row(left, r, "f'c bê tông (MPa)", self._var("fc_mpa", "30"), 12); r += 1
        self._add_row(left, r, "fy thép (MPa)", self._var("fy_mpa", "400"), 12); r += 1
        self._add_row(left, r, "Số thanh thép", self._var("n_rebars", "0"), 12); r += 1
        self._add_row(left, r, "ĐK thép chủ (mm)", self._var("rebar_dia", "0"), 12); r += 1
        self._add_row(left, r, "Loại đai 1=xoắn,2=thường", self._var("stirrup_type", "1"), 12, ["1", "2"]); r += 1
        self._add_row(left, r, "Bỏ ma sát đầu cọc khoan (m)", self._var("exclude_top", "1.5"), 12); r += 1
        self._add_row(left, r, "ER SPT (%)", self._var("spt_er", "60"), 12); r += 1
        self._add_row(left, r, "SPT nhập vào", self._var("spt_input_mode", "Nₕₜ"), 18, ["Nₕₜ", "N60"]); r += 1
        self.vars["include_downdrag"] = tk.BooleanVar(value=False)
        ttk.Checkbutton(left, text="Xét đến ma sát âm", variable=self.vars["include_downdrag"]).grid(row=r, column=0, columnspan=2, sticky="w", padx=6, pady=4); r += 1

        r = 0
        self._add_row(right, r, "Mô hình σ'p đất rời", self._var("sand_mode", "sand"), 12, ["sand", "gravel"]); r += 1
        self._add_row(right, r, "Hệ số thành phần cát m", self._var("sand_m", "0.6"), 12, ["0.6", "0.8"]); r += 1
        joint_widget = self._add_row(right, r, "Điều kiện khe nứt đá", self._var("rock_joint_condition", "Khe nứt hở hoặc có mùn"), 28, ["Khe nứt khép kín", "Khe nứt hở hoặc có mùn"])
        r += 1
        self._add_row(right, r, "Điều kiện thi công", self._var("rock_construction_condition", "Có chống đỡ"), 18, ["Có chống đỡ", "Không chống đỡ"]); r += 1
        self.vars["include_rock_tip"] = tk.BooleanVar(value=False)
        ttk.Checkbutton(right, text="Tính Qp mũi đá", variable=self.vars["include_rock_tip"]).grid(row=r, column=0, columnspan=2, sticky="w", padx=6, pady=4); r += 1
        self.vars["allow_rock_tip_exceed_25qu"] = tk.BooleanVar(value=False)
        ttk.Checkbutton(right, text="Cho phép Qp đá > 2.5qu (có thử tải)", variable=self.vars["allow_rock_tip_exceed_25qu"]).grid(row=r, column=0, columnspan=2, sticky="w", padx=6, pady=2); r += 1
        self.vars["clay_use_c_phi"] = tk.BooleanVar(value=False)
        ttk.Checkbutton(right, text="Tính sét theo C, φ", variable=self.vars["clay_use_c_phi"]).grid(row=r, column=0, columnspan=2, sticky="w", padx=6, pady=4); r += 1
        self.vars["ignore_group_igm_rock"] = tk.BooleanVar(value=True)
        ttk.Checkbutton(right, text="Bỏ qua hệ số nhóm cọc với IGM, đá", variable=self.vars["ignore_group_igm_rock"]).grid(row=r, column=0, columnspan=2, sticky="w", padx=6, pady=4); r += 1
        self.vars["allow_geology_extrapolation"] = tk.BooleanVar(value=False)
        ttk.Checkbutton(right, text="Cho phép ngoại suy lớp cuối tới mũi cọc", variable=self.vars["allow_geology_extrapolation"]).grid(row=r, column=0, columnspan=2, sticky="w", padx=6, pady=4); r += 1
        igm_alpha_widget = self._add_row(right, r, "Hệ số thực nghiệm alpha cho IGM (theo bảng B1)", self._var("igm_alpha", "0.25"), 12)
        r += 1
        igm_joint_widget = self._add_row(right, r, "Hệ số điều chỉnh mức độ khe nối IGM", self._var("igm_joint_factor", "0.45"), 12)
        r += 1
        self._add_row(right, r, "Khoảng cách các khe nứt sd (mm)", self._var("crack_spacing", "50"), 12); r += 1
        self._add_row(right, r, "Chiều rộng khe nứt td (mm)", self._var("crack_width", "5"), 12); r += 1

    def _build_items_tab(self):
        top = ttk.Frame(self.tab_items)
        top.pack(fill=tk.X, pady=(0, 6))
        ttk.Label(top, text="Nhập thông số cọc theo từng hạng mục/mố/trụ. Loại cọc: 1 = Khoan nhồi; 2/3 = Đóng/Ép vuông; 2T/3T = Đóng/Ép tròn; 2O/3O = Cọc ống (cột Ds = D trong).", foreground=self.pal["muted"]).pack(side=tk.LEFT, padx=4)
        # Các thông tin phụ theo hạng mục được đưa lên góc phải cạnh Template để bảng chính gọn hơn.
        ttk.Button(top, text="Template dữ liệu", command=self.show_item_template).pack(side=tk.RIGHT, padx=4)
        ttk.Button(top, text="Điền lực nhổ", command=self.show_uplift_item_dialog).pack(side=tk.RIGHT, padx=4)
        ttk.Button(top, text="Thông số tính ma sát âm", command=self.show_downdrag_item_dialog).pack(side=tk.RIGHT, padx=4)
        table_frame = ttk.Frame(self.tab_items)
        table_frame.pack(fill=tk.BOTH, expand=True)
        cols = [
            ("item", "Hạng mục", 110), ("pile_type", "Loại cọc", 78), ("pile_count", "Số cọc", 62),
            ("row_count", "Số hàng cọc", 85),
            ("D", "D (mm)", 70), ("Ds", "Ds (mm)", 70), ("spacing", "S (m)", 65),
            ("Bx", "Bx bệ (m)", 78), ("By", "By bệ (m)", 78), ("Cz", "Cz (m)", 68),
            ("ground", "CĐ mặt đất", 82), ("cap", "CĐ đáy bệ", 82), ("tip", "CĐ mũi", 82), ("water", "Mực nước ngầm", 105),
            ("Pu_cd", "Pu CĐ/cọc (kN)", 105), ("Pu_db", "Pu ĐB/cọc (kN)", 105),
            ("Ncap_cd", "N đáy bệ CĐ (kN)", 118), ("Ncap_db", "N đáy bệ ĐB (kN)", 118),
            ("note", "Ghi chú", 180),
        ]
        self.item_table = EditableTree(table_frame, cols, height=15, combo_columns={"pile_type": ["1", "2", "2T", "2O", "3", "3T", "3O"]}, copy_suffix_col=0)
        btns = ttk.Frame(self.tab_items)
        btns.pack(fill=tk.X, pady=6)
        ttk.Button(btns, text='Import từ TS-PILE/MCOC', command=self.import_from_n2d_pile_result).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Paste từ Clipboard", command=lambda: self.item_table.paste_from_clipboard(self.root)).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Import từ CSV/TXT", command=self.import_items_file).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Thêm dòng", command=self._add_item_rows_dialog).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Copy dòng", command=self.item_table.copy_selected_row).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Xóa dòng", command=self.item_table.delete_selected_row).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Xóa toàn bộ", command=self.item_table.clear).pack(side=tk.LEFT, padx=3)

    def _build_layers_tab(self):
        top_opts = ttk.Frame(self.tab_layers)
        top_opts.pack(fill=tk.X, pady=(0, 4))
        self.vars["common_geology"] = tk.BooleanVar(value=False)
        ttk.Checkbutton(top_opts, text="Một LK địa chất dùng chung cho toàn bộ hạng mục", variable=self.vars["common_geology"]).pack(side=tk.LEFT, padx=4)
        ttk.Button(top_opts, text="Template dữ liệu", command=self.show_geology_template).pack(side=tk.RIGHT, padx=4)
        note = ttk.Label(self.tab_layers, text="Quy ước lớp đất đá: 0 là KK/Karst, 1 là Cát, 2 là Sét, 3 là Đá nguyên gốc, 4 là Đá phong hóa, 5 là IGM, 6 là Sỏi Cuội.", foreground=self.pal["muted"], wraplength=1080, justify=tk.LEFT)
        note.pack(anchor=tk.W, pady=(0, 8))
        table_frame = ttk.Frame(self.tab_layers)
        table_frame.pack(fill=tk.BOTH, expand=True)
        cols = [
            ("geo_item", "Lỗ Khoan", 95), ("name", "Lớp", 70), ("bottom", "CĐ đáy", 80), ("type", "Loại", 58),
            ("n", "SPT", 65), ("gamma", "γ", 65), ("c", "C MPa", 70),
            ("phi", "φ độ", 65), ("qu", "qu MPa", 70), ("rqd", "RQD", 60),
            ("gsi", "GSI", 60), ("mi", "mi", 60), ("dist", "D", 55), ("comment", "Ghi chú", 180)
        ]
        self.layer_table = EditableTree(table_frame, cols, height=18, visual_grid=False)
        # Không dùng mô phỏng kẻ đứng/kẻ ngang vì Treeview hiển thị xấu trên một số theme.
        # Vẫn giữ gộp hiển thị cột hạng mục và các chức năng điền dữ liệu nhanh.
        self.layer_table.tree.bind("<ButtonRelease-1>", self._on_layer_tree_click, add="+")
        self.layer_table.extra_context_menu_commands = [
            ("Copy lỗ khoan địa chất (_copy)", self.copy_context_geology_item_auto),
            ("Thêm 1 dòng phía trên", self.insert_context_geology_row_above),
            ("Thêm 1 dòng phía dưới", self.insert_context_geology_row_below),
            ("Xóa lỗ khoan địa chất", self.delete_context_geology_item),
        ]
        btns = ttk.Frame(self.tab_layers)
        btns.pack(fill=tk.X, pady=6)
        # Thứ tự nút địa chất theo luồng nhập dữ liệu: nhập/import -> quản lý hạng mục -> quản lý dòng -> chỉnh loại đất.
        ttk.Button(btns, text="Nhập dữ liệu từ ảnh/PDF", command=self.import_geology_from_borehole_image).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Paste từ Clipboard", command=self._paste_geology_clipboard).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Import từ CSV/TXT", command=self.import_geology_file).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Thêm LK địa chất", command=self._add_geology_item_dialog).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Copy LK địa chất", command=self.copy_selected_geology_item).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Khớp LK địa chất", command=self.open_item_geology_mapping_dialog).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Thêm dòng", command=self._add_layer_rows_dialog).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Xóa dòng", command=self.layer_table.delete_selected_row).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Xóa toàn bộ", command=self.layer_table.clear).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Chỉnh loại đất", command=self._bulk_set_layer_type_dialog).pack(side=tk.LEFT, padx=3)

    def show_item_template(self):
        win = tk.Toplevel(self.root)
        win.title("Template dữ liệu cọc tính toán")
        win.geometry("1280x430")
        try:
            win.transient(self.root)
            win.grab_set()
        except Exception:
            pass
        frm = ttk.Frame(win, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frm, text="Template dữ liệu cọc tính toán", font=("Arial", 16, "bold"), foreground=self.pal["accent_dark"]).pack(anchor=tk.W, pady=(0, 8))
        ttk.Label(frm, text="Copy hàng tiêu đề và các dòng mẫu sang Excel, nhập dữ liệu rồi paste lại vào bảng Thông số cọc tính toán. Quy ước Loại cọc: 1 = Cọc khoan nhồi; 2 = Cọc đóng vuông; 3 = Cọc ép vuông; 2T/3T = Cọc đóng/ép tròn; 2O/3O = Cọc ống (khi đó cột Ds = đường kính trong).", style="Muted.TLabel", wraplength=1220, justify=tk.LEFT).pack(anchor=tk.W, pady=(0, 8))
        sample = (
            "Hạng mục\tLoại cọc (1/2/3/2T/3T/2O/3O)\tSố cọc\tSố hàng cọc\tD (mm)\tDs (mm; ống: D trong)\tS (m)\tBx bệ (m)\tBy bệ (m)\tCz (m)\tCĐ mặt đất\tCĐ đáy bệ\tCĐ mũi\tMực nước ngầm\tPu CĐ/cọc (kN)\tPu ĐB/cọc (kN)\tLực nhổ CĐ/cọc (kN)\tLực nhổ ĐB/cọc (kN)\tN đáy bệ CĐ (kN)\tN đáy bệ ĐB (kN)\tGhi chú\n"
            "P1\t1\t8\t2\t1200\t1200\t3.6\t\t\t\t1.5\t1.5\t-40.0\t1.5\t\t\t\t\t\t\t\n"
            "P2\t1\t10\t2\t1500\t1500\t4.0\t\t\t\t2.0\t1.0\t-55.0\t1.0\t\t\t\t\t\t\t"
        )
        txt = tk.Text(frm, height=12, wrap=tk.NONE, font=("Consolas", 10), bg=self.pal.get("entry_bg", "white"), fg=self.pal.get("entry_fg", "black"))
        txt.pack(fill=tk.BOTH, expand=True)
        txt.insert("1.0", sample)
        def copy_template():
            self.root.clipboard_clear()
            self.root.clipboard_append(sample)
            self._set_status("Đã copy template dữ liệu cọc tính toán")
        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(btns, text="Copy template", style="Accent.TButton", command=copy_template).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(btns, text="Đóng", command=win.destroy).pack(side=tk.RIGHT)
        self._center_window(win)

    def show_geology_template(self):
        win = tk.Toplevel(self.root)
        win.title("Template dữ liệu LK địa chất")
        win.geometry("1060x430")
        try:
            win.transient(self.root)
            win.grab_set()
        except Exception:
            pass
        frm = ttk.Frame(win, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frm, text="Template dữ liệu LK địa chất", font=("Arial", 16, "bold"), foreground=self.pal["accent_dark"]).pack(anchor=tk.W, pady=(0, 8))
        ttk.Label(frm, text="Có thể copy bảng mẫu dưới đây ra Excel, nhập dữ liệu rồi paste lại vào bảng Thông số địa chất.", style="Muted.TLabel").pack(anchor=tk.W, pady=(0, 8))
        sample = (
            "Lỗ khoan\tLớp\tCĐ đáy\tLoại\tSPT\tγ\tC MPa\tφ độ\tqu MPa\tRQD\tGSI\tmi\tD\tGhi chú\n"
            "HK1\t1\t-2.40\t1\t12\t18.5\t\t30\t\t\t\t\t\tCát pha\n"
            "HK1\t2\t-15.00\t2\t18\t17.8\t\t\t\t\t\t\t\tSét cứng\n"
            "HK1\t3\t-32.00\t5\t100\t20.0\t\t\t4.5\t\t\t\t\tIGM\n"
            "HK1\t4\t-40.00\t4\t100\t22.0\t\t\t18\t50\t35\t8\t0.5\tĐá phong hóa"
        )
        txt = tk.Text(frm, height=12, wrap=tk.NONE, font=("Consolas", 10), bg=self.pal.get("entry_bg", "white"), fg=self.pal.get("entry_fg", "black"))
        txt.pack(fill=tk.BOTH, expand=True)
        txt.insert("1.0", sample)
        def copy_template():
            self.root.clipboard_clear()
            self.root.clipboard_append(sample)
            self._set_status("Đã copy template dữ liệu LK địa chất")
        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(btns, text="Copy template", style="Accent.TButton", command=copy_template).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(btns, text="Đóng", command=win.destroy).pack(side=tk.RIGHT)
        self._center_window(win)

    def _build_result_tab(self):
        tools = ttk.LabelFrame(self.tab_result, text="Lọc kết quả / tối ưu chiều dài cọc")
        tools.pack(fill=tk.X, pady=(0, 8))

        metric_labels = self._fos_metric_labels()
        self.result_filter_metric = tk.StringVar(value=metric_labels[0] if metric_labels else "FOS nén cọc đơn CĐ")
        self.result_filter_mode = tk.StringVar(value="Từ a đến b")
        self.result_filter_a = tk.StringVar(value="")
        self.result_filter_b = tk.StringVar(value="")

        ttk.Label(tools, text="Lọc theo:").grid(row=0, column=0, sticky=tk.W, padx=(8, 4), pady=4)
        ttk.Combobox(tools, textvariable=self.result_filter_metric, values=metric_labels, state="readonly", width=24).grid(row=0, column=1, sticky=tk.W, padx=4, pady=4)
        ttk.Combobox(tools, textvariable=self.result_filter_mode, values=["Từ a đến b", "Lớn hơn a", "Nhỏ hơn b"], state="readonly", width=12).grid(row=0, column=2, sticky=tk.W, padx=4, pady=4)
        ttk.Label(tools, text="a:").grid(row=0, column=3, sticky=tk.E, padx=(8, 2), pady=4)
        ttk.Entry(tools, textvariable=self.result_filter_a, width=8).grid(row=0, column=4, sticky=tk.W, padx=2, pady=4)
        ttk.Label(tools, text="b:").grid(row=0, column=5, sticky=tk.E, padx=(8, 2), pady=4)
        ttk.Entry(tools, textvariable=self.result_filter_b, width=8).grid(row=0, column=6, sticky=tk.W, padx=2, pady=4)
        ttk.Button(tools, text="Áp dụng lọc", command=self.apply_result_fos_filter).grid(row=0, column=7, sticky=tk.W, padx=(10, 3), pady=4)
        ttk.Button(tools, text="Bỏ lọc", command=self.clear_result_fos_filter).grid(row=0, column=8, sticky=tk.W, padx=3, pady=4)
        ttk.Button(tools, text="Tối ưu chiều dài theo FOS", style="Accent.TButton", command=self.open_pile_length_optimization_dialog).grid(row=0, column=9, sticky=tk.E, padx=(16, 8), pady=4)
        tools.grid_columnconfigure(9, weight=1)

        top = ttk.Frame(self.tab_result)
        top.pack(fill=tk.X, pady=(0, 8))
        self.summary_text = tk.Text(top, height=10, wrap=tk.WORD, font=("Consolas", 10), bg=self.pal.get("entry_bg", "white"), fg=self.pal.get("entry_fg", "black"))
        self.summary_text.pack(fill=tk.X, expand=False)
        table_frame = ttk.Frame(self.tab_result)
        table_frame.pack(fill=tk.BOTH, expand=True)
        cols = [
            ("item", "Hạng mục", 95), ("layer", "Lớp", 70), ("top", "CĐ trên", 72), ("bottom", "CĐ dưới", 72), ("t", "dày", 60),
            ("skin", "L ma sát", 72), ("dd_len", "L MS âm", 72), ("dd", "DD kN", 72), ("type", "Loại", 110), ("fg", "fg nhóm", 65), ("N", "Nₕₜ", 55), ("N60", "N60", 55), ("(N1)60", "(N1)60", 70),
            ("sig", "σ'v MPa", 82), ("c_su", "C/Su", 65), ("ab", "α/β", 65),
            ("phi_cd", "φ CĐ", 55), ("qs_cd", "qs CĐ", 74), ("Qs_cd", "Qs CĐ", 78),
            ("phi_db", "φ ĐB", 55), ("qs_db", "qs ĐB", 74), ("Qs_db", "Qs ĐB", 78),
            ("note", "Ghi chú", 260)
        ]
        self.result_table = EditableTree(table_frame, cols, height=14)


    def _merge_first_col_visual(self, editable_tree: EditableTree):
        """Mô phỏng merge ô cột đầu và kẻ/tô phân nhóm cho bảng địa chất.

        Treeview của Tkinter không có cell-merge thật như Excel, nên phần mềm giữ dữ liệu
        bằng cách để tên hạng mục ở dòng đầu mỗi nhóm, các dòng sau để trống trong cột
        hạng mục. Khi tính toán, hàm đọc bảng vẫn tự hiểu các dòng trống thuộc hạng mục
        ngay phía trên. Các tag màu giúp nhìn rõ ranh giới ngang giữa các lớp/nhóm.
        """
        rows = editable_tree.get_rows()
        # cấu hình tag hiển thị; bọc try để vẫn chạy trên mọi theme Tk.
        try:
            editable_tree.tree.tag_configure("geo_group_start", background="#E6F0FF")
            editable_tree.tree.tag_configure("geo_row_even", background="#FFFFFF")
            editable_tree.tree.tag_configure("geo_row_odd", background="#F7FAFC")
        except Exception:
            pass
        editable_tree.clear()
        last_key = ""
        for idx, r in enumerate(rows):
            rr = list(r)
            tags = ["geo_row_even" if idx % 2 == 0 else "geo_row_odd"]
            if rr:
                cur = _display_item_name(rr[0])
                key = _normalize_item_name(cur)
                if key:
                    if key == last_key:
                        rr[0] = ""
                    else:
                        last_key = key
                        tags.append("geo_group_start")
            editable_tree._insert_raw_row(tk.END, rr, tags=tuple(tags))


    def _bulk_set_layer_type_dialog(self):
        """Chỉnh cột Loại hàng loạt theo tên lớp địa chất.

        Ví dụ: nhập "4" và chọn "bắt đầu bằng" để các lớp 4, 4a, 4b... cùng chuyển sang đá;
        nhập "6" để chuyển toàn bộ lớp số 6 sang sét.
        """
        dlg = tk.Toplevel(self.root)
        dlg.title("Chỉnh loại đất")
        dlg.transient(self.root)
        dlg.grab_set()
        dlg.geometry("560x340")
        ttk.Label(
            dlg,
            text="Chọn các tên lớp cần đổi loại. Có thể nhập nhiều mã lớp, cách nhau bằng dấu phẩy.\n"
                 "Ví dụ: nhập 4 và chọn 'Bắt đầu bằng' để áp dụng cho 4, 4a, 4b...",
            wraplength=480,
            justify=tk.LEFT,
        ).pack(anchor=tk.W, padx=14, pady=(12, 8))

        form = ttk.Frame(dlg)
        form.pack(fill=tk.X, padx=14, pady=4)
        ttk.Label(form, text="Tên lớp:", width=16).grid(row=0, column=0, sticky=tk.W, pady=4)
        layer_var = tk.StringVar(value="4")
        ttk.Entry(form, textvariable=layer_var, width=34).grid(row=0, column=1, sticky="ew", pady=4)
        match_prefix_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(form, text="Bắt đầu bằng mã lớp", variable=match_prefix_var).grid(row=1, column=1, sticky=tk.W, pady=4)
        ttk.Label(form, text="Loại mới:", width=16).grid(row=2, column=0, sticky=tk.W, pady=4)
        choices = [
            "0 - Không khí/Hang karst",
            "1 - Cát",
            "2 - Sét",
            "3 - Đá nguyên khối",
            "4 - Đá nứt vỡ/phong hóa",
            "5 - IGM",
            "6 - Cuội sỏi",
        ]
        type_var = tk.StringVar(value="4 - Đá nứt vỡ/phong hóa")
        ttk.Combobox(form, textvariable=type_var, values=choices, state="readonly", width=31).grid(row=2, column=1, sticky="ew", pady=4)
        form.grid_columnconfigure(1, weight=1)

        quick = ttk.Frame(dlg)
        quick.pack(fill=tk.X, padx=14, pady=(8, 4))
        def quick_set(layer_code: str, type_choice: str, prefix: bool = True):
            layer_var.set(layer_code)
            type_var.set(type_choice)
            match_prefix_var.set(prefix)
        ttk.Button(quick, text="Lớp 4* → Đá", command=lambda: quick_set("4", "4 - Đá nứt vỡ/phong hóa", True)).pack(side=tk.LEFT, padx=3)
        ttk.Button(quick, text="Lớp 6 → Sét", command=lambda: quick_set("6", "2 - Sét", False)).pack(side=tk.LEFT, padx=3)

        result_var = tk.StringVar(value="")
        ttk.Label(dlg, textvariable=result_var, foreground=self.pal.get("accent_dark", "#1D4ED8")).pack(anchor=tk.W, padx=14, pady=(6, 0))

        def apply_changes(close_after: bool = False):
            tokens = [t.strip() for t in re.split(r"[,;]+", layer_var.get() or "") if t.strip()]
            # Hỗ trợ thói quen gõ cách nếu người dùng không dùng dấu phẩy, nhưng dấu phẩy vẫn là khuyến nghị chính.
            if len(tokens) == 1 and " " in tokens[0].strip():
                tokens = [t.strip() for t in tokens[0].split() if t.strip()]
            if not tokens:
                messagebox.showwarning("Chỉnh loại đất", "Cần nhập ít nhất một tên/mã lớp.", parent=dlg)
                return
            m = re.match(r"\s*(\d+)", type_var.get() or "")
            if not m:
                messagebox.showwarning("Chỉnh loại đất", "Chưa chọn loại đất mới.", parent=dlg)
                return
            new_type = m.group(1)
            prefix = bool(match_prefix_var.get())
            rows = self.layer_table.get_rows()
            changed = 0
            tokens_norm = [_strip_accents(t).lower() for t in tokens]
            for rr in rows:
                while len(rr) < len(self.layer_table.columns):
                    rr.append("")
                lname = _strip_accents(str(rr[1] or "")).lower().strip()
                if not lname:
                    continue
                ok = False
                for tok in tokens_norm:
                    if prefix:
                        if lname.startswith(tok):
                            ok = True
                            break
                    else:
                        if lname == tok:
                            ok = True
                            break
                if ok and str(rr[3]).strip() != new_type:
                    rr[3] = new_type
                    changed += 1
            if changed:
                self.layer_table.set_rows(rows, record_undo=True)
                self._merge_first_col_visual(self.layer_table)
            result_var.set(f"Đã chỉnh loại đất cho {changed} dòng: {', '.join(tokens)} → loại {new_type}.")
            self._set_status(result_var.get())
            if close_after:
                dlg.destroy()

        btns = ttk.Frame(dlg)
        btns.pack(fill=tk.X, padx=14, pady=14)
        ttk.Button(btns, text="Apply", command=lambda: apply_changes(False)).pack(side=tk.RIGHT, padx=4)
        ttk.Button(btns, text="OK", style="Accent.TButton", command=lambda: apply_changes(True)).pack(side=tk.RIGHT, padx=4)
        ttk.Button(btns, text="Hủy", command=dlg.destroy).pack(side=tk.RIGHT, padx=4)
        self.root.wait_window(dlg)

    def _paste_geology_clipboard(self):
        self.layer_table.paste_from_clipboard(self.root)
        self._merge_first_col_visual(self.layer_table)

    def _add_layer_rows_dialog(self):
        n = simpledialog.askinteger("Thêm dòng", "Số dòng cần thêm:", initialvalue=10, minvalue=1, maxvalue=500)
        if n:
            self.layer_table.add_blank_rows(n)

    def _add_geology_item_dialog(self):
        item = simpledialog.askstring("Thêm LK địa chất", "Tên lỗ khoan địa chất:", parent=self.root)
        if item is None:
            return
        item = item.strip()
        if not item:
            messagebox.showwarning("Thiếu LK địa chất", "Cần nhập tên lỗ khoan địa chất.")
            return
        n = simpledialog.askinteger("Thêm LK địa chất", f"Số lớp đất/đá của lỗ khoan {item}:", initialvalue=5, minvalue=1, maxvalue=200, parent=self.root)
        if not n:
            return
        existing = self.layer_table.get_rows()
        new_rows = []
        for i in range(int(n)):
            # Treeview không hỗ trợ merge cell thật; tool mô phỏng bằng cách chỉ ghi tên hạng mục ở dòng đầu,
            # các dòng sau để trống và parser sẽ tự hiểu thuộc cùng hạng mục phía trên.
            new_rows.append([item if i == 0 else "", f"L{i+1}", "", "1", "", "18", "", "", "", "", "", "", "", ""])
        self.layer_table.set_rows(existing + new_rows, record_undo=True)
        self._merge_first_col_visual(self.layer_table)

    def copy_selected_geology_item(self):
        """Copy toàn bộ nhóm địa chất của hạng mục chứa dòng đang chọn.

        Bảng địa chất đang mô phỏng merge ô hạng mục: chỉ dòng đầu có tên hạng mục,
        các dòng sau có thể trống. Hàm này xác định hạng mục hiệu lực của dòng được chọn,
        copy toàn bộ các lớp thuộc hạng mục đó và cho phép đặt tên hạng mục mới.
        """
        sel = list(self.layer_table.tree.selection())
        if not sel:
            messagebox.showwarning("Copy LK địa chất", "Chọn một dòng thuộc lỗ khoan địa chất cần copy.")
            return
        children = self.layer_table.data_children()
        first = next((rowid for rowid in children if rowid in sel), sel[0])
        all_rows = self.layer_table.get_rows()
        try:
            first_idx = children.index(first)
        except ValueError:
            first_idx = 0

        # Tìm tên hạng mục hiệu lực tại dòng chọn.
        current_item = ""
        for i in range(0, min(first_idx, len(all_rows)) + 1):
            if i < len(all_rows):
                name = _display_item_name(all_rows[i][0] if all_rows[i] else "")
                if name:
                    current_item = name
        if not current_item:
            messagebox.showwarning("Copy LK địa chất", "Không xác định được tên lỗ khoan của dòng đang chọn.")
            return
        key = _normalize_item_name(current_item)

        # Lấy toàn bộ dòng thuộc cùng hạng mục hiệu lực.
        group_rows: List[List[str]] = []
        group_indices: List[int] = []
        eff = ""
        for i, row in enumerate(all_rows):
            rr = list(row)
            if not any(str(x).strip() for x in rr):
                continue
            name = _display_item_name(rr[0] if rr else "")
            if name:
                eff = name
            if _normalize_item_name(eff) == key:
                group_rows.append(rr)
                group_indices.append(i)
        if not group_rows:
            messagebox.showwarning("Copy LK địa chất", "Không tìm thấy dòng địa chất thuộc lỗ khoan đã chọn.")
            return

        default_name = current_item + " (1)"
        new_name = simpledialog.askstring("Copy LK địa chất", "Tên lỗ khoan địa chất mới:", initialvalue=default_name, parent=self.root)
        if new_name is None:
            return
        new_name = new_name.strip() or default_name

        copied: List[List[Any]] = []
        for i, row in enumerate(group_rows):
            rr = list(row)
            while len(rr) < len(self.layer_table.columns):
                rr.append("")
            rr[0] = new_name if i == 0 else ""
            copied.append(rr)

        insert_idx = max(group_indices) + 1 if group_indices else len(all_rows)
        self.layer_table.set_rows(all_rows[:insert_idx] + copied + all_rows[insert_idx:], record_undo=True)
        self._merge_first_col_visual(self.layer_table)

    def _context_layer_row_index(self) -> Optional[int]:
        """Index dòng dữ liệu địa chất đang được chuột phải/chọn, bỏ qua dòng kẻ mô phỏng nếu có."""
        if not hasattr(self, "layer_table"):
            return None
        children = self.layer_table.data_children()
        rowid = getattr(self.layer_table, "_context_rowid", "") or ""
        if rowid not in children:
            sel = list(self.layer_table.tree.selection())
            rowid = next((r for r in children if r in sel), "")
        if not rowid:
            return None
        try:
            return children.index(rowid)
        except ValueError:
            return None

    def _geology_item_bounds_for_index(self, rows: List[List[Any]], idx: int) -> Optional[Tuple[str, int, int]]:
        """Tìm hạng mục hiệu lực và khoảng dòng liên tục [start, end] chứa idx.

        Bảng địa chất đang mô phỏng merge cột Hạng mục: dòng đầu nhóm có tên, các dòng sau để trống.
        Hàm này chỉ lấy nhóm liên tục chứa dòng đang bấm, không gom các nhóm cùng tên nằm rải rác.
        """
        if idx is None or idx < 0 or idx >= len(rows):
            return None
        start = idx
        item = ""
        for i in range(idx, -1, -1):
            name = _display_item_name(rows[i][0] if rows[i] else "")
            if name:
                start = i
                item = name
                break
        if not item:
            return None
        end = len(rows) - 1
        for j in range(start + 1, len(rows)):
            name = _display_item_name(rows[j][0] if rows[j] else "")
            if name:
                end = j - 1
                break
        return item, start, end

    def _unique_geology_copy_name(self, base_name: str) -> str:
        base = str(base_name or "LK").strip() or "LK"
        candidate = f"{base}_copy"
        existing = set()
        for row in self.layer_table.get_rows():
            name = _display_item_name(row[0] if row else "")
            if name:
                existing.add(_normalize_item_name(name))
        if _normalize_item_name(candidate) not in existing:
            return candidate
        k = 2
        while True:
            candidate = f"{base}_copy{k}"
            if _normalize_item_name(candidate) not in existing:
                return candidate
            k += 1

    def copy_context_geology_item_auto(self):
        """Menu chuột phải: copy toàn bộ hạng mục chứa dòng đang bấm xuống dưới, tên mới thêm _copy."""
        idx = self._context_layer_row_index()
        rows = self.layer_table.get_rows()
        bounds = self._geology_item_bounds_for_index(rows, idx if idx is not None else -1)
        if not bounds:
            messagebox.showwarning("Copy lỗ khoan địa chất", "Không xác định được lỗ khoan địa chất của dòng đang chọn.")
            return
        item, start, end = bounds
        new_name = self._unique_geology_copy_name(item)
        copied: List[List[Any]] = []
        for k, row in enumerate(rows[start:end + 1]):
            rr = list(row)
            while len(rr) < len(self.layer_table.columns):
                rr.append("")
            rr[0] = new_name if k == 0 else ""
            copied.append(rr)
        insert_idx = end + 1
        self.layer_table.set_rows(rows[:insert_idx] + copied + rows[insert_idx:], record_undo=True)
        self._merge_first_col_visual(self.layer_table)
        self._set_status(f"Đã copy lỗ khoan địa chất {item} thành {new_name}")

    def _insert_context_geology_row(self, above: bool):
        idx = self._context_layer_row_index()
        rows = self.layer_table.get_rows()
        if idx is None:
            messagebox.showwarning("Thêm dòng", "Hãy bấm chuột phải vào dòng địa chất cần chèn.")
            return
        ncol = len(self.layer_table.columns)
        new_row = ["" for _ in range(ncol)]
        bounds = self._geology_item_bounds_for_index(rows, idx)
        insert_idx = idx if above else idx + 1
        if bounds:
            item, start, end = bounds
            # Nếu chèn phía trên dòng đầu nhóm, dòng mới phải giữ tên hạng mục để nhóm không bị tách.
            if above and idx == start:
                new_row[0] = item
                rr = list(rows[start])
                while len(rr) < ncol:
                    rr.append("")
                rr[0] = ""
                rows[start] = rr
        self.layer_table.set_rows(rows[:insert_idx] + [new_row] + rows[insert_idx:], record_undo=True)
        self._merge_first_col_visual(self.layer_table)

    def insert_context_geology_row_above(self):
        self._insert_context_geology_row(above=True)

    def insert_context_geology_row_below(self):
        self._insert_context_geology_row(above=False)

    def delete_context_geology_item(self):
        """Menu chuột phải: xóa toàn bộ hạng mục chứa dòng đang bấm, không ảnh hưởng hạng mục khác."""
        idx = self._context_layer_row_index()
        rows = self.layer_table.get_rows()
        bounds = self._geology_item_bounds_for_index(rows, idx if idx is not None else -1)
        if not bounds:
            messagebox.showwarning("Xóa lỗ khoan địa chất", "Không xác định được lỗ khoan địa chất của dòng đang chọn.")
            return
        item, start, end = bounds
        n = end - start + 1
        if not messagebox.askyesno("Xóa lỗ khoan địa chất", f"Xóa toàn bộ {n} dòng của lỗ khoan '{item}'?"):
            return
        self.layer_table.set_rows(rows[:start] + rows[end + 1:], record_undo=True)
        self._merge_first_col_visual(self.layer_table)
        self._set_status(f"Đã xóa lỗ khoan địa chất {item}")

    def _current_item_names(self) -> List[str]:
        """Danh sách hạng mục đang có trong bảng Thông tin cọc tính toán."""
        names: List[str] = []
        seen = set()
        for idx, row in enumerate(self.item_table.get_rows() if hasattr(self, "item_table") else []):
            name = str(row[0] if row else "").strip() or f"HM{idx+1}"
            key = _normalize_item_name(name)
            if key and key not in seen:
                names.append(name)
                seen.add(key)
        return names

    def _dialog_tree_rows_by_item(self, data_map: Dict[str, Dict[str, str]], cols_kind: str) -> List[List[str]]:
        rows: List[List[str]] = []
        for name in self._current_item_names():
            key = _normalize_item_name(name)
            d = data_map.get(key, {})
            if cols_kind == "dd":
                rows.append([name, d.get("top", ""), d.get("bottom", ""), d.get("gamma", "Tự động") or "Tự động"])
            else:
                rows.append([name, d.get("uplift_cd", ""), d.get("uplift_db", "")])
        # Giữ lại dữ liệu của hạng mục đã nhập nhưng hiện không còn trong bảng, để người dùng thấy và quyết định xóa/sửa.
        current_keys = {_normalize_item_name(n) for n in self._current_item_names()}
        for key, d in list(data_map.items()):
            if key in current_keys:
                continue
            name = d.get("item", key)
            if cols_kind == "dd":
                rows.append([name, d.get("top", ""), d.get("bottom", ""), d.get("gamma", "Tự động") or "Tự động"])
            else:
                rows.append([name, d.get("uplift_cd", ""), d.get("uplift_db", "")])
        return rows

    def show_downdrag_item_dialog(self):
        """Nhập vùng ma sát âm riêng theo từng hạng mục.

        Tab Thông số chung chỉ bật/tắt xét ma sát âm; vùng ma sát âm phụ thuộc địa chất/cao độ từng mố trụ
        nên được nhập ở tab Thông tin cọc tính toán qua hộp thoại này.
        """
        if not self._current_item_names():
            messagebox.showwarning("Thông số tính ma sát âm", "Chưa có hạng mục trong bảng Thông tin cọc tính toán.")
            return
        win = tk.Toplevel(self.root)
        win.title("Thông số tính ma sát âm theo hạng mục")
        win.geometry("760x430")
        try:
            win.transient(self.root)
            win.grab_set()
        except Exception:
            pass
        frm = ttk.Frame(win, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frm, text="Thông số tính ma sát âm theo từng hạng mục", font=("Arial", 15, "bold"), foreground=self.pal["accent_dark"]).pack(anchor=tk.W, pady=(0, 6))
        ttk.Label(
            frm,
            text="Nhập cao độ đỉnh/đáy vùng ma sát âm cho từng hạng mục. γDD để trống hoặc ghi 'Tự động' thì lấy theo TCVN 11823-3, Bảng 4.",
            style="Muted.TLabel", wraplength=720, justify=tk.LEFT,
        ).pack(anchor=tk.W, pady=(0, 8))
        table_frame = ttk.Frame(frm)
        table_frame.pack(fill=tk.BOTH, expand=True)
        cols = [("item", "Hạng mục", 150), ("top", "CĐ đỉnh MS âm", 140), ("bottom", "CĐ đáy MS âm", 140), ("gamma", "γDD", 120)]
        tree = EditableTree(table_frame, cols, height=10)
        tree.set_rows(self._dialog_tree_rows_by_item(self.item_downdrag_data, "dd"))
        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(10, 0))
        def apply_close():
            new_map: Dict[str, Dict[str, str]] = {}
            for row in tree.get_rows():
                vals = (list(row) + [""] * 4)[:4]
                item, top, bottom, gamma = [str(x).strip() for x in vals]
                if not item:
                    continue
                # Chỉ lưu khi có khai báo vùng MS âm. γDD riêng mà không có vùng thì không có ý nghĩa tính toán.
                if top or bottom:
                    new_map[_normalize_item_name(item)] = {
                        "item": item,
                        "top": top,
                        "bottom": bottom,
                        "gamma": gamma or "Tự động",
                    }
            self.item_downdrag_data = new_map
            self._set_status(f"Đã cập nhật thông số ma sát âm cho {len(new_map)} hạng mục")
            win.destroy()
        ttk.Button(btns, text="Thêm dòng", command=lambda: tree.add_blank_rows(1)).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Xóa dòng", command=tree.delete_selected_row).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="OK", style="Accent.TButton", command=apply_close).pack(side=tk.RIGHT, padx=3)
        ttk.Button(btns, text="Hủy", command=win.destroy).pack(side=tk.RIGHT, padx=3)
        self._center_window(win)

    def show_uplift_item_dialog(self):
        """Nhập lực nhổ riêng theo từng hạng mục, không kéo dài bảng cọc tính toán."""
        if not self._current_item_names():
            messagebox.showwarning("Điền lực nhổ", "Chưa có hạng mục trong bảng Thông tin cọc tính toán.")
            return
        win = tk.Toplevel(self.root)
        win.title("Điền lực nhổ theo hạng mục")
        win.geometry("660x390")
        try:
            win.transient(self.root)
            win.grab_set()
        except Exception:
            pass
        frm = ttk.Frame(win, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frm, text="Điền lực nhổ theo từng hạng mục", font=("Arial", 15, "bold"), foreground=self.pal["accent_dark"]).pack(anchor=tk.W, pady=(0, 6))
        ttk.Label(frm, text="Nhập lực nhổ lớn nhất của 1 cọc theo từng trạng thái giới hạn. Để trống nếu hạng mục không có lực nhổ.", style="Muted.TLabel", wraplength=620, justify=tk.LEFT).pack(anchor=tk.W, pady=(0, 8))
        table_frame = ttk.Frame(frm)
        table_frame.pack(fill=tk.BOTH, expand=True)
        cols = [("item", "Hạng mục", 170), ("uplift_cd", "Tu CĐ/1 cọc (kN)", 160), ("uplift_db", "Tu ĐB/1 cọc (kN)", 160)]
        tree = EditableTree(table_frame, cols, height=10)
        tree.set_rows(self._dialog_tree_rows_by_item(self.item_uplift_data, "uplift"))
        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(10, 0))
        def apply_close():
            new_map: Dict[str, Dict[str, str]] = {}
            for row in tree.get_rows():
                vals = (list(row) + [""] * 3)[:3]
                item, u_cd, u_db = [str(x).strip() for x in vals]
                if not item:
                    continue
                if u_cd or u_db:
                    new_map[_normalize_item_name(item)] = {"item": item, "uplift_cd": u_cd, "uplift_db": u_db}
            self.item_uplift_data = new_map
            self._set_status(f"Đã cập nhật lực nhổ cho {len(new_map)} hạng mục")
            win.destroy()
        ttk.Button(btns, text="Thêm dòng", command=lambda: tree.add_blank_rows(1)).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="Xóa dòng", command=tree.delete_selected_row).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="OK", style="Accent.TButton", command=apply_close).pack(side=tk.RIGHT, padx=3)
        ttk.Button(btns, text="Hủy", command=win.destroy).pack(side=tk.RIGHT, padx=3)
        self._center_window(win)

    def _add_item_rows_dialog(self):
        n = simpledialog.askinteger("Thêm dòng", "Số dòng hạng mục cần thêm:", initialvalue=5, minvalue=1, maxvalue=500)
        if n:
            self.item_table.add_blank_rows(n)

    def _item_names_from_table(self) -> List[str]:
        names: List[str] = []
        seen = set()
        for row in self.item_table.get_rows():
            if not any(str(x).strip() for x in row):
                continue
            name = _display_item_name(row[0] if row else "")
            if not name:
                continue
            key = _normalize_item_name(name)
            if key and key not in seen:
                seen.add(key)
                names.append(name)
        return names

    def _geology_rows_with_effective_items(self) -> List[Tuple[str, List[str]]]:
        """Trả về các dòng địa chất kèm hạng mục hiệu lực.

        Bảng địa chất mô phỏng merge ô: chỉ dòng đầu của một hạng mục có tên,
        các dòng sau có thể để trống. Hàm này truyền tên hạng mục gần nhất xuống các dòng trống.
        """
        out: List[Tuple[str, List[str]]] = []
        current = ""
        for row in self.layer_table.get_rows():
            if not any(str(x).strip() for x in row):
                continue
            rr = list(row)
            geo_item = _display_item_name(rr[0] if rr else "")
            if geo_item:
                current = geo_item
            effective = geo_item or current
            out.append((effective, rr))
        return out

    def _geology_item_names_from_table(self) -> List[str]:
        names: List[str] = []
        seen = set()
        for effective, _row in self._geology_rows_with_effective_items():
            name = _display_item_name(effective)
            key = _normalize_item_name(name)
            if key and key not in seen:
                seen.add(key)
                names.append(name)
        return names

    def _matching_diagnostics(self) -> Tuple[List[str], List[str], List[str]]:
        """Kiểm tra khớp tên giữa tab hạng mục và tab địa chất.

        Trả về (missing_geology_for_items, extra_geology_items, duplicate_geo_keys).
        Khớp theo key chuẩn hóa: bỏ dấu tiếng Việt, bỏ khoảng trắng/ký tự đặc biệt, không phân biệt hoa thường.
        """
        if bool(self.vars.get("common_geology") and self.vars["common_geology"].get()):
            return [], [], []
        item_names = self._item_names_from_table()
        geo_names = self._geology_item_names_from_table()
        item_keys = {_normalize_item_name(n): n for n in item_names if _normalize_item_name(n)}
        geo_key_to_names: Dict[str, List[str]] = {}
        for n in geo_names:
            k = _normalize_item_name(n)
            if k:
                geo_key_to_names.setdefault(k, []).append(n)
        missing = [name for k, name in item_keys.items() if k not in geo_key_to_names and k not in ("", "*", "all", "chung")]
        extra = [names[0] for k, names in geo_key_to_names.items() if k not in item_keys and k not in ("", "*", "all", "chung")]
        duplicate = [" / ".join(names) for k, names in geo_key_to_names.items() if len(names) > 1]
        return missing, extra, duplicate

    def _show_item_geology_mapping_dialog(self, missing: List[str], extra: List[str]) -> bool:
        """Cho phép khớp hạng mục tính toán với lỗ khoan địa chất khác.

        Dùng khi số hạng mục tính toán và số lỗ khoan địa chất không trùng nhau, ví dụ 4 hạng mục
        nhưng chỉ có 3 hố khoan/địa chất, một lỗ khoan địa chất dùng chung cho 2 hạng mục.
        """
        item_names = self._item_names_from_table()
        geo_names = self._geology_item_names_from_table()
        if not item_names or not geo_names:
            messagebox.showerror("Khớp hạng mục/địa chất", "Chưa có đủ hạng mục tính toán hoặc lỗ khoan địa chất để khớp.")
            return False
        win = tk.Toplevel(self.root)
        win.title("Khớp hạng mục tính toán với địa chất")
        win.geometry("760x420")
        win.transient(self.root)
        win.grab_set()
        ttk.Label(win, text="Chọn lỗ khoan địa chất dùng cho từng hạng mục tính toán", font=("Arial", 10, "bold")).pack(anchor="w", padx=12, pady=(10, 4))
        ttk.Label(win, text="Cột trái là hạng mục cần kiểm toán. Cột phải là lỗ khoan địa chất trong tab Thông số địa chất. Có thể chọn cùng một lỗ khoan địa chất cho nhiều hạng mục.", wraplength=720).pack(anchor="w", padx=12, pady=(0, 8))
        frame = ttk.Frame(win)
        frame.pack(fill=tk.BOTH, expand=True, padx=12, pady=6)
        ttk.Label(frame, text="Hạng mục cần tính toán", font=("Arial", 9, "bold")).grid(row=0, column=0, sticky="ew", padx=4, pady=4)
        ttk.Label(frame, text="Lỗ khoan địa chất", font=("Arial", 9, "bold")).grid(row=0, column=1, sticky="ew", padx=4, pady=4)
        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(1, weight=1)
        combos: Dict[str, ttk.Combobox] = {}
        geo_key_to_display = {_normalize_item_name(g): g for g in geo_names}
        for i, item in enumerate(item_names, start=1):
            ttk.Label(frame, text=item).grid(row=i, column=0, sticky="ew", padx=4, pady=3)
            cb = ttk.Combobox(frame, values=geo_names, state="readonly")
            ik = _normalize_item_name(item)
            current = self.item_geo_map_display.get(ik, "")
            if current in geo_names:
                cb.set(current)
            elif ik in geo_key_to_display:
                cb.set(geo_key_to_display[ik])
            elif geo_names:
                cb.set(geo_names[0])
            cb.grid(row=i, column=1, sticky="ew", padx=4, pady=3)
            combos[item] = cb
        result = {"ok": False}
        def apply_mapping():
            self.item_geo_map.clear()
            self.item_geo_map_display.clear()
            for item, cb in combos.items():
                ik = _normalize_item_name(item)
                gv = cb.get().strip()
                gk = _normalize_item_name(gv)
                if ik and gk:
                    self.item_geo_map[ik] = gk
                    self.item_geo_map_display[ik] = gv
            result["ok"] = True
            win.destroy()
        btns = ttk.Frame(win)
        btns.pack(fill=tk.X, padx=12, pady=10)
        ttk.Button(btns, text="Áp dụng khớp nối", command=apply_mapping).pack(side=tk.RIGHT, padx=4)
        ttk.Button(btns, text="Hủy", command=win.destroy).pack(side=tk.RIGHT, padx=4)
        self.root.wait_window(win)
        return bool(result["ok"])

    def open_item_geology_mapping_dialog(self) -> bool:
        """Mở lại bảng khớp hạng mục - lỗ khoan địa chất để người dùng chỉnh sửa bất kỳ lúc nào.

        Khác với check_item_geology_matching(), hàm này không chỉ báo "đã khớp" khi tên trùng,
        mà luôn mở bảng chọn để có thể đổi lại các cặp khớp nhầm như HM1-LK2, HM2-LK1.
        """
        common = bool(self.vars.get("common_geology") and self.vars["common_geology"].get())
        if common:
            messagebox.showinfo("Khớp hạng mục/địa chất", "Đang bật tùy chọn dùng chung một LK địa chất cho toàn bộ hạng mục. Tắt tùy chọn này nếu muốn khớp riêng từng hạng mục.")
            return True
        missing, extra, _duplicate = self._matching_diagnostics()
        return self._show_item_geology_mapping_dialog(missing, extra)

    def check_item_geology_matching(self, quiet: bool = False) -> bool:
        common = bool(self.vars.get("common_geology") and self.vars["common_geology"].get())
        if common:
            msg = "Đang bật tùy chọn: Một LK địa chất dùng chung cho toàn bộ hạng mục. Chương trình sẽ dùng cùng một bảng địa chất cho tất cả hạng mục."
            if not quiet:
                messagebox.showinfo("Khớp hạng mục/địa chất", msg)
            return True
        missing, extra, duplicate = self._matching_diagnostics()
        # Nếu đã có map thủ công thì chỉ cần mỗi hạng mục có một key địa chất tồn tại.
        if missing and self.item_geo_map:
            geo_keys = {_normalize_item_name(n) for n in self._geology_item_names_from_table()}
            missing = [m for m in missing if self.item_geo_map.get(_normalize_item_name(m), "") not in geo_keys]
        if not missing and not duplicate:
            if not quiet:
                msg = "Đã khớp: mỗi hạng mục trong tab 2 có lỗ khoan địa chất tương ứng."
                if self.item_geo_map_display:
                    msg += "\n\nKhớp nối thủ công:\n" + "\n".join(f"- {k}: {v}" for k, v in self.item_geo_map_display.items())
                messagebox.showinfo("Khớp hạng mục/địa chất", msg)
            return True
        if quiet:
            return False
        lines = []
        if missing:
            lines.append("Hạng mục chưa có lỗ khoan địa chất trùng tên:")
            lines.extend([f"- {x}" for x in missing])
        if extra:
            if lines:
                lines.append("")
            lines.append("Lỗ khoan địa chất có tên nhưng không có trong tab thông tin riêng:")
            lines.extend([f"- {x}" for x in extra])
        if duplicate:
            if lines:
                lines.append("")
            lines.append("Tên lỗ khoan địa chất bị trùng sau khi chuẩn hóa:")
            lines.extend([f"- {x}" for x in duplicate])
        lines.append("")
        lines.append("Có mở bảng khớp nối để chọn lỗ khoan địa chất tương ứng cho từng hạng mục không?")
        if messagebox.askyesno("Chưa khớp hạng mục/địa chất", "\n".join(lines)):
            return self._show_item_geology_mapping_dialog(missing, extra)
        return False

    def _auto_mode_from_layers(self, layers: List[SoilLayer], pile_type: str) -> str:
        if "đóng" in str(pile_type).lower() or "ép" in str(pile_type).lower() or "ep" in str(pile_type).lower():
            return "Cọc đóng"
        if layers and int(layers[-1].soil_type) in (3, 4):
            return "Cọc khoan trong đá"
        return "Cọc khoan trong đất"

    def _parse_layer_rows(self, only_item: str = "") -> List[SoilLayer]:
        layers: List[SoilLayer] = []
        common = bool(self.vars.get("common_geology") and self.vars["common_geology"].get())
        wanted_key = _normalize_item_name(only_item)
        # Nếu người dùng đã khớp thủ công hạng mục tính toán với địa chất khác, dùng key địa chất đó.
        target_geo_key = self.item_geo_map.get(wanted_key, wanted_key)
        for effective_item, row in self._geology_rows_with_effective_items():
            # V0.2.8: luôn khớp theo tên hạng mục đã chuẩn hóa.
            # Dòng địa chất trống ô hạng mục được hiểu thuộc hạng mục gần nhất phía trên.
            effective_key = _normalize_item_name(effective_item)
            if (not common) and wanted_key and effective_key not in (target_geo_key, "", "all", "chung"):
                continue
            geo_item, name, bottom, stype, n_in, gamma, cval, phi, qu, rqd, gsi, mi, dist, comment = (list(row) + [""] * 14)[:14]
            try:
                stype_i = _safe_int(stype, 1)
                n_raw = str(n_in or "").strip()
                gamma_raw = str(gamma or "").strip()
                qu_raw = str(qu or "").strip()
                rqd_raw = str(rqd or "").strip()
                rqd_val = _safe_float(rqd, 50.0)
                # QA-GEO RQD: với lớp đá phong hóa/nứt vỡ, RQD<20 ứng xử như IGM/đá mềm (loại 5);
                # RQD>=20 mới giữ là đá phong hóa/nứt vỡ (loại 4). Chỉ tự đổi khi người dùng có nhập RQD.
                if stype_i in (4, 5) and rqd_raw != "":
                    stype_i = 5 if rqd_val < 20.0 else 4
                gsi_raw = str(gsi or "").strip()
                mi_raw = str(mi or "").strip()
                dist_raw = str(dist or "").strip()
                ly = SoilLayer(
                    name=str(name).strip(), bottom_elev_m=_safe_float(bottom, 0.0), soil_type=stype_i,
                    n_spt=(0.0 if stype_i == 0 else _safe_float(n_in, 0.0)), gamma_kN_m3=(0.0 if stype_i == 0 else _safe_float(gamma, 18.0)),
                    c_mpa=(0.0 if stype_i == 0 else _safe_float(cval, 0.0)), phi_deg=(0.0 if stype_i == 0 else _safe_float(phi, 0.0)), qu_mpa=(0.0 if stype_i == 0 else _safe_float(qu, 0.0)),
                    rqd=rqd_val, gsi=_safe_float(gsi, 30.0), mi=_safe_float(mi, 9.0),
                    disturbance=_safe_float(dist, 0.5), comment=str(comment).strip()
                )
                # Gắn cờ thiếu dữ liệu từ raw cell để bước tính toán có thể cảnh báo nhưng vẫn giữ giá trị mặc định nội bộ.
                setattr(ly, "_missing_spt", stype_i != 0 and n_raw == "")
                setattr(ly, "_missing_gamma", stype_i != 0 and gamma_raw == "")
                setattr(ly, "_missing_rock_data", stype_i in (3, 4) and (qu_raw == "" or rqd_raw == "" or gsi_raw == "" or mi_raw == "" or dist_raw == "" or _safe_float(qu_raw, 0.0) <= 0.0))
                layers.append(ly)
            except Exception:
                continue
        return layers

    def _parse_downdrag_factor_text(self, value: Any, inp: PileInput) -> float:
        """Đọc γDD từ giao diện; để trống/Tự động thì lấy theo TCVN 11823-3 Bảng 4."""
        txt = _strip_accents(str(value or "")).strip().lower()
        if (not txt) or ("tu dong" in txt) or ("auto" in txt):
            return SCTCalculator.default_downdrag_load_factor(inp)
        val = _safe_float(txt, 0.0)
        if val <= 0.0:
            return SCTCalculator.default_downdrag_load_factor(inp)
        return val

    def _base_input_from_vars(self) -> PileInput:
        """Thông tin chung dùng cho mọi hạng mục.

        Các dữ liệu hình học/cao độ/nội lực/loại cọc không đọc ở đây nữa;
        chúng được lấy từ bảng thông tin riêng từng hạng mục.
        """
        inp = PileInput(
            project=self.vars["project"].get(),
            concrete_gamma_kN_m3=_safe_float(self.vars["concrete_gamma"].get(), 24.5),
            fc_mpa=_safe_float(self.vars["fc_mpa"].get(), 30.0),
            fy_mpa=_safe_float(self.vars["fy_mpa"].get(), 400.0),
            n_rebars=_safe_int(self.vars["n_rebars"].get(), 0),
            rebar_dia_mm=_safe_float(self.vars["rebar_dia"].get(), 0.0),
            stirrup_type=_safe_int(self.vars["stirrup_type"].get(), 1),
            exclude_top_bored_m=_safe_float(self.vars["exclude_top"].get(), 1.5),
            spt_er_percent=_safe_float(self.vars.get("spt_er", tk.StringVar(value="60")).get(), 60.0),
            spt_input_mode=(self.vars.get("spt_input_mode", tk.StringVar(value="Nₕₜ")).get()),
            sand_preconsolidation_mode=self.vars["sand_mode"].get(),
            sand_m=_safe_float(self.vars["sand_m"].get(), 0.6),
            igm_alpha=_safe_float(self.vars.get("igm_alpha", tk.StringVar(value="0.25")).get(), 0.25),
            igm_joint_factor=_safe_float(self.vars.get("igm_joint_factor", tk.StringVar(value="0.45")).get(), 0.45),
            igm_missing_qu_policy="require_qu",
            rock_side_method=(self.vars.get("rock_side_method").get() if self.vars.get("rock_side_method") else "fractured"),
            rock_construction_condition=(self.vars.get("rock_construction_condition").get() if self.vars.get("rock_construction_condition") else "Có chống đỡ"),
            rock_joint_condition=(self.vars.get("rock_joint_condition").get() if self.vars.get("rock_joint_condition") else "Khe nứt hở hoặc có mùn"),
            rock_open_joint=(False if "khép" in str(self.vars.get("rock_joint_condition", tk.StringVar(value="")).get()).lower() else True),
            rock_tip_condition=(self.vars.get("rock_tip_condition").get() if self.vars.get("rock_tip_condition") else "fractured"),
            include_rock_tip=bool(self.vars.get("include_rock_tip") and self.vars["include_rock_tip"].get()),
            allow_rock_tip_exceed_25qu=bool(self.vars.get("allow_rock_tip_exceed_25qu") and self.vars["allow_rock_tip_exceed_25qu"].get()),
            clay_use_c_phi=bool(self.vars.get("clay_use_c_phi") and self.vars["clay_use_c_phi"].get()),
            ignore_group_igm_rock=bool(self.vars.get("ignore_group_igm_rock") and self.vars["ignore_group_igm_rock"].get()),
            allow_geology_extrapolation=bool(self.vars.get("allow_geology_extrapolation") and self.vars["allow_geology_extrapolation"].get()),
            include_downdrag=bool(self.vars.get("include_downdrag") and self.vars["include_downdrag"].get()),
            downdrag_top_elev_m=0.0,
            downdrag_bottom_elev_m=0.0,
            downdrag_factor=0.0,
            crack_spacing_mm=_safe_float(self.vars["crack_spacing"].get(), 50.0),
            crack_width_mm=_safe_float(self.vars["crack_width"].get(), 5.0),
        )
        return inp

    def _read_multi_inputs(self) -> List[PileInput]:
        base = self._base_input_from_vars()
        items: List[PileInput] = []
        for row in self.item_table.get_rows():
            if not any(str(x).strip() for x in row):
                continue
            raw = list(row)
            # V0.2.13: format chuẩn gồm Cz và lực nhổ.
            vals = (raw + [""] * 19)[:19]
            (item, pile_type, pile_count, layout, D, Ds, spacing, bx, by, cz,
             ground, cap, tip, water, pu_cd, pu_db, ncap_cd, ncap_db, note) = vals
            # Tương thích nhẹ bảng V0.2.12 chưa có Cz/lực nhổ: nếu ô Cz trông giống cao độ mặt đất
            # và các ô cuối chưa có lực nhổ, người dùng nên import lại. Không tự đoán nhiều để tránh lệch cột.
            item_name = str(item).strip() or f"HM{len(items)+1}"
            inp = PileInput(**{**base.__dict__})
            inp.layers = []
            inp.item = item_name
            inp.pile_type = normalize_pile_type_choice(pile_type)
            inp.pile_count_in_group = _safe_int(pile_count, 1)
            inp.group_layout = str(layout).strip() or "2"
            inp.diameter_mm = _safe_float(D, 0.0)
            if inp.driven_shape == "ONG":
                # QA fix P1: với cọc ống (2O/3O) cột Ds được hiểu là ĐƯỜNG KÍNH TRONG; mũi không mở rộng.
                inp.driven_inner_dia_mm = _safe_float(Ds, 0.0)
                inp.tip_diameter_mm = inp.diameter_mm
            else:
                inp.driven_inner_dia_mm = 0.0
                inp.tip_diameter_mm = _safe_float(Ds, inp.diameter_mm if inp.diameter_mm > 0 else 0.0)
            inp.spacing_m = _safe_float(spacing, 3.0)
            inp.cap_width_m = _safe_float(bx, 0.0)
            inp.cap_length_m = _safe_float(by, 0.0)
            inp.cap_thickness_m = _safe_float(cz, 0.0)
            inp.cap_bottom_elev_m = _safe_float(cap, 0.0)
            # Nếu không nhập cao độ thiên nhiên/mặt đất, hiểu bằng cao độ đáy bệ.
            ground_txt = str(ground or "").strip()
            inp.ground_elev_m = _safe_float(ground, inp.cap_bottom_elev_m) if ground_txt else inp.cap_bottom_elev_m
            inp.pile_tip_elev_m = _safe_float(tip, -30.0)
            inp.water_elev_m = _safe_float(water, inp.ground_elev_m)
            inp.force_cd_kn = _safe_float(pu_cd, 0.0)
            inp.force_db_kn = _safe_float(pu_db, 0.0)
            udata = self.item_uplift_data.get(_normalize_item_name(item_name), {})
            inp.uplift_cd_kn = _safe_float(udata.get("uplift_cd", ""), 0.0)
            inp.uplift_db_kn = _safe_float(udata.get("uplift_db", ""), 0.0)
            inp.cap_force_cd_kn = _safe_float(ncap_cd, 0.0)
            inp.cap_force_db_kn = _safe_float(ncap_db, 0.0)
            inp.layers = self._parse_layer_rows(item_name)
            inp.mode = self._auto_mode_from_layers(inp.layers, inp.pile_type)
            inp.analysis_type_auto = inp.mode
            if getattr(base, "include_downdrag", False):
                dd = self.item_downdrag_data.get(_normalize_item_name(item_name), {})
                top_txt = str(dd.get("top", "")).strip()
                bot_txt = str(dd.get("bottom", "")).strip()
                if top_txt or bot_txt:
                    inp.include_downdrag = True
                    inp.downdrag_top_elev_m = _safe_float(top_txt, inp.ground_elev_m)
                    inp.downdrag_bottom_elev_m = _safe_float(bot_txt, inp.pile_tip_elev_m)
                    inp.downdrag_factor = self._parse_downdrag_factor_text(dd.get("gamma", "Tự động"), inp)
                else:
                    inp.include_downdrag = False
                    inp.downdrag_factor = 0.0
            else:
                inp.include_downdrag = False
                inp.downdrag_factor = 0.0
            items.append(inp)
        return items

    def _read_input(self) -> PileInput:
        inputs = self._read_multi_inputs()
        if not inputs:
            raise ValueError("Chưa có hạng mục trong bảng thông tin riêng từng hạng mục.")
        return inputs[0]

    def _find_igm_missing_qu_cases(self, inputs: List[PileInput]) -> List[Tuple[str, str, float]]:
        cases: List[Tuple[str, str, float]] = []
        for inp in inputs:
            for ly in inp.layers:
                if int(getattr(ly, "soil_type", 0) or 0) != 5:
                    continue
                n60 = SCTCalculator._n60_from_input(inp, ly)
                if n60 >= 100.0 and float(getattr(ly, "qu_mpa", 0.0) or 0.0) <= 0.0:
                    cases.append((str(inp.item or ""), str(ly.name or ""), n60))
        return cases

    def _ask_igm_missing_qu_policy(self, cases: List[Tuple[str, str, float]]) -> str:
        """Trả về 'input_qu', 'use_spt' hoặc 'cancel'."""
        dlg = tk.Toplevel(self.root)
        dlg.title("IGM thiếu qu")
        dlg.transient(self.root)
        dlg.grab_set()
        try:
            dlg.resizable(False, False)
        except Exception:
            pass
        frame = ttk.Frame(dlg, padding=14)
        frame.pack(fill=tk.BOTH, expand=True)
        msg = "Có lớp IGM có SPT ≥ 100 nhưng chưa nhập qu.\nChọn cách xử lý:"
        ttk.Label(frame, text=msg, justify="left").pack(anchor="w", pady=(0, 8))
        preview = "\n".join(f"- {item}: lớp {layer}, N60={n:.1f}" for item, layer, n in cases[:8])
        if len(cases) > 8:
            preview += f"\n... và {len(cases)-8} lớp khác"
        ttk.Label(frame, text=preview, justify="left", foreground=self.pal.get("muted", "#555555")).pack(anchor="w", pady=(0, 10))
        result = {"value": "cancel"}
        def choose(value: str):
            result["value"] = value
            dlg.destroy()
        btns = ttk.Frame(frame)
        btns.pack(anchor="e", pady=(8, 0))
        ttk.Button(btns, text="Nhập qu", command=lambda: choose("input_qu")).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Tính theo SPT", command=lambda: choose("use_spt")).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Hủy", command=lambda: choose("cancel")).pack(side=tk.LEFT, padx=4)
        self.root.wait_window(dlg)
        return result.get("value", "cancel")

    def _apply_igm_missing_qu_policy(self, inputs: List[PileInput]) -> bool:
        cases = self._find_igm_missing_qu_cases(inputs)
        if not cases:
            return True
        choice = self._ask_igm_missing_qu_policy(cases)
        if choice == "use_spt":
            for inp in inputs:
                inp.igm_missing_qu_policy = "use_spt"
            return True
        if choice == "input_qu":
            messagebox.showinfo("Nhập qu", "Hãy nhập qu cho các lớp IGM có SPT ≥ 100 rồi bấm Tính toán lại.")
        return False


    # ------------------------------------------------------------------
    # Lọc kết quả và tối ưu chiều dài cọc theo FOS
    # ------------------------------------------------------------------
    def _fos_metric_options(self) -> List[Tuple[str, str]]:
        return [
            ("FOS quyết định chiều dài cọc", "governing"),
            ("FOS nén cọc đơn CĐ", "single_cd"),
            ("FOS nén cọc đơn ĐB", "single_db"),
            ("FOS nén nhóm CĐ", "group_cd"),
            ("FOS nén nhóm ĐB", "group_db"),
            ("FOS nhổ cọc đơn CĐ", "uplift_cd"),
            ("FOS nhổ cọc đơn ĐB", "uplift_db"),
            ("FOS nhổ nhóm CĐ", "uplift_group_cd"),
            ("FOS nhổ nhóm ĐB", "uplift_group_db"),
        ]

    def _fos_metric_labels(self) -> List[str]:
        return [label for label, _ in self._fos_metric_options()]

    def _fos_metric_key_from_label(self, label: Any) -> str:
        label_s = str(label or "").strip()
        for lb, key in self._fos_metric_options():
            if lb == label_s:
                return key
        return "single_cd"

    def _fos_label_from_key(self, key: str) -> str:
        for lb, kk in self._fos_metric_options():
            if kk == key:
                return lb
        return str(key or "")

    def _fos_from_dcr_value(self, dcr: Any) -> float:
        try:
            d = float(dcr or 0.0)
        except Exception:
            d = 0.0
        if d <= 0.0:
            return float("inf")
        return 1.0 / d

    def _base_fos_metric_keys(self) -> List[str]:
        return [
            "single_cd", "single_db", "group_cd", "group_db",
            "uplift_cd", "uplift_db", "uplift_group_cd", "uplift_group_db",
        ]

    def _normalize_fos_metric_keys(self, metric_keys: Any) -> List[str]:
        if metric_keys is None:
            return ["single_cd"]
        if isinstance(metric_keys, str):
            raw = [metric_keys]
        else:
            try:
                raw = list(metric_keys)
            except Exception:
                raw = ["single_cd"]
        out: List[str] = []
        base = self._base_fos_metric_keys()
        for k in raw:
            kk = str(k or "").strip()
            if not kk:
                continue
            if kk == "governing":
                out.extend(base)
            else:
                out.append(kk)
        # Giữ thứ tự, bỏ trùng và bỏ key lạ.
        valid = set(base)
        uniq: List[str] = []
        for k in out:
            if k in valid and k not in uniq:
                uniq.append(k)
        return uniq or ["single_cd"]

    def _fos_value_for_result(self, res: CapacityResult, metric_key: str = "single_cd") -> float:
        if str(metric_key or "") == "governing":
            return self._governing_fos_for_result(res)
        c = self._report_check_values(res)
        mapping = {
            "single_cd": "dcr_single_cd",
            "single_db": "dcr_single_db",
            "group_cd": "dcr_group_cd",
            "group_db": "dcr_group_db",
            "uplift_cd": "dcr_uplift_cd",
            "uplift_db": "dcr_uplift_db",
            "uplift_group_cd": "dcr_uplift_group_cd",
            "uplift_group_db": "dcr_uplift_group_db",
        }
        return self._fos_from_dcr_value(c.get(mapping.get(metric_key, "dcr_single_cd"), 0.0))

    def _governing_fos_for_result(self, res: CapacityResult, metric_keys: Any = None) -> float:
        keys = self._normalize_fos_metric_keys(metric_keys if metric_keys is not None else ["governing"])
        vals: List[float] = []
        for key in keys:
            try:
                v = float(self._fos_value_for_result(res, key))
            except Exception:
                continue
            if math.isnan(v) or v <= 0.0:
                continue
            vals.append(v)
        if not vals:
            return float("inf")
        finite = [v for v in vals if not math.isinf(v)]
        return min(finite) if finite else float("inf")

    def _fos_label_for_keys(self, metric_keys: Any) -> str:
        keys = self._normalize_fos_metric_keys(metric_keys)
        if len(keys) == 1:
            return self._fos_label_from_key(keys[0])
        return "FOS quyết định theo các tiêu chí đã chọn"

    def _fmt_fos_value(self, value: Any, nd: int = 3) -> str:
        try:
            v = float(value)
        except Exception:
            return "-"
        if math.isinf(v):
            return "∞"
        return _fmt(v, nd)

    def _show_result_collection(self, results: List[CapacityResult], status_text: str = ""):
        results = [r for r in (results or []) if r]
        if not results:
            self.summary_text.config(state=tk.NORMAL)
            self.summary_text.delete("1.0", tk.END)
            self.summary_text.insert(tk.END, "Không có kết quả để hiển thị.")
            self.summary_text.config(state=tk.DISABLED)
            self.result_table.set_rows([])
            if status_text:
                self._set_status(status_text)
            return
        if len(results) == 1:
            self._show_result(results[0])
        else:
            self._show_results(results)
        if status_text:
            self._set_status(status_text)

    def apply_result_fos_filter(self):
        results = list(getattr(self, "last_results", None) or ([self.last_result] if self.last_result else []))
        results = [r for r in results if r]
        if not results:
            messagebox.showwarning("Lọc FOS", "Chưa có kết quả. Hãy bấm Tính toán trước.")
            return
        key = self._fos_metric_key_from_label(getattr(self, "result_filter_metric", tk.StringVar(value="")).get())
        mode = str(getattr(self, "result_filter_mode", tk.StringVar(value="Từ a đến b")).get() or "Từ a đến b")
        a_txt = str(getattr(self, "result_filter_a", tk.StringVar(value="")).get() or "").strip()
        b_txt = str(getattr(self, "result_filter_b", tk.StringVar(value="")).get() or "").strip()
        a = _safe_float(a_txt, float("nan")) if a_txt else float("nan")
        b = _safe_float(b_txt, float("nan")) if b_txt else float("nan")
        if mode == "Từ a đến b" and (math.isnan(a) or math.isnan(b)):
            messagebox.showwarning("Lọc FOS", "Chế độ 'Từ a đến b' cần nhập cả a và b.")
            return
        if mode == "Lớn hơn a" and math.isnan(a):
            messagebox.showwarning("Lọc FOS", "Chế độ 'Lớn hơn a' cần nhập a.")
            return
        if mode == "Nhỏ hơn b" and math.isnan(b):
            messagebox.showwarning("Lọc FOS", "Chế độ 'Nhỏ hơn b' cần nhập b.")
            return
        if not math.isnan(a) and not math.isnan(b) and a > b:
            a, b = b, a
        filtered: List[CapacityResult] = []
        for res in results:
            fos = self._fos_value_for_result(res, key)
            ok = False
            if mode == "Từ a đến b":
                ok = (fos >= a and fos <= b)
            elif mode == "Lớn hơn a":
                ok = (fos >= a)
            elif mode == "Nhỏ hơn b":
                ok = (fos <= b)
            if ok:
                filtered.append(res)
        if not filtered:
            self.summary_text.config(state=tk.NORMAL)
            self.summary_text.delete("1.0", tk.END)
            self.summary_text.insert(tk.END, f"Không có hạng mục thỏa điều kiện lọc {self._fos_label_from_key(key)} - {mode}.")
            self.summary_text.config(state=tk.DISABLED)
            self.result_table.set_rows([])
            self._set_status("Lọc FOS: không có kết quả phù hợp")
            return
        self._show_result_collection(filtered, f"Đang lọc {len(filtered)}/{len(results)} hạng mục theo {self._fos_label_from_key(key)}")

    def clear_result_fos_filter(self):
        results = list(getattr(self, "last_results", None) or ([self.last_result] if self.last_result else []))
        results = [r for r in results if r]
        if not results:
            messagebox.showwarning("Bỏ lọc", "Chưa có kết quả để hiển thị.")
            return
        self._show_result_collection(results, f"Đã bỏ lọc, hiển thị {len(results)} hạng mục")

    def open_pile_length_optimization_dialog(self):
        """Tối ưu cao độ mũi/chiều dài cọc theo điều kiện FOS.

V1.0.62: tối ưu chiều dài cọc chỉ xét các điều kiện FOS; bỏ hẳn giao diện ràng buộc lớp mũi/chiều dài ngập tối thiểu."""
        try:
            if not self._validate_required_pile_geometry_before_run():
                return
            if not self.check_item_geology_matching(quiet=True):
                if not self.check_item_geology_matching(quiet=False):
                    return
            inputs = self._read_multi_inputs()
            if not inputs:
                messagebox.showwarning("Tối ưu chiều dài", "Chưa có hạng mục trong bảng thông tin cọc tính toán.")
                return
            if not self._apply_igm_missing_qu_policy(inputs):
                return
        except Exception as exc:
            messagebox.showerror("Tối ưu chiều dài", str(exc))
            return

        dlg = tk.Toplevel(self.root)
        dlg.title("Tối ưu chiều dài cọc")
        dlg.transient(self.root)
        dlg.grab_set()
        try:
            dlg.geometry("980x620")
            dlg.minsize(900, 560)
        except Exception:
            pass
        frm = ttk.Frame(dlg, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)

        # ------------------------------------------------------------------
        # 1) Chọn hạng mục theo kiểu 2 cột.
        # ------------------------------------------------------------------
        ttk.Label(frm, text="1. Các hạng mục cần điều chỉnh", font=("Arial", 11, "bold")).pack(anchor=tk.W, pady=(0, 6))
        select_frame = ttk.Frame(frm)
        select_frame.pack(fill=tk.BOTH, expand=False)
        select_frame.columnconfigure(0, weight=1)
        select_frame.columnconfigure(2, weight=1)

        ttk.Label(select_frame, text="Toàn bộ hạng mục").grid(row=0, column=0, sticky=tk.W, padx=(0, 4))
        ttk.Label(select_frame, text="Hạng mục cần điều chỉnh").grid(row=0, column=2, sticky=tk.W, padx=(4, 0))

        left_frame = ttk.Frame(select_frame)
        left_frame.grid(row=1, column=0, sticky=tk.NSEW, padx=(0, 6))
        mid_frame = ttk.Frame(select_frame)
        mid_frame.grid(row=1, column=1, sticky=tk.NS, padx=4)
        right_frame = ttk.Frame(select_frame)
        right_frame.grid(row=1, column=2, sticky=tk.NSEW, padx=(6, 0))

        lb_all = tk.Listbox(left_frame, selectmode=tk.EXTENDED, height=8, exportselection=False)
        all_ysb = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=lb_all.yview)
        lb_all.configure(yscrollcommand=all_ysb.set)
        lb_all.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        all_ysb.pack(side=tk.RIGHT, fill=tk.Y)

        lb_sel = tk.Listbox(right_frame, selectmode=tk.EXTENDED, height=8, exportselection=False)
        sel_ysb = ttk.Scrollbar(right_frame, orient=tk.VERTICAL, command=lb_sel.yview)
        lb_sel.configure(yscrollcommand=sel_ysb.set)
        lb_sel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sel_ysb.pack(side=tk.RIGHT, fill=tk.Y)

        item_display: List[str] = []
        for i, inp in enumerate(inputs):
            try:
                res0 = SCTCalculator.calculate(inp)
                fos0 = self._fmt_fos_value(self._governing_fos_for_result(res0))
            except Exception:
                fos0 = "-"
            toe_name, toe_embed = self._current_toe_layer_text(inp)
            txt = f"{i+1}. {inp.item} | L={_fmt(inp.pile_length_m,2)} m | CĐ mũi={_fmt(inp.pile_tip_elev_m,2)} | FOSqđ={fos0} | Mũi: {toe_name}, ngập {toe_embed}"
            item_display.append(txt)
            lb_all.insert(tk.END, txt)

        selected_indices: List[int] = []
        constraint_vars: Dict[str, Dict[str, Any]] = {}
        constraint_rows_holder: Dict[str, Any] = {}

        def refresh_selected_list():
            lb_sel.delete(0, tk.END)
            for idx in selected_indices:
                if 0 <= idx < len(item_display):
                    lb_sel.insert(tk.END, item_display[idx])
            # V1.0.61: bỏ phần ràng buộc lớp mũi khi tối ưu.

        def add_selected():
            for idx in lb_all.curselection():
                if idx not in selected_indices:
                    selected_indices.append(int(idx))
            refresh_selected_list()

        def add_all():
            selected_indices[:] = list(range(len(inputs)))
            refresh_selected_list()

        def remove_selected():
            remove_pos = sorted([int(i) for i in lb_sel.curselection()], reverse=True)
            for pos in remove_pos:
                if 0 <= pos < len(selected_indices):
                    selected_indices.pop(pos)
            refresh_selected_list()

        def remove_all():
            selected_indices.clear()
            refresh_selected_list()

        ttk.Button(mid_frame, text=">", width=5, command=add_selected).pack(pady=(18, 4))
        ttk.Button(mid_frame, text="<", width=5, command=remove_selected).pack(pady=4)
        ttk.Button(mid_frame, text=">>", width=5, command=add_all).pack(pady=(14, 4))
        ttk.Button(mid_frame, text="<<", width=5, command=remove_all).pack(pady=4)
        lb_all.bind("<Double-1>", lambda _e: add_selected())
        lb_sel.bind("<Double-1>", lambda _e: remove_selected())

        # ------------------------------------------------------------------
        # 2) Mục tiêu điều chỉnh.
        # ------------------------------------------------------------------
        opts = ttk.LabelFrame(frm, text="2. Mục tiêu điều chỉnh")
        opts.pack(fill=tk.BOTH, expand=False, pady=(12, 6))
        opts.columnconfigure(0, weight=1)
        opts.columnconfigure(2, weight=1)

        step_var = tk.StringVar(value="0.5")

        # Chọn tiêu chí FOS dùng 2 cột. Mỗi tiêu chí được xét có điều kiện riêng.
        metric_options = self._fos_metric_options()
        metric_all_keys = [key for _lb, key in metric_options]
        selected_metric_keys: List[str] = []
        metric_target_vars: Dict[str, Dict[str, tk.StringVar]] = {}

        metric_select_frame = ttk.Frame(opts)
        metric_select_frame.grid(row=0, column=0, sticky=tk.NSEW, padx=6, pady=6)
        metric_select_frame.columnconfigure(0, weight=1)
        metric_select_frame.columnconfigure(2, weight=1)
        ttk.Label(metric_select_frame, text="Toàn bộ tiêu chí FOS").grid(row=0, column=0, sticky=tk.W, padx=(0, 4))
        ttk.Label(metric_select_frame, text="Tiêu chí được xét").grid(row=0, column=2, sticky=tk.W, padx=(4, 0))

        metric_left_frame = ttk.Frame(metric_select_frame)
        metric_left_frame.grid(row=1, column=0, sticky=tk.NSEW, padx=(0, 6))
        metric_mid_frame = ttk.Frame(metric_select_frame)
        metric_mid_frame.grid(row=1, column=1, sticky=tk.NS, padx=4)
        metric_right_frame = ttk.Frame(metric_select_frame)
        metric_right_frame.grid(row=1, column=2, sticky=tk.NSEW, padx=(6, 0))

        metric_lb_all = tk.Listbox(metric_left_frame, selectmode=tk.EXTENDED, height=8, exportselection=False, width=32)
        metric_all_scroll = ttk.Scrollbar(metric_left_frame, orient=tk.VERTICAL, command=metric_lb_all.yview)
        metric_lb_all.configure(yscrollcommand=metric_all_scroll.set)
        metric_lb_all.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        metric_all_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        metric_lb_sel = tk.Listbox(metric_right_frame, selectmode=tk.EXTENDED, height=8, exportselection=False, width=32)
        metric_sel_scroll = ttk.Scrollbar(metric_right_frame, orient=tk.VERTICAL, command=metric_lb_sel.yview)
        metric_lb_sel.configure(yscrollcommand=metric_sel_scroll.set)
        metric_lb_sel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        metric_sel_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        for lb, key in metric_options:
            metric_lb_all.insert(tk.END, lb)

        target_rows_frame = ttk.LabelFrame(opts, text="Điều kiện riêng cho từng tiêu chí FOS")
        target_rows_frame.grid(row=0, column=1, sticky=tk.NSEW, padx=(8, 6), pady=6)
        target_rows_frame.columnconfigure(0, weight=1)
        target_rows_frame.columnconfigure(1, weight=0)
        target_rows_frame.columnconfigure(2, weight=0)
        target_rows_frame.columnconfigure(3, weight=0)
        target_rows_frame.columnconfigure(4, weight=0)

        def _ensure_metric_target_vars(key: str) -> Dict[str, tk.StringVar]:
            if key not in metric_target_vars:
                metric_target_vars[key] = {
                    "mode": tk.StringVar(value="Từ a đến b"),
                    "a": tk.StringVar(value="1.10"),
                    "b": tk.StringVar(value="1.20"),
                }
            return metric_target_vars[key]

        def rebuild_metric_target_rows():
            for child in target_rows_frame.winfo_children():
                child.destroy()
            ttk.Label(target_rows_frame, text="Tiêu chí", font=("Arial", 9, "bold")).grid(row=0, column=0, sticky=tk.W, padx=4, pady=3)
            ttk.Label(target_rows_frame, text="Dạng", font=("Arial", 9, "bold")).grid(row=0, column=1, sticky=tk.W, padx=4, pady=3)
            ttk.Label(target_rows_frame, text="a", font=("Arial", 9, "bold")).grid(row=0, column=2, sticky=tk.W, padx=4, pady=3)
            ttk.Label(target_rows_frame, text="b", font=("Arial", 9, "bold")).grid(row=0, column=3, sticky=tk.W, padx=4, pady=3)
            if not selected_metric_keys:
                ttk.Label(target_rows_frame, text="Chọn tiêu chí FOS ở bên trái rồi bấm >", foreground=self.pal.get("muted", "#666666")).grid(row=1, column=0, columnspan=4, sticky=tk.W, padx=4, pady=6)
                return
            for row_no, key in enumerate(selected_metric_keys, start=1):
                vars_for_metric = _ensure_metric_target_vars(key)
                ttk.Label(target_rows_frame, text=self._fos_label_from_key(key), width=27).grid(row=row_no, column=0, sticky=tk.W, padx=4, pady=2)
                ttk.Combobox(target_rows_frame, textvariable=vars_for_metric["mode"], values=["Gần giá trị", "Từ a đến b", "Lớn hơn b"], state="readonly", width=13).grid(row=row_no, column=1, sticky=tk.W, padx=4, pady=2)
                ttk.Entry(target_rows_frame, textvariable=vars_for_metric["a"], width=8).grid(row=row_no, column=2, sticky=tk.W, padx=4, pady=2)
                ttk.Entry(target_rows_frame, textvariable=vars_for_metric["b"], width=8).grid(row=row_no, column=3, sticky=tk.W, padx=4, pady=2)

        def refresh_metric_selected_list():
            metric_lb_sel.delete(0, tk.END)
            for key in selected_metric_keys:
                metric_lb_sel.insert(tk.END, self._fos_label_from_key(key))
            rebuild_metric_target_rows()

        def add_metric_selected():
            for pos in metric_lb_all.curselection():
                if 0 <= int(pos) < len(metric_all_keys):
                    key = metric_all_keys[int(pos)]
                    if key not in selected_metric_keys:
                        selected_metric_keys.append(key)
                        _ensure_metric_target_vars(key)
            refresh_metric_selected_list()

        def add_metric_all():
            selected_metric_keys[:] = list(metric_all_keys)
            for key in selected_metric_keys:
                _ensure_metric_target_vars(key)
            refresh_metric_selected_list()

        def remove_metric_selected():
            remove_pos = sorted([int(i) for i in metric_lb_sel.curselection()], reverse=True)
            for pos in remove_pos:
                if 0 <= pos < len(selected_metric_keys):
                    selected_metric_keys.pop(pos)
            refresh_metric_selected_list()

        def remove_metric_all():
            selected_metric_keys.clear()
            refresh_metric_selected_list()

        ttk.Button(metric_mid_frame, text=">", width=5, command=add_metric_selected).pack(pady=(18, 4))
        ttk.Button(metric_mid_frame, text="<", width=5, command=remove_metric_selected).pack(pady=4)
        ttk.Button(metric_mid_frame, text=">>", width=5, command=add_metric_all).pack(pady=(14, 4))
        ttk.Button(metric_mid_frame, text="<<", width=5, command=remove_metric_all).pack(pady=4)
        metric_lb_all.bind("<Double-1>", lambda _e: add_metric_selected())
        metric_lb_sel.bind("<Double-1>", lambda _e: remove_metric_selected())

        # Mặc định xét FOS quyết định chiều dài cọc với khoảng 1.10-1.20.
        selected_metric_keys[:] = ["governing"]
        _ensure_metric_target_vars("governing")
        refresh_metric_selected_list()

        bottom_target_bar = ttk.Frame(opts)
        bottom_target_bar.grid(row=1, column=0, columnspan=2, sticky=tk.EW, padx=6, pady=(0, 6))
        ttk.Label(bottom_target_bar, text="Bước L (m):").pack(side=tk.LEFT, padx=(0, 4))
        ttk.Entry(bottom_target_bar, textvariable=step_var, width=10).pack(side=tk.LEFT, padx=(0, 12))
        ttk.Label(bottom_target_bar, text="Ví dụ: FOS cọc đơn CĐ đặt 1.10-1.20, FOS nhóm CĐ đặt 1.05-1.15.", foreground=self.pal.get("muted", "#666666")).pack(side=tk.LEFT, padx=4)

        def build_metric_target_specs_from_ui() -> Tuple[List[Dict[str, Any]], List[str]]:
            specs: List[Dict[str, Any]] = []
            warnings: List[str] = []
            for key in selected_metric_keys:
                vars_for_metric = _ensure_metric_target_vars(key)
                mode = str(vars_for_metric["mode"].get() or "Từ a đến b")
                a_txt = str(vars_for_metric["a"].get() or "").strip()
                b_txt = str(vars_for_metric["b"].get() or "").strip()
                a = _safe_float(a_txt, float("nan")) if a_txt else float("nan")
                b = _safe_float(b_txt, float("nan")) if b_txt else float("nan")
                label = self._fos_label_from_key(key)
                if mode == "Gần giá trị":
                    if math.isnan(a) or a <= 0.0:
                        warnings.append(f"{label}: dạng 'Gần giá trị' cần nhập a/giá trị FOS > 0.")
                        continue
                    b = a
                elif mode == "Từ a đến b":
                    if math.isnan(a) or math.isnan(b) or a <= 0.0 or b <= 0.0:
                        warnings.append(f"{label}: dạng 'Từ a đến b' cần nhập cả a và b > 0.")
                        continue
                    if a > b:
                        a, b = b, a
                elif mode == "Lớn hơn b":
                    if math.isnan(b) or b <= 0.0:
                        b = a
                    if math.isnan(b) or b <= 0.0:
                        warnings.append(f"{label}: dạng 'Lớn hơn b' cần nhập b > 0.")
                        continue
                    a = b
                specs.append({"key": key, "mode": mode, "a": float(a), "b": float(b), "label": label})
            return specs, warnings

        # V1.0.62: bỏ hẳn phần giao diện ràng buộc lớp mũi/chiều dài ngập tối thiểu.
        # Tối ưu chiều dài cọc chỉ xét các điều kiện FOS đã chọn.
        def rebuild_constraint_rows():
            return

        def build_constraints_from_ui() -> Tuple[Dict[str, Dict[str, Any]], List[str]]:
            return {}, []

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(10, 0))
        def run_opt():
            if not selected_indices:
                messagebox.showwarning("Tối ưu chiều dài", "Cần chọn ít nhất một hạng mục ở cột bên phải.", parent=dlg)
                return
            step = _safe_float(step_var.get(), 0.5)
            if step <= 0.0:
                messagebox.showwarning("Tối ưu chiều dài", "Bước chiều dài phải > 0.", parent=dlg)
                return
            if not selected_metric_keys:
                messagebox.showwarning("Tối ưu chiều dài", "Cần chọn ít nhất một tiêu chí FOS ở cột 'Tiêu chí được xét'.", parent=dlg)
                return
            target_specs, target_warnings = build_metric_target_specs_from_ui()
            if target_warnings:
                messagebox.showwarning("Mục tiêu FOS", "\n".join(target_warnings[:12]), parent=dlg)
                return
            if not target_specs:
                messagebox.showwarning("Tối ưu chiều dài", "Cần khai báo ít nhất một điều kiện FOS hợp lệ.", parent=dlg)
                return
            # V1.0.61: bỏ ràng buộc lớp mũi/chiều dài ngập tối thiểu khỏi quá trình tối ưu.
            constraints = {}
            chosen_indices = list(selected_indices)
            dlg.destroy()
            self.optimize_pile_lengths_by_fos(inputs, chosen_indices, target_specs, "", 0.0, 0.0, step, constraints)
        ttk.Button(btns, text="Tối ưu", style="Accent.TButton", command=run_opt).pack(side=tk.RIGHT, padx=4)
        ttk.Button(btns, text="Hủy", command=dlg.destroy).pack(side=tk.RIGHT, padx=4)
        ttk.Button(btns, text="Chọn tất cả", command=add_all).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Bỏ chọn tất cả", command=remove_all).pack(side=tk.LEFT, padx=4)

        # Khởi tạo rỗng theo đúng nguyên tắc chọn bằng mũi tên.
        refresh_selected_list()
        self.root.wait_window(dlg)

    def _candidate_lengths_for_optimization(self, inp: PileInput, step: float) -> List[float]:
        """Tạo miền tìm kiếm chiều dài cọc cho tối ưu.

        Lỗi cũ của V59 là miền tìm kiếm đôi khi chỉ dừng ở chiều dài cọc hiện tại,
        dù bảng lỗ khoan còn sâu hơn. Nguyên nhân thường gặp là dữ liệu lớp trong
        kết quả tính chỉ phản ánh đoạn cọc đang làm việc, hoặc một vài giá trị 0.00
        bị xử lý bằng toán tử ``or`` làm mất cao độ thật. Hàm này đọc lại đáy địa
        chất trực tiếp từ bảng lỗ khoan đã khớp với hạng mục và KHÔNG dùng ``or``
        cho cao độ, vì 0.00 là một cao độ hợp lệ.
        """
        step = max(float(step or 0.5), 0.001)
        cap_elev = float(getattr(inp, "cap_bottom_elev_m", 0.0) or 0.0)
        old_len = max(cap_elev - float(getattr(inp, "pile_tip_elev_m", cap_elev) or cap_elev), 0.0)
        min_len = step

        bottoms: List[float] = []

        # 1) Lấy từ inp.layers hiện có. Không dùng ``or`` vì bottom_elev_m = 0.00 là hợp lệ.
        for ly in getattr(inp, "layers", []) or []:
            try:
                v = float(getattr(ly, "bottom_elev_m"))
                if not math.isnan(v):
                    bottoms.append(v)
            except Exception:
                pass

        # 2) Đọc lại trực tiếp từ bảng lỗ khoan địa chất theo khớp nối hiện tại.
        #    Nhánh này giúp tối ưu vẫn thấy được toàn bộ lỗ khoan, kể cả các lớp nằm dưới mũi cọc hiện tại.
        try:
            common = bool(self.vars.get("common_geology") and self.vars["common_geology"].get())
            wanted_key = _normalize_item_name(getattr(inp, "item", ""))
            target_geo_key = self.item_geo_map.get(wanted_key, wanted_key)
            for effective_item, row in self._geology_rows_with_effective_items():
                effective_key = _normalize_item_name(effective_item)
                if (not common) and wanted_key and effective_key not in (target_geo_key, "", "all", "chung"):
                    continue
                rr = list(row) + [""] * 4
                raw_bottom = str(rr[2] if len(rr) > 2 else "").strip().replace(",", ".")
                if not raw_bottom:
                    continue
                try:
                    v = float(raw_bottom)
                except Exception:
                    continue
                if not math.isnan(v):
                    bottoms.append(v)
        except Exception:
            pass

        if bottoms:
            deepest = min(bottoms)
            max_geo_len = max(cap_elev - deepest, old_len)
        else:
            max_geo_len = old_len

        if bool(getattr(inp, "allow_geology_extrapolation", False)):
            max_len = max(max_geo_len, old_len + 30.0)
        else:
            max_len = max_geo_len
        max_len = max(max_len, min_len)

        n0 = int(math.ceil(min_len / step - 1e-9))
        n1 = int(math.floor(max_len / step + 1e-9))
        vals = [round(i * step, 6) for i in range(n0, n1 + 1) if i * step > 0]

        # Luôn đưa chiều dài hiện tại vào miền tìm kiếm để có phương án so sánh.
        old_grid = round(round(old_len / step) * step, 6)
        if old_grid > 0 and old_grid not in vals:
            vals.append(old_grid)

        # Nếu đáy khảo sát không đúng bội số step, thêm chiều dài đến đáy địa chất để tránh thiếu đoạn cuối.
        max_grid = round(max_len, 6)
        if max_grid > 0 and max_grid not in vals:
            vals.append(max_grid)

        vals = sorted(v for v in set(vals) if v > 0)
        return vals

    def _current_toe_layer_text(self, inp: PileInput) -> Tuple[str, str]:
        ly, embed, _top, _bot = self._toe_layer_and_embedment_for_tip(inp, getattr(inp, "pile_tip_elev_m", 0.0))
        if ly is None:
            return "-", "-"
        return str(getattr(ly, "name", "") or "?"), _fmt(embed, 2) + " m"

    def _toe_layer_and_embedment_for_tip(self, inp: PileInput, tip_elev: float) -> Tuple[Optional[SoilLayer], float, float, float]:
        layers = SCTCalculator._sorted_layers(inp)
        if not layers:
            return None, 0.0, 0.0, 0.0
        top = float(getattr(inp, "ground_elev_m", 0.0) or 0.0)
        tip = float(tip_elev)
        for ly in layers:
            bottom = float(getattr(ly, "bottom_elev_m", top) or top)
            if top + 1e-9 >= tip >= bottom - 1e-9:
                return ly, max(top - tip, 0.0), top, bottom
            top = bottom
        if bool(getattr(inp, "allow_geology_extrapolation", False)):
            last = layers[-1]
            layer_top = layers[-2].bottom_elev_m if len(layers) >= 2 else float(getattr(inp, "ground_elev_m", 0.0) or 0.0)
            if tip <= float(getattr(last, "bottom_elev_m", tip) or tip) + 1e-9:
                return last, max(float(layer_top) - tip, 0.0), float(layer_top), tip
        return None, 0.0, 0.0, 0.0

    def _tip_layer_option_records(self, inp: PileInput) -> List[Dict[str, Any]]:
        """Danh sách lớp địa chất theo lỗ khoan đã khớp với hạng mục.

        Mỗi record có tên lớp và khoảng cao độ để người dùng chọn đúng lớp khi tên lớp lặp lại
        do nhập theo từng điểm SPT hoặc do có hang/đá xen kẹp.
        """
        records: List[Dict[str, Any]] = []
        layers = SCTCalculator._sorted_layers(inp)
        if not layers:
            return records
        top = float(getattr(inp, "ground_elev_m", 0.0) or 0.0)
        used_display: Dict[str, int] = {}
        for ly in layers:
            bottom = float(getattr(ly, "bottom_elev_m", top) or top)
            name = str(getattr(ly, "name", "") or "").strip() or "?"
            thick = max(top - bottom, 0.0)
            display = f"{name}  ({_fmt(top,2)} → {_fmt(bottom,2)}, d={_fmt(thick,2)}m)"
            if display in used_display:
                used_display[display] += 1
                display = f"{display} #{used_display[display]}"
            else:
                used_display[display] = 1
            records.append({"name": name, "top": float(top), "bottom": float(bottom), "display": display})
            top = bottom
        return records

    def _tip_layer_record_from_display(self, inp: PileInput, display: Any) -> Optional[Dict[str, Any]]:
        text = str(display or "").strip()
        if not text:
            return None
        for rec in self._tip_layer_option_records(inp):
            if rec.get("display") == text:
                return rec
        # Tương thích nếu người dùng nhập/đọc lại dữ liệu cũ chỉ có tên lớp.
        target_key = _normalize_item_name(text.split("(")[0].strip())
        for rec in self._tip_layer_option_records(inp):
            if _normalize_item_name(rec.get("name", "")) == target_key:
                return rec
        return None

    def _find_first_layer_interval_by_name(self, inp: PileInput, layer_name: Any) -> Tuple[Optional[SoilLayer], float, float]:
        target_key = _normalize_item_name(layer_name)
        layers = SCTCalculator._sorted_layers(inp)
        top = float(getattr(inp, "ground_elev_m", 0.0) or 0.0)
        for ly in layers:
            bottom = float(getattr(ly, "bottom_elev_m", top) or top)
            if _normalize_item_name(getattr(ly, "name", "")) == target_key:
                return ly, float(top), float(bottom)
            top = bottom
        return None, 0.0, 0.0

    def _parse_tip_constraints_text(self, raw_text: str) -> Tuple[Dict[str, Dict[str, Any]], List[str]]:
        constraints: Dict[str, Dict[str, Any]] = {}
        warnings: List[str] = []
        for line_no, line in enumerate(str(raw_text or "").splitlines(), start=1):
            raw = line.strip()
            if not raw:
                continue
            if "\t" in raw:
                parts = [p.strip() for p in raw.split("\t")]
            else:
                parts = [p.strip() for p in re.split(r"[;,]", raw)]
            if not parts or not parts[0]:
                continue
            head = _strip_accents(parts[0]).lower()
            if "hang muc" in head or "hạng mục" in parts[0].lower():
                continue
            item = parts[0]
            layer = parts[1] if len(parts) >= 2 else ""
            embed_txt = parts[2] if len(parts) >= 3 else ""
            if not str(layer).strip():
                continue
            embed = _safe_float(embed_txt, 0.0) if str(embed_txt).strip() else 0.0
            if embed < 0.0:
                warnings.append(f"Dòng {line_no}: chiều dài ngập tối thiểu không được âm.")
                continue
            constraints[_normalize_item_name(item)] = {"layer": str(layer).strip(), "embed_min": float(embed), "item": item}
        return constraints, warnings

    def _tip_constraint_ok(self, inp: PileInput, constraint: Optional[Dict[str, Any]]) -> Tuple[bool, str]:
        if not constraint or not str(constraint.get("layer", "")).strip():
            return True, ""
        target_layer = str(constraint.get("layer", "")).strip()
        min_embed = max(float(constraint.get("embed_min", 0.0) or 0.0), 0.0)
        tip = float(getattr(inp, "pile_tip_elev_m", 0.0) or 0.0)
        toe, embed, toe_top, toe_bottom = self._toe_layer_and_embedment_for_tip(inp, tip)
        if toe is None:
            return False, f"không xác định được lớp tại mũi; yêu cầu lớp {target_layer}"

        # Nếu người dùng chọn một record cụ thể trong combobox thì dùng đúng khoảng cao độ đó.
        layer_top = constraint.get("layer_top", None)
        layer_bottom = constraint.get("layer_bottom", None)
        if layer_top is None or layer_bottom is None:
            target_ly, layer_top, layer_bottom = self._find_first_layer_interval_by_name(inp, target_layer)
            if target_ly is None:
                return False, f"không tìm thấy lớp {target_layer} trong lỗ khoan đã khớp"
        layer_top = float(layer_top)
        layer_bottom = float(layer_bottom)
        toe_name = str(getattr(toe, "name", "") or "?")

        if min_embed > 0.0:
            # Chiều dài ngập tính từ mặt trên lớp được chọn. Nếu lớp được chọn mỏng hơn yêu cầu,
            # mũi cọc được phép đi tiếp xuống lớp dưới, miễn tổng chiều dài từ mặt trên lớp chọn đạt yêu cầu.
            actual_embed = layer_top - tip
            if actual_embed + 1e-9 < min_embed:
                return False, f"ngập tính từ lớp {target_layer} {_fmt(actual_embed,2)} m < yêu cầu {_fmt(min_embed,2)} m"
            return True, f"ngập tính từ lớp {target_layer} {_fmt(actual_embed,2)} m"

        # Nếu không nhập chiều dài ngập, hiểu là chỉ yêu cầu mũi nằm trong đúng lớp đã chọn.
        if not (layer_top + 1e-9 >= tip >= layer_bottom - 1e-9):
            return False, f"mũi nằm trong lớp {toe_name}; yêu cầu lớp {target_layer}"
        return True, f"mũi trong lớp {target_layer}"

    def _fos_target_description(self, mode: str, a: float, b: float) -> str:
        mode = str(mode or "Gần giá trị")
        if mode == "Từ a đến b":
            lo, hi = sorted([float(a), float(b)])
            return f"{_fmt(lo,3)} ≤ FOS ≤ {_fmt(hi,3)}"
        if mode == "Lớn hơn b":
            return f"FOS ≥ {_fmt(float(b),3)}"
        return f"FOS gần {_fmt(float(a),3)}"

    def _fos_candidate_score(self, fos: float, length_m: float, mode: str, a: float, b: float) -> Tuple[bool, Tuple[float, ...], Tuple[float, ...]]:
        mode = str(mode or "Gần giá trị")
        L = float(length_m)
        f = float(fos)
        if mode == "Từ a đến b":
            lo, hi = sorted([float(a), float(b)])
            center = (lo + hi) / 2.0
            if lo <= f <= hi:
                return True, (L, abs(f - center)), (0.0, L, abs(f - center))
            dist = lo - f if f < lo else f - hi
            return False, (float("inf"),), (abs(dist), L)
        if mode == "Lớn hơn b":
            threshold = float(b)
            if f >= threshold:
                return True, (L, abs(f - threshold)), (0.0, L, abs(f - threshold))
            return False, (float("inf"),), (threshold - f, L)
        target = float(a)
        ok = f + 1e-9 >= target
        score = (abs(f - target), L)
        return ok, score, score

    def _valid_fos_metric_keys_for_dialog(self) -> set:
        return set(["governing"] + self._base_fos_metric_keys())

    def _normalize_fos_target_specs(self, target_specs: Any, target_mode: str = "Gần giá trị", target_a: float = 1.10, target_b: float = 1.10) -> List[Dict[str, Any]]:
        """Chuẩn hóa mục tiêu tối ưu FOS.

        Dạng mới: list[dict] với từng tiêu chí có mode/a/b riêng.
        Vẫn giữ tương thích dạng cũ: metric_keys + một điều kiện chung.
        """
        valid = self._valid_fos_metric_keys_for_dialog()
        specs: List[Dict[str, Any]] = []
        if isinstance(target_specs, (list, tuple)) and target_specs and all(isinstance(x, dict) for x in target_specs):
            raw_specs = list(target_specs)
        else:
            # Tương thích lời gọi cũ. Không dùng _normalize_fos_metric_keys ở đây để giữ được key "governing".
            if isinstance(target_specs, str):
                keys = [target_specs]
            else:
                try:
                    keys = list(target_specs)
                except Exception:
                    keys = ["single_cd"]
            raw_specs = [{"key": k, "mode": target_mode, "a": target_a, "b": target_b} for k in keys]

        for sp in raw_specs:
            key = str(sp.get("key", "") or "").strip()
            if key not in valid:
                continue
            mode = str(sp.get("mode", target_mode) or target_mode or "Gần giá trị")
            a = _safe_float(sp.get("a", target_a), float("nan"))
            b = _safe_float(sp.get("b", target_b), float("nan"))
            if mode == "Từ a đến b":
                if math.isnan(a) or math.isnan(b):
                    continue
                if a > b:
                    a, b = b, a
            elif mode == "Lớn hơn b":
                if math.isnan(b) or b <= 0.0:
                    b = a
                a = b
            else:
                mode = "Gần giá trị"
                if math.isnan(a):
                    continue
                b = a
            specs.append({"key": key, "mode": mode, "a": float(a), "b": float(b), "label": self._fos_label_from_key(key)})

        # Bỏ trùng key, giữ dòng cuối cùng để người dùng có thể sửa lại điều kiện.
        out_by_key: Dict[str, Dict[str, Any]] = {}
        order: List[str] = []
        for sp in specs:
            key = sp["key"]
            if key not in out_by_key:
                order.append(key)
            out_by_key[key] = sp
        return [out_by_key[k] for k in order] or [{"key": "single_cd", "mode": "Gần giá trị", "a": 1.10, "b": 1.10, "label": self._fos_label_from_key("single_cd")}]

    def _fos_target_description_for_specs(self, target_specs: Any) -> str:
        specs = self._normalize_fos_target_specs(target_specs)
        parts: List[str] = []
        for sp in specs:
            parts.append(f"{sp.get('label') or self._fos_label_from_key(sp.get('key',''))}: {self._fos_target_description(sp.get('mode'), sp.get('a'), sp.get('b'))}")
        return "; ".join(parts)

    def _selected_fos_min_for_result(self, res: CapacityResult, target_specs: Any) -> float:
        vals: List[float] = []
        for sp in self._normalize_fos_target_specs(target_specs):
            try:
                v = float(self._fos_value_for_result(res, sp.get("key", "single_cd")))
            except Exception:
                continue
            if math.isnan(v) or v <= 0.0:
                continue
            vals.append(v)
        if not vals:
            return float("inf")
        finite = [v for v in vals if not math.isinf(v)]
        return min(finite) if finite else float("inf")

    def _multi_fos_candidate_score(self, res: CapacityResult, length_m: float, target_specs: Any) -> Tuple[bool, Tuple[float, ...], Tuple[float, ...], float]:
        specs = self._normalize_fos_target_specs(target_specs)
        L = float(length_m)
        failures = 0
        total_deficit = 0.0
        total_closeness = 0.0
        has_near_target = False
        vals: List[float] = []
        for sp in specs:
            key = sp.get("key", "single_cd")
            mode = str(sp.get("mode", "Gần giá trị") or "Gần giá trị")
            a = float(sp.get("a", 0.0) or 0.0)
            b = float(sp.get("b", a) or a)
            try:
                f = float(self._fos_value_for_result(res, key))
            except Exception:
                f = float("nan")
            if math.isnan(f) or f <= 0.0 or math.isinf(f):
                failures += 1
                total_deficit += 1e6
                total_closeness += 1e6
                continue
            vals.append(f)
            if mode == "Từ a đến b":
                lo, hi = sorted([a, b])
                center = (lo + hi) / 2.0
                if f < lo:
                    failures += 1
                    total_deficit += lo - f
                elif f > hi:
                    failures += 1
                    total_deficit += f - hi
                total_closeness += abs(f - center)
            elif mode == "Lớn hơn b":
                th = b if b > 0 else a
                if f < th:
                    failures += 1
                    total_deficit += th - f
                total_closeness += abs(f - th)
            else:
                has_near_target = True
                target = a
                if f + 1e-9 < target:
                    failures += 1
                    total_deficit += target - f
                total_closeness += abs(f - target)
        governing = min(vals) if vals else float("inf")
        ok = failures == 0
        if ok:
            # Với mục tiêu gần giá trị, ưu tiên gần mục tiêu; với khoảng/lớn hơn, ưu tiên chiều dài ngắn hơn.
            ok_score = (total_closeness, L) if has_near_target else (L, total_closeness)
        else:
            ok_score = (float("inf"),)
        any_score = (failures, total_deficit, total_closeness, L)
        return ok, ok_score, any_score, governing

    def optimize_pile_lengths_by_fos(self, inputs: List[PileInput], selected_indices: List[int], metric_keys: Any, target_mode: str, target_a: float, target_b: float, step_m: float, tip_constraints: Optional[Dict[str, Dict[str, Any]]] = None):
        selected_set = set(int(i) for i in selected_indices)
        updates: Dict[str, Dict[str, Any]] = {}
        messages: List[str] = []
        final_inputs: List[PileInput] = []
        tip_constraints = tip_constraints or {}
        target_specs = self._normalize_fos_target_specs(metric_keys, target_mode, target_a, target_b)

        for idx, inp in enumerate(inputs):
            if idx not in selected_set:
                final_inputs.append(copy.deepcopy(inp))
                continue
            item_key = _normalize_item_name(inp.item)
            constraint = tip_constraints.get(item_key)
            try:
                res0 = SCTCalculator.calculate(inp)
                old_fos = self._selected_fos_min_for_result(res0, target_specs)
            except Exception as exc:
                messages.append(f"{inp.item}: không tính được FOS ban đầu ({exc}).")
                final_inputs.append(copy.deepcopy(inp))
                continue
            if math.isinf(old_fos):
                messages.append(f"{inp.item}: FOS = ∞ do chưa có tải kiểm toán hoặc sức kháng/tải không hợp lệ; bỏ qua tối ưu.")
                final_inputs.append(copy.deepcopy(inp))
                continue

            candidates = self._candidate_lengths_for_optimization(inp, step_m)
            if not candidates:
                messages.append(f"{inp.item}: không tạo được miền tìm kiếm chiều dài.")
                final_inputs.append(copy.deepcopy(inp))
                continue

            best_any = None
            best_ok = None
            skipped_by_tip = 0
            last_tip_note = ""
            for L in candidates:
                cand = copy.deepcopy(inp)
                cand.pile_tip_elev_m = float(cand.cap_bottom_elev_m) - float(L)
                tip_ok, tip_note = self._tip_constraint_ok(cand, constraint)
                if not tip_ok:
                    skipped_by_tip += 1
                    last_tip_note = tip_note
                    continue
                try:
                    cres = SCTCalculator.calculate(cand)
                    ok, ok_score, any_score, fos = self._multi_fos_candidate_score(cres, L, target_specs)
                except Exception:
                    continue
                if math.isinf(fos):
                    continue
                if best_any is None or any_score < best_any[0]:
                    best_any = (any_score, L, fos, cand.pile_tip_elev_m, tip_note)
                if ok:
                    if best_ok is None or ok_score < best_ok[0]:
                        best_ok = (ok_score, L, fos, cand.pile_tip_elev_m, tip_note)
            chosen = best_ok or best_any
            if chosen is None:
                if constraint:
                    messages.append(f"{inp.item}: không có chiều dài hợp lệ thỏa ràng buộc lớp mũi ({last_tip_note or 'kiểm tra tên lớp/chiều sâu ngập'}).")
                else:
                    messages.append(f"{inp.item}: không tìm được phương án tính hợp lệ.")
                final_inputs.append(copy.deepcopy(inp))
                continue
            _, new_len, new_fos, new_tip, tip_note = chosen
            inp2 = copy.deepcopy(inp)
            inp2.pile_tip_elev_m = float(new_tip)
            final_inputs.append(inp2)
            updates[item_key] = {
                "item": inp.item,
                "old_len": inp.pile_length_m,
                "new_len": new_len,
                "old_tip": inp.pile_tip_elev_m,
                "new_tip": new_tip,
                "old_fos": old_fos,
                "new_fos": new_fos,
                "ok": bool(best_ok is not None),
                "tip_note": tip_note,
                "has_tip_constraint": bool(constraint),
            }
            if constraint and skipped_by_tip and not tip_note:
                messages.append(f"{inp.item}: đã bỏ qua {skipped_by_tip} chiều dài không thỏa ràng buộc lớp mũi.")

        if updates:
            self._apply_optimized_tips_to_item_table(updates)
            try:
                results = [SCTCalculator.calculate(inp) for inp in final_inputs]
                self.last_results = results
                self.last_result = results[0] if results else None
                self._show_result_collection(results, f"Đã tối ưu chiều dài cho {len(updates)} hạng mục")
                self.nb.select(self.tab_result)
            except Exception as exc:
                messagebox.showerror("Tối ưu chiều dài", f"Đã cập nhật CĐ mũi nhưng tính lại kết quả bị lỗi:\n{exc}")

        desc = self._fos_target_description_for_specs(target_specs)
        lines = [f"Điều kiện FOS: {desc}", f"Bước L = {_fmt(step_m,3)} m", ""]
        for u in updates.values():
            flag = "" if u.get("ok") else " (ngoài điều kiện, lấy gần nhất)"
            tip_note = f"; {u.get('tip_note')}" if u.get("has_tip_constraint") and u.get("tip_note") else ""
            lines.append(f"{u['item']}: L {_fmt(u['old_len'],2)} → {_fmt(u['new_len'],2)} m; CĐ mũi {_fmt(u['old_tip'],2)} → {_fmt(u['new_tip'],2)}; FOS {self._fmt_fos_value(u['old_fos'])} → {self._fmt_fos_value(u['new_fos'])}{tip_note}{flag}")
        if messages:
            lines.append("")
            lines.append("Cảnh báo/bỏ qua:")
            lines.extend("- " + m for m in messages)
        if not updates and not messages:
            lines.append("Không có hạng mục nào được tối ưu.")
        messagebox.showinfo("Tối ưu chiều dài cọc", "\n".join(lines[:45]))

    def _apply_optimized_tips_to_item_table(self, updates: Dict[str, Dict[str, Any]]):
        rows = self.item_table.get_rows()
        for rr in rows:
            while len(rr) < len(self.item_table.columns):
                rr.append("")
            key = _normalize_item_name(rr[0])
            if key in updates:
                rr[12] = _fmt(updates[key]["new_tip"], 2)
        self.item_table.set_rows(rows, record_undo=True)

    def _validate_required_pile_geometry_before_run(self) -> bool:
        """Không cho chạy nếu thiếu các thông số hình học tối thiểu của cọc.

        Dòng trống được bỏ qua. Cao độ đáy bệ và cao độ mũi cọc được kiểm tra theo
        nội dung ô nhập, vì giá trị 0.00 là hợp lệ và không được xem là thiếu.
        """
        try:
            rows = self.item_table.get_rows() if getattr(self, "item_table", None) else []
        except Exception:
            rows = []
        bad_items: List[str] = []
        for idx, row in enumerate(rows, start=1):
            rr = list(row or [])
            if not any(str(x).strip() for x in rr):
                continue
            while len(rr) < 13:
                rr.append("")
            item = str(rr[0] or "").strip() or f"Dòng {idx}"
            d_txt = str(rr[4] or "").strip()
            cap_txt = str(rr[11] or "").strip()
            tip_txt = str(rr[12] or "").strip()
            missing = False
            if not d_txt or _safe_float(d_txt, 0.0) <= 0.0:
                missing = True
            if not cap_txt:
                missing = True
            if not tip_txt:
                missing = True
            if missing:
                bad_items.append(item)
        if bad_items:
            msg = "Thêm đường kính cọc, cao độ cọc để chạy chương trình"
            if len(bad_items) <= 8:
                msg += "\n" + "Hạng mục thiếu: " + ", ".join(bad_items)
            else:
                msg += "\n" + "Có " + str(len(bad_items)) + " hạng mục đang thiếu dữ liệu."
            messagebox.showwarning("Thiếu số liệu cọc", msg)
            return False
        return True

    def _fatal_input_check_before_calculate(self, inputs: List[PileInput]) -> bool:
        """Kiểm tra các lỗi bắt buộc dừng tính trước khi gọi lõi tính toán."""
        missing_gamma: List[str] = []
        missing_rock: List[str] = []
        exceed_items: List[Tuple[PileInput, float]] = []
        for inp in inputs:
            item = str(getattr(inp, "item", "") or "?").strip() or "?"
            layers = SCTCalculator._sorted_layers(inp)
            if layers:
                last_bottom = min(float(getattr(ly, "bottom_elev_m", 0.0) or 0.0) for ly in layers)
                if float(getattr(inp, "pile_tip_elev_m", 0.0) or 0.0) < last_bottom - 1e-9:
                    exceed_items.append((inp, last_bottom))
            for idx, ly in enumerate(layers):
                st = int(getattr(ly, "soil_type", 0) or 0)
                lname = str(getattr(ly, "name", "") or f"L{idx+1}").strip() or f"L{idx+1}"
                label = f"{item} - {lname}"
                if st != 0 and bool(getattr(ly, "_missing_gamma", False)):
                    missing_gamma.append(label)
                if st in (3, 4) and bool(getattr(ly, "_missing_rock_data", False)):
                    missing_rock.append(label)

        if missing_gamma:
            msg = "Thiếu trọng lượng riêng các lớp đất."
            msg += "\n" + ", ".join(missing_gamma[:20])
            if len(missing_gamma) > 20:
                msg += f"\n... và {len(missing_gamma)-20} lớp khác."
            messagebox.showerror("Thiếu số liệu địa chất", msg)
            return False

        if missing_rock:
            msg = "Thiếu số liệu lớp đá."
            msg += "\nCác lớp đá loại 3, 4 phải có qu, RQD, GSI, mi và D."
            msg += "\n" + ", ".join(missing_rock[:20])
            if len(missing_rock) > 20:
                msg += f"\n... và {len(missing_rock)-20} lớp khác."
            messagebox.showerror("Thiếu số liệu lớp đá", msg)
            return False

        exceed_not_allowed = [(inp, lb) for inp, lb in exceed_items if not bool(getattr(inp, "allow_geology_extrapolation", False))]
        if exceed_not_allowed:
            lines = ["Chiều dài cọc vượt quá lỗ khoan địa chất!", "Có muốn ngoại suy lớp đất cuối không?"]
            for inp, last_bottom in exceed_not_allowed[:12]:
                lines.append(f"- {getattr(inp, 'item', '?')}: mũi cọc {float(getattr(inp, 'pile_tip_elev_m', 0.0) or 0.0):g}, đáy lỗ khoan {last_bottom:g}")
            if len(exceed_not_allowed) > 12:
                lines.append(f"... và {len(exceed_not_allowed)-12} hạng mục khác.")
            if not messagebox.askyesno("Ngoại suy địa chất", "\n".join(lines)):
                return False
            for inp, _last_bottom in exceed_not_allowed:
                inp.allow_geology_extrapolation = True
            try:
                if "allow_geology_extrapolation" in self.vars:
                    self.vars["allow_geology_extrapolation"].set(True)
            except Exception:
                pass
        return True

    def _collect_result_warnings(self, results: List[CapacityResult]) -> List[Tuple[str, str]]:
        rows: List[Tuple[str, str]] = []
        seen = set()
        for res in results or []:
            item = str(getattr(res.pile_input, "item", "") or "?").strip() or "?"
            for w in getattr(res, "warnings", []) or []:
                text = str(w or "").strip()
                if not text:
                    continue
                key = (item, text)
                if key in seen:
                    continue
                seen.add(key)
                rows.append(key)
        return rows

    def _update_warning_button(self) -> None:
        try:
            n = len(getattr(self, "last_warning_rows", []) or [])
            if getattr(self, "warning_button", None) is not None:
                self.warning_button.configure(text=(f"Warning ({n})" if n else "Warning"), state=(tk.NORMAL if n else tk.DISABLED))
        except Exception:
            pass

    def show_warning_panel(self):
        rows = list(getattr(self, "last_warning_rows", []) or [])
        if not rows:
            messagebox.showinfo("Warning", "Không có cảnh báo.")
            return
        dlg = tk.Toplevel(self.root)
        dlg.title("Warning")
        dlg.geometry("820x420")
        dlg.transient(self.root)
        dlg.grab_set()
        ttk.Label(dlg, text=f"Tổng số cảnh báo: {len(rows)}", font=("Arial", 10, "bold")).pack(anchor=tk.W, padx=10, pady=(10, 4))
        frame = ttk.Frame(dlg, padding=(10, 0, 10, 8))
        frame.pack(fill=tk.BOTH, expand=True)
        tree = ttk.Treeview(frame, columns=("item", "warning"), show="headings", height=14)
        tree.heading("item", text="Hạng mục")
        tree.heading("warning", text="Warning")
        tree.column("item", width=160, stretch=False, anchor=tk.W)
        tree.column("warning", width=620, stretch=True, anchor=tk.W)
        vsb = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=vsb.set)
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        for item, warning in rows:
            tree.insert("", tk.END, values=(item, warning))
        ttk.Button(dlg, text="Đóng", command=dlg.destroy).pack(anchor=tk.E, padx=10, pady=(0, 10))

    def calculate(self):
        try:
            if not self._validate_required_pile_geometry_before_run():
                return
            # V0.2.8: bắt buộc kiểm tra khớp tên hạng mục trước khi tính,
            # để Trụ T1 dùng đúng địa chất Trụ T1, T2 dùng đúng địa chất T2.
            if not self.check_item_geology_matching(quiet=True):
                if not self.check_item_geology_matching(quiet=False):
                    return
            inputs = self._read_multi_inputs()
            if not inputs:
                raise ValueError("Chưa có hạng mục trong bảng thông tin riêng từng hạng mục.")
            self.last_warning_rows = []
            self._update_warning_button()
            if not self._fatal_input_check_before_calculate(inputs):
                return
            if not self._apply_igm_missing_qu_policy(inputs):
                return
            # QA-UX U3+: phần tính thuần chạy thread nền. Không gọi Tkinter trực tiếp từ worker;
            # worker chỉ đẩy message vào Queue, main thread poll bằng root.after để tránh lỗi Tk ngẫu nhiên.
            if getattr(self, "_calc_busy", False):
                messagebox.showinfo("Đang tính toán", "Một phiên tính toán đang chạy; chờ xong rồi bấm lại.")
                return
            self._calc_busy = True
            prog = tk.Toplevel(self.root)
            prog.title("Đang tính toán")
            prog.transient(self.root)
            prog.resizable(False, False)
            try:
                prog.protocol("WM_DELETE_WINDOW", lambda: None)
            except Exception:
                pass
            pf = ttk.Frame(prog, padding=14)
            pf.pack(fill=tk.BOTH, expand=True)
            ptxt = tk.StringVar(value=f"Đang tính 0/{len(inputs)} hạng mục...")
            ttk.Label(pf, textvariable=ptxt, wraplength=420).pack(anchor=tk.W, pady=(0, 8))
            pbar = ttk.Progressbar(pf, mode="determinate", maximum=max(len(inputs), 1), length=420)
            pbar.pack(fill=tk.X)
            try:
                self._center_window(prog)
            except Exception:
                pass
            safe_lift_window(prog)

            calc_queue: "queue.Queue[Tuple[str, Any]]" = queue.Queue()

            def _calc_worker():
                results: List[CapacityResult] = []
                try:
                    for _i, _inp in enumerate(inputs, start=1):
                        calc_queue.put(("progress", _i, str(getattr(_inp, "item", "") or "")))
                        results.append(SCTCalculator.calculate(_inp))
                    calc_queue.put(("finish", results, ""))
                except Exception:
                    calc_queue.put(("finish", results, traceback.format_exc()))

            def _poll_calc_queue():
                if not self._root_alive():
                    return
                try:
                    while True:
                        msg = calc_queue.get_nowait()
                        kind = msg[0]
                        if kind == "progress":
                            _i, _name = int(msg[1]), str(msg[2] or "")
                            try:
                                ptxt.set(f"Đang tính {_i}/{len(inputs)} hạng mục: {_name}")
                                pbar.configure(value=max(_i - 1, 0))
                            except Exception:
                                pass
                        elif kind == "finish":
                            results, err_txt = msg[1], str(msg[2] or "")
                            self._calc_busy = False
                            try:
                                pbar.configure(value=max(len(inputs), 1))
                            except Exception:
                                pass
                            try:
                                prog.destroy()
                            except Exception:
                                pass
                            safe_release_grabs(self.root)
                            if err_txt:
                                sys.stderr.write(err_txt + "\n")
                                messagebox.showerror("Lỗi tính toán", err_txt.strip().splitlines()[-1])
                                return
                            self.last_result = results[0] if results else None
                            self.last_results = results
                            self.last_warning_rows = self._collect_result_warnings(results)
                            self._update_warning_button()
                            if len(results) == 1:
                                self._show_result(results[0])
                                self._set_status("Đã tính toán cho 1 hạng mục")
                            else:
                                self._show_results(results)
                                self._set_status(f"Đã tính toán cho {len(results)} hạng mục")
                            self.nb.select(self.tab_result)
                            if self.last_warning_rows:
                                self.root.after(80, self.show_warning_panel)
                            return
                except queue.Empty:
                    pass
                except Exception as exc:
                    self._calc_busy = False
                    try:
                        prog.destroy()
                    except Exception:
                        pass
                    traceback.print_exc()
                    messagebox.showerror("Lỗi tính toán", str(exc))
                    return
                self.root.after(60, _poll_calc_queue)

            threading.Thread(target=_calc_worker, daemon=True).start()
            self.root.after(60, _poll_calc_queue)
        except Exception as exc:
            self._calc_busy = False
            traceback.print_exc()
            messagebox.showerror("Lỗi tính toán", str(exc))


    def _limit_state_check_text(self, res: CapacityResult, cap: LimitStateCapacity, force_single: float, force_group: float, label: str) -> str:
        n = max(int(res.pile_input.pile_count_in_group), 1)
        single_cap = cap.compression_single_net_kn
        if cap.material_pr_kn > 0:
            single_cap = min(single_cap, cap.material_pr_kn)
        group_cap = cap.compression_group_total_kn
        if cap.material_pr_kn > 0:
            group_cap = min(group_cap, cap.material_pr_kn * n)
        if force_group <= 0 and force_single > 0:
            force_group = force_single * n
        dd = max(float(getattr(cap, "downdrag_kn", 0.0) or 0.0), 0.0)
        design_single = max(float(force_single or 0.0), 0.0) + dd
        design_group = max(float(force_group or 0.0), 0.0) + dd * n
        dcr_single = design_single / single_cap if design_single > 0 and single_cap > 0 else 0.0
        dcr_group = design_group / group_cap if design_group > 0 and group_cap > 0 else 0.0
        st_single = "ĐẠT" if dcr_single <= 1.0 else "KHÔNG ĐẠT"
        st_group = "ĐẠT" if dcr_group <= 1.0 else "KHÔNG ĐẠT"
        fos_single = self._report_fos_text(dcr_single)
        fos_group = self._report_fos_text(dcr_group)
        dd_txt = f", DD={_fmt(dd,2)} kN/cọc" if dd > 0 else ""
        return (f"{label}: cọc đơn [Q]={_fmt(single_cap,2)} kN, Pu+DD={_fmt(design_single,2)} kN{dd_txt}, "
                f"FOS={fos_single} -> {st_single}; nhóm fg={_fmt(cap.group_factor,3)}, "
                f"[Q] nhóm={_fmt(group_cap,2)} kN, N đáy bệ+DD={_fmt(design_group,2)} kN, "
                f"FOS nhóm={fos_group} -> {st_group}")

    def _limit_state_uplift_check_text(self, res: CapacityResult, cap: LimitStateCapacity, uplift_single: float, label: str) -> str:
        n = max(int(res.pile_input.pile_count_in_group), 1)
        t_single = max(float(uplift_single or 0.0), 0.0)
        t_group = t_single * n
        cap_single = max(float(cap.uplift_single_magnitude_kn or 0.0), 0.0)
        cap_group = max(float(cap.uplift_group_total_kn or 0.0), 0.0)
        dcr_single = t_single / cap_single if t_single > 0 and cap_single > 0 else 0.0
        dcr_group = t_group / cap_group if t_group > 0 and cap_group > 0 else 0.0
        st_single = "ĐẠT" if dcr_single <= 1.0 else "KHÔNG ĐẠT"
        st_group = "ĐẠT" if dcr_group <= 1.0 else "KHÔNG ĐẠT"
        fos_single = self._report_fos_text(dcr_single)
        fos_group = self._report_fos_text(dcr_group)
        return (f"{label} nhổ: cọc đơn [T]={_fmt(cap_single,2)} kN, Tu/cọc={_fmt(t_single,2)} kN, "
                f"FOS={fos_single} -> {st_single}; nhóm [T]={_fmt(cap_group,2)} kN, "
                f"Tu nhóm={_fmt(t_group,2)} kN, FOS nhóm={fos_group} -> {st_group}")

    def _show_results(self, results: List[CapacityResult]):
        self.summary_text.config(state=tk.NORMAL)
        self.summary_text.delete("1.0", tk.END)
        lines = [f"Tổng số hạng mục: {len(results)}", ""]
        for res in results:
            lines.append(f"[{res.pile_input.item}] {res.pile_input.mode}; loại cọc: {res.pile_input.pile_type}; số hàng cọc: {res.pile_input.group_layout}; D={_fmt(res.pile_input.diameter_mm,0)} mm; S/D={_fmt(res.pile_input.spacing_m / max(res.pile_input.diameter_m, 1e-9),3)}; L={_fmt(res.pile_input.pile_length_m,2)} m; FOS quyết định={self._fmt_fos_value(self._governing_fos_for_result(res))}")
            lines.append("  " + self._limit_state_check_text(res, res.strength, res.pile_input.force_cd_kn, res.pile_input.cap_force_cd_kn, "TTGHCĐ"))
            lines.append("  " + self._limit_state_uplift_check_text(res, res.strength, res.pile_input.uplift_cd_kn, "TTGHCĐ"))
            lines.append("  " + self._limit_state_check_text(res, res.extreme, res.pile_input.force_db_kn, res.pile_input.cap_force_db_kn, "TTGHĐB"))
            lines.append("  " + self._limit_state_uplift_check_text(res, res.extreme, res.pile_input.uplift_db_kn, "TTGHĐB"))
        self.summary_text.insert(tk.END, "\n".join(lines))
        self.summary_text.config(state=tk.DISABLED)
        rows = []
        for res in results:
            first_layer = True
            for lr in res.layers:
                rows.append([
                    res.pile_input.item if first_layer else "",
                    lr.name, _fmt(lr.top_elev_m,2), _fmt(lr.bottom_elev_m,2), _fmt(lr.thickness_m,2), _fmt(lr.skin_length_m,2),
                    _fmt(lr.downdrag_length_m,2), _fmt(lr.downdrag_kn,2), lr.soil_label, _fmt(lr.group_factor,3), _fmt(lr.n_spt,2), _fmt(lr.n60,2), _fmt(lr.n1_60,2), _fmt(lr.sigma_v_eff_mpa,4),
                    _fmt(lr.c_mpa if lr.c_mpa > 0 else lr.su_mpa,4), _fmt(lr.alpha_or_beta,4), _fmt(lr.phi_strength,3), _fmt(lr.qs_factored_kpa,2), _fmt(lr.qs_factored_kn,2),
                    _fmt(lr.phi_extreme,3), _fmt(lr.qs_extreme_kpa,2), _fmt(lr.qs_extreme_kn,2), lr.note
                ])
                first_layer = False
        self.result_table.set_rows(rows)

    def _show_result(self, res: CapacityResult):
        self.summary_text.config(state=tk.NORMAL)
        self.summary_text.delete("1.0", tk.END)
        lines = []
        lines.append(f"Dự án: {res.pile_input.project} | Hạng mục: {res.pile_input.item}")
        lines.append(f"Loại bài toán: {res.pile_input.mode}; D = {_fmt(res.pile_input.diameter_mm,0)} mm; L = {_fmt(res.pile_input.pile_length_m,3)} m")
        lines.append(f"W cọc khô = {_fmt(res.strength.pile_weight_dry_kn,2)} kN; W' hữu hiệu = {_fmt(res.strength.pile_weight_effective_kn,2)} kN; Đẩy nổi = {_fmt(res.strength.buoyancy_kn,2)} kN")
        lines.append(f"FOS quyết định chiều dài cọc = {self._fmt_fos_value(self._governing_fos_for_result(res))}")
        lines.append(f"Mũi cọc: {res.toe_info.get('note','')}")
        lines.append("")
        for cap in (res.strength, res.extreme):
            lines.append(f"{cap.label}:")
            lines.append(f"  ΣQs = {_fmt(cap.qshaft_kn,2)} kN; Qp = {_fmt(cap.qtip_kn,2)} kN; Qr gross = {_fmt(cap.compression_single_gross_kn,2)} kN")
            if cap.downdrag_kn > 0:
                lines.append(f"  Ma sát âm DD = {_fmt(cap.downdrag_kn,2)} kN/cọc, γDD={_fmt(res.pile_input.downdrag_factor,2)}; DD cộng vào tải nén khi kiểm toán")
            lines.append(f"  [Q] cọc đơn sau trừ W' = {_fmt(cap.compression_single_net_kn,2)} kN")
            lines.append(f"  Nhổ cọc đơn |Qr,uplift| = {_fmt(cap.uplift_single_magnitude_kn,2)} kN; nhóm nhổ = {_fmt(cap.uplift_group_total_kn,2)} kN; dạng dấu = {_fmt(cap.uplift_single_signed_kn,2)} kN")
            lines.append(f"  fg = {_fmt(cap.group_factor,3)}; [Q] nhóm quy đổi/cọc = {_fmt(cap.compression_group_single_net_kn,2)} kN; Tổng nhóm = {_fmt(cap.compression_group_total_kn,2)} kN")
            lines.append(f"  Pr vật liệu = {_fmt(cap.material_pr_kn,2)} kN; Giá trị khống chế = {_fmt(cap.governing_kn,2)} kN")
            if cap.code == LIMIT_STATE_STRENGTH:
                lines.append("  " + self._limit_state_check_text(res, cap, res.pile_input.force_cd_kn, res.pile_input.cap_force_cd_kn, "Kiểm TTGHCĐ"))
                lines.append("  " + self._limit_state_uplift_check_text(res, cap, res.pile_input.uplift_cd_kn, "Kiểm TTGHCĐ"))
            else:
                lines.append("  " + self._limit_state_check_text(res, cap, res.pile_input.force_db_kn, res.pile_input.cap_force_db_kn, "Kiểm TTGHĐB"))
                lines.append("  " + self._limit_state_uplift_check_text(res, cap, res.pile_input.uplift_db_kn, "Kiểm TTGHĐB"))
        self.summary_text.insert(tk.END, "\n".join(lines))
        self.summary_text.config(state=tk.DISABLED)
        rows = []
        first_layer = True
        for lr in res.layers:
            rows.append([
                res.pile_input.item if first_layer else "",
                lr.name, _fmt(lr.top_elev_m,2), _fmt(lr.bottom_elev_m,2), _fmt(lr.thickness_m,2), _fmt(lr.skin_length_m,2),
                _fmt(lr.downdrag_length_m,2), _fmt(lr.downdrag_kn,2), lr.soil_label, _fmt(lr.group_factor,3), _fmt(lr.n_spt,2), _fmt(lr.n60,2), _fmt(lr.n1_60,2), _fmt(lr.sigma_v_eff_mpa,4), _fmt(lr.c_mpa if lr.c_mpa > 0 else lr.su_mpa,4),
                _fmt(lr.alpha_or_beta,4), _fmt(lr.phi_strength,3), _fmt(lr.qs_factored_kpa,2), _fmt(lr.qs_factored_kn,2),
                _fmt(lr.phi_extreme,3), _fmt(lr.qs_extreme_kpa,2), _fmt(lr.qs_extreme_kn,2), lr.note
            ])
            first_layer = False
        self.result_table.set_rows(rows)

    def _sanitize_filename(self, name: Any) -> str:
        """Làm sạch tên file Windows: <Hạng mục>_SCT."""
        return safe_filename(name, "HANG_MUC")

    def _unique_path(self, folder: str, stem: str, ext: str) -> str:
        path = os.path.join(folder, stem + ext)
        if not os.path.exists(path):
            return path
        i = 1
        while True:
            p = os.path.join(folder, f"{stem}_{i}{ext}")
            if not os.path.exists(p):
                return p
            i += 1

    def _report_check_values(self, res: CapacityResult) -> Dict[str, Any]:
        """Tính các giá trị kiểm toán nén/nhổ cọc đơn và nhóm để dùng chung DOCX/PDF."""
        n = max(int(res.pile_input.pile_count_in_group), 1)
        single_cd = min(res.strength.compression_single_net_kn, res.strength.material_pr_kn) if res.strength.material_pr_kn > 0 else res.strength.compression_single_net_kn
        single_db = min(res.extreme.compression_single_net_kn, res.extreme.material_pr_kn) if res.extreme.material_pr_kn > 0 else res.extreme.compression_single_net_kn
        group_cd = min(res.strength.compression_group_total_kn, res.strength.material_pr_kn * n) if res.strength.material_pr_kn > 0 else res.strength.compression_group_total_kn
        group_db = min(res.extreme.compression_group_total_kn, res.extreme.material_pr_kn * n) if res.extreme.material_pr_kn > 0 else res.extreme.compression_group_total_kn
        pu_cd = float(res.pile_input.force_cd_kn or 0.0)
        pu_db = float(res.pile_input.force_db_kn or 0.0)
        dd_cd = max(float(res.strength.downdrag_kn or 0.0), 0.0)
        dd_db = max(float(res.extreme.downdrag_kn or 0.0), 0.0)
        pu_cd_design = pu_cd + dd_cd
        pu_db_design = pu_db + dd_db
        ncap_cd = float(res.pile_input.cap_force_cd_kn or 0.0) if float(res.pile_input.cap_force_cd_kn or 0.0) > 0 else pu_cd * n
        ncap_db = float(res.pile_input.cap_force_db_kn or 0.0) if float(res.pile_input.cap_force_db_kn or 0.0) > 0 else pu_db * n
        ncap_cd_design = ncap_cd + dd_cd * n
        ncap_db_design = ncap_db + dd_db * n
        uplift_cd = float(res.pile_input.uplift_cd_kn or 0.0)
        uplift_db = float(res.pile_input.uplift_db_kn or 0.0)
        uplift_group_cd = uplift_cd * n
        uplift_group_db = uplift_db * n
        cap_uplift_cd = max(float(res.strength.uplift_single_magnitude_kn or 0.0), 0.0)
        cap_uplift_db = max(float(res.extreme.uplift_single_magnitude_kn or 0.0), 0.0)
        cap_uplift_group_cd = max(float(res.strength.uplift_group_total_kn or 0.0), 0.0)
        cap_uplift_group_db = max(float(res.extreme.uplift_group_total_kn or 0.0), 0.0)
        dcr_single_cd = pu_cd_design / single_cd if pu_cd_design > 0 and single_cd > 0 else 0.0
        dcr_single_db = pu_db_design / single_db if pu_db_design > 0 and single_db > 0 else 0.0
        dcr_group_cd = ncap_cd_design / group_cd if ncap_cd_design > 0 and group_cd > 0 else 0.0
        dcr_group_db = ncap_db_design / group_db if ncap_db_design > 0 and group_db > 0 else 0.0
        dcr_uplift_cd = uplift_cd / cap_uplift_cd if uplift_cd > 0 and cap_uplift_cd > 0 else 0.0
        dcr_uplift_db = uplift_db / cap_uplift_db if uplift_db > 0 and cap_uplift_db > 0 else 0.0
        dcr_uplift_group_cd = uplift_group_cd / cap_uplift_group_cd if uplift_group_cd > 0 and cap_uplift_group_cd > 0 else 0.0
        dcr_uplift_group_db = uplift_group_db / cap_uplift_group_db if uplift_group_db > 0 and cap_uplift_group_db > 0 else 0.0
        dcr_governing_cd = max(dcr_single_cd, dcr_group_cd, dcr_uplift_cd, dcr_uplift_group_cd)
        dcr_governing_db = max(dcr_single_db, dcr_group_db, dcr_uplift_db, dcr_uplift_group_db)
        dcr_governing_all = max(dcr_governing_cd, dcr_governing_db)
        return {
            "single_cd": single_cd, "single_db": single_db, "group_cd": group_cd, "group_db": group_db,
            "pu_cd": pu_cd, "pu_db": pu_db, "dd_cd": dd_cd, "dd_db": dd_db,
            "pu_cd_design": pu_cd_design, "pu_db_design": pu_db_design,
            "ncap_cd": ncap_cd, "ncap_db": ncap_db,
            "ncap_cd_design": ncap_cd_design, "ncap_db_design": ncap_db_design,
            "dcr_single_cd": dcr_single_cd, "dcr_single_db": dcr_single_db,
            "dcr_group_cd": dcr_group_cd, "dcr_group_db": dcr_group_db,
            "uplift_cd": uplift_cd, "uplift_db": uplift_db,
            "uplift_group_cd": uplift_group_cd, "uplift_group_db": uplift_group_db,
            "cap_uplift_cd": cap_uplift_cd, "cap_uplift_db": cap_uplift_db,
            "cap_uplift_group_cd": cap_uplift_group_cd, "cap_uplift_group_db": cap_uplift_group_db,
            "dcr_uplift_cd": dcr_uplift_cd, "dcr_uplift_db": dcr_uplift_db,
            "dcr_uplift_group_cd": dcr_uplift_group_cd, "dcr_uplift_group_db": dcr_uplift_group_db,
            "dcr_governing_cd": dcr_governing_cd, "dcr_governing_db": dcr_governing_db, "dcr_governing_all": dcr_governing_all,
            "status_single_cd": "ĐẠT" if dcr_single_cd <= 1.0 else "KHÔNG ĐẠT",
            "status_single_db": "ĐẠT" if dcr_single_db <= 1.0 else "KHÔNG ĐẠT",
            "status_group_cd": "ĐẠT" if dcr_group_cd <= 1.0 else "KHÔNG ĐẠT",
            "status_group_db": "ĐẠT" if dcr_group_db <= 1.0 else "KHÔNG ĐẠT",
            "status_uplift_cd": "ĐẠT" if dcr_uplift_cd <= 1.0 else "KHÔNG ĐẠT",
            "status_uplift_db": "ĐẠT" if dcr_uplift_db <= 1.0 else "KHÔNG ĐẠT",
            "status_uplift_group_cd": "ĐẠT" if dcr_uplift_group_cd <= 1.0 else "KHÔNG ĐẠT",
            "status_uplift_group_db": "ĐẠT" if dcr_uplift_group_db <= 1.0 else "KHÔNG ĐẠT",
        }

    def _capacity_report_title(self, res: CapacityResult) -> str:
        ptype = str(res.pile_input.pile_type or "").lower()
        if "đóng" in ptype or "dong" in ptype:
            return "SỨC CHỊU TẢI CỌC ĐÓNG"
        if "ép" in ptype or "ep" in ptype:
            return "SỨC CHỊU TẢI CỌC ÉP"
        return "SỨC CHỊU TẢI CỌC KHOAN NHỒI"

    def _soil_code_name(self, soil_type: int) -> str:
        return SOIL_TYPE_LABELS.get(int(soil_type or 0), "-")

    def _short_soil_name(self, soil_type: int, zero_is_top: bool = False) -> str:
        st = int(soil_type or 0)
        return {0:("Không khí" if zero_is_top else "Hang karst"),1:"Cát",2:"Sét",3:"Đá nguyên khối",4:"Đá nứt vỡ",5:"IGM",6:"Cuội sỏi"}.get(st,"-")

    def _zero_is_top_layer(self, res: CapacityResult, lr: LayerCalc) -> bool:
        seen_nonzero = False
        for x in res.layers:
            if x is lr:
                return not seen_nonzero
            if int(getattr(x, 'soil_type', 0) or 0) != 0:
                seen_nonzero = True
        return False

    def _blank_repeated_cells(self, rows: List[List[Any]], cols: List[int]) -> List[List[Any]]:
        out=[]
        prev={c:None for c in cols}
        for i,row in enumerate(rows):
            row=list(row)
            if i>0:
                for c in cols:
                    val = row[c] if c < len(row) else None
                    if val == prev.get(c) and str(val or '').strip():
                        row[c]=''
                    else:
                        prev[c]=val
            out.append(row)
        return out

    def _ton(self, kn: float) -> float:
        return float(kn or 0.0) / 9.80665

    def _docx_set_cell_shading(self, cell, fill: str):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        tc_pr.append(shd)

    def _report_blank_zero_text(self, text: Any) -> str:
        """Không in các giá trị bằng 0 trong bảng báo cáo; giữ nguyên chuỗi mô tả có chữ."""
        if text is None:
            return ""
        if isinstance(text, (int, float)):
            try:
                return "" if abs(float(text)) < 1e-12 else str(text)
            except Exception:
                return str(text)
        raw = str(text)
        t = raw.strip().replace("\u00a0", " ").replace(",", ".")
        if not t:
            return ""
        # Chỉ xoá nếu ô là một số 0 thuần túy hoặc số 0 kèm đơn vị; không động vào nhãn như "CĐ IGM N<100".
        if re.fullmatch(r"[-+]?0+(?:\.0+)?(?:\s*(?:kN|Tấn|tấn|MPa|Mpa|kPa|m|mm|%|độ))?", t, flags=re.IGNORECASE):
            return ""
        return raw

    def _report_symbol_subscript_runs(self, text: Any) -> List[Tuple[str, bool]]:
        """Tách chuỗi ký hiệu kỹ thuật thành các đoạn chữ chính/chữ phụ để in báo cáo.

        Không thay đổi dữ liệu tính toán; chỉ định dạng hiển thị trong DOCX/PDF.
        Các ký hiệu được chuẩn hóa theo dạng thường gặp trong TCVN 11823-10:2017,
        ví dụ C_N, (N_1)_60, N_60, N_160, q_s, Q_s, q_p, Q_p, S_u.
        """
        raw = self._report_blank_zero_text(text)
        if raw is None:
            return [("", False)]
        txt = str(raw)
        # Mẫu dài đặt trước mẫu ngắn để không tách nhầm: N160 trước N60/N1; sigma_vb trước sigma_v.
        specs: List[Tuple[str, List[Tuple[str, bool]]]] = [
            ("(N1)60", [("(", False), ("N", False), ("1", True), (")", False), ("60", True)]),
            ("N160", [("N", False), ("160", True)]),
            ("N60", [("N", False), ("60", True)]),
            ("Nₕₜ", [("N", False), ("ht", True)]),
            ("CN", [("C", False), ("N", True)]),
            ("Pmax", [("P", False), ("max", True)]),
            ("γDD", [("γ", False), ("DD", True)]),
            ("γc", [("γ", False), ("c", True)]),
            ("σ'vb", [("σ'", False), ("vb", True)]),
            ("σ'v", [("σ'", False), ("v", True)]),
            ("σ'p", [("σ'", False), ("p", True)]),
            ("f'c", [("f'", False), ("c", True)]),
            ("qBN", [("q", False), ("BN", True)]),
            ("qλ", [("q", False), ("λ", True)]),
            ("Ksp", [("K", False), ("sp", True)]),
            ("Ld", [("L", False), ("d", True)]),
            ("Db", [("D", False), ("b", True)]),
            ("Ap", [("A", False), ("p", True)]),
            ("Dp", [("D", False), ("p", True)]),
            ("Pr", [("P", False), ("r", True)]),
            ("Qr", [("Q", False), ("r", True)]),
            ("Qs", [("Q", False), ("s", True)]),
            ("Qp", [("Q", False), ("p", True)]),
            ("qs", [("q", False), ("s", True)]),
            ("qp", [("q", False), ("p", True)]),
            ("Su", [("S", False), ("u", True)]),
            ("Nc", [("N", False), ("c", True)]),
            ("Nq", [("N", False), ("q", True)]),
            ("pa", [("p", False), ("a", True)]),
            ("φf", [("φ", False), ("f", True)]),
            ("αE", [("α", False), ("E", True)]),
            ("fsn", [("f", False), ("sn", True)]),
            ("mb", [("m", False), ("b", True)]),
            ("mi", [("m", False), ("i", True)]),
        ]
        out: List[Tuple[str, bool]] = []
        i = 0
        while i < len(txt):
            hit = None
            for pat, runs in specs:
                if txt.startswith(pat, i):
                    # Không đổi CN nằm trong chuỗi chữ dài kiểu TCVN; chỉ đổi khi đứng độc lập hoặc sau ký tự không phải chữ/số.
                    if pat == "CN":
                        prev = txt[i-1] if i > 0 else " "
                        nxt = txt[i+len(pat)] if i + len(pat) < len(txt) else " "
                        if (prev.isalnum() or nxt.isalnum()):
                            continue
                    hit = (pat, runs)
                    break
            if hit:
                for seg, sub in hit[1]:
                    if seg:
                        out.append((seg, sub))
                i += len(hit[0])
            else:
                out.append((txt[i], False))
                i += 1
        # Gộp các đoạn liên tiếp cùng kiểu để giảm số run.
        merged: List[Tuple[str, bool]] = []
        for seg, sub in out:
            if not seg:
                continue
            if merged and merged[-1][1] == sub:
                merged[-1] = (merged[-1][0] + seg, sub)
            else:
                merged.append((seg, sub))
        return merged or [("", False)]

    def _report_pdf_markup(self, text: Any) -> str:
        """Chuỗi an toàn cho ReportLab Paragraph, có subscript cho ký hiệu kỹ thuật."""
        import html
        clean = self._report_blank_zero_text(text)
        lines = str(clean).split("\n")
        rendered: List[str] = []
        for line in lines:
            parts: List[str] = []
            for seg, sub in self._report_symbol_subscript_runs(line):
                esc = html.escape(seg, quote=False)
                parts.append(f"<sub>{esc}</sub>" if sub else esc)
            rendered.append("".join(parts))
        return "<br/>".join(rendered)

    def _report_row_emphasis(self, row: List[Any], row_index: int = 0, header_rows: int = 1) -> bool:
        if row_index < header_rows or not row:
            return False
        label = _strip_accents(str(row[0] if row else "")).lower()
        return ("fos" in label) or label.startswith("kiem tra") or (" dat" in label) or ("khong dat" in label)

    def _docx_set_cell_text(self, cell, text: Any, bold: bool = False, color: Optional[str] = None, size: int = 9, align_center: bool = False, no_wrap: bool = False):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        cell.text = ""
        p = cell.paragraphs[0]
        try:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        except Exception:
            pass
        if align_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        clean_text = self._report_blank_zero_text(text)
        parts = str(clean_text).split("\n")
        r = None
        for idx, part in enumerate(parts):
            if idx > 0:
                br = p.add_run()
                br.add_break()
            for seg, is_sub in self._report_symbol_subscript_runs(part):
                r = p.add_run(seg)
                r.bold = bool(bold)
                r.font.name = "Times New Roman"
                r.font.size = Pt(size)
                r.font.subscript = bool(is_sub)
        if r is None:
            r = p.add_run("")
            r.bold = bool(bold)
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
        if color:
            color = color.replace("#", "")
            rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
            for _run in p.runs:
                _run.font.color.rgb = rgb
        try:
            cell.vertical_alignment = 1
        except Exception:
            pass
        if no_wrap:
            self._docx_set_cell_no_wrap(cell)

    def _docx_set_table_width(self, table, width_cm: float):
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            tbl_pr = table._tbl.tblPr
            tbl_w = tbl_pr.find(qn('w:tblW'))
            if tbl_w is None:
                tbl_w = OxmlElement('w:tblW')
                tbl_pr.append(tbl_w)
            tbl_w.set(qn('w:type'), 'dxa')
            tbl_w.set(qn('w:w'), str(int(float(width_cm) * 567)))
            # Cố định layout bảng để Word không tự co cột gây tách chữ.
            tbl_layout = tbl_pr.find(qn('w:tblLayout'))
            if tbl_layout is None:
                tbl_layout = OxmlElement('w:tblLayout')
                tbl_pr.append(tbl_layout)
            tbl_layout.set(qn('w:type'), 'fixed')
        except Exception:
            pass

    def _docx_set_cell_margins(self, cell, top: int = 18, start: int = 35, bottom: int = 18, end: int = 35):
        """Giảm padding trong bảng DOCX. Đơn vị twips; 35 twips xấp xỉ 0.06 cm."""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn('w:tcMar'))
            if tc_mar is None:
                tc_mar = OxmlElement('w:tcMar')
                tc_pr.append(tc_mar)
            for m, v in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
                node = tc_mar.find(qn(f'w:{m}'))
                if node is None:
                    node = OxmlElement(f'w:{m}')
                    tc_mar.append(node)
                node.set(qn('w:w'), str(int(v)))
                node.set(qn('w:type'), 'dxa')
        except Exception:
            pass

    def _report_auto_table_font_size(self, rows: List[List[Any]], widths: Optional[List[float]], default_size: int = 10) -> int:
        try:
            ncol = max(len(r) for r in rows) if rows else 0
            # Các bảng rất nhiều cột chỉ đủ đẹp khi giảm về 9 pt.
            if ncol >= 8:
                return min(int(default_size), 9)
            return int(default_size)
        except Exception:
            return int(default_size)

    def _docx_make_table(self, parent, rows: List[List[Any]], header_rows: int = 1, widths: Optional[List[float]] = None, font_size: int = 10, merge_repeat_cols: Optional[List[int]] = None, min_row_height_in: float = 0.22, table_align: str = "left", left_cols: Optional[List[int]] = None):
        from docx.shared import Cm, Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
        if not rows:
            return None
        ncol = max(len(r) for r in rows)
        font_size = self._report_auto_table_font_size(rows, widths, font_size)
        table = parent.add_table(rows=len(rows), cols=ncol)
        table.style = "Table Grid"
        align_key = str(table_align or "left").strip().lower()
        left_cols_set = set(int(c) for c in (left_cols or []))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER if align_key == "center" else WD_TABLE_ALIGNMENT.LEFT
        try:
            table.autofit = False
            table.allow_autofit = False
        except Exception:
            pass
        if widths:
            try:
                self._docx_set_table_width(table, sum(float(w) for w in widths))
                for j, w_cm in enumerate(widths[:ncol]):
                    table.columns[j].width = Cm(float(w_cm))
                    for c in table.columns[j].cells:
                        c.width = Cm(float(w_cm))
            except Exception:
                pass
        for i, row in enumerate(rows):
            tr = table.rows[i]
            row_emphasis = self._report_row_emphasis(row, i, header_rows)
            try:
                tr.height = Inches(min_row_height_in)
                tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            except Exception:
                pass
            for j in range(ncol):
                cell = table.cell(i, j)
                self._docx_set_cell_margins(cell)
                value = row[j] if j < len(row) else ""
                is_header = i < header_rows
                align_center = True if is_header else (False if j in left_cols_set or j == 0 else True)
                self._docx_set_cell_text(
                    cell, value,
                    bold=(is_header or row_emphasis),
                    color=("C00000" if row_emphasis and not is_header else None),
                    size=font_size,
                    align_center=align_center,
                    no_wrap=self._report_cell_should_nowrap(value),
                )
                if is_header:
                    self._docx_set_cell_shading(cell, "D9EAF7")
                elif i % 2 == 0:
                    self._docx_set_cell_shading(cell, "F8FAFC")
                if widths and j < len(widths):
                    try:
                        cell.width = Cm(float(widths[j]))
                    except Exception:
                        pass
        if merge_repeat_cols:
            for c in merge_repeat_cols:
                i = header_rows
                while i < len(rows):
                    val = rows[i][c] if c < len(rows[i]) else ""
                    j = i
                    while j + 1 < len(rows):
                        nxt = rows[j + 1][c] if c < len(rows[j + 1]) else ""
                        if str(val).strip() and nxt == val:
                            j += 1
                        else:
                            break
                    if j > i:
                        try:
                            table.cell(i, c).merge(table.cell(j, c))
                            self._docx_set_cell_text(table.cell(i, c), val, size=font_size, align_center=(False if c in left_cols_set or c == 0 else True), no_wrap=self._report_cell_should_nowrap(val))
                        except Exception:
                            pass
                    i = j + 1
        return table

    def _docx_heading(self, doc, text: str, level: int = 1):
        from docx.shared import Pt, Inches
        p = doc.add_paragraph()
        try:
            p.paragraph_format.space_before = Inches(0.10 if level == 1 else 0.08)
            p.paragraph_format.space_after = Inches(0.04)
            p.paragraph_format.left_indent = Inches(0)
        except Exception:
            pass
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12 if level == 1 else 10)
        return p

    def _report_content_width_cm(self) -> float:
        try:
            page_w_cm, _page_h_cm = self._report_docx_page_cm()
            return max(float(page_w_cm) - 2.70, 12.0)
        except Exception:
            return 18.30

    def _scale_widths_to_total(self, widths: List[float], total_cm: Optional[float] = None) -> List[float]:
        vals = [float(w) for w in (widths or []) if float(w) > 0]
        if not vals:
            return []
        total = float(total_cm if total_cm is not None else self._report_content_width_cm())
        s = sum(vals) or 1.0
        return [max(0.45, total * w / s) for w in vals]

    def _report_cell_should_nowrap(self, text: Any) -> bool:
        t = self._report_blank_zero_text(text).strip()
        if not t:
            return False
        # Không cho Word bẻ đôi các nhãn/ngắn hoặc giá trị số; câu dài vẫn được xuống dòng theo khoảng trắng.
        if re.search(r"\s", t):
            return False
        return len(t) <= 24

    def _docx_set_cell_no_wrap(self, cell):
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            tc_pr = cell._tc.get_or_add_tcPr()
            if tc_pr.find(qn('w:noWrap')) is None:
                tc_pr.append(OxmlElement('w:noWrap'))
        except Exception:
            pass

    def _docx_widths_from_sample(self, kind: str) -> List[float]:
        """Bề rộng cột theo mẫu 2222_SCT.docx người dùng gửi."""
        preset = {
            # Bảng thông số đầu vào đặt trong cột trái cạnh hình sơ họa.
            # Tổng bề rộng phải nhỏ hơn cột trái để hình bên phải không lấn mất cột Đơn vị.
            "input": [5.40, 1.80, 3.10, 1.55],
            "phi": [2.40, 1.75, 1.75, 1.75, 1.75, 1.75, 1.55, 1.55, 1.55, 1.55],
            "factor": [9.051, 9.051],
            "shaft_formula": [4.159, 7.909, 6.034],
            "tip_formula": [3.683, 8.386, 6.034],
            "tip_summary": [9.051, 9.051],
            "capacity": [9.051, 4.525, 4.525],
            "check": [9.051, 4.525, 4.525],
        }
        vals = list(preset.get(kind, []))
        if vals and kind != "input":
            return self._scale_widths_to_total(vals, self._report_content_width_cm())
        return vals

    def _layer_table_widths_cm(self, rows: List[List[Any]], total_cm: float = 18.10) -> List[float]:
        # Đúng theo bảng mẫu 2222_SCT.docx: cột σ'v và Qs rộng hơn, Lớp/SPT hẹp.
        fixed = []
        if rows and len(rows[0]) == len(fixed):
            return fixed
        if not rows:
            return []
        header = [str(x).strip() for x in rows[0]]
        weights = []
        for h in header:
            h_key = re.sub(r"\s+", " ", str(h).replace("\n", " ")).strip()
            if h_key in ("Tên lớp", "Lớp"):
                w = 1.08
            elif h_key == "Dày":
                w = 1.02
            elif h_key == "Loại":
                w = 1.10
            elif h_key == "Nₕₜ":
                w = 0.88
            elif h_key == "(N1)60":
                w = 1.08
            elif h_key == "σ'v (MPa)":
                w = 1.42
            elif h_key == "Su (kPa)":
                w = 1.12
            elif h_key in ("β", "α", "αE"):
                w = 0.72
            elif h_key == "φf (độ)":
                w = 1.08
            elif h_key in ("qu đá (MPa)", "qu (MPa)"):
                w = 1.12
            elif h_key == "RQD (%)":
                w = 1.12
            elif h_key in ("qs CĐ (kPa)", "qs ĐB (kPa)"):
                w = 1.22
            elif h_key in ("Qs CĐ (kN)", "Qs ĐB (kN)"):
                w = 1.40
            elif h_key in ("Bề dày (m)",):
                w = 1.08
            else:
                w = 1.00
            weights.append(w)
        s = sum(weights) or 1.0
        return [max(0.50, total_cm * w / s) for w in weights]

    def _docx_add_kv_table(self, parent, rows: List[Tuple[str, str, Any, str]], table_align: str = "left"):
        data = [["Nội dung", "Ký hiệu", "Giá trị", "Đơn vị"]]
        for r in rows:
            data.append([r[0], r[1], r[2], r[3]])
        return self._docx_make_table(parent, data, header_rows=1, widths=self._docx_widths_from_sample("input"), font_size=10, min_row_height_in=0.22, table_align=table_align)

    def _report_input_rows(self, res: CapacityResult) -> List[Tuple[str, str, Any, str]]:
        inp = res.pile_input
        rows = [
            ("Tên hạng mục", "", inp.item, ""),
            ("Loại cọc", "", inp.pile_type, ""),
            ("Tiết diện cọc", "", ({"VUONG": "Vuông đặc", "TRON": "Tròn đặc",
                                    "ONG": f"Ống, D trong = {_fmt(inp.inner_diameter_m, 3)} m"}.get(inp.driven_shape, "Tròn đặc")
                                   if SCTCalculator._is_driven(inp) else "Tròn đặc"), ""),
            ("Chiều dài cọc", "L", _fmt(inp.pile_length_m, 3), "(m)"),
            ("Đường kính cọc", "D", _fmt(inp.diameter_m, 3), "(m)"),
            ("Đường kính mũi cọc", "Dp", _fmt(inp.tip_diameter_m, 3), "(m)"),
            ("Chu vi cọc", "P", _fmt(inp.perimeter_m, 3), "(m)"),
            ("Diện tích mũi cọc", "Ap", _fmt(inp.area_m2, 3), "(m²)"),
            ("S (khoảng cách cọc)", "S", _fmt(inp.spacing_m, 3), "(m)"),
            ("Cao độ thiên nhiên", "", _fmt(inp.ground_elev_m, 3), "(m)"),
            ("Cao độ đáy bệ", "", _fmt(inp.cap_bottom_elev_m, 3), "(m)"),
            ("Cao độ mũi cọc", "", _fmt(inp.pile_tip_elev_m, 3), "(m)"),
            ("Số cọc", "n", _fmt(inp.pile_count_in_group, 0), "(cọc)"),
            ("Số hàng cọc", "", inp.group_layout, "(hàng)"),
            ("Cường độ chịu nén bê tông cọc", "f'c", _fmt(inp.fc_mpa, 3), "(MPa)"),
            ("Khối lượng riêng bê tông", "γc", _fmt(inp.concrete_gamma_kN_m3, 2), "(kN/m³)"),
            ("Nₕₜ nhập vào", "", inp.spt_input_mode, ""),
            ("Hiệu quả búa SPT", "ER", _fmt(inp.spt_er_percent, 0), "(%)"),
        ]
        if bool(getattr(inp, "include_downdrag", False)):
            rows.append(("Hệ số tải trọng ma sát âm", "γDD", _fmt(inp.downdrag_factor, 2), ""))
        return rows

    def _report_phi_heading(self, res: CapacityResult) -> str:
        if SCTCalculator._is_driven(res.pile_input):
            return "2.1. Hệ số sức kháng (Bảng 9: TCVN 11823-10:2017)"
        return "2.1. Hệ số sức kháng (Bảng 10: TCVN 11823-10:2017)"

    def _report_phi_widths_cm(self) -> List[float]:
        return self._scale_widths_to_total([6.20, 3.00, 3.00, 2.20], self._report_content_width_cm())

    def _report_phi_rows(self, res: CapacityResult) -> List[List[Any]]:
        """Bảng hệ số sức kháng in trong báo cáo, tách theo loại cọc.

        Cọc khoan dùng nhóm hệ số của Bảng 10; cọc đóng/ép dùng nhóm hệ số
        phân tích tĩnh của Bảng 9 cho đất dính theo α và đất rời theo Meyerhof SPT.
        Chỉ in các nhóm đất/đá có xuất hiện trong hạng mục để tránh gây nhầm lẫn.
        """
        stypes = self._soil_types_in_result(res)
        is_driven = SCTCalculator._is_driven(res.pile_input)
        rows: List[List[Any]] = [["Nhóm đất/đá / phương pháp", "φ bên - TTGHCĐ", "φ mũi - TTGHCĐ", "φ - TTGHĐB"]]
        if is_driven:
            if 2 in stypes:
                rows.append(["Đất dính - phương pháp α", _fmt(PHI_DRIVEN_SHAFT_CLAY, 2), _fmt(PHI_DRIVEN_TIP_CLAY, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
            if any(t in stypes for t in (1, 6)):
                rows.append(["Đất rời - Meyerhof SPT", _fmt(PHI_DRIVEN_SHAFT_SAND, 2), _fmt(PHI_DRIVEN_TIP_SAND, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
            if 5 in stypes:
                rows.append(["IGM N<100", _fmt(PHI_BORE_SHAFT_IGM_SPT_LT100, 2), _fmt(PHI_BORE_SHAFT_IGM_SPT_LT100, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
                rows.append(["IGM N≥100", _fmt(PHI_BORE_SHAFT_IGM_SPT_GE100, 2), _fmt(PHI_BORE_SHAFT_IGM_SPT_GE100, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
            if any(t in stypes for t in (3, 4)):
                rows.append(["Đá", _fmt(PHI_ROCK_SIDE, 2), _fmt(PHI_ROCK_TIP, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
        else:
            if 2 in stypes:
                rows.append(["Đất dính", _fmt(PHI_BORE_SHAFT_CLAY, 2), _fmt(PHI_BORE_TIP_CLAY, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
            if any(t in stypes for t in (1, 6)):
                rows.append(["Đất rời", _fmt(PHI_BORE_SHAFT_SAND, 2), _fmt(PHI_BORE_TIP_SAND, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
            if 5 in stypes:
                rows.append(["IGM N<100", _fmt(PHI_BORE_SHAFT_IGM_SPT_LT100, 2), _fmt(PHI_BORE_SHAFT_IGM_SPT_LT100, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
                rows.append(["IGM N≥100", _fmt(PHI_BORE_SHAFT_IGM_SPT_GE100, 2), _fmt(PHI_BORE_SHAFT_IGM_SPT_GE100, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
            if any(t in stypes for t in (3, 4)):
                rows.append(["Đá", _fmt(PHI_ROCK_SIDE, 2), _fmt(PHI_ROCK_TIP, 2), _fmt(PHI_EXTREME_DEFAULT, 2)])
        if len(rows) == 1:
            rows.append(["Không có lớp tính sức kháng", "", "", ""])
        return rows

    def _report_tip_summary_note(self, res: CapacityResult) -> str:
        """Ghi chú ngắn, trung tính cho bảng tóm tắt mũi cọc trong báo cáo."""
        toe_type = int(res.toe_info.get("soil_type", 0) or 0)
        is_driven = SCTCalculator._is_driven(res.pile_input)
        if toe_type == 2:
            if is_driven and not res.pile_input.clay_use_c_phi:
                return "Mũi cọc trong đất dính: qp = 9·Su."
            if res.pile_input.clay_use_c_phi:
                return "Mũi cọc trong đất dính: tính theo thông số c, φ."
            return "Mũi cọc trong đất dính: qp = Nc·Su."
        if toe_type in (1, 6):
            if is_driven:
                return "Mũi cọc trong đất rời: Meyerhof SPT."
            return "Mũi cọc trong đất rời: qp = 0.057·N60."
        if toe_type == 5:
            return "Mũi cọc trong IGM: tính theo nhánh IGM."
        if toe_type in (3, 4):
            return "Mũi cọc trong đá: tính theo điều kiện khai báo."
        if toe_type == 0:
            return "Mũi cọc nằm trong lớp rỗng; không tính sức kháng mũi."
        return "Kiểm tra lại lớp đất/đá tại mũi cọc."

    def _report_factor_rows(self, res: CapacityResult) -> List[List[Any]]:
        inp = res.pile_input
        s_over_d = inp.spacing_m / max(inp.diameter_m, 1e-9)
        # Lấy fg theo từng nhóm vật liệu có xuất hiện trong lớp đất.
        cats: List[Tuple[str, List[int]]] = [("Đất sét", [2]), ("Đất cát/đất rời", [1, 6]), ("IGM/đá", [3, 4, 5])]
        rows = [["Nhóm đất/đá", "Hệ số nhóm fg"]]
        for name, codes in cats:
            if any(int(lr.soil_type) in codes and lr.skin_length_m > 0 for lr in res.layers):
                # Lấy hệ số đầu tiên cùng nhóm, vì cùng S/D và số hàng. Nếu option bỏ qua IGM/đá thì có thể =1.
                fg = None
                for lr in res.layers:
                    if int(lr.soil_type) in codes and lr.skin_length_m > 0:
                        fg = lr.group_factor
                        break
                rows.append([name, _fmt(fg if fg is not None else 1.0, 3)])
        return rows

    def _soil_types_in_result(self, res: CapacityResult) -> set:
        return {int(lr.soil_type) for lr in res.layers if lr.skin_length_m > 0 or abs(lr.thickness_m) > 0}

    def _report_shaft_formula_rows(self, res: CapacityResult) -> List[List[Any]]:
        stypes = self._soil_types_in_result(res)
        rows = [["Mục", "Công thức", "Nguồn"]]
        is_driven = SCTCalculator._is_driven(res.pile_input)
        if 2 in stypes:
            if res.pile_input.clay_use_c_phi:
                rows += [
                    ["Sét", "qs = c + σ'v·tanφ", ""],
                    ["Trong đó", "c là lực dính; φ là góc nội ma sát; σ'v là ứng suất hữu hiệu tại giữa lớp.", ""],
                ]
            elif is_driven:
                rows += [
                    ["Sét - cọc đóng/ép", "qs = α·Su", "TCVN 11823-10:2017, Điều 7.3.8.6.2"],
                    ["Hệ số α", "α lấy theo Tomlinson.", "TCVN 11823-10:2017, Điều 7.3.8.6.2, Hình 18"],
                    ["Su", "Su = 0.006·N60 MPa", ""],
                ]
            else:
                rows += [
                    ["Sét", "qs = α·Su", "TCVN 11823-10:2017, Điều 8.3.5.1.2"],
                    ["Hệ số α", "α = 0.55 khi Su/pa ≤ 1.5", "TCVN 11823-10:2017, Điều 8.3.5.1.2"],
                    ["Hệ số α", "α = 0.55 - 0.1·(Su/pa - 1.5) khi 1.5 < Su/pa ≤ 2.5", "TCVN 11823-10:2017, Điều 8.3.5.1.2"],
                    ["Su", "Su = 0.006·N60 MPa", ""],
                    ["pa", "pa = 0.101 MPa là áp suất khí quyển.", ""],
                ]
        if any(t in stypes for t in (1, 6)):
            if is_driven:
                qs_factor = "0.96" if _is_nondisplacement_driven_pile(res.pile_input.pile_type) else "1.9"
                rows += [
                    ["Cát - cọc đóng/ép", f"qs = {qs_factor}·N160 kPa", "TCVN 11823-10:2017, Điều 7.3.8.6.7"],
                    ["N160", "N160 = (N1)60", "TCVN 11823-10:2017, Điều 7.3.8.6.7"],
                    ["(N1)60", "(N1)60 = CN·N60; CN = 0.77·log10(1.92/σ'v) ≤ 2.0", "TCVN 11823-10:2017, Điều 4.6.2.4"],
                ]
            else:
                rows += [
                    ["Cát", "qs = β·σ'v", "TCVN 11823-10:2017, Điều 8.3.5.2.2"],
                    ["Hệ số β", "β = (1 - sinφf)·(σ'p/σ'v)^sinφf·tanφf", "TCVN 11823-10:2017, Điều 8.3.5.2.2"],
                    ["Góc ma sát φf", "φf = 27.5 + 9.2·log10((N1)60)", "TCVN 11823-10:2017, Điều 8.3.5.2.2"],
                    ["(N1)60", "(N1)60 = CN·N60; CN = 0.77·log10(1.92/σ'v) ≤ 2.0", "TCVN 11823-10:2017, Điều 4.6.2.4"],
                    ["N60", "N60 = (ER/60%)·Nₕₜ", ""],
                    ["Ứng suất tiền cố kết - cát", "σ'p = 0.47·N60^m·pa", "TCVN 11823-10:2017, Điều 8.3.5.2.2"],
                    ["Ứng suất tiền cố kết - cuội sỏi", "σ'p = 0.15·N60·pa", "TCVN 11823-10:2017, Điều 8.3.5.2.2"],
                ]
        if 5 in stypes:
            rows += [
                ["IGM loại 1", "SPT < 100: fsn = 0.0036·N MPa", "TCVN 11823-10:2017, Phụ lục B"],
                ["IGM loại 2", "SPT ≥ 100 và qu < 4.7 MPa: fsn = 0.098√qu MPa", "TCVN 11823-10:2017, Phụ lục B"],
                ["IGM loại 3", "4.7 ≤ qu < 23.9 MPa: fsn = 0.81·qu^0.51 MPa", "TCVN 11823-10:2017, Phụ lục B"],
            ]
        if any(t in stypes for t in (3, 4)):
            rows += [
                ["Đá - thi công không chống đỡ", "qs = C·pa·(min(qu,f'c)/pa)^0.5, với thiết bị khoan thông thường lấy C=1.0", "TCVN 11823-10:2017, Điều 8.3.5.4.2, Công thức (96)"],
                ["Đá - thi công có chống đỡ", "qs = 0.65·αE·pa·(min(qu,f'c)/pa)^0.5", "TCVN 11823-10:2017, Điều 8.3.5.4.2, Công thức (97)"],
                ["Trong đó", "αE nội suy theo RQD và điều kiện khe nứt đá; qu là cường độ nén dọc trục của đá.", "TCVN 11823-10:2017, Điều 8.3.5.4.2"],
            ]
        return rows

    def _report_tip_formula_rows(self, res: CapacityResult) -> List[List[Any]]:
        toe_type = int(res.toe_info.get("soil_type", 0) or 0)
        rows = [["Mục", "Công thức", "Nguồn"]]
        is_driven = SCTCalculator._is_driven(res.pile_input)
        if toe_type == 2:
            if is_driven and not res.pile_input.clay_use_c_phi:
                rows += [["Sét - cọc đóng/ép", "qp = 9·Su", "TCVN 11823-10:2017, Điều 7.3.8.6.5"]]
            elif res.pile_input.clay_use_c_phi:
                rows += [
                    ["Sét", "qp = c·Nc + σ'v·Nq", ""],
                    ["Trong đó", "Nc, Nq là hệ số sức chịu tải theo φ; c là lực dính; σ'v là ứng suất hữu hiệu tại mũi.", ""],
                ]
            else:
                rows += [
                    ["Sét", "qp = Nc·Su ≤ 4.0 MPa", "TCVN 11823-10:2017, Điều 8.3.5.1.3"],
                    ["Hệ số Nc", "Nc = 6·[1 + 0.2·(Z/D)] ≤ 9", "TCVN 11823-10:2017, Điều 8.3.5.1.3"],
                    ["Điều kiện Su nhỏ", "Nếu Su < 0.024 MPa trong phạm vi 2D dưới mũi thì Nc nhân 0.67.", "TCVN 11823-10:2017, Điều 8.3.5.1.3"],
                ]
        elif toe_type in (1, 6):
            if is_driven:
                rows += [
                    ["Cát - cọc đóng/ép", "qp = 0.038·N160·Db/D ≤ qλ", "TCVN 11823-10:2017, Điều 7.3.8.6.7"],
                    ["Trong đó", "Db là chiều dài cọc ngập trong tầng đất chịu lực; D là bề rộng hoặc đường kính cọc.", "TCVN 11823-10:2017, Điều 7.3.8.6.7"],
                    ["Giới hạn qλ", "qλ = 0.4·N160 MPa đối với cát; qλ = 0.3·N160 MPa đối với cát bột không pha sét.", "TCVN 11823-10:2017, Điều 7.3.8.6.7"],
                ]
            else:
                rows += [
                    ["Cát", "qp = 0.057·N60 ≤ 3.0 MPa", "TCVN 11823-10:2017, Điều 8.3.5.2.3"],
                    ["Trong đó", "N60 là số búa đã hiệu chỉnh theo năng lượng búa.", "TCVN 11823-10:2017, Điều 8.3.5.2.3"],
                ]
        elif toe_type == 5:
            rows += [
                ["IGM loại 1", "SPT < 100: qBN = 0.0439·N MPa", "TCVN 11823-10:2017, Phụ lục B"],
                ["IGM loại 2", "SPT ≥ 100 và qu < 4.7 MPa: qBN = 0.057(1+Ld/D)√qu MPa", "TCVN 11823-10:2017, Phụ lục B"],
                ["IGM", "qBN = 3·qu·Ksp·d", ""],
            ]
        elif toe_type == 3:
            if res.pile_input.include_rock_tip:
                rows += [
                    ["Đá tốt", "Khi thỏa điều kiện áp dụng: qp = 2.5·qu", "TCVN 11823-10:2017, Điều 8.3.5.4.3"],
                    ["Đá không đủ điều kiện 2.5qu", "qp theo Hoek-Brown/GSI và giới hạn qp ≤ 2.5·qu", "TCVN 11823-10:2017, Điều 8.3.5.4.3"],
                    ["Thông số Hoek-Brown", "A = σ'vb + qu[mb(σ'vb/qu)]^a; qp = A + qu[mb(A/qu)+s]^a", "TCVN 11823-10:2017, Điều 8.3.5.4.3, Công thức (99)-(100)"],
                ]
            else:
                rows += [["Đá", "Bỏ qua sức kháng mũi, chỉ lấy sức kháng thân", ""]]
        elif toe_type == 4:
            if res.pile_input.include_rock_tip:
                rows += [
                    ["Đá nứt vỡ/phong hóa", "qp theo Hoek-Brown/GSI và giới hạn qp ≤ 2.5·qu", "TCVN 11823-10:2017, Điều 8.3.5.4.3"],
                    ["Thông số Hoek-Brown", "A = σ'vb + qu[mb(σ'vb/qu)]^a; qp = A + qu[mb(A/qu)+s]^a", "TCVN 11823-10:2017, Điều 8.3.5.4.3, Công thức (99)-(100)"],
                    ["Thông số GSI", "s = e^((GSI-100)/(9-3D)); a = 1/2 + 1/6·(e^(-GSI/15)-e^(-20/3)); mb = mi·e^((GSI-100)/(28-14D))", "TCVN 11823-10:2017, Điều 4.6.4"],
                ]
            else:
                rows += [["Đá nứt vỡ/phong hóa", "Bỏ qua sức kháng mũi, chỉ lấy sức kháng thân", ""]]
        elif toe_type == 0:
            rows += [["Không khí/Hang karst", "Không tính sức kháng mũi trong lớp rỗng.", ""]]
        else:
            rows += [["Mũi cọc", "Không xác định được lớp đất/đá tại mũi cọc hoặc không tính Qp.", ""]]
        return rows

    def _report_layer_rows(self, res: CapacityResult) -> List[List[Any]]:
        stypes = self._soil_types_in_result(res)
        has_clay = 2 in stypes
        has_sand = any(t in stypes for t in (1, 6))
        has_igm = 5 in stypes
        has_rock = any(t in stypes for t in (3, 4))
        header = ["Lớp", "Dày", "Loại", "Nₕₜ", "(N1)60", "σ'v\n(MPa)"]
        if has_clay: header += ["Su\n(kPa)", "α"]
        if has_sand: header += ["φf\n(độ)", "β"]
        if has_igm: header += ["qu\n(MPa)"]
        if has_rock: header += ["qu đá\n(MPa)", "RQD\n(%)", "αE"]
        header += ["qs CĐ\n(kPa)", "Qs CĐ\n(kN)", "qs ĐB\n(kPa)", "Qs ĐB\n(kN)"]
        rows=[header]
        for lr in res.layers:
            st=int(lr.soil_type)
            row=[lr.name, _fmt(lr.thickness_m,1), self._short_soil_name(st, self._zero_is_top_layer(res, lr)), _fmt(lr.n_spt,0), _fmt(lr.n1_60,0), _fmt(lr.sigma_v_eff_mpa,3)]
            if has_clay: row += [_fmt((lr.su_mpa if st == 2 else 0.0)*1000.0,1) if st==2 else "", _fmt(lr.alpha_or_beta,3) if st==2 else ""]
            if has_sand: row += [_fmt(lr.phi_deg,2) if st in (1,6) else "", _fmt(lr.alpha_or_beta,2) if st in (1,6) else ""]
            if has_igm: row += [_fmt(getattr(lr,'qu_mpa',0.0),3) if st==5 else ""]
            if has_rock: row += [_fmt(getattr(lr,'qu_mpa',0.0),3) if st in (3,4) else "", _fmt(getattr(lr,'rqd',0.0),1) if st in (3,4) else "", _fmt(lr.alpha_or_beta,3) if st in (3,4) else ""]
            row += [_fmt(lr.qs_factored_kpa,2), _fmt(lr.qs_factored_kn,2), _fmt(lr.qs_extreme_kpa,2), _fmt(lr.qs_extreme_kn,2)]
            rows.append(row)
        if len(rows) > 1:
            total = ["Tổng cộng", _fmt(sum(float(getattr(lr, 'thickness_m', 0.0) or 0.0) for lr in res.layers), 2), "", "", "", ""]
            if has_clay: total += ["", ""]
            if has_sand: total += ["", ""]
            if has_igm: total += [""]
            if has_rock: total += ["", "", ""]
            total += ["", _fmt(sum(float(getattr(lr, 'qs_factored_kn', 0.0) or 0.0) for lr in res.layers), 2), "", _fmt(sum(float(getattr(lr, 'qs_extreme_kn', 0.0) or 0.0) for lr in res.layers), 2)]
            rows.append(total)
        return rows

    def _report_include_ton(self) -> bool:
        return "tấn" in str(getattr(self, "_report_result_units", "kN") or "").lower()

    def _report_capacity_rows(self, res: CapacityResult) -> List[List[Any]]:
        st = res.strength
        ex = res.extreme
        gov_st = min(st.compression_single_net_kn, st.material_pr_kn)
        gov_ex = min(ex.compression_single_net_kn, ex.material_pr_kn)
        include_ton = self._report_include_ton()
        rows = [["Nội dung", "TTGHCĐ", "TTGHĐB"]]

        def add_kn(label: str, a: float, b: float):
            rows.append([f"{label} (kN)", _fmt(a, 2), _fmt(b, 2)])
            if include_ton:
                rows.append([f"{label} (Tấn)", _fmt(self._ton(a), 2), _fmt(self._ton(b), 2)])

        add_kn("Sức kháng vật liệu cọc Pr", st.material_pr_kn, ex.material_pr_kn)
        add_kn("Trọng lượng hữu hiệu cọc W'", -st.pile_weight_effective_kn, -ex.pile_weight_effective_kn)
        add_kn("Sức kháng bên Qs", st.qshaft_kn, ex.qshaft_kn)
        add_kn("Sức kháng mũi Qp", st.qtip_kn, ex.qtip_kn)
        if bool(getattr(res.pile_input, "include_downdrag", False)):
            rows.append(["Ma sát âm DD đã nhân γDD cộng vào tải nén (kN/cọc)", _fmt(st.downdrag_kn, 2), _fmt(ex.downdrag_kn, 2)])
        add_kn("Sức kháng theo đất nền Qr=(Qs+Qp)-W'", st.compression_single_net_kn, ex.compression_single_net_kn)
        add_kn("Sức chịu tải cọc đơn", gov_st, gov_ex)
        add_kn("Sức chịu tải nhóm cọc", st.compression_group_total_kn, ex.compression_group_total_kn)
        add_kn("Sức kháng nhổ cọc đơn", st.uplift_single_magnitude_kn, ex.uplift_single_magnitude_kn)
        add_kn("Sức kháng nhổ nhóm", st.uplift_group_total_kn, ex.uplift_group_total_kn)
        return rows

    def _report_fos_text(self, dcr: float) -> str:
        try:
            d = float(dcr or 0.0)
        except Exception:
            d = 0.0
        if d <= 0.0:
            return "∞"
        return _fmt(1.0 / d, 3)

    def _report_check_rows(self, res: CapacityResult) -> List[List[Any]]:
        c = self._report_check_values(res)
        use_dd = bool(getattr(res.pile_input, "include_downdrag", False))
        rows = [
            ["Nội dung", "TTGHCĐ", "TTGHĐB"],
            ["Nội lực đầu cọc lớn nhất Pmax (kN)", _fmt(c['pu_cd'],2), _fmt(c['pu_db'],2)],
        ]
        if use_dd:
            rows += [
                ["Ma sát âm DD đã nhân γDD cộng vào tải nén (kN/cọc)", _fmt(c['dd_cd'],2), _fmt(c['dd_db'],2)],
                ["Pmax + DD dùng kiểm SCT cọc đơn (kN)", _fmt(c['pu_cd_design'],2), _fmt(c['pu_db_design'],2)],
            ]
        rows += [
            ["Sức chịu tải cọc đơn (kN)", _fmt(c['single_cd'],2), _fmt(c['single_db'],2)],
            ["FOS nén cọc đơn", self._report_fos_text(c["dcr_single_cd"]), self._report_fos_text(c["dcr_single_db"])],
            ["Kiểm tra SCT cọc đơn", c["status_single_cd"], c["status_single_db"]],
            ["Lực nhổ lớn nhất 1 cọc (kN)", _fmt(c['uplift_cd'],2), _fmt(c['uplift_db'],2)],
            ["Sức kháng nhổ cọc đơn (kN)", _fmt(c['cap_uplift_cd'],2), _fmt(c['cap_uplift_db'],2)],
            ["FOS nhổ cọc đơn", self._report_fos_text(c["dcr_uplift_cd"]), self._report_fos_text(c["dcr_uplift_db"])],
            ["Kiểm tra nhổ cọc đơn", c["status_uplift_cd"], c["status_uplift_db"]],
            ["Lực đứng Max tại đáy bệ (kN)", _fmt(c['ncap_cd'],2), _fmt(c['ncap_db'],2)],
        ]
        if use_dd:
            rows.append(["Lực đứng Max tại đáy bệ + DD nhóm (kN)", _fmt(c['ncap_cd_design'],2), _fmt(c['ncap_db_design'],2)])
        rows += [
            ["Sức chịu tải nhóm cọc (kN)", _fmt(c['group_cd'],2), _fmt(c['group_db'],2)],
            ["FOS nén nhóm", self._report_fos_text(c["dcr_group_cd"]), self._report_fos_text(c["dcr_group_db"])],
            ["Kiểm tra SCT theo nhóm cọc", c["status_group_cd"], c["status_group_db"]],
            ["Lực nhổ max tại đáy bệ (kN)", _fmt(c['uplift_group_cd'],2), _fmt(c['uplift_group_db'],2)],
            ["Sức kháng nhổ nhóm (kN)", _fmt(c['cap_uplift_group_cd'],2), _fmt(c['cap_uplift_group_db'],2)],
            ["FOS nhổ nhóm", self._report_fos_text(c["dcr_uplift_group_cd"]), self._report_fos_text(c["dcr_uplift_group_db"])],
            ["Kiểm tra nhổ nhóm", c["status_uplift_group_cd"], c["status_uplift_group_db"]],
            ["FOS quyết định chiều dài cọc", self._report_fos_text(c["dcr_governing_cd"]), self._report_fos_text(c["dcr_governing_db"])],
        ]
        return rows

    def _collapse_sketch_layer_dicts(self, layers: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Gộp các đoạn liên tiếp cùng tên lớp/loại đất chỉ để vẽ sơ họa.

        Khi import theo từng điểm SPT, bảng địa chất có thể bị chia thành nhiều đoạn
        nhỏ cùng một lớp để giữ đủ N-SPT. Hình sơ họa chỉ nên biểu diễn một lớp địa
        chất liên tục; biểu đồ SPT vẫn dùng dữ liệu điểm gốc và không bị gộp.
        """
        out: List[Dict[str, Any]] = []
        for ly in layers or []:
            cur = dict(ly)
            name_key = _normalize_item_name(cur.get("name", ""))
            st = int(_safe_int(cur.get("soil_type", 0), 0))
            cur["soil_type"] = st
            cur["_name_key"] = name_key
            if out:
                prev = out[-1]
                same_name = str(prev.get("_name_key", "")) == name_key and bool(name_key)
                same_type = int(_safe_int(prev.get("soil_type", 0), 0)) == st
                touching = abs(_safe_float(prev.get("bottom_elev_m"), 0.0) - _safe_float(cur.get("top_elev_m"), 0.0)) <= 1e-6
                if same_name and same_type and touching:
                    prev["bottom_elev_m"] = cur.get("bottom_elev_m", prev.get("bottom_elev_m"))
                    vals = list(prev.get("_spt_values", []) or [])
                    nv = _safe_float(cur.get("n_spt", 0.0), 0.0)
                    if nv > 0:
                        vals.append(nv)
                    prev["_spt_values"] = vals
                    if vals:
                        prev["n_spt"] = sum(vals) / len(vals)
                    continue
            nv = _safe_float(cur.get("n_spt", 0.0), 0.0)
            cur["_spt_values"] = [nv] if nv > 0 else []
            out.append(cur)
        for ly in out:
            ly.pop("_name_key", None)
        return out

    def _make_geology_sketch(self, res: CapacityResult, path: str):
        try:
            from PIL import Image, ImageDraw, ImageFont
        except Exception:
            return None
        w, h = 820, 1120
        img = Image.new('RGB', (w, h), 'white')
        dr = ImageDraw.Draw(img)
        try:
            font_small = ImageFont.truetype('DejaVuSans.ttf', 44)
            font = ImageFont.truetype('DejaVuSans.ttf', 48)
            font_med = ImageFont.truetype('DejaVuSans.ttf', 54)
            font_big = ImageFont.truetype('DejaVuSans-Bold.ttf', 58)
        except Exception:
            font_small = font = font_med = font_big = ImageFont.load_default()
        colors = {0:'#F8FAFC',1:'#F4C37D',2:'#DCC8A6',3:'#C6B7A7',4:'#9B8B79',5:'#B6C9B0',6:'#E6D287'}
        inp = res.pile_input
        margin_l, margin_r, top_y, bot_y = 22, 18, 112, h - 24
        axis_x = 360
        sketch_left = 70
        sketch_right = 600
        spt_left = axis_x + 18
        spt_right = w - margin_r - 12
        dr.text((margin_l, 16), 'Sơ họa cọc - địa chất - biểu đồ SPT (Nₕₜ)', font=font_big, fill='black')
        tops = [getattr(lr, 'top_elev_m', 0.0) for lr in res.layers]
        bottoms = [getattr(lr, 'bottom_elev_m', 0.0) for lr in res.layers]
        top_elev = max(tops + [inp.ground_elev_m, inp.cap_bottom_elev_m])
        bot_elev = min(bottoms + [inp.pile_tip_elev_m])
        if abs(top_elev - bot_elev) < 1e-9:
            top_elev += 1.0
            bot_elev -= 1.0
        top_elev += 0.6
        bot_elev -= 0.6
        def y_from_elev(z: float) -> float:
            return top_y + (top_elev - z) / max(top_elev - bot_elev, 1e-9) * (bot_y - top_y)
        ground_y = y_from_elev(inp.ground_elev_m)
        dr.line([sketch_left - 10, ground_y, sketch_right + 8, ground_y], fill='#64748B', width=3)
        ground_txt = f'CĐTN: +{_fmt(inp.ground_elev_m,2)}' if inp.ground_elev_m >= 0 else f'CĐTN: {_fmt(inp.ground_elev_m,2)}'
        dr.text((sketch_left - 2, max(ground_y - 52, top_y - 8)), ground_txt, font=font_small, fill='#334155')
        raw_sketch_layers = []
        for lr in res.layers:
            raw_sketch_layers.append({
                'name': str(getattr(lr, 'name', '') or '-'),
                'top_elev_m': float(getattr(lr, 'top_elev_m', 0.0) or 0.0),
                'bottom_elev_m': float(getattr(lr, 'bottom_elev_m', 0.0) or 0.0),
                'soil_type': int(getattr(lr, 'soil_type', 0) or 0),
                'n_spt': _safe_float(getattr(lr, 'n_spt', 0.0), 0.0),
            })
        draw_layers = self._collapse_sketch_layer_dicts(raw_sketch_layers)

        centers = []
        for lr in raw_sketch_layers:
            nval = _safe_float(lr.get('n_spt', 0.0), 0.0)
            if nval > 0:
                centers.append(((y_from_elev(lr['top_elev_m']) + y_from_elev(lr['bottom_elev_m'])) / 2.0, nval))

        seen_nonzero = False
        circle_r = 46
        for lr in draw_layers:
            y1 = y_from_elev(lr.get('top_elev_m', 0.0))
            y2 = y_from_elev(lr.get('bottom_elev_m', 0.0))
            if y2 < y1:
                y1, y2 = y2, y1
            st = int(lr.get('soil_type', 0) or 0)
            dr.rectangle([sketch_left, y1, sketch_right, y2], fill=colors.get(st, '#E5E7EB'), outline='#64748B', width=2)
            zero_top = not seen_nonzero
            if st != 0:
                seen_nonzero = True
            cy = (y1 + y2) / 2
            cx = sketch_left + 62
            dr.ellipse([cx-circle_r, cy-circle_r, cx+circle_r, cy+circle_r], fill='white', outline='black', width=2)
            lab = str(lr.get('name', '') or '-')
            tw, th = self._pillow_text_size(dr, lab, font)
            dr.text((cx - tw/2, cy - th/2 - 1), lab, font=font, fill='black')
            soil_name = self._short_soil_name(st, zero_top)
            vals = list(lr.get('_spt_values', []) or [])
            sp_text = f'SPT={_fmt(vals[0],0)}' if len(vals) <= 1 and vals else ('SPT: xem biểu đồ' if len(vals) > 1 else '')
            tx = cx + circle_r + 18
            if y2 - y1 > 110:
                dr.text((tx, cy - 42), soil_name, font=font_small, fill='black')
                if sp_text:
                    dr.text((tx, cy + 6), sp_text, font=font_small, fill='black')
            else:
                txt_line = f'{soil_name}; {sp_text}' if sp_text else soil_name
                dr.text((tx, cy - 20), txt_line, font=font_small, fill='black')
            zb = float(lr.get('bottom_elev_m', 0.0) or 0.0)
            elev_txt = f'+{_fmt(zb,2)}' if zb >= 0 else _fmt(zb,2)
            ew, eh = self._pillow_text_size(dr, elev_txt, font_small)
            dr.text((sketch_right - ew - 8, y2 - eh - 3), elev_txt, font=font_small, fill='#111827')
        dr.line([axis_x, top_y - 10, axis_x, bot_y], fill='#334155', width=2)
        pile_top = y_from_elev(inp.cap_bottom_elev_m)
        pile_tip_y = y_from_elev(inp.pile_tip_elev_m)
        pile_w = 30
        if SCTCalculator._is_driven(inp):
            point_h = max(22.0, min(52.0, 0.08 * max(pile_tip_y - pile_top, 1.0)))
            point_base_y = max(pile_top + 12.0, pile_tip_y - point_h)
            dr.rectangle([axis_x - pile_w/2, pile_top, axis_x + pile_w/2, point_base_y], fill='#E5E7EB', outline='#111827', width=3)
            dr.polygon([(axis_x - pile_w/2, point_base_y), (axis_x + pile_w/2, point_base_y), (axis_x, pile_tip_y)], fill='#E5E7EB', outline='#111827')
        else:
            dr.rectangle([axis_x - pile_w/2, pile_top, axis_x + pile_w/2, pile_tip_y], fill='#E5E7EB', outline='#111827', width=3)
        dr.text((axis_x - 72, max(pile_top - 60, top_y - 2)), 'CỌC', font=font_med, fill='#111827')
        dr.text((axis_x - 124, pile_tip_y + 4), f'D={_fmt(inp.diameter_m,2)} m', font=font_small, fill='#374151')
        pile_top_txt = f'CĐ đầu cọc: +{_fmt(inp.cap_bottom_elev_m,2)}' if inp.cap_bottom_elev_m >= 0 else f'CĐ đầu cọc: {_fmt(inp.cap_bottom_elev_m,2)}'
        tip_txt = f'CĐ mũi: +{_fmt(inp.pile_tip_elev_m,2)}' if inp.pile_tip_elev_m >= 0 else f'CĐ mũi: {_fmt(inp.pile_tip_elev_m,2)}'
        dr.text((sketch_left, pile_top + 4), pile_top_txt, font=font_small, fill='#111827')
        dr.text((sketch_left, pile_tip_y - 52), tip_txt, font=font_small, fill='#111827')
        dr.text((spt_left + 10, top_y - 66), 'Biểu đồ SPT', font=font_med, fill='black')
        max_n = max([pt[1] for pt in centers] + [10.0])
        spt_max = max(10, int((max_n + 9) // 10 * 10))
        usable = max(60, spt_right - spt_left)
        for gv in range(0, spt_max + 1, 10):
            x = spt_left + usable * gv / max(spt_max, 1)
            dr.line([x, top_y, x, bot_y], fill='#E5E7EB', width=1)
            label = str(gv)
            tw, th = self._pillow_text_size(dr, label, font_small)
            dr.text((x - tw/2, top_y - th - 8), label, font=font_small, fill='#6B7280')
        pts = []
        for cy, n1 in centers:
            x = spt_left + usable * min(max(n1, 0.0), spt_max) / max(spt_max, 1)
            pts.append((x, cy))
        for p1, p2 in zip(pts[:-1], pts[1:]):
            dr.line([p1[0], p1[1], p2[0], p2[1]], fill='black', width=2)
        for (x, cy), (_, n1) in zip(pts, centers):
            dr.ellipse([x - 8, cy - 8, x + 8, cy + 8], fill='black', outline='black')
            dr.text((x + 10, cy - 20), _fmt(n1,0), font=font_small, fill='black')
        img.save(path)
        return path
    def _make_simple_geology_sketch(self, res: CapacityResult, path: str):
        """Vẽ sơ họa địa chất nhỏ để chèn vào bảng đầu vào của báo cáo.

        V1.0.5:
        - Vẽ toàn bộ các lớp địa chất theo số liệu khai báo, không cắt ở mũi cọc.
        - Hình cọc được làm lại theo phong cách kỹ thuật: thân cọc rõ ràng, chân cọc loe nhẹ/tapered.
        - Giữ biểu đồ SPT dạng đường nối các điểm và màu các lớp phân biệt tốt hơn.
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
        except Exception:
            return None

        w, h = 560, 1160
        img = Image.new('RGB', (w, h), 'white')
        dr = ImageDraw.Draw(img)

        def _load_font(size: int, bold: bool = False):
            candidates = []
            if bold:
                candidates += [
                    r"C:\Windows\Fonts\timesbd.ttf",
                    r"C:\Windows\Fonts\arialbd.ttf",
                    "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf",
                    "/usr/share/fonts/truetype/liberation/LiberationSerif-Bold.ttf",
                    "/usr/share/fonts/truetype/tinos/Tinos-Bold.ttf",
                    "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
                ]
            else:
                candidates += [
                    r"C:\Windows\Fonts\times.ttf",
                    r"C:\Windows\Fonts\arial.ttf",
                    "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
                    "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
                    "/usr/share/fonts/truetype/tinos/Tinos-Regular.ttf",
                    "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
                ]
            for fp in candidates:
                try:
                    if os.path.exists(fp):
                        return ImageFont.truetype(fp, size)
                except Exception:
                    pass
            try:
                return ImageFont.truetype('times.ttf' if not bold else 'timesbd.ttf', size)
            except Exception:
                return ImageFont.load_default()

        font_label = _load_font(28, bold=True)
        font_small = _load_font(23)
        font_elev = _load_font(26)
        font_elev_bold = _load_font(28, bold=True)
        font_spt_title = _load_font(26, bold=True)
        font_spt_value = _load_font(23, bold=True)
        font_dim = _load_font(21)

        layer_palette = [
            '#B8E0B6', '#B7DDE8', '#EDC08C', '#E7EA8A', '#CFCFCF', '#D8C3E6',
            '#F7D9C4', '#BEE3DB', '#D0E1F9', '#F6D6E6', '#D7F0B3', '#E6D8B8'
        ]
        type_edge = {
            0: '#94A3B8', 1: '#3B7A57', 2: '#2B6CB0', 3: '#6D28D9',
            4: '#57534E', 5: '#166534', 6: '#8B5A2B'
        }

        inp = res.pile_input

        # Dùng toàn lỗ khoan địa chất gốc để hình không bị cắt ở mũi cọc.
        raw_layers = SCTCalculator._sorted_layers(inp)
        sketch_layers = []
        current_top = float(inp.ground_elev_m)
        for ly in raw_layers:
            bottom = float(getattr(ly, 'bottom_elev_m', current_top) or current_top)
            top = current_top
            if bottom > top:
                bottom = top
            sketch_layers.append({
                'name': str(getattr(ly, 'name', '') or f'L{len(sketch_layers) + 1}'),
                'top_elev_m': top,
                'bottom_elev_m': bottom,
                'soil_type': int(getattr(ly, 'soil_type', 0) or 0),
                'n_spt': _safe_float(getattr(ly, 'n_spt', 0.0), 0.0),
            })
            current_top = bottom

        if not sketch_layers:
            # fallback từ res.layers nếu thiếu địa chất gốc
            for i, lr in enumerate(res.layers):
                sketch_layers.append({
                    'name': str(getattr(lr, 'name', '') or f'L{i + 1}'),
                    'top_elev_m': float(getattr(lr, 'top_elev_m', inp.ground_elev_m) or inp.ground_elev_m),
                    'bottom_elev_m': float(getattr(lr, 'bottom_elev_m', inp.pile_tip_elev_m) or inp.pile_tip_elev_m),
                    'soil_type': int(getattr(lr, 'soil_type', 0) or 0),
                    'n_spt': _safe_float(getattr(lr, 'n_spt', 0.0), 0.0),
                })

        # Giữ dữ liệu SPT chi tiết để vẽ biểu đồ, nhưng gộp các đoạn cùng tên lớp khi tô sơ họa địa tầng.
        spt_layers = [dict(x) for x in sketch_layers]
        sketch_layers = self._collapse_sketch_layer_dicts(sketch_layers)

        tops = [float(x['top_elev_m']) for x in sketch_layers] if sketch_layers else [inp.ground_elev_m]
        bottoms = [float(x['bottom_elev_m']) for x in sketch_layers] if sketch_layers else [inp.pile_tip_elev_m]
        top_elev = max(tops + [inp.ground_elev_m, inp.cap_bottom_elev_m])
        bot_elev = min(bottoms + [inp.pile_tip_elev_m])
        if abs(top_elev - bot_elev) < 1e-9:
            top_elev += 1.0
            bot_elev -= 1.0
        top_elev += 0.65
        bot_elev -= 0.65

        x_left, x_right = 20, w - 18
        top_y, bot_y = 88, h - 52
        geo_left, geo_right = 66, 330
        pile_x = 212
        label_x = 118
        shaft_half = 22
        flare_half = 32
        spt_left, spt_right = 360, w - 18

        def y_from_elev(z: float) -> float:
            return top_y + (top_elev - float(z)) / max(top_elev - bot_elev, 1e-9) * (bot_y - top_y)

        def elev_text(z: Any, prefix: str = 'CĐ') -> str:
            val = _safe_float(z, 0.0)
            body = ('+' if val >= 0 else '') + _fmt(val, 3)
            return f'{prefix}: {body}'

        def elev_num_text(z: Any, ndigits: int = 2) -> str:
            val = _safe_float(z, 0.0)
            return ('+' if val >= 0 else '') + _fmt(val, ndigits)

        def draw_text_box(x: float, y: float, text: str, font_obj, fill='#111827', anchor='lt', outline='#CBD5E1'):
            tw, th = self._pillow_text_size(dr, text, font_obj)
            if anchor == 'rt':
                x -= tw
            elif anchor == 'ct':
                x -= tw / 2
            padx, pady = 6, 4
            try:
                dr.rounded_rectangle([x - padx, y - pady, x + tw + padx, y + th + pady], radius=6, fill='white', outline=outline, width=1)
            except Exception:
                dr.rectangle([x - padx, y - pady, x + tw + padx, y + th + pady], fill='white', outline=outline, width=1)
            dr.text((x, y), text, font=font_obj, fill=fill)
            return tw, th

        def _text_box_rect(x: float, y: float, text: str, font_obj, anchor='lt', pad_x: int = 6, pad_y: int = 4):
            tw, th = self._pillow_text_size(dr, text, font_obj)
            if anchor == 'rt':
                x -= tw
            elif anchor == 'ct':
                x -= tw / 2
            return [x - pad_x, y - pad_y, x + tw + pad_x, y + th + pad_y], x, y, tw, th

        def _rect_overlap(a, b, gap: float = 4.0) -> bool:
            return not (a[2] + gap <= b[0] or b[2] + gap <= a[0] or a[3] + gap <= b[1] or b[3] + gap <= a[1])

        occupied_boxes = []

        def place_text_box(preferred_positions, text: str, font_obj, fill='#111827', outline='#CBD5E1'):
            candidates = []
            for item in preferred_positions:
                if len(item) == 2:
                    candidates.append((item[0], item[1], 'lt'))
                else:
                    candidates.append((item[0], item[1], item[2]))
            # thử các vị trí ưu tiên trước
            for px, py, anc in candidates:
                rect, dx, dy, _tw, _th = _text_box_rect(px, py, text, font_obj, anchor=anc)
                if any(_rect_overlap(rect, r) for r in occupied_boxes):
                    continue
                draw_text_box(px, py, text, font_obj, fill=fill, anchor=anc, outline=outline)
                occupied_boxes.append(rect)
                return rect
            # nếu vẫn đụng nhau, tự dò xuống dưới/qua phải
            if candidates:
                px, py, anc = candidates[0]
            else:
                px, py, anc = (x_left + 4, top_y + 4, 'lt')
            for dy_try in [0, 16, 32, 48, 64, -16, 80]:
                for dx_try in [0, 12, 24, 36, 48, -12, -24]:
                    rect, _dx, _dy, _tw, _th = _text_box_rect(px + dx_try, py + dy_try, text, font_obj, anchor=anc)
                    if any(_rect_overlap(rect, r) for r in occupied_boxes):
                        continue
                    draw_text_box(px + dx_try, py + dy_try, text, font_obj, fill=fill, anchor=anc, outline=outline)
                    occupied_boxes.append(rect)
                    return rect
            # bất đắc dĩ thì vẽ luôn vị trí đầu
            rect, _dx, _dy, _tw, _th = _text_box_rect(px, py, text, font_obj, anchor=anc)
            draw_text_box(px, py, text, font_obj, fill=fill, anchor=anc, outline=outline)
            occupied_boxes.append(rect)
            return rect

        def add_fixed_text_box(x: float, y: float, text: str, font_obj, fill='#111827', anchor='lt', outline='#CBD5E1'):
            rect, _dx, _dy, _tw, _th = _text_box_rect(x, y, text, font_obj, anchor=anchor)
            draw_text_box(x, y, text, font_obj, fill=fill, anchor=anchor, outline=outline)
            occupied_boxes.append(rect)
            return rect

        def draw_layer_name(center_x: float, center_y: float, text: str, edge: str, layer_top: float, layer_bottom: float):
            tw, th = self._pillow_text_size(dr, text, font_label)
            radius = max(24, min(34, int(max(tw, th) / 2 + 12)))
            base_cy = max(layer_top + radius + 4, min(center_y, layer_bottom - radius - 4))
            # Ưu tiên dịch tên lớp sang phải để giữ nguyên vị trí các cao độ.
            candidates = [
                (center_x + 0, base_cy),
                (center_x + 20, base_cy),
                (center_x + 34, base_cy),
                (center_x + 20, base_cy - 14),
                (center_x + 20, base_cy + 14),
                (center_x + 38, base_cy - 12),
                (center_x + 38, base_cy + 12),
                (center_x - 10, base_cy),
            ]
            best = None
            for cx, cy in candidates:
                cy = max(layer_top + radius + 4, min(cy, layer_bottom - radius - 4))
                rect = [cx - radius, cy - radius, cx + radius, cy + radius]
                if rect[0] < geo_left + 4 or rect[2] > pile_x - shaft_half - 8:
                    continue
                if any(_rect_overlap(rect, r, gap=2.0) for r in occupied_boxes):
                    continue
                best = (rect, cx, cy)
                break
            if best is None:
                cx, cy = candidates[-1]
                cy = max(layer_top + radius + 4, min(cy, layer_bottom - radius - 4))
                rect = [cx - radius, cy - radius, cx + radius, cy + radius]
            else:
                rect, cx, cy = best
            dr.ellipse(rect, fill='white', outline=edge, width=3)
            dr.text((cx - tw / 2, cy - th / 2 - 1), text, font=font_label, fill='#111827')
            occupied_boxes.append(rect)
            return rect

        def layer_fill_color(idx: int, soil_type: int) -> str:
            if int(soil_type or 0) == 0:
                return '#F8FAFC'
            return layer_palette[idx % len(layer_palette)]

        # Khung và nền.
        dr.rectangle([geo_left, top_y, geo_right, bot_y], outline='#64748B', width=2)
        dr.rectangle([spt_left, top_y, spt_right, bot_y], outline='#94A3B8', width=2)

        point_data = []
        label_boxes = []
        layer_label_data = []
        for idx, ly in enumerate(sketch_layers):
            y1 = y_from_elev(ly['top_elev_m'])
            y2 = y_from_elev(ly['bottom_elev_m'])
            if y2 < y1:
                y1, y2 = y2, y1
            st = int(ly['soil_type'] or 0)
            fill = layer_fill_color(idx, st)
            edge = type_edge.get(st, '#64748B')
            dr.rectangle([geo_left, y1, geo_right, y2], fill=fill, outline=edge, width=2)
            cy = (y1 + y2) / 2.0
            label_boxes.append((float(y2), elev_num_text(ly['bottom_elev_m'], 2)))
            layer_label_data.append((ly['name'], edge, cy, y1, y2))

        for idx, ly in enumerate(spt_layers):
            nval = _safe_float(ly.get('n_spt', 0.0), 0.0)
            if nval > 0:
                cy = (y_from_elev(ly['top_elev_m']) + y_from_elev(ly['bottom_elev_m'])) / 2.0
                point_data.append((cy, nval, idx))

        ground_y = y_from_elev(inp.ground_elev_m)
        dr.line([x_left, ground_y, x_right, ground_y], fill='#334155', width=4)
        gy_label_y = max(8, min(ground_y - 40, top_y - 36))
        add_fixed_text_box(x_left + 6, gy_label_y, elev_text(inp.ground_elev_m, 'CĐ mặt đất'), font_elev_bold, fill='#0F172A')

        # Hình cọc: cọc khoan nhồi giữ mũi loe; cọc đóng/ép dùng mũi tam giác nhỏ.
        pile_top_y = y_from_elev(inp.cap_bottom_elev_m)
        pile_tip_y = y_from_elev(inp.pile_tip_elev_m)
        if pile_tip_y < pile_top_y:
            pile_top_y, pile_tip_y = pile_tip_y, pile_top_y
        pile_top_y = max(top_y, min(pile_top_y, bot_y))
        pile_tip_y = max(top_y, min(pile_tip_y, bot_y))
        is_driven_or_pressed = SCTCalculator._is_driven(inp)
        if is_driven_or_pressed:
            point_h = max(26.0, min(62.0, 0.08 * max(pile_tip_y - pile_top_y, 1.0)))
            point_base_y = max(pile_top_y + 18.0, pile_tip_y - point_h)
            pile_outer_half = shaft_half
            dr.rectangle([pile_x - shaft_half, pile_top_y, pile_x + shaft_half, point_base_y], fill='#F2F2F2', outline='#111827', width=2)
            dr.polygon([(pile_x - shaft_half, point_base_y), (pile_x + shaft_half, point_base_y), (pile_x, pile_tip_y)], fill='#F2F2F2', outline='#111827')
            for xi in range(int(pile_x - shaft_half) + 1, int(pile_x + shaft_half)):
                rel = abs((xi - pile_x) / max(shaft_half, 1))
                shade = int(246 - 58 * rel)
                shade = max(184, min(248, shade))
                dr.line([xi, pile_top_y + 1, xi, point_base_y - 1], fill=(shade, shade, shade), width=1)
            dr.line([pile_x - shaft_half, pile_top_y, pile_x - shaft_half, point_base_y], fill='#A3A3A3', width=2)
            dr.line([pile_x + shaft_half, pile_top_y, pile_x + shaft_half, point_base_y], fill='#6B7280', width=2)
            dr.line([pile_x - shaft_half, point_base_y, pile_x, pile_tip_y], fill='#A3A3A3', width=2)
            dr.line([pile_x + shaft_half, point_base_y, pile_x, pile_tip_y], fill='#6B7280', width=2)
        else:
            taper_h = max(28.0, min(70.0, 0.09 * (pile_tip_y - pile_top_y)))
            taper_y = max(pile_top_y + 16.0, pile_tip_y - taper_h)
            pile_outer_half = flare_half
            pts = [
                (pile_x - shaft_half, pile_top_y),
                (pile_x + shaft_half, pile_top_y),
                (pile_x + shaft_half, taper_y),
                (pile_x + flare_half, pile_tip_y),
                (pile_x - flare_half, pile_tip_y),
                (pile_x - shaft_half, taper_y),
            ]
            dr.polygon(pts, fill='#F2F2F2', outline='#111827')
            # Tạo cảm giác 3D nhẹ bằng các dải xám dọc.
            left = int(pile_x - flare_half)
            right = int(pile_x + flare_half)
            for xi in range(left + 1, right):
                rel = abs((xi - pile_x) / max(flare_half, 1))
                shade = int(246 - 70 * rel)
                shade = max(176, min(248, shade))
                dr.line([xi, pile_top_y + 1, xi, taper_y], fill=(shade, shade, shade), width=1)
                if abs(xi - pile_x) <= flare_half:
                    ratio = abs(xi - pile_x) / max(flare_half, 1)
                    y_start = taper_y + (pile_tip_y - taper_y) * max(0.0, (ratio - shaft_half / max(flare_half,1)) / max(1 - shaft_half / max(flare_half,1), 1e-6)) if abs(xi - pile_x) > shaft_half else taper_y
                    dr.line([xi, y_start, xi, pile_tip_y - 1], fill=(shade, shade, shade), width=1)
            dr.line([pile_x - shaft_half, pile_top_y, pile_x - shaft_half, taper_y], fill='#A3A3A3', width=2)
            dr.line([pile_x + shaft_half, pile_top_y, pile_x + shaft_half, taper_y], fill='#6B7280', width=2)
            dr.line([pile_x - shaft_half, taper_y, pile_x - flare_half, pile_tip_y], fill='#A3A3A3', width=2)
            dr.line([pile_x + shaft_half, taper_y, pile_x + flare_half, pile_tip_y], fill='#6B7280', width=2)
        head_y = max(top_y + 4, pile_top_y - 34)
        add_fixed_text_box(pile_x, head_y, elev_num_text(inp.cap_bottom_elev_m, 2), font_elev, fill='#374151', anchor='ct')
        toe_y = min(bot_y - 34, pile_tip_y - 42)
        add_fixed_text_box(pile_x + pile_outer_half + 12, toe_y, elev_text(inp.pile_tip_elev_m, 'CĐ mũi'), font_elev, fill='#374151')

        # Kích thước tham khảo: mũi tên L đặt gần cọc hơn; L là chiều dài thực tế của cọc (m).
        dim_x = pile_x + pile_outer_half + 10
        dr.line([dim_x, pile_top_y, dim_x, pile_tip_y], fill='black', width=2)
        dr.polygon([(dim_x, pile_top_y), (dim_x - 5, pile_top_y + 10), (dim_x + 5, pile_top_y + 10)], fill='black')
        dr.polygon([(dim_x, pile_tip_y), (dim_x - 5, pile_tip_y - 10), (dim_x + 5, pile_tip_y - 10)], fill='black')
        pile_len_txt = f"L = {_fmt(abs(inp.cap_bottom_elev_m - inp.pile_tip_elev_m), 2)} m"
        add_fixed_text_box(dim_x + 10, (pile_top_y + pile_tip_y) / 2 - 14, pile_len_txt, font_dim, fill='black')
        dim_y = min(bot_y - 10, pile_tip_y + 42)
        dr.line([pile_x - pile_outer_half, dim_y, pile_x + pile_outer_half, dim_y], fill='black', width=2)
        dr.polygon([(pile_x - pile_outer_half, dim_y), (pile_x - pile_outer_half + 10, dim_y - 5), (pile_x - pile_outer_half + 10, dim_y + 5)], fill='black')
        dr.polygon([(pile_x + pile_outer_half, dim_y), (pile_x + pile_outer_half - 10, dim_y - 5), (pile_x + pile_outer_half - 10, dim_y + 5)], fill='black')
        tw, th = self._pillow_text_size(dr, 'Dₚ', font_dim)
        dr.text((pile_x - tw / 2, dim_y - th - 6), 'Dₚ', font=font_dim, fill='black')

        # Cao độ đáy lớp: sắp xếp tránh đè nhau.
        label_boxes_sorted = []
        prev_y = -10**9
        min_gap = 34
        for yb, txt in sorted(label_boxes, key=lambda t: t[0]):
            y_lbl = max(top_y + 6, min(yb - 24, bot_y - 28))
            if y_lbl - prev_y < min_gap:
                y_lbl = prev_y + min_gap
            y_lbl = min(y_lbl, bot_y - 28)
            label_boxes_sorted.append((y_lbl, txt))
            prev_y = y_lbl
        for _y_lbl, _txt in label_boxes_sorted:
            add_fixed_text_box(x_left + 4, _y_lbl, _txt, font_elev, fill='#374151')

        # Sau khi khóa toàn bộ cao độ, mới dịch tên lớp để tránh chồng lấn và vẫn giữ đúng bản chất cao độ.
        pile_block_rect = [pile_x - pile_outer_half - 4, pile_top_y - 2, pile_x + pile_outer_half + 4, pile_tip_y + 2]
        occupied_boxes.append(pile_block_rect)
        for lab, edge, cy, y1, y2 in layer_label_data:
            draw_layer_name(label_x, cy, lab, edge, y1, y2)

        # Biểu đồ SPT theo toàn bộ lớp địa chất.
        dr.text((spt_left + 4, top_y - 36), 'SPT', font=font_spt_title, fill='#111827')
        max_n = max([n for (_cy, n, _idx) in point_data] + [10.0])
        spt_max = max(10, int((max_n + 9) // 10 * 10))
        usable = max(60, spt_right - spt_left - 12)
        for gv in range(0, spt_max + 1, 10):
            x = spt_left + usable * gv / max(spt_max, 1)
            dr.line([x, top_y, x, bot_y], fill='#E5E7EB', width=1)
            lab = str(gv)
            tw, th = self._pillow_text_size(dr, lab, font_small)
            dr.text((x - tw / 2, top_y + 4), lab, font=font_small, fill='#6B7280')
        spt_pts = []
        for cy, n1, idx in point_data:
            x = spt_left + usable * min(max(n1, 0.0), spt_max) / max(spt_max, 1)
            spt_pts.append((x, cy, n1, idx))
        for p1, p2 in zip(spt_pts[:-1], spt_pts[1:]):
            dr.line([p1[0], p1[1], p2[0], p2[1]], fill='#1D4ED8', width=4)
        for x, cy, n1, idx in spt_pts:
            dr.ellipse([x - 8, cy - 8, x + 8, cy + 8], fill='white', outline='#111827', width=2)
            dr.ellipse([x - 5, cy - 5, x + 5, cy + 5], fill='#1D4ED8', outline='#1D4ED8')
            val_txt = _fmt(n1, 0)
            tw, th = self._pillow_text_size(dr, val_txt, font_spt_value)
            tx = x + 12
            anchor = 'lt'
            if tx + tw + 12 > spt_right:
                tx = x - 12
                anchor = 'rt'
            draw_text_box(tx, cy - th / 2 - 2, val_txt, font_spt_value, fill='#111827', anchor=anchor, outline='#93C5FD')

        dr.rectangle([x_left, top_y, x_right, bot_y], outline='#475569', width=2)
        img.save(path)
        return path

    def _pillow_text_size(self, draw, text: str, font) -> Tuple[int, int]:
        try:
            box = draw.textbbox((0, 0), str(text), font=font)
            return max(1, box[2] - box[0]), max(1, box[3] - box[1])
        except Exception:
            try:
                return draw.textsize(str(text), font=font)
            except Exception:
                return (8 * len(str(text)), 12)

    def _report_page_setup(self) -> Tuple[str, str]:
        size = str(getattr(self, "_report_page_size", "A4") or "A4").upper()
        orient = str(getattr(self, "_report_orientation", "Dọc") or "Dọc")
        return size, orient

    def _report_docx_page_cm(self) -> Tuple[float, float]:
        size, orient = self._report_page_setup()
        if size == "A3":
            w, h = 29.7, 42.0
        else:
            w, h = 21.0, 29.7
        if "ngang" in orient.lower() or "land" in orient.lower():
            w, h = h, w
        return w, h

    def _report_pdf_pagesize(self):
        try:
            from reportlab.lib.pagesizes import A4, A3, landscape
        except Exception:
            return None
        size, orient = self._report_page_setup()
        page = A3 if size == "A3" else A4
        if "ngang" in orient.lower() or "land" in orient.lower():
            page = landscape(page)
        return page

    def _ask_report_export_options(self, results: List[CapacityResult]) -> Optional[Dict[str, Any]]:
        dlg = tk.Toplevel(self.root)
        dlg.title('Tùy chọn xuất báo cáo TS-CAP')
        dlg.transient(self.root)
        dlg.grab_set()
        try:
            dlg.resizable(False, False)
        except Exception:
            pass
        frame = ttk.Frame(dlg, padding=14)
        frame.pack(fill=tk.BOTH, expand=True)
        fmt_var = tk.StringVar(value="DOC")
        scope_var = tk.StringVar(value="Tất cả hạng mục")
        page_var = tk.StringVar(value=str(getattr(self, "_report_page_size", "A4")))
        orient_var = tk.StringVar(value=str(getattr(self, "_report_orientation", "Dọc")))
        unit_var = tk.StringVar(value=str(getattr(self, "_report_result_units", "kN")))
        suffix_var = tk.StringVar(value="_CAP")
        naming_var = tk.StringVar(value="Theo hạng mục")
        custom_var = tk.StringVar(value="")

        # Danh sách hạng mục đã có kết quả, giữ đúng thứ tự tính toán và bỏ trùng tên.
        item_names: List[str] = []
        seen_items = set()
        for r in results:
            name = str(getattr(r.pile_input, "item", "") or "HANG_MUC").strip() or "HANG_MUC"
            key = _normalize_item_name(name)
            if key not in seen_items:
                item_names.append(name)
                seen_items.add(key)

        ttk.Label(frame, text="Định dạng xuất").grid(row=0, column=0, sticky="w", pady=4)
        ttk.Combobox(frame, textvariable=fmt_var, values=["DOC", "PDF", "BOTH"], state="readonly", width=20).grid(row=0, column=1, sticky="w", pady=4)
        ttk.Label(frame, text="Phạm vi in").grid(row=1, column=0, sticky="w", pady=4)
        scope_cb = ttk.Combobox(frame, textvariable=scope_var, values=["Tất cả hạng mục", "Hạng mục đang chọn", "Chọn hạng mục"], state="readonly", width=24)
        scope_cb.grid(row=1, column=1, sticky="w", pady=4)

        ttk.Label(frame, text="Chọn hạng mục").grid(row=2, column=0, sticky="nw", pady=4)
        select_box = ttk.Frame(frame)
        select_box.grid(row=2, column=1, sticky="w", pady=4)
        item_list = tk.Listbox(select_box, selectmode=tk.EXTENDED, width=38, height=min(8, max(3, len(item_names))))
        item_scroll = ttk.Scrollbar(select_box, orient="vertical", command=item_list.yview)
        item_list.configure(yscrollcommand=item_scroll.set)
        item_list.grid(row=0, column=0, sticky="nsew")
        item_scroll.grid(row=0, column=1, sticky="ns")
        for name in item_names:
            item_list.insert(tk.END, name)
        if item_names:
            item_list.selection_set(0, tk.END)
        select_btns = ttk.Frame(select_box)
        select_btns.grid(row=1, column=0, columnspan=2, sticky="w", pady=(4,0))
        ttk.Button(select_btns, text="Chọn tất cả", command=lambda: item_list.selection_set(0, tk.END)).pack(side=tk.LEFT, padx=(0,4))
        ttk.Button(select_btns, text="Bỏ chọn", command=lambda: item_list.selection_clear(0, tk.END)).pack(side=tk.LEFT, padx=4)

        ttk.Label(frame, text="Khổ giấy").grid(row=3, column=0, sticky="w", pady=4)
        ttk.Combobox(frame, textvariable=page_var, values=["A4", "A3"], state="readonly", width=20).grid(row=3, column=1, sticky="w", pady=4)
        ttk.Label(frame, text="Hướng giấy").grid(row=4, column=0, sticky="w", pady=4)
        ttk.Combobox(frame, textvariable=orient_var, values=["Dọc", "Ngang"], state="readonly", width=20).grid(row=4, column=1, sticky="w", pady=4)
        ttk.Label(frame, text="Đơn vị kết quả").grid(row=5, column=0, sticky="w", pady=4)
        ttk.Combobox(frame, textvariable=unit_var, values=["kN", "kN và Tấn"], state="readonly", width=20).grid(row=5, column=1, sticky="w", pady=4)
        ttk.Label(frame, text="Cách đặt tên file").grid(row=6, column=0, sticky="nw", pady=4)
        name_box = ttk.Frame(frame)
        name_box.grid(row=6, column=1, sticky="w", pady=4)
        ttk.Radiobutton(name_box, text="Theo hạng mục", variable=naming_var, value="Theo hạng mục").pack(anchor="w")
        ttk.Radiobutton(name_box, text="Tên tự đặt", variable=naming_var, value="Tên tự đặt").pack(anchor="w")
        ttk.Label(frame, text="Tên tự đặt").grid(row=7, column=0, sticky="w", pady=4)
        custom_entry = ttk.Entry(frame, textvariable=custom_var, width=26)
        custom_entry.grid(row=7, column=1, sticky="w", pady=4)
        ttk.Label(frame, text="Hậu tố tên file").grid(row=8, column=0, sticky="w", pady=4)
        ttk.Entry(frame, textvariable=suffix_var, width=26).grid(row=8, column=1, sticky="w", pady=4)
        note = ttk.Label(frame, text="Có thể xuất tất cả, hạng mục đang chọn, hoặc chọn nhiều hạng mục trong danh sách. Mặc định tên file theo dạng: <Tên hạng mục><hậu tố>. Nếu chọn Tên tự đặt: 1 hạng mục -> <Tên tự đặt><hậu tố>; nhiều hạng mục -> <Tên tự đặt>_<Tên hạng mục><hậu tố>.", wraplength=560)
        note.grid(row=9, column=0, columnspan=2, sticky="w", pady=(10,4))

        def toggle_custom(*_):
            st = "normal" if naming_var.get() == "Tên tự đặt" else "disabled"
            try:
                custom_entry.configure(state=st)
            except Exception:
                pass
        def toggle_scope(*_):
            enabled = scope_var.get() == "Chọn hạng mục"
            try:
                item_list.configure(state=("normal" if enabled else "disabled"))
            except Exception:
                pass
        naming_var.trace_add("write", toggle_custom)
        scope_var.trace_add("write", toggle_scope)
        toggle_custom()
        toggle_scope()

        result: Dict[str, Any] = {}
        def ok():
            if naming_var.get() == "Tên tự đặt" and not str(custom_var.get()).strip():
                messagebox.showwarning("Tên file", "Hãy nhập Tên tự đặt hoặc chuyển về chế độ Theo hạng mục.", parent=dlg)
                return
            selected_item_names: List[str] = []
            if scope_var.get() == "Chọn hạng mục":
                selected_item_names = [item_names[i] for i in item_list.curselection()]
                if not selected_item_names:
                    messagebox.showwarning("Chọn hạng mục", "Hãy chọn ít nhất một hạng mục để xuất báo cáo.", parent=dlg)
                    return
            result.update({
                "format": fmt_var.get(),
                "scope": scope_var.get(),
                "selected_item_names": selected_item_names,
                "selected_item_keys": [_normalize_item_name(x) for x in selected_item_names],
                "page_size": page_var.get(),
                "orientation": orient_var.get(),
                "result_units": unit_var.get(),
                "suffix": suffix_var.get() or "_CAP",
                "naming_mode": naming_var.get(),
                "custom_name": custom_var.get().strip(),
            })
            dlg.destroy()
        def cancel():
            result.clear(); dlg.destroy()
        btns = ttk.Frame(frame)
        btns.grid(row=10, column=0, columnspan=2, sticky="e", pady=16)
        ttk.Button(btns, text="OK", command=ok).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Hủy", command=cancel).pack(side=tk.LEFT, padx=4)
        self.root.wait_window(dlg)
        return result or None

    def _write_item_docx_report(self, res: CapacityResult, path: str):
        try:
            from docx import Document
            from docx.shared import Cm, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_SECTION
        except Exception as exc:
            raise RuntimeError(f"Thiếu thư viện python-docx. Cài bằng: pip install python-docx. Chi tiết: {exc}")
        doc = Document()
        sec = doc.sections[0]
        try:
            pw, ph = self._report_docx_page_cm()
            sec.page_width = Cm(pw)
            sec.page_height = Cm(ph)
        except Exception:
            pass
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.4)
        sec.left_margin = Cm(1.35)
        sec.right_margin = Cm(1.35)
        # Font mặc định
        try:
            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(10)
        except Exception:
            pass
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self._capacity_report_title(res))
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0, 0, 200)
        self._docx_heading(doc, "1. SỐ LIỆU ĐẦU VÀO", 1)
        
        from docx.shared import Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def _set_cell_margins(cell, top=60, start=60, bottom=60, end=60):
            """Giảm margin ô layout để hình và bảng không chen lấn nhau.

            Đơn vị là twentieths of a point theo XML của Word.
            """
            try:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_mar = tc_pr.first_child_found_in("w:tcMar")
                if tc_mar is None:
                    tc_mar = OxmlElement('w:tcMar')
                    tc_pr.append(tc_mar)
                for m_name, m_val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
                    node = tc_mar.find(qn(f"w:{m_name}"))
                    if node is None:
                        node = OxmlElement(f"w:{m_name}")
                        tc_mar.append(node)
                    node.set(qn('w:w'), str(int(m_val)))
                    node.set(qn('w:type'), 'dxa')
            except Exception:
                pass

        layout_table = doc.add_table(rows=1, cols=2)
        try:
            layout_table.autofit = False
            layout_table.allow_autofit = False
            layout_table.style = "Table Grid"
        except Exception:
            pass

        # A4 dọc còn 18.1 cm sau margin. Cột trái đủ rộng cho 4 cột thông số;
        # cột phải chỉ giữ hình sơ họa, hình chèn theo WIDTH để không lấn bảng.
        try:
            page_w_cm, _page_h_cm = self._report_docx_page_cm()
            content_w_cm = max(page_w_cm - 2.70, 12.0)
        except Exception:
            content_w_cm = 18.1
        right_cm = min(5.30, max(4.85, content_w_cm * 0.29))
        left_cm = max(10.80, content_w_cm - right_cm)
        img_w_cm = max(4.40, min(right_cm - 0.28, 5.00))

        cell_left = layout_table.cell(0, 0)
        cell_right = layout_table.cell(0, 1)
        try:
            layout_table.columns[0].width = Cm(left_cm)
            layout_table.columns[1].width = Cm(right_cm)
            for c in layout_table.columns[0].cells:
                c.width = Cm(left_cm)
            for c in layout_table.columns[1].cells:
                c.width = Cm(right_cm)
            cell_left.width = Cm(left_cm)
            cell_right.width = Cm(right_cm)
        except Exception:
            pass
        _set_cell_margins(cell_left, start=40, end=40)
        _set_cell_margins(cell_right, start=40, end=40)

        input_rows_data = self._report_input_rows(res)
        self._docx_add_kv_table(cell_left, input_rows_data, table_align="left")

        img_path = self._unique_path(os.path.dirname(path), f"sketch_{res.pile_input.item}", ".png")
        self._make_simple_geology_sketch(res, img_path)

        if os.path.exists(img_path):
            p_img = cell_right.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_img.paragraph_format.space_before = Pt(0)
                p_img.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            run = p_img.add_run()
            try:
                from PIL import Image
                with Image.open(img_path) as _im:
                    im_w_px, im_h_px = _im.size
                # Fit theo chiều đứng của bảng đầu vào trước; sau đó khống chế thêm theo bề rộng cột phải.
                row_count = 1 + len(input_rows_data)  # gồm 1 dòng header của bảng key-value
                target_h_cm = min(11.60, max(8.60, row_count * 0.62))
                max_w_cm = max(4.35, min(right_cm - 0.18, 5.05))
                width_for_target_h = target_h_cm * (im_w_px / max(im_h_px, 1))
                if width_for_target_h <= max_w_cm:
                    run.add_picture(img_path, height=Cm(target_h_cm))
                else:
                    run.add_picture(img_path, width=Cm(max_w_cm))
            except Exception:
                try:
                    run.add_picture(img_path, width=Cm(img_w_cm))
                except Exception:
                    pass

        self._docx_heading(doc, "2. TÍNH TOÁN", 1)
        self._docx_heading(doc, self._report_phi_heading(res), 2)
        self._docx_make_table(doc, self._report_phi_rows(res), header_rows=1, widths=self._report_phi_widths_cm(), font_size=10, table_align="left")
        self._docx_heading(doc, "2.2. Hệ số chiết giảm nhóm của cọc", 2)
        self._docx_make_table(doc, self._report_factor_rows(res), header_rows=1, widths=self._docx_widths_from_sample("factor"), font_size=10, table_align="left")
        self._docx_heading(doc, "2.3. Sức kháng bên", 2)
        self._docx_make_table(doc, self._report_shaft_formula_rows(res), header_rows=1, widths=self._docx_widths_from_sample("shaft_formula"), font_size=10, merge_repeat_cols=[0,2], min_row_height_in=0.22, table_align="left", left_cols=[1,2])
        self._docx_heading(doc, "Bảng kết quả sức kháng bên theo từng lớp", 2)
        layer_rows = self._report_layer_rows(res)
        self._docx_make_table(doc, layer_rows, header_rows=1, widths=self._layer_table_widths_cm(layer_rows, total_cm=self._report_content_width_cm()), font_size=10, min_row_height_in=0.22, table_align="left")
        self._docx_heading(doc, "2.4. Sức kháng mũi cọc", 2)
        self._docx_heading(doc, "Công thức áp dụng", 2)
        self._docx_make_table(doc, self._report_tip_formula_rows(res), header_rows=1, widths=self._docx_widths_from_sample("tip_formula"), font_size=10, merge_repeat_cols=[0,2], min_row_height_in=0.22, table_align="left", left_cols=[1,2])
        self._docx_heading(doc, "Nội dung tính toán", 2)
        tip_rows = [["Nội dung", "Giá trị"], ["Sức kháng mũi cọc", self._report_tip_summary_note(res)], ["Qp CĐ", f"{_fmt(res.strength.qtip_kn,0)} kN"], ["Qp ĐB", f"{_fmt(res.extreme.qtip_kn,0)} kN"]]
        self._docx_make_table(doc, tip_rows, header_rows=1, widths=self._docx_widths_from_sample("tip_summary"), font_size=10, table_align="left")
        self._docx_heading(doc, "2.5. Sức chịu tải của cọc", 2)
        self._docx_make_table(doc, self._report_capacity_rows(res), header_rows=1, widths=self._docx_widths_from_sample("capacity"), font_size=10, table_align="left")
        self._docx_heading(doc, "2.6. Kiểm toán", 2)
        self._docx_make_table(doc, self._report_check_rows(res), header_rows=1, widths=self._docx_widths_from_sample("check"), font_size=10, table_align="left")
        if res.warnings:
            self._docx_heading(doc, "Cảnh báo", 2)
            for w in res.warnings:
                doc.add_paragraph(str(w), style=None)
        # Footer
        for section in doc.sections:
            footer = section.footer.paragraphs[0]
            footer.text = f"File: {os.path.basename(path)}"
            footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.save(path)

    def _pdf_font_names(self):
        """Đăng ký font có hỗ trợ tiếng Việt cho ReportLab."""
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except Exception:
            return "Helvetica", "Helvetica-Bold"
        candidates_regular = [
            r"C:\Windows\Fonts\times.ttf", r"C:\Windows\Fonts\arial.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        ]
        candidates_bold = [
            r"C:\Windows\Fonts\timesbd.ttf", r"C:\Windows\Fonts\arialbd.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf",
        ]
        reg = next((p for p in candidates_regular if os.path.exists(p)), None)
        bold = next((p for p in candidates_bold if os.path.exists(p)), None)
        if reg:
            try:
                pdfmetrics.registerFont(TTFont("N2DRegular", reg))
                if bold:
                    pdfmetrics.registerFont(TTFont("N2DBold", bold))
                else:
                    pdfmetrics.registerFont(TTFont("N2DBold", reg))
                return "N2DRegular", "N2DBold"
            except Exception:
                pass
        return "Helvetica", "Helvetica-Bold"

    def _pdf_table(self, data, col_widths=None, font="Helvetica", bold_font="Helvetica-Bold", font_size=10, header_bg="#D9EAF7", grid_color="#666666", table_align: str = "CENTER", left_cols: Optional[List[int]] = None):
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        align_key = str(table_align or "CENTER").upper()
        left_cols_set = set(int(c) for c in (left_cols or []))
        try:
            font_size = self._report_auto_table_font_size(data, col_widths, font_size)
        except Exception:
            pass
        lead = max(font_size + 2, 10)
        style = ParagraphStyle("cell", fontName=font, fontSize=font_size, leading=lead, alignment=0, splitLongWords=0)
        hstyle = ParagraphStyle("head", fontName=bold_font, fontSize=font_size, leading=lead, alignment=1, splitLongWords=0)
        estyle = ParagraphStyle("emphasis", fontName=bold_font, fontSize=font_size, leading=lead, alignment=0, textColor=colors.HexColor("#C00000"), splitLongWords=0)
        pdata = []
        for i, row in enumerate(data):
            out = []
            row_emphasis = self._report_row_emphasis(row, i, 1)
            for j, val in enumerate(row):
                txt = self._report_blank_zero_text(val)
                if i == 0:
                    align_style = hstyle
                elif row_emphasis:
                    align_style = estyle if (j in left_cols_set or j == 0) else ParagraphStyle(f"emph_{i}_{j}", parent=estyle, alignment=1, splitLongWords=0)
                else:
                    align_style = style if j in left_cols_set or j == 0 else ParagraphStyle(f"cell_{i}_{j}", parent=style, alignment=1, splitLongWords=0)
                out.append(Paragraph(self._report_pdf_markup(txt), align_style))
            pdata.append(out)
        t = Table(pdata, colWidths=col_widths, repeatRows=1, hAlign="LEFT" if align_key == "LEFT" else "CENTER")
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(grid_color)),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(header_bg)),
            ("FONTNAME", (0, 0), (-1, 0), bold_font),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ("TOPPADDING", (0, 0), (-1, -1), 1.6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1.6),
            ("LEFTPADDING", (0, 0), (-1, -1), 1.8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 1.8),
        ]))
        return t

    def _pdf_layer_widths(self, rows, cm, total_cm: float = 19.0):
        vals = self._layer_table_widths_cm(rows, total_cm=total_cm)
        return [v * cm for v in vals]

    def _write_item_pdf_report(self, res: CapacityResult, path: str):
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, KeepTogether
            from reportlab.lib.pagesizes import A4, A3, landscape
            from reportlab.lib import colors
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.units import cm
        except Exception as exc:
            raise RuntimeError(f"Thiếu thư viện reportlab. Cài bằng: pip install reportlab. Chi tiết: {exc}")
        font, bold_font = self._pdf_font_names()
        pagesize = self._report_pdf_pagesize() or A4
        doc = SimpleDocTemplate(path, pagesize=pagesize, rightMargin=1.35*cm, leftMargin=1.35*cm, topMargin=1.35*cm, bottomMargin=1.25*cm)
        title_style = ParagraphStyle("title", fontName=bold_font, fontSize=15, leading=18, alignment=1, textColor=colors.blue)
        sec_style = ParagraphStyle("section", fontName=bold_font, fontSize=10, leading=13, spaceBefore=8, spaceAfter=3)
        body_style = ParagraphStyle("body", fontName=font, fontSize=10, leading=12)
        small_style = ParagraphStyle("small", fontName=font, fontSize=10, leading=12)
        story = []
        story.append(Paragraph(self._capacity_report_title(res), title_style))
        story.append(Paragraph("1. SỐ LIỆU ĐẦU VÀO", sec_style))
        input_data = [["Nội dung", "Ký hiệu", "Giá trị", "Đơn vị"]] + [[a, b, c, d] for a, b, c, d in self._report_input_rows(res)]
        story.append(self._pdf_table(input_data, col_widths=[w*cm for w in self._scale_widths_to_total([7.30, 2.10, 3.20, 1.60], self._report_content_width_cm())], font=font, bold_font=bold_font, font_size=10, table_align="LEFT"))
        story.append(Paragraph("2. TÍNH TOÁN", sec_style))
        story.append(Paragraph(self._report_phi_heading(res), sec_style))
        story.append(self._pdf_table(self._report_phi_rows(res), col_widths=[w*cm for w in self._report_phi_widths_cm()], font=font, bold_font=bold_font, font_size=10, table_align="LEFT"))
        story.append(Paragraph("2.2. Hệ số chiết giảm nhóm của cọc", sec_style))
        story.append(self._pdf_table(self._report_factor_rows(res), col_widths=[w*cm for w in self._scale_widths_to_total([9.051, 9.051], self._report_content_width_cm())], font=font, bold_font=bold_font, font_size=10, table_align="LEFT"))
        story.append(Paragraph("2.3. Sức kháng bên", sec_style))
        story.append(self._pdf_table(self._blank_repeated_cells(self._report_shaft_formula_rows(res), [0,2]), col_widths=[w*cm for w in self._scale_widths_to_total([4.159, 7.909, 6.034], self._report_content_width_cm())], font=font, bold_font=bold_font, font_size=10.0, table_align="LEFT", left_cols=[1,2]))
        layer_rows = self._report_layer_rows(res)
        story.append(self._pdf_table(layer_rows, col_widths=self._pdf_layer_widths(layer_rows, cm, total_cm=self._report_content_width_cm()), font=font, bold_font=bold_font, font_size=10, table_align="LEFT"))
        story.append(Paragraph("2.4. Sức kháng mũi cọc", sec_style))
        story.append(Paragraph("Công thức áp dụng", sec_style))
        story.append(self._pdf_table(self._blank_repeated_cells(self._report_tip_formula_rows(res), [0,2]), col_widths=[w*cm for w in self._scale_widths_to_total([3.683, 8.386, 6.034], self._report_content_width_cm())], font=font, bold_font=bold_font, font_size=10.0, table_align="LEFT", left_cols=[1,2]))
        story.append(Paragraph("Nội dung tính toán", sec_style))
        story.append(self._pdf_table([["Nội dung", "Giá trị"], ["Sức kháng mũi cọc", self._report_tip_summary_note(res)], ["Qp CĐ", f"{_fmt(res.strength.qtip_kn,0)} kN"], ["Qp ĐB", f"{_fmt(res.extreme.qtip_kn,0)} kN"]], col_widths=[w*cm for w in self._scale_widths_to_total([9.051, 9.051], self._report_content_width_cm())], font=font, bold_font=bold_font, font_size=10, table_align="LEFT"))
        story.append(Paragraph("2.5. Sức chịu tải của cọc", sec_style))
        story.append(self._pdf_table(self._report_capacity_rows(res), col_widths=[w*cm for w in self._scale_widths_to_total([9.051, 4.525, 4.525], self._report_content_width_cm())], font=font, bold_font=bold_font, font_size=10, table_align="LEFT"))
        story.append(Paragraph("2.6. Kiểm toán", sec_style))
        story.append(self._pdf_table(self._report_check_rows(res), col_widths=[w*cm for w in self._scale_widths_to_total([9.051, 4.525, 4.525], self._report_content_width_cm())], font=font, bold_font=bold_font, font_size=10, table_align="LEFT"))
        if res.warnings:
            story.append(Paragraph("Cảnh báo", sec_style))
            for w in res.warnings:
                story.append(Paragraph("- " + str(w), small_style))
        def page_cb(canvas, doc_obj):
            canvas.saveState()
            canvas.setFont(font, 7)
            canvas.drawString(doc_obj.leftMargin, 0.7*cm, f"File: {os.path.basename(path)}")
            canvas.drawRightString(doc_obj.pagesize[0]-doc_obj.rightMargin, 0.7*cm, f"Page: {canvas.getPageNumber()}")
            canvas.restoreState()
        doc.build(story, onFirstPage=page_cb, onLaterPages=page_cb)

    def export_report(self):
        'Xuất báo cáo riêng từng hạng mục, tham khảo cách tổ chức option của TS-COL.'
        if getattr(self, "_export_busy", False):
            messagebox.showinfo("Đang xuất báo cáo", "Một phiên xuất báo cáo đang chạy; chờ xong rồi bấm lại.")
            return
        results = getattr(self, "last_results", None) or ([self.last_result] if self.last_result else [])
        results = [r for r in results if r]
        if not results:
            messagebox.showwarning("Chưa có kết quả", "Hãy bấm Tính toán trước khi xuất báo cáo.")
            return
        opts = self._ask_report_export_options(results)
        if not opts:
            return
        fmt_u = str(opts.get("format") or "DOC").strip().upper()
        if fmt_u in ("DOCX", "DOC"):
            export_docx, export_pdf = True, False
        elif fmt_u == "PDF":
            export_docx, export_pdf = False, True
        else:
            export_docx, export_pdf = True, True
        self._report_page_size = str(opts.get("page_size") or "A4")
        self._report_orientation = str(opts.get("orientation") or "Dọc")
        self._report_result_units = str(opts.get("result_units") or "kN")
        suffix = str(opts.get("suffix") or "_CAP")
        naming_mode = str(opts.get("naming_mode") or "Theo hạng mục")
        custom_name = str(opts.get("custom_name") or "").strip()
        selected_results = list(results)
        scope_text = str(opts.get("scope") or "")
        if scope_text.startswith("Chọn"):
            selected_keys = set(str(k) for k in (opts.get("selected_item_keys") or []))
            selected_results = [r for r in results if _normalize_item_name(r.pile_input.item) in selected_keys]
            if not selected_results:
                messagebox.showwarning("Chọn hạng mục", "Không tìm thấy kết quả tính cho các hạng mục đã chọn.")
                return
        elif scope_text.startswith("Hạng mục"):
            sel_name = ""
            try:
                sel = self.item_table.tree.selection()
                if sel:
                    vals = list(self.item_table.tree.item(sel[0], "values"))
                    sel_name = str(vals[0] if vals else "").strip()
            except Exception:
                sel_name = ""
            if not sel_name:
                messagebox.showwarning("Chọn hạng mục", "Chưa chọn hạng mục trong tab Thông tin riêng. Chương trình sẽ xuất tất cả hạng mục.")
            else:
                key = _normalize_item_name(sel_name)
                selected_results = [r for r in results if _normalize_item_name(r.pile_input.item) == key]
                if not selected_results:
                    messagebox.showwarning("Chọn hạng mục", f"Không tìm thấy kết quả tính cho hạng mục: {sel_name}")
                    return
        out_dir = filedialog.askdirectory(title='Chọn thư mục lưu báo cáo TS-CAP')
        if not out_dir:
            return
        self._export_busy = True
        saved: List[str] = []
        errors: List[str] = []
        multi_items = len(selected_results) > 1
        # QA-UX U4: xuất từng hạng mục theo từng nhịp event-loop (root.after) kèm hộp tiến
        # trình; giữa hai hạng mục giao diện được vẽ lại nên không rơi vào "Not responding".
        # Hàm ghi DOCX/PDF giữ nguyên main thread (không rủi ro Tkinter đa luồng).
        prog = tk.Toplevel(self.root)
        prog.title("Đang xuất báo cáo")
        prog.transient(self.root)
        prog.resizable(False, False)
        try:
            prog.protocol("WM_DELETE_WINDOW", lambda: None)
        except Exception:
            pass
        pf = ttk.Frame(prog, padding=14)
        pf.pack(fill=tk.BOTH, expand=True)
        ptxt = tk.StringVar(value=f"Chuẩn bị xuất {len(selected_results)} hạng mục...")
        ttk.Label(pf, textvariable=ptxt, wraplength=430).pack(anchor=tk.W, pady=(0, 8))
        pbar = ttk.Progressbar(pf, mode="determinate", maximum=max(len(selected_results), 1), length=430)
        pbar.pack(fill=tk.X)
        try:
            self._center_window(prog)
        except Exception:
            pass
        safe_lift_window(prog)

        def _finish_export():
            self._export_busy = False
            try:
                prog.destroy()
            except Exception:
                pass
            if saved:
                self._set_status(f"Đã xuất {len(saved)} file báo cáo TS-CAP vào {os.path.basename(out_dir)}")
            msg = f"Đã xuất {len(saved)} file vào:\n{out_dir}"
            if errors:
                msg += "\n\nMột số lỗi:\n" + "\n".join(errors[:10])
            messagebox.showinfo("Xuất báo cáo", msg)

        def _export_one(idx: int):
            if not self._root_alive():
                self._export_busy = False
                return
            if idx >= len(selected_results):
                _finish_export()
                return
            res = selected_results[idx]
            try:
                ptxt.set(f"Đang xuất {idx + 1}/{len(selected_results)}: {res.pile_input.item}")
                pbar.configure(value=idx)
                prog.update_idletasks()
            except Exception:
                pass
            if naming_mode == "Tên tự đặt" and custom_name:
                base = custom_name if not multi_items else f"{custom_name}_{res.pile_input.item or 'HANG_MUC'}"
                stem = self._sanitize_filename(base + suffix)
            else:
                stem = self._sanitize_filename((res.pile_input.item or "HANG_MUC") + suffix)
            if export_docx:
                docx_path = self._unique_path(out_dir, stem, ".docx")
                try:
                    self._write_item_docx_report(res, docx_path)
                    saved.append(docx_path)
                except Exception as exc:
                    errors.append(f"{res.pile_input.item} DOCX: {exc}")
            if export_pdf:
                pdf_path = self._unique_path(out_dir, stem, ".pdf")
                try:
                    self._write_item_pdf_report(res, pdf_path)
                    saved.append(pdf_path)
                except Exception as exc:
                    errors.append(f"{res.pile_input.item} PDF: {exc}")
            self.root.after(10, lambda: _export_one(idx + 1))

        self.root.after(10, lambda: _export_one(0))

    def export_csv(self):
        results = getattr(self, "last_results", None) or ([self.last_result] if self.last_result else [])
        if not results:
            messagebox.showwarning("Chưa có kết quả", "Hãy bấm Tính toán trước khi xuất CSV.")
            return
        path = filedialog.asksaveasfilename(title='Lưu CSV kết quả TS-CAP', defaultextension=".csv", filetypes=[("CSV", "*.csv")])
        if not path:
            return
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow([APP_NAME, APP_SUBTITLE])
            writer.writerow([])
            writer.writerow(["SUMMARY"])
            writer.writerow(["Project", "Item", "PileType", "Mode", "GroupLayout", "D_mm", "L_m", "fg_CD", "DD_CD", "PuDD_CD_perPile_kN", "Q_single_CD_kN", "FOS_single_CD", "NcapDD_CD_kN", "Q_group_CD_kN", "FOS_group_CD", "Tu_CD_perPile_kN", "T_single_CD_kN", "FOS_uplift_CD", "T_group_CD_kN", "FOS_uplift_group_CD", "DD_DB", "PuDD_DB_perPile_kN", "Q_single_DB_kN", "FOS_single_DB", "NcapDD_DB_kN", "Q_group_DB_kN", "FOS_group_DB", "Tu_DB_perPile_kN", "T_single_DB_kN", "FOS_uplift_DB", "T_group_DB_kN", "FOS_uplift_group_DB"])
            for res in results:
                c = self._report_check_values(res)
                writer.writerow([
                    res.pile_input.project, res.pile_input.item, res.pile_input.pile_type, res.pile_input.mode,
                    res.pile_input.group_layout, res.pile_input.diameter_mm, res.pile_input.pile_length_m, res.strength.group_factor,
                    c['dd_cd'], c['pu_cd_design'], c['single_cd'], self._report_fos_text(c['dcr_single_cd']), c['ncap_cd_design'], c['group_cd'], self._report_fos_text(c['dcr_group_cd']),
                    c['uplift_cd'], c['cap_uplift_cd'], self._report_fos_text(c['dcr_uplift_cd']), c['cap_uplift_group_cd'], self._report_fos_text(c['dcr_uplift_group_cd']),
                    c['dd_db'], c['pu_db_design'], c['single_db'], self._report_fos_text(c['dcr_single_db']), c['ncap_db_design'], c['group_db'], self._report_fos_text(c['dcr_group_db']),
                    c['uplift_db'], c['cap_uplift_db'], self._report_fos_text(c['dcr_uplift_db']), c['cap_uplift_group_db'], self._report_fos_text(c['dcr_uplift_group_db']),
                ])
            for res in results:
                writer.writerow([])
                writer.writerow(["DETAIL", res.pile_input.item])
                writer.writerow(["Trong luong coc", "Gia tri", "Don vi"])
                writer.writerow(["W coc kho", res.strength.pile_weight_dry_kn, "kN"])
                writer.writerow(["W coc huu hieu", res.strength.pile_weight_effective_kn, "kN"])
                writer.writerow(["Luc day noi", res.strength.buoyancy_kn, "kN"])
                writer.writerow([])
                writer.writerow(["TTGH", "Qs", "Qp", "Qr gross", "W' coc", "DD", "[Q] coc don net", "Qshaft uplift", "|Qr uplift|", "Tong uplift nhom", "Qr uplift signed", "fg", "[Q] nhom/coc", "Tong [Q] nhom", "Pr vat lieu", "Khong che", "Don vi"])
                for cap in (res.strength, res.extreme):
                    writer.writerow([cap.label, cap.qshaft_kn, cap.qtip_kn, cap.compression_single_gross_kn, cap.pile_weight_effective_kn, cap.downdrag_kn, cap.compression_single_net_kn, cap.qshaft_uplift_kn, cap.uplift_single_magnitude_kn, cap.uplift_group_total_kn, cap.uplift_single_signed_kn, cap.group_factor, cap.compression_group_single_net_kn, cap.compression_group_total_kn, cap.material_pr_kn, cap.governing_kn, "kN"])
                writer.writerow([])
                writer.writerow(["Item", "Layer", "Top", "Bottom", "Thickness", "SkinLen", "DowndragLen", "Downdrag", "SoilType", "fg", "Nₕₜ", "N60", "(N1)60", "gamma_eff", "sigma_v_eff", "C/Su", "alpha/beta", "phi_CD", "qs_CD_kPa", "Qs_CD_kN", "phi_DB", "qs_DB_kPa", "Qs_DB_kN", "qs_uplift_CD_kPa", "Qs_uplift_CD_kN", "qs_nominal_kPa", "Qs_nominal_kN", "Note"])
                for lr in res.layers:
                    writer.writerow([res.pile_input.item, lr.name, lr.top_elev_m, lr.bottom_elev_m, lr.thickness_m, lr.skin_length_m, lr.downdrag_length_m, lr.downdrag_kn, lr.soil_label, lr.group_factor, lr.n_spt, lr.n60, lr.n1_60, lr.gamma_eff_kN_m3, lr.sigma_v_eff_mpa, (lr.c_mpa if lr.c_mpa > 0 else lr.su_mpa), lr.alpha_or_beta, lr.phi_strength, lr.qs_factored_kpa, lr.qs_factored_kn, lr.phi_extreme, lr.qs_extreme_kpa, lr.qs_extreme_kn, lr.qs_uplift_factored_kpa, lr.qs_uplift_factored_kn, lr.qs_nominal_kpa, lr.qs_nominal_kn, lr.note])
                if res.warnings:
                    writer.writerow([])
                    writer.writerow(["Canh bao"])
                    for w in res.warnings:
                        writer.writerow([w])
        self._set_status(f"Đã xuất CSV: {os.path.basename(path)}")
        messagebox.showinfo("Xuất CSV", f"Đã lưu:\n{path}")

    def _read_delimited_file(self, path: str) -> List[List[str]]:
        ext = os.path.splitext(str(path))[1].lower()
        if ext in (".xlsx", ".xlsm"):
            try:
                import openpyxl
            except Exception as exc:
                raise RuntimeError(f"Thiếu openpyxl để đọc Excel: {exc}")
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                vals = ["" if v is None else str(v).strip() for v in row]
                if any(vals):
                    rows.append(vals)
            return rows
        with open(path, "r", encoding="utf-8-sig", errors="ignore") as f:
            text = f.read()
        dialect = csv.Sniffer().sniff(text[:2048], delimiters=",;\t") if text.strip() else csv.excel
        return [row for row in csv.reader(text.splitlines(), dialect) if any(str(x).strip() for x in row)]


    def _parse_combo_range_text(self, text: str) -> set:
        out = set()
        text = str(text or "").strip()
        if not text:
            return out
        for part in re.split(r"[,;\s]+", text):
            if not part:
                continue
            if "-" in part:
                a, b = part.split("-", 1)
                try:
                    ia, ib = int(float(a)), int(float(b))
                    if ia <= ib:
                        out.update(str(i) for i in range(ia, ib + 1))
                    else:
                        out.update(str(i) for i in range(ib, ia + 1))
                except Exception:
                    out.add(part.strip())
            else:
                try:
                    out.add(str(int(float(part))))
                except Exception:
                    out.add(part.strip())
        return out

    def _cluster_count(self, values: List[float], tol: float = 0.05) -> int:
        vals = sorted(float(v) for v in values)
        if not vals:
            return 0
        groups = [vals[0]]
        for v in vals[1:]:
            if abs(v - groups[-1]) > tol:
                groups.append(v)
        return len(groups)

    def _min_pile_spacing(self, piles: List[Dict[str, Any]], fallback: float = 3.0) -> float:
        best = None
        for i, p1 in enumerate(piles):
            for p2 in piles[i+1:]:
                d = math.hypot(float(p1.get("X", 0.0)) - float(p2.get("X", 0.0)), float(p1.get("Y", 0.0)) - float(p2.get("Y", 0.0)))
                if d > 1e-9 and (best is None or d < best):
                    best = d
        return float(best if best is not None else fallback)

    def _row_count_from_pile_coords(self, piles: List[Dict[str, Any]]) -> int:
        """Số hàng cọc để xét hệ số nhóm: lấy theo phương dọc cầu X.

        Ví dụ có 3 tọa độ X khác nhau thì xem là 3 hàng cọc; không xét số vị trí Y.
        """
        if not piles:
            return 1
        nx = self._cluster_count([float(p.get("X", 0.0)) for p in piles])
        return max(int(nx if nx > 0 else 1), 1)

    def _parse_foundation_input_file_for_sct(self, filepath: str) -> Dict[str, Any]:
        'Đọc file INPUT MCOC hoặc file text có cùng chuỗi số như parser TS-PILE.\n\n        Trả về data tối thiểu gồm Global, Load_Combos, Piles. Không giải nội lực.\n        '
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            raw_text = f.read()
        raw_norm = _strip_accents(raw_text).lower()
        # Chặn lỗi rất nguy hiểm: output classic của TS-PILE/MCOC cũng chứa rất nhiều số,
        # nếu đem đọc bằng parser INPUT sẽ bị lệch hoàn toàn n_cọc/n_TH/Bx/By/N đáy bệ.
        # File output phải đi qua _parse_mcoc_output_text_for_sct(), không parse như INPUT.
        if ("chuong trinh tinh mong coc" in raw_norm or
            "cac to hop tai trong tinh toan" in raw_norm or
            "noi luc dau coc" in raw_norm or
            "bang tong ket noi luc" in raw_norm):
            raise ValueError('Đây là file output TS-PILE/MCOC, không phải file INPUT số thuần.')
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        if not lines:
            raise ValueError("File rỗng")
        data = {"Global": {}, "Load_Combos": [], "Piles": []}
        data["Global"]["Project_ID"] = lines[0].strip()
        content = " ".join(lines[1:])
        tokens = []
        for t in re.split(r"\s+", content):
            try:
                tokens.append(float(t.replace(",", ".")))
            except Exception:
                pass
        if len(tokens) < 18:
            raise ValueError("Không nhận dạng được file INPUT MCOC: thiếu 18 số đầu.")
        n_piles, n_combos = int(tokens[0]), int(tokens[1])
        expected = 18 + 6 * n_combos + 16 * n_piles
        if len(tokens) < expected:
            raise ValueError(f"File INPUT MCOC thiếu dữ liệu: cần {expected} số, chỉ có {len(tokens)}.")
        data["Global"].update({
            "Kn": tokens[6], "Bx": tokens[8], "By": tokens[9], "Cz": tokens[10],
            "EI_uon": tokens[11], "Er_uon": tokens[12], "EA_nen": tokens[13], "Er_nen": tokens[14],
            "md": tokens[15], "mq": tokens[16], "m": tokens[17]
        })
        idx = 18
        for i in range(n_combos):
            data["Load_Combos"].append({"Name": str(i+1), "Hx": tokens[idx], "Hy": tokens[idx+1], "N_load": tokens[idx+2], "Mx": tokens[idx+3], "My": tokens[idx+4], "Mz": tokens[idx+5]})
            idx += 6
        for i in range(n_piles):
            data["Piles"].append({"Name": str(i+1), "Lo": tokens[idx], "H": tokens[idx+1], "Bpx": tokens[idx+2], "Bpy": tokens[idx+3], "d_ngoai": tokens[idx+4], "d_trong": tokens[idx+5], "day_vo": tokens[idx+6], "Area": tokens[idx+7], "J_xy": tokens[idx+8], "Po": tokens[idx+9], "Co": tokens[idx+10], "Ct": tokens[idx+11], "X": tokens[idx+12], "Y": tokens[idx+13], "Phi": tokens[idx+14], "Xi": tokens[idx+15]})
            idx += 16
        return data

    def _parse_forces_csv_for_sct(self, path: str) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        rows = self._read_delimited_file(path)
        if not rows:
            return [], {}
        header = [h.strip() for h in rows[0]]
        hlow = [h.lower() for h in header]
        forces = []
        combos = {}
        def idx_any(names):
            for name in names:
                if name.lower() in hlow:
                    return hlow.index(name.lower())
            return None
        ipile = idx_any(["T.C", "TC", "Pile", "Cọc", "Coc"])
        icombo = idx_any(["T.H", "TH", "Combo", "Tổ hợp", "To hop"])
        incol = idx_any(["N", "N_tf", "Pu", "Pu_tf"])
        iitem = idx_any(["Item", "Hạng mục", "Hang muc", "SourceFile"])
        if incol is None:
            return [], {}
        for r in rows[1:]:
            vals = r + [""] * len(header)
            forces.append({
                "item": vals[iitem] if iitem is not None else "",
                "pile": vals[ipile] if ipile is not None else "",
                "combo": vals[icombo] if icombo is not None else "",
                "N_ton": _safe_float(vals[incol], 0.0),
            })
        return forces, combos

    def _extract_mcoc_project_name_for_sct(self, lines: List[str], fallback: str = "") -> str:
        'Dò tên hạng mục/công trình trong output MCOC/TS-PILE.\n\n        Logic bám theo TS-COL: ưu tiên dòng "Công trình : ..." để tên hạng mục\n        dùng cho khớp dữ liệu SCT, địa chất và báo cáo.\n        '
        patterns = [
            r"c[oô]ng\s*tr[iì]nh\s*[:：]\s*(.+)$",
            r"cong\s*trinh\s*[:：]\s*(.+)$",
            r"h[aạ]ng\s*m[uụ]c\s*[:：]\s*(.+)$",
            r"hang\s*muc\s*[:：]\s*(.+)$",
        ]
        for raw in lines:
            line = str(raw or "").strip()
            if not line:
                continue
            low = _strip_accents(line).lower()
            if "chuong trinh" in low or "bo giao thong" in low or "module" in low:
                continue
            for pat in patterns:
                m = re.search(pat, line, flags=re.I)
                if m:
                    name = re.sub(r"\s+", " ", m.group(1).strip()).strip("-–— ")
                    if name:
                        return name
        return fallback

    def _parse_mcoc_output_text_for_sct(self, path: str) -> Dict[str, Any]:
        'Đọc output MCOC/TS-PILE theo cùng nguyên tắc import MCOC của TS-COL.\n\n        Output chuẩn có các vùng chính:\n        - CÁC TỔ HỢP TẢI TRỌNG TÍNH TOÁN: T.T, Hx, Hy, P, Mx, My, Mz.\n          Với module SCT, lực đáy bệ lấy từ cột P/N của bảng tải trọng ban đầu.\n        - THÔNG SỐ CỌC: T.C, Lo, H, Bpx, Bpy, A/D, B/dtr, Cday, Fo, Io, Po, Co, Ct.\n          Dòng "n t" được hiểu là cọc sau dùng lại thông số hình học dòng trước.\n        - TOẠ ĐỘ ĐẦU CỌC: T.C, X, Y, Phi, Xi. Phần này dùng để lấy số cọc,\n          số hàng cọc và khoảng cách cọc nhỏ nhất.\n        - NỘI LỰC ĐẦU CỌC: T.C, T.H, N, Q2, Q3, M1, M2, M3.\n          TS-COL lấy đủ 6 thành phần; SCT chỉ lấy N để tìm Pu/cọc lớn nhất.\n\n        Đơn vị output MCOC: lực tấn, moment tấn.m. Hàm này chỉ lưu N_ton/P_ton,\n        việc đổi sang kN thực hiện ở bước import hạng mục.\n        '
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()
        except Exception:
            return {}

        source = os.path.splitext(os.path.basename(path))[0]
        item = self._extract_mcoc_project_name_for_sct(lines, fallback=source)
        forces: List[Dict[str, Any]] = []
        combos: List[Dict[str, Any]] = []
        piles_by_name: Dict[str, Dict[str, Any]] = {}
        g: Dict[str, Any] = {}

        def norm(txt: str) -> str:
            return _strip_accents(str(txt or "")).lower()

        num_pat = r"[-+]?\d+(?:[\.,]\d+)?"
        def nums_in(line: str) -> List[str]:
            return re.findall(num_pat, str(line or ""))
        def as_name(x: Any) -> str:
            s = str(x).replace(',', '.').strip()
            try:
                f = float(s)
                if abs(f - round(f)) < 1e-9:
                    return str(int(round(f)))
            except Exception:
                pass
            return str(x).strip()

        # 0) Dò thông số chung in trên đầu output.
        # Output chuẩn có dòng: Kn = ... Ax = ... By = ... Cz = ...
        # Trong TS-PILE/SCT, Ax của output MCOC được hiểu là Bx bệ.
        for raw in lines[:120]:
            line = raw.strip()
            low = norm(line)
            if not line:
                continue
            m = re.search(r"\bkn\s*=\s*(%s).*?\ba[x]?\s*=\s*(%s).*?\bby\s*=\s*(%s).*?\bcz\s*=\s*(%s)" % (num_pat, num_pat, num_pat, num_pat), low, flags=re.I)
            if m:
                g["Kn"] = _safe_float(m.group(1), 0.0)
                g["Bx"] = _safe_float(m.group(2), 0.0)
                g["By"] = _safe_float(m.group(3), 0.0)
                g["Cz"] = _safe_float(m.group(4), 0.0)
                continue
            m = re.search(r"e\s*v\.?uon\s*=\s*(%s).*?e\s*r\.?uon\s*=\s*(%s).*?e\s*v\.?nen\s*=\s*(%s).*?e\s*r\.?nen\s*=\s*(%s)" % (num_pat, num_pat, num_pat, num_pat), low, flags=re.I)
            if m:
                g["EI_uon"] = _safe_float(m.group(1), 0.0)
                g["Er_uon"] = _safe_float(m.group(2), 0.0)
                g["EA_nen"] = _safe_float(m.group(3), 0.0)
                g["Er_nen"] = _safe_float(m.group(4), 0.0)
                continue
            # Dò nhãn lẻ, tránh bắt sai các bảng tính phía dưới.
            if "chieu dai be" in low or "bx be" in low:
                ns = nums_in(line)
                if ns:
                    g.setdefault("Bx", _safe_float(ns[-1], 0.0))
            if "chieu rong be" in low or "by be" in low:
                ns = nums_in(line)
                if ns:
                    g.setdefault("By", _safe_float(ns[-1], 0.0))
            if "chieu cao be" in low or re.search(r"\bcz\b", low):
                ns = nums_in(line)
                if ns:
                    g.setdefault("Cz", _safe_float(ns[-1], 0.0))

        # 1) Bảng tải trọng ban đầu: T.T Hx Hy P/N Mx My Mz.
        # Dùng cột P/N làm N đáy bệ cho kiểm toán nhóm cọc.
        in_combo = False
        header_seen = False
        for raw in lines:
            line = raw.strip()
            low = norm(line)
            if not line:
                continue
            if ("cac to hop" in low and ("tai" in low or "load" in low)) or "load combinations" in low:
                in_combo = True
                header_seen = False
                continue
            if in_combo and ("thong so coc" in low or "toa do" in low or "chuyen vi" in low or "noi luc dau coc" in low):
                break
            if not in_combo:
                continue
            if re.search(r"\b(t\.t|tt|t\.h|th)\b", low) and ("hx" in low and "hy" in low):
                header_seen = True
                continue
            ns = nums_in(line)
            if len(ns) >= 7:
                # Cho phép đọc cả khi header bị mất, nhưng chỉ trong vùng combo.
                try:
                    combos.append({
                        "Name": as_name(ns[0]),
                        "Hx": _safe_float(ns[1], 0.0),
                        "Hy": _safe_float(ns[2], 0.0),
                        "N_load": _safe_float(ns[3], 0.0),
                        "Mx": _safe_float(ns[4], 0.0),
                        "My": _safe_float(ns[5], 0.0),
                        "Mz": _safe_float(ns[6], 0.0),
                    })
                except Exception:
                    continue

        # 2) Thông số cọc. Xử lý đúng dòng "n t" trong output MCOC:
        # cọc sau dùng lại thông số hình học dòng trước.
        # Lưu ý form classic có thể bị xuống dòng ở cuối dòng thông số cọc:
        #   1  0.00 14.00 ... Po Co
        #   Ct
        # nên parser phải ghép dòng Ct kế tiếp trước khi đọc 13 số.
        pile_section: List[str] = []
        in_pile = False
        for raw in lines:
            st = raw.rstrip("\n")
            low = norm(st.strip())
            if "thong so coc" in low:
                in_pile = True
                continue
            if in_pile and ("toa do" in low or "chuyen vi" in low or "noi luc dau coc" in low or "bang tong ket" in low):
                break
            if in_pile:
                pile_section.append(st)

        last_geom: Optional[Dict[str, Any]] = None
        i = 0
        while i < len(pile_section):
            st = pile_section[i].strip()
            low = norm(st)
            i += 1
            if not st:
                continue
            # Bỏ qua dòng tiêu đề/câu chú thích trong vùng thông số cọc.
            if ("t.c" in low and "lo" in low and "bpx" in low) or "coc tron" in low or "d ngoai" in low or "d trong" in low:
                continue
            ns = nums_in(st)
            if not ns:
                continue
            # Dòng "2   n t" nghĩa là cọc 2 dùng thông số hình học của cọc 1.
            if last_geom is not None and (re.search(r"\bn\s*t\b", low) or re.search(r"\bnt\b", low)):
                name = as_name(ns[0])
                piles_by_name[name] = {"Name": name, **last_geom}
                continue
            # Form classic đôi khi bị wrap mất Ct sang dòng kế tiếp. Nếu dòng hiện tại mới có 12 số,
            # thử lấy thêm đúng 1 số của dòng kế tiếp làm Ct.
            if len(ns) == 12:
                j = i
                while j < len(pile_section):
                    nxt = pile_section[j].strip()
                    low_next = norm(nxt)
                    if not nxt:
                        j += 1
                        continue
                    ns2 = nums_in(nxt)
                    if (len(ns2) == 1 and
                        not ("t.c" in low_next and "lo" in low_next) and
                        not re.search(r"\bn\s*t\b", low_next) and
                        "coc tron" not in low_next and "toa do" not in low_next):
                        ns = ns + ns2
                        i = j + 1
                    break
            if len(ns) >= 13:
                try:
                    name = as_name(ns[0])
                    geom = {
                        "Name": name,
                        "Lo": _safe_float(ns[1], 0.0),
                        "H": _safe_float(ns[2], 0.0),
                        "Bpx": _safe_float(ns[3], 0.0),
                        "Bpy": _safe_float(ns[4], 0.0),
                        "d_ngoai": _safe_float(ns[5], 0.0),
                        "d_trong": _safe_float(ns[6], 0.0),
                        "day_vo": _safe_float(ns[7], 0.0),
                        "Area": _safe_float(ns[8], 0.0),
                        "J_xy": _safe_float(ns[9], 0.0),
                        "Po": _safe_float(ns[10], 0.0),
                        "Co": _safe_float(ns[11], 0.0),
                        "Ct": _safe_float(ns[12], 0.0),
                    }
                    piles_by_name[name] = geom
                    last_geom = {k: v for k, v in geom.items() if k != "Name"}
                except Exception:
                    continue
        # 3) Tọa độ đầu cọc. Dừng trước bảng CHUYỂN VỊ BỆ CỌC để tránh nhầm
        # chuyển vị bệ thành tọa độ cọc.
        in_coord = False
        header_seen = False
        for raw in lines:
            line = raw.strip()
            low = norm(line)
            if not line:
                continue
            if "toa do" in low and "coc" in low:
                in_coord = True
                header_seen = False
                continue
            if in_coord and ("chuyen vi" in low or "noi luc dau coc" in low or "bang tong ket" in low or "tinh toan" in low):
                break
            if not in_coord:
                continue
            if re.search(r"\bt\.?c\b", low) and "x" in low and "y" in low:
                header_seen = True
                continue
            ns = nums_in(line)
            if len(ns) >= 5:
                try:
                    name = as_name(ns[0])
                    p = piles_by_name.setdefault(name, {"Name": name})
                    p.update({
                        "X": _safe_float(ns[1], 0.0),
                        "Y": _safe_float(ns[2], 0.0),
                        "Phi": _safe_float(ns[3], 0.0),
                        "Xi": _safe_float(ns[4], 0.0),
                    })
                except Exception:
                    continue

        # 4) Nội lực đầu cọc: giống TS-COL, nhưng chỉ lấy N.
        in_table = False
        header_seen = False
        current_pile = None
        for raw in lines:
            line = raw.strip()
            low = norm(line)
            if "noi luc dau coc" in low or "noi luc tai dau coc" in low:
                in_table = True
                header_seen = False
                current_pile = None
                continue
            if not in_table:
                continue
            if "bang tong ket" in low or "tinh toan kiem tra" in low or ("so sanh" in low and forces):
                break
            if not line:
                continue
            if re.search(r"\bT\.?C\b", line, flags=re.I) and re.search(r"\bT\.?H\b", line, flags=re.I) and re.search(r"\bQ2\b", line, flags=re.I):
                header_seen = True
                continue
            if not header_seen:
                continue
            ns = nums_in(line)
            if len(ns) >= 8:
                pile = ns[0]
                combo = ns[1]
                nval = ns[2]
                current_pile = pile
            elif len(ns) >= 7 and current_pile is not None:
                pile = current_pile
                combo = ns[0]
                nval = ns[1]
            else:
                continue
            try:
                forces.append({
                    "item": item,
                    "pile": as_name(pile),
                    "combo": as_name(combo),
                    "N_ton": _safe_float(nval, 0.0),
                })
            except Exception:
                continue

        if not (forces or combos or piles_by_name):
            return {}
        return {
            "Global": {"Project_ID": item, **g},
            "Load_Combos": combos,
            "Piles": list(piles_by_name.values()),
            "forces": forces,
            "source": source,
        }

    def _try_parse_text_report_for_sct(self, path: str) -> Dict[str, Any]:
        'Đọc file output MCOC/TS-PILE.\n\n        Giữ tên hàm cũ để tương thích, nhưng logic mới bám theo parser TS-COL:\n        đọc đúng bảng nội lực đầu cọc 6 thành phần và dùng riêng N cho SCT.\n        '
        return self._parse_mcoc_output_text_for_sct(path)

    def _item_row_from_n2d_data(self, data: Dict[str, Any], forces: List[Dict[str, Any]], cd_set: set, db_set: set, default_item: str) -> List[Any]:
        piles = data.get("Piles", []) or []
        g = data.get("Global", {}) or {}
        n_piles = len(piles) or 1
        first = piles[0] if piles else {}
        # Không dùng mặc định 1.2m/cọc khoan nhồi khi file MCOC không có thông số rõ ràng.
        # Để trống các ô này để người dùng tự xác nhận, tránh nhầm lẫn khi tính SCT.
        D_m = _safe_float(first.get("d_ngoai", ""), 0.0)
        dtr = _safe_float(first.get("d_trong", ""), 0.0)
        day_vo = _safe_float(first.get("day_vo", ""), 0.0)
        has_pile_geom = bool(first) and D_m > 0.0
        pile_type = ""
        if has_pile_geom:
            pile_type = "1" if abs(day_vo) < 1e-8 and dtr <= 1e-8 else "2"
        spacing = self._min_pile_spacing(piles, fallback=(max(3.0 * D_m, 3.0) if D_m > 0 else 0.0))
        # Số hàng cọc phục vụ hệ số nhóm lấy theo số tọa độ X khác nhau.
        row_count = self._row_count_from_pile_coords(piles)
        # Output classic TS-PILE/MCOC không có cao độ tuyệt đối EL mặt đất/đáy bệ/mũi cọc/MNN.
        ground_elev = ""
        cap_elev = ""
        tip_elev = ""
        water_elev = ""
        def combo_in(c, st):
            if not st:
                return False
            try:
                key = str(int(float(c)))
            except Exception:
                key = str(c).strip()
            return key in st
        def max_compression_force(st):
            vals = []
            for f in forces or []:
                if combo_in(f.get("combo", ""), st):
                    vals.append(max(_safe_float(f.get("N_ton", 0.0), 0.0), 0.0) * 9.80665)
            return max(vals) if vals else 0.0
        def max_uplift_force(st):
            vals = []
            for f in forces or []:
                if combo_in(f.get("combo", ""), st):
                    nton = _safe_float(f.get("N_ton", 0.0), 0.0)
                    if nton < 0.0:
                        vals.append(abs(nton) * 9.80665)
            return max(vals) if vals else 0.0
        pu_cd = max_compression_force(cd_set)
        pu_db = max_compression_force(db_set)
        uplift_cd = max_uplift_force(cd_set)
        uplift_db = max_uplift_force(db_set)
        try:
            self.item_uplift_data[_normalize_item_name(default_item)] = {"item": default_item, "uplift_cd": _fmt(uplift_cd, 0), "uplift_db": _fmt(uplift_db, 0)}
        except Exception:
            pass
        def max_cap(st):
            vals = []
            for c in data.get("Load_Combos", []):
                if combo_in(c.get("Name", ""), st):
                    vals.append(max(_safe_float(c.get("N_load", 0.0), 0.0), 0.0) * 9.80665)
            return max(vals) if vals else 0.0
        ncap_cd = max_cap(cd_set)
        ncap_db = max_cap(db_set)
        if ncap_cd <= 0 and pu_cd > 0:
            ncap_cd = pu_cd * n_piles
        if ncap_db <= 0 and pu_db > 0:
            ncap_db = pu_db * n_piles
        d_txt = _fmt(D_m * 1000.0, 0) if D_m > 0.0 else ""
        ds_txt = d_txt
        spacing_txt = _fmt(spacing, 3) if _safe_float(spacing, 0.0) > 0.0 else ""
        note_bits = [
            f"Import từ TS-PILE/MCOC: {default_item}",
            "lấy hạng mục từ dòng Công trình",
            "số cọc/số hàng/S từ tọa độ nếu đọc được",
            "Bx/By/Cz và N đáy bệ từ tải ban đầu",
            "lực nhổ đã đưa vào nút Điền lực nhổ",
            "cao độ/MNN nhập thủ công",
        ]
        if not has_pile_geom:
            note_bits.append("loại cọc và D/Ds để trống vì file không có thông số cọc rõ ràng")
        return [
            default_item, pile_type, str(n_piles), str(row_count), d_txt, ds_txt, spacing_txt,
            _fmt(g.get("Bx", ""), 3), _fmt(g.get("By", ""), 3), _fmt(g.get("Cz", ""), 3),
            ground_elev, cap_elev, tip_elev, water_elev,
            _fmt(pu_cd, 0), _fmt(pu_db, 0), _fmt(ncap_cd, 0), _fmt(ncap_db, 0),
            "; ".join(note_bits) + "."
        ]

    def _merge_n2d_input_and_output_data(self, input_data: Dict[str, Any], output_data: Dict[str, Any]) -> Dict[str, Any]:
        'Ghép dữ liệu INPUT MCOC và OUTPUT TS-PILE/MCOC nếu người dùng chọn file có đủ cả hai.\n\n        - Hình học cọc, Bx/By/Cz, tổ hợp tải trọng ban đầu ưu tiên lấy từ INPUT.\n        - Nội lực đầu cọc ưu tiên lấy từ OUTPUT.\n        - Nếu OUTPUT có đủ hình học/tải trọng thì vẫn dùng được độc lập.\n        '
        input_data = input_data or {}
        output_data = output_data or {}
        g_in = dict(input_data.get("Global", {}) or {})
        g_out = dict(output_data.get("Global", {}) or {})
        # Nếu file là output classic, output_data đã đọc đúng form. Không để parser INPUT
        # nhầm số trong report rồi ghi đè Bx/By/N đáy bệ/hình học.
        # Nếu là file INPUT số thuần thì output_data rỗng và input_data được dùng.
        if output_data.get("Load_Combos") or output_data.get("Piles") or output_data.get("forces"):
            g = {**g_in, **g_out}
            combos = output_data.get("Load_Combos") or input_data.get("Load_Combos") or []
            piles = output_data.get("Piles") or input_data.get("Piles") or []
        else:
            g = {**g_out, **g_in}
            combos = input_data.get("Load_Combos") or output_data.get("Load_Combos") or []
            piles = input_data.get("Piles") or output_data.get("Piles") or []
        if not str(g.get("Project_ID", "")).strip():
            g["Project_ID"] = g_out.get("Project_ID", "") or g_in.get("Project_ID", "") or input_data.get("source", "") or output_data.get("source", "")
        return {
            "Global": g,
            "Load_Combos": combos,
            "Piles": piles,
            "forces": output_data.get("forces", []) or input_data.get("forces", []),
            "source": output_data.get("source", "") or input_data.get("source", ""),
        }

    def _try_parse_input_file_quiet(self, path: str) -> Dict[str, Any]:
        try:
            return self._parse_foundation_input_file_for_sct(path)
        except Exception:
            return {}

    def import_from_n2d_pile_result(self):
        paths = filedialog.askopenfilenames(
            title='Import từ file kết quả/INPUT TS-PILE/MCOC',
            filetypes=[('TS-PILE/MCOC output/input', "*.out *.doc *.txt *.dat *.csv *.xlsx *.xlsm *.xls"), ("File không đuôi", "*"), ("All files", "*.*")]
        )
        if not paths:
            return
        cd_text = simpledialog.askstring("Khoảng tổ hợp TTGHCĐ", "Từ tổ hợp nào đến tổ hợp nào là CƯỜNG ĐỘ?\nVí dụ: 1-12 hoặc 1,2,3,4", initialvalue="1-999", parent=self.root)
        if cd_text is None:
            return
        db_text = simpledialog.askstring("Khoảng tổ hợp TTGHĐB", "Từ tổ hợp nào đến tổ hợp nào là ĐẶC BIỆT?\nVí dụ: 13-20 hoặc 13,14,15", initialvalue="", parent=self.root)
        if db_text is None:
            return
        cd_set = self._parse_combo_range_text(cd_text)
        db_set = self._parse_combo_range_text(db_text)
        mapped = []
        warn = []
        for path in paths:
            stem = os.path.splitext(os.path.basename(path))[0]
            ext = os.path.splitext(path)[1].lower()
            try:
                if ext in (".csv", ".xlsx", ".xlsm"):
                    # CSV/XLSX xuất từ TS-COL/TS-PILE: dùng cột N/Pu, Combo, Item nếu có.
                    # Không có hình học nên chỉ lấy Pu/cọc; user bổ sung hình học/cao độ.
                    forces, _ = self._parse_forces_csv_for_sct(path)
                    if not forces:
                        warn.append(f"{stem}: không đọc được cột lực dọc N/Pu từ CSV/XLSX.")
                        continue
                    def combo_key(c):
                        try:
                            return str(int(float(c)))
                        except Exception:
                            return str(c).strip()
                    def max_force_from_csv(st):
                        vals = []
                        for f in forces:
                            if combo_key(f.get("combo", "")) in st:
                                vals.append(max(_safe_float(f.get("N_ton"), 0.0), 0.0) * 9.80665)
                        return max(vals) if vals else 0.0
                    pu_cd = max_force_from_csv(cd_set)
                    pu_db = max_force_from_csv(db_set)
                    item_name = _display_item_name(forces[0].get("item", "")) or stem
                    mapped.append([item_name, "", "", "", "", "", "", "", "", "", "", "", "", "", _fmt(pu_cd,0), _fmt(pu_db,0), "", "", "Import CSV/XLSX lực: bổ sung loại cọc/hình học/cao độ/N đáy bệ."])
                    warn.append(f"{stem}: chỉ đọc được Pu/cọc từ bảng lực; chưa có hình học và N đáy bệ ban đầu.")
                    continue

                # TXT/DOC/OUT/file không đuôi: ưu tiên đọc OUTPUT classic theo form TS-PILE.
                # Chỉ khi không nhận ra output mới thử đọc như INPUT số thuần.
                output_data = self._try_parse_text_report_for_sct(path) if ext in (".txt", ".doc", ".out", ".dat", "") else {}
                input_data = {}
                if not (output_data.get("Load_Combos") or output_data.get("Piles") or output_data.get("forces")):
                    input_data = self._try_parse_input_file_quiet(path) if ext in (".txt", ".doc", ".out", ".dat", "") else {}
                data = self._merge_n2d_input_and_output_data(input_data, output_data)
                forces = data.get("forces", []) or []
                item_name = str(data.get("Global", {}).get("Project_ID", "") or output_data.get("source", "") or input_data.get("source", "") or stem).strip()
                if not (data.get("Piles") or forces or data.get("Load_Combos")):
                    warn.append(f"{stem}: không nhận dạng được INPUT/OUTPUT MCOC.")
                    continue
                mapped.append(self._item_row_from_n2d_data(data, forces, cd_set, db_set, item_name))
                if not forces:
                    warn.append(f"{stem}: chưa đọc được bảng NỘI LỰC ĐẦU CỌC; Pu/cọc để 0 hoặc cần import output có bảng nội lực.")
                if not data.get("Load_Combos"):
                    warn.append(f"{stem}: chưa đọc được bảng tổ hợp tải ban đầu; N đáy bệ để 0 hoặc lấy tạm từ Pu*cọc.")
                if not data.get("Piles"):
                    warn.append(f"{stem}: chưa đọc được thông số/tọa độ cọc; cần bổ sung hình học thủ công.")
            except Exception as exc:
                warn.append(f"{stem}: không import được ({exc})")
        if not mapped:
            messagebox.showerror('Import từ TS-PILE/MCOC', "Không đọc được hạng mục nào.\n" + "\n".join(warn[:12]))
            return
        existing_rows = list(self.item_table.get_rows() or [])
        self.item_table.set_rows(existing_rows + mapped, record_undo=True)
        msg = f"Đã thêm {len(mapped)} hạng mục từ TS-PILE/MCOC."
        if warn:
            msg += "\n\nCảnh báo:\n" + "\n".join(warn[:12])
        messagebox.showinfo('Import từ TS-PILE/MCOC', msg)
        self._set_status(f"Đã thêm {len(mapped)} hạng mục từ TS-PILE/MCOC")

    def import_items_file(self):
        path = filedialog.askopenfilename(title="Import hạng mục CSV/TXT", filetypes=[("CSV/TXT/Excel", "*.csv *.txt *.xlsx *.xlsm"), ("All files", "*.*")])
        if not path:
            return
        rows = self._read_delimited_file(path)
        if not rows:
            return
        header = [h.strip().lower() for h in rows[0]]
        data = rows[1:] if any(k in "|".join(header) for k in ["hang", "hạng", "item", "d", "tip"]) else rows
        mapped = []
        joined_header = "|".join(header)
        for r in data:
            raw = list(r)
            # Tương thích template cũ có 2 cột lực nhổ nằm sau Pu ĐB/cọc.
            if len(raw) >= 21 or ("nhổ" in joined_header or "nho" in _strip_accents(joined_header)):
                old_vals = (raw + [""] * 21)[:21]
                item_name = str(old_vals[0]).strip()
                if item_name:
                    self.item_uplift_data[_normalize_item_name(item_name)] = {
                        "item": item_name,
                        "uplift_cd": str(old_vals[16]).strip(),
                        "uplift_db": str(old_vals[17]).strip(),
                    }
                vals = old_vals[:16] + old_vals[18:21]
            else:
                vals = (raw + [""] * 19)[:19]
            mapped.append(vals)
        existing_rows = list(self.item_table.get_rows() or [])
        self.item_table.set_rows(existing_rows + mapped, record_undo=True)
        self._set_status(f"Đã thêm {len(mapped)} hạng mục từ {os.path.basename(path)}")

    def _borehole_item_name_from_path(self, path: str) -> str:
        base = os.path.splitext(os.path.basename(str(path or "")))[0]
        return base.strip() or "Lo_khoan"


    def _borehole_page_text(self, text: str, page_no: int) -> str:
        """Lấy phần OCR của một trang từ text OCR toàn PDF/ảnh nhiều trang.

        Nếu text đã có marker --- PAGE n --- mà không tìm thấy trang yêu cầu thì trả rỗng,
        không trả toàn bộ text. Cách cũ có thể làm trang 2 bị lấy nhầm nội dung trang 1 khi OCR
        của trang 2 thất bại hoặc bị bỏ qua.
        """
        s = str(text or "")
        m = re.search(rf"---\s*PAGE\s+{int(page_no)}\s*---\s*(.*?)(?=\n---\s*PAGE\s+\d+\s*---|\Z)", s, re.I | re.S)
        if m:
            return m.group(1)
        if re.search(r"---\s*PAGE\s+\d+\s*---", s, re.I):
            return ""
        return s

    def _borehole_file_page_count(self, path: str) -> int:
        """Đếm số trang/frame của PDF/TIFF/ảnh để kiểm soát import nhiều trang."""
        ext = os.path.splitext(str(path or ""))[1].lower()
        if ext == ".pdf":
            try:
                import fitz  # type: ignore
                doc = fitz.open(path)
                n = int(doc.page_count)
                doc.close()
                return max(n, 1)
            except Exception:
                try:
                    from pdf2image import pdfinfo_from_path  # type: ignore
                    info = pdfinfo_from_path(path)
                    return max(int(info.get("Pages", 1)), 1)
                except Exception:
                    return 1
        try:
            from PIL import Image
            img = Image.open(path)
            return max(int(getattr(img, "n_frames", 1) or 1), 1)
        except Exception:
            return 1

    def _make_ocr_rows_for_one_borehole_layer(
        self,
        item_name: str,
        lname: str,
        top_depth: float,
        bottom_depth: float,
        top_elev: Optional[float],
        stype: Any,
        spt_points: List[Tuple[float, float]],
        desc: str,
        bottom_elev_override: Optional[float] = None,
    ) -> List[List[Any]]:
        """Tạo dòng địa chất từ một lớp OCR.

        - Chế độ trung bình: một dòng/lớp, N là trung bình các điểm SPT trong lớp.
        - Chế độ từng điểm: chia lớp thành các đoạn đại diện cho từng điểm SPT theo khoảng cách nhập.
          Vì bảng địa chất chỉ có cột cao độ đáy lớp, ranh giới đoạn được lấy tại trung điểm giữa
          các điểm SPT; dòng cuối luôn khép về đúng đáy lớp gốc.
        """
        top_depth = float(top_depth or 0.0)
        bottom_depth = float(bottom_depth or top_depth)
        thickness = max(bottom_depth - top_depth, 0.0)
        lname_s = str(lname or "").strip()
        desc_s = str(desc or "")
        desc_low = _strip_accents(desc_s).lower()
        is_fill = lname_s.upper() == "D" or "dat dap" in desc_low or "dat phu" in desc_low

        def elev_at_depth(d: float) -> float:
            # Nếu đã đọc được trực tiếp cao độ đáy lớp từ cột bảng, dùng giá trị đó cho đáy lớp.
            # Không dịch cả trang theo cao độ miệng lỗ khoan vì OCR header có thể sai 10.00 -> 1.00.
            try:
                if bottom_elev_override is not None and abs(float(d) - float(bottom_depth)) <= 0.05:
                    return float(bottom_elev_override)
            except Exception:
                pass
            return (float(top_elev) - float(d)) if top_elev is not None else -float(d)

        def base_row(bottom_d: float, n_value: Any, note: str = "") -> List[Any]:
            comment = desc_s
            n_value2, spt_note = self._spt_import_row_value(n_value)
            if note:
                comment = (comment + " | " + note).strip(" |")
            if spt_note:
                comment = (comment + " | " + spt_note).strip(" |")
            n_txt = ""
            try:
                if str(n_value2).strip() != "":
                    n_txt = str(int(round(float(n_value2))))
            except Exception:
                n_txt = str(n_value2 or "").strip()
            return [
                item_name, lname_s, _fmt(elev_at_depth(bottom_d), 2), str(stype),
                n_txt,
                "", "", "", "", "", "", "", "", comment,
            ]

        # V1.0.27: chọn SPT theo độ sâu điểm thực đọc được, không sinh các điểm lý thuyết
        # rồi gán nearest. Cách cũ làm chế độ "lấy đủ điểm SPT" bị mất/đổi điểm và sau đó
        # còn bị cleanup gộp lại thành một dòng/lớp. Cho phép dung sai nhỏ ở đáy lớp để giữ
        # các mẫu SPT nằm đúng tại ranh giới đáy lỗ/lớp.
        layer_points = sorted(
            [(float(d), float(n)) for d, n in (spt_points or []) if (top_depth + 1e-6) < float(d) <= (bottom_depth + 0.30)],
            key=lambda x: x[0],
        )
        if is_fill:
            return [base_row(bottom_depth, "")]

        mode = str(getattr(self, "_borehole_spt_mode", "avg") or "avg")
        if mode != "points":
            vals = [n for d, n in layer_points if d <= bottom_depth + 0.30]
            n_avg = round(sum(vals) / len(vals)) if vals else ""
            return [base_row(bottom_depth, n_avg)]

        try:
            spacing = float(getattr(self, "_borehole_spt_spacing_m", 2.0) or 2.0)
        except Exception:
            spacing = 2.0
        if spacing <= 0:
            spacing = 2.0
        if thickness <= 0.05 or not layer_points:
            vals = [n for _, n in layer_points]
            n_avg = round(sum(vals) / len(vals)) if vals else ""
            return [base_row(bottom_depth, n_avg)]

        # Một dòng cho mỗi điểm SPT thật trong lớp. Ranh giới đáy từng đoạn lấy theo trung
        # điểm giữa hai điểm SPT liên tiếp; đoạn cuối khép về đúng đáy lớp gốc để không làm
        # thay đổi địa tầng. Khoảng cách nhập dùng để giới hạn ảnh hưởng về phía dưới khi
        # hai điểm cách nhau bất thường xa nhau.
        point_depths = [max(top_depth + 1e-4, min(float(d), bottom_depth)) for d, _n in layer_points]
        vals = [n for _d, n in layer_points]
        boundaries: List[float] = []
        last_bd = top_depth
        for i, pd in enumerate(point_depths):
            if i < len(point_depths) - 1:
                mid = (pd + point_depths[i + 1]) / 2.0
                bd = min(bottom_depth, mid, pd + spacing / 2.0)
                if bd <= last_bd + 0.03:
                    bd = min(bottom_depth, max(last_bd + 0.03, mid))
            else:
                bd = bottom_depth
            boundaries.append(float(bd))
            last_bd = float(bd)

        out: List[List[Any]] = []
        for i, bd in enumerate(boundaries):
            out.append(base_row(bd, vals[i] if i < len(vals) else "", ""))
        return out

    def _normalize_ocr_text(self, text: Any) -> str:
        # OCR có thể trả về nhiều kiểu ký tự gần giống nhau; chuẩn hóa nhẹ để parse ổn định hơn.
        s = str(text or "")
        s = s.replace("–", "-").replace("—", "-").replace("−", "-")
        s = s.replace("O", "0").replace("o", "0") if False else s  # không thay chữ o toàn cục để tránh hỏng tiếng Việt.
        s = re.sub(r"(?<=\d),(?=\d)", ".", s)
        s = re.sub(r"[ \t]+", " ", s)
        return s

    def _safe_float_ocr(self, value: Any, default: float = 0.0) -> float:
        s = str(value or "").strip()
        s = s.replace(",", ".")
        s = re.sub(r"[^0-9+\-.]", "", s)
        if not s or s in {"-", "+", "."}:
            return float(default)
        try:
            return float(s)
        except Exception:
            return float(default)

    def _safe_int_ocr(self, value: Any, default: int = 0) -> int:
        s = str(value or "").strip()
        if ">" in s and "100" in s:
            return 100
        nums = re.findall(r"-?\d+", s)
        if not nums:
            return int(default)
        try:
            return int(nums[-1])
        except Exception:
            return int(default)

    def _extract_borehole_top_elev(self, text: str) -> Optional[float]:
        """Đọc cao độ miệng lỗ khoan từ OCR header/text PDF.

        OCR của các form lỗ khoan rất hay dính chữ và dấu, ví dụ:
        ``CAO ĐỘ LỖ KHOAN :.16.100M`` hoặc ``CAO DO LO KHOAN:35.77``.
        Bản cũ bắt regex quá chặt nên nhiều ảnh bị mất cao độ miệng, kéo theo cao độ đáy
        bị suy sai khi cột cao độ không OCR được. Hàm này chuẩn hóa bỏ dấu và cho phép
        ký tự nhiễu ngắn giữa nhãn và số.
        """
        raw = str(text or "")
        if not raw.strip():
            return None
        norm = _strip_accents(raw)
        norm = norm.replace("\u00a0", " ")
        norm = re.sub(r"[，；;]", ":", norm)
        # Không xóa dấu chấm trước số vì OCR thường trả ":.16.100"; xử lý bằng regex dưới.
        pats = [
            r"cao\s*do\s*lo\s*khoan\s*[^0-9+\-]{0,12}([+\-]?(?:\d{1,3}(?:[\.,]\d{1,3})?|[\.,]\d{1,3}))\s*m?\b",
            r"ca[o0]\s*d[o0]\s*l[o0]\s*khoan\s*[^0-9+\-]{0,12}([+\-]?(?:\d{1,3}(?:[\.,]\d{1,3})?|[\.,]\d{1,3}))\s*m?\b",
            r"cao\s*do\s*lk\s*[^0-9+\-]{0,12}([+\-]?(?:\d{1,3}(?:[\.,]\d{1,3})?|[\.,]\d{1,3}))\s*m?\b",
        ]
        for pat in pats:
            m = re.search(pat, norm, flags=re.I)
            if not m:
                continue
            token = str(m.group(1) or "").strip()
            # OCR có thể cho :.16.100M, regex lấy ".16". Nếu ngay sau match còn ".100",
            # nối lại để thành 16.100 thay vì 0.16.
            tail = norm[m.end(1):m.end(1)+6]
            if token.startswith("."):
                # Tìm toàn bộ cụm số ngay sau nhãn để bắt trường hợp :.16.100
                prefix = norm[max(0, m.start(1)-3):m.end(1)+8]
                m2 = re.search(r"[+\-]?\.?\d{1,3}(?:[\.,]\d{1,3})?", prefix)
                if m2:
                    token = m2.group(0).lstrip(".")
            token = token.replace(",", ".")
            try:
                val = float(token)
                if -100.0 <= val <= 300.0:
                    return val
            except Exception:
                pass
        # Fallback: tìm dòng có cả nhãn và lấy số hợp lý cuối dòng.
        for line in norm.splitlines():
            low = re.sub(r"\s+", " ", line.lower())
            if "cao" in low and "khoan" in low and ("lo" in low or "l0" in low):
                nums = self._ocr_numeric_candidates(line)
                vals = [float(v) for v in nums if -100.0 <= float(v) <= 300.0]
                if vals:
                    # Tránh lấy tỷ lệ 1/100; ưu tiên số có phần thập phân hoặc nằm trước chữ M.
                    vals2 = [v for v in vals if abs(v) > 2.0]
                    return float(vals2[0] if vals2 else vals[0])
        return None

    def _rqd_value_from_borehole_text(self, text: Any) -> Optional[float]:
        """Đọc RQD (%) từ mô tả OCR nếu có.

        Hỗ trợ các dạng thường gặp: ``RQD=0``, ``RQD = 15%``, ``RQD<20``, ``R.Q.D: 25``.
        Chỉ dùng giá trị trong khoảng 0..100 để tránh nhầm số độ sâu/SPT.
        """
        raw = _strip_accents(str(text or "")).upper().replace("%", " ")
        raw = raw.replace("O", "0")
        pats = [
            r"R\s*\.?\s*Q\s*\.?\s*D\s*(?:=|:|<|>|≤|≥|<=|>=)?\s*([0-9]{1,3}(?:[\.,][0-9]+)?)",
            r"RQD\s*(?:=|:|<|>|≤|≥|<=|>=)?\s*([0-9]{1,3}(?:[\.,][0-9]+)?)",
        ]
        for pat in pats:
            m = re.search(pat, raw, re.I)
            if not m:
                continue
            try:
                val = float(str(m.group(1)).replace(",", "."))
                if 0.0 <= val <= 100.0:
                    return val
            except Exception:
                pass
        return None

    def _infer_soil_type_from_borehole_text(self, layer_name: Any, description: Any) -> int:
        """Gợi ý loại đất đá, không phân biệt hoa/thường/dấu; người dùng vẫn duyệt lại ở bảng preview.

        V1.0.11 ưu tiên ký hiệu lớp và cụm mô tả địa chất "Đá..." trước khi xét các từ
        như "sét", "dăm sạn" bị OCR lẫn trong phần mô tả lớp kề bên. Với log lỗ khoan
        kiểu BRITEC, các lớp 8a/8b mô tả "Đá vôi, phong hóa, nứt nẻ..." được đưa về
        loại 4 = Đá nứt vỡ/phong hóa, không tự ép thành đất dính hoặc IGM.
        """
        lname_norm = re.sub(r"\s+", "", _strip_accents(layer_name or "").lower())
        raw = f"{layer_name or ''} {description or ''}"
        s_desc = re.sub(r"\s+", " ", _strip_accents(description or "").lower()).strip()
        s = _strip_accents(raw).lower()
        s = re.sub(r"\s+", " ", s).strip()
        rqd_val = self._rqd_value_from_borehole_text(raw)

        # QA-GEO RQD: nếu mô tả OCR đọc được RQD của lớp đá phong hóa thì dùng RQD để phân loại.
        # RQD<20: IGM/đá mềm (loại 5). RQD>=20: đá phong hóa/nứt vỡ (loại 4).
        has_weathered_rock_by_text = any(k in s for k in ["da ", "da voi", "rock", "limestone", "phong hoa", "nut vo", "nut ne", "da set bot ket", "cat ket"])
        if rqd_val is not None and has_weathered_rock_by_text:
            if any(k in s for k in ["nguyen khoi", "da nguyen", "intact"]) and rqd_val >= 20.0:
                return 3
            return 5 if rqd_val < 20.0 else 4

        # Ký hiệu lớp có độ tin cậy cao hơn mô tả OCR bị lẫn dòng kế bên.
        if "hang" in lname_norm or "karst" in lname_norm:
            return 0
        if lname_norm in ("dc", "d", "1"):
            return 2
        if lname_norm in ("cs",):
            return 1
        # Quy ước lớp thường gặp trong bộ log LKBC/LKKN/R3.
        # Đặt trước các từ khóa mô tả như "rất chặt" để không nhầm cát bụi 5B thành đá.
        if lname_norm in ("5a",):
            return 2
        if lname_norm in ("5b",):
            return 1
        if lname_norm in ("6", "6a", "6b", "7", "7a", "7b"):
            if rqd_val is not None:
                return 5 if rqd_val < 20.0 else 4
            return 4
        if lname_norm in ("4", "4a", "4b") and any(k in s for k in ["cuoi", "soi", "gravel", "cobble"]):
            return 6
        # Form GEOTEST/BRITEC trong hồ sơ cầu thường dùng 8/8a/8b cho đá vôi phong hóa/nứt nẻ.
        # Ưu tiên ký hiệu lớp vì mô tả OCR của các lớp mỏng dễ bị dính sang lớp sét/hang bên cạnh.
        if lname_norm in ("8", "8a", "8b"):
            if rqd_val is not None:
                return 5 if rqd_val < 20.0 else 4
            return 4
        # Một số form AutoCAD cầu miền Tây dùng 6A/6B/7 cho đá/sét bột kết, cát kết phong hóa.
        # Khi mô tả OCR không đọc được thì vẫn không nên rơi về đất sét mặc định.
        if lname_norm in ("6", "6a", "6b", "7", "7a", "7b") and not any(k in s_desc for k in ["dat", "cat", "set", "clay", "sand", "cuoi", "soi"]):
            return 4
        if lname_norm in ("4", "4a", "4b") and not s_desc:
            return 6
        if lname_norm in ("5a",) and not s_desc:
            return 2
        if lname_norm in ("5b",) and not s_desc:
            return 1
        if lname_norm in ("2", "2a", "3") and any(k in s for k in ["set", "clay", "bui", "silt", "bun"]):
            return 2

        # Các lớp đá theo ký hiệu phổ biến trong log cầu đường. Chỉ cần có mô tả đá/phong hóa/nứt
        # là đưa về loại 4; người dùng vẫn có thể chỉnh hàng loạt sang IGM nếu dự án quy ước khác.
        is_named_rock_layer = bool(re.fullmatch(r"(?:\d+[a-z]?|[a-z]?\d+[a-z]?)", lname_norm or "")) and any(
            k in s for k in ["da ", "da voi", "phong hoa", "nut", "nứt", "rock", "limestone"]
        )
        if is_named_rock_layer or re.match(r"^\s*(da|rock|limestone)\b", s_desc):
            if any(k in s for k in ["nguyen khoi", "da nguyen", "intact"]):
                return 3
            return 4

        # Ưu tiên cụm đầu mô tả: nhiều form ghi "Sét ... đôi chỗ lẫn dăm sạn"; không được nhầm thành IGM.
        if re.match(r"^(cat|sand)", s_desc) or re.match(r"^cat\s+set", s_desc):
            return 1
        if re.match(r"^(set|clay|bun|bui|silt)", s_desc):
            return 2

        # loại 0: rỗng/không khí/karst/mất dung dịch/hang
        if any(k in s for k in ["khong khi", "hang", "karst", "rong", "lo rong", "mat dung dich", "mat nuoc", "cavity"]):
            return 0

        # đá/cuội sỏi xét trước cát/sét để không bị nhầm do mô tả hỗn hợp.
        if any(k in s for k in ["cuoi", "soi", "gravel", "cobble", "cuoi ket", "dam ket"]):
            return 6
        if any(k in s for k in ["da nguyen", "nguyen khoi", "granite", "diorite", "basalt", "quartzite", "rhyolite"]):
            return 3

        # Có chữ "đá" rõ ràng thì giữ là đá nứt vỡ/phong hóa. Chỉ coi IGM khi mô tả thiên về đất đá/dăm sạn/spph
        # mà không phải một lớp đá vôi/đá gốc rõ ràng.
        has_explicit_rock = any(k in s for k in ["da voi", "da ", "rock", "limestone", "phong hoa", "nut vo", "nut ne"])
        if has_explicit_rock:
            if any(k in s for k in ["spph", "dat da phong hoa", "dam san", "da dam"]) and not any(k in s for k in ["da voi", "limestone"]):
                return 5
            return 4

        if any(k in s for k in ["da nut", "nut vo", "phong hoa", "dam manh", "da dam", "rat chat", "spph", "dam san", "igneous"]):
            if any(k in s for k in ["dam san", "spph", "da dam"]):
                return 5
            return 4

        # đất hạt rời/dính
        has_sand = any(k in s for k in ["cat", "sand"])
        has_clay = any(k in s for k in ["set", "clay", "bun", "silt", "bui"])
        if has_sand and not has_clay:
            return 1
        if has_clay and not has_sand:
            return 2
        if has_sand and has_clay:
            # Mô tả "cát sét" thường vẫn ứng xử như đất rời lẫn hạt mịn; để type=1, người dùng có thể chỉnh.
            return 1
        if any(k in s for k in ["dat dap", "dat phu", "nen mat duong", "mat duong", "topsoil", "fill"]):
            return 2
        return 2


    def _pdf_extractable_text_pages(self, path: str, min_chars: int = 30) -> List[str]:
        """Trích text trực tiếp từ PDF nếu là PDF digital/text-based.

        Mục tiêu: không render mọi PDF thành ảnh rồi OCR. Với PDF xuất từ CAD/Excel/Word,
        text trực tiếp thường chính xác hơn OCR và giữ tốt các số cao độ/SPT. Nếu một trang
        không có text đáng kể thì trả chuỗi rỗng để các bước ảnh/OCR vẫn xử lý trang đó.
        """
        pages: List[str] = []
        ext = os.path.splitext(str(path or ""))[1].lower()
        if ext != ".pdf":
            return pages
        try:
            import fitz  # type: ignore
            doc = fitz.open(path)
            for page in doc:
                txt = page.get_text("text") or ""
                pages.append(txt if len(txt.strip()) >= min_chars else "")
            doc.close()
            if any(t.strip() for t in pages):
                return pages
        except Exception:
            pages = []
        try:
            import pdfplumber  # type: ignore
            out: List[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    txt = page.extract_text() or ""
                    out.append(txt if len(txt.strip()) >= min_chars else "")
            return out
        except Exception:
            return pages


    def _pdf_vector_tokenize_borehole_text(self, text: str) -> List[str]:
        """Tách token từ text PDF vector mà không phụ thuộc Unicode tiếng Việt.

        Các PDF xuất từ CAD có thể bị encoding kiểu VNI/TCVN3 nên chỉ dùng token số,
        layer name và nhãn SPT. Không dùng mô tả dài để tránh lệch cột.
        """
        s = self._normalize_ocr_text(text)
        s = s.replace("−", "-").replace("–", "-").replace("—", "-")
        # Giữ token >50, -22.82, 5A, K1, HK5-1, SPT31. Bỏ phần lớn dấu tiếng Việt vì không cần cho 3 cột lõi.
        return re.findall(r"\d+[A-Za-zĐđ]|[A-Za-zĐđ_][A-Za-zĐđ0-9_\-/]{0,20}|>?[-+]?\d+(?:[\.,]\d+)?", s)

    def _pdf_vector_is_number_token(self, token: Any) -> bool:
        return bool(re.fullmatch(r">?[-+]?\d+(?:[\.,]\d+)?", str(token or "").strip()))

    def _pdf_vector_float_token(self, token: Any, default: float = 0.0) -> float:
        s = str(token or "").strip().replace(",", ".")
        s = s.replace(">", "")
        return self._safe_float_ocr(s, default)

    def _pdf_vector_is_layer_name_token(self, token: Any) -> bool:
        """Nhận token có thể là ký hiệu lớp địa chất trong text stream PDF.

        Cho phép lớp dạng số 1/2/3, 5A, 6B, K1, D, DC, CS, Hang..., nhưng loại nhãn SPT,
        tên mẫu HK5-1 và các token đơn vị/header rõ ràng.
        """
        t = str(token or "").strip()
        if not t:
            return False
        low = _strip_accents(t).lower()
        if low.startswith("spt") or re.fullmatch(r"hk\d+[-_/]\d+", low):
            return False
        if low in {"m", "cm", "mm", "n", "x", "y", "ty", "le", "cao", "do", "sau", "lop", "be", "day"}:
            return False
        if re.fullmatch(r"\d{1,2}[A-Za-zĐđ]?", t):
            return True
        return bool(re.fullmatch(r"[A-Za-zĐđ][A-Za-zĐđ0-9_/\-]{0,8}", t))

    def _parse_spt_points_from_pdf_vector_text(self, text: str) -> List[Tuple[float, float]]:
        """Đọc điểm SPT từ text PDF vector dạng:
        1.8 - 2.0  HK5-1  2.0 - 2.45  0 0 0 0 0  SPT1

        Trả về (độ sâu đại diện, N/30). Với >50 lấy 50, >100 lấy 100.
        """
        s = self._normalize_ocr_text(text)
        num = r">?\d+(?:[\.,]\d+)?"
        pat = re.compile(
            rf"({num})\s*-\s*({num})\s+[^\s]+\s+({num})\s*-\s*({num})\s+((?:{num}\s+){{2,6}})SPT\s*\d+",
            re.I,
        )
        points: List[Tuple[float, float]] = []
        for m in pat.finditer(s):
            try:
                d0 = self._pdf_vector_float_token(m.group(3), 0.0)
                d1 = self._pdf_vector_float_token(m.group(4), d0)
                vals = re.findall(num, m.group(5) or "")
                if not vals:
                    continue
                last = vals[-1]
                if ">" in str(last):
                    nval = 100.0 if self._pdf_vector_float_token(last, 0.0) >= 100.0 else self._pdf_vector_float_token(last, 50.0)
                else:
                    nval = self._pdf_vector_float_token(last, 0.0)
                # V1.0.27: với PDF vector, cột SPT đã có khoảng "2.0 - 2.45".
                # Dùng độ sâu bắt đầu mẫu SPT để không bỏ mất điểm nằm tại đáy lớp/lỗ khoan
                # như mẫu 70.0 - 70.45; dùng tâm mẫu trước đây làm điểm 70.22 nên bị loại
                # khỏi lớp có đáy 70.0.
                depth = float(d0)
                if 0.0 <= depth <= 220.0 and 0.0 <= nval <= 500.0:
                    points.append((round(depth, 3), float(nval)))
            except Exception:
                continue
        dedup: Dict[float, float] = {}
        for d, n in points:
            dedup[round(d, 2)] = n
        return sorted(dedup.items(), key=lambda x: x[0])

    def _parse_borehole_layers_from_pdf_vector_text(self, text: str, item_name: str) -> List[Dict[str, Any]]:
        """Đọc lớp địa chất từ text stream PDF vector theo chuỗi token.

        Nhánh này phục vụ các PDF CAD/digital mà text không nằm theo từng dòng bảng,
        nhưng trình tự token vẫn là:
            top_elev, top_depth, layer, thickness, bottom_elev, bottom_depth,
            layer, thickness, bottom_elev, bottom_depth, ...

        Chỉ lấy 3 dữ liệu lõi: tên lớp, cao độ đáy, độ sâu đáy/bề dày. Không lấy mô tả.
        """
        toks = self._pdf_vector_tokenize_borehole_text(text)
        if len(toks) < 8:
            return []

        def isnum(i: int) -> bool:
            return 0 <= i < len(toks) and self._pdf_vector_is_number_token(toks[i])

        def f(i: int, default: float = 0.0) -> float:
            return self._pdf_vector_float_token(toks[i], default)

        start = -1
        top_elev = None
        # Tìm record đầu: Cao độ đầu, độ sâu 0.00, tên lớp, bề dày, cao độ đáy, độ sâu đáy.
        for i in range(0, min(len(toks) - 5, 250)):
            if not (isnum(i) and isnum(i + 1) and self._pdf_vector_is_layer_name_token(toks[i + 2]) and isnum(i + 3) and isnum(i + 4) and isnum(i + 5)):
                continue
            te = f(i)
            td = f(i + 1)
            th = f(i + 3)
            be = f(i + 4)
            bd = f(i + 5)
            if abs(td) > 0.20:
                continue
            if not (0.01 <= th <= 80.0 and 0.01 <= bd <= 200.0):
                continue
            # Kiểm tra quan hệ cao độ = cao độ đầu - độ sâu. Nới dung sai cho PDF làm tròn.
            if abs((te - bd) - be) <= max(0.30, 0.015 * max(abs(bd), 1.0)):
                start = i
                top_elev = te
                break
        if start < 0 or top_elev is None:
            return []

        layers: List[Dict[str, Any]] = []
        name = str(toks[start + 2]).strip()
        thickness = f(start + 3)
        bottom_elev = f(start + 4)
        bottom_depth = f(start + 5)
        layers.append({
            "item": item_name,
            "name": name,
            "bottom_elev": bottom_elev,
            "bottom_depth": bottom_depth,
            "thickness": thickness,
            "desc": "",
        })
        prev_depth = bottom_depth
        j = start + 6
        # Các lớp sau: tên lớp, bề dày, cao độ đáy, độ sâu đáy.
        while j + 3 < len(toks):
            if not (self._pdf_vector_is_layer_name_token(toks[j]) and isnum(j + 1) and isnum(j + 2) and isnum(j + 3)):
                break
            lname = str(toks[j]).strip()
            th = f(j + 1)
            be = f(j + 2)
            bd = f(j + 3)
            if not (0.01 <= th <= 100.0 and prev_depth + 0.01 <= bd <= 220.0):
                break
            # Dung sai bề dày và cao độ để chấp nhận làm tròn, nhưng không cho parse lấn sang bảng SPT.
            thick_ok = abs((bd - prev_depth) - th) <= max(0.35, 0.03 * max(th, 1.0))
            elev_ok = abs((float(top_elev) - bd) - be) <= max(0.50, 0.02 * max(abs(bd), 1.0))
            if not (thick_ok or elev_ok):
                break
            layers.append({
                "item": item_name,
                "name": lname,
                "bottom_elev": be,
                "bottom_depth": bd,
                "thickness": th,
                "desc": "",
            })
            prev_depth = bd
            j += 4
        if len(layers) < 2:
            return []
        layers.sort(key=lambda x: float(x.get("bottom_depth") or 0.0))
        return layers

    def _try_make_geology_rows_from_pdf_vector_text(self, path: str, item_name: str) -> Tuple[List[List[Any]], str, str]:
        """Nhánh import riêng cho PDF digital/vector.

        Nếu đọc được lớp bằng text stream PDF, ưu tiên nhánh này trước render ảnh/OCR vì số liệu
        cao độ, độ sâu và SPT thường chính xác hơn OCR. Không dùng mô tả dài để tránh sai lệch cột.
        """
        if os.path.splitext(str(path or ""))[1].lower() != ".pdf":
            return [], "", ""
        try:
            import fitz  # type: ignore
        except Exception:
            return [], "", ""
        rows: List[List[Any]] = []
        text_pages: List[str] = []
        warnings: List[str] = []
        try:
            doc = fitz.open(path)
            page_count = int(doc.page_count)
            for page_idx, page in enumerate(doc, start=1):
                txt = page.get_text("text") or ""
                text_pages.append(txt)
                if len(txt.strip()) < 80:
                    continue
                page_item_name = item_name if page_count <= 1 else f"{item_name}_p{page_idx:02d}"
                layers = self._parse_borehole_layers_from_pdf_vector_text(txt, page_item_name)
                if not layers:
                    continue
                spt_points = self._parse_spt_points_from_pdf_vector_text(txt)
                prev_depth = 0.0
                page_rows: List[List[Any]] = []
                page_top_elev_cands: List[float] = []
                for lay in layers:
                    try:
                        bottom_depth = float(lay.get("bottom_depth") or 0.0)
                        bottom_elev = float(lay.get("bottom_elev") or 0.0)
                    except Exception:
                        continue
                    lname = str(lay.get("name") or "").strip()
                    stype = self._infer_soil_type_from_borehole_text(lname, "")
                    top_elev = bottom_elev + bottom_depth
                    page_top_elev_cands.append(float(top_elev))
                    # Không đưa mô tả dài từ PDF vector vào bảng vì text theo cột dễ tràn/lệch.
                    page_rows.extend(self._make_ocr_rows_for_one_borehole_layer(
                        page_item_name, lname, prev_depth, bottom_depth, top_elev, stype, spt_points, "", bottom_elev_override=bottom_elev
                    ))
                    prev_depth = bottom_depth
                if page_rows:
                    # QA-OCR v3: repair + lọc cao độ NGAY TỪNG TRANG với cao độ miệng của chính
                    # trang đó. PDF nhiều trang = nhiều lỗ khoan có cao độ miệng khác nhau; nếu
                    # lọc toàn bộ bằng top_elev của trang cuối sẽ làm mất lớp của các trang trước.
                    # Dùng max các ứng viên trong trang: nếu một lớp parse lỗi thì phễu lọc chỉ
                    # LỎNG hơn chứ không chặt hơn (thiên về bảo toàn dữ liệu).
                    page_top_elev = max(page_top_elev_cands) if page_top_elev_cands else None
                    page_rows = self._repair_borehole_row_sequence(page_rows, top_elev=page_top_elev)
                    rows.extend(page_rows)
                    if not spt_points:
                        warnings.append(f"Trang {page_idx}: đọc được lớp từ PDF vector nhưng chưa nhận được SPT.")
            doc.close()
        except Exception as exc:
            return [], "", ""
        if not rows:
            return [], "", self._join_borehole_page_texts(text_pages, label="PDF VECTOR TEXT") if text_pages else ""
        # QA-OCR v3: KHÔNG repair lần hai trên rows gộp nhiều trang — mỗi trang đã được
        # repair với đúng cao độ miệng của nó; repair toàn cục sẽ sort theo cao độ xuyên
        # hạng mục và có nguy cơ merge dòng cùng tên lớp của hai lỗ khoan khác nhau.
        has_spt = any(str(r[4] or "").strip() for r in rows)
        warn = "Đã nhận dạng bằng nhánh PDF vector theo tọa độ/text stream: ưu tiên Tên lớp, CĐ đáy và SPT; mô tả dài được bỏ qua để tránh lệch cột. Anh kiểm tra/sửa loại đất ở preview nếu cần."
        if not has_spt:
            warn = "Đã nhận dạng lớp bằng nhánh PDF vector, nhưng chưa nhận được SPT; cột SPT đang để trống."
        if warnings:
            warn = (warn + "\n" + "\n".join(warnings)).strip()
        return rows, warn, self._join_borehole_page_texts(text_pages, label="PDF VECTOR TEXT")

    def _opencv_preprocess_for_ocr(self, img_obj: Any, *, scale: float = 1.0, numeric: bool = False) -> Optional[Any]:
        """Tiền xử lý ảnh bằng OpenCV nếu có cài cv2; lỗi thì trả None để fallback PIL.

        Gồm: grayscale, deskew nhẹ, adaptive threshold và phóng lớn crop. Cách này ổn định
        hơn ImageOps.autocontrast với ảnh scan lệch/bóng mờ. Không bắt buộc cv2 để giữ tool
        chạy được trên máy chưa cài OpenCV.
        """
        try:
            import cv2  # type: ignore
            import numpy as np  # type: ignore
            from PIL import Image
            arr = np.array(img_obj.convert("RGB") if hasattr(img_obj, "convert") else img_obj)
            if arr.ndim == 3:
                gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
            else:
                gray = arr.copy()
            # Nền tối/chữ sáng thì đảo để thống nhất nền trắng chữ đen.
            if float(np.mean(gray)) < 128.0:
                gray = cv2.bitwise_not(gray)
            # Deskew bằng các pixel tối đáng kể, chỉ xoay khi lệch đủ lớn để tránh méo crop nhỏ.
            mask = gray < 230
            coords = np.column_stack(np.where(mask)) if np.any(mask) else np.empty((0, 2))
            if coords.shape[0] > 40:
                angle = cv2.minAreaRect(coords)[-1]
                if angle < -45:
                    angle = -(90 + angle)
                else:
                    angle = -angle
                if 0.4 < abs(angle) < 8.0:
                    h, w = gray.shape[:2]
                    M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
                    gray = cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
            if scale and abs(float(scale) - 1.0) > 1e-9:
                gray = cv2.resize(gray, None, fx=float(scale), fy=float(scale), interpolation=cv2.INTER_CUBIC)
            block = 31 if numeric else 35
            if min(gray.shape[:2]) < 80:
                block = 21
            if block % 2 == 0:
                block += 1
            thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block, 2)
            return Image.fromarray(thresh)
        except Exception:
            return None

    def _opencv_table_line_centers(self, img_obj: Any) -> Tuple[List[int], List[int]]:
        """Dò đường kẻ bảng bằng OpenCV morphology, fallback dùng projection cũ khi lỗi.

        Cách này ít phụ thuộc vào ngưỡng pixel đơn lẻ hơn, chịu được đường kẻ đứt/nhiễu nhẹ.
        """
        try:
            import cv2  # type: ignore
            import numpy as np  # type: ignore
            arr = np.array(img_obj.convert("RGB") if hasattr(img_obj, "convert") else img_obj)
            gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY) if arr.ndim == 3 else arr.copy()
            if float(np.mean(gray)) < 128.0:
                gray = cv2.bitwise_not(gray)
            bw = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 35, 8)
            h, w = bw.shape[:2]
            hk = max(25, int(w * 0.035))
            vk = max(25, int(h * 0.035))
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (hk, 1))
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, vk))
            hlines = cv2.morphologyEx(bw, cv2.MORPH_OPEN, horizontal_kernel, iterations=1)
            vlines = cv2.morphologyEx(bw, cv2.MORPH_OPEN, vertical_kernel, iterations=1)
            def centers_from_projection(mask, axis: int, threshold_ratio: float) -> List[int]:
                proj = np.sum(mask > 0, axis=axis)
                lim = max(8, int((h if axis == 1 else w) * threshold_ratio))
                idx = np.where(proj >= lim)[0].tolist()
                if not idx:
                    return []
                centers: List[int] = []
                st = prev = idx[0]
                for v in idx[1:]:
                    if v > prev + 2:
                        centers.append((st + prev) // 2)
                        st = v
                    prev = v
                centers.append((st + prev) // 2)
                return centers
            xs = centers_from_projection(vlines, axis=0, threshold_ratio=0.22)
            ys = centers_from_projection(hlines, axis=1, threshold_ratio=0.22)
            if len(xs) >= 4 and len(ys) >= 3:
                return xs, ys
        except Exception:
            pass
        return [], []

    def _auto_orient_borehole_image(self, img_obj: Any) -> Any:
        """Tự xoay trang lỗ khoan AutoCAD nếu nội dung nằm ngang trong trang PDF dọc.

        Một số PDF in từ AutoCAD có page dọc nhưng toàn bộ khung hình trụ lỗ khoan bị
        xoay 90 độ. Khi đó parser cũ nhìn nhầm chiều sâu/lớp và bỏ rất nhiều dòng.
        Dựa vào bbox pixel tối: nếu vùng nội dung rộng hơn cao rõ rệt trong trang dọc
        thì xoay 90 độ để đưa form về dạng đứng trước khi dò lưới.
        """
        try:
            from PIL import ImageOps
            import numpy as np  # type: ignore
            img = img_obj.convert("RGB") if hasattr(img_obj, "convert") else img_obj
            w, h = img.size
            g = ImageOps.grayscale(img)
            arr = np.array(g)
            mask = arr < 230
            if not bool(mask.any()):
                return img_obj
            ys, xs = np.where(mask)
            bw = int(xs.max() - xs.min() + 1)
            bh = int(ys.max() - ys.min() + 1)
            ratio = bw / max(bh, 1)
            # Trang PDF dọc nhưng nội dung là một log nằm ngang: bbox thường có ratio > 1.25.
            # Bỏ điều kiện này với các ảnh/PDF landscape thật để không xoay sai.
            if h > w and ratio > 1.25:
                return img.rotate(90, expand=True)
            return img_obj
        except Exception:
            return img_obj


    def _render_pdf_pages_for_ocr(self, path: str) -> List[Any]:
        pages: List[Any] = []
        # Ưu tiên PyMuPDF nếu có.
        try:
            import fitz  # type: ignore
            from PIL import Image
            doc = fitz.open(path)
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                pages.append(self._auto_orient_borehole_image(img))
            doc.close()
            return pages
        except Exception:
            pass
        # Fallback pdf2image nếu máy người dùng đã cài poppler.
        try:
            from pdf2image import convert_from_path  # type: ignore
            pages = list(convert_from_path(path, dpi=220))
            return pages
        except Exception as exc:
            raise RuntimeError(
                "Chưa đọc được PDF. Cần cài một trong hai bộ thư viện: PyMuPDF (pip install pymupdf) "
                "hoặc pdf2image + Poppler. Chi tiết: " + str(exc)
            )

    def _images_from_borehole_file(self, path: str) -> List[Any]:
        """Mở ảnh/PDF lỗ khoan, có cache ngắn hạn để tránh render PDF lặp lại.

        V1.0.12 bị chậm do cùng một PDF bị render thành ảnh nhiều lần: một lần để OCR
        toàn trang và một lần cho parser lưới. Cache theo path/mtime/size giúp giảm rõ
        thời gian và tránh máy bị lag khi PDF nhiều trang.
        """
        ext = os.path.splitext(path)[1].lower()
        try:
            st = os.stat(path)
            key = (os.path.abspath(path), float(st.st_mtime), int(st.st_size))
        except Exception:
            key = (os.path.abspath(str(path or "")), 0.0, 0)
        cache: Dict[Any, List[Any]] = getattr(self, "_borehole_image_cache", {})
        if key in cache:
            return cache[key]

        if ext == ".pdf":
            frames = self._render_pdf_pages_for_ocr(path)
        else:
            try:
                from PIL import Image
                img = Image.open(path)
                frames = []
                try:
                    # hỗ trợ TIFF nhiều trang nếu có
                    i = 0
                    while True:
                        img.seek(i)
                        frames.append(self._auto_orient_borehole_image(img.convert("RGB").copy()))
                        i += 1
                except EOFError:
                    pass
                frames = frames or [self._auto_orient_borehole_image(img.convert("RGB"))]
            except Exception as exc:
                raise RuntimeError("Không mở được ảnh/PDF: " + str(exc))

        # Chỉ giữ 1-2 file gần nhất để không chiếm RAM khi import nhiều PDF lớn.
        cache[key] = frames
        try:
            while len(cache) > 2:
                first_key = next(iter(cache))
                if first_key != key:
                    cache.pop(first_key, None)
                else:
                    break
        except Exception:
            pass
        self._borehole_image_cache = cache
        return frames

    def _common_tesseract_candidates(self) -> List[str]:
        """Các vị trí Tesseract OCR hay gặp, đặc biệt trên Windows khi chưa đưa vào PATH."""
        candidates: List[str] = []
        for key in ("TESSERACT_CMD", "TESSERACT_EXE"):
            val = os.environ.get(key, "").strip().strip('"')
            if val:
                candidates.append(val)
        found = shutil.which("tesseract")
        if found:
            candidates.append(found)
        if sys.platform.startswith("win"):
            roots = []
            for key in ("PROGRAMFILES", "PROGRAMFILES(X86)", "LOCALAPPDATA"):
                val = os.environ.get(key)
                if val:
                    roots.append(val)
            userprofile = os.environ.get("USERPROFILE")
            if userprofile:
                roots.append(os.path.join(userprofile, "AppData", "Local", "Programs"))
            for root in roots:
                candidates.extend([
                    os.path.join(root, "Tesseract-OCR", "tesseract.exe"),
                    os.path.join(root, "Tesseract", "tesseract.exe"),
                ])
            candidates.extend([
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            ])
        else:
            candidates.extend(["/usr/bin/tesseract", "/usr/local/bin/tesseract", "/opt/homebrew/bin/tesseract"])
        # bỏ trùng, giữ thứ tự
        out: List[str] = []
        seen = set()
        for c in candidates:
            c = str(c).strip().strip('"')
            if not c or c in seen:
                continue
            seen.add(c)
            out.append(c)
        return out

    def _get_tesseract_cmd(self, ask_user: bool = False) -> Optional[str]:
        """Tìm chương trình Tesseract OCR. Nếu ask_user=True thì cho chọn tesseract.exe."""
        # QA fix O1: tuyệt đối không mở messagebox/filedialog từ thread OCR chạy nền
        # (Tkinter không thread-safe; trên máy chưa cài Tesseract sẽ treo/crash giữa lúc import).
        # Đường dẫn được hỏi 1 lần trên main thread ngay khi bấm Import (xem import_geology_from_borehole_image).
        if ask_user and threading.current_thread() is not threading.main_thread():
            ask_user = False
        cached = getattr(self, "_tesseract_cmd", "")
        if cached:
            if os.path.isfile(cached) or shutil.which(cached):
                return cached
        for cand in self._common_tesseract_candidates():
            if os.path.isfile(cand):
                self._tesseract_cmd = cand
                return cand
            found = shutil.which(cand)
            if found:
                self._tesseract_cmd = found
                return found
        if ask_user:
            msg = (
                "Không tìm thấy chương trình Tesseract OCR trong PATH hoặc các thư mục mặc định.\n\n"
                "Nếu anh đã cài Tesseract OCR, bấm OK để chọn file tesseract.exe thủ công.\n"
                "Thường nằm tại: C:\\Program Files\\Tesseract-OCR\\tesseract.exe\n\n"
                "Nếu chưa cài Tesseract OCR thì bấm Cancel và cài Tesseract OCR trước."
            )
            if messagebox.askokcancel("OCR lỗ khoan - chọn Tesseract", msg):
                exe = filedialog.askopenfilename(
                    title="Chọn tesseract.exe",
                    filetypes=[("Tesseract OCR", "tesseract.exe"), ("Executable", "*.exe"), ("All files", "*.*")],
                )
                if exe and os.path.isfile(exe):
                    self._tesseract_cmd = exe
                    return exe
        return None

    def _ocr_borehole_file(self, path: str) -> str:
        """OCR ảnh/PDF lỗ khoan.

        Ưu tiên dùng thư viện pytesseract nếu đã cài. Nếu máy chưa cài pytesseract
        nhưng có chương trình Tesseract OCR, tự gọi Tesseract CLI. Trên Windows,
        bản này tự tìm tesseract.exe trong Program Files và cho chọn thủ công nếu
        chương trình đã cài nhưng chưa đưa vào PATH.
        """
        pytesseract_mod = None
        pytesseract_err: Optional[Exception] = None
        try:
            import pytesseract as pytesseract_mod  # type: ignore
        except Exception as exc:
            pytesseract_err = exc
            pytesseract_mod = None

        # Tìm engine Tesseract ngay từ đầu. Nếu có pytesseract nhưng tesseract.exe không nằm trong PATH,
        # vẫn gán đường dẫn này để pytesseract chạy được.
        tess_cmd = self._get_tesseract_cmd(ask_user=(pytesseract_mod is None))
        if pytesseract_mod is not None and tess_cmd:
            try:
                pytesseract_mod.pytesseract.tesseract_cmd = tess_cmd  # type: ignore[attr-defined]
            except Exception:
                pass

        images = self._images_from_borehole_file(path)
        texts: List[str] = []
        pdf_text_pages = self._pdf_extractable_text_pages(path) if os.path.splitext(str(path or ""))[1].lower() == ".pdf" else []

        def ocr_by_cli(img_obj: Any, lang: str) -> str:
            cmd_path = self._get_tesseract_cmd(ask_user=True)
            if not cmd_path:
                raise RuntimeError(
                    "Máy chưa cài chương trình Tesseract OCR hoặc chưa chọn được file tesseract.exe. "
                    "Cần cài Tesseract OCR trước; nếu muốn gọi qua Python thì cài thêm: pip install pytesseract pillow. "
                    f"Chi tiết import pytesseract: {pytesseract_err}"
                )
            tmp_path = ""
            try:
                fd, tmp_path = tempfile.mkstemp(prefix="ts_sct_ocr_", suffix=".png")
                os.close(fd)
                img_obj.save(tmp_path)
                cmd = [cmd_path, tmp_path, "stdout", "-l", lang, "--psm", "6"]
                cp = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=90, **_no_window_kwargs())
                if cp.returncode != 0:
                    return ""
                return cp.stdout or ""
            except Exception:
                return ""
            finally:
                if tmp_path:
                    try:
                        os.remove(tmp_path)
                    except Exception:
                        pass

        for idx, img in enumerate(images):
            # Với PDF text-based, lấy trực tiếp text, không OCR lại trang đó để tránh sai số.
            if idx < len(pdf_text_pages) and str(pdf_text_pages[idx] or "").strip():
                texts.append(f"\n--- PAGE {idx+1} ---\n--- PDF TEXT ---\n" + str(pdf_text_pages[idx]))
                continue
            # tiền xử lý: ưu tiên OpenCV deskew + adaptive threshold; fallback PIL nếu chưa cài cv2.
            try:
                w, h = img.size
                scale = 1.5 if max(w, h) < 2200 else 1.0
                img_for_ocr = self._opencv_preprocess_for_ocr(img, scale=scale, numeric=False)
                if img_for_ocr is None:
                    from PIL import ImageOps, ImageFilter
                    if scale != 1.0:
                        img = img.resize((int(w*scale), int(h*scale)))
                    gray = ImageOps.grayscale(img)
                    gray = ImageOps.autocontrast(gray)
                    gray = gray.filter(ImageFilter.SHARPEN)
                    img_for_ocr = gray
            except Exception:
                img_for_ocr = img
            page_texts: List[str] = []
            for lang in ("vie+eng", "eng"):
                for psm in (6, 11):
                    t = ""
                    cfg = f"--psm {psm}"
                    if pytesseract_mod is not None:
                        try:
                            t = pytesseract_mod.image_to_string(img_for_ocr, lang=lang, config=cfg)
                        except Exception:
                            # Nếu pytesseract có nhưng engine/language lỗi, thử gọi CLI trực tiếp.
                            t = ocr_by_cli(img_for_ocr, lang)
                    else:
                        t = ocr_by_cli(img_for_ocr, lang)
                    if t and len(t.strip()) > 20:
                        page_texts.append(f"--- OCR lang={lang} psm={psm} ---\n" + t)
                if page_texts and sum(len(x) for x in page_texts) > 500:
                    break
            if page_texts:
                texts.append(f"\n--- PAGE {idx+1} ---\n" + "\n".join(page_texts))
        if not texts:
            extra = ""
            if pytesseract_mod is None and pytesseract_err is not None:
                extra = f" Pytesseract chưa cài: {pytesseract_err}."
            raise RuntimeError(
                "OCR không đọc được nội dung. Hãy thử ảnh rõ hơn hoặc kiểm tra Tesseract OCR. "
                "Nếu dùng tiếng Việt, nên cài thêm gói ngôn ngữ vie cho Tesseract."
                + extra
            )
        return "\n".join(texts)

    def _join_borehole_page_texts(self, page_texts: List[str], label: str = "FAST OCR") -> str:
        parts: List[str] = []
        for i, txt in enumerate(page_texts or [], start=1):
            parts.append(f"\n--- PAGE {i} ---\n--- {label} ---\n" + str(txt or ""))
        return "\n".join(parts)

    def _ocr_borehole_header_text_from_image(self, img_obj: Any, page_no: int = 1) -> str:
        """OCR nhanh phần header để lấy cao độ lỗ khoan, không OCR toàn trang.

        Đây là chìa khóa tăng tốc: parser lưới chỉ cần cao độ đầu lỗ khoan để đổi
        độ sâu sang cao độ. Phần lớp/ranh giới/SPT được đọc theo crop nhỏ.
        """
        try:
            w, h = img_obj.size
            # Header thường nằm phía trên 18-22% ảnh. Nới nhẹ để bắt được dòng "CAO ĐỘ LỖ KHOAN".
            crop = img_obj.crop((0, 0, w, max(40, int(h * 0.22))))
            proc = self._preprocess_crop_for_ocr(crop, scale=2, numeric=False)
            txt = self._ocr_pil_crop(proc, lang="vie+eng", psm=6, timeout=8)
            if not str(txt or "").strip():
                txt = self._ocr_pil_crop(proc, lang="eng", psm=6, timeout=8)
            return str(txt or "")
        except Exception:
            return ""

    def _ocr_borehole_full_page_text_fast(self, img_obj: Any, page_no: int = 1) -> str:
        """OCR toàn trang nhưng chỉ dùng khi parser lưới/text PDF thất bại.

        Bản cũ thử nhiều tổ hợp lang/psm cho mọi trang nên rất chậm. Hàm này chỉ chạy
        tối đa 1-2 lượt OCR cho trang thiếu dữ liệu, đủ để fallback text mà không làm
        treo UI quá lâu.
        """
        try:
            w, h = img_obj.size
            max_side = max(w, h)
            # Giảm kích thước ảnh rất lớn để đỡ ngốn CPU/RAM; ảnh nhỏ thì phóng nhẹ.
            if max_side > 3600:
                scale = 3200.0 / float(max_side)
            elif max_side < 1800:
                scale = 1.35
            else:
                scale = 1.0
            img_for_ocr = self._opencv_preprocess_for_ocr(img_obj, scale=scale, numeric=False)
            if img_for_ocr is None:
                from PIL import ImageOps, ImageFilter
                im = img_obj
                if scale and abs(scale - 1.0) > 1e-9:
                    im = im.resize((max(1, int(w * scale)), max(1, int(h * scale))))
                gray = ImageOps.grayscale(im)
                gray = ImageOps.autocontrast(gray)
                img_for_ocr = gray.filter(ImageFilter.SHARPEN)
            txt = self._ocr_pil_crop(img_for_ocr, lang="vie+eng", psm=6, timeout=45)
            # Nếu thiếu gói tiếng Việt hoặc text quá ít, thử eng/psm 11 một lần.
            if len(str(txt or "").strip()) < 80:
                txt2 = self._ocr_pil_crop(img_for_ocr, lang="eng", psm=11, timeout=35)
                if len(str(txt2 or "")) > len(str(txt or "")):
                    txt = txt2
            return str(txt or "")
        except Exception:
            return ""

    def _present_borehole_pages_from_rows(self, rows: List[List[Any]], item_name: str, page_count: int) -> set:
        """Xác định các trang PDF/TIFF đã có dữ liệu trong rows."""
        if page_count <= 1:
            return {1} if rows else set()
        present_pages = set()
        prefix = f"{item_name}_p"
        for r in rows or []:
            hm = str(r[0] if r else "")
            if hm.startswith(prefix):
                m = re.search(r"_p(\d+)\b", hm)
                if m:
                    try:
                        present_pages.add(int(m.group(1)))
                    except Exception:
                        pass
        return present_pages

    def _make_geology_rows_from_borehole_file_fast(self, path: str) -> Tuple[str, List[List[Any]], str, str]:
        """Luồng import nhanh cho ảnh/PDF lỗ khoan.

        Thứ tự xử lý:
        1) PDF digital: lấy text trực tiếp, không OCR.
        2) Ảnh/scan: OCR nhanh header từng trang để lấy cao độ lỗ khoan.
        3) Parser lưới đọc lớp/SPT từ crop nhỏ.
        4) Chỉ trang nào parser lưới không nhận được mới OCR toàn trang để fallback.
        """
        item_name = self._borehole_item_name_from_path(path)
        is_pdf = os.path.splitext(str(path or ""))[1].lower() == ".pdf"

        # V1.0.26/V1.0.27: PDF digital/vector phải được thử trước khi render ảnh.
        # Nếu thành công, tránh cả bước render và OCR, đồng thời giữ số cao độ/SPT chính xác hơn.
        if is_pdf:
            vector_rows, vector_warn, vector_text = self._try_make_geology_rows_from_pdf_vector_text(path, item_name)
            if vector_rows:
                return item_name, vector_rows, vector_warn, vector_text

        images = self._images_from_borehole_file(path)
        page_count = max(len(images), self._borehole_file_page_count(path))
        pdf_text_pages = self._pdf_extractable_text_pages(path) if is_pdf else []

        page_texts: List[str] = []
        for i in range(page_count):
            txt = ""
            if i < len(pdf_text_pages) and str(pdf_text_pages[i] or "").strip():
                txt = str(pdf_text_pages[i])
            elif i < len(images):
                txt = self._ocr_borehole_header_text_from_image(images[i], i + 1)
            page_texts.append(txt)
        ocr_text = self._join_borehole_page_texts(page_texts, label="PDF TEXT / FAST HEADER OCR")

        grid_rows, grid_warn = self._try_make_geology_rows_from_borehole_image_grid(path, item_name, ocr_text)
        rows: List[List[Any]] = list(grid_rows or [])
        warnings2: List[str] = []
        has_spt_any = any(str(r[4] or "").strip() for r in rows)

        missing_pages = sorted(set(range(1, page_count + 1)) - self._present_borehole_pages_from_rows(rows, item_name, page_count))
        for page_no in missing_pages:
            page_item_name = item_name if page_count <= 1 else f"{item_name}_p{page_no:02d}"
            page_text = page_texts[page_no - 1] if 0 <= page_no - 1 < len(page_texts) else ""
            # Nếu chỉ mới OCR header thì không đủ fallback. Lúc này mới OCR toàn trang.
            needs_full_ocr = len(str(page_text or "").strip()) < 220 or not re.search(r"\d+(?:[\.,]\d+).{0,30}\d+(?:[\.,]\d+)", str(page_text or ""), re.S)
            if needs_full_ocr and page_no - 1 < len(images):
                self._set_status(f"Đang OCR trang {page_no}/{page_count}: {os.path.basename(path)}...")
                self._update_ui_idle()
                full_txt = self._ocr_borehole_full_page_text_fast(images[page_no - 1], page_no)
                if full_txt:
                    page_text = (str(page_text or "") + "\n" + full_txt).strip()
                    if 0 <= page_no - 1 < len(page_texts):
                        page_texts[page_no - 1] = page_text
            if not str(page_text or "").strip():
                warnings2.append(f"Trang {page_no}: chưa đọc được dữ liệu text/OCR.")
                continue
            fb_rows, fb_has_spt = self._rows_from_borehole_text_layers(page_item_name, page_text)
            if fb_rows:
                # QA-OCR v4: trước đây dòng fallback KHÔNG qua repair nên tên lớp dính số
                # (2004/325330S/375400425...) và CĐ đáy vô lý lọt thẳng vào bảng.
                # Repair theo từng trang với cao độ miệng của chính trang đó
                # (ràng buộc CĐ đáy ≈ CĐ miệng − độ sâu).
                try:
                    fb_top_elev = self._extract_borehole_top_elev(page_text)
                except Exception:
                    fb_top_elev = None
                fb_rows = self._repair_borehole_row_sequence(fb_rows, top_elev=fb_top_elev)
            if fb_rows:
                rows.extend(fb_rows)
                has_spt_any = has_spt_any or fb_has_spt
                warnings2.append(f"Trang {page_no}: đã đọc theo text/OCR.")
            elif page_no in missing_pages:
                warnings2.append(f"Trang {page_no}: chưa nhận được lớp địa chất; cần kiểm tra trong Xem text OCR hoặc tách trang thành ảnh rõ hơn.")

        ocr_text = self._join_borehole_page_texts(page_texts, label="PDF TEXT / FAST OCR")

        warn = ""
        if not rows:
            warn = "OCR chưa nhận được lớp địa chất theo cấu trúc bảng. Có thể ảnh/PDF quá mờ hoặc form khác mẫu; hãy dùng nút Xem text OCR để copy/sửa thủ công."
        elif not has_spt_any:
            warn = "Đã nhận được lớp nhưng chưa nhận được bảng SPT; cột Nₕₜ đang để trống."
        else:
            warn = grid_warn or "Đã nhận dạng lớp và lấy SPT trung bình theo từng lớp từ cột N/30. Anh vẫn nên kiểm tra lại lớp, SPT và loại đất trong bảng preview trước khi import."
        if page_count > 1:
            got_pages = self._present_borehole_pages_from_rows(rows, item_name, page_count)
            if len(got_pages) < page_count:
                warnings2.insert(0, f"PDF/TIFF có {page_count} trang, mới nhận được {len(got_pages)} trang có dữ liệu.")
            else:
                warnings2.insert(0, f"PDF/TIFF có {page_count} trang, đã nhận được đủ {len(got_pages)} trang.")
        if warnings2:
            warn = (warn + "\n" + "\n".join(warnings2)).strip()
        return item_name, rows, warn, ocr_text

    def _parse_spt_points_from_ocr_text(self, text: str) -> List[Tuple[float, float]]:
        """Trả về list (độ sâu m, N/30 hoặc >100). Cố gắng bắt các dòng dạng: 3.00 1 2 3 5."""
        norm = self._normalize_ocr_text(text)
        points: List[Tuple[float, float]] = []
        for line in norm.splitlines():
            s = line.strip()
            if not s:
                continue
            # Bỏ các dòng mô tả quá dài không giống bảng SPT.
            if len(s) > 120 and not re.search(r"\d+\.\d+", s):
                continue
            toks = re.findall(r">?\d+(?:\.\d+)?", s)
            if len(toks) < 2:
                continue
            depth = self._safe_float_ocr(toks[0], -1.0)
            if not (0.0 <= depth <= 150.0):
                continue
            # Ưu tiên token cuối là N/30; nếu dòng chứa >100 thì lấy 100.
            if any(">100" in t or "> 100" in s for t in toks):
                nval = 100.0
            else:
                nval = self._safe_float_ocr(toks[-1], 0.0)
            if 0.0 <= nval <= 500.0:
                # Loại bớt các dòng cao độ lớp có dạng name depth thickness nếu token thứ 2 quá lớn? vẫn chấp nhận vì chỉ là gợi ý.
                points.append((depth, nval))
        # lọc trùng gần cùng độ sâu: lấy giá trị lớn nhất/ cuối cùng
        dedup: Dict[float, float] = {}
        for d, n in points:
            key = round(d, 2)
            dedup[key] = n
        return sorted(dedup.items(), key=lambda x: x[0])

    def _parse_borehole_layers_from_ocr_text(self, text: str, item_name: str) -> List[Dict[str, Any]]:
        norm = self._normalize_ocr_text(text)
        top_elev = self._extract_borehole_top_elev(norm)
        layers: List[Dict[str, Any]] = []
        # Regex chính cho các dòng có đủ: lớp, cao độ đáy, độ sâu đáy, bề dày, mô tả.
        # Không phân biệt 4a/4A do đưa tên lớp về nguyên dạng và chuẩn hóa khi so sánh.
        layer_line = re.compile(
            r"^\s*([A-Za-zĐđ0-9][A-Za-zĐđ0-9\-_/]{0,8})\s+"
            r"([-+]?\d+(?:[\.,]\d+)?)\s+"
            r"(\d+(?:[\.,]\d+)?)\s+"
            r"(\d+(?:[\.,]\d+)?)\s+(.{3,})$",
            flags=re.I,
        )
        for line in norm.splitlines():
            s = line.strip()
            if not s:
                continue
            m = layer_line.match(s)
            if not m:
                continue
            name, bottom_elev_s, bottom_depth_s, thick_s, desc = m.groups()
            # Loại các dòng SPT vô tình match: tên lớp thường không phải số thuần dài.
            if re.fullmatch(r"\d+(?:[\.,]\d+)?", name):
                continue
            bottom_depth = self._safe_float_ocr(bottom_depth_s, 0.0)
            thickness = self._safe_float_ocr(thick_s, 0.0)
            if bottom_depth <= 0 or thickness <= 0 or thickness > 100:
                continue
            bottom_elev = self._safe_float_ocr(bottom_elev_s, (top_elev or 0.0) - bottom_depth)
            layers.append({
                "item": item_name,
                "name": name.strip(),
                "bottom_elev": bottom_elev,
                "bottom_depth": bottom_depth,
                "thickness": thickness,
                "desc": desc.strip(),
            })
        # Fallback: bắt dòng dạng "CS ... 22.00 ... mô tả" nếu OCR không đưa đủ cột.
        if not layers:
            simple = re.compile(r"^\s*([A-Za-zĐđ0-9][A-Za-zĐđ0-9\-_/]{0,8})\s+.*?(\d+(?:[\.,]\d+)?)\s+(.{6,})$", re.I)
            for line in norm.splitlines():
                s = line.strip()
                m = simple.match(s)
                if not m:
                    continue
                name, depth_s, desc = m.groups()
                if len(_strip_accents(desc)) < 8:
                    continue
                depth = self._safe_float_ocr(depth_s, 0.0)
                if 0 < depth <= 120:
                    bottom_elev = (top_elev - depth) if top_elev is not None else -depth
                    layers.append({"item": item_name, "name": name.strip(), "bottom_elev": bottom_elev, "bottom_depth": depth, "thickness": 0.0, "desc": desc.strip()})
        # Sắp theo độ sâu và tính lại thickness nếu cần.
        layers.sort(key=lambda x: float(x.get("bottom_depth") or 0.0))
        prev = 0.0
        for d in layers:
            bd = float(d.get("bottom_depth") or 0.0)
            if not d.get("thickness"):
                d["thickness"] = max(bd - prev, 0.0)
            prev = max(prev, bd)
        return layers


    def _pil_mean_gray(self, img_obj: Any) -> float:
        try:
            from PIL import ImageOps
            g = ImageOps.grayscale(img_obj)
            hist = g.histogram()
            total = sum(hist) or 1
            return sum(i * v for i, v in enumerate(hist)) / total
        except Exception:
            return 255.0

    def _borehole_small_image_cache_key(self, img_obj: Any, prefix: str, extra: Tuple[Any, ...] = ()) -> Optional[Tuple[Any, ...]]:
        """Tạo khóa cache cho các crop OCR nhỏ."""
        try:
            w, h = img_obj.size
            if int(w) * int(h) > 900000:
                return None
            mode = getattr(img_obj, "mode", "")
            crc = zlib.crc32(img_obj.tobytes()) & 0xFFFFFFFF
            return (prefix, int(w), int(h), str(mode), int(crc)) + tuple(extra or ())
        except Exception:
            return None

    def _borehole_cache_get(self, cache_name: str, key: Optional[Tuple[Any, ...]]) -> Any:
        if key is None:
            return None
        try:
            return getattr(self, cache_name, {}).get(key, None)
        except Exception:
            return None

    def _borehole_cache_set(self, cache_name: str, key: Optional[Tuple[Any, ...]], value: Any, max_items: int = 600) -> Any:
        if key is None:
            return value
        try:
            cache = getattr(self, cache_name, {})
            cache[key] = value
            while len(cache) > int(max_items):
                cache.pop(next(iter(cache)), None)
            setattr(self, cache_name, cache)
        except Exception:
            pass
        return value

    def _borehole_lite_rows_high_confidence(self, rows: List[List[Any]], image_count: int = 1) -> bool:
        """Cho phép bỏ qua parser chi tiết khi parser nhanh đã đủ tin cậy."""
        try:
            if len(rows or []) < 2:
                return False
            names = [str((r + [""] * 14)[1] or "").strip() for r in rows]
            good = 0
            for nm in names:
                n = re.sub(r"[^a-z0-9]", "", _strip_accents(nm).lower())
                if n and not re.fullmatch(r"l\d+", n) and not self._is_ambiguous_borehole_layer_name(nm):
                    good += 1
            ratio = good / max(len(names), 1)
            if ratio < (0.90 if len(names) <= 4 else 0.78):
                return False
            elevs = []
            for r in rows:
                try:
                    elevs.append(float(str((r + [""] * 14)[2]).replace(",", ".")))
                except Exception:
                    pass
            if len(elevs) >= 2:
                decreases = sum(1 for a, b in zip(elevs, elevs[1:]) if b <= a + 0.05)
                if decreases < max(1, len(elevs) - 2):
                    return False
            return True
        except Exception:
            return False

    def _preprocess_crop_for_ocr(self, img_obj: Any, scale: int = 3, numeric: bool = False) -> Any:
        # QA-OCR v5: mặc định quay lại ảnh xám + resize hình học tuyệt đối.
        # Benchmark 16 ảnh lỗ khoan cho thấy pipeline deskew/adaptive-threshold cũ làm mất
        # nhiều ranh giới ở nhóm R3/JPG. Ảnh xám giữ tọa độ TSV ổn định hơn cho bài toán
        # fit trục sâu.  Có thể bật _ocr_use_legacy_preprocess=True để thử lại pipeline cũ.
        # Bất biến: KHÔNG thêm viền/padding, vì tọa độ OCR được scale ngược về crop gốc.
        if not bool(getattr(self, "_ocr_use_legacy_preprocess", False)):
            try:
                import cv2  # type: ignore
                import numpy as _np  # type: ignore
                from PIL import Image as _PILImage
                arr = _np.array(img_obj.convert("L") if hasattr(img_obj, "convert") else img_obj)
                if arr.ndim == 3:
                    arr = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
                if float(_np.mean(arr)) < 128.0:
                    arr = 255 - arr
                sc = max(float(scale or 1), 1.0)
                if sc > 1.01:
                    arr = cv2.resize(arr, None, fx=sc, fy=sc, interpolation=cv2.INTER_CUBIC)
                return _PILImage.fromarray(arr)
            except Exception:
                pass
        cv_img = self._opencv_preprocess_for_ocr(img_obj, scale=float(scale or 1), numeric=numeric)
        if cv_img is not None:
            return cv_img
        try:
            from PIL import ImageOps, ImageEnhance, ImageFilter
            im = ImageOps.grayscale(img_obj)
            if self._pil_mean_gray(im) < 128.0:
                im = ImageOps.invert(im)
            im = ImageOps.autocontrast(im)
            if scale and scale != 1:
                im = im.resize((max(1, im.width * scale), max(1, im.height * scale)))
            im = ImageEnhance.Contrast(im).enhance(2.4 if numeric else 1.8)
            im = im.filter(ImageFilter.SHARPEN)
            return im
        except Exception:
            return img_obj

    def _borehole_ocr_engine_key(self) -> str:
        try:
            # Khi OCR chạy trong worker thread, không đọc trực tiếp Tk StringVar.
            if not self._is_main_thread() and hasattr(self, "_borehole_ocr_engine_runtime"):
                return _normalize_borehole_ocr_engine(getattr(self, "_borehole_ocr_engine_runtime", "TESSERACT"))
            return _normalize_borehole_ocr_engine(self.borehole_ocr_engine.get())
        except Exception:
            return "TESSERACT"

    def _get_paddleocr_reader(self, lang: str = "vie+eng") -> Any:
        """Khởi tạo PaddleOCR theo kiểu lazy-load.

        PaddleOCR là tùy chọn cho ảnh/PDF scan. Nếu máy chưa cài paddleocr/paddlepaddle
        hoặc API PaddleOCR khác phiên bản, hàm trả None để code fallback về Tesseract.
        """
        paddle_lang = "vi" if ("vie" in str(lang).lower() or "vi" in str(lang).lower()) else "en"
        cache = getattr(self, "_paddleocr_reader_cache", {}) or {}
        if paddle_lang in cache:
            return cache.get(paddle_lang)
        try:
            from paddleocr import PaddleOCR  # type: ignore
        except Exception as exc:
            if not getattr(self, "_paddleocr_missing_warned", False):
                self._paddleocr_missing_warned = True
                try:
                    self._set_status("PaddleOCR chưa sẵn sàng; dùng Tesseract")
                except Exception:
                    pass
            return None
        reader = None
        init_errors = []
        # PaddleOCR có thay đổi tham số giữa các phiên bản; thử vài cấu hình phổ biến.
        init_kwargs_list = [
            # Không dùng angle classifier mặc định để giảm thời gian khởi tạo và OCR.
            {"use_angle_cls": False, "lang": paddle_lang, "show_log": False},
            {"use_angle_cls": False, "lang": paddle_lang},
            {"lang": paddle_lang, "show_log": False},
            {"lang": paddle_lang},
        ]
        for kwargs in init_kwargs_list:
            try:
                # Một số bản PaddleOCR vẫn in log dù show_log=False; chặn stdout/stderr để CMD không bị tràn log.
                with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                    reader = PaddleOCR(**kwargs)
                break
            except Exception as exc:
                init_errors.append(exc)
                reader = None
        if reader is None:
            if not getattr(self, "_paddleocr_init_warned", False):
                self._paddleocr_init_warned = True
                try:
                    self._set_status("Không khởi tạo được PaddleOCR; dùng Tesseract")
                except Exception:
                    pass
            return None
        cache[paddle_lang] = reader
        self._paddleocr_reader_cache = cache
        return reader

    @staticmethod
    def _paddle_box_to_rect(box: Any) -> Tuple[int, int, int, int]:
        try:
            pts = box
            # box dạng [x1,y1,x2,y2]
            if isinstance(pts, (list, tuple)) and len(pts) == 4 and all(isinstance(v, (int, float)) for v in pts):
                x1, y1, x2, y2 = [float(v) for v in pts]
                left, top = min(x1, x2), min(y1, y2)
                return int(left), int(top), int(abs(x2 - x1)), int(abs(y2 - y1))
            xs, ys = [], []
            for pnt in pts:
                if isinstance(pnt, (list, tuple)) and len(pnt) >= 2:
                    xs.append(float(pnt[0])); ys.append(float(pnt[1]))
            if xs and ys:
                left, right = min(xs), max(xs)
                top, bottom = min(ys), max(ys)
                return int(left), int(top), int(max(right - left, 1)), int(max(bottom - top, 1))
        except Exception:
            pass
        return 0, 0, 0, 0

    def _parse_paddleocr_result_items(self, result: Any) -> List[Dict[str, Any]]:
        """Chuẩn hóa kết quả PaddleOCR về list {text,left,top,width,height,conf}."""
        items: List[Dict[str, Any]] = []

        def add_item(text: Any, box: Any = None, conf: Any = 0.0):
            txt = str(text or "").strip()
            if not txt:
                return
            left, top, width, height = self._paddle_box_to_rect(box) if box is not None else (0, 0, 0, 0)
            try:
                cf = float(conf)
            except Exception:
                cf = 0.0
            # Một số phiên bản trả score 0-1, đổi về 0-100 cho đồng bộ TSV.
            if 0.0 <= cf <= 1.0:
                cf *= 100.0
            items.append({"text": txt, "left": left, "top": top, "width": width, "height": height, "conf": cf})

        def looks_like_box(x: Any) -> bool:
            try:
                if isinstance(x, (list, tuple)) and len(x) >= 4 and all(isinstance(v, (int, float)) for v in x[:4]):
                    return True
                if isinstance(x, (list, tuple)) and len(x) >= 4 and all(isinstance(p, (list, tuple)) and len(p) >= 2 for p in x[:4]):
                    return True
            except Exception:
                return False
            return False

        def walk(node: Any):
            if node is None:
                return
            if isinstance(node, dict):
                texts = node.get("rec_texts") or node.get("texts") or node.get("text")
                scores = node.get("rec_scores") or node.get("scores") or node.get("score")
                boxes = node.get("dt_polys") or node.get("rec_polys") or node.get("boxes") or node.get("rec_boxes")
                if isinstance(texts, list):
                    for i, txt in enumerate(texts):
                        box = boxes[i] if isinstance(boxes, list) and i < len(boxes) else None
                        conf = scores[i] if isinstance(scores, list) and i < len(scores) else 0.0
                        add_item(txt, box, conf)
                    return
                for v in node.values():
                    walk(v)
                return
            if isinstance(node, (list, tuple)):
                # Dạng cổ điển: [box, (text, score)] hoặc [box, [text, score]]
                if len(node) >= 2 and looks_like_box(node[0]) and isinstance(node[1], (list, tuple)) and len(node[1]) >= 1:
                    txt = node[1][0]
                    conf = node[1][1] if len(node[1]) > 1 else 0.0
                    add_item(txt, node[0], conf)
                    return
                # Dạng [box, text, score]
                if len(node) >= 2 and looks_like_box(node[0]) and isinstance(node[1], str):
                    conf = node[2] if len(node) > 2 else 0.0
                    add_item(node[1], node[0], conf)
                    return
                for v in node:
                    walk(v)
                return

        walk(result)
        # Loại trùng thô và sắp xếp theo vị trí.
        dedup: Dict[Tuple[str, int, int], Dict[str, Any]] = {}
        for it in items:
            key = (str(it.get("text", "")), int(it.get("left", 0)) // 3, int(it.get("top", 0)) // 3)
            dedup[key] = it
        return sorted(dedup.values(), key=lambda r: (int(r.get("top", 0)), int(r.get("left", 0))))

    def _paddleocr_items_from_pil(self, img_obj: Any, *, lang: str = "vie+eng", timeout: int = 60) -> List[Dict[str, Any]]:
        cache_key = self._borehole_small_image_cache_key(img_obj, "paddle_data", (str(lang),))
        cached = self._borehole_cache_get("_borehole_paddle_ocr_cache", cache_key)
        if cached is not None:
            return [dict(x) for x in (cached or [])]
        # PaddleOCR CPU rất chậm với crop dài/toàn trang. Chỉ dùng cho crop nhỏ;
        # crop lớn, cột SPT dài hoặc full page sẽ fallback Tesseract để giữ tốc độ.
        try:
            w, h = getattr(img_obj, "size", (0, 0))
            w, h = int(w or 0), int(h or 0)
            if max(w, h) > 2200 or (w * h) > 1600000:
                return []
        except Exception:
            pass
        reader = self._get_paddleocr_reader(lang=lang)
        if reader is None:
            return []
        try:
            import numpy as np  # type: ignore
            img_rgb = img_obj.convert("RGB") if hasattr(img_obj, "convert") else img_obj
            arr = np.array(img_rgb)
            with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                try:
                    result = reader.ocr(arr, cls=False)
                except TypeError:
                    result = reader.ocr(arr)
            items = self._parse_paddleocr_result_items(result)
            return self._borehole_cache_set("_borehole_paddle_ocr_cache", cache_key, items, max_items=300)
        except Exception:
            return []

    def _paddleocr_text_from_pil(self, img_obj: Any, *, lang: str = "vie+eng", whitelist: str = "") -> str:
        items = self._paddleocr_items_from_pil(img_obj, lang=lang)
        if not items:
            return ""
        texts: List[str] = []
        for it in items:
            txt = str(it.get("text", "") or "").strip()
            if not txt:
                continue
            if whitelist:
                txt = "".join(ch for ch in txt if ch in whitelist)
            if txt:
                texts.append(txt)
        return "\n".join(texts)

    def _get_rapidocr_reader(self) -> Any:
        """QA-OCR: khởi tạo RapidOCR (PP-OCR/onnxruntime) lazy-load; thiếu thư viện thì trả None."""
        cached = getattr(self, "_rapidocr_reader", None)
        if cached is not None:
            return cached if cached is not False else None
        try:
            from rapidocr_onnxruntime import RapidOCR  # type: ignore
        except Exception:
            self._rapidocr_reader = False
            if not getattr(self, "_rapidocr_missing_warned", False):
                self._rapidocr_missing_warned = True
                try:
                    self._set_status("RapidOCR chưa cài (pip install rapidocr-onnxruntime); dùng Tesseract")
                except Exception:
                    pass
            return None
        try:
            with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                reader = RapidOCR()
        except Exception:
            self._rapidocr_reader = False
            return None
        self._rapidocr_reader = reader
        return reader

    @staticmethod
    def _parse_rapidocr_result_items(result: Any) -> List[Dict[str, Any]]:
        """Đổi kết quả RapidOCR ([box, text, score]) về format items dùng chung với Paddle."""
        items: List[Dict[str, Any]] = []
        for rec in (result or []):
            try:
                box, txt = rec[0], str(rec[1] or "").strip()
                conf = float(rec[2]) if len(rec) > 2 else 0.0
            except Exception:
                continue
            if not txt:
                continue
            left, top, width, height = SCTApp._paddle_box_to_rect(box)
            items.append({"text": txt, "left": left, "top": top, "width": width, "height": height, "conf": conf})
        return sorted(items, key=lambda r: (int(r.get("top", 0)), int(r.get("left", 0))))

    def _rapidocr_items_from_pil(self, img_obj: Any, *, lang: str = "vie+eng", timeout: int = 60) -> List[Dict[str, Any]]:
        """OCR bằng RapidOCR, trả items {text,left,top,width,height,conf} theo tọa độ ảnh GỐC."""
        cache_key = self._borehole_small_image_cache_key(img_obj, "rapid_data", ())
        cached = self._borehole_cache_get("_borehole_rapid_ocr_cache", cache_key)
        if cached is not None:
            return [dict(x) for x in (cached or [])]
        reader = self._get_rapidocr_reader()
        if reader is None:
            return []
        try:
            import numpy as np  # type: ignore
            img_rgb = img_obj.convert("RGB") if hasattr(img_obj, "convert") else img_obj
            w, h = img_rgb.size
            if max(w, h) > 4200:
                return []
            # Ảnh/crop độ phân giải thấp: phóng to để chữ đạt cỡ model đọc tốt (benchmark: x2-x3).
            sc = 1.0
            if max(w, h) < 1400:
                sc = min(2.5, 1400.0 / max(max(w, h), 1))
            if sc > 1.01:
                from PIL import Image as _PILImage
                resample = getattr(getattr(_PILImage, "Resampling", _PILImage), "LANCZOS")
                img_rgb = img_rgb.resize((max(1, int(w * sc)), max(1, int(h * sc))), resample)
            arr = np.array(img_rgb)
            with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                result = reader(arr)
            if isinstance(result, tuple):
                result = result[0]
            items = self._parse_rapidocr_result_items(result)
            if sc > 1.01:
                for it in items:
                    for k in ("left", "top", "width", "height"):
                        it[k] = int(round(float(it.get(k, 0)) / sc))
            return self._borehole_cache_set("_borehole_rapid_ocr_cache", cache_key, items, max_items=300)
        except Exception:
            return []

    def _rapidocr_text_from_pil(self, img_obj: Any, *, lang: str = "vie+eng", whitelist: str = "") -> str:
        """Text RapidOCR đã gom theo DÒNG (các box cùng cao độ y ghép chung một dòng)."""
        items = self._rapidocr_items_from_pil(img_obj, lang=lang)
        if not items:
            return ""
        rows: List[List[Dict[str, Any]]] = []
        for it in items:
            txt = str(it.get("text", "") or "").strip()
            if whitelist:
                txt = "".join(ch for ch in txt if ch in whitelist).strip()
            if not txt:
                continue
            it = dict(it); it["text"] = txt
            cy = float(it.get("top", 0)) + 0.5 * float(it.get("height", 0))
            hh = max(float(it.get("height", 0)), 6.0)
            placed = False
            for row in rows:
                r0 = row[0]
                cy0 = float(r0.get("top", 0)) + 0.5 * float(r0.get("height", 0))
                if abs(cy - cy0) <= 0.6 * max(hh, float(r0.get("height", 0)) or 6.0):
                    row.append(it); placed = True; break
            if not placed:
                rows.append([it])
        lines: List[str] = []
        for row in rows:
            row.sort(key=lambda r: int(r.get("left", 0)))
            lines.append(" ".join(str(r.get("text", "")) for r in row))
        return "\n".join(lines)

    def _ocr_pil_crop(self, img_obj: Any, *, lang: str = "vie+eng", psm: int = 6, whitelist: str = "", timeout: int = 20) -> str:
        """OCR một crop nhỏ.

        Engine mặc định là Tesseract/OpenCV. Nếu Settings chọn PaddleOCR thì thử PaddleOCR
        trước cho ảnh/PDF scan; nếu chưa cài hoặc nhận rỗng sẽ fallback về Tesseract.
        """
        engine_key = self._borehole_ocr_engine_key()
        cache_key = self._borehole_small_image_cache_key(img_obj, "ocr", (engine_key, str(lang), int(psm), str(whitelist or "")))
        cached = self._borehole_cache_get("_borehole_ocr_crop_cache", cache_key)
        if cached is not None:
            return str(cached or "")
        if engine_key in ("PADDLE", "RAPID"):
            # V1.0.51/QA-OCR: engine scan (Paddle/Rapid) được thử trước; rỗng thì fallback Tesseract.
            if engine_key == "RAPID":
                txt_alt = self._rapidocr_text_from_pil(img_obj, lang=lang, whitelist=whitelist)
            else:
                txt_alt = self._paddleocr_text_from_pil(img_obj, lang=lang, whitelist=whitelist)
            if str(txt_alt or "").strip():
                return self._borehole_cache_set("_borehole_ocr_crop_cache", cache_key, txt_alt)
        pytesseract_mod = None
        try:
            import pytesseract as pytesseract_mod  # type: ignore
        except Exception:
            pytesseract_mod = None
        tess_cmd = self._get_tesseract_cmd(ask_user=(pytesseract_mod is None))
        config = f"--psm {int(psm)}"
        if whitelist:
            config += f" -c tessedit_char_whitelist={whitelist}"
        if pytesseract_mod is not None:
            if tess_cmd:
                try:
                    pytesseract_mod.pytesseract.tesseract_cmd = tess_cmd  # type: ignore[attr-defined]
                except Exception:
                    pass
            try:
                txt_out = pytesseract_mod.image_to_string(img_obj, lang=lang, config=config, timeout=timeout) or ""
                return self._borehole_cache_set("_borehole_ocr_crop_cache", cache_key, txt_out)
            except Exception:
                pass
        if not tess_cmd:
            return ""
        tmp_path = ""
        try:
            fd, tmp_path = tempfile.mkstemp(prefix="ts_sct_crop_ocr_", suffix=".png")
            os.close(fd)
            img_obj.save(tmp_path)
            cmd = [tess_cmd, tmp_path, "stdout", "-l", lang, "--psm", str(int(psm))]
            if whitelist:
                cmd.extend(["-c", f"tessedit_char_whitelist={whitelist}"])
            cp = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=timeout, **_no_window_kwargs())
            txt_out = cp.stdout or ""
            return self._borehole_cache_set("_borehole_ocr_crop_cache", cache_key, txt_out)
        except Exception:
            return ""
        finally:
            if tmp_path:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass


    def _ocr_pil_crop_data(self, img_obj: Any, *, lang: str = "eng", psm: int = 6, whitelist: str = "", timeout: int = 20) -> List[Dict[str, Any]]:
        """OCR dạng TSV để lấy cả vị trí chữ/số trong crop.

        Dùng cho cột N/30 của bảng SPT: lấy giá trị và tọa độ y, rồi nội suy
        theo chiều sâu lỗ khoan để tính trung bình SPT trong từng lớp.
        """
        engine_key = self._borehole_ocr_engine_key()
        cache_key = self._borehole_small_image_cache_key(img_obj, "ocr_data", (engine_key, str(lang), int(psm), str(whitelist or "")))
        cached = self._borehole_cache_get("_borehole_ocr_crop_data_cache", cache_key)
        if cached is not None:
            return [dict(x) for x in (cached or [])]
        if engine_key in ("PADDLE", "RAPID"):
            # V1.0.51/QA-OCR: engine scan xử lý cả crop số nhỏ; rỗng thì fallback Tesseract.
            if engine_key == "RAPID":
                rows_paddle = self._rapidocr_items_from_pil(img_obj, lang=lang)
            else:
                rows_paddle = self._paddleocr_items_from_pil(img_obj, lang=lang)
            if whitelist and rows_paddle:
                filtered = []
                for r in rows_paddle:
                    rr = dict(r)
                    txt = "".join(ch for ch in str(rr.get("text", "") or "") if ch in whitelist).strip()
                    if txt:
                        rr["text"] = txt
                        filtered.append(rr)
                rows_paddle = filtered
            if rows_paddle:
                return self._borehole_cache_set("_borehole_ocr_crop_data_cache", cache_key, rows_paddle)
        pytesseract_mod = None
        try:
            import pytesseract as pytesseract_mod  # type: ignore
        except Exception:
            pytesseract_mod = None
        tess_cmd = self._get_tesseract_cmd(ask_user=(pytesseract_mod is None))
        config = f"--psm {int(psm)}"
        if whitelist:
            config += f" -c tessedit_char_whitelist={whitelist}"
        rows: List[Dict[str, Any]] = []
        if pytesseract_mod is not None:
            if tess_cmd:
                try:
                    pytesseract_mod.pytesseract.tesseract_cmd = tess_cmd  # type: ignore[attr-defined]
                except Exception:
                    pass
            try:
                data = pytesseract_mod.image_to_data(img_obj, lang=lang, config=config, output_type=pytesseract_mod.Output.DICT, timeout=timeout)
                n = len(data.get("text", []))
                for i in range(n):
                    txt = str(data.get("text", [""])[i] or "").strip()
                    if not txt:
                        continue
                    rows.append({
                        "text": txt,
                        "left": int(data.get("left", [0])[i] or 0),
                        "top": int(data.get("top", [0])[i] or 0),
                        "width": int(data.get("width", [0])[i] or 0),
                        "height": int(data.get("height", [0])[i] or 0),
                        "conf": float(data.get("conf", [-1])[i] or -1),
                    })
                return rows
            except Exception:
                rows = []
        if not tess_cmd:
            return rows
        tmp_path = ""
        try:
            fd, tmp_path = tempfile.mkstemp(prefix="ts_sct_tsv_ocr_", suffix=".png")
            os.close(fd)
            img_obj.save(tmp_path)
            cmd = [tess_cmd, tmp_path, "stdout", "-l", lang, "--psm", str(int(psm)), "tsv"]
            if whitelist:
                cmd.extend(["-c", f"tessedit_char_whitelist={whitelist}"])
            cp = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=timeout, **_no_window_kwargs())
            lines = (cp.stdout or "").splitlines()
            if not lines:
                return rows
            header = lines[0].split("\t")
            idx = {name: i for i, name in enumerate(header)}
            for line in lines[1:]:
                parts = line.split("\t")
                if len(parts) < len(header):
                    continue
                txt = parts[idx.get("text", len(parts)-1)].strip() if "text" in idx else parts[-1].strip()
                if not txt:
                    continue
                def _ival(key: str, default: int = 0) -> int:
                    try:
                        return int(float(parts[idx[key]])) if key in idx else default
                    except Exception:
                        return default
                def _fval(key: str, default: float = -1.0) -> float:
                    try:
                        return float(parts[idx[key]]) if key in idx else default
                    except Exception:
                        return default
                rows.append({"text": txt, "left": _ival("left"), "top": _ival("top"), "width": _ival("width"), "height": _ival("height"), "conf": _fval("conf")})
        except Exception:
            return rows
        finally:
            if tmp_path:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
        return rows

    def _parse_spt_n30_value_from_ocr_token(self, token: Any) -> Optional[float]:
        """Chuẩn hóa giá trị N/30 từ OCR.

        QA-OCR SPT50:
        - Ký hiệu refusal dạng ``>50`` trong hồ sơ lỗ khoan KHÔNG được đổi thành 100,
          vì cột SPT đang nhập là N/30 dùng cho tính SCT; tự nâng 50 -> 100 sẽ làm tăng
          sức kháng và có thể chuyển nhầm lớp sang nhánh IGM/đá theo N>=100.
        - Chỉ quy về 100 khi token thể hiện rõ ``100`` hoặc ``>100``.
        """
        raw = str(token or "").strip()
        if not raw:
            return None
        compact = raw.replace(" ", "")
        compact_norm = compact.upper().replace("O", "0").replace("S", "5")
        # Các biến thể OCR của >50: >50, >S0, 50>, /50, > 50...
        if re.search(r">\s*5[0O]", compact_norm) or compact_norm.startswith("50>") or compact_norm in {"/50", "\\50"}:
            return 50.0
        m_gt = re.search(r">\s*(\d+)", compact_norm)
        if m_gt:
            try:
                gv = float(m_gt.group(1))
                if gv <= 55.0:
                    return max(gv, 50.0)
                if gv >= 95.0:
                    return 100.0
                return min(gv, 100.0)
            except Exception:
                pass
        # Chỉ nhận 100 khi token có cụm 100 rõ ràng, tránh các chuỗi dính như 10.00 trong cột khác.
        if re.search(r"(^|[^0-9])100([^0-9]|$)", compact_norm):
            return 100.0
        # Nhận cả SPT=0 khi token là số 0 rõ ràng trong bảng. Trước đây 0 bị coi là missing,
        # sau đó các rule "cứu dữ liệu" có thể tự nâng thành 50/100.
        nums = re.findall(r"\d+(?:\.\d+)?", compact_norm)
        if not nums:
            return None
        try:
            vals = [float(x) for x in nums]
            v = max(vals)
        except Exception:
            return None
        if v < 0:
            return None
        if v > 100:
            # QA V1.0.29: số trơn > 100 (vd 340, 2460) là OCR đọc nhầm toạ độ/độ sâu/cột khác,
            # KHÔNG phải refusal N=100. Refusal thật ghi rõ "100" hoặc ">100" đã bắt ở trên.
            # Trả None (ô không đọc được) để KHÔNG bịa SPT=100 từ số rác.
            return None
        return v


    def _spt_graph_fallback_value(self, value: Any) -> float:
        """Giá trị SPT lấy từ ĐỒ THỊ chỉ dùng làm fallback bảo thủ.

        Đường cong SPT trên ảnh scan thường không phân biệt chắc ``>50`` với thang 0-100.
        Vì vậy không dùng đồ thị để tự nâng N lên 100. Giá trị OCR trực tiếp ở cột N/30
        vẫn được giữ nguyên nếu token ghi rõ >65, 75, 87, 100...
        """
        try:
            return max(0.0, min(50.0, float(value)))
        except Exception:
            return 0.0


    def _spt_combo_value_from_n2n3(self, n2: Optional[float], n3: Optional[float]) -> Optional[float]:
        """Ghép N2+N3 nhưng không biến refusal >50 thành 100 khi OCR chưa chắc chắn."""
        if n2 is None or n3 is None:
            return None
        n2f = float(n2); n3f = float(n3)
        if 45.0 <= n2f <= 55.0 and 45.0 <= n3f <= 55.0:
            return 50.0
        if n2f >= 95.0 or n3f >= 95.0:
            return 100.0
        return max(0.0, min(100.0, n2f + n3f))


    def _spt_import_row_value(self, value: Any) -> Tuple[Any, str]:
        """Giá trị SPT ghi ra bảng sau OCR ảnh/PDF scan.

        Không hạ mọi giá trị >50 về 50. Nếu token OCR/nguồn text ghi rõ >65, 75, 87, 100...
        thì giữ đúng số. Riêng đồ thị SPT fallback đã bị giới hạn bảo thủ trong
        ``_spt_graph_fallback_value`` nên không cần hạ thêm ở bước ghi bảng.
        """
        try:
            txt = str(value or "").strip()
            if not txt:
                return value, ""
            v = float(value)
            if v < 0.0:
                return "", ""
            return v, ""
        except Exception:
            return value, ""


    def _extract_spt_points_from_graph_curve(self, img_obj: Any, xlines: List[int], x_desc1: int, data_top: float, data_bottom: float, total_depth: float) -> List[Tuple[float, float]]:
        """Fallback đồ thị SPT - mặc định TẮT để tránh sinh SPT giả.

        QA-OCR SPT-SAFE:
        Đồ thị SPT chỉ là hình minh họa nên rất dễ nhầm trục 0, đường lưới, đường cong
        hoặc ký hiệu refusal ``>50`` thành giá trị 50/100. Vì cột SPT trong TS-CAP được
        dùng trực tiếp để tính SCT, mặc định KHÔNG dùng đồ thị để tự sinh hoặc nâng giá
        trị SPT. Chỉ bật lại khi debug bằng biến môi trường:
            TS_CAP_ENABLE_SPT_GRAPH_FALLBACK=1
        Khi tắt, phần mềm chỉ lấy SPT từ số OCR đọc trực tiếp trong bảng SPT/N/30 hoặc
        N2+N3; nếu không đọc chắc thì để trống để người dùng soát.
        """
        try:
            if str(os.environ.get("TS_CAP_ENABLE_SPT_GRAPH_FALLBACK", "")).strip().lower() not in {"1", "true", "yes", "on"}:
                return []
            if data_bottom <= data_top or total_depth <= 0:
                return []
            w, h = img_obj.size
            pix = img_obj.load()
            right_lines = sorted([x for x in xlines if x > x_desc1 + 2])
            if len(right_lines) < 5:
                return []
            x_graph_guess = right_lines[4]
            xs: List[int] = []
            ys: List[int] = []
            for yy in range(max(0, int(data_top) - 10), min(h, int(data_bottom) + 10)):
                for xx in range(max(0, int(x_graph_guess) - 2), min(w, int(w * 0.96))):
                    r, g, b = pix[xx, yy]
                    if g > 80 and g > r + 20 and g > b + 8:
                        xs.append(xx)
                        ys.append(yy)
            if not xs:
                return []
            # Ưu tiên biên đồ thị theo đường kẻ dọc chính, ổn định hơn min/max pixel xanh
            # vì lưới xanh/đường cong đôi khi không bắt đủ mép 0 hoặc 200.
            x0, x1 = min(xs), max(xs)
            try:
                if len(right_lines) >= 6:
                    gx0 = float(right_lines[4])
                    gx1 = float(right_lines[5])
                    if gx1 > gx0 + 30:
                        x0, x1 = gx0, gx1
            except Exception:
                pass
            if x1 <= x0 + 10:
                return []
            import statistics
            pts: List[Tuple[float, float]] = []
            for yy in range(max(0, int(data_top)), min(h, int(data_bottom)) + 1):
                bxs: List[int] = []
                for xx in range(int(round(x0)), int(round(x1)) + 1):
                    r, g, b = pix[xx, yy]
                    # Đường cong thường màu xanh/lam; loại phần lưới xanh bằng điều kiện B trội hơn G.
                    if b > 75 and b > r + 20 and b > g + 10:
                        bxs.append(xx)
                if not bxs:
                    continue
                xx = float(statistics.median(bxs))
                depth = (yy - data_top) / max(data_bottom - data_top, 1.0) * total_depth
                # Đồ thị chỉ là fallback vị trí/giá trị gần đúng. Không dùng đồ thị để biến
                # ký hiệu refusal >50 thành N=100 trong bảng địa chất.
                nval = 200.0 * (xx - x0) / max(x1 - x0, 1.0)
                nval = self._spt_graph_fallback_value(nval)
                if 0.0 <= depth <= total_depth + 0.5:
                    pts.append((float(depth), float(nval)))
            # Rút gọn theo bước ~0.25m để không tạo quá nhiều điểm.
            out: List[Tuple[float, float]] = []
            bucket: Dict[int, List[float]] = {}
            for d, n in pts:
                key = int(round(d / 0.25))
                bucket.setdefault(key, []).append(n)
            for key in sorted(bucket):
                vals = bucket[key]
                if vals:
                    out.append((key * 0.25, sum(vals) / len(vals)))
            return out
        except Exception:
            return []

    def _sample_spt_graph_points_by_spacing(self, graph_pts: List[Tuple[float, float]], total_depth: float, spacing: Optional[float] = None) -> List[Tuple[float, float]]:
        """Rút gọn đường cong SPT thành các điểm SPT theo khoảng cách người dùng nhập.

        Với ảnh scan, OCR cột N/30 có thể chỉ đọc được vài số do chữ nhỏ/màu xanh/đường kẻ.
        Khi đó fallback theo đồ thị SPT sẽ cho rất nhiều điểm liên tục theo pixel. Hàm này
        lấy lại một điểm đại diện theo bước SPT thực tế, tránh sinh hàng trăm dòng địa chất.
        """
        try:
            pts0 = sorted((float(d), float(n)) for d, n in (graph_pts or []) if 0.0 <= float(d) <= float(total_depth) + 0.75 and 0.0 <= float(n) <= 500.0)
        except Exception:
            pts0 = []
        if not pts0:
            return []
        try:
            sp = float(spacing if spacing is not None else getattr(self, "_borehole_spt_spacing_m", 2.0) or 2.0)
        except Exception:
            sp = 2.0
        if sp <= 0.10:
            sp = 2.0
        min_d = min(d for d, _n in pts0)
        max_d = min(max(d for d, _n in pts0), float(total_depth) + 0.30)
        # Làm tròn xuống theo 0.5m để bắt được mẫu đầu dạng 2.0/3.0 dù đường cong bắt đầu lệch 0.2-0.3m.
        start = max(0.0, math.floor((min_d + 1e-9) * 2.0) / 2.0)
        if start <= 0.05 and sp >= 1.0:
            start = sp
        # Không bắt đầu quá cao so với điểm đồ thị đầu tiên; nếu bị lệch thì dịch lên điểm gần nhất.
        if min_d - start > max(0.65, sp * 0.45):
            start = min_d
        out: List[Tuple[float, float]] = []
        half_win = max(0.32, min(0.80, sp * 0.22))
        d = start
        guard = 0
        while d <= max_d + 0.20 and guard < 500:
            guard += 1
            local = [n for dd, n in pts0 if abs(dd - d) <= half_win]
            if not local:
                nearest = min(pts0, key=lambda t: abs(t[0] - d))
                if abs(nearest[0] - d) <= max(0.80, sp * 0.55):
                    local = [nearest[1]]
            if local:
                local = sorted(local)
                nval = local[len(local) // 2]
                # Loại điểm 0 cô lập giữa vùng rất lớn, thường do bắt nhầm trục/lưới ở cuối đồ thị.
                if nval <= 1.0 and out:
                    near_future = [n for dd, n in pts0 if 0.0 < dd - d <= sp and n > 20.0]
                    if near_future and out[-1][1] > 20.0:
                        nval = self._spt_graph_fallback_value(max(out[-1][1], sorted(near_future)[len(near_future)//2]))
                out.append((round(float(d), 3), self._spt_graph_fallback_value(nval)))
            d += sp
        # Bỏ trùng rất gần và giữ thứ tự.
        dedup: List[Tuple[float, float]] = []
        for d0, n0 in out:
            if dedup and abs(d0 - dedup[-1][0]) < 0.20:
                dedup[-1] = (dedup[-1][0], n0)
            else:
                dedup.append((d0, n0))
        return dedup


    def _extract_spt_points_from_n30_column_paddle(self, img_obj: Any, xlines: List[int], x_desc1: int, data_top: float, data_bottom: float, total_depth: float) -> List[Tuple[float, float]]:
        """Đọc cột N/30 bằng PaddleOCR khi người dùng chọn PaddleOCR.

        V1.0.48: trước đây PaddleOCR không tác động đến cột SPT vì toàn bộ crop số
        đều bị ép dùng Tesseract cho nhanh. Hàm này chỉ chạy khi Settings chọn PaddleOCR,
        crop riêng các cột N2/N3/N/30 rất hẹp rồi OCR bằng Paddle để tạo khác biệt thực sự
        với nhánh Tesseract. Nếu Paddle không đọc đủ tốt thì trả rỗng để nhánh Tesseract/
        vị trí đồ thị xử lý tiếp.
        """
        try:
            engine_key_n30 = self._borehole_ocr_engine_key()
            if engine_key_n30 not in ("PADDLE", "RAPID"):
                return []
            if data_bottom <= data_top or total_depth <= 0:
                return []
            right_lines = sorted([x for x in xlines if x > x_desc1 + 2])
            if len(right_lines) < 5:
                return []
            reader = self._get_rapidocr_reader() if engine_key_n30 == "RAPID" else self._get_paddleocr_reader(lang="eng")
            if reader is None:
                return []
            w, h = img_obj.size
            y0 = max(0, int(data_top) - 6)
            y1 = min(h, int(data_bottom) + 6)

            def _paddle_col_points(x0_col: int, x1_col: int, *, scale: float = 4.0) -> List[Tuple[float, float]]:
                pts_local: List[Tuple[float, float]] = []
                try:
                    x0 = max(0, int(x0_col) - 4)
                    x1 = min(w, int(x1_col) + 4)
                    if x1 <= x0 + 4 or y1 <= y0 + 4:
                        return []
                    crop = img_obj.crop((x0, y0, x1, y1)).convert("RGB")
                    # Cột số rất hẹp; cho phép phóng lớn nhưng giữ cạnh dài trong giới hạn hợp lý.
                    max_side = max(crop.size)
                    sc = float(scale)
                    if max_side * sc > 3600:
                        sc = max(1.0, 3600.0 / max(max_side, 1))
                    if sc > 1.01:
                        from PIL import Image as _PILImage
                        crop = crop.resize((max(1, int(crop.width * sc)), max(1, int(crop.height * sc))), _PILImage.Resampling.LANCZOS)
                    import numpy as _np  # type: ignore
                    arr = _np.array(crop)
                    with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                        if engine_key_n30 == "RAPID":
                            result = reader(arr)
                            if isinstance(result, tuple):
                                result = result[0]
                        else:
                            try:
                                result = reader.ocr(arr, cls=False)
                            except TypeError:
                                result = reader.ocr(arr)
                    items = (self._parse_rapidocr_result_items(result) if engine_key_n30 == "RAPID"
                             else self._parse_paddleocr_result_items(result))
                    for rec in items:
                        txt_raw = str(rec.get("text", "") or "")
                        # Paddle đôi khi gộp vài số trong cùng một box; tách token nhưng giữ cùng cao độ.
                        token_parts = re.findall(r">\s*\d+|\d+\s*/\s*\d+|\d+", txt_raw)
                        if not token_parts:
                            token_parts = [txt_raw]
                        cy_scaled = float(rec.get("top", 0) or 0) + 0.5 * float(rec.get("height", 0) or 0)
                        cy_img = y0 + cy_scaled / max(sc, 1e-6)
                        if cy_img < data_top - 8 or cy_img > data_bottom + 8:
                            continue
                        depth = (cy_img - data_top) / max(data_bottom - data_top, 1.0) * total_depth
                        for tok in token_parts:
                            val = self._parse_spt_n30_value_from_ocr_token(tok)
                            if val is None:
                                continue
                            if 0.0 <= depth <= total_depth + 0.5:
                                pts_local.append((float(depth), float(val)))
                except Exception:
                    return []
                pts_local.sort(key=lambda t: t[0])
                # Gộp các token rất sát cùng một cao độ.
                grouped_local: List[Tuple[float, float]] = []
                for d, n in pts_local:
                    if not grouped_local or abs(d - grouped_local[-1][0]) > 0.22:
                        grouped_local.append((d, n))
                    else:
                        od, on = grouped_local[-1]
                        # N/30 thì ưu tiên giá trị lớn hơn; với >100 giữ 100.
                        grouped_local[-1] = ((od + d) / 2.0, max(float(on), float(n)))
                return grouped_local

            direct = _paddle_col_points(right_lines[3], right_lines[4], scale=4.5)
            n2_pts = _paddle_col_points(right_lines[1], right_lines[2], scale=4.5)
            n3_pts = _paddle_col_points(right_lines[2], right_lines[3], scale=4.5)

            combo: List[Tuple[float, float]] = []
            used3: set[int] = set()
            for d2, n2 in n2_pts:
                candidates = [(j, abs(d3 - d2), d3, n3) for j, (d3, n3) in enumerate(n3_pts) if j not in used3]
                if not candidates:
                    continue
                j, gap, d3, n3 = min(candidates, key=lambda t: t[1])
                if gap <= 0.45:
                    used3.add(j)
                    nval = self._spt_combo_value_from_n2n3(n2, n3)
                    if nval is None:
                        continue
                    combo.append(((float(d2) + float(d3)) / 2.0, nval))

            pts_tagged: List[Tuple[float, float, str]] = [(d, n, "direct") for d, n in direct]
            for d, n in combo:
                near_i = None
                near_gap = 999.0
                for i0, (d0, n0, tag0) in enumerate(pts_tagged):
                    if tag0 != "direct":
                        continue
                    g = abs(float(d) - float(d0))
                    if g < near_gap:
                        near_i, near_gap = i0, g
                if near_i is not None and near_gap <= 0.45:
                    d0, n0, tag0 = pts_tagged[near_i]
                    # Nếu direct và N2+N3 cùng cao độ chênh nhau ít, lấy trung bình/giá trị hợp lý.
                    if abs(float(n0) - float(n)) <= max(5.0, 0.18 * max(float(n0), float(n), 1.0)):
                        pts_tagged[near_i] = ((float(d0) + float(d)) / 2.0, round((float(n0) + float(n)) / 2.0), "paddle_merge")
                    elif 0.0 < n <= 100.0 and (n0 <= 0 or n0 > 100 or abs(n0 - n) > 25):
                        pts_tagged[near_i] = ((float(d0) + float(d)) / 2.0, float(n), "combo_fix")
                else:
                    pts_tagged.append((d, n, "combo"))
            pts_tagged.sort(key=lambda t: t[0])
            grouped: List[Tuple[float, float, str]] = []
            for d, n, tag in pts_tagged:
                if not grouped or abs(float(d) - float(grouped[-1][0])) > 0.25:
                    grouped.append((float(d), float(n), tag))
                else:
                    od, on, otag = grouped[-1]
                    if tag in ("combo_fix", "paddle_merge"):
                        grouped[-1] = ((od + float(d)) / 2.0, float(n), tag)
                    elif otag in ("combo_fix", "paddle_merge"):
                        grouped[-1] = ((od + float(d)) / 2.0, float(on), otag)
                    else:
                        grouped[-1] = ((od + float(d)) / 2.0, max(float(on), float(n)), tag)
            out = [(d, max(0.0, min(100.0, n))) for d, n, _tag in grouped]
            try:
                out = self._densify_spt_points_by_spacing(out, total_depth, getattr(self, "_borehole_spt_spacing_m", 2.0))
            except Exception:
                pass
            return out
        except Exception:
            return []


    def _densify_spt_points_by_spacing(self, pts: List[Tuple[float, float]], total_depth: float, spacing: float) -> List[Tuple[float, float]]:
        """Không tự chèn/nội suy điểm SPT mặc định.

        SPT là dữ liệu thí nghiệm dùng trực tiếp cho tính SCT. Việc nội suy hoặc kéo dài
        chuỗi 100 xuống đáy có thể tạo dữ liệu giả. Chỉ bật lại khi debug bằng biến môi trường
        TS_CAP_ENABLE_SPT_DENSIFY=1.
        """
        try:
            raw = sorted((float(d), float(n)) for d, n in (pts or []) if 0.0 <= float(d) <= float(total_depth) + 0.50)
        except Exception:
            raw = list(pts or [])
        if str(os.environ.get("TS_CAP_ENABLE_SPT_DENSIFY", "")).strip().lower() not in {"1", "true", "yes", "on"}:
            return raw
        try:
            mode = str(getattr(self, "_borehole_spt_mode", "avg") or "avg")
            if mode != "points":
                return raw
            sp = float(spacing or getattr(self, "_borehole_spt_spacing_m", 2.0) or 2.0)
            if sp <= 0.10:
                sp = 2.0
            if len(raw) < 2:
                return raw
            out: List[Tuple[float, float]] = []
            for i, (d0, n0) in enumerate(raw):
                if not out or abs(d0 - out[-1][0]) > 0.20:
                    out.append((d0, n0))
                else:
                    od, on = out[-1]
                    out[-1] = ((od + d0) / 2.0, on if abs(on) <= abs(n0) else n0)
                if i >= len(raw) - 1:
                    continue
                d1, n1 = raw[i + 1]
                gap = d1 - d0
                if gap <= sp * 1.45 or gap > sp * 8.0:
                    continue
                k = 1
                while d0 + k * sp < d1 - sp * 0.45 and k < 30:
                    dn = d0 + k * sp
                    if dn <= d0 + 0.20 or dn >= d1 - 0.20:
                        k += 1
                        continue
                    t = (dn - d0) / max(d1 - d0, 1e-9)
                    nv = n0 + (n1 - n0) * t
                    out.append((round(dn, 3), max(0.0, min(100.0, float(nv)))))
                    k += 1
            out.sort(key=lambda t: t[0])
            return out
        except Exception:
            return raw

    def _merge_spt_ocr_values_to_graph_depths(self, ocr_pts: List[Tuple[float, float]], graph_pts: List[Tuple[float, float]], spacing: float) -> List[Tuple[float, float]]:
        """Ghép giá trị OCR với vị trí/giá trị đồ thị SPT.

        Ưu tiên số OCR ở cột N/30 hoặc N2+N3 nếu có điểm gần cùng cao độ. Nếu OCR không
        đọc được số tại vị trí đó, dùng lại giá trị đọc từ đường cong SPT thay vì ghi 0.
        Cách này tránh lỗi V1.0.48 tạo hàng loạt SPT=0 khi ảnh scan quá nhỏ/màu xanh.
        """
        try:
            ocr = sorted((float(d), float(n)) for d, n in (ocr_pts or []) if float(n) > 0)
            graph = sorted((float(d), float(n)) for d, n in (graph_pts or []))
            if not graph:
                return ocr
            tol = max(0.55, min(1.0, float(spacing or 2.0) * 0.35))
            out: List[Tuple[float, float]] = []
            used: set[int] = set()
            for gd, _gn in graph:
                match_i = None
                match_gap = 999.0
                for i, (od, on) in enumerate(ocr):
                    if i in used:
                        continue
                    g = abs(od - gd)
                    if g < match_gap:
                        match_i, match_gap = i, g
                if match_i is not None and match_gap <= tol:
                    used.add(match_i)
                    out.append((gd, ocr[match_i][1]))
                else:
                    # Không để SPT=0 hàng loạt khi OCR không đọc được số; dùng giá trị đồ thị
                    # làm fallback để người dùng có dữ liệu gần đúng thay vì bảng trống.
                    out.append((gd, self._spt_graph_fallback_value(_gn)))
            # Thêm các điểm OCR thật không có trong graph.
            for i, (od, on) in enumerate(ocr):
                if i not in used:
                    out.append((od, on))
            out.sort(key=lambda t: t[0])
            # Gộp trùng rất gần.
            dedup: List[Tuple[float, float]] = []
            for d, n in out:
                if dedup and abs(d - dedup[-1][0]) < 0.25:
                    od, on = dedup[-1]
                    dedup[-1] = ((od + d) / 2.0, on if on > 0 else n)
                else:
                    dedup.append((d, n))
            try:
                max_depth = max([d for d, _n in graph] + [d for d, _n in ocr]) if (graph or ocr) else 0.0
                dedup = self._densify_spt_points_by_spacing(dedup, max_depth, spacing)
            except Exception:
                pass
            return dedup
        except Exception:
            return ocr_pts or []


    def _spt_numeric_ocr_variants(self, img_obj: Any, *, scale: int = 8) -> List[Any]:
        """Tạo các biến thể OCR cho ô/cột số SPT scan, không phụ thuộc màu chữ.

        Nguyên tắc: OCR bảng scan phải dựa vào hình học và tách nhiễu đường kẻ,
        không dựa vào màu cụ thể như xanh/đỏ/đen. Hàm này tạo nhiều ảnh nhị phân:
        - ngưỡng thích nghi xám, loại đường kẻ dài;
        - ngưỡng Otsu xám, loại đường kẻ dài;
        - mặt nạ chữ có độ bão hòa màu cao nhưng không khóa vào hue nào.
        Các ảnh giữ nguyên kích thước gốc rồi mới phóng to, để tọa độ OCR còn map
        được về cao độ/độ sâu.
        """
        variants: List[Any] = []
        try:
            import cv2  # type: ignore
            import numpy as np  # type: ignore
            from PIL import Image as _PILImage, ImageOps, ImageEnhance
            pil_rgb = img_obj.convert("RGB") if hasattr(img_obj, "convert") else img_obj
            arr_rgb = np.array(pil_rgb)
            if arr_rgb.size == 0:
                return []
            gray = cv2.cvtColor(arr_rgb, cv2.COLOR_RGB2GRAY)

            def _remove_lines_and_make(mask: Any) -> Optional[Any]:
                try:
                    mh, mw = mask.shape[:2]
                    # loại đường kẻ dài tương đối theo kích thước crop
                    vlen = max(10, int(mh * 0.22))
                    hlen = max(10, int(mw * 0.45))
                    vert = cv2.morphologyEx(mask, cv2.MORPH_OPEN, cv2.getStructuringElement(cv2.MORPH_RECT, (1, vlen)))
                    horiz = cv2.morphologyEx(mask, cv2.MORPH_OPEN, cv2.getStructuringElement(cv2.MORPH_RECT, (hlen, 1)))
                    clean = cv2.subtract(mask, cv2.bitwise_or(vert, horiz))
                    num, labels, stats, _cent = cv2.connectedComponentsWithStats(clean, 8)
                    comp = np.zeros_like(clean)
                    for lab in range(1, num):
                        x, y, w0, h0, area = stats[lab]
                        # Giữ nét chữ/số/dấu >, /; bỏ nhiễu lớn, đường dài.
                        if area >= 2 and h0 >= 2 and w0 <= max(24, int(mw * 0.80)) and h0 <= max(26, int(mh * 0.65)):
                            comp[labels == lab] = 255
                    if int(comp.sum()) <= 0:
                        return None
                    # Nối nhẹ nét đứt nhưng không làm dính sang ô bên cạnh.
                    comp = cv2.dilate(comp, cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1)), iterations=1)
                    out = 255 - comp
                    im = _PILImage.fromarray(out).convert("L")
                    sc = max(1, int(scale or 1))
                    try:
                        im = im.resize((max(1, im.width * sc), max(1, im.height * sc)), _PILImage.Resampling.LANCZOS)
                    except Exception:
                        im = im.resize((max(1, im.width * sc), max(1, im.height * sc)))
                    return im
                except Exception:
                    return None

            # 1. Adaptive threshold: tốt cho scan mờ/không đều sáng.
            for block, cval in ((15, 8), (21, 10), (31, 12)):
                try:
                    mask = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, block, cval)
                    im = _remove_lines_and_make(mask)
                    if im is not None:
                        variants.append(im)
                except Exception:
                    pass
            # 2. Otsu/autocontrast: tốt cho PDF scan rõ nét.
            try:
                _thr, mask = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
                im = _remove_lines_and_make(mask)
                if im is not None:
                    variants.append(im)
            except Exception:
                pass
            # 3. Mặt nạ mực màu bất kỳ: không khóa hue; dùng độ bão hòa/khác biệt kênh màu.
            try:
                arr = arr_rgb.astype(np.int16)
                mx = arr.max(axis=2); mn = arr.min(axis=2)
                color_delta = mx - mn
                gray_like = gray.astype(np.int16)
                mask_color = ((color_delta > 18) & (gray_like < 245)).astype(np.uint8) * 255
                im = _remove_lines_and_make(mask_color)
                if im is not None:
                    variants.append(im)
            except Exception:
                pass
            # 4. Bản gốc xám tăng tương phản, để cứu các ký tự bị biến mất khi tách lưới.
            try:
                im0 = ImageOps.grayscale(pil_rgb)
                im0 = ImageOps.autocontrast(im0)
                im0 = ImageEnhance.Contrast(im0).enhance(2.2)
                sc = max(1, int(scale or 1))
                im0 = im0.resize((max(1, im0.width * sc), max(1, im0.height * sc)))
                variants.append(im0)
            except Exception:
                pass
            # Giới hạn số biến thể để không làm chậm quá mức.
            return variants[:6]
        except Exception:
            return []

    def _ocr_generic_spt_cell_candidates(self, img_obj: Any) -> List[float]:
        """OCR một ô số SPT scan theo cách tổng quát, không phụ thuộc màu chữ.

        Hàm trả về nhiều ứng viên để bước sau kiểm tra bằng N = N2 + N3 và/hoặc
        giá trị N/30 trực tiếp. Không dùng đồ thị để bịa số khi OCR ô không chắc.
        """
        candidates: List[float] = []
        try:
            seen_txt: set[str] = set()
            variants = self._spt_numeric_ocr_variants(img_obj, scale=10)
            if not variants:
                variants = [self._preprocess_crop_for_ocr(img_obj, scale=10, numeric=True)]
            for imv in variants[:6]:
                for psm in (7, 8, 10):
                    try:
                        txt = self._ocr_pil_crop(imv, lang="eng", psm=psm, whitelist="0123456789/>", timeout=5)
                    except Exception:
                        txt = ""
                    for part in re.split(r"[\s\n\r]+", str(txt or "")):
                        part = part.strip()
                        if not part or part in seen_txt:
                            continue
                        seen_txt.add(part)
                        val2 = self._parse_spt_n30_value_from_ocr_token(part)
                        if val2 is not None and 0.0 < float(val2) <= 100.0:
                            candidates.append(float(val2))
            return candidates
        except Exception:
            return []


    def _ocr_blue_spt_cell_candidates(self, img_obj: Any) -> List[float]:
        """Tương thích tên hàm cũ: OCR một ô số SPT scan bằng phương pháp tổng quát.

        V1.0.51 không còn khóa vào màu xanh. Trước hết dùng tách lưới/mực tổng quát;
        nhánh màu chỉ là fallback cho các bản scan cũ.
        """
        try:
            gen = self._ocr_generic_spt_cell_candidates(img_obj)
            if gen:
                return gen
        except Exception:
            pass
        candidates: List[float] = []
        try:
            import cv2  # type: ignore
            import numpy as np  # type: ignore
            from PIL import Image as _PILImage
            arr = np.array(img_obj.convert("RGB") if hasattr(img_obj, "convert") else img_obj)
            if arr.size == 0:
                return []
            hsv = cv2.cvtColor(arr, cv2.COLOR_RGB2HSV)
            h, sat, val = hsv[:, :, 0], hsv[:, :, 1], hsv[:, :, 2]
            mask = ((h > 85) & (h < 135) & (sat > 32) & (val > 45)).astype(np.uint8) * 255
            if int(mask.sum()) <= 0:
                return []
            # Bỏ nhiễu nhỏ nhưng giữ dấu '>' và '/'.
            num_labels, labels, stats, _cent = cv2.connectedComponentsWithStats(mask, 8)
            clean = np.zeros_like(mask)
            for lab in range(1, num_labels):
                x, y, w, hh, area = stats[lab]
                if area >= 2 and hh >= 2:
                    clean[labels == lab] = 255
            if int(clean.sum()) <= 0:
                return []
            variants = []
            for kx, ky in ((1, 1), (2, 1), (1, 2)):
                mm = cv2.dilate(clean, np.ones((ky, kx), np.uint8), iterations=1)
                im = _PILImage.fromarray(255 - mm).convert("L")
                aa = np.array(im)
                ys, xs = np.where(aa < 210)
                if len(xs) == 0:
                    continue
                im = im.crop((max(0, int(xs.min()) - 2), max(0, int(ys.min()) - 2), min(im.width, int(xs.max()) + 3), min(im.height, int(ys.max()) + 3)))
                for sc in (10, 14):
                    try:
                        imr = im.resize((max(1, int(im.width * sc)), max(1, int(im.height * sc))), _PILImage.Resampling.LANCZOS)
                    except Exception:
                        imr = im.resize((max(1, int(im.width * sc)), max(1, int(im.height * sc))))
                    canvas = _PILImage.new("L", (imr.width + 44, imr.height + 44), 255)
                    canvas.paste(imr, (22, 22))
                    variants.append(canvas)
            seen_txt = set()
            for imv in variants[:10]:
                for psm in (7, 8, 10):
                    try:
                        txt = self._ocr_pil_crop(imv, lang="eng", psm=psm, whitelist="0123456789/>", timeout=4)
                    except Exception:
                        txt = ""
                    txt = str(txt or "").strip()
                    if not txt or txt in seen_txt:
                        continue
                    seen_txt.add(txt)
                    val2 = self._parse_spt_n30_value_from_ocr_token(txt)
                    if val2 is not None and 0.0 < float(val2) <= 100.0:
                        candidates.append(float(val2))
            # Gộp trùng và sắp theo tần suất. Không chọn ở đây vì cần so với N2+N3.
            return candidates
        except Exception:
            return []

    def _choose_spt_value_from_candidates(self, direct_vals: List[float], n2_vals: List[float], n3_vals: List[float]) -> Optional[float]:
        """Chọn N/30 từ OCR trực tiếp và/hoặc N2+N3, không tự nâng 0 lên số lớn.

        Nguyên tắc an toàn:
        - 0 là giá trị hợp lệ nếu đọc trực tiếp trong bảng N2/N3/N/30.
        - Nếu N2+N3 ≈ 0 nhưng cột N/30 bị OCR thành 10/47/50/100 thì ưu tiên N2+N3.
        - Không dùng đồ thị để sửa/nâng giá trị.
        """
        def _best(vals: List[float]) -> Optional[float]:
            vals2 = [float(v) for v in (vals or []) if 0.0 <= float(v) <= 100.0]
            if not vals2:
                return None
            counts: Dict[int, int] = {}
            for v in vals2:
                counts[int(round(v))] = counts.get(int(round(v)), 0) + 1
            # Nếu có nhiều giá trị khác nhau, ưu tiên tần suất; không ưu tiên số lớn một cách mù quáng.
            best = sorted(counts.items(), key=lambda kv: (kv[1], -abs(kv[0]), -kv[0]), reverse=True)[0][0]
            return float(best)
        direct = _best(direct_vals)
        n2 = _best(n2_vals)
        n3 = _best(n3_vals)
        combo: Optional[float] = None
        if n2 is not None and n3 is not None:
            combo = self._spt_combo_value_from_n2n3(n2, n3)
        if direct is not None and combo is not None:
            # N2+N3 cùng rất nhỏ/zero là bằng chứng mạnh hơn một OCR N/30 bị dính nét.
            if combo <= 2.0 and direct >= 8.0:
                return combo
            if 45.0 <= direct <= 55.0 and combo >= 90.0:
                return direct
            if abs(direct - combo) <= max(4.0, 0.18 * max(direct, combo, 1.0)):
                return combo
            if direct < 10.0 and combo >= 10.0:
                return combo
            if combo >= 0.0 and abs(direct - combo) > 25.0:
                return combo
            return direct
        if combo is not None:
            return combo
        return direct

    def _extract_spt_points_from_blue_text_rows(self, img_obj: Any, xlines: List[int], x_desc1: int, data_top: float, data_bottom: float, total_depth: float) -> List[Tuple[float, float]]:
        """Đọc SPT từ các hàng chữ xanh trong cụm N1/N2/N3/N/30.

        Đây là nhánh riêng cho ảnh scan màu: lấy màu xanh để xác định từng hàng SPT,
        OCR từng ô nhỏ N2, N3, N/30 rồi chọn giá trị. Không dùng đồ thị để gán số nếu
        OCR ô đã đọc được.
        """
        try:
            import cv2  # type: ignore
            import numpy as np  # type: ignore
            if data_bottom <= data_top or total_depth <= 0:
                return []
            right_lines = sorted([int(x) for x in xlines if x > x_desc1 + 2])
            if len(right_lines) < 5:
                return []
            w, h = img_obj.size
            y0 = max(0, int(data_top) - 8)
            y1 = min(h, int(data_bottom) + 8)
            # N1..N/30: từ sau cột độ sâu đến trước đồ thị.
            x0 = max(0, int(right_lines[0]) - 3)
            x1 = min(w, int(right_lines[4]) + 3)
            if x1 <= x0 + 10 or y1 <= y0 + 10:
                return []
            crop = img_obj.crop((x0, y0, x1, y1)).convert("RGB")
            arr = np.array(crop)
            hsv = cv2.cvtColor(arr, cv2.COLOR_RGB2HSV)
            hh, ss, vv = hsv[:, :, 0], hsv[:, :, 1], hsv[:, :, 2]
            mask = ((hh > 85) & (hh < 135) & (ss > 32) & (vv > 45)).astype(np.uint8)
            # Xóa các pixel xanh của lưới/đồ thị nếu có, chỉ crop cột số nên chủ yếu còn chữ số.
            row_sum = mask.sum(axis=1)
            ys = np.where(row_sum > 2)[0]
            if len(ys) < 3:
                return []
            groups: List[Tuple[int, int]] = []
            st = int(ys[0]); prev = int(ys[0])
            for yy in ys[1:]:
                yy = int(yy)
                if yy <= prev + 2:
                    prev = yy
                else:
                    if prev - st >= 2:
                        groups.append((st, prev))
                    st = prev = yy
            if prev - st >= 2:
                groups.append((st, prev))
            if len(groups) < 3:
                return []
            cols_abs = [
                (right_lines[1], right_lines[2]),  # N2
                (right_lines[2], right_lines[3]),  # N3
                (right_lines[3], right_lines[4]),  # N/30
            ]
            pts: List[Tuple[float, float]] = []
            for a, b in groups:
                cy = y0 + (float(a) + float(b)) / 2.0
                depth = (cy - float(data_top)) / max(float(data_bottom) - float(data_top), 1.0) * float(total_depth)
                if not (0.0 <= depth <= float(total_depth) + 0.50):
                    continue
                yy0 = max(0, y0 + int(a) - 7)
                yy1 = min(h, y0 + int(b) + 8)
                vals_cols: List[List[float]] = []
                for xa, xb in cols_abs:
                    cell = img_obj.crop((max(0, int(xa) - 4), yy0, min(w, int(xb) + 4), yy1))
                    vals_cols.append(self._ocr_blue_spt_cell_candidates(cell))
                nval = self._choose_spt_value_from_candidates(vals_cols[2], vals_cols[0], vals_cols[1])
                if nval is not None and nval > 0:
                    pts.append((round(float(depth), 3), max(0.0, min(100.0, float(nval)))))
            # Lọc trùng theo độ sâu gần nhau.
            pts.sort(key=lambda t: t[0])
            out: List[Tuple[float, float]] = []
            for d, n in pts:
                if out and abs(d - out[-1][0]) < 0.25:
                    od, on = out[-1]
                    out[-1] = ((od + d) / 2.0, max(on, n))
                else:
                    out.append((d, n))
            return out
        except Exception:
            return []


    def _extract_spt_points_from_blue_column_ocr(self, img_obj: Any, xlines: List[int], x_desc1: int, data_top: float, data_bottom: float, total_depth: float) -> List[Tuple[float, float]]:
        """Đọc cột số SPT màu xanh bằng OCR theo cột, sau đó map theo hàng.

        Hàm này nhanh hơn OCR từng ô: mỗi cột N2/N3/N/30 chỉ OCR một lần. Các token
        được gán về hàng SPT theo tọa độ y của chữ xanh. Nếu cột N/30 thiếu số tại một
        hàng thì dùng N2+N3 hoặc giá trị đồ thị gần cùng cao độ để tránh sinh 0/trống.
        """
        try:
            import cv2  # type: ignore
            import numpy as np  # type: ignore
            from PIL import Image as _PILImage
            right_lines = sorted([int(x) for x in xlines if x > x_desc1 + 2])
            if len(right_lines) < 5 or data_bottom <= data_top or total_depth <= 0:
                return []
            w, h = img_obj.size
            y0 = max(0, int(data_top) - 8)
            y1 = min(h, int(data_bottom) + 8)
            # Dò hàng theo toàn bộ cụm N1..N/30.
            xg0 = max(0, int(right_lines[0]) - 3)
            xg1 = min(w, int(right_lines[4]) + 3)
            crop_all = img_obj.crop((xg0, y0, xg1, y1)).convert("RGB")
            # V1.0.51: dò hàng số bằng mặt nạ mực tổng quát, không khóa vào màu xanh.
            mask_all = None
            try:
                variants0 = self._spt_numeric_ocr_variants(crop_all, scale=1)
                if variants0:
                    arr0 = np.array(variants0[0].convert("L") if hasattr(variants0[0], "convert") else variants0[0])
                    mask_all = (arr0 < 210).astype(np.uint8)
            except Exception:
                mask_all = None
            if mask_all is None or int(mask_all.sum()) <= 0:
                # Fallback cũ cho bản scan có chữ số màu bão hòa. Không phụ thuộc hue xanh,
                # chỉ chọn pixel có độ bão hòa/khác biệt kênh màu rõ.
                arr_all = np.array(crop_all)
                arr_i = arr_all.astype(np.int16)
                mx = arr_i.max(axis=2); mn = arr_i.min(axis=2)
                gray0 = cv2.cvtColor(arr_all, cv2.COLOR_RGB2GRAY)
                mask_all = (((mx - mn) > 18) & (gray0 < 245)).astype(np.uint8)
            row_sum = mask_all.sum(axis=1)
            ys = np.where(row_sum > 2)[0]
            if len(ys) < 3:
                return []
            row_groups: List[Tuple[int, int]] = []
            st = int(ys[0]); prev = int(ys[0])
            for yy0 in ys[1:]:
                yy = int(yy0)
                if yy <= prev + 2:
                    prev = yy
                else:
                    if prev - st >= 2:
                        row_groups.append((st, prev))
                    st = prev = yy
            if prev - st >= 2:
                row_groups.append((st, prev))
            if not row_groups:
                return []
            row_centers = [y0 + (float(a) + float(b)) / 2.0 for a, b in row_groups]

            def _ocr_blue_column(xa: int, xb: int) -> List[Tuple[float, float, str]]:
                try:
                    x0 = max(0, int(xa) - 4); x1 = min(w, int(xb) + 4)
                    if x1 <= x0 + 4:
                        return []
                    crop = img_obj.crop((x0, y0, x1, y1)).convert("RGB")
                    variants = self._spt_numeric_ocr_variants(crop, scale=8)
                    if not variants:
                        return []
                    # Dùng biến thể đã lọc lưới đầu tiên cho OCR theo cột; các biến thể còn lại
                    # đã được dùng trong OCR từng ô nếu cần.
                    im = variants[0]
                    canvas = _PILImage.new("L", (im.width + 50, im.height + 50), 255)
                    canvas.paste(im, (25, 25))
                    data = self._ocr_pil_crop_data(canvas, lang="eng", psm=6, whitelist="0123456789/>", timeout=8)
                    toks: List[Tuple[float, float, str]] = []
                    for rec in data:
                        val2 = self._parse_spt_n30_value_from_ocr_token(rec.get("text"))
                        if val2 is None:
                            continue
                        cy_scaled = float(rec.get("top", 0) or 0) + 0.5 * float(rec.get("height", 0) or 0)
                        cy_img = y0 + (cy_scaled - 25.0) / 8.0
                        if cy_img < data_top - 10 or cy_img > data_bottom + 10:
                            continue
                        toks.append((float(cy_img), max(0.0, min(100.0, float(val2))), str(rec.get("text", ""))))
                    toks.sort(key=lambda t: t[0])
                    return toks
                except Exception:
                    return []

            n2_toks = _ocr_blue_column(right_lines[1], right_lines[2])
            n3_toks = _ocr_blue_column(right_lines[2], right_lines[3])
            n30_toks = _ocr_blue_column(right_lines[3], right_lines[4])
            graph_sampled: List[Tuple[float, float]] = []
            try:
                sp = float(getattr(self, "_borehole_spt_spacing_m", 2.0) or 2.0)
            except Exception:
                sp = 2.0
            if sp <= 0.10:
                sp = 2.0
            try:
                graph_pts = self._extract_spt_points_from_graph_curve(img_obj, xlines, x_desc1, data_top, data_bottom, total_depth)
                graph_sampled = self._sample_spt_graph_points_by_spacing(graph_pts, total_depth, sp)
            except Exception:
                graph_sampled = []

            def _nearest_token(toks: List[Tuple[float, float, str]], cy: float, tol_px: float = 9.0) -> Optional[float]:
                if not toks:
                    return None
                best = min(toks, key=lambda t: abs(float(t[0]) - float(cy)))
                if abs(float(best[0]) - float(cy)) <= tol_px:
                    return float(best[1])
                return None

            def _nearest_graph(depth: float) -> Optional[float]:
                if not graph_sampled:
                    return None
                best = min(graph_sampled, key=lambda t: abs(float(t[0]) - float(depth)))
                if abs(float(best[0]) - float(depth)) <= max(0.80, sp * 0.55):
                    return float(best[1])
                return None

            out: List[Tuple[float, float]] = []
            for cy in row_centers:
                depth = (float(cy) - float(data_top)) / max(float(data_bottom) - float(data_top), 1.0) * float(total_depth)
                if not (0.0 <= depth <= float(total_depth) + 0.50):
                    continue
                direct = _nearest_token(n30_toks, cy)
                n2 = _nearest_token(n2_toks, cy)
                n3 = _nearest_token(n3_toks, cy)
                combo = None
                if n2 is not None and n3 is not None:
                    combo = self._spt_combo_value_from_n2n3(n2, n3)
                gval = _nearest_graph(depth)
                val_final: Optional[float] = None
                if direct is not None and combo is not None:
                    if abs(direct - combo) <= max(4.0, 0.18 * max(direct, combo, 1.0)):
                        val_final = combo
                    elif direct < 10.0 and combo >= 10.0:
                        val_final = combo
                    elif abs(direct - combo) > 25.0:
                        val_final = combo
                    else:
                        val_final = direct
                elif direct is not None:
                    val_final = direct
                elif combo is not None:
                    val_final = combo
                # Đồ thị SPT mặc định tắt; không dùng đồ thị để tự nâng/sinh SPT nếu OCR bảng không chắc.
                if gval is not None and str(os.environ.get("TS_CAP_ENABLE_SPT_GRAPH_FALLBACK", "")).strip().lower() in {"1", "true", "yes", "on"}:
                    if val_final is None:
                        val_final = self._spt_graph_fallback_value(gval)
                    elif val_final < 10.0 and gval >= 12.0:
                        val_final = self._spt_graph_fallback_value(gval)
                    elif abs(val_final - gval) > max(20.0, 0.45 * max(val_final, gval, 1.0)) and combo is None:
                        val_final = self._spt_graph_fallback_value(gval)
                if val_final is not None:
                    out.append((round(float(depth), 3), max(0.0, min(100.0, float(val_final)))))
            out.sort(key=lambda t: t[0])
            # Gộp trùng gần nhau.
            dedup: List[Tuple[float, float]] = []
            for d, n in out:
                if dedup and abs(d - dedup[-1][0]) < 0.25:
                    od, on = dedup[-1]
                    dedup[-1] = ((od + d) / 2.0, max(on, n))
                else:
                    dedup.append((d, n))
            return dedup
        except Exception:
            return []


    def _refine_spt_points_with_graph_check(self, pts: List[Tuple[float, float]], img_obj: Any, xlines: List[int], x_desc1: int, data_top: float, data_bottom: float, total_depth: float, spacing: float) -> List[Tuple[float, float]]:
        """Dùng đồ thị SPT để kiểm tra các giá trị OCR rõ ràng bất thường.

        Không thay toàn bộ bằng đồ thị; chỉ sửa các trường hợp OCR mất chữ số hàng chục
        hoặc đọc >100 thành 10/52... trong vùng đường cong cho giá trị lớn rõ ràng.
        """
        try:
            raw = sorted((float(d), float(n)) for d, n in (pts or []) if 0.0 <= float(d) <= float(total_depth) + 0.50)
            if not raw:
                return []
            graph_pts = self._extract_spt_points_from_graph_curve(img_obj, xlines, x_desc1, data_top, data_bottom, total_depth)
            graph_sampled = self._sample_spt_graph_points_by_spacing(graph_pts, total_depth, spacing)
            if not graph_sampled:
                return raw
            tol = max(0.80, float(spacing or 2.0) * 0.55)
            out: List[Tuple[float, float]] = []
            for d, n in raw:
                gd, gn = min(graph_sampled, key=lambda t: abs(float(t[0]) - float(d)))
                if abs(float(gd) - float(d)) > tol:
                    out.append((d, n)); continue
                g = float(gn)
                v = float(n)
                # Các điều kiện sửa rất bảo thủ: chỉ nâng giá trị khi OCR có dấu hiệu mất chữ số.
                if v <= 0.0:
                    v = g
                elif v <= 10.0 and g >= v + 14.0:
                    # Chỉ sửa khi OCR rõ ràng mất chữ số hàng chục, ví dụ 54 đọc thành 5.
                    v = g
                elif g >= 90.0 and v < 20.0:
                    # Vùng xuyên đá/rất chặt: OCR thường mất dấu >100 hoặc chỉ còn 10.
                    v = g
                out.append((d, max(0.0, min(100.0, v))))
            return out
        except Exception:
            return pts or []

    def _extract_spt_points_from_n30_column(self, img_obj: Any, xlines: List[int], x_desc1: int, data_top: float, data_bottom: float, total_depth: float) -> List[Tuple[float, float]]:
        """Đọc cột N/30 trong vùng thí nghiệm SPT bằng OCR theo tọa độ.

        V1.0.11: vẫn ưu tiên cột N/30, nhưng nếu OCR bỏ sót nhiều số thì đọc thêm N2 và N3
        rồi lấy N=N2+N3 theo định nghĩa SPT. Cách này giảm lỗi với ảnh hẹp, nơi Tesseract chỉ
        đọc được vài giá trị ở cột N/30.
        """
        def _col_points(x0_col: int, x1_col: int, *, scale: int = 8) -> List[Tuple[float, float]]:
            pts_local: List[Tuple[float, float]] = []
            try:
                w, h = img_obj.size
                y0 = max(0, int(data_top) - 4)
                y1 = min(h, int(data_bottom) + 4)
                # Nới nhẹ hai bên nhưng không lấn quá nhiều sang cột kế bên.
                x0 = max(0, int(x0_col) - 2)
                x1 = min(w, int(x1_col) + 2)
                if x1 <= x0 + 5 or y1 <= y0 + 5:
                    return []
                col = img_obj.crop((x0, y0, x1, y1))
                # V1.0.51: OCR nhiều biến thể đã tách đường kẻ, không phụ thuộc màu chữ.
                variants = self._spt_numeric_ocr_variants(col, scale=scale) or [self._preprocess_crop_for_ocr(col, scale=scale, numeric=True)]
                seen_keys: set[Tuple[int, int]] = set()
                for proc in variants[:5]:
                    data = self._ocr_pil_crop_data(proc, lang="eng", psm=6, whitelist="0123456789.,>/<", timeout=10)
                    for rec in data:
                        val = self._parse_spt_n30_value_from_ocr_token(rec.get("text"))
                        if val is None:
                            continue
                        cy_scaled = float(rec.get("top", 0) or 0) + 0.5 * float(rec.get("height", 0) or 0)
                        cy_img = y0 + cy_scaled / float(scale)
                        if cy_img < data_top - 6 or cy_img > data_bottom + 6:
                            continue
                        depth = (cy_img - data_top) / max(data_bottom - data_top, 1.0) * total_depth
                        if 0.0 <= depth <= total_depth + 0.5:
                            key = (int(round(float(depth) * 10)), int(round(float(val))))
                            if key in seen_keys:
                                continue
                            seen_keys.add(key)
                            pts_local.append((float(depth), float(val)))
            except Exception:
                return []
            pts_local.sort(key=lambda t: t[0])
            return pts_local

        try:
            if data_bottom <= data_top or total_depth <= 0:
                return []
            right_lines = sorted([x for x in xlines if x > x_desc1 + 2])
            # Form BRITEC: sau cột mô tả lần lượt là Độ sâu, N1, N2, N3, N/30, rồi tới đồ thị.
            # x_desc1 là biên trái cột Độ sâu; right_lines[0] là biên phải cột Độ sâu.
            if len(right_lines) < 5:
                return []
            try:
                sp_for_expected = float(getattr(self, "_borehole_spt_spacing_m", 2.0) or 2.0)
            except Exception:
                sp_for_expected = 2.0
            if sp_for_expected <= 0.10:
                sp_for_expected = 2.0
            mode_for_expected = str(getattr(self, "_borehole_spt_mode", "avg") or "avg")
            expected_for_ocr = max(3, int(max(float(total_depth), 0.0) / max(sp_for_expected, 0.1) * (0.45 if mode_for_expected == "points" else 0.25)))

            # Ảnh scan: đọc từng hàng số N2/N3/N/30 theo lưới, không phụ thuộc màu.
            blue_row_pts = self._extract_spt_points_from_blue_column_ocr(img_obj, xlines, x_desc1, data_top, data_bottom, total_depth)

            # Nếu người dùng chọn PaddleOCR, thử đọc riêng các cột số SPT bằng Paddle trước.
            # Nếu Paddle không đọc đủ thì vẫn ghép với blue-row/Tesseract, không dùng kết quả 0.
            paddle_pts_for_merge: List[Tuple[float, float]] = []
            if self._borehole_ocr_engine_key() in ("PADDLE", "RAPID"):
                paddle_pts_for_merge = self._extract_spt_points_from_n30_column_paddle(img_obj, xlines, x_desc1, data_top, data_bottom, total_depth)
                if len(paddle_pts_for_merge) >= expected_for_ocr:
                    return paddle_pts_for_merge

            # OCR trực tiếp cột N/30.
            direct = _col_points(right_lines[3], right_lines[4], scale=8)
            # Đưa ứng viên đọc theo hàng màu xanh và Paddle vào danh sách để ghép với đồ thị/cleanup.
            if blue_row_pts:
                direct = list(direct or []) + list(blue_row_pts)
            if paddle_pts_for_merge:
                direct = list(direct or []) + list(paddle_pts_for_merge)

            # Fallback bổ trợ: đọc N2 và N3, ghép theo cùng cao độ, N/30 = N2 + N3.
            # V1.0.17 tăng tolerance nhẹ vì PDF AutoCAD render ra chữ rất sát đường kẻ;
            # đồng thời dùng N2+N3 để sửa các token N/30 bị dính thành 24/45... khi combo rõ ràng nhỏ hơn.
            n2_pts = _col_points(right_lines[1], right_lines[2], scale=8)
            n3_pts = _col_points(right_lines[2], right_lines[3], scale=8)
            combo: List[Tuple[float, float]] = []
            used3: set[int] = set()
            for d2, n2 in n2_pts:
                if not n3_pts:
                    break
                candidates = [(j, abs(d3 - d2), d3, n3) for j, (d3, n3) in enumerate(n3_pts) if j not in used3]
                if not candidates:
                    continue
                j, gap, d3, n3 = min(candidates, key=lambda t: t[1])
                if gap <= 0.45:
                    used3.add(j)
                    nval = self._spt_combo_value_from_n2n3(n2, n3)
                    if nval is None:
                        continue
                    combo.append(((float(d2) + float(d3)) / 2.0, nval))

            # Ghép: ưu tiên N/30 đọc trực tiếp, nhưng nếu direct quá lớn bất thường so với
            # N2+N3 gần cùng cao độ thì thay bằng combo. Điều này sửa lỗi OCR đọc '2' thành '24'.
            pts_tagged: List[Tuple[float, float, str]] = [(d, n, "direct") for d, n in direct]
            for d, n in combo:
                near_i = None
                near_gap = 999.0
                for i0, (d0, n0, tag0) in enumerate(pts_tagged):
                    g = abs(float(d) - float(d0))
                    if tag0 == "direct" and g < near_gap:
                        near_i, near_gap = i0, g
                if near_i is not None and near_gap <= 0.45:
                    d0, n0, tag0 = pts_tagged[near_i]
                    # Nếu N2+N3 rất nhỏ/zero nhưng N/30 OCR ra 10/47/50..., coi N/30 là dính nét.
                    if 0.0 <= n <= 2.0 and 8.0 <= n0 <= 100.0:
                        pts_tagged[near_i] = ((float(d0) + float(d)) / 2.0, float(n), "combo_fix")
                    # Giữ 100/>100; chỉ sửa khi direct lớn hơn combo quá nhiều và combo hợp lý.
                    elif 0.0 < n <= 25.0 and 20.0 < n0 < 80.0 and n0 >= max(n * 2.5, n + 12.0):
                        pts_tagged[near_i] = ((float(d0) + float(d)) / 2.0, float(n), "combo_fix")
                else:
                    pts_tagged.append((d, n, "combo"))
            pts_tagged.sort(key=lambda t: t[0])

            # QA-OCR SPT-NO-GUESS: loại số N/30 trực tiếp lớn nhưng không có N2+N3 hỗ trợ.
            # Đây là nguồn tạo các số 47/50/100 từ nét 0 + đường kẻ trong ảnh scan.
            try:
                combo_depths_all = [float(cd) for cd, _cn in combo]
                filtered_tagged: List[Tuple[float, float, str]] = []
                for d0, n0, tag0 in pts_tagged:
                    if tag0 == "direct" and float(n0) >= 8.0:
                        if not any(abs(float(d0) - cd) <= 0.50 for cd in combo_depths_all):
                            continue
                    filtered_tagged.append((d0, n0, tag0))
                pts_tagged = filtered_tagged
            except Exception:
                pass

            # Gộp các token trùng gần cùng cao độ: ưu tiên combo_fix; không ưu tiên số lớn một cách mù quáng.
            grouped_tagged: List[Tuple[float, float, str]] = []
            for d, n, tag in pts_tagged:
                if not grouped_tagged or abs(d - grouped_tagged[-1][0]) > 0.20:
                    grouped_tagged.append((d, n, tag))
                else:
                    od, on, otag = grouped_tagged[-1]
                    if tag == "combo_fix" or otag == "combo_fix":
                        keep_n = n if tag == "combo_fix" else on
                        keep_tag = "combo_fix"
                    elif otag == "combo" or tag == "combo":
                        keep_n = on if otag == "combo" else n
                        keep_tag = "combo"
                    else:
                        # Cùng một cao độ, nếu không có bằng chứng combo thì giữ giá trị nhỏ hơn;
                        # tránh lấy max làm 0/1 bị biến thành 50/100 do OCR dính nét.
                        keep_n = min(on, n)
                        keep_tag = tag if n <= on else otag
                    grouped_tagged[-1] = ((od + d) / 2.0, float(keep_n), keep_tag)
            grouped: List[Tuple[float, float]] = [(d, n) for d, n, _tag in grouped_tagged]

            # V1.0.17: sửa spike đơn lẻ do OCR dính nét, ví dụ N/30 = 2 bị đọc thành 24
            # ở cuối trang. Chỉ sửa khi nó đứng cạnh các giá trị rất nhỏ để không làm hỏng N=24 thật.
            if len(grouped) >= 2:
                fixed_grouped: List[Tuple[float, float]] = []
                for i_pt, (d_pt, n_pt) in enumerate(grouped):
                    n_int = int(round(float(n_pt)))
                    neigh: List[float] = []
                    if i_pt > 0 and abs(float(d_pt) - float(grouped[i_pt - 1][0])) <= 4.0:
                        neigh.append(float(grouped[i_pt - 1][1]))
                    if i_pt + 1 < len(grouped) and abs(float(grouped[i_pt + 1][0]) - float(d_pt)) <= 4.0:
                        neigh.append(float(grouped[i_pt + 1][1]))
                    if 20 < n_int < 80 and neigh and max(neigh) <= 10 and (n_int % 10) in (4, 5):
                        lead = n_int // 10
                        if 0 < lead <= 10 and lead <= max(neigh) + 3:
                            fixed_grouped.append((d_pt, float(lead)))
                            continue
                    fixed_grouped.append((d_pt, n_pt))
                grouped = fixed_grouped

            try:
                sp = float(getattr(self, "_borehole_spt_spacing_m", 2.0) or 2.0)
            except Exception:
                sp = 2.0
            if sp <= 0.10:
                sp = 2.0
            mode = str(getattr(self, "_borehole_spt_mode", "avg") or "avg")
            # Ảnh scan nhỏ/mờ: OCR cột N/30 có thể chỉ đọc được vài số.
            # Nếu số điểm đọc được ít bất thường so với chiều sâu/khoảng cách SPT, dùng đồ thị SPT
            # và rút gọn theo đúng khoảng cách người dùng nhập.
            expected = max(3, int(max(float(total_depth), 0.0) / max(sp, 0.1) * (0.35 if mode == "points" else 0.25)))
            if len(grouped) < expected and str(os.environ.get("TS_CAP_ENABLE_SPT_GRAPH_FALLBACK", "")).strip().lower() in {"1", "true", "yes", "on"}:
                graph_pts = self._extract_spt_points_from_graph_curve(img_obj, xlines, x_desc1, data_top, data_bottom, total_depth)
                graph_sampled = self._sample_spt_graph_points_by_spacing(graph_pts, total_depth, sp)
                if len(graph_sampled) > len(grouped):
                    if mode == "points":
                        # Đồ thị chỉ đủ tốt để định vị các điểm SPT; không tự lấy N từ đồ thị
                        # vì ảnh scan nhỏ/màu xanh có thể làm sai lớn giá trị SPT.
                        return self._merge_spt_ocr_values_to_graph_depths(grouped, graph_sampled, sp)
                    return graph_sampled
                if len(grouped) < 3 and graph_pts:
                    if mode == "points":
                        return self._merge_spt_ocr_values_to_graph_depths(grouped, graph_sampled or graph_pts, sp)
                    return graph_sampled or graph_pts
            # V1.0.51: ở chế độ lấy đủ điểm SPT, không tự chèn điểm theo khoảng cách
            # khi đã có đủ điểm OCR từ bảng; việc chèn làm sinh điểm giả 0m/2m hoặc trùng hàng.
            try:
                if mode != "points" or len(grouped) < expected:
                    grouped = self._densify_spt_points_by_spacing(grouped, total_depth, sp)
            except Exception:
                pass
            try:
                # QA-OCR STRICT: mặc định không dùng đồ thị để sửa SPT; chỉ bật khi debug.
                if str(os.environ.get("TS_CAP_ENABLE_SPT_GRAPH_FALLBACK", "")).strip().lower() in {"1", "true", "yes", "on"}:
                    grouped = self._refine_spt_points_with_graph_check(grouped, img_obj, xlines, x_desc1, data_top, data_bottom, total_depth, sp)
            except Exception:
                pass
            try:
                min_valid_depth = max(0.75, float(sp) * 0.45)
                # Không nhận các hàng giả sinh từ header N1/N2/N3 hoặc đường kẻ ngay miệng bảng.
                grouped = [(d, n) for d, n in grouped if float(d) >= min_valid_depth]
            except Exception:
                pass
            return grouped
        except Exception:
            graph_pts = self._extract_spt_points_from_graph_curve(img_obj, xlines, x_desc1, data_top, data_bottom, total_depth)
            try:
                sp = float(getattr(self, "_borehole_spt_spacing_m", 2.0) or 2.0)
            except Exception:
                sp = 2.0
            graph_sampled = self._sample_spt_graph_points_by_spacing(graph_pts, total_depth, sp) or graph_pts or []
            if str(getattr(self, "_borehole_spt_mode", "avg") or "avg") == "points":
                return [(float(d), max(0.0, min(100.0, float(_n)))) for d, _n in graph_sampled]
            return graph_sampled


    def _image_line_projection_groups(self, img_obj: Any) -> Tuple[List[int], List[int]]:
        """Nhận dạng các đường kẻ dài của bảng lỗ khoan.

        V1.0.20: cache theo id trang ảnh vì lite parser và detailed parser gọi lại cùng
        một trang; kết quả lưới không đổi trong một lần import.

        V1.0.10: không trả riêng kết quả OpenCV nếu OpenCV chỉ bắt được các đường
        rất dài. Với log địa chất, nhiều ranh giới lớp mỏng chỉ kẻ ở các cột trái
        và không chạy hết qua cột mô tả/SPT. Vì vậy hàm này hợp nhất:
        - OpenCV morphology: bắt lưới chính, ít nhiễu.
        - Projection pixel: bắt thêm các ranh giới ngắn của lớp mỏng.
        """
        try:
            cache_key = ("line_groups", id(img_obj), getattr(img_obj, "size", None))
            cached = self._borehole_cache_get("_borehole_line_group_cache", cache_key)
            if cached is not None:
                return list(cached[0]), list(cached[1])
        except Exception:
            cache_key = None

        def _merge_centers(vals: List[int], tol: int = 3) -> List[int]:
            vals = sorted(int(v) for v in vals if v is not None)
            if not vals:
                return []
            clusters: List[List[int]] = []
            for v in vals:
                if not clusters or v - clusters[-1][-1] > tol:
                    clusters.append([v])
                else:
                    clusters[-1].append(v)
            return [int(round(sum(cl) / len(cl))) for cl in clusters]

        xs_cv: List[int] = []
        ys_cv: List[int] = []
        try:
            xs_cv, ys_cv = self._opencv_table_line_centers(img_obj)
        except Exception:
            xs_cv, ys_cv = [], []

        xs_proj: List[int] = []
        ys_proj: List[int] = []
        try:
            from PIL import ImageOps
            im = ImageOps.grayscale(img_obj)
            if self._pil_mean_gray(im) < 128.0:
                im = ImageOps.invert(im)
            w, h = im.size
            pix = im.load()
            thr = 105
            col_counts = [0] * w
            row_counts = [0] * h
            for y in range(h):
                cnt = 0
                for x in range(w):
                    if pix[x, y] < thr:
                        cnt += 1
                        col_counts[x] += 1
                row_counts[y] = cnt
            min_v = max(45, int(h * 0.35))
            # Hạ ngưỡng đường ngang một chút để giữ các ranh giới lớp mỏng chỉ kẻ
            # ở cụm cột trái, nhưng vẫn đủ lớn để loại phần chữ thông thường.
            min_h = max(55, int(w * 0.23))
            xs_raw = [i for i, c in enumerate(col_counts) if c >= min_v]
            ys_raw = [i for i, c in enumerate(row_counts) if c >= min_h]

            def group_centers(vals: List[int]) -> List[int]:
                if not vals:
                    return []
                out: List[int] = []
                start = prev = vals[0]
                for v in vals[1:]:
                    if v > prev + 2:
                        out.append((start + prev) // 2)
                        start = v
                    prev = v
                out.append((start + prev) // 2)
                return out
            xs_proj = group_centers(xs_raw)
            ys_proj = group_centers(ys_raw)
        except Exception:
            xs_proj, ys_proj = [], []

        xs = _merge_centers(xs_cv + xs_proj, tol=3)
        ys = _merge_centers(ys_cv + ys_proj, tol=3)
        return self._borehole_cache_set("_borehole_line_group_cache", cache_key, (xs, ys), max_items=80)

    def _ocr_numeric_cell(self, img_obj: Any, *, psm: int = 7) -> str:
        crop = self._preprocess_crop_for_ocr(img_obj, scale=5, numeric=True)
        return self._ocr_pil_crop(crop, lang="eng", psm=psm, whitelist="0123456789.,-+>/<", timeout=8).strip()

    def _ocr_text_cell(self, img_obj: Any, *, psm: int = 6, scale: int = 3) -> str:
        crop = self._preprocess_crop_for_ocr(img_obj, scale=scale, numeric=False)
        txt = self._ocr_pil_crop(crop, lang="vie+eng", psm=psm, whitelist="", timeout=15)
        return re.sub(r"[ \t]+", " ", str(txt or "")).strip()

    def _ocr_numeric_candidates(self, txt: str) -> List[float]:
        out: List[float] = []
        # Cho phép số dạng .50 / -.50 vì Tesseract hay làm mất số 0 đầu.
        # Bản cũ regex chỉ bắt "50" trong ".50", làm 0.50 bị hiểu thành 50.
        pattern = r"[-+]?(?:\d+(?:[\.,]\d+)?|[\.,]\d+)"
        for raw in re.findall(pattern, str(txt or "")):
            t = raw.replace(",", ".")
            if t.startswith("."):
                t = "0" + t
            elif t.startswith("-."):
                t = "-0" + t[1:]
            elif t.startswith("+."):
                t = "+0" + t[1:]
            try:
                if "." not in t and "," not in raw:
                    sign = -1.0 if t.startswith("-") else 1.0
                    digs = re.sub(r"[^0-9]", "", t)
                    if len(digs) == 4:
                        val = sign * (int(digs) / 100.0)
                    elif len(digs) == 3 and digs.startswith("0"):
                        val = sign * (int(digs) / 100.0)
                    elif len(digs) == 3 and int(digs) > 120:
                        val = sign * (int(digs) / 100.0)
                    else:
                        val = sign * float(int(digs)) if digs else 0.0
                else:
                    val = float(t)
                if abs(val) < 2000:
                    out.append(val)
            except Exception:
                continue
        return out


    def _infer_borehole_top_elev_from_layer_markers(
        self,
        img_obj: Any,
        candidate_rows: List[Dict[str, Any]],
        x_elev0: int,
        x_elev1: int,
        x_depth0: int,
        x_depth1: int,
    ) -> Optional[float]:
        """Suy luận cao độ miệng lỗ khoan khi OCR header không đọc được.

        Với một số trang PDF scan, dòng "CAO ĐỘ LỖ KHOAN" bị OCR rỗng, làm code đổi
        độ sâu sang cao độ âm giả. Khi bảng có cả cột Cao độ đáy lớp và Độ sâu đáy lớp,
        có thể lấy: cao độ lỗ khoan = cao độ đáy + độ sâu đáy. Chỉ dùng vài ranh giới
        đầu để tránh các số OCR sâu dưới đáy bị dính dấu hoặc mất dấu âm.
        """
        try:
            w, h = img_obj.size
        except Exception:
            return None
        tops: List[float] = []
        for r in (candidate_rows or [])[:4]:
            try:
                yy = int(round(float(r.get("y1"))))
            except Exception:
                continue
            windows = [
                (max(0, yy - 18), min(h, yy + 2)),
                (max(0, yy - 12), min(h, yy + 8)),
                (max(0, yy - 6), min(h, yy + 14)),
            ]
            elev_vals: List[float] = []
            depth_vals: List[float] = []
            for ya, yb in windows:
                if yb <= ya + 4:
                    continue
                try:
                    txt_e = self._ocr_numeric_cell(img_obj.crop((max(0, int(x_elev0) + 1), ya, min(w, int(x_elev1) - 1), yb)), psm=7)
                    elev_vals.extend([float(v) for v in self._ocr_numeric_candidates(txt_e) if -300.0 <= float(v) <= 300.0])
                except Exception:
                    pass
                try:
                    txt_d = self._ocr_numeric_cell(img_obj.crop((max(0, int(x_depth0) + 1), ya, min(w, int(x_depth1) - 1), yb)), psm=7)
                    depth_vals.extend([abs(float(v)) for v in self._ocr_numeric_candidates(txt_d) if 0.0 <= abs(float(v)) <= 160.0])
                except Exception:
                    pass
            for ev in elev_vals:
                for dep in depth_vals:
                    top = float(ev) + float(dep)
                    # Cao độ lỗ khoan cầu đường thường nằm trong miền này; đủ rộng cho dự án khác.
                    if -100.0 <= top <= 300.0:
                        tops.append(top)
        if not tops:
            return None
        # V1.0.17: lọc miền hợp lý và vote theo 0.1m. Giá trị này chỉ là fallback cục bộ
        # khi header/cột cao độ không đọc được; không dùng để tịnh tiến toàn bộ trang PDF.
        try:
            from collections import Counter
            valid_tops = [float(t) for t in tops if 0.0 < float(t) < 300.0]
            if valid_tops:
                rounded = [round(t, 1) for t in valid_tops]
                best, support = Counter(rounded).most_common(1)[0]
                if support >= 2 or len(valid_tops) == 1:
                    return float(best)
        except Exception:
            pass
        # Fallback cũ: nhóm theo 0.05m và lấy cụm xuất hiện nhiều nhất.
        buckets: Dict[float, List[float]] = {}
        for v in tops:
            key = round(float(v) / 0.05) * 0.05
            buckets.setdefault(key, []).append(float(v))
        best_key, best_vals = max(buckets.items(), key=lambda kv: (len(kv[1]), -abs(kv[0])))
        if len(best_vals) >= 1:
            best_vals = sorted(best_vals)
            return float(best_vals[len(best_vals) // 2])
        return None


    def _read_borehole_depth_markers_from_boundaries(
        self,
        img_obj: Any,
        ys: List[int],
        x_elev0: int,
        x_elev1: int,
        x_depth0: int,
        x_depth1: int,
        top_elev: Optional[float],
    ) -> List[Tuple[int, float, str]]:
        """Đọc các mốc sâu tại đúng đường ranh giới lớp.

        Không đọc cả ô dài vì trong ô có thể chứa nhiều số SPT/mô tả. Ta crop một dải mỏng
        quanh đường ranh giới ở cột "Cao độ đáy lớp" và "Độ sâu đáy lớp", sau đó dùng các
        mốc này để hiệu chỉnh trục sâu bằng RANSAC/least-square. Cách này tránh lỗi kiểu
        đáy 50 m bị kéo thành 70-80 m khi OCR lấy nhầm chiều cao vùng log.
        """
        out: List[Tuple[int, float, str]] = []
        try:
            w, h = img_obj.size
        except Exception:
            return out

        # V1.0.13 speed-up: OCR nguyên cột cao độ/độ sâu một lần bằng TSV rồi gán về
        # đường ranh giới gần nhất. Cách cũ OCR 3 cửa sổ x 2 cột x số lớp, rất chậm.
        try:
            if ys:
                y_min = max(0, int(min(ys)) - 24)
                y_max = min(h, int(max(ys)) + 24)
                y_targets = [int(round(float(y))) for y in ys]
                fast_out: List[Tuple[int, float, str]] = []

                def _fast_col(xa: int, xb: int, src: str, from_elev: bool = False) -> None:
                    if xb <= xa + 4 or y_max <= y_min + 10:
                        return
                    scale = 5
                    crop = img_obj.crop((max(0, int(xa) + 1), y_min, min(w, int(xb) - 1), y_max))
                    proc = self._preprocess_crop_for_ocr(crop, scale=scale, numeric=True)
                    data = self._ocr_pil_crop_data(proc, lang="eng", psm=6, whitelist="0123456789.,-+", timeout=10)
                    for rec in data:
                        vals = self._ocr_numeric_candidates(str(rec.get("text", "")))
                        if not vals:
                            continue
                        cy = y_min + (float(rec.get("top", 0) or 0) + 0.5 * float(rec.get("height", 0) or 0)) / float(scale)
                        yy = min(y_targets, key=lambda y0: abs(float(y0) - cy))
                        if abs(float(yy) - cy) > 20:
                            continue
                        for v in vals:
                            if from_elev:
                                if top_elev is None:
                                    continue
                                dep = float(top_elev) - float(v)
                            else:
                                dep = abs(float(v))
                            if 0.0 <= dep <= 160.0:
                                fast_out.append((int(yy), float(dep), src))

                _fast_col(x_depth0, x_depth1, "depth_col_fast", False)
                if top_elev is not None:
                    _fast_col(x_elev0, x_elev1, "elev_col_fast", True)
                # Cần tối thiểu 3 mốc ở 3 ranh giới khác nhau để fit trục sâu ổn định.
                if len(set(y for y, _d, _src in fast_out)) >= 3:
                    return fast_out
        except Exception:
            pass

        def uniq(vals: List[float], tol: float = 0.03) -> List[float]:
            res: List[float] = []
            for v in sorted(vals):
                if not res or abs(v - res[-1]) > tol:
                    res.append(v)
            return res

        for yy0 in ys or []:
            try:
                yy = int(round(float(yy0)))
            except Exception:
                continue
            # Ba cửa sổ hơi lệch nhau để không bị đường kẻ cắt ngang chữ/số.
            windows = [
                (max(0, yy - 12), min(h, yy + 8)),
                (max(0, yy - 18), min(h, yy + 2)),
                (max(0, yy - 6), min(h, yy + 14)),
            ]
            depth_vals: List[float] = []
            elev_vals: List[float] = []
            for ya, yb in windows:
                if yb <= ya + 3:
                    continue
                try:
                    txt_d = self._ocr_numeric_cell(img_obj.crop((max(0, int(x_depth0) + 1), ya, min(w, int(x_depth1) - 1), yb)), psm=7)
                    for v in self._ocr_numeric_candidates(txt_d):
                        av = abs(float(v))
                        if 0.0 <= av <= 160.0:
                            depth_vals.append(av)
                except Exception:
                    pass
                if top_elev is not None:
                    try:
                        txt_e = self._ocr_numeric_cell(img_obj.crop((max(0, int(x_elev0) + 1), ya, min(w, int(x_elev1) - 1), yb)), psm=7)
                        for v in self._ocr_numeric_candidates(txt_e):
                            dep = float(top_elev) - float(v)
                            if 0.0 <= dep <= 160.0:
                                elev_vals.append(dep)
                    except Exception:
                        pass

            for d in uniq(depth_vals):
                out.append((yy, float(d), "depth_col"))
            for d in uniq(elev_vals):
                out.append((yy, float(d), "elev_col"))
        return out

    def _fit_borehole_depth_axis(self, markers: List[Tuple[int, float, str]]) -> Optional[Tuple[float, float, int]]:
        """Fit tuyến tính depth = a*y + b từ các mốc sâu, loại outlier OCR.

        Trả (a, b, số inlier). Chỉ dùng khi có ít nhất hai mốc tin cậy và trục sâu tăng theo y.
        """
        pts0: List[Tuple[float, float]] = []
        for yy, dd, _src in markers or []:
            try:
                y = float(yy)
                d = float(dd)
            except Exception:
                continue
            if 0.0 <= d <= 160.0:
                pts0.append((y, d))
        # Gộp các mốc gần trùng nhau nhưng giữ mọi ứng viên để RANSAC tự loại outlier.
        if len(pts0) < 2:
            return None
        max_depth = max(d for _y, d in pts0)
        tol = max(0.45, 0.018 * max(max_depth, 1.0))
        best_inliers: List[Tuple[float, float]] = []
        best_score = -1.0
        n = len(pts0)
        for i in range(n):
            y1, d1 = pts0[i]
            for j in range(i + 1, n):
                y2, d2 = pts0[j]
                if abs(y2 - y1) < 12:
                    continue
                a = (d2 - d1) / (y2 - y1)
                if not (0.002 <= a <= 0.50):
                    continue
                b = d1 - a * y1
                inliers = [(y, d) for y, d in pts0 if abs((a * y + b) - d) <= tol]
                # Ưu tiên nhiều inlier, sau đó ưu tiên phủ được chiều sâu lớn.
                span = (max((d for _y, d in inliers), default=0.0) - min((d for _y, d in inliers), default=0.0))
                score = len(inliers) * 1000.0 + span
                if score > best_score:
                    best_score = score
                    best_inliers = inliers
        if len(best_inliers) < 2:
            return None
        sy = sum(y for y, _d in best_inliers)
        sd = sum(d for _y, d in best_inliers)
        syy = sum(y * y for y, _d in best_inliers)
        syd = sum(y * d for y, d in best_inliers)
        m = float(len(best_inliers))
        den = m * syy - sy * sy
        if abs(den) < 1e-9:
            return None
        a = (m * syd - sy * sd) / den
        b = (sd - a * sy) / m
        if not (0.002 <= a <= 0.50):
            return None
        return float(a), float(b), len(best_inliers)

    def _depth_from_axis(self, axis: Optional[Tuple[float, float, int]], y: float) -> Optional[float]:
        if not axis:
            return None
        try:
            a, b, _n = axis
            d = a * float(y) + b
            if 0.0 <= d <= 180.0:
                return float(d)
        except Exception:
            pass
        return None

    def _lookup_depth_marker_near(
        self,
        markers: List[Tuple[int, float, str]],
        y: float,
        axis: Optional[Tuple[float, float, int]] = None,
        tol_y: int = 4,
    ) -> Optional[float]:
        """Lấy mốc OCR gần đường ranh giới, ưu tiên giá trị khớp trục sâu robust."""
        cands: List[float] = []
        yy = float(y)
        for y0, d, _src in markers or []:
            if abs(float(y0) - yy) <= tol_y and 0.0 <= float(d) <= 160.0:
                cands.append(float(d))
        if not cands:
            return None
        pred = self._depth_from_axis(axis, yy)
        if pred is not None:
            cands.sort(key=lambda d: abs(d - pred))
            # Chỉ nhận nếu mốc OCR không quá lệch trục fit; tránh lấy nhầm số 17.40 ở cột bề dày.
            if abs(cands[0] - pred) <= max(0.55, 0.025 * max(pred, 1.0)):
                return cands[0]
            return None
        cands.sort()
        return cands[len(cands) // 2]


    def _read_borehole_elev_markers_from_boundaries(
        self,
        img_obj: Any,
        ys: List[int],
        x_elev0: int,
        x_elev1: int,
    ) -> List[Tuple[int, float, str]]:
        """Đọc trực tiếp cao độ đáy lớp tại các đường ranh giới.

        Nguyên tắc V1.0.16: cao độ đáy là dữ liệu gốc của bảng, nên phải ưu tiên đọc trực tiếp
        theo từng ranh giới. Không suy ra toàn bộ cao độ bằng cách lấy cao độ miệng lỗ khoan trừ
        độ sâu nếu trong ô cao độ đã đọc được số hợp lý. Cách này tránh lỗi chỉ vì OCR header đọc
        sai một chữ số mà toàn bộ trang bị tịnh tiến cao độ.
        """
        out: List[Tuple[int, float, str]] = []
        try:
            w, h = img_obj.size
        except Exception:
            return out
        y_targets = [int(round(float(y))) for y in (ys or [])]
        if not y_targets:
            return out

        def _add_near(cy_img: float, vals: List[float], src: str, tol_y: float = 22.0) -> None:
            if not vals:
                return
            yy = min(y_targets, key=lambda y0: abs(float(y0) - float(cy_img)))
            if abs(float(yy) - float(cy_img)) > tol_y:
                return
            for v in vals:
                fv = float(v)
                if -300.0 <= fv <= 300.0:
                    out.append((int(yy), fv, src))

        # Đọc nhanh cả cột bằng TSV, rồi gán số về ranh giới gần nhất.
        try:
            y_min = max(0, min(y_targets) - 26)
            y_max = min(h, max(y_targets) + 26)
            if y_max > y_min + 10 and int(x_elev1) > int(x_elev0) + 4:
                scale = 6
                crop = img_obj.crop((max(0, int(x_elev0) + 1), y_min, min(w, int(x_elev1) - 1), y_max))
                proc = self._preprocess_crop_for_ocr(crop, scale=scale, numeric=True)
                data = self._ocr_pil_crop_data(proc, lang="eng", psm=6, whitelist="0123456789.,-+", timeout=10)
                for rec in data:
                    vals = [float(v) for v in self._ocr_numeric_candidates(str(rec.get("text", ""))) if -300.0 <= float(v) <= 300.0]
                    if not vals:
                        continue
                    cy = y_min + (float(rec.get("top", 0) or 0) + 0.5 * float(rec.get("height", 0) or 0)) / float(scale)
                    _add_near(cy, vals, "elev_col_fast")
        except Exception:
            pass

        # Fallback từng ranh giới cho các số bị đường kẻ cắt ngang hoặc không xuất hiện trong TSV.
        def uniq(vals: List[float], tol: float = 0.03) -> List[float]:
            res: List[float] = []
            for v in sorted(vals):
                if not res or abs(v - res[-1]) > tol:
                    res.append(v)
            return res

        for yy0 in y_targets:
            if any(abs(int(yy0) - int(y_exist)) <= 2 for y_exist, _e, _src in out):
                continue
            vals_here: List[float] = []
            windows = [
                (max(0, int(yy0) - 18), min(h, int(yy0) + 2)),
                (max(0, int(yy0) - 12), min(h, int(yy0) + 8)),
                (max(0, int(yy0) - 6), min(h, int(yy0) + 14)),
            ]
            for ya, yb in windows:
                if yb <= ya + 4:
                    continue
                try:
                    txt = self._ocr_numeric_cell(img_obj.crop((max(0, int(x_elev0) + 1), ya, min(w, int(x_elev1) - 1), yb)), psm=7)
                    vals_here.extend([float(v) for v in self._ocr_numeric_candidates(txt) if -300.0 <= float(v) <= 300.0])
                except Exception:
                    pass
            for v in uniq(vals_here):
                out.append((int(yy0), float(v), "elev_col"))
        return out

    def _lookup_elev_marker_near(
        self,
        markers: List[Tuple[int, float, str]],
        y: float,
        predicted_elev: Optional[float] = None,
        tol_y: int = 5,
    ) -> Optional[float]:
        """Lấy cao độ đáy đọc trực tiếp gần ranh giới.

        Nếu có cao độ dự đoán thì chỉ dùng để chọn trong các ứng viên và sửa lỗi mất dấu âm cục bộ.
        Không dùng nó để dịch toàn bộ trang.
        """
        vals: List[float] = []
        yy = float(y)
        for y0, ev, _src in markers or []:
            if abs(float(y0) - yy) <= tol_y and -300.0 <= float(ev) <= 300.0:
                vals.append(float(ev))
        if not vals:
            return None
        pred = None
        try:
            if predicted_elev is not None:
                pred = float(predicted_elev)
        except Exception:
            pred = None
        if pred is not None:
            expanded: List[float] = []
            for v in vals:
                expanded.append(v)
                # OCR hay làm mất dấu âm ở cột cao độ; chỉ xét đảo dấu khi thật sự gần giá trị dự đoán.
                if v > 0:
                    expanded.append(-v)
            expanded.sort(key=lambda v: abs(v - pred))
            best = expanded[0]
            # Cột cao độ rất quan trọng; nếu đã có dự đoán cục bộ mà OCR lệch quá xa
            # thì bỏ giá trị OCR đó, dùng fallback top-depth. Không lấy trung vị tùy tiện.
            if abs(best - pred) <= max(0.25, 0.025 * max(abs(pred), 1.0)):
                return float(best)
            return None
        vals.sort()
        return float(vals[len(vals) // 2])

    def _infer_borehole_top_elev_from_marker_pairs(
        self,
        elev_markers: List[Tuple[int, float, str]],
        depth_markers: List[Tuple[int, float, str]],
    ) -> Optional[float]:
        """Suy cao độ miệng lỗ khoan từ các cặp cùng ranh giới, chỉ khi đủ bằng chứng.

        Phải có ít nhất 2 ranh giới độc lập cùng cho một giá trị top gần nhau. Không còn thay top
        chỉ vì một cặp số tình cờ lệch vài mét.
        """
        tops: List[float] = []
        for ye, ev, _se in elev_markers or []:
            near_depths = [float(d) for yd, d, _sd in depth_markers or [] if abs(float(yd) - float(ye)) <= 5 and 0.0 <= float(d) <= 160.0]
            if not near_depths:
                continue
            # Nếu cao độ đọc bị mất dấu âm, cả ev+d và -ev+d sẽ cạnh tranh theo cụm sau.
            cand_elevs = [float(ev)] + ([-float(ev)] if float(ev) > 0.0 else [])
            for ee in cand_elevs:
                for dd in near_depths:
                    top = ee + dd
                    if -100.0 <= top <= 300.0:
                        tops.append(float(top))
        if len(tops) < 2:
            return None
        buckets: Dict[float, List[float]] = {}
        for v in tops:
            key = round(float(v) / 0.10) * 0.10
            buckets.setdefault(key, []).append(float(v))
        best_key, best_vals = max(buckets.items(), key=lambda kv: (len(kv[1]), -abs(kv[0])))
        if len(best_vals) < 2:
            return None
        if max(best_vals) - min(best_vals) > 0.25:
            return None
        best_vals = sorted(best_vals)
        return float(best_vals[len(best_vals) // 2])


    def _ocr_layer_symbol_cell(self, img_obj: Any, x0: int, x1: int, y0: int, y1: int) -> str:
        """OCR nhanh riêng ô ký hiệu lớp.

        Ô ký hiệu lớp thường chỉ có 1-4 ký tự nằm giữa ô rất cao. OCR cả ô cao dễ trả rỗng
        hoặc nhầm sang đường kẻ. Bản này chỉ OCR 1-2 crop nhỏ để giữ tốc độ.
        """
        try:
            w, h = img_obj.size
            xa = max(0, int(x0) - 2)
            xb = min(w, max(int(x1) - 4, int(x0) + 4))
            height = int(y1) - int(y0)
            yc = int(round((float(y0) + float(y1)) / 2.0))
            crops: List[Tuple[int, int, int]] = []
            if height > 45:
                # Với ô ký hiệu cao, nhãn có thể không đúng giữa ô. Thử crop giữa và crop toàn ô
                # ở scale vừa phải để đọc được CS/4a/5B/6A mà không lấy mô tả.
                crops.append((yc - 42, yc + 42, 4))
                crops.append((int(y0) + 2, int(y1) - 2, 6))
                crops.append((yc - 36, yc + 36, 8))
            elif height <= 24:
                crops.append((int(y0) - 8, int(y1) + 18, 6))
                crops.append((int(y0) - 18, int(y1) + 6, 8))
            else:
                crops.append((int(y0) + 1, int(y1) - 1, 6))
                crops.append((int(y0) + 1, int(y1) - 1, 8))
            # Full ô làm phương án sau cùng, dùng scale vừa để bắt nhãn lệch tâm.
            crops.append((int(y0) + 2, int(y1) - 2, 4))
            candidates: List[str] = []
            whitelist = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz"
            for ya, yb, scale in crops:
                ya = max(0, ya); yb = min(h, yb)
                if yb <= ya + 3 or xb <= xa + 3:
                    continue
                crop = img_obj.crop((xa, ya, xb, yb))
                proc = self._preprocess_crop_for_ocr(crop, scale=scale, numeric=False)
                # PSM 11 đọc tốt hơn các ký hiệu nằm lẻ giữa ô cao như 2A/5B/6A.
                for psm_try in (11, 10, 7):
                    txt = self._ocr_pil_crop(proc, lang="eng", psm=psm_try, whitelist=whitelist, timeout=6)
                    if str(txt or "").strip():
                        cand_txt = str(txt).strip()
                        candidates.append(cand_txt)
                        ns_early = re.sub(r"[^A-Z0-9]", "", _strip_accents(cand_txt).upper())
                        # Nếu đã đọc được ký hiệu chắc chắn từ chính ô ký hiệu, trả ngay để giảm số lần OCR.
                        if ns_early in ("D", "DC", "HANG", "8A", "8B", "CS", "CL", "CH", "ML", "MH", "SP", "SC", "SM", "SW", "GP", "GC", "GM") or re.fullmatch(r"\d+[A-Z]", ns_early or "") or re.search(r"\d+[A-Z]", ns_early or ""):
                            return cand_txt
                        break
            def norm(s: str) -> str:
                return re.sub(r"[^A-Z0-9]", "", _strip_accents(s).upper())
            def score(s: str) -> Tuple[int, int]:
                ns = norm(s)
                if not ns:
                    return (-100, 0)
                if ns in ("D", "DC", "HANG", "8A", "8B", "CS", "CL", "CH", "ML", "MH", "SP", "SC", "SM", "SW", "GP", "GC", "GM"):
                    return (100, -len(ns))
                if ns.startswith("D") and len(ns) <= 3:
                    return (90, -len(ns))
                if "HANG" in ns or ns in ("HAN", "HNG", "HAANG"):
                    return (88, -len(ns))
                if re.fullmatch(r"\d+[A-Z]", ns):
                    return (90, -len(ns))
                if re.fullmatch(r"\d+", ns):
                    return (84, -len(ns))
                # Candidate có thêm nhiễu như 6A1 vẫn tốt hơn chữ rời kiểu SB.
                if re.match(r"\d+[A-Z]", ns):
                    return (82, -len(ns))
                if re.search(r"\d+[A-Z]", ns):
                    return (80, -len(ns))
                if re.match(r"\d+", ns):
                    return (72, -len(ns))
                if re.search(r"\d+", ns):
                    return (68, -len(ns))
                if re.search(r"8[A-Z]?", ns):
                    return (70, -len(ns))
                # SA/SB rất hay là OCR nhầm của 5A/5B hoặc 8A/8B, không coi là ký hiệu chắc chắn.
                if ns in ("SA", "SB"):
                    return (20, -len(ns))
                if len(ns) <= 3:
                    return (30, -len(ns))
                return (0, -len(ns))
            if candidates:
                return max(candidates, key=score)
        except Exception:
            pass
        return ""


    def _clean_layer_name_from_ocr(self, raw: str, desc: str, idx: int, prev_names: List[str]) -> str:
        """Làm sạch ký hiệu lớp từ đúng ô Ký hiệu lớp, không suy đoán theo mô tả dài.

        V1.0.22: mô tả địa tầng trong log có thể tràn sang vùng lớp khác, nên hàm này
        chỉ tin vào crop cột ký hiệu lớp. Nếu OCR không rõ, giữ nhãn L<n> để người dùng
        biết cần hiệu chỉnh, thay vì đoán theo mô tả hoặc theo một chuỗi lỗ khoan cụ thể.
        """
        s0 = str(raw or "").strip()
        s_fix = s0.replace("$", "8")
        s = _strip_accents(s_fix).upper()
        s = re.sub(r"[^A-Z0-9]", "", s)

        if not s:
            return "?"
        # Các lỗi OCR rất phổ biến của ký hiệu lớp, chỉ sửa khi dựa trên chính ký hiệu.
        if s in ("D", "DD", "1D", "ID", "LD") or (s.startswith("D") and len(s) <= 3 and s not in ("DC",)):
            return "D"
        if s in ("DC", "DL"):
            return "DC"
        if "HANG" in s or re.fullmatch(r"H[A-Z]*NG", s or "") or s in ("HAN", "HNG", "HAANG"):
            return "Hang"
        # OCR có thể nhận 8A/8B thành BA/BB hoặc B8A; ưu tiên pattern có số.
        m8 = re.search(r"8\s*([AB])", s)
        if m8:
            return "8" + m8.group(1).upper()
        if s in ("8A", "8B"):
            return s
        # Ký hiệu đất kiểu CS, CL... chỉ lấy nếu nằm gọn trong ô ký hiệu, không lấy từ mô tả.
        m = re.fullmatch(r"(CS|CH|CL|ML|MH|SP|SC|SM|SW|GP|GC|GM|DC)", s)
        if m:
            return m.group(1)
        # Ký hiệu lớp một chữ cái hợp lệ (vd lớp "A"); không nhận nhiễu nhiều chữ.
        if re.fullmatch(r"[AB]", s):
            return s
        # QA V1.0.28: chuỗi dính >=3 chữ số là OCR nối cột (2004, 325330S, 375400425)
        # -> KHÔNG bịa tên lớp, trả "?" để tầng preview đánh dấu cần soát.
        if re.search(r"\d{3,}", s):
            return "?"
        _ALLOWED_SUF = "ABCD"
        m = re.fullmatch(r"(\d{1,2})([A-Z]?)", s)
        if m:
            suf = m.group(2).upper()
            return m.group(1) + (suf if suf in _ALLOWED_SUF else "")
        m = re.search(r"(\d{1,2})([A-Z]?)", s)
        if m:
            suf = m.group(2).upper()
            return m.group(1) + (suf if suf in _ALLOWED_SUF else "")
        return "?"

    def _find_borehole_desc_xrange(self, xlines: List[int], w: int, x4: int) -> Tuple[int, int]:
        """Tìm cột mô tả ĐỊA TẦNG theo các đường dọc.

        Bản V1.0.18 dùng điều kiện x > x4 + 0.10*w nên với nhiều form LKBC/LKKN
        đã nhảy thẳng sang cột SPT, khiến OCR mô tả toàn ký tự hatch/đồ thị và suy loại đất sai.
        Cột mô tả thường là khoảng rộng nhất giữa hai đường dọc sau cụm Trụ cắt/Lỗ khoan và trước SPT.
        """
        xs = sorted(int(x) for x in (xlines or []) if 0 <= int(x) <= int(w))
        # Các đường sau cụm cột bề dày/trụ cắt, giới hạn trước vùng mép phải.
        cands = [x for x in xs if x > int(x4) + max(12, int(0.012 * max(w, 1))) and x < int(w * 0.86)]
        best: Optional[Tuple[int, int, int]] = None
        for a, b in zip(cands, cands[1:]):
            gap = int(b) - int(a)
            if gap >= max(75, int(0.12 * max(w, 1))):
                if best is None or gap > best[0]:
                    best = (gap, int(a), int(b))
        if best is not None:
            return best[1], best[2]
        # Fallback: lấy các đường gần sau x4, bỏ qua đường đôi trụ cắt nếu có.
        if len(cands) >= 3:
            return cands[1], cands[2]
        if len(cands) >= 2:
            return cands[0], cands[1]
        x0 = min(max(int(x4) + int(0.10 * max(w, 1)), 0), max(int(w) - 1, 0))
        x1 = min(int(w) - 1, x0 + int(0.30 * max(w, 1)))
        return x0, x1

    def _is_ambiguous_borehole_layer_name(self, name: Any) -> bool:
        s = re.sub(r"[^a-z0-9]", "", _strip_accents(str(name or "")).lower())
        if not s:
            return True
        if re.fullmatch(r"l\d+", s):
            return True
        if s in ("i", "ii", "iii", "iv", "v", "vi", "sa", "sb"):
            return True
        return False

    def _borehole_rows_quality(self, rows: List[List[Any]]) -> Dict[str, float]:
        """Chấm điểm sơ bộ kết quả parser lỗ khoan để quyết định dùng/merge.

        Parser lưới nhanh rất tốt để không bỏ ranh giới lớp mỏng, nhưng với ảnh scan thường
        nó hay trả tên L6/L11 và mô tả rỗng. Không được để kết quả này ghi đè parser chi tiết.
        """
        n = len(rows or [])
        if n <= 0:
            return {"n": 0.0, "comments": 0.0, "ambiguous": 0.0, "spt": 0.0, "score": -999.0}
        comments = 0
        ambiguous = 0
        spt = 0
        for r in rows:
            name = r[1] if len(r) > 1 else ""
            if self._is_ambiguous_borehole_layer_name(name):
                ambiguous += 1
            if len(r) > 13 and str(r[13] or "").strip():
                comments += 1
            if len(r) > 4 and str(r[4] or "").strip():
                spt += 1
        score = n * 2.0 + comments * 1.5 + spt * 0.3 - ambiguous * 2.5
        return {"n": float(n), "comments": float(comments), "ambiguous": float(ambiguous), "spt": float(spt), "score": float(score)}

    def _borehole_row_bottom_elev(self, row: Any) -> Optional[float]:
        """Lấy cao độ đáy lớp từ dòng OCR; trả None nếu không đọc được."""
        try:
            r = list((row + [""] * 14)[:14]) if isinstance(row, list) else list(row)
            return float(str(r[2]).strip().replace(",", "."))
        except Exception:
            return None

    def _borehole_row_name_norm(self, row_or_name: Any) -> str:
        """Chuẩn hóa tên lớp OCR để so sánh chuỗi lớp."""
        try:
            if isinstance(row_or_name, (list, tuple)):
                value = row_or_name[1] if len(row_or_name) > 1 else ""
            else:
                value = row_or_name
            return re.sub(r"[^a-z0-9]", "", _strip_accents(str(value or "")).lower())
        except Exception:
            return ""

    def _sort_borehole_rows_by_elevation(self, rows: List[List[Any]]) -> List[List[Any]]:
        """Sắp xếp các dòng lỗ khoan theo cao độ đáy giảm dần trong từng hạng mục.

        Lỗi thường gặp ở V1.0.19/V1.0.20 là parser chi tiết bắt được lớp phủ D rất mỏng
        nhưng do merge theo parser lite, dòng D bị append xuống cuối. Khi đó pass repair đã
        vào vùng 8A/Hang và sửa nhầm D thành 8A/Hang, tạo dòng đáy có cao độ lớn hơn lớp đầu.
        Hàm này chỉ sắp xếp theo cột CĐ đáy, không sửa số liệu.
        """
        if not rows:
            return []
        groups: Dict[str, List[List[Any]]] = {}
        order: List[str] = []
        for r0 in rows:
            r = list((r0 + [""] * 14)[:14]) if isinstance(r0, list) else list(r0)
            while len(r) < 14:
                r.append("")
            key = str(r[0] or "")
            if key not in groups:
                groups[key] = []
                order.append(key)
            groups[key].append(r)
        out: List[List[Any]] = []
        for key in order:
            grp = groups.get(key, [])
            vals = [self._borehole_row_bottom_elev(r) for r in grp]
            valid = [v for v in vals if v is not None]
            if len(valid) >= 2 and (max(valid) - min(valid)) > 0.05:
                with_idx = list(enumerate(grp))
                def _sort_key(kv: Tuple[int, List[Any]]) -> Tuple[int, float, int]:
                    ev = self._borehole_row_bottom_elev(kv[1])
                    return (0 if ev is not None else 1, -float(ev) if ev is not None else 1e18, kv[0])
                with_idx.sort(key=_sort_key)
                out.extend([r for _i, r in with_idx])
            else:
                out.extend(grp)
        return out

    def _borehole_row_semantic_score(self, row: List[Any]) -> float:
        """Chấm điểm dòng OCR khi cần chọn giữa hai dòng gần trùng cao độ."""
        r = list((row + [""] * 14)[:14]) if isinstance(row, list) else list(row)
        name = str(r[1] or "").strip()
        norm = self._borehole_row_name_norm(name)
        desc = str(r[13] or "")
        dlow = _strip_accents(desc).lower()
        score = 0.0
        if not self._is_ambiguous_borehole_layer_name(name):
            score += 3.0
        if desc.strip():
            score += 2.0
        if str(r[4] or "").strip():
            score += 0.5
        rock_words = ["da voi", "da set", "da cat", "phong hoa", "nut ne", "nứt nẻ", "limestone", "rock"]
        hang_words = ["vat chat", "lap nhet", "lấp nhét", "hang", "karst", "chay", "chảy"]
        has_rock = any(k in dlow for k in rock_words)
        has_hang = any(k in dlow for k in hang_words)
        if norm in ("8", "8a", "8b", "3", "4", "6a", "6b", "7") and has_rock:
            score += 2.5
        if norm == "hang" and has_hang:
            score += 2.5
        if norm == "hang" and has_rock and not has_hang:
            score -= 3.0
        if norm in ("d", "1") and any(k in dlow for k in ["dat dap", "dat phu", "nen mat", "mat duong", "thổ nhưỡng", "tho nhuong"]):
            score += 2.0
        if norm == "8b" and any(k in dlow for k in ["phong hoa vua", "hoa vua", "mạnh - vừa", "manh - vua"]):
            score += 1.5
        return score

    def _dedup_borehole_rows_by_elevation(self, rows: List[List[Any]]) -> List[List[Any]]:
        """Bỏ dòng gần trùng cao độ do border/footer hoặc line kép.

        Dùng tolerance theo chiều sâu toàn log, nhưng giới hạn 0.22m để không gộp nhầm
        các lớp karst thật chỉ dày 0.30m.
        """
        rows2 = self._sort_borehole_rows_by_elevation(rows)
        elevs = [self._borehole_row_bottom_elev(r) for r in rows2]
        valid = [float(v) for v in elevs if v is not None]
        elev_range = (max(valid) - min(valid)) if len(valid) >= 2 else 0.0
        tol = max(0.045, min(0.22, 0.0045 * max(elev_range, 1.0)))
        out: List[List[Any]] = []
        for r in rows2:
            ev = self._borehole_row_bottom_elev(r)
            dup_i = None
            if ev is not None:
                for i, old in enumerate(out):
                    oe = self._borehole_row_bottom_elev(old)
                    if oe is not None and abs(float(ev) - float(oe)) <= tol:
                        dup_i = i
                        break
            if dup_i is None:
                out.append(r)
                continue
            old = out[dup_i]
            old_score = self._borehole_row_semantic_score(old)
            new_score = self._borehole_row_semantic_score(r)
            # Nếu một dòng là Hang nhưng mô tả lại là đá vôi/phong hóa, ưu tiên dòng đá.
            old_norm = self._borehole_row_name_norm(old)
            new_norm = self._borehole_row_name_norm(r)
            if old_norm == "hang" and new_norm in ("8", "8a", "8b"):
                new_score += 2.0
            if new_norm == "hang" and old_norm in ("8", "8a", "8b"):
                old_score += 2.0
            if new_score > old_score:
                # Giữ mô tả/SPT tốt nhất khi thay dòng.
                if not str(r[13] or "").strip() and str(old[13] or "").strip():
                    r[13] = old[13]
                if not str(r[4] or "").strip() and str(old[4] or "").strip():
                    r[4] = old[4]
                out[dup_i] = r
            else:
                if not str(old[13] or "").strip() and str(r[13] or "").strip():
                    old[13] = r[13]
                if not str(old[4] or "").strip() and str(r[4] or "").strip():
                    old[4] = r[4]
        return out


    def _borehole_spt_is_100_or_more(self, value: Any) -> bool:
        """Nhận diện SPT >=100 thật sự. Không coi >50 là 100/refusal."""
        txt = str(value or "").strip().replace(" ", "")
        if not txt:
            return False
        txt_up = txt.upper().replace("O", "0").replace(",", ".")
        if txt_up.startswith(">"):
            nums = re.findall(r"\d+(?:\.\d+)?", txt_up)
            if not nums:
                return False
            try:
                return float(nums[0]) >= 100.0
            except Exception:
                return False
        try:
            return float(txt_up) >= 99.5
        except Exception:
            return False

    def _borehole_spt_missing_or_zero_after_refusal(self, value: Any) -> bool:
        """Chỉ blank mới là thiếu. SPT=0 là số liệu thật, không được xem là missing."""
        return str(value or "").strip() == ""

    def _fill_borehole_spt_after_refusal(self, rows: List[List[Any]]) -> List[List[Any]]:
        """Không tự điền 100 sau refusal.

        Bản cũ tự điền 100 cho các ô trống/0 phía dưới sau khi gặp >50/>100. Điều này
        làm SPT=0 hoặc ô OCR không chắc bị biến thành 100. Để an toàn thiết kế, giữ nguyên
        dữ liệu đọc được; thiếu thì để trống để người dùng soát.
        """
        return [list((r + [""] * 14)[:14]) if isinstance(r, list) else list(r) for r in (rows or [])]

    def _cleanup_borehole_rows_after_repair(self, rows: List[List[Any]]) -> List[List[Any]]:
        """Dọn chuỗi sau OCR theo nguyên tắc bảo toàn số liệu lõi.

        Chỉ sort/dedup và gộp dòng cùng tên liền kề. Không dùng mô tả để đổi tên lớp,
        vì mô tả dài có thể tràn qua ô của lớp khác.
        """
        rows2 = self._dedup_borehole_rows_by_elevation(rows)
        out: List[List[Any]] = []
        preserve_spt_segments = str(getattr(self, "_borehole_spt_mode", "avg") or "avg") == "points"
        for r in rows2:
            if out and not preserve_spt_segments:
                prev = out[-1]
                pn = self._borehole_row_name_norm(prev)
                rn = self._borehole_row_name_norm(r)
                pe = self._borehole_row_bottom_elev(prev)
                re_ = self._borehole_row_bottom_elev(r)
                if pn and rn and pn == rn and pe is not None and re_ is not None and float(re_) < float(pe) - 0.03:
                    merged = list(r)
                    # Ưu tiên giữ SPT nếu dòng sâu hơn chưa có.
                    if not str(merged[4] or "").strip() and str(prev[4] or "").strip():
                        merged[4] = prev[4]
                    # Không merge mô tả dài; để trống cho an toàn.
                    merged[13] = ""
                    out[-1] = merged
                    continue
            r[13] = ""
            out.append(r)
        # QA-OCR v5: không xóa cứng dòng cuối chỉ vì tên mơ hồ.
        # Trong các log đá/R3, lớp cuối thường OCR tên kém và không có SPT nhưng là lớp quyết định
        # chiều sâu mũi cọc. Nếu nghi ngờ thì giữ để người dùng soát trong preview.
        # Chỉ xóa khi dòng cuối gần trùng cao độ với dòng trước (khả năng là đường kẻ/footer kép).
        while len(out) >= 3:
            last = out[-1]
            if self._is_ambiguous_borehole_layer_name(last[1]) and not str(last[4] or "").strip():
                le = self._borehole_row_bottom_elev(last)
                pe = self._borehole_row_bottom_elev(out[-2])
                if le is not None and pe is not None and abs(float(le) - float(pe)) <= 0.30:
                    out.pop()
                    continue
            break
        return self._fill_borehole_spt_after_refusal(out)

    def _repair_borehole_row_sequence(self, rows: List[List[Any]], top_elev: Optional[float] = None) -> List[List[Any]]:
        """Hậu xử lý bảo thủ cho import lỗ khoan.

        V1.0.22 không còn sửa tên lớp theo mô tả hoặc theo một chuỗi lỗ khoan cụ thể.
        Thuật toán chỉ bảo toàn ba cột lõi: tên lớp đọc từ cột ký hiệu, cao độ đáy đọc từ
        cột số, và SPT đọc từ cột N/30. Loại đất suy tạm từ tên lớp, người dùng có thể sửa.

        QA-OCR v2 (theo góp ý test thực tế):
        - Tên lớp chỉ nhận mẫu hợp lệ (1, 1B, 3A, 8a, 10, D, CS, Hang, mã USCS...). Chuỗi dính
          nhiều chữ số kiểu 325330/375400425 là các số độ sâu tràn cột -> thay bằng L<n>,
          KHÔNG dùng để suy loại đất.
        - Nếu biết cao độ miệng lỗ khoan (top_elev): loại dòng có CĐ đáy cao hơn miệng
          hoặc sâu bất thường (>160 m) — đây là số đọc nhầm từ cột khác.
        """
        rows = self._sort_borehole_rows_by_elevation(rows or [])
        out: List[List[Any]] = []
        for i, r0 in enumerate(rows or []):
            r = list((r0 + [""] * 14)[:14]) if isinstance(r0, list) else list(r0)
            while len(r) < 14:
                r.append("")
            # QA-OCR v2: lọc CĐ đáy vô lý khi biết cao độ miệng lỗ khoan.
            if top_elev is not None:
                try:
                    elev_val = float(str(r[2]).replace(",", "."))
                    if elev_val > float(top_elev) + 0.30 or elev_val < float(top_elev) - 160.0:
                        continue
                except Exception:
                    pass
            name = str(r[1] or "").strip()
            # Chuẩn hóa hoa/thường từ chính ký hiệu lớp, không dựa vào mô tả.
            nn = re.sub(r"[^a-z0-9]", "", _strip_accents(name).lower())
            digit_count = sum(1 for ch in nn if ch.isdigit())
            if nn == "hang":
                name = "Hang"
            elif nn in ("dc", "cs", "cl", "ch", "ml", "mh", "sp", "sc", "sm", "sw", "gp", "gc", "gm", "tk", "d"):
                name = nn.upper()
            elif re.fullmatch(r"\d{1,2}[a-d]\d?", nn or ""):
                # 1B, 3A, 8a, 5A1...: viết hoa phần chữ, giữ nguyên phần số.
                name = "".join(ch.upper() if ch.isalpha() else ch for ch in nn)
            elif re.fullmatch(r"\d{1,2}", nn or ""):
                name = nn
            elif re.fullmatch(r"\d{1,2}[a-z]", nn or ""):
                # QA-OCR v4: hậu tố ngoài a-d (vd 4F) thường là OCR nhận nhầm; giữ phần số.
                name = nn[:-1]
            elif digit_count >= 3 and digit_count >= 0.6 * max(len(nn), 1):
                # QA-OCR v2: chuỗi dính số (vd 325330) không phải tên lớp.
                name = "?"
            elif not name:
                name = "?"
            r[1] = name
            # Loại đất chỉ suy tạm từ ký hiệu lớp; không dùng mô tả lệch cột.
            try:
                r[3] = str(self._infer_soil_type_from_borehole_text(r[1], ""))
            except Exception:
                pass
            r[13] = ""
            out.append(r)
        return self._cleanup_borehole_rows_after_repair(out)

    def _merge_borehole_grid_rows(self, detailed_rows: List[List[Any]], lite_rows: List[List[Any]], top_elev: Optional[float] = None) -> List[List[Any]]:
        """Gộp parser chi tiết và parser lưới nhanh theo cao độ đáy.

        - Parser chi tiết: mô tả/tên lớp tốt hơn nhưng có thể bỏ lớp mỏng.
        - Parser nhanh: giữ ranh giới lớp tốt hơn nhưng tên/mô tả có thể rỗng.
        Kết quả ưu tiên dòng chi tiết khi cùng cao độ, bổ sung dòng nhanh khi chi tiết bỏ sót.
        """
        if not detailed_rows:
            return self._repair_borehole_row_sequence(list(lite_rows or []), top_elev=top_elev)
        if not lite_rows:
            return self._repair_borehole_row_sequence(list(detailed_rows or []), top_elev=top_elev)
        # Khi nhập đủ điểm SPT, parser nhanh có thể đã tách một lớp thành nhiều đoạn SPT
        # nhưng chưa đủ ranh giới địa tầng, nếu trộn lại sẽ sinh các dòng trung gian sai.
        # Vì vậy ưu tiên parser chi tiết; điểm SPT đã được lấy trực tiếp từ cột N/30 hoặc đồ thị.
        if str(getattr(self, "_borehole_spt_mode", "avg") or "avg") == "points":
            return self._repair_borehole_row_sequence(list(detailed_rows or []), top_elev=top_elev)

        def elev_of(r: List[Any]) -> Optional[float]:
            try:
                return float(str(r[2]).replace(",", "."))
            except Exception:
                return None

        detailed = [list((r + [""] * 14)[:14]) for r in detailed_rows]
        lite = [list((r + [""] * 14)[:14]) for r in lite_rows]
        used = set()
        merged: List[List[Any]] = []
        tol = 0.08
        for lr in lite:
            le = elev_of(lr)
            best_j = None
            best_de = 999.0
            for j, dr in enumerate(detailed):
                if j in used:
                    continue
                de = elev_of(dr)
                if le is None or de is None:
                    continue
                diff = abs(float(le) - float(de))
                if diff < best_de:
                    best_de = diff
                    best_j = j
            if best_j is not None and best_de <= tol:
                dr = detailed[best_j]
                used.add(best_j)
                # Nếu dòng chi tiết vẫn tên mơ hồ mà dòng nhanh có tên rõ hơn thì lấy tên nhanh.
                if self._is_ambiguous_borehole_layer_name(dr[1]) and not self._is_ambiguous_borehole_layer_name(lr[1]):
                    dr[1] = lr[1]
                # Nếu dòng chi tiết thiếu SPT mà dòng nhanh có thì lấy SPT nhanh.
                if not str(dr[4] or "").strip() and str(lr[4] or "").strip():
                    dr[4] = lr[4]
                merged.append(dr)
            else:
                merged.append(lr)
        # Thêm các dòng chi tiết không có trong lite, đúng vị trí tương đối theo cao độ.
        for j, dr in enumerate(detailed):
            if j not in used:
                merged.append(dr)
        # Sort + dedup trước khi repair để lớp phủ đầu rất mỏng không bị append xuống cuối
        # và bị sửa nhầm thành 8A/Hang sau khi đã vào vùng karst.
        dedup = self._dedup_borehole_rows_by_elevation(merged)
        return self._repair_borehole_row_sequence(dedup, top_elev=top_elev)

    def _try_make_geology_rows_from_borehole_image_grid_lite_legacy(self, path: str, item_name: str, ocr_text: str) -> Tuple[List[List[Any]], str]:
        """Pipeline cũ - parser nhanh cho PDF AutoCAD/vector có lưới rõ.

        Khác parser đầy đủ: không OCR mô tả dài từng lớp trước khi xác định ranh giới.
        Dòng lớp được giữ theo ranh giới lưới + cột số Cao độ/Độ sâu; mô tả có thể để trống.
        Mục tiêu là không bỏ lớp khi form có nhiều lớp nhưng OCR mô tả yếu hoặc trang bị xoay 90°.
        """
        rows: List[List[Any]] = []
        try:
            images = self._images_from_borehole_file(path)
        except Exception:
            return [], ""
        multi_page = len(images) > 1
        large_batch = len(images) > 4

        def _merge_ys(vals: List[int], tol: int = 3) -> List[int]:
            vals = sorted(int(v) for v in vals if v is not None)
            if not vals:
                return []
            out: List[int] = []
            st = pv = vals[0]
            for v in vals[1:]:
                if v > pv + tol:
                    out.append((st + pv) // 2)
                    st = v
                pv = v
            out.append((st + pv) // 2)
            return out

        for page_idx, img in enumerate(images, start=1):
            try:
                page_text = self._borehole_page_text(ocr_text, page_idx)
                if (not large_batch) and not str(page_text or '').strip():
                    try:
                        page_text = self._ocr_borehole_header_text_from_image(img, page_idx)
                    except Exception:
                        page_text = ""
                top_elev = self._extract_borehole_top_elev(page_text)
                page_item_name = item_name if not multi_page else f"{item_name}_p{page_idx:02d}"
                w, h = img.size
                xlines, ylines0 = self._image_line_projection_groups(img)
                xlines = sorted([x for x in xlines if 0 <= x <= w])
                left_lines = [x for x in xlines if x < w * 0.42]
                if len(left_lines) < 5:
                    continue
                scale_col_layout = bool(len(left_lines) >= 6 and (w > h * 1.08 or left_lines[0] > w * 0.22))
                # QA-OCR STRICT: một số ảnh/screenshot có đường biên trang rất sát mép trái
                # (x≈0..8 px) đứng trước cột "Lớp" thật. Bản cũ lấy nhầm đường này làm
                # biên cột ký hiệu lớp, khiến cột Cao độ bị đọc nhầm thành Độ sâu và sinh
                # hàng loạt lớp/SPT sai. Chỉ bỏ qua cột đầu khi nó sát mép VÀ rất hẹp;
                # không áp cho form R3 vì cột đầu của R3 rộng ~40 px là cột thật.
                leading_edge_col_layout = bool(
                    len(left_lines) >= 6
                    and int(left_lines[0]) <= max(10, int(w * 0.015))
                    and int(left_lines[1] - left_lines[0]) <= max(35, int(w * 0.055))
                )
                if scale_col_layout or leading_edge_col_layout:
                    x0, x1, x2, x3, x4 = left_lines[1:6]
                else:
                    x0, x1, x2, x3, x4 = left_lines[:5]
                x_desc0, x_desc1 = self._find_borehole_desc_xrange(xlines, w, x4)

                from PIL import ImageOps
                import numpy as np  # type: ignore
                g = ImageOps.grayscale(img)
                if self._pil_mean_gray(g) < 128.0:
                    g = ImageOps.invert(g)
                arr = np.array(g)
                y_candidates: List[int] = list(ylines0 or [])
                # Dò thêm đường ranh giới từ các vùng sạch. Với form thường, cột trái đủ;
                # với form Cái Cám sau khi xoay, nhiều ranh giới chỉ hiện rõ nếu tính tới vùng trụ cắt/mô tả.
                regions: List[Tuple[int, int, float]] = []
                # Form đứng chuẩn: chỉ dùng cụm cột trái để tránh hatch đá tạo hàng trăm y giả.
                regions.append((max(0, x0), min(w - 1, x4), 0.42))
                # Form có cột tỷ lệ sau khi tự xoay: một số ranh giới lớp chỉ kéo dài qua
                # vùng trụ cắt/mô tả, nên bổ sung vùng rộng hơn nhưng chỉ cho layout này.
                if scale_col_layout:
                    regions.append((max(0, left_lines[0]), min(w - 1, x_desc0), 0.34))
                for xa, xb, ratio in regions:
                    if xb <= xa + 8:
                        continue
                    cnts = (arr[:, xa:xb] < 105).sum(axis=1)
                    thresh = max(50, int((xb - xa) * ratio))
                    min_y = int(h * (0.17 if scale_col_layout else 0.20))
                    idx = np.where((cnts >= thresh) & (np.arange(len(cnts)) > min_y) & (np.arange(len(cnts)) < h - 2))[0].tolist()
                    y_candidates.extend(_merge_ys(idx, tol=2))
                ys = _merge_ys(y_candidates, tol=4)
                # Lọc đường quá gần phía header nhưng không xóa các lớp mặt rất mỏng.
                ys = [y for y in ys if y > h * (0.17 if scale_col_layout else 0.20) and y < h - 1]
                if len(ys) < 3:
                    continue

                elev_markers = self._read_borehole_elev_markers_from_boundaries(img, ys, x1, x2)
                depth_markers = self._read_borehole_depth_markers_from_boundaries(img, ys, x1, x2, x2, x3, None)
                inferred_top = self._infer_borehole_top_elev_from_marker_pairs(elev_markers, depth_markers)
                if top_elev is None and inferred_top is not None:
                    top_elev = float(inferred_top)
                # Đường đầu tiên của vùng dữ liệu xem như depth=0 để fit trục hình học.
                if ys:
                    depth_markers.append((int(ys[0]), 0.0, "data_top_lite"))
                depth_axis = self._fit_borehole_depth_axis(depth_markers)
                if depth_axis is None and len(depth_markers) < 2:
                    continue
                data_top = ys[0]
                # data_bottom tạm lấy theo mốc sâu lớn nhất đọc được, tránh lấy border cuối trang.
                marker_by_y = {int(y): float(d) for y, d, _src in depth_markers if 0.0 <= float(d) <= 180.0}
                data_bottom = max([y for y, d in marker_by_y.items() if d > 0.0] or [ys[-1]])
                total_depth = max([d for _y, d, _src in depth_markers if 0.0 <= float(d) <= 180.0] or [0.0])
                if total_depth <= 0 and depth_axis:
                    dd = self._depth_from_axis(depth_axis, data_bottom)
                    total_depth = float(dd or 0.0)
                spt_points_use = []
                # PDF nhiều trang rất dễ lag nếu OCR cột SPT cho từng trang. Với batch lớn,
                # ưu tiên giữ đủ lớp/cao độ trước; người dùng có thể bổ sung SPT sau ở preview.
                if len(images) <= 4 and total_depth > 0:
                    try:
                        spt_points_use = self._extract_spt_points_from_n30_column(img, xlines, x_desc1, data_top, data_bottom, total_depth)
                    except Exception:
                        spt_points_use = []
                if not spt_points_use:
                    # QA-OCR STRICT: không lấy SPT từ OCR toàn trang. OCR toàn trang dễ nhầm
                    # tỷ lệ 1/250, tọa độ, tiêu đề trục đồ thị hoặc số độ sâu thành SPT.
                    # SPT chỉ được lấy từ cột N/30 hoặc N2+N3 theo tọa độ bảng.
                    # Có thể bật lại để debug bằng TS_CAP_ENABLE_SPT_TEXT_FALLBACK=1.
                    if str(os.environ.get("TS_CAP_ENABLE_SPT_TEXT_FALLBACK", "")).strip().lower() in {"1", "true", "yes", "on"}:
                        spt_points_use = self._parse_spt_points_from_ocr_text(page_text)
                    else:
                        spt_points_use = []

                page_rows: List[List[Any]] = []
                names: List[str] = []
                prev_depth = 0.0
                for idx in range(len(ys) - 1):
                    y0, y1 = float(ys[idx]), float(ys[idx + 1])
                    # Bỏ các khoảng sau đáy hố khoan/khung ngoài.
                    if y0 >= data_bottom + 4:
                        continue
                    marker_depth_direct = self._lookup_depth_marker_near(depth_markers, y1, None, tol_y=5)
                    marker_depth = self._lookup_depth_marker_near(depth_markers, y1, depth_axis, tol_y=5)
                    bottom_depth = marker_depth
                    if bottom_depth is None:
                        bottom_depth = self._depth_from_axis(depth_axis, y1)
                    if bottom_depth is None or bottom_depth <= prev_depth + 0.03:
                        continue
                    if total_depth > 0 and bottom_depth > total_depth + max(0.5, 0.03 * total_depth):
                        continue
                    pred_elev = (float(top_elev) - float(bottom_depth)) if top_elev is not None else None
                    marker_elev = self._lookup_elev_marker_near(elev_markers, y1, pred_elev, tol_y=6)
                    if marker_elev is not None:
                        bottom_elev = float(marker_elev)
                        # OCR hay mất dấu âm ở cột Cao độ đáy. Nếu cao độ đọc được mâu thuẫn lớn
                        # với quan hệ CĐ đáy = CĐ miệng lỗ khoan - độ sâu, dùng giá trị suy từ độ sâu.
                        if pred_elev is not None and (bottom_elev > float(top_elev) + 0.50 or abs(bottom_elev - float(pred_elev)) > max(0.80, 0.06 * max(abs(float(bottom_depth)), 1.0))):
                            bottom_elev = float(pred_elev)
                    else:
                        bottom_elev = (float(top_elev) - float(bottom_depth) if top_elev is not None else -float(bottom_depth))
                    if large_batch:
                        raw_name = ""
                    else:
                        raw_name = self._ocr_layer_symbol_cell(img, max(0, x0 - 3), min(w, x1 + 8), int(y0) + 1, int(y1) - 1)
                    raw_sym_check = re.sub(r"[^A-Z0-9]", "", _strip_accents(str(raw_name or "")).upper())
                    raw_valid = (
                        raw_sym_check in ("D", "DC", "HANG", "K")
                        or bool(re.fullmatch(r"\d+[A-Z]?", raw_sym_check or ""))
                        or bool(re.fullmatch(r"[A-Z]?\d+[A-Z]?", raw_sym_check or ""))
                    )
                    # Nếu ranh giới này không có số depth/elev trực tiếp và ô tên cũng không rõ,
                    # nhiều khả năng đó là leader/mô tả/footer chứ không phải lớp địa chất.
                    if marker_depth_direct is None and not raw_valid:
                        # Batch nhiều trang: không OCR tên lớp; vẫn giữ khoảng nếu chiều dày hình học đủ rõ.
                        if not large_batch or bottom_depth is None or (float(bottom_depth) - float(prev_depth)) < 0.25:
                            continue
                    lname = ("?" if large_batch and not str(raw_name or "").strip() else self._clean_layer_name_from_ocr(raw_name, "", idx, names))
                    # Nếu ô tên hoàn toàn rỗng và tên suy luận chỉ là L*, bỏ qua để không nhập footer.
                    if (not str(raw_name or '').strip()) and re.fullmatch(r"L\d+", str(lname or "")):
                        continue
                    # Với lớp đầu rất mỏng, nếu OCR đọc sai ký hiệu thì vẫn giữ tên đọc được, không ép D cho mọi form.
                    names.append(lname)
                    stype = self._infer_soil_type_from_borehole_text(lname, "")
                    page_rows.extend(self._make_ocr_rows_for_one_borehole_layer(
                        page_item_name, lname, prev_depth, float(bottom_depth), top_elev, stype, spt_points_use,
                        "", bottom_elev_override=float(bottom_elev)
                    ))
                    prev_depth = float(bottom_depth)
                if len(page_rows) >= 2:
                    rows.extend(page_rows)
            except Exception:
                continue
        warn = ""
        if rows:
            warn = "Đã nhận dạng nhanh từ lưới; ưu tiên đủ ranh giới lớp, tên lớp, CĐ đáy và SPT. Mô tả dài để trống nhằm tránh lệch cột."
        return rows, warn



    def _try_make_geology_rows_from_borehole_image_grid_legacy(self, path: str, item_name: str, ocr_text: str) -> Tuple[List[List[Any]], str]:
        """Parser chuyên cho form log lỗ khoan dạng bảng như ảnh BRITEC.

        Điểm khác OCR cũ: không cố đọc cả bảng thành dòng text hoàn chỉnh. Hàm này nhận diện
        đường kẻ ngang/dọc, suy ra chiều dày lớp từ vị trí đường kẻ theo tỷ lệ sâu, rồi OCR từng ô
        mô tả nhỏ. Nhờ vậy các ảnh nền tối, chữ nhỏ, tiêu đề xoay dọc sẽ ổn định hơn.
        """
        lite_rows: List[List[Any]] = []
        lite_warn = ""
        try:
            lite_rows, lite_warn = self._try_make_geology_rows_from_borehole_image_grid_lite_legacy(path, item_name, ocr_text)
        except Exception:
            lite_rows, lite_warn = [], ""
        rows: List[List[Any]] = []
        warn = ""
        try:
            images = self._images_from_borehole_file(path)
        except Exception:
            return self._repair_borehole_row_sequence(lite_rows), lite_warn
        # V1.0.22: nếu parser nhanh đã đủ tin cậy thì dùng ngay để tránh OCR mô tả dài.
        # Nếu chưa đủ tin cậy, chạy parser chi tiết nhưng vẫn chỉ đọc cột ký hiệu/cao độ/SPT.
        try:
            spt_points_mode = str(getattr(self, "_borehole_spt_mode", "avg") or "avg") == "points"
            if (not spt_points_mode) and lite_rows and (len(images) > 4 or len(lite_rows) >= 9 or (len(lite_rows) >= 4 and self._borehole_lite_rows_high_confidence(lite_rows, len(images)))):
                fast_warn = lite_warn or "Đã nhận dạng nhanh theo 3 cột lõi: Tên lớp, CĐ đáy và SPT. Mô tả dài để trống nhằm tránh lệch cột."
                return self._repair_borehole_row_sequence(lite_rows), fast_warn
        except Exception:
            pass
        multi_page = len(images) > 1
        lite_has_spt_any = any(str((r + [""] * 14)[4] or "").strip() for r in (lite_rows or []))
        for page_idx, img in enumerate(images, start=1):
            try:
                page_text = self._borehole_page_text(ocr_text, page_idx)
                if not str(page_text or '').strip():
                    try:
                        page_text = self._ocr_borehole_header_text_from_image(img, page_idx)
                    except Exception:
                        page_text = ""
                top_elev = self._extract_borehole_top_elev(page_text)
                # SPT vẫn lấy từ OCR text trang tương ứng; người dùng duyệt lại trong preview.
                spt_points = self._parse_spt_points_from_ocr_text(page_text)
                page_item_name = item_name if not multi_page else f"{item_name}_p{page_idx:02d}"
                w, h = img.size
                xlines, ylines = self._image_line_projection_groups(img)
                xlines = sorted([x for x in xlines if 0 <= x <= w])
                ylines = sorted([y for y in ylines if 0 <= y <= h])
                left_lines = [x for x in xlines if x < w * 0.42]
                if len(left_lines) < 5 or len(ylines) < 4:
                    continue
                # Form chuẩn: Ký hiệu | Cao độ | Độ sâu | Bề dày | Trụ cắt...
                # Một số PDF AutoCAD khác: Tỷ lệ | Tên lớp | Cao độ | Độ sâu | Bề dày | Trụ cắt...
                # Nếu có cột tỷ lệ ở ngoài cùng trái, dịch bộ cột sang phải một nhịp để không đọc
                # nhầm thang tỷ lệ thành tên lớp và không bỏ các lớp ở trang dạng Cái Cám.
                scale_col_layout = bool(len(left_lines) >= 6 and (w > h * 1.08 or left_lines[0] > w * 0.22))
                # QA-OCR STRICT: một số ảnh/screenshot có đường biên trang rất sát mép trái
                # (x≈0..8 px) đứng trước cột "Lớp" thật. Bản cũ lấy nhầm đường này làm
                # biên cột ký hiệu lớp, khiến cột Cao độ bị đọc nhầm thành Độ sâu và sinh
                # hàng loạt lớp/SPT sai. Chỉ bỏ qua cột đầu khi nó sát mép VÀ rất hẹp;
                # không áp cho form R3 vì cột đầu của R3 rộng ~40 px là cột thật.
                leading_edge_col_layout = bool(
                    len(left_lines) >= 6
                    and int(left_lines[0]) <= max(10, int(w * 0.015))
                    and int(left_lines[1] - left_lines[0]) <= max(35, int(w * 0.055))
                )
                if scale_col_layout or leading_edge_col_layout:
                    x0, x1, x2, x3, x4 = left_lines[1:6]
                else:
                    x0, x1, x2, x3, x4 = left_lines[:5]
                # Tìm cột mô tả theo khoảng rộng nhất sau cụm trụ cắt/lỗ khoan.
                x_desc0, x_desc1 = self._find_borehole_desc_xrange(xlines, w, x4)
                if x_desc1 <= x_desc0 + 20:
                    continue
                # Các đường ngang dữ liệu thường nằm từ khoảng 1/4 chiều cao ảnh trở xuống.
                # Lọc thêm theo số pixel tối qua vùng từ cột ký hiệu tới hết mô tả để loại các
                # đường/ngang giả chỉ thuộc đồ thị SPT hoặc chữ mô tả.
                try:
                    from PIL import ImageOps
                    line_img = ImageOps.grayscale(img)
                    if self._pil_mean_gray(line_img) < 128.0:
                        line_img = ImageOps.invert(line_img)
                    pix_line = line_img.load()
                    # Với form có cột tỷ lệ / trang AutoCAD xoay, ranh giới lớp thường không
                    # đủ dài để lọt qua ngưỡng projection chung. Bổ sung yline bằng cách chỉ đếm
                    # trong cụm cột trái sạch (Tên lớp/Cao độ/Độ sâu/Bề dày), tránh nhiễu hatch.
                    if scale_col_layout:
                        try:
                            extra_ys: List[int] = []
                            left_only_span = max(1, x4 - x0)
                            for yy2 in range(max(0, int(h * 0.18)), min(h - 1, h)):
                                cnt2 = 0
                                for xx2 in range(max(0, x0), min(w - 1, x4) + 1):
                                    if pix_line[xx2, yy2] < 105:
                                        cnt2 += 1
                                if cnt2 >= max(55, int(left_only_span * 0.42)):
                                    extra_ys.append(yy2)
                            if extra_ys:
                                grouped_extra: List[int] = []
                                st2 = pv2 = extra_ys[0]
                                for vv2 in extra_ys[1:]:
                                    if vv2 > pv2 + 2:
                                        grouped_extra.append((st2 + pv2) // 2)
                                        st2 = vv2
                                    pv2 = vv2
                                grouped_extra.append((st2 + pv2) // 2)
                                ylines = sorted(set(list(ylines) + grouped_extra))
                        except Exception:
                            pass
                    span_len = max(1, x_desc1 - x0)
                    y_keep: List[Tuple[int, int]] = []
                    for yy in ylines:
                        min_y_data = h * (0.18 if scale_col_layout else 0.20)
                        if not (yy > min_y_data and yy < h - 2):
                            continue
                        cnt = 0
                        for xx in range(max(0, x0), min(w - 1, x_desc1) + 1):
                            if pix_line[xx, yy] < 105:
                                cnt += 1
                        if cnt >= max(70, int(span_len * 0.38)):
                            y_keep.append((yy, cnt))
                    # Ưu tiên các đường ranh giới có nét trong cụm cột trái
                    # (ký hiệu lớp/cao độ/độ sâu/bề dày). Các đường ngang trong cột mô tả
                    # hoặc đồ thị SPT có thể là đường leader/gạch nền, không phải ranh giới lớp.
                    left_span = max(1, x4 - x0)
                    left_y_keep: List[Tuple[int, int]] = []
                    for yy, cnt_total in y_keep:
                        cnt_left = 0
                        for xx in range(max(0, x0), min(w - 1, x4) + 1):
                            if pix_line[xx, yy] < 105:
                                cnt_left += 1
                        if cnt_left >= max(35, int(left_span * 0.28)):
                            left_y_keep.append((yy, cnt_left))
                    primary = left_y_keep if len(left_y_keep) >= 3 else y_keep
                    ys = []
                    for yy, cnt in primary:
                        # Chỉ gộp các đường rất sát nhau do bề dày nét scan.
                        # Không gộp cụm 2-3 đường cách nhau 6-15 px vì đó thường là lớp phủ mỏng.
                        if not ys or yy - ys[-1] > 4:
                            ys.append(yy)
                        else:
                            ys[-1] = yy
                    ys = sorted(set(ys))
                except Exception:
                    ys = [y for y in ylines if y > h * 0.20 and y < h - 2]
                if len(ys) < 3:
                    continue
                # Xác định đáy lỗ khoan bằng giá trị lớn nhất trong cột độ sâu đáy lớp.
                depth_vals: List[float] = []
                for i in range(len(ys) - 1):
                    y0, y1 = ys[i], ys[i + 1]
                    if y1 - y0 < 5:
                        continue
                    pad = 2
                    txt = self._ocr_numeric_cell(img.crop((max(0, x2 + pad), y0 + pad, min(w, x3 - pad), y1 - pad)), psm=7)
                    for v in self._ocr_numeric_candidates(txt):
                        if 1.0 <= abs(v) <= 150.0:
                            depth_vals.append(abs(v))
                total_depth = max(depth_vals) if depth_vals else 0.0
                # V1.0.11: chưa loại trang ngay nếu OCR cột độ sâu đọc sai. Ta sẽ hiệu chỉnh lại
                # bằng các mốc số ở đường ranh giới lớp sau khi xác định candidate_rows.
                if not (3.0 <= total_depth <= 150.0):
                    total_depth = 0.0
                candidate_rows: List[Dict[str, Any]] = []
                # V1.0.22: không OCR mô tả dài để quyết định lớp vì mô tả có thể tràn sang ô khác.
                # Candidate lớp được lấy từ ranh giới lưới ở cụm cột trái; tên lớp chỉ đọc từ cột Ký hiệu.
                for i in range(len(ys) - 1):
                    y0, y1 = ys[i], ys[i + 1]
                    if y1 - y0 < 5:
                        continue
                    raw_name = self._ocr_layer_symbol_cell(img, max(0, x0 - 3), min(w, x1 + 8), int(y0) + 2, int(y1) - 2)
                    rlow = _strip_accents(raw_name).lower()
                    raw_sym = re.sub(r"[^A-Z0-9]", "", _strip_accents(raw_name).upper())
                    header_like = any(k in rlow for k in [
                        "ky hieu", "cao do", "do sau", "be day", "lo khoan", "mo ta",
                        "dia tang", "thi nghiem", "so bua", "do thi", "spt"
                    ])
                    if header_like:
                        continue
                    name_like = (
                        raw_sym in ("D", "DC", "DD", "HANG", "8", "8A", "8B", "4A", "4B", "CS")
                        or "HANG" in raw_sym
                        or re.fullmatch(r"\d+[A-Z]?", raw_sym or "") is not None
                    )
                    # Giữ lớp đầu rất mỏng nếu nó nằm ngay đầu vùng dữ liệu, nhưng không ép tên thành D.
                    maybe_thin_top = (not candidate_rows) and (y1 - y0 <= max(18, int(h * 0.018))) and y0 > h * 0.18
                    # Với đường ranh giới đã lọc từ cụm cột trái, giữ cả dòng chưa OCR rõ tên nếu chiều cao hợp lý.
                    # Dòng chưa rõ tên sẽ là L<n> để người dùng sửa, thay vì đoán theo mô tả.
                    geometry_like = (y1 - y0) >= max(8, int(h * 0.004))
                    if not (name_like or maybe_thin_top or geometry_like):
                        continue
                    candidate_rows.append({"y0": y0, "y1": y1, "desc": "", "raw_name": raw_name})
                if len(candidate_rows) < 2:
                    continue
                # V1.0.16: không tự thay cao độ toàn trang theo sai khác header.
                # Đọc trực tiếp cột Cao độ đáy lớp và cột Độ sâu đáy lớp tại từng ranh giới;
                # top_elev chỉ là fallback cho những ranh giới không đọc được cao độ.
                data_top = min(r["y0"] for r in candidate_rows)
                data_bottom = max(r["y1"] for r in candidate_rows)
                if data_bottom <= data_top:
                    continue

                boundary_ys = sorted(set([int(round(data_top))] + [int(round(float(r["y1"]))) for r in candidate_rows]))
                elev_markers = self._read_borehole_elev_markers_from_boundaries(img, boundary_ys, x1, x2)
                # Trục sâu robust chỉ lấy từ cột Độ sâu đáy lớp. Không dùng top_elev để biến cột cao độ
                # thành độ sâu, vì nếu OCR header sai thì trục sâu cũng bị kéo sai.
                depth_markers = self._read_borehole_depth_markers_from_boundaries(
                    img, boundary_ys, x1, x2, x2, x3, None
                )
                inferred_top_elev = self._infer_borehole_top_elev_from_marker_pairs(elev_markers, depth_markers)
                top_elev_fallback = top_elev
                if top_elev_fallback is None and inferred_top_elev is not None:
                    top_elev_fallback = float(inferred_top_elev)
                elif top_elev is not None and inferred_top_elev is not None and abs(float(inferred_top_elev) - float(top_elev)) > 0.75:
                    # Chỉ dùng giá trị suy từ bảng cho các ô bị thiếu cao độ; không dịch các ô đã đọc được.
                    top_elev_fallback = float(inferred_top_elev)
                depth_markers.append((int(round(data_top)), 0.0, "data_top"))
                depth_axis = self._fit_borehole_depth_axis(depth_markers)
                axis_bottom = self._depth_from_axis(depth_axis, data_bottom)
                if axis_bottom is not None and 3.0 <= axis_bottom <= 150.0:
                    # Nếu OCR toàn cột cho total_depth lệch lớn, dùng trục robust.
                    if not (3.0 <= total_depth <= 150.0) or abs(float(total_depth) - float(axis_bottom)) > max(0.75, 0.03 * axis_bottom):
                        total_depth = float(axis_bottom)
                if not (3.0 <= total_depth <= 150.0):
                    continue

                # Đọc SPT từ đúng cột N/30 của bảng, lấy tọa độ y để quy đổi theo độ sâu.
                # V1.0.20: nếu parser lite đã có SPT trung bình, không OCR lại cột SPT ở parser chi tiết;
                # bước merge theo cao độ sẽ lấy SPT từ lite_rows. Đây là phần tiết kiệm thời gian lớn
                # mà không làm mất ranh giới lớp/mô tả.
                points_mode = str(getattr(self, "_borehole_spt_mode", "avg") or "avg") == "points"
                if lite_has_spt_any and not points_mode:
                    spt_points_img = []
                    spt_points_use = []
                else:
                    spt_points_img = self._extract_spt_points_from_n30_column(img, xlines, x_desc1, data_top, data_bottom, total_depth)
                    spt_points_use = spt_points_img if spt_points_img else spt_points
                prev_depth = 0.0
                prev_bottom_elev_value: Optional[float] = None
                names: List[str] = []
                page_rows: List[List[Any]] = []
                for idx, r in enumerate(candidate_rows):
                    y0, y1 = float(r["y0"]), float(r["y1"])
                    axis_bottom_depth = self._depth_from_axis(depth_axis, y1)
                    axis_top_depth = self._depth_from_axis(depth_axis, y0)
                    geom_bottom_depth = (y1 - data_top) / max(data_bottom - data_top, 1.0) * total_depth
                    geom_top_depth = (y0 - data_top) / max(data_bottom - data_top, 1.0) * total_depth
                    if axis_bottom_depth is not None and axis_top_depth is not None:
                        geom_bottom_depth = axis_bottom_depth
                        geom_top_depth = axis_top_depth
                    geom_thickness = max(geom_bottom_depth - geom_top_depth, 0.0)
                    # Nếu OCR đọc rõ số độ sâu tại ranh giới lớp và khớp trục robust thì dùng số OCR,
                    # nếu không thì dùng giá trị nội suy hình học đã hiệu chỉnh.
                    marker_bottom_depth = self._lookup_depth_marker_near(depth_markers, y1, depth_axis)
                    bottom_depth = marker_bottom_depth if marker_bottom_depth is not None else geom_bottom_depth
                    # Dự đoán cao độ chỉ để chọn/sửa cục bộ ứng viên OCR ở cột Cao độ đáy.
                    pred_bottom_elev = (float(top_elev_fallback) - float(bottom_depth)) if top_elev_fallback is not None else None
                    marker_bottom_elev = self._lookup_elev_marker_near(elev_markers, y1, pred_bottom_elev)
                    thickness = max(bottom_depth - prev_depth, geom_thickness)
                    # V1.0.13: không OCR lại từng ô bề dày/độ sâu nếu đã có marker/trục sâu.
                    # Đây là nguồn gây chậm lớn ở V1.0.12. Chỉ fallback OCR ô độ sâu khi chưa fit
                    # được trục và chưa có marker gần ranh giới.
                    dep_vals: List[float] = []
                    if depth_axis is None and marker_bottom_depth is None:
                        depth_txt = self._ocr_numeric_cell(img.crop((max(0, x2 + 2), int(y0) + 2, min(w, x3 - 2), int(y1) - 2)), psm=7)
                        dep_vals = [abs(v) for v in self._ocr_numeric_candidates(depth_txt) if prev_depth < abs(v) <= total_depth + 0.5]
                    used_ocr = False
                    if (not used_ocr) and dep_vals:
                        dep = dep_vals[-1]
                        geom_close = abs(dep - geom_bottom_depth) <= max(0.45, 0.025 * max(total_depth, 1.0))
                        if dep > prev_depth + 0.03 and geom_close:
                            bottom_depth = dep
                            thickness = bottom_depth - prev_depth
                            used_ocr = True
                    if marker_bottom_elev is not None:
                        bottom_elev = float(marker_bottom_elev)
                        if pred_bottom_elev is not None and top_elev_fallback is not None and (bottom_elev > float(top_elev_fallback) + 0.50 or abs(bottom_elev - float(pred_bottom_elev)) > max(0.80, 0.06 * max(abs(float(bottom_depth)), 1.0))):
                            bottom_elev = float(pred_bottom_elev)
                        local_top_elev = float(bottom_elev) + float(bottom_depth)
                    else:
                        local_top_elev = top_elev_fallback
                        bottom_elev = (float(local_top_elev) - bottom_depth) if local_top_elev is not None else -bottom_depth
                    # Không lấy mô tả dài vào bảng import; mô tả dễ lệch cột khi lớp mỏng.
                    desc = ""
                    lname = self._clean_layer_name_from_ocr(str(r.get("raw_name") or ""), "", idx, names)
                    # Không ép lớp đầu thành D và không sửa theo chuỗi D->2->3->8A/Hang.
                    # Nếu OCR ký hiệu chưa rõ, giữ L<n> để người dùng hiệu chỉnh trong preview.
                    names.append(lname)
                    points = [n for d, n in spt_points_use if prev_depth < d <= bottom_depth]
                    n_avg = round(sum(points) / len(points)) if points else 0
                    stype = self._infer_soil_type_from_borehole_text(lname, "")
                    # Lớp phủ ký hiệu D thường không có chỉ tiêu SPT thiết kế trong log.
                    if str(lname).strip().upper() == "D":
                        n_avg = 0
                    # Ảnh/log lỗ khoan chỉ có tên lớp, cao độ đáy, phân loại lớp và SPT.
                    # Không tự sinh γ, C, φ, qu, RQD, GSI, mi, D vì trong ảnh không có các chỉ tiêu này.
                    gamma = ""
                    c_mpa = ""
                    qu_mpa = ""
                    rqd = ""
                    gsi = ""
                    mi = ""
                    dist = ""
                    page_rows.extend(self._make_ocr_rows_for_one_borehole_layer(
                        page_item_name, lname, prev_depth, bottom_depth, local_top_elev, stype, spt_points_use, desc, bottom_elev_override=bottom_elev
                    ))
                    prev_depth = bottom_depth
                    try:
                        prev_bottom_elev_value = float(bottom_elev)
                    except Exception:
                        prev_bottom_elev_value = None
                if len(page_rows) >= 2:
                    rows.extend(page_rows)
                    # Không break: PDF/TIFF nhiều trang cần đọc tiếp các trang sau, mỗi trang là một hạng mục riêng.
                    continue
            except Exception:
                continue
        if rows:
            # Gộp thêm các ranh giới lớp mỏng mà parser nhanh bắt được nhưng parser chi tiết bỏ sót.
            if lite_rows:
                rows = self._merge_borehole_grid_rows(rows, lite_rows)
            else:
                rows = self._repair_borehole_row_sequence(rows)
            if not any(str(r[4] or "").strip() for r in rows):
                warn = "Đã nhận được lớp bằng parser lưới ảnh, nhưng chưa nhận được cột N/30 SPT; cột SPT đang để trống. Anh kiểm tra/sửa ở bảng preview trước khi import."
            else:
                warn = "Đã nhận dạng theo 3 cột lõi: Tên lớp, CĐ đáy và SPT. Mô tả dài được bỏ qua để tránh lệch cột; anh kiểm tra/sửa loại đất ở preview nếu cần."
            return rows, warn
        if lite_rows:
            return self._repair_borehole_row_sequence(lite_rows), (lite_warn or "Đã nhận dạng nhanh từ lưới ảnh; anh kiểm tra lại tên lớp, loại đất và SPT ở preview.")
        return rows, warn


    def _try_make_geology_rows_from_borehole_image_grid_lite(self, path: str, item_name: str, ocr_text: str) -> Tuple[List[List[Any]], str]:
        """Parser nhanh cho PDF AutoCAD/vector có lưới rõ.

        Khác parser đầy đủ: không OCR mô tả dài từng lớp trước khi xác định ranh giới.
        Dòng lớp được giữ theo ranh giới lưới + cột số Cao độ/Độ sâu; mô tả có thể để trống.
        Mục tiêu là không bỏ lớp khi form có nhiều lớp nhưng OCR mô tả yếu hoặc trang bị xoay 90°.
        """
        rows: List[List[Any]] = []
        try:
            images = self._images_from_borehole_file(path)
        except Exception:
            return [], ""
        multi_page = len(images) > 1
        large_batch = len(images) > 4

        def _merge_ys(vals: List[int], tol: int = 3) -> List[int]:
            vals = sorted(int(v) for v in vals if v is not None)
            if not vals:
                return []
            out: List[int] = []
            st = pv = vals[0]
            for v in vals[1:]:
                if v > pv + tol:
                    out.append((st + pv) // 2)
                    st = v
                pv = v
            out.append((st + pv) // 2)
            return out

        for page_idx, img in enumerate(images, start=1):
            try:
                page_text = self._borehole_page_text(ocr_text, page_idx)
                if (not large_batch) and not str(page_text or '').strip():
                    try:
                        page_text = self._ocr_borehole_header_text_from_image(img, page_idx)
                    except Exception:
                        page_text = ""
                top_elev = self._extract_borehole_top_elev(page_text)
                page_item_name = item_name if not multi_page else f"{item_name}_p{page_idx:02d}"
                w, h = img.size
                xlines, ylines0 = self._image_line_projection_groups(img)
                xlines = sorted([x for x in xlines if 0 <= x <= w])
                left_lines = [x for x in xlines if x < w * 0.42]
                if len(left_lines) < 5:
                    continue
                scale_col_layout = bool(len(left_lines) >= 6 and (w > h * 1.08 or left_lines[0] > w * 0.22))
                # QA-OCR STRICT: một số ảnh/screenshot có đường biên trang rất sát mép trái
                # (x≈0..8 px) đứng trước cột "Lớp" thật. Bản cũ lấy nhầm đường này làm
                # biên cột ký hiệu lớp, khiến cột Cao độ bị đọc nhầm thành Độ sâu và sinh
                # hàng loạt lớp/SPT sai. Chỉ bỏ qua cột đầu khi nó sát mép VÀ rất hẹp;
                # không áp cho form R3 vì cột đầu của R3 rộng ~40 px là cột thật.
                leading_edge_col_layout = bool(
                    len(left_lines) >= 6
                    and int(left_lines[0]) <= max(10, int(w * 0.015))
                    and int(left_lines[1] - left_lines[0]) <= max(35, int(w * 0.055))
                )
                if scale_col_layout or leading_edge_col_layout:
                    x0, x1, x2, x3, x4 = left_lines[1:6]
                else:
                    x0, x1, x2, x3, x4 = left_lines[:5]
                x_desc0, x_desc1 = self._find_borehole_desc_xrange(xlines, w, x4)

                from PIL import ImageOps
                import numpy as np  # type: ignore
                g = ImageOps.grayscale(img)
                if self._pil_mean_gray(g) < 128.0:
                    g = ImageOps.invert(g)
                arr = np.array(g)
                y_candidates: List[int] = list(ylines0 or [])
                # Dò thêm đường ranh giới từ các vùng sạch. Với form thường, cột trái đủ;
                # với form Cái Cám sau khi xoay, nhiều ranh giới chỉ hiện rõ nếu tính tới vùng trụ cắt/mô tả.
                regions: List[Tuple[int, int, float]] = []
                # Form đứng chuẩn: chỉ dùng cụm cột trái để tránh hatch đá tạo hàng trăm y giả.
                regions.append((max(0, x0), min(w - 1, x4), 0.42))
                # Form có cột tỷ lệ sau khi tự xoay: một số ranh giới lớp chỉ kéo dài qua
                # vùng trụ cắt/mô tả, nên bổ sung vùng rộng hơn nhưng chỉ cho layout này.
                if scale_col_layout:
                    regions.append((max(0, left_lines[0]), min(w - 1, x_desc0), 0.34))
                for xa, xb, ratio in regions:
                    if xb <= xa + 8:
                        continue
                    cnts = (arr[:, xa:xb] < 105).sum(axis=1)
                    thresh = max(50, int((xb - xa) * ratio))
                    min_y = int(h * (0.17 if scale_col_layout else 0.20))
                    idx = np.where((cnts >= thresh) & (np.arange(len(cnts)) > min_y) & (np.arange(len(cnts)) < h - 2))[0].tolist()
                    y_candidates.extend(_merge_ys(idx, tol=2))
                ys = _merge_ys(y_candidates, tol=4)
                # Lọc đường quá gần phía header nhưng không xóa các lớp mặt rất mỏng.
                ys = [y for y in ys if y > h * (0.17 if scale_col_layout else 0.20) and y < h - 1]
                if len(ys) < 3:
                    continue

                elev_markers = self._read_borehole_elev_markers_from_boundaries(img, ys, x1, x2)
                depth_markers = self._read_borehole_depth_markers_from_boundaries(img, ys, x1, x2, x2, x3, None)
                inferred_top = self._infer_borehole_top_elev_from_marker_pairs(elev_markers, depth_markers)
                if top_elev is None and inferred_top is not None:
                    top_elev = float(inferred_top)
                # Đường đầu tiên của vùng dữ liệu xem như depth=0 để fit trục hình học.
                if ys:
                    depth_markers.append((int(ys[0]), 0.0, "data_top_lite"))
                depth_axis = self._fit_borehole_depth_axis(depth_markers)
                if depth_axis is None and len(depth_markers) < 2:
                    continue
                data_top = ys[0]
                # data_bottom tạm lấy theo mốc sâu lớn nhất đọc được, tránh lấy border cuối trang.
                marker_by_y = {int(y): float(d) for y, d, _src in depth_markers if 0.0 <= float(d) <= 180.0}
                data_bottom = max([y for y, d in marker_by_y.items() if d > 0.0] or [ys[-1]])
                total_depth = max([d for _y, d, _src in depth_markers if 0.0 <= float(d) <= 180.0] or [0.0])
                if total_depth <= 0 and depth_axis:
                    dd = self._depth_from_axis(depth_axis, data_bottom)
                    total_depth = float(dd or 0.0)
                spt_points_use = []
                # PDF nhiều trang rất dễ lag nếu OCR cột SPT cho từng trang. Với batch lớn,
                # ưu tiên giữ đủ lớp/cao độ trước; người dùng có thể bổ sung SPT sau ở preview.
                if len(images) <= 4 and total_depth > 0:
                    try:
                        spt_points_use = self._extract_spt_points_from_n30_column(img, xlines, x_desc1, data_top, data_bottom, total_depth)
                    except Exception:
                        spt_points_use = []
                if not spt_points_use:
                    # QA-OCR STRICT: không lấy SPT từ OCR toàn trang. OCR toàn trang dễ nhầm
                    # tỷ lệ 1/250, tọa độ, tiêu đề trục đồ thị hoặc số độ sâu thành SPT.
                    # SPT chỉ được lấy từ cột N/30 hoặc N2+N3 theo tọa độ bảng.
                    # Có thể bật lại để debug bằng TS_CAP_ENABLE_SPT_TEXT_FALLBACK=1.
                    if str(os.environ.get("TS_CAP_ENABLE_SPT_TEXT_FALLBACK", "")).strip().lower() in {"1", "true", "yes", "on"}:
                        spt_points_use = self._parse_spt_points_from_ocr_text(page_text)
                    else:
                        spt_points_use = []

                page_rows: List[List[Any]] = []
                names: List[str] = []
                prev_depth = 0.0
                for idx in range(len(ys) - 1):
                    y0, y1 = float(ys[idx]), float(ys[idx + 1])
                    # Bỏ các khoảng sau đáy hố khoan/khung ngoài.
                    if y0 >= data_bottom + 4:
                        continue
                    marker_depth_direct = self._lookup_depth_marker_near(depth_markers, y1, None, tol_y=5)
                    marker_depth = self._lookup_depth_marker_near(depth_markers, y1, depth_axis, tol_y=5)
                    bottom_depth = marker_depth
                    if bottom_depth is None:
                        bottom_depth = self._depth_from_axis(depth_axis, y1)
                    if bottom_depth is None or bottom_depth <= prev_depth + 0.03:
                        continue
                    if total_depth > 0 and bottom_depth > total_depth + max(0.5, 0.03 * total_depth):
                        continue
                    pred_elev = (float(top_elev) - float(bottom_depth)) if top_elev is not None else None
                    marker_elev = self._lookup_elev_marker_near(elev_markers, y1, pred_elev, tol_y=6)
                    if marker_elev is not None:
                        bottom_elev = float(marker_elev)
                        # OCR hay mất dấu âm ở cột Cao độ đáy. Nếu cao độ đọc được mâu thuẫn lớn
                        # với quan hệ CĐ đáy = CĐ miệng lỗ khoan - độ sâu, dùng giá trị suy từ độ sâu.
                        if pred_elev is not None and (bottom_elev > float(top_elev) + 0.50 or abs(bottom_elev - float(pred_elev)) > max(0.80, 0.06 * max(abs(float(bottom_depth)), 1.0))):
                            bottom_elev = float(pred_elev)
                    else:
                        bottom_elev = (float(top_elev) - float(bottom_depth) if top_elev is not None else -float(bottom_depth))
                    if large_batch:
                        raw_name = ""
                    else:
                        raw_name = self._ocr_layer_symbol_cell(img, max(0, x0 - 3), min(w, x1 + 8), int(y0) + 1, int(y1) - 1)
                    raw_sym_check = re.sub(r"[^A-Z0-9]", "", _strip_accents(str(raw_name or "")).upper())
                    raw_valid = (
                        raw_sym_check in ("D", "DC", "HANG", "K")
                        or bool(re.fullmatch(r"\d+[A-Z]?", raw_sym_check or ""))
                        or bool(re.fullmatch(r"[A-Z]?\d+[A-Z]?", raw_sym_check or ""))
                    )
                    # Nếu ranh giới này không có số depth/elev trực tiếp và ô tên cũng không rõ,
                    # nhiều khả năng đó là leader/mô tả/footer chứ không phải lớp địa chất.
                    # QA-OCR v5: bảo toàn ranh giới lớp theo hình học nhưng không mở quá rộng.
                    # Bản v4 bỏ lớp thật khi ô ký hiệu OCR rỗng; bản thử đầu v5 lại giữ quá nhiều
                    # đường giả. Quy tắc cân bằng:
                    #   (1) giữ lớp phủ đầu nếu đáy nông <= 1.0 m;
                    #   (2) giữ khoảng rất dày >= 6 m vì thường là lớp đá/lớp cuối bị OCR tên kém;
                    #   (3) các khoảng mỏng/không có marker trực tiếp vẫn bỏ để tránh footer/đường chia giả.
                    thickness_by_depth = float(bottom_depth) - float(prev_depth)
                    preserve_geom_without_name = (
                        (float(prev_depth) <= 0.05 and float(bottom_depth) <= 1.00)
                        or (thickness_by_depth >= 6.00)
                    )
                    if marker_depth_direct is None and not raw_valid:
                        if bottom_depth is None or (not preserve_geom_without_name):
                            continue
                    lname = ("?" if large_batch and not str(raw_name or "").strip() else self._clean_layer_name_from_ocr(raw_name, "", idx, names))
                    # Nếu ô tên rỗng nhưng ranh giới hình học thuộc nhóm bảo toàn thì giữ L<n> để người dùng soát.
                    if (not str(raw_name or '').strip()) and re.fullmatch(r"L\d+", str(lname or "")) and (not preserve_geom_without_name):
                        continue
                    # Với lớp đầu rất mỏng, nếu OCR đọc sai ký hiệu thì vẫn giữ tên đọc được, không ép D cho mọi form.
                    names.append(lname)
                    stype = self._infer_soil_type_from_borehole_text(lname, "")
                    page_rows.extend(self._make_ocr_rows_for_one_borehole_layer(
                        page_item_name, lname, prev_depth, float(bottom_depth), top_elev, stype, spt_points_use,
                        "", bottom_elev_override=float(bottom_elev)
                    ))
                    prev_depth = float(bottom_depth)
                if len(page_rows) >= 2:
                    # QA-OCR v5: lite parser cũng phải qua repair trước khi trả về.
                    # Trước đây hàm này trả thẳng tên OCR thô như 325330/375400425;
                    # nhánh full có repair nhưng benchmark/luồng fast-lite có thể dùng trực tiếp.
                    # Không truyền top_elev ở bước này để tránh xóa nhầm nếu OCR header sai;
                    # các nhánh full/PDF sẽ lọc theo top_elev sau nếu cần.
                    page_rows = self._repair_borehole_row_sequence(page_rows, top_elev=None)
                    rows.extend(page_rows)
            except Exception:
                continue
        warn = ""
        if rows:
            warn = "Đã nhận dạng nhanh từ lưới; ưu tiên đủ ranh giới lớp, tên lớp, CĐ đáy và SPT. Mô tả dài để trống nhằm tránh lệch cột."
        return rows, warn



    def _borehole_parallel_parser_score(self, rows: List[List[Any]]) -> float:
        """Chấm điểm kết quả import ảnh lỗ khoan khi chạy song song pipeline cũ/hybrid.

        Điểm này KHÔNG phải đánh giá thiết kế; chỉ dùng để chọn parser ít rủi ro hơn:
        - ưu tiên có nhiều ranh giới lớp hợp lý;
        - phạt tên lớp mơ hồ L<n> hoặc chuỗi OCR dính số;
        - phạt cao độ đáy không giảm dần trong cùng hạng mục;
        - cộng điểm nhẹ cho dòng có SPT.
        """
        try:
            rows2 = [list((r + [""] * 14)[:14]) if isinstance(r, list) else list(r) for r in (rows or [])]
        except Exception:
            rows2 = list(rows or [])
        if not rows2:
            return -9999.0
        q = self._borehole_rows_quality(rows2)
        score = float(q.get("score", 0.0))
        score += min(float(len(rows2)), 16.0) * 1.20
        groups: Dict[str, List[List[Any]]] = {}
        for r in rows2:
            key = str(r[0] if r else "")
            groups.setdefault(key, []).append(r)
        for _key, grp in groups.items():
            elevs: List[float] = []
            for r in grp:
                ev = self._borehole_row_bottom_elev(r)
                if ev is not None:
                    elevs.append(float(ev))
            # CĐ đáy hợp lệ phải giảm dần khi đi xuống. Không bắt quá chặt vì có file ghép trang.
            if len(elevs) >= 2:
                inversions = sum(1 for a, b in zip(elevs, elevs[1:]) if b > a + 0.05)
                score -= inversions * 8.0
                # Dải chiều sâu quá nhỏ mà nhiều dòng thường là đọc nhầm lưới/footer.
                if max(elevs) - min(elevs) < 1.0 and len(elevs) >= 3:
                    score -= 10.0
            score += min(len(elevs), 12) * 0.60
        # Phạt tỷ lệ mơ hồ quá cao, nhưng không loại hẳn vì hybrid dùng L<n> để bảo toàn lớp.
        amb = float(q.get("ambiguous", 0.0))
        n = max(float(q.get("n", len(rows2))), 1.0)
        amb_ratio = amb / n
        if amb_ratio > 0.55:
            score -= 10.0 * amb_ratio * n / max(n, 1.0)
        return score

    def _select_borehole_parallel_rows(
        self,
        legacy_rows: List[List[Any]], legacy_warn: str,
        hybrid_rows: List[List[Any]], hybrid_warn: str,
        top_elev: Optional[float] = None,
    ) -> Tuple[List[List[Any]], str]:
        """Chọn/gộp kết quả hai pipeline OCR ảnh lỗ khoan.

        Nguyên tắc kỹ thuật: không để pipeline mới ghi đè pipeline cũ khi nó kém rõ,
        nhưng cũng không bỏ các ranh giới lớp mà hybrid bắt được. Nếu hai pipeline cùng hợp lý,
        gộp theo cao độ đáy và giữ dòng có thông tin tốt hơn.
        """
        legacy_rows = list(legacy_rows or [])
        hybrid_rows = list(hybrid_rows or [])
        if not legacy_rows and not hybrid_rows:
            return [], (legacy_warn or hybrid_warn or "")
        if not legacy_rows:
            rows = self._repair_borehole_row_sequence(hybrid_rows, top_elev=top_elev)
            return rows, (hybrid_warn or "Đã nhận dạng bằng pipeline hybrid bảo toàn lớp.") + "\n[OCR AUTO] Chỉ pipeline hybrid có dữ liệu."
        if not hybrid_rows:
            rows = self._repair_borehole_row_sequence(legacy_rows, top_elev=top_elev)
            return rows, (legacy_warn or "Đã nhận dạng bằng pipeline cũ.") + "\n[OCR AUTO] Chỉ pipeline cũ có dữ liệu."

        legacy_rows = self._repair_borehole_row_sequence(legacy_rows, top_elev=top_elev)
        hybrid_rows = self._repair_borehole_row_sequence(hybrid_rows, top_elev=top_elev)
        ls = self._borehole_parallel_parser_score(legacy_rows)
        hs = self._borehole_parallel_parser_score(hybrid_rows)

        # Nếu một pipeline vượt rõ rệt, ưu tiên pipeline đó nhưng vẫn thử gộp để bổ sung lớp mỏng.
        if hs > ls + 8.0:
            primary, secondary = hybrid_rows, legacy_rows
            primary_name = "hybrid"
            base_warn = hybrid_warn
        elif ls > hs + 8.0:
            primary, secondary = legacy_rows, hybrid_rows
            primary_name = "cũ"
            base_warn = legacy_warn
        else:
            # Điểm gần nhau: lấy pipeline nhiều lớp hơn làm primary để giảm rủi ro mất lớp thật.
            if len(hybrid_rows) >= len(legacy_rows):
                primary, secondary = hybrid_rows, legacy_rows
                primary_name = "hybrid+gộp"
                base_warn = hybrid_warn or legacy_warn
            else:
                primary, secondary = legacy_rows, hybrid_rows
                primary_name = "cũ+gộp"
                base_warn = legacy_warn or hybrid_warn

        try:
            merged = self._merge_borehole_grid_rows(primary, secondary, top_elev=top_elev)
        except Exception:
            merged = list(primary)
        ms = self._borehole_parallel_parser_score(merged)
        best_rows = primary
        best_score = max(ls, hs)
        if merged and ms >= best_score - 4.0:
            best_rows = merged
            primary_name = primary_name + "+merge"
            best_score = ms

        warn_lines = []
        if base_warn:
            warn_lines.append(str(base_warn))
        warn_lines.append(
            f"[OCR AUTO] Đã chạy song song pipeline cũ ({len(legacy_rows)} dòng, score={ls:.1f}) "
            f"và hybrid bảo toàn lớp ({len(hybrid_rows)} dòng, score={hs:.1f}); "
            f"chọn {primary_name} ({len(best_rows)} dòng, score={best_score:.1f})."
        )
        if any(self._is_ambiguous_borehole_layer_name((r[1] if len(r) > 1 else "")) for r in best_rows):
            warn_lines.append("[OCR AUTO] Có dòng tên lớp dạng L<n>/mơ hồ; cần rà lại trong bảng preview trước khi import chính thức.")
        return best_rows, "\n".join(warn_lines).strip()

    def _try_make_geology_rows_from_borehole_image_grid_hybrid(self, path: str, item_name: str, ocr_text: str) -> Tuple[List[List[Any]], str]:
        """Parser chuyên cho form log lỗ khoan dạng bảng như ảnh BRITEC.

        Điểm khác OCR cũ: không cố đọc cả bảng thành dòng text hoàn chỉnh. Hàm này nhận diện
        đường kẻ ngang/dọc, suy ra chiều dày lớp từ vị trí đường kẻ theo tỷ lệ sâu, rồi OCR từng ô
        mô tả nhỏ. Nhờ vậy các ảnh nền tối, chữ nhỏ, tiêu đề xoay dọc sẽ ổn định hơn.
        """
        lite_rows: List[List[Any]] = []
        lite_warn = ""
        try:
            lite_rows, lite_warn = self._try_make_geology_rows_from_borehole_image_grid_lite(path, item_name, ocr_text)
        except Exception:
            lite_rows, lite_warn = [], ""
        rows: List[List[Any]] = []
        warn = ""
        try:
            images = self._images_from_borehole_file(path)
        except Exception:
            return self._repair_borehole_row_sequence(lite_rows), lite_warn
        # V1.0.22: nếu parser nhanh đã đủ tin cậy thì dùng ngay để tránh OCR mô tả dài.
        # Nếu chưa đủ tin cậy, chạy parser chi tiết nhưng vẫn chỉ đọc cột ký hiệu/cao độ/SPT.
        try:
            spt_points_mode = str(getattr(self, "_borehole_spt_mode", "avg") or "avg") == "points"
            if (not spt_points_mode) and lite_rows and (len(images) > 4 or len(lite_rows) >= 9 or (len(lite_rows) >= 4 and self._borehole_lite_rows_high_confidence(lite_rows, len(images)))):
                fast_warn = lite_warn or "Đã nhận dạng nhanh theo 3 cột lõi: Tên lớp, CĐ đáy và SPT. Mô tả dài để trống nhằm tránh lệch cột."
                return self._repair_borehole_row_sequence(lite_rows), fast_warn
        except Exception:
            pass
        multi_page = len(images) > 1
        lite_has_spt_any = any(str((r + [""] * 14)[4] or "").strip() for r in (lite_rows or []))
        for page_idx, img in enumerate(images, start=1):
            try:
                page_text = self._borehole_page_text(ocr_text, page_idx)
                if not str(page_text or '').strip():
                    try:
                        page_text = self._ocr_borehole_header_text_from_image(img, page_idx)
                    except Exception:
                        page_text = ""
                top_elev = self._extract_borehole_top_elev(page_text)
                # SPT vẫn lấy từ OCR text trang tương ứng; người dùng duyệt lại trong preview.
                spt_points = self._parse_spt_points_from_ocr_text(page_text)
                page_item_name = item_name if not multi_page else f"{item_name}_p{page_idx:02d}"
                w, h = img.size
                xlines, ylines = self._image_line_projection_groups(img)
                xlines = sorted([x for x in xlines if 0 <= x <= w])
                ylines = sorted([y for y in ylines if 0 <= y <= h])
                left_lines = [x for x in xlines if x < w * 0.42]
                if len(left_lines) < 5 or len(ylines) < 4:
                    continue
                # Form chuẩn: Ký hiệu | Cao độ | Độ sâu | Bề dày | Trụ cắt...
                # Một số PDF AutoCAD khác: Tỷ lệ | Tên lớp | Cao độ | Độ sâu | Bề dày | Trụ cắt...
                # Nếu có cột tỷ lệ ở ngoài cùng trái, dịch bộ cột sang phải một nhịp để không đọc
                # nhầm thang tỷ lệ thành tên lớp và không bỏ các lớp ở trang dạng Cái Cám.
                scale_col_layout = bool(len(left_lines) >= 6 and (w > h * 1.08 or left_lines[0] > w * 0.22))
                # QA-OCR STRICT: một số ảnh/screenshot có đường biên trang rất sát mép trái
                # (x≈0..8 px) đứng trước cột "Lớp" thật. Bản cũ lấy nhầm đường này làm
                # biên cột ký hiệu lớp, khiến cột Cao độ bị đọc nhầm thành Độ sâu và sinh
                # hàng loạt lớp/SPT sai. Chỉ bỏ qua cột đầu khi nó sát mép VÀ rất hẹp;
                # không áp cho form R3 vì cột đầu của R3 rộng ~40 px là cột thật.
                leading_edge_col_layout = bool(
                    len(left_lines) >= 6
                    and int(left_lines[0]) <= max(10, int(w * 0.015))
                    and int(left_lines[1] - left_lines[0]) <= max(35, int(w * 0.055))
                )
                if scale_col_layout or leading_edge_col_layout:
                    x0, x1, x2, x3, x4 = left_lines[1:6]
                else:
                    x0, x1, x2, x3, x4 = left_lines[:5]
                # Tìm cột mô tả theo khoảng rộng nhất sau cụm trụ cắt/lỗ khoan.
                x_desc0, x_desc1 = self._find_borehole_desc_xrange(xlines, w, x4)
                if x_desc1 <= x_desc0 + 20:
                    continue
                # Các đường ngang dữ liệu thường nằm từ khoảng 1/4 chiều cao ảnh trở xuống.
                # Lọc thêm theo số pixel tối qua vùng từ cột ký hiệu tới hết mô tả để loại các
                # đường/ngang giả chỉ thuộc đồ thị SPT hoặc chữ mô tả.
                try:
                    from PIL import ImageOps
                    line_img = ImageOps.grayscale(img)
                    if self._pil_mean_gray(line_img) < 128.0:
                        line_img = ImageOps.invert(line_img)
                    pix_line = line_img.load()
                    # Với form có cột tỷ lệ / trang AutoCAD xoay, ranh giới lớp thường không
                    # đủ dài để lọt qua ngưỡng projection chung. Bổ sung yline bằng cách chỉ đếm
                    # trong cụm cột trái sạch (Tên lớp/Cao độ/Độ sâu/Bề dày), tránh nhiễu hatch.
                    if scale_col_layout:
                        try:
                            extra_ys: List[int] = []
                            left_only_span = max(1, x4 - x0)
                            for yy2 in range(max(0, int(h * 0.18)), min(h - 1, h)):
                                cnt2 = 0
                                for xx2 in range(max(0, x0), min(w - 1, x4) + 1):
                                    if pix_line[xx2, yy2] < 105:
                                        cnt2 += 1
                                if cnt2 >= max(55, int(left_only_span * 0.42)):
                                    extra_ys.append(yy2)
                            if extra_ys:
                                grouped_extra: List[int] = []
                                st2 = pv2 = extra_ys[0]
                                for vv2 in extra_ys[1:]:
                                    if vv2 > pv2 + 2:
                                        grouped_extra.append((st2 + pv2) // 2)
                                        st2 = vv2
                                    pv2 = vv2
                                grouped_extra.append((st2 + pv2) // 2)
                                ylines = sorted(set(list(ylines) + grouped_extra))
                        except Exception:
                            pass
                    span_len = max(1, x_desc1 - x0)
                    y_keep: List[Tuple[int, int]] = []
                    for yy in ylines:
                        min_y_data = h * (0.18 if scale_col_layout else 0.20)
                        if not (yy > min_y_data and yy < h - 2):
                            continue
                        cnt = 0
                        for xx in range(max(0, x0), min(w - 1, x_desc1) + 1):
                            if pix_line[xx, yy] < 105:
                                cnt += 1
                        if cnt >= max(70, int(span_len * 0.38)):
                            y_keep.append((yy, cnt))
                    # Ưu tiên các đường ranh giới có nét trong cụm cột trái
                    # (ký hiệu lớp/cao độ/độ sâu/bề dày). Các đường ngang trong cột mô tả
                    # hoặc đồ thị SPT có thể là đường leader/gạch nền, không phải ranh giới lớp.
                    left_span = max(1, x4 - x0)
                    left_y_keep: List[Tuple[int, int]] = []
                    for yy, cnt_total in y_keep:
                        cnt_left = 0
                        for xx in range(max(0, x0), min(w - 1, x4) + 1):
                            if pix_line[xx, yy] < 105:
                                cnt_left += 1
                        if cnt_left >= max(35, int(left_span * 0.28)):
                            left_y_keep.append((yy, cnt_left))
                    primary = left_y_keep if len(left_y_keep) >= 3 else y_keep
                    ys = []
                    for yy, cnt in primary:
                        # Chỉ gộp các đường rất sát nhau do bề dày nét scan.
                        # Không gộp cụm 2-3 đường cách nhau 6-15 px vì đó thường là lớp phủ mỏng.
                        if not ys or yy - ys[-1] > 4:
                            ys.append(yy)
                        else:
                            ys[-1] = yy
                    ys = sorted(set(ys))
                except Exception:
                    ys = [y for y in ylines if y > h * 0.20 and y < h - 2]
                if len(ys) < 3:
                    continue
                # Xác định đáy lỗ khoan bằng giá trị lớn nhất trong cột độ sâu đáy lớp.
                depth_vals: List[float] = []
                for i in range(len(ys) - 1):
                    y0, y1 = ys[i], ys[i + 1]
                    if y1 - y0 < 5:
                        continue
                    pad = 2
                    txt = self._ocr_numeric_cell(img.crop((max(0, x2 + pad), y0 + pad, min(w, x3 - pad), y1 - pad)), psm=7)
                    for v in self._ocr_numeric_candidates(txt):
                        if 1.0 <= abs(v) <= 150.0:
                            depth_vals.append(abs(v))
                total_depth = max(depth_vals) if depth_vals else 0.0
                # V1.0.11: chưa loại trang ngay nếu OCR cột độ sâu đọc sai. Ta sẽ hiệu chỉnh lại
                # bằng các mốc số ở đường ranh giới lớp sau khi xác định candidate_rows.
                if not (3.0 <= total_depth <= 150.0):
                    total_depth = 0.0
                candidate_rows: List[Dict[str, Any]] = []
                # V1.0.22: không OCR mô tả dài để quyết định lớp vì mô tả có thể tràn sang ô khác.
                # Candidate lớp được lấy từ ranh giới lưới ở cụm cột trái; tên lớp chỉ đọc từ cột Ký hiệu.
                for i in range(len(ys) - 1):
                    y0, y1 = ys[i], ys[i + 1]
                    if y1 - y0 < 5:
                        continue
                    raw_name = self._ocr_layer_symbol_cell(img, max(0, x0 - 3), min(w, x1 + 8), int(y0) + 2, int(y1) - 2)
                    rlow = _strip_accents(raw_name).lower()
                    raw_sym = re.sub(r"[^A-Z0-9]", "", _strip_accents(raw_name).upper())
                    header_like = any(k in rlow for k in [
                        "ky hieu", "cao do", "do sau", "be day", "lo khoan", "mo ta",
                        "dia tang", "thi nghiem", "so bua", "do thi", "spt"
                    ])
                    if header_like:
                        continue
                    name_like = (
                        raw_sym in ("D", "DC", "DD", "HANG", "8", "8A", "8B", "4A", "4B", "CS")
                        or "HANG" in raw_sym
                        or re.fullmatch(r"\d+[A-Z]?", raw_sym or "") is not None
                    )
                    # Giữ lớp đầu rất mỏng nếu nó nằm ngay đầu vùng dữ liệu, nhưng không ép tên thành D.
                    maybe_thin_top = (not candidate_rows) and (y1 - y0 <= max(18, int(h * 0.018))) and y0 > h * 0.18
                    # Với đường ranh giới đã lọc từ cụm cột trái, giữ cả dòng chưa OCR rõ tên nếu chiều cao hợp lý.
                    # Dòng chưa rõ tên sẽ là L<n> để người dùng sửa, thay vì đoán theo mô tả.
                    geometry_like = (y1 - y0) >= max(8, int(h * 0.004))
                    if not (name_like or maybe_thin_top or geometry_like):
                        continue
                    candidate_rows.append({"y0": y0, "y1": y1, "desc": "", "raw_name": raw_name})
                if len(candidate_rows) < 2:
                    continue
                # V1.0.16: không tự thay cao độ toàn trang theo sai khác header.
                # Đọc trực tiếp cột Cao độ đáy lớp và cột Độ sâu đáy lớp tại từng ranh giới;
                # top_elev chỉ là fallback cho những ranh giới không đọc được cao độ.
                data_top = min(r["y0"] for r in candidate_rows)
                data_bottom = max(r["y1"] for r in candidate_rows)
                if data_bottom <= data_top:
                    continue

                boundary_ys = sorted(set([int(round(data_top))] + [int(round(float(r["y1"]))) for r in candidate_rows]))
                elev_markers = self._read_borehole_elev_markers_from_boundaries(img, boundary_ys, x1, x2)
                # Trục sâu robust chỉ lấy từ cột Độ sâu đáy lớp. Không dùng top_elev để biến cột cao độ
                # thành độ sâu, vì nếu OCR header sai thì trục sâu cũng bị kéo sai.
                depth_markers = self._read_borehole_depth_markers_from_boundaries(
                    img, boundary_ys, x1, x2, x2, x3, None
                )
                inferred_top_elev = self._infer_borehole_top_elev_from_marker_pairs(elev_markers, depth_markers)
                top_elev_fallback = top_elev
                if top_elev_fallback is None and inferred_top_elev is not None:
                    top_elev_fallback = float(inferred_top_elev)
                elif top_elev is not None and inferred_top_elev is not None and abs(float(inferred_top_elev) - float(top_elev)) > 0.75:
                    # Chỉ dùng giá trị suy từ bảng cho các ô bị thiếu cao độ; không dịch các ô đã đọc được.
                    top_elev_fallback = float(inferred_top_elev)
                depth_markers.append((int(round(data_top)), 0.0, "data_top"))
                depth_axis = self._fit_borehole_depth_axis(depth_markers)
                axis_bottom = self._depth_from_axis(depth_axis, data_bottom)
                if axis_bottom is not None and 3.0 <= axis_bottom <= 150.0:
                    # Nếu OCR toàn cột cho total_depth lệch lớn, dùng trục robust.
                    if not (3.0 <= total_depth <= 150.0) or abs(float(total_depth) - float(axis_bottom)) > max(0.75, 0.03 * axis_bottom):
                        total_depth = float(axis_bottom)
                if not (3.0 <= total_depth <= 150.0):
                    continue

                # Đọc SPT từ đúng cột N/30 của bảng, lấy tọa độ y để quy đổi theo độ sâu.
                # V1.0.20: nếu parser lite đã có SPT trung bình, không OCR lại cột SPT ở parser chi tiết;
                # bước merge theo cao độ sẽ lấy SPT từ lite_rows. Đây là phần tiết kiệm thời gian lớn
                # mà không làm mất ranh giới lớp/mô tả.
                points_mode = str(getattr(self, "_borehole_spt_mode", "avg") or "avg") == "points"
                if lite_has_spt_any and not points_mode:
                    spt_points_img = []
                    spt_points_use = []
                else:
                    spt_points_img = self._extract_spt_points_from_n30_column(img, xlines, x_desc1, data_top, data_bottom, total_depth)
                    spt_points_use = spt_points_img if spt_points_img else spt_points
                prev_depth = 0.0
                prev_bottom_elev_value: Optional[float] = None
                names: List[str] = []
                page_rows: List[List[Any]] = []
                for idx, r in enumerate(candidate_rows):
                    y0, y1 = float(r["y0"]), float(r["y1"])
                    axis_bottom_depth = self._depth_from_axis(depth_axis, y1)
                    axis_top_depth = self._depth_from_axis(depth_axis, y0)
                    geom_bottom_depth = (y1 - data_top) / max(data_bottom - data_top, 1.0) * total_depth
                    geom_top_depth = (y0 - data_top) / max(data_bottom - data_top, 1.0) * total_depth
                    if axis_bottom_depth is not None and axis_top_depth is not None:
                        geom_bottom_depth = axis_bottom_depth
                        geom_top_depth = axis_top_depth
                    geom_thickness = max(geom_bottom_depth - geom_top_depth, 0.0)
                    # Nếu OCR đọc rõ số độ sâu tại ranh giới lớp và khớp trục robust thì dùng số OCR,
                    # nếu không thì dùng giá trị nội suy hình học đã hiệu chỉnh.
                    marker_bottom_depth = self._lookup_depth_marker_near(depth_markers, y1, depth_axis)
                    bottom_depth = marker_bottom_depth if marker_bottom_depth is not None else geom_bottom_depth
                    # Dự đoán cao độ chỉ để chọn/sửa cục bộ ứng viên OCR ở cột Cao độ đáy.
                    pred_bottom_elev = (float(top_elev_fallback) - float(bottom_depth)) if top_elev_fallback is not None else None
                    marker_bottom_elev = self._lookup_elev_marker_near(elev_markers, y1, pred_bottom_elev)
                    thickness = max(bottom_depth - prev_depth, geom_thickness)
                    # V1.0.13: không OCR lại từng ô bề dày/độ sâu nếu đã có marker/trục sâu.
                    # Đây là nguồn gây chậm lớn ở V1.0.12. Chỉ fallback OCR ô độ sâu khi chưa fit
                    # được trục và chưa có marker gần ranh giới.
                    dep_vals: List[float] = []
                    if depth_axis is None and marker_bottom_depth is None:
                        depth_txt = self._ocr_numeric_cell(img.crop((max(0, x2 + 2), int(y0) + 2, min(w, x3 - 2), int(y1) - 2)), psm=7)
                        dep_vals = [abs(v) for v in self._ocr_numeric_candidates(depth_txt) if prev_depth < abs(v) <= total_depth + 0.5]
                    used_ocr = False
                    if (not used_ocr) and dep_vals:
                        dep = dep_vals[-1]
                        geom_close = abs(dep - geom_bottom_depth) <= max(0.45, 0.025 * max(total_depth, 1.0))
                        if dep > prev_depth + 0.03 and geom_close:
                            bottom_depth = dep
                            thickness = bottom_depth - prev_depth
                            used_ocr = True
                    if marker_bottom_elev is not None:
                        bottom_elev = float(marker_bottom_elev)
                        if pred_bottom_elev is not None and top_elev_fallback is not None and (bottom_elev > float(top_elev_fallback) + 0.50 or abs(bottom_elev - float(pred_bottom_elev)) > max(0.80, 0.06 * max(abs(float(bottom_depth)), 1.0))):
                            bottom_elev = float(pred_bottom_elev)
                        local_top_elev = float(bottom_elev) + float(bottom_depth)
                    else:
                        local_top_elev = top_elev_fallback
                        bottom_elev = (float(local_top_elev) - bottom_depth) if local_top_elev is not None else -bottom_depth
                    # Không lấy mô tả dài vào bảng import; mô tả dễ lệch cột khi lớp mỏng.
                    desc = ""
                    lname = self._clean_layer_name_from_ocr(str(r.get("raw_name") or ""), "", idx, names)
                    # Không ép lớp đầu thành D và không sửa theo chuỗi D->2->3->8A/Hang.
                    # Nếu OCR ký hiệu chưa rõ, giữ L<n> để người dùng hiệu chỉnh trong preview.
                    names.append(lname)
                    points = [n for d, n in spt_points_use if prev_depth < d <= bottom_depth]
                    n_avg = round(sum(points) / len(points)) if points else 0
                    stype = self._infer_soil_type_from_borehole_text(lname, "")
                    # Lớp phủ ký hiệu D thường không có chỉ tiêu SPT thiết kế trong log.
                    if str(lname).strip().upper() == "D":
                        n_avg = 0
                    # Ảnh/log lỗ khoan chỉ có tên lớp, cao độ đáy, phân loại lớp và SPT.
                    # Không tự sinh γ, C, φ, qu, RQD, GSI, mi, D vì trong ảnh không có các chỉ tiêu này.
                    gamma = ""
                    c_mpa = ""
                    qu_mpa = ""
                    rqd = ""
                    gsi = ""
                    mi = ""
                    dist = ""
                    page_rows.extend(self._make_ocr_rows_for_one_borehole_layer(
                        page_item_name, lname, prev_depth, bottom_depth, local_top_elev, stype, spt_points_use, desc, bottom_elev_override=bottom_elev
                    ))
                    prev_depth = bottom_depth
                    try:
                        prev_bottom_elev_value = float(bottom_elev)
                    except Exception:
                        prev_bottom_elev_value = None
                if len(page_rows) >= 2:
                    rows.extend(page_rows)
                    # Không break: PDF/TIFF nhiều trang cần đọc tiếp các trang sau, mỗi trang là một hạng mục riêng.
                    continue
            except Exception:
                continue
        if rows:
            # QA-OCR v2: chỉ áp lọc theo cao độ miệng khi mọi dòng cùng 1 hạng mục
            # (file nhiều trang = nhiều hạng mục có cao độ miệng khác nhau).
            try:
                sanity_top_elev = local_top_elev if len(set(str(r[0]) for r in rows)) <= 1 else None
            except Exception:
                sanity_top_elev = None
            # Gộp thêm các ranh giới lớp mỏng mà parser nhanh bắt được nhưng parser chi tiết bỏ sót.
            if lite_rows:
                rows = self._merge_borehole_grid_rows(rows, lite_rows, top_elev=sanity_top_elev)
            else:
                rows = self._repair_borehole_row_sequence(rows, top_elev=sanity_top_elev)
            if not any(str(r[4] or "").strip() for r in rows):
                warn = "Đã nhận được lớp bằng parser lưới ảnh, nhưng chưa nhận được cột N/30 SPT; cột SPT đang để trống. Anh kiểm tra/sửa ở bảng preview trước khi import."
            else:
                warn = "Đã nhận dạng theo 3 cột lõi: Tên lớp, CĐ đáy và SPT. Mô tả dài được bỏ qua để tránh lệch cột; anh kiểm tra/sửa loại đất ở preview nếu cần."
            return rows, warn
        if lite_rows:
            return self._repair_borehole_row_sequence(lite_rows), (lite_warn or "Đã nhận dạng nhanh từ lưới ảnh; anh kiểm tra lại tên lớp, loại đất và SPT ở preview.")
        return rows, warn


    def _try_make_geology_rows_from_borehole_image_grid(self, path: str, item_name: str, ocr_text: str) -> Tuple[List[List[Any]], str]:
        """OCR ảnh lỗ khoan chế độ tích hợp song song: pipeline cũ + hybrid bảo toàn lớp.

        Mặc định chạy SONG SONG hai parser lưới nhanh:
        - legacy lite: giống logic cũ, ổn định và nhanh;
        - hybrid lite: bảo toàn ranh giới lớp tốt hơn cho ảnh khó/R3.
        Chỉ khi cả hai parser lưới không có dữ liệu mới fallback sang parser chi tiết chậm.
        Có thể ép chạy chi tiết đầy đủ bằng self._borehole_parser_mode = "PARALLEL_FULL".
        """
        top_elev = None
        try:
            top_elev = self._extract_borehole_top_elev(ocr_text)
        except Exception:
            top_elev = None
        force_full = str(getattr(self, "_borehole_parser_mode", "AUTO") or "AUTO").upper() in ("PARALLEL", "PARALLEL_FULL", "AUTO_FULL")

        legacy_rows: List[List[Any]] = []
        hybrid_rows: List[List[Any]] = []
        legacy_warn = ""
        hybrid_warn = ""
        if not force_full:
            try:
                legacy_rows, legacy_warn = self._try_make_geology_rows_from_borehole_image_grid_lite_legacy(path, item_name, ocr_text)
                legacy_rows = self._repair_borehole_row_sequence(legacy_rows, top_elev=top_elev)
            except Exception as exc:
                legacy_rows, legacy_warn = [], f"Pipeline cũ-lite lỗi: {exc}"
            try:
                hybrid_rows, hybrid_warn = self._try_make_geology_rows_from_borehole_image_grid_lite(path, item_name, ocr_text)
                hybrid_rows = self._repair_borehole_row_sequence(hybrid_rows, top_elev=top_elev)
            except Exception as exc:
                hybrid_rows, hybrid_warn = [], f"Pipeline hybrid-lite lỗi: {exc}"

            if legacy_rows or hybrid_rows:
                rows, warn = self._select_borehole_parallel_rows(legacy_rows, legacy_warn, hybrid_rows, hybrid_warn, top_elev=top_elev)
                warn = (warn or "") + "\n[OCR AUTO] Chế độ nhanh: đã chạy song song legacy-lite và hybrid-lite; parser chi tiết chậm chỉ dùng khi hai parser nhanh không có dữ liệu."
                return rows, warn.strip()

        # Fallback chậm: chỉ dùng khi parser nhanh không có dữ liệu hoặc người dùng ép PARALLEL_FULL.
        try:
            legacy_rows, legacy_warn = self._try_make_geology_rows_from_borehole_image_grid_legacy(path, item_name, ocr_text)
        except Exception as exc:
            legacy_rows, legacy_warn = [], f"Pipeline cũ lỗi: {exc}"
        try:
            hybrid_rows, hybrid_warn = self._try_make_geology_rows_from_borehole_image_grid_hybrid(path, item_name, ocr_text)
        except Exception as exc:
            hybrid_rows, hybrid_warn = [], f"Pipeline hybrid lỗi: {exc}"
        rows, warn = self._select_borehole_parallel_rows(legacy_rows, legacy_warn, hybrid_rows, hybrid_warn, top_elev=top_elev)
        return rows, ((warn or "") + "\n[OCR AUTO] Fallback chi tiết: parser nhanh không đủ dữ liệu hoặc đang ép PARALLEL_FULL.").strip()

    def _rows_from_borehole_text_layers(self, item_name: str, text: str) -> Tuple[List[List[Any]], bool]:
        """Fallback tạo dòng địa chất từ text OCR/PDF của một trang hoặc một file.

        Trả thêm bool has_spt để sinh cảnh báo đúng hơn. Hàm này tách riêng để PDF nhiều trang
        có thể fallback theo từng trang, thay vì chỉ cần parser lưới thành công ở trang 1 là bỏ qua
        các trang còn lại.
        """
        layers = self._parse_borehole_layers_from_ocr_text(text, item_name)
        spt_points = self._parse_spt_points_from_ocr_text(text)
        rows: List[List[Any]] = []
        prev_depth = 0.0
        for lay in layers:
            bottom_depth = float(lay.get("bottom_depth") or 0.0)
            points = [n for d, n in spt_points if prev_depth < d <= bottom_depth]
            n_avg = round(sum(points) / len(points)) if points else 0
            desc = str(lay.get("desc") or "")
            stype = self._infer_soil_type_from_borehole_text(lay.get("name"), desc)
            lname_tmp = str(lay.get("name") or "").strip().upper()
            if lname_tmp == "D" or "dat dap" in _strip_accents(str(desc or "")).lower() or "dat phu" in _strip_accents(str(desc or "")).lower():
                n_avg = 0
            lname = str(lay.get("name") or "").upper() if re.fullmatch(r"[a-z]", str(lay.get("name") or "")) else str(lay.get("name") or "")
            top_elev = None
            try:
                top_elev = float(lay.get("bottom_elev") or 0.0) + bottom_depth
            except Exception:
                top_elev = None
            rows.extend(self._make_ocr_rows_for_one_borehole_layer(
                item_name, lname, prev_depth, bottom_depth, top_elev, stype, spt_points, desc
            ))
            prev_depth = bottom_depth
        return rows, bool(spt_points)

    def _make_geology_rows_from_borehole_ocr(self, path: str, text: str) -> Tuple[str, List[List[Any]], str]:
        item_name = self._borehole_item_name_from_path(path)
        if os.path.splitext(str(path or ""))[1].lower() == ".pdf":
            vector_rows, vector_warn, _vector_text = self._try_make_geology_rows_from_pdf_vector_text(path, item_name)
            if vector_rows:
                return item_name, vector_rows, vector_warn
        grid_rows, grid_warn = self._try_make_geology_rows_from_borehole_image_grid(path, item_name, text)
        rows: List[List[Any]] = list(grid_rows or [])
        warnings2: List[str] = []
        has_spt_any = any(str(r[4] or "").strip() for r in rows)

        # V1.0.12: với PDF/TIFF nhiều trang, parser lưới có thể đọc được trang 1 nhưng bỏ trang 2
        # do form lệch/scan khác. Không được return ngay; phải kiểm từng trang còn thiếu và fallback
        # bằng text OCR/PDF text riêng của trang đó.
        page_count = self._borehole_file_page_count(path)
        if page_count > 1:
            present_pages = set()
            prefix = f"{item_name}_p"
            for r in rows:
                hm = str(r[0] if r else "")
                if hm.startswith(prefix):
                    m = re.search(r"_p(\d+)\b", hm)
                    if m:
                        try:
                            present_pages.add(int(m.group(1)))
                        except Exception:
                            pass
            for page_no in range(1, page_count + 1):
                if page_no in present_pages:
                    continue
                page_text = self._borehole_page_text(text, page_no)
                page_item_name = f"{item_name}_p{page_no:02d}"
                if not str(page_text or "").strip():
                    warnings2.append(f"Trang {page_no}: chưa đọc được dữ liệu text/OCR.")
                    continue
                fb_rows, fb_has_spt = self._rows_from_borehole_text_layers(page_item_name, page_text)
                if fb_rows:
                    # QA-OCR v4: fallback text cũng phải qua repair theo trang (như luồng fast).
                    try:
                        fb_top_elev = self._extract_borehole_top_elev(page_text)
                    except Exception:
                        fb_top_elev = None
                    fb_rows = self._repair_borehole_row_sequence(fb_rows, top_elev=fb_top_elev)
                if fb_rows:
                    rows.extend(fb_rows)
                    has_spt_any = has_spt_any or fb_has_spt
                    warnings2.append(f"Trang {page_no}: đã đọc theo text/OCR.")
                else:
                    warnings2.append(f"Trang {page_no}: chưa nhận được lớp địa chất; cần kiểm tra trong Xem text OCR hoặc tách trang thành ảnh rõ hơn.")
        elif not rows:
            rows, has_spt_any = self._rows_from_borehole_text_layers(item_name, text)
            if rows:
                # QA-OCR v4: fallback 1 trang cũng qua repair với cao độ miệng đọc từ text.
                try:
                    fb_top_elev = self._extract_borehole_top_elev(text)
                except Exception:
                    fb_top_elev = None
                rows = self._repair_borehole_row_sequence(rows, top_elev=fb_top_elev)

        warn = ""
        if not rows:
            warn = "OCR chưa nhận được lớp địa chất theo cấu trúc bảng. Có thể ảnh/PDF quá mờ hoặc form khác mẫu; hãy dùng nút Xem text OCR để copy/sửa thủ công."
        elif not has_spt_any:
            warn = "Đã nhận được lớp nhưng chưa nhận được bảng SPT; cột Nₕₜ đang để trống."
        else:
            warn = grid_warn or "Đã nhận dạng lớp và lấy SPT trung bình theo từng lớp từ cột N/30. Anh vẫn nên kiểm tra lại lớp, SPT và loại đất trong bảng preview trước khi import."
        if page_count > 1:
            got_pages = set()
            prefix = f"{item_name}_p"
            for r in rows:
                hm = str(r[0] if r else "")
                if hm.startswith(prefix):
                    m = re.search(r"_p(\d+)\b", hm)
                    if m:
                        try:
                            got_pages.add(int(m.group(1)))
                        except Exception:
                            pass
            if len(got_pages) < page_count:
                warnings2.insert(0, f"PDF/TIFF có {page_count} trang, mới nhận được {len(got_pages)} trang có dữ liệu.")
            else:
                warnings2.insert(0, f"PDF/TIFF có {page_count} trang, đã nhận được đủ {len(got_pages)} trang.")
        if warnings2:
            warn = (warn + "\n" + "\n".join(warnings2)).strip()
        return item_name, rows, warn

    def _show_ocr_preview_dialog(self, title: str, rows: List[List[Any]], ocr_text: str) -> Optional[List[List[Any]]]:
        dlg = tk.Toplevel(self.root)
        dlg.title(title)
        dlg.transient(self.root)
        dlg.grab_set()
        dlg.geometry("1180x620")
        ttk.Label(dlg, text="Kiểm tra/sửa dữ liệu OCR trước khi đổ vào bảng địa chất. OCR lỗ khoan chỉ tự điền: Lớp, CĐ đáy, Loại và SPT theo tùy chọn đã chọn; các chỉ tiêu γ, C, φ, qu, RQD, GSI, mi, D để trống vì ảnh không có dữ liệu này. Double-click để sửa ô. Loại đất: 0=không khí/karst, 1=cát, 2=sét, 3=đá nguyên khối, 4=đá nứt vỡ, 5=IGM, 6=cuội sỏi.", wraplength=1120).pack(anchor=tk.W, padx=10, pady=(10, 4))
        frame = ttk.Frame(dlg)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=6)
        cols = [
            ("geo_item", "Lỗ Khoan", 120), ("name", "Lớp", 70), ("bottom", "CĐ đáy", 80), ("type", "Loại", 58),
            ("n", "Nₕₜ", 65), ("gamma", "γ", 65), ("c", "C MPa", 70), ("phi", "φ độ", 65),
            ("qu", "qu MPa", 70), ("rqd", "RQD", 60), ("gsi", "GSI", 60), ("mi", "mi", 60), ("dist", "D", 55), ("comment", "Mô tả/Ghi chú", 260)
        ]
        table = EditableTree(frame, cols, height=18)
        table.set_rows(rows)
        result: Dict[str, Any] = {"ok": False}
        btns = ttk.Frame(dlg)
        btns.pack(fill=tk.X, padx=10, pady=8)
        def show_text():
            tw = tk.Toplevel(dlg)
            tw.title("Text OCR thô")
            tw.geometry("900x600")
            txt = tk.Text(tw, wrap=tk.WORD, font=("Consolas", 10))
            txt.pack(fill=tk.BOTH, expand=True)
            txt.insert("1.0", ocr_text)
            txt.focus_set()
        def ok():
            result["ok"] = True
            result["rows"] = table.get_rows()
            dlg.destroy()
        def cancel():
            result.clear()
            dlg.destroy()
        ttk.Button(btns, text="Xem text OCR", command=show_text).pack(side=tk.LEFT, padx=3)
        ttk.Button(btns, text="OK - Đổ vào bảng địa chất", command=ok).pack(side=tk.RIGHT, padx=3)
        ttk.Button(btns, text="Hủy", command=cancel).pack(side=tk.RIGHT, padx=3)
        self.root.wait_window(dlg)
        return result.get("rows") if result.get("ok") else None

    def _ask_borehole_spt_import_options(self) -> Optional[Tuple[str, float]]:
        dlg = tk.Toplevel(self.root)
        dlg.title("Tùy chọn nhập OCR lỗ khoan")
        dlg.transient(self.root)
        dlg.grab_set()
        dlg.geometry("500x250")
        ttk.Label(dlg, text="Cách lấy SPT khi nhập từ ảnh/PDF lỗ khoan:", wraplength=420).pack(anchor=tk.W, padx=14, pady=(14, 6))
        ttk.Label(dlg, text="V1.0.16 dùng chế độ nhanh: ưu tiên đọc trực tiếp cột Cao độ đáy theo từng ranh giới; không dịch cao độ toàn trang theo sai khác header.", foreground="#475569", wraplength=420).pack(anchor=tk.W, padx=22, pady=(0, 4))
        mode = tk.StringVar(value=str(getattr(self, "_borehole_spt_mode", "avg") or "avg"))
        spacing_var = tk.StringVar(value=str(getattr(self, "_borehole_spt_spacing_m", 2.0) or 2.0))
        ttk.Radiobutton(dlg, text="Lấy SPT trung bình trong từng lớp", variable=mode, value="avg").pack(anchor=tk.W, padx=22, pady=3)
        ttk.Radiobutton(dlg, text="Lấy đủ SPT từng điểm và chia lớp theo khoảng cách điểm", variable=mode, value="points").pack(anchor=tk.W, padx=22, pady=3)
        row = ttk.Frame(dlg)
        row.pack(fill=tk.X, padx=22, pady=(8, 4))
        ttk.Label(row, text="Khoảng cách điểm SPT (m):").pack(side=tk.LEFT)
        ttk.Entry(row, textvariable=spacing_var, width=10).pack(side=tk.LEFT, padx=8)
        result: Dict[str, Any] = {}
        def ok():
            try:
                sp = float(str(spacing_var.get()).replace(",", "."))
            except Exception:
                sp = 2.0
            if sp <= 0:
                sp = 2.0
            result["mode"] = mode.get()
            result["spacing"] = sp
            dlg.destroy()
        def cancel():
            result.clear()
            dlg.destroy()
        btns = ttk.Frame(dlg)
        btns.pack(fill=tk.X, padx=14, pady=12)
        ttk.Button(btns, text="OK", command=ok).pack(side=tk.RIGHT, padx=4)
        ttk.Button(btns, text="Hủy", command=cancel).pack(side=tk.RIGHT, padx=4)
        self.root.wait_window(dlg)
        if not result:
            return None
        return str(result.get("mode") or "avg"), float(result.get("spacing") or 2.0)

    def import_geology_from_borehole_image(self):
        paths = filedialog.askopenfilenames(
            title="Chọn một hoặc nhiều ảnh/PDF hình trụ lỗ khoan",
            filetypes=[
                ("Ảnh/PDF", "*.png *.jpg *.jpeg *.tif *.tiff *.bmp *.webp *.pdf"),
                ("All files", "*.*"),
            ],
        )
        if not paths:
            return
        opts = self._ask_borehole_spt_import_options()
        if opts is None:
            return
        self._borehole_spt_mode, self._borehole_spt_spacing_m = opts
        # Chụp engine hiện tại trước khi chạy worker; worker không đọc trực tiếp Tk StringVar.
        self._borehole_ocr_engine_runtime = self._borehole_ocr_engine_key()
        # QA fix O1: xác định đường dẫn Tesseract NGAY TRÊN MAIN THREAD (được phép hỏi người dùng).
        # Worker nền sau đó chỉ dùng kết quả đã cache, không bao giờ mở hộp thoại từ thread phụ.
        if not self._get_tesseract_cmd(ask_user=False):
            self._get_tesseract_cmd(ask_user=True)

        progress_dlg = tk.Toplevel(self.root)
        progress_dlg.title("Đang import địa chất")
        progress_dlg.transient(self.root)
        progress_dlg.resizable(False, False)
        frm = ttk.Frame(progress_dlg, padding=14)
        frm.pack(fill=tk.BOTH, expand=True)
        progress_text = tk.StringVar(value="Chuẩn bị quét ảnh/PDF...")
        ttk.Label(frm, textvariable=progress_text, wraplength=520).pack(anchor=tk.W, pady=(0, 8))
        pbar = ttk.Progressbar(frm, mode="determinate", maximum=max(len(paths), 1), value=0, length=520)
        pbar.pack(fill=tk.X, pady=(0, 10))
        note = ttk.Label(
            frm,
            text="OCR đang chạy nền. Giao diện có thể thao tác nhẹ, nhưng không nên sửa bảng địa chất cho đến khi import xong.",
            style="Muted.TLabel",
            wraplength=520,
        )
        note.pack(anchor=tk.W)
        cancel_flag = {"cancel": False}
        def request_cancel():
            cancel_flag["cancel"] = True
            progress_text.set("Đang yêu cầu dừng sau file hiện tại...")
        ttk.Button(frm, text="Dừng sau file hiện tại", command=request_cancel).pack(anchor=tk.E, pady=(10, 0))
        self._center_window(progress_dlg)

        def ui_progress(idx: int, total: int, base: str):
            try:
                progress_text.set(f"Đang quét ({idx}/{total}): {base}")
                pbar.configure(value=max(0, idx - 1))
                self._set_status(f"Đang OCR ảnh/PDF lỗ khoan ({idx}/{total}): {base}...")
            except Exception:
                pass

        def worker():
            parsed_rows: List[List[Any]] = []
            parsed_file_count = 0
            warns: List[str] = []
            ocr_text_parts: List[str] = []
            total = len(paths)
            for idx, path in enumerate(paths, start=1):
                if cancel_flag.get("cancel"):
                    warns.append("Người dùng đã yêu cầu dừng import sau file trước đó.")
                    break
                base = os.path.basename(path)
                self.root.after(0, lambda i=idx, t=total, b=base: ui_progress(i, t, b))
                try:
                    self._borehole_ocr_crop_cache = {}
                    self._borehole_ocr_crop_data_cache = {}
                    self._borehole_line_group_cache = {}
                    item_name, rows, warn, ocr_text = self._make_geology_rows_from_borehole_file_fast(path)
                    if warn:
                        warns.append(f"{base}: {warn}")
                    if ocr_text:
                        ocr_text_parts.append(f"===== {base} =====\n{ocr_text}")
                    if rows:
                        parsed_rows.extend(rows)
                        parsed_file_count += 1
                    else:
                        warns.append(f"{base}: không đọc được dòng địa chất nào")
                except Exception as exc:
                    warns.append(f"{base}: {exc}")
            result = {
                "parsed_rows": parsed_rows,
                "parsed_file_count": parsed_file_count,
                "warns": warns,
                "ocr_text_parts": ocr_text_parts,
                "total": len(paths),
                "cancelled": bool(cancel_flag.get("cancel")),
            }
            self.root.after(0, lambda: self._finish_borehole_import_background(progress_dlg, pbar, progress_text, result))

        threading.Thread(target=worker, daemon=True).start()

    def _finish_borehole_import_background(self, progress_dlg: Any, pbar: Any, progress_text: Any, result: Dict[str, Any]):
        try:
            pbar.configure(value=result.get("total", 0))
            progress_text.set("Đã quét xong, đang mở preview...")
            progress_dlg.destroy()
        except Exception:
            pass
        parsed_rows = list(result.get("parsed_rows") or [])
        parsed_file_count = int(result.get("parsed_file_count") or 0)
        warns = list(result.get("warns") or [])
        ocr_text_parts = list(result.get("ocr_text_parts") or [])
        total = int(result.get("total") or 0)
        existing_rows = list(self.layer_table.get_rows() or [])
        if parsed_rows:
            combined_text = "\n\n".join(ocr_text_parts)
            title = f"Preview OCR - {parsed_file_count}/{total} file"
            edited = self._show_ocr_preview_dialog(title, parsed_rows, combined_text)
            if edited is None:
                self._set_status("Đã hủy nhập OCR lỗ khoan")
                return
            existing_rows.extend(edited)
            self.layer_table.set_rows(existing_rows, record_undo=True)
            self._merge_first_col_visual(self.layer_table)
            self.vars["common_geology"].set(False)
            self._set_status(f"Đã OCR và thêm {len(edited)} lớp địa chất từ {parsed_file_count} file")
            msg = (
                f"Đã quét xong {total} file.\n"
                f"Đã thêm {len(edited)} lớp địa chất từ {parsed_file_count} file đọc được.\n\n"
                "Hãy rà soát lại 3 cột Lớp, CĐ đáy và SPT trước khi tính toán."
            )
            if warns:
                msg += "\n\nCảnh báo tóm tắt:\n" + "\n".join(warns[:6])
                if len(warns) > 6:
                    msg += f"\n... và {len(warns) - 6} cảnh báo khác."
            messagebox.showinfo("OCR lỗ khoan", msg)
        else:
            self._set_status("OCR lỗ khoan không thành công")
            msg = "Không đọc được dòng địa chất nào từ các file đã chọn."
            if warns:
                msg += "\n\n" + "\n".join(warns[:8])
                if len(warns) > 8:
                    msg += f"\n... và {len(warns) - 8} cảnh báo khác."
            messagebox.showwarning("OCR lỗ khoan", msg)

    def import_geology_file(self):
        path = filedialog.askopenfilename(title="Import địa chất CSV/TXT", filetypes=[("CSV/TXT/Excel", "*.csv *.txt *.xlsx *.xlsm"), ("All files", "*.*")])
        if not path:
            return
        rows = self._read_delimited_file(path)
        if not rows:
            return
        header = [h.strip().lower() for h in rows[0]]
        data = rows[1:] if any(k in "|".join(header) for k in ["lop", "lớp", "bottom", "cao"]) else rows
        mapped = []
        for r in data:
            vals = (r + [""] * 14)[:14]
            mapped.append(vals)
        existing_rows = list(self.layer_table.get_rows() or [])
        self.layer_table.set_rows(existing_rows + mapped, record_undo=True)
        self._merge_first_col_visual(self.layer_table)
        self._set_status(f"Đã thêm {len(mapped)} lớp địa chất từ {os.path.basename(path)}")



    def _load_example_bored_soil(self):
        # Ví dụ lấy theo sheet CKN trong Dat trong sct.xlsx, đã làm gọn cột.
        vals = {
            "project": "km9+900", "item": "Trụ T7T", "mode": "Cọc khoan trong đất", "pile_count": "8",
            "spacing_m": "3.26", "ground_elev_m": "1.5", "cap_bottom_elev_m": "1.5", "pile_tip_elev_m": "-55.5", "water_elev_m": "1.5",
            "diameter_mm": "1200", "tip_diameter_mm": "1200", "concrete_gamma": "24.5", "fc_mpa": "30", "fy_mpa": "400",
            "n_rebars": "20", "rebar_dia": "28", "stirrup_type": "1", "exclude_top": "1.5", "spt_er": "60", "spt_input_mode": "Nₕₜ", "sand_mode": "sand", "sand_m": "0.6"
        }
        for k, v in vals.items():
            if k in self.vars:
                self.vars[k].set(v)
        self.item_table.set_rows([["Trụ T7T", "1", "8", "2", "1200", "1200", "3.26", "", "", "", "1.5", "1.5", "-55.5", "1.5", "", "", "", "", "Ví dụ"]])
        rows = [
            ["Trụ T7T", "1B", "0.38", "2", "1", "17.8", "0.006", "", "", "", "", "", "", ""],
            ["Trụ T7T", "1", "0", "2", "1", "17.3", "0.006", "", "", "", "", "", "", ""],
            ["Trụ T7T", "1", "-5.52", "2", "2", "17.3", "0.012", "", "", "", "", "", "", ""],
            ["Trụ T7T", "3", "-7.32", "2", "6", "19.5", "0.036", "", "", "", "", "", "", ""],
            ["Trụ T7T", "4a", "-9.22", "2", "10", "15.6", "0.060", "", "", "", "", "", "", ""],
            ["Trụ T7T", "5a", "-21.42", "1", "9", "18.92", "", "", "", "", "", "", "", ""],
            ["Trụ T7T", "9", "-25.02", "2", "11", "19.8", "0.066", "", "", "", "", "", "", ""],
            ["Trụ T7T", "13", "-51.02", "1", "23", "19.1", "", "", "", "", "", "", "", ""],
            ["Trụ T7T", "17", "-55.5", "1", "31", "19.23", "", "", "", "", "", "", "", ""],
        ]
        self.layer_table.set_rows(rows)
        self.vars["common_geology"].set(False)
        self._set_status("Đã nạp ví dụ cọc khoan trong đất")

    def _load_example_bored_rock(self):
        vals = {
            "project": "Ví dụ cọc khoan đá", "item": "Trụ mẫu", "mode": "Cọc khoan trong đá", "pile_count": "8",
            "spacing_m": "3.6", "ground_elev_m": "1.5", "cap_bottom_elev_m": "1.5", "pile_tip_elev_m": "-43.5", "water_elev_m": "1.5",
            "diameter_mm": "1200", "tip_diameter_mm": "1200", "concrete_gamma": "24.5", "fc_mpa": "30", "fy_mpa": "400",
            "n_rebars": "24", "rebar_dia": "28", "stirrup_type": "1", "exclude_top": "1.5", "spt_er": "60", "rock_construction_condition": "Có chống đỡ", "rock_joint_condition": "Khe nứt hở hoặc có mùn",
            "crack_spacing": "50", "crack_width": "5"
        }
        for k, v in vals.items():
            if k in self.vars:
                self.vars[k].set(v)
        self.item_table.set_rows([["Trụ mẫu", "1", "8", "2", "1200", "1200", "3.6", "", "", "", "1.5", "1.5", "-43.5", "1.5", "", "", "", "", "Ví dụ đá"]])
        rows = [
            ["Trụ mẫu", "1", "-6.72", "2", "7.5", "19", "0.045", "", "", "", "", "", "", ""],
            ["Trụ mẫu", "2", "-9.02", "2", "10.5", "19", "0.063", "", "", "", "", "", "", ""],
            ["Trụ mẫu", "10", "-27.22", "2", "7.3", "19", "0.0438", "", "", "", "", "", "", ""],
            ["Trụ mẫu", "11", "-35.32", "1", "24.2", "19", "", "", "", "", "", "", "", ""],
            ["Trụ mẫu", "18", "-43.5", "4", "24", "26", "", "", "30", "30", "10", "9", "0.5", "Đá nứt vỡ/phong hóa"],
        ]
        self.layer_table.set_rows(rows)
        self.vars["common_geology"].set(False)
        self._set_status("Đã nạp ví dụ cọc khoan trong đá")

    def run(self):
        self.root.mainloop()


def main():
    # QA-UX U1: chế độ không GUI cho kiểm thử tự động/tool phân tích code —
    # tránh cửa sổ tự bật ra và tránh treo tiến trình ở mainloop.
    if "--nogui" in sys.argv or str(os.environ.get("TS_CAP_NO_GUI", "")).strip() == "1":
        print("TS-CAP: chạy chế độ --nogui/TS_CAP_NO_GUI=1 — không mở giao diện.")
        return

    # QA-UX U2+: chỉ tạo MỘT Tk root. Root này ban đầu là splash, sau khi license OK
    # sẽ được tái dùng làm cửa sổ chính để tránh rủi ro tạo/destroy nhiều Tk root.
    root = tk.Tk()
    root.title(APP_NAME)
    root.resizable(False, False)
    try:
        root.protocol("WM_DELETE_WINDOW", lambda: None)  # không đóng splash giữa lúc thread license đang chạy
    except Exception:
        pass
    frm = ttk.Frame(root, padding=18)
    frm.pack(fill=tk.BOTH, expand=True)
    ttk.Label(frm, text=APP_NAME, font=("Arial", 14, "bold")).pack(anchor=tk.W)
    splash_status = tk.StringVar(value="Đang kiểm tra bản quyền và chuẩn bị môi trường...")
    ttk.Label(frm, textvariable=splash_status, wraplength=390).pack(anchor=tk.W, pady=(6, 10))
    pb = ttk.Progressbar(frm, mode="indeterminate", length=390)
    pb.pack(fill=tk.X)
    pb.start(12)
    try:
        apply_app_icon(root, "ts_cap")
    except Exception:
        pass
    try:
        root.update_idletasks()
        x = (root.winfo_screenwidth() - root.winfo_reqwidth()) // 2
        y = (root.winfo_screenheight() - root.winfo_reqheight()) // 3
        root.geometry(f"+{max(x, 0)}+{max(y, 0)}")
    except Exception:
        pass
    safe_lift_window(root)
    main_state = {"license_failed": False}

    def _show_blocked_and_exit():
        main_state["license_failed"] = True
        try:
            pb.stop()
        except Exception:
            pass
        try:
            root.withdraw()
        except Exception:
            pass
        _notify_license_blocked()
        try:
            root.destroy()
        except Exception:
            pass

    def _open_main_app(days_left: Any):
        try:
            pb.stop()
        except Exception:
            pass
        try:
            for child in list(root.winfo_children()):
                child.destroy()
        except Exception:
            pass
        try:
            root.resizable(True, True)
            root.protocol("WM_DELETE_WINDOW", root.destroy)
        except Exception:
            pass
        app = SCTApp(root=root)
        try:
            setattr(app, "license_days_left", days_left)
        except Exception:
            pass
        # giữ tham chiếu để không bị GC khi mainloop đang chạy
        root._ts_cap_app = app

    def _license_worker():
        try:
            ok, days = check_server_trial()
        except Exception:
            ok, days = False, 0
        try:
            root.after(0, lambda: _open_main_app(days) if ok else _show_blocked_and_exit())
        except Exception:
            pass

    threading.Thread(target=_license_worker, daemon=True).start()
    root.mainloop()
    if main_state.get("license_failed"):
        sys.exit()


if __name__ == "__main__":
    main()
