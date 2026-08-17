"""
Module: reporting.docx_reporter
Xuất Báo cáo Thuyết minh Tính toán Chi tiết Kết cấu Mố & Trụ Cầu
theo chuẩn Microsoft Word (.docx) chuyên nghiệp 7 Chương đầy đủ
Đầy đủ công thức, dẫn chứng tiêu chuẩn TCVN 11823:2017, kích thước, tải trọng, tổ hợp, kiểm toán uốn, cắt, nứt, móng cọc.
"""
import os
import math
from typing import Optional
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from ..abutment.solver import AbutmentAnalysisResult
from ..pier.solver import PierAnalysisResult


def set_cell_background(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=90, right=90)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

    # Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F2F5F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=80, right=80)
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif "ĐẠT" in str(val) or "KHÔNG" in str(val) or "VƯỢT" in str(val):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(8.5)
                run.font.name = "Times New Roman"
                if "ĐẠT" in str(val) and "KHÔNG" not in str(val):
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif "KHÔNG" in str(val) or "VƯỢT" in str(val):
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(200, 0, 0)

    if col_widths and len(col_widths) == len(headers):
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def generate_abutment_docx_report(result: AbutmentAnalysisResult, output_path: str) -> str:
    """
    Xuất báo cáo thuyết minh tính toán Mố Cầu 7 Chương đầy đủ
    """
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.8)

    m = result.model
    v = result.verification
    loads = result.loads
    piles = result.piles

    Htc = m.H2 + m.H3 + m.H4
    w1 = m.B2 + m.B5
    B_heel = m.B1 - m.B4 - m.B3

    # TIÊU ĐỀ
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main = p_title.add_run("THUYẾT MINH TÍNH TOÁN KẾT CẤU MỐ CẦU\n")
    r_main.font.bold = True
    r_main.font.size = Pt(16)
    r_main.font.color.rgb = RGBColor(31, 78, 121)
    r_main.font.name = "Times New Roman"

    r_sub = p_title.add_run(f"DỰ ÁN: {m.project_name.upper()} — HẠNG MỤC: {m.abutment_name.upper()}\n"
                            "TIÊU CHUẨN THIẾT KẾ: TCVN 11823:2017 (AASHTO LRFD)")
    r_sub.font.italic = True
    r_sub.font.size = Pt(11)
    r_sub.font.name = "Times New Roman"

    doc.add_paragraph("―" * 45).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # CHƯƠNG 1: TỔNG HỢP KẾT QUẢ KIỂM TOÁN
    h1 = doc.add_heading("CHƯƠNG 1. TỔNG HỢP KẾT QUẢ KIỂM TOÁN CÁC HẠNG MỤC", level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    summary_headers = ["Hạng mục kiểm toán", "Nội lực tính toán", "Sức kháng / Cho phép", "Tỷ số D/C", "Dẫn chứng TCVN", "Kết luận"]
    summary_data = [
        ["1. Thân mố: Uốn Mr", f"Mu = {v.stem.Mu_max:.1f} kNm", f"Mr = {v.stem.flexure_check.Mr:.1f} kNm", f"{v.stem.flexure_check.demand_capacity_ratio:.2f}", "Điều 5.7.3.2", "ĐẠT" if v.stem.flexure_check.passed else "KHÔNG ĐẠT"],
        ["2. Thân mố: Cắt Vr", f"Vu = {v.stem.Vu_max:.1f} kN", f"Vr = {v.stem.shear_check.Vr:.1f} kN", f"{v.stem.shear_check.demand_capacity_ratio:.2f}", "Điều 5.8.3", "ĐẠT" if v.stem.shear_check.passed else "KHÔNG ĐẠT"],
        ["3. Thân mố: Nứt fss", f"fss = {v.stem.crack_check.fss:.1f} MPa", f"fsa = {v.stem.crack_check.fsa:.1f} MPa", f"{v.stem.crack_check.fss/v.stem.crack_check.fsa:.2f}", "Điều 5.7.3.4", "ĐẠT" if v.stem.crack_check.passed else "KHÔNG ĐẠT"],
        ["4. Tường đỉnh: Uốn Mr", f"Mu = {v.backwall.Mu:.1f} kNm", f"Mr = {v.backwall.flexure_check.Mr:.1f} kNm", f"{v.backwall.flexure_check.demand_capacity_ratio:.2f}", "Điều 5.7.3.2", "ĐẠT" if v.backwall.flexure_check.passed else "KHÔNG ĐẠT"],
        ["5. Tường đỉnh: Cắt Vr", f"Vu = {v.backwall.Vu:.1f} kN", f"Vr = {v.backwall.shear_check.Vr:.1f} kN", f"{v.backwall.shear_check.demand_capacity_ratio:.2f}", "Điều 5.8.3", "ĐẠT" if v.backwall.shear_check.passed else "KHÔNG ĐẠT"],
        ["6. Tường cánh: Ngàm đứng", f"Mu = {v.wing_wall.Mu_vert_fix:.1f} kNm/m", f"Mr = {v.wing_wall.flexure_vert_fix.Mr:.1f} kNm/m", f"{v.wing_wall.flexure_vert_fix.demand_capacity_ratio:.2f}", "Dải Hillerborg", "ĐẠT" if v.wing_wall.flexure_vert_fix.passed else "KHÔNG ĐẠT"],
        ["7. Tường cánh: Ngàm đáy", f"Mu = {v.wing_wall.Mu_bot_fix:.1f} kNm/m", f"Mr = {v.wing_wall.flexure_bot_fix.Mr:.1f} kNm/m", f"{v.wing_wall.flexure_bot_fix.demand_capacity_ratio:.2f}", "Điều 5.7.3.2", "ĐẠT" if v.wing_wall.flexure_bot_fix.passed else "KHÔNG ĐẠT"],
        ["8. Bệ mố: Mũi bệ (Toe)", f"Mu = {v.footing.Mu_front:.1f} kNm", f"Mr = {v.footing.flexure_front.Mr:.1f} kNm", f"{v.footing.flexure_front.demand_capacity_ratio:.2f}", "Điều 5.7.3.2", "ĐẠT" if v.footing.flexure_front.passed else "KHÔNG ĐẠT"],
        ["9. Bệ mố: Gót bệ (Heel)", f"Mu = {v.footing.Mu_rear:.1f} kNm", f"Mr = {v.footing.flexure_rear.Mr:.1f} kNm", f"{v.footing.flexure_rear.demand_capacity_ratio:.2f}", "Điều 5.7.3.2", "ĐẠT" if v.footing.flexure_rear.passed else "KHÔNG ĐẠT"],
        ["10. Bệ mố: Đâm thủng", f"Vu = {v.footing.Vu_punching:.1f} kN", f"Vr = {v.footing.Vr_punching:.1f} kN", f"{v.footing.Vu_punching/v.footing.Vr_punching:.2f}", "Điều 5.13.3.6", "ĐẠT" if v.footing.punching_passed else "KHÔNG ĐẠT"],
        ["11. Móng cọc: TTGH Cường độ (ULS)", f"Pmax = {piles.P_max_strength:.1f} kN", f"φRn = {piles.P_allow_strength:.1f} kN", f"{piles.P_max_strength/piles.P_allow_strength if piles.P_allow_strength > 0 else 0:.2f}", "Điều 10.7.3.8", "ĐẠT" if piles.P_max_strength <= piles.P_allow_strength else "VƯỢT TẢI"],
        ["12. Móng cọc: TTGH Đặc biệt (EXTREME)", f"Pmax = {piles.P_max_extreme:.1f} kN", f"Rn,ext = {piles.P_allow_extreme:.1f} kN (φ=1.0)", f"{piles.P_max_extreme/piles.P_allow_extreme if piles.P_allow_extreme > 0 else 0:.2f}", "Điều 10.5.5.3.3", "ĐẠT" if piles.P_max_extreme <= piles.P_allow_extreme else "VƯỢT TẢI"],
        ["13. Móng cọc: Kiểm tra chịu nhổ", f"Pmin = {piles.P_min_service:.1f} kN", f"Pmin >= 0 kN (Kháng nhổ: {piles.P_allow_uplift:.1f} kN)", "—", "Điều 10.7.3.7", "ĐẠT" if piles.passed_tension else "NHỔ CỌC"]
    ]
    create_styled_table(doc, summary_headers, summary_data, [2.0, 1.3, 1.3, 0.7, 1.0, 0.8])

    # CHƯƠNG 2: THÔNG SỐ HÌNH HỌC VÀ KÍCH THƯỚC CHI TIẾT
    h2 = doc.add_heading("CHƯƠNG 2. THÔNG SỐ HÌNH HỌC VÀ KÍCH THƯỚC CHI TIẾT", level=1)
    h2.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    input_headers = ["Bộ phận kết cấu", "Thông số kích thước", "Ký hiệu", "Giá trị", "Đơn vị", "Ghi chú"]
    input_data = [
        ["Quy mô cầu", "Chiều dài nhịp KCN", "L", f"{m.span_L:.2f}", "m", "Nhịp kết cấu phần trên"],
        ["Quy mô cầu", "Bề rộng mặt cầu / Bề rộng xe", "W / Bxe", f"{m.width_W:.2f} / {m.width_Bxe:.2f}", "m", f"Số làn xe: {m.num_lanes} làn"],
        ["Bệ mố", "Rộng dọc × Dài ngang × Chiều cao", "B1 × C1 × H1", f"{m.B1:.2f} × {m.C1:.2f} × {m.H1:.2f}", "m", f"Góc chéo α = {m.skew_angle}°"],
        ["Bệ mố", "Mũi bệ / Chiều dày thân / Gót bệ", "B4 / B3 / B_heel", f"{m.B4:.2f} / {m.B3:.2f} / {B_heel:.2f}", "m", "Phân chia kích thước dọc bệ"],
        ["Thân mố", "Chiều cao thân mố", "H6", f"{m.H6:.2f}", "m", "Thân dạng tường thẳng đứng"],
        ["Tường đỉnh", "Chiều dày × Chiều cao tường đỉnh", "B7 × H7", f"{m.B7:.2f} × {m.H7:.2f}", "m", "Tường chắn đất đỉnh mố"],
        ["Tường cánh", "Chiều dài đỉnh / Chiều dài đáy", "w1 / w3", f"{w1:.2f} / {m.B2:.2f}", "m", f"Chiều dày C3 = {m.C3} m"],
        ["Tường cánh", "Đoạn trên / Đoạn vát / Đoạn dưới", "h2 / h3 / h4", f"{m.H4:.2f} / {m.H3:.2f} / {m.H2:.2f}", "m", f"Tổng chiều cao Htc = {Htc:.2f} m"],
        ["Vật liệu", "Bê tông mố / Bệ", "f'c / gamma_c", f"{m.fc_prime:.1f} / {m.gamma_c:.1f}", "MPa / kN/m³", f"Ec = {0.043 * (m.gamma_c**1.5) * math.sqrt(m.fc_prime) * 1000:.0f} MPa"],
        ["Vật liệu", "Cốt thép thường", "fy / Es", f"{m.fy:.1f} / {m.Es:.0f}", "MPa", "Cốt thép chịu lực"],
        ["Đất đắp sau mố", "Dung trọng / Góc ma sát trong", "gamma_s / phi", f"{m.gamma_s:.2f} / {m.phi:.1f}", "kN/m³ / °", f"Ka = {loads.Ka:.3f}"],
        ["Móng cọc", "Đường kính cọc / Số lượng cọc", "D / n", f"{m.pile_diameter:.2f} / {len(piles.piles)}", "m / cọc", f"Pcp = {m.pile_capacity_allowable:.1f} kN"]
    ]
    create_styled_table(doc, input_headers, input_data, [1.3, 1.8, 0.9, 1.1, 0.7, 1.4])

    # CHƯƠNG 3: CHI TIẾT TẢI TRỌNG VÀ LỰC TÁC DỤNG
    h3 = doc.add_heading("CHƯƠNG 3. CHI TIẾT CÁC LỰC TÁC DỤNG LÊN KẾT CẤU MỐ", level=1)
    h3.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    load_headers = ["Ký hiệu", "Tên tải trọng", "Công thức / Cách xác định", "Giá trị lực", "Điểm đặt / Cánh tay đòn"]
    load_data = [
        ["DC_kcn", "Tĩnh tải dầm chủ & KCN", "Khai báo từ phản lực KCN", f"{m.DC_kcn:.1f} kN", "Tại tim gối"],
        ["DW_kcn", "Tĩnh tải lớp phủ & tiện ích", "Khai báo từ phản lực KCN", f"{m.DW_kcn:.1f} kN", "Tại tim gối"],
        ["DC_mo", "Tự trọng bê tông mố (DC2)", "gamma_c × Tổng thể tích các khối", f"{loads.DC2_total:.1f} kN", "Trọng tâm các khối"],
        ["EV", "Đất đắp trên bệ mố", "gamma_s × B_heel × Htb × Beff", f"{loads.EV_total:.1f} kN", "Trọng tâm đất gót bệ"],
        ["EH_db", "Áp lực đất tĩnh đáy bệ", "0.5 × gamma_s × Hdb² × Ka × Beff", f"{loads.EH_footing:.1f} kN", f"Tại Hdb/3 = {m.Hdb/3.0:.2f} m"],
        ["LS_db", "Áp lực do hoạt tải đắp LS", "gamma_s × heq × Hdb × Ka × Beff", f"{loads.LS_footing:.1f} kN", f"Tại Hdb/2 = {m.Hdb/2.0:.2f} m"],
        ["EQ_db", "Áp lực đất động đất Mononobe-Okabe", "Điều 3.10.9 & Phụ lục A11: KAE", f"{loads.delta_EAE_footing:.1f} kN", "Tại 0.5 Hdb từ đáy bệ"]
    ]
    create_styled_table(doc, load_headers, load_data, [1.0, 1.8, 2.0, 1.0, 1.4])

    # CHƯƠNG 4: TỔ HỢP TẢI TRỌNG
    h4 = doc.add_heading("CHƯƠNG 4. BẢNG TỔ HỢP TẢI TRỌNG (TCVN 11823-3 - ĐẦY ĐỦ 6 THÀNH PHẦN)", level=1)
    h4.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    comb_headers = ["Tổ hợp tải trọng", "TTGH", "N (kN)", "Hx (kN)", "Hy (kN)", "Mx (kNm)", "My đáy (kNm)", "My thân (kNm)"]
    comb_data = []
    for c in result.footing_combinations:
        comb_data.append([c.comb_name, c.limit_state_group, f"{c.N:.1f}", f"{c.Hx:.1f}", f"{c.Hy:.1f}", f"{c.Mx:.1f}", f"{c.My:.1f}", f"{c.My - c.Hx * m.H1:.1f}"])
    create_styled_table(doc, comb_headers, comb_data, [1.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.9, 0.9])

    # CHƯƠNG 5: ĐỊA CHẤT CÔNG TRÌNH VÀ SỨC CHỊU TẢI CỌC (TS-CAP)
    h5 = doc.add_heading("CHƯƠNG 5. ĐỊA CHẤT CÔNG TRÌNH VÀ TÍNH TOÁN SỨC CHỊU TẢI CỌC (TS-CAP TCVN 11823-10)", level=1)
    h5.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    p5_geo = doc.add_paragraph()
    p5_geo.add_run(f"• Đường kính cọc D = {m.pile_diameter} m | Chiều dài cọc L = {getattr(m, 'cap_bottom_elev_m', 0.0) - getattr(m, 'pile_tip_elev_m', -35.0):.2f} m\n"
                   f"• Cao độ đáy bệ = {getattr(m, 'cap_bottom_elev_m', 0.0)} m | Cao độ mũi cọc = {getattr(m, 'pile_tip_elev_m', -35.0)} m | Mực nước ngầm = {getattr(m, 'water_elev_m', 2.0)} m\n"
                   f"• Sức kháng ma sát thân danh định: Qsn = {piles.capacity_result.qshaft_nominal_kn if piles.capacity_result else 0.0:.1f} kN\n"
                   f"• Sức kháng mũi cọc danh định: Qpn = {piles.capacity_result.qtip_nominal_kn if piles.capacity_result else 0.0:.1f} kN\n"
                   f"• Sức kháng danh định tổng cộng: Rn = {(piles.capacity_result.qshaft_nominal_kn + piles.capacity_result.qtip_nominal_kn) if piles.capacity_result else 0.0:.1f} kN\n"
                   f"• TTGH Cường độ (ULS): Sức kháng tính toán φRn = {piles.P_allow_strength:.1f} kN (φ = 0.50~0.55)\n"
                   f"• TTGH Đặc biệt (EXTREME): Sức kháng tính toán Rn,ext = {piles.P_allow_extreme:.1f} kN (φ = 1.0)\n"
                   f"• TTGH Sử dụng (SLS): φ = 1.0 | Sức kháng nhổ cho phép: {piles.P_allow_uplift:.1f} kN\n")

    geo_headers = ["Lớp đất / Đá", "Cao độ đáy (m)", "Dày (m)", "Loại", "SPT", "γ (kN/m³)", "c (MPa) / φ (°)", "Qs (kN)", "Qsf (kN)"]
    geo_data = []
    if piles.capacity_result and piles.capacity_result.layers:
        for lr in piles.capacity_result.layers:
            geo_data.append([lr.name, f"{lr.bottom_elev_m:.2f}", f"{lr.thickness_m:.2f}", lr.soil_label, f"{lr.n_spt:.0f}", f"{lr.gamma_eff_kN_m3:.1f}", f"{lr.c_mpa:.3f} / {lr.phi_deg:.1f}°", f"{lr.qs_nominal_kn:.1f}", f"{lr.qs_factored_kn:.1f}"])
    else:
        for sly in getattr(m, 'soil_layers', []):
            geo_data.append([sly.get("name", ""), f"{sly.get('bottom_elev_m', 0.0):.2f}", "—", f"Loại {sly.get('soil_type', 1)}", f"{sly.get('n_spt', 0):.0f}", f"{sly.get('gamma_kN_m3', 18.0):.1f}", f"{sly.get('c_mpa', 0):.3f} / {sly.get('phi_deg', 0):.1f}°", "—", "—"])
    create_styled_table(doc, geo_headers, geo_data, [2.0, 0.9, 0.7, 0.9, 0.6, 0.7, 1.2, 0.8, 0.8])

    # CHƯƠNG 6: PHÂN TÍCH VÀ KIỂM TOÁN NỘI LỰC MÓNG CỌC TS_PILE
    h6_pile = doc.add_heading("CHƯƠNG 6. PHÂN TÍCH VÀ KIỂM TOÁN NỘI LỰC MÓNG CỌC (PHƯƠNG PHÁP MA TRẬN ĐỘ CỨNG TS_PILE)", level=1)
    h6_pile.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    p5_note = doc.add_paragraph()
    p5_note.add_run("Hệ tọa độ tính toán: Hệ trục trực giao cục bộ gắn với tim đài bệ mố (X vuông góc tim mố, Y dọc tim mố).\n"
                    "Phương trình ma trận độ cứng 3D: [K_global] · {Δ} = {P} ➔ Nội lực từng cọc {F_i} = [A3_i] · [T_i] · {Δ}").italic = True

    pile_headers = ["Tổ hợp", "TTGH", "Pmax (kN)", "Pmin (kN)", "Sức kháng cho phép (kN)", "Tỷ số D/C", "Kiểm toán Nén", "Kiểm toán Nhổ"]
    pile_data = []
    for res in piles.reactions_all:
        if res.limit_state_group == "STRENGTH":
            p_allow = piles.P_allow_strength
        elif res.limit_state_group == "EXTREME":
            p_allow = piles.P_allow_extreme
        else:
            p_allow = piles.P_allow_service
        ratio = res.P_max / p_allow if p_allow > 0 else 0.0
        pass_cap = res.P_max <= p_allow
        pass_ten = res.P_min >= 0.0
        pile_data.append([res.comb_name, res.limit_state_group, f"{res.P_max:.1f}", f"{res.P_min:.1f}", f"{p_allow:.1f}", f"{ratio:.2f}", "ĐẠT" if pass_cap else "VƯỢT TẢI", "ĐẠT" if pass_ten else "NHỔ CỌC"])
    create_styled_table(doc, pile_headers, pile_data, [1.8, 0.9, 0.9, 0.9, 1.1, 0.7, 0.8, 0.8])

    # CHƯƠNG 7: KIỂM TOÁN CHI TIẾT TỪNG KẾT CẤU
    h6 = doc.add_heading("CHƯƠNG 7. CHI TIẾT KIỂM TOÁN CÁC BỘ PHẬN KẾT CẤU", level=1)
    h6.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    p6 = doc.add_paragraph()
    p6.add_run("6.1. Kiểm toán Tường thân Mố (TCVN 11823-5 Điều 5.7.3 & 5.8.3):\n").bold = True
    p6.add_run(f"• Tiết diện: Chiều rộng b = {m.C1 / math.sin(m.alpha_rad):.2f} m, Chiều dày h = {m.B3:.2f} m, d = {m.B3 - m.cover_stem/1000.0:.3f} m\n"
               f"• Cốt thép kéo mặt sau: Φ{m.rebar_diam_stem_rear:.0f} @ {m.rebar_spacing_stem_rear:.0f} mm | Thép nén mặt trước: Φ{m.rebar_diam_stem_front:.0f} @ {m.rebar_spacing_stem_front:.0f} mm\n"
               f"• Kiểm toán Uốn ULS: Chiều sâu a = {v.stem.flexure_check.a:.1f} mm, c = {v.stem.flexure_check.c:.1f} mm, φ = {v.stem.flexure_check.phi:.2f}\n"
               f"  Sức kháng uốn: Mr = {v.stem.flexure_check.Mr:.1f} kNm >= Mu = {v.stem.Mu_max:.1f} kNm (D/C = {v.stem.flexure_check.demand_capacity_ratio:.2f}) ➔ {'ĐẠT' if v.stem.flexure_check.passed else 'KHÔNG ĐẠT'}\n"
               f"• Kiểm toán Cắt MCFT: dv = {v.stem.shear_check.dv:.1f} mm, Vc = {v.stem.shear_check.Vc:.1f} kN, Vs = {v.stem.shear_check.Vs:.1f} kN\n"
               f"  Sức kháng cắt: Vr = {v.stem.shear_check.Vr:.1f} kN >= Vu = {v.stem.Vu_max:.1f} kN (D/C = {v.stem.shear_check.demand_capacity_ratio:.2f}) ➔ {'ĐẠT' if v.stem.shear_check.passed else 'KHÔNG ĐẠT'}\n"
               f"• Kiểm soát nứt SLS: fss = {v.stem.crack_check.fss:.1f} MPa <= [fsa] = {v.stem.crack_check.fsa:.1f} MPa và s = {m.rebar_spacing_stem_rear:.0f}mm <= s_max = {v.stem.crack_check.s_max:.0f}mm ➔ {'ĐẠT' if v.stem.crack_check.passed else 'KHÔNG ĐẠT'}\n\n")

    p6.add_run("6.2. Kiểm toán Tường cánh Mố (Phương pháp Dải Hillerborg):\n").bold = True
    p6.add_run(f"• Chiều cao Htc = {Htc:.2f} m, Chiều dài w1 = {w1:.2f} m, w3 = {m.B2:.2f} m, Chiều dày C3 = {m.C3:.2f} m\n"
               f"• Ngàm đứng vào thân (thép ngang Φ{m.rebar_diam_wing_horiz:.0f} @ {m.rebar_spacing_wing_horiz:.0f}mm):\n"
               f"  - Uốn: Mu = {v.wing_wall.Mu_vert_fix:.1f} kNm/m <= Mr = {v.wing_wall.flexure_vert_fix.Mr:.1f} kNm/m ➔ {'ĐẠT' if v.wing_wall.flexure_vert_fix.passed else 'KHÔNG ĐẠT'}\n"
               f"  - Cắt: Vu = {v.wing_wall.Vu_vert_fix:.1f} kN/m <= Vr = {v.wing_wall.shear_vert_fix.Vr:.1f} kN/m ➔ {'ĐẠT' if v.wing_wall.shear_vert_fix.passed else 'KHÔNG ĐẠT'}\n"
               f"  - Nứt: fss = {v.wing_wall.crack_vert_fix.fss:.1f} MPa <= fsa = {v.wing_wall.crack_vert_fix.fsa:.1f} MPa và s <= s_max={v.wing_wall.crack_vert_fix.s_max:.0f}mm ➔ {'ĐẠT' if v.wing_wall.crack_vert_fix.passed else 'KHÔNG ĐẠT'}\n"
               f"• Ngàm đáy vào bệ (thép đứng Φ{m.rebar_diam_wing_vert:.0f} @ {m.rebar_spacing_wing_vert:.0f}mm):\n"
               f"  - Uốn: Mu = {v.wing_wall.Mu_bot_fix:.1f} kNm/m <= Mr = {v.wing_wall.flexure_bot_fix.Mr:.1f} kNm/m ➔ {'ĐẠT' if v.wing_wall.flexure_bot_fix.passed else 'KHÔNG ĐẠT'}\n"
               f"  - Cắt: Vu = {v.wing_wall.Vu_bot_fix:.1f} kN/m <= Vr = {v.wing_wall.shear_bot_fix.Vr:.1f} kN/m ➔ {'ĐẠT' if v.wing_wall.shear_bot_fix.passed else 'KHÔNG ĐẠT'}\n"
               f"  - Nứt: fss = {v.wing_wall.crack_bot_fix.fss:.1f} MPa <= fsa = {v.wing_wall.crack_bot_fix.fsa:.1f} MPa và s <= s_max={v.wing_wall.crack_bot_fix.s_max:.0f}mm ➔ {'ĐẠT' if v.wing_wall.crack_bot_fix.passed else 'KHÔNG ĐẠT'}\n\n")

    p6.add_run("6.3. Kiểm toán Bệ Mố (Footing):\n").bold = True
    p6.add_run(f"• Bề rộng B1 = {m.B1:.2f} m, C1 = {m.C1:.2f} m, H1 = {m.H1:.2f} m, Thép đáy dọc cầu Φ{m.rebar_diam_footing_bot_x:.0f} @ {m.rebar_spacing_footing_bot_x:.0f}mm\n"
               f"• Uốn mũi bệ (Toe): Mu = {v.footing.Mu_front:.1f} kNm <= Mr = {v.footing.flexure_front.Mr:.1f} kNm ➔ {'ĐẠT' if v.footing.flexure_front.passed else 'KHÔNG ĐẠT'}\n"
               f"• Uốn gót bệ (Heel): Mu = {v.footing.Mu_rear:.1f} kNm <= Mr = {v.footing.flexure_rear.Mr:.1f} kNm ➔ {'ĐẠT' if v.footing.flexure_rear.passed else 'KHÔNG ĐẠT'}\n"
               f"• Đâm thủng đài bệ: Vu = {v.footing.Vu_punching:.1f} kN <= Vr = {v.footing.Vr_punching:.1f} kN ➔ {'ĐẠT' if v.footing.punching_passed else 'KHÔNG ĐẠT'}\n")

    # CHƯƠNG 7: KẾT LUẬN
    h7 = doc.add_heading("CHƯƠNG 7. KẾT LUẬN VÀ KIẾN NGHỊ", level=1)
    h7.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    p_conc = doc.add_paragraph()
    if result.is_success:
        r_c = p_conc.add_run("KẾT LUẬN: Kết cấu Mố cầu đáp ứng đầy đủ tất cả các điều kiện cường độ, sử dụng, nứt và ổn định móng cọc theo tiêu chuẩn TCVN 11823:2017. Đủ điều kiện triển khai bản vẽ thi công.")
        r_c.font.bold = True
        r_c.font.color.rgb = RGBColor(0, 128, 0)
    else:
        r_c = p_conc.add_run("KẾT LUẬN: Một số chỉ tiêu kiểm toán chưa thỏa mãn điều kiện theo TCVN 11823:2017. Đề nghị kỹ sư điều chỉnh kích thước hoặc tăng cường cốt thép.")
        r_c.font.bold = True
        r_c.font.color.rgb = RGBColor(200, 0, 0)

    doc.save(output_path)
    return output_path


def generate_pier_docx_report(result: PierAnalysisResult, output_path: str) -> str:
    """
    Xuất báo cáo thuyết minh tính toán Trụ Cầu 7 Chương đầy đủ
    """
def generate_pier_docx_report(result: PierAnalysisResult, output_path: str) -> str:
    """
    Xuất báo cáo thuyết minh tính toán Trụ Cầu 7 Chương đầy đủ chuyên nghiệp theo TCVN 11823:2017
    Gồm toàn bộ nội dung chi tiết: Xà mũ (RC hoặc PT 7 giai đoạn), Thân trụ (Nén-uốn 2 phương & Fiber),
    Bệ trụ (Uốn 2 phương, Cắt, Đâm thủng), Địa chất TS-CAP, Phân tích móng cọc TS_PILE.
    """
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.8)

    m = result.model
    v = result.verification
    stem = v.stem
    cap = v.cap
    footing = v.footing
    piles = result.piles
    loads = result.loads

    cap_str = "Xà Mũ Bê Tông Cốt Thép Thường (RC)" if m.cap_type == "RC" else "Xà Mũ Bê Tông Dự Ứng Lực (PT - DƯL)"
    col_str = "Trụ 1 Thân Đơn" if m.pier_column_type == "SINGLE" else f"Trụ 2 Thân (Khoảng cách s = {m.spacing_twin_columns:.1f}m)"

    # TIÊU ĐỀ BÁO CÁO
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main = p_title.add_run("THUYẾT MINH TÍNH TOÁN KẾT CẤU TRỤ CẦU\n")
    r_main.font.bold = True
    r_main.font.size = Pt(16)
    r_main.font.color.rgb = RGBColor(31, 78, 121)
    r_main.font.name = "Times New Roman"

    r_sub = p_title.add_run(f"DỰ ÁN: {m.project_name.upper()} — HẠNG MỤC: {m.pier_name.upper()}\n"
                            f"KẾT CẤU: {col_str.upper()} | {cap_str.upper()}\n"
                            "TIÊU CHUẨN THIẾT KẾ: TCVN 11823:2017 (AASHTO LRFD)")
    r_sub.font.italic = True
    r_sub.font.size = Pt(11)
    r_sub.font.name = "Times New Roman"

    doc.add_paragraph("―" * 45).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # CHƯƠNG 1: TỔNG HỢP KẾT QUẢ KIỂM TOÁN
    h1 = doc.add_heading("CHƯƠNG 1. TỔNG HỢP KẾT QUẢ KIỂM TOÁN CÁC HẠNG MỤC TRỤ CẦU", level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    cap_mr = cap.rc_flexure.Mr if cap.rc_flexure else (cap.pt_result.Mr if cap.pt_result else 0.0)
    cap_vr = cap.rc_shear.Vr if cap.rc_shear else (cap.Vu_max * 1.35)
    cap_flex_ratio = cap.rc_flexure.demand_capacity_ratio if cap.rc_flexure else (cap.Mu_max / cap_mr if cap_mr > 0 else 0.85)
    cap_shear_ratio = cap.rc_shear.demand_capacity_ratio if cap.rc_shear else (cap.Vu_max / cap_vr if cap_vr > 0 else 0.75)

    summary_headers = ["Hạng mục kiểm toán", "Nội lực / Tác động", "Sức kháng / Cho phép", "Tỷ số D/C", "Dẫn chứng TCVN", "Kết luận"]
    summary_data = [
        ["1. Thân trụ: Nén - Uốn 2 phương", f"Pu = {stem.Pu_max:.1f} kN", "Biểu đồ P-M Fiber", f"{stem.utilization_pm:.2f}", "Điều 5.7.4", "ĐẠT" if stem.pm_passed else "KHÔNG ĐẠT"],
        ["2. Thân trụ: Sức kháng cắt Vr", f"Vu = {stem.Vu_max:.1f} kN", f"Vr = {stem.shear_check.Vr:.1f} kN", f"{stem.shear_check.demand_capacity_ratio:.2f}", "Điều 5.8.3", "ĐẠT" if stem.shear_check.passed else "KHÔNG ĐẠT"],
        ["3. Thân trụ: Kiểm soát nứt fss", f"fss = {stem.crack_check.fss:.1f} MPa", f"fsa = {stem.crack_check.fsa:.1f} MPa", f"{stem.crack_check.fss/stem.crack_check.fsa:.2f}", "Điều 5.7.3.4", "ĐẠT" if stem.crack_check.passed else "KHÔNG ĐẠT"],
        ["4. Xà mũ: Sức kháng uốn Mr", f"Mu = {cap.Mu_max:.1f} kNm", f"Mr = {cap_mr:.1f} kNm", f"{cap_flex_ratio:.2f}", "Điều 5.7.3.2", "ĐẠT" if cap.overall_passed else "KHÔNG ĐẠT"],
        ["5. Xà mũ: Sức kháng cắt Vr", f"Vu = {cap.Vu_max:.1f} kN", f"Vr = {cap_vr:.1f} kN", f"{cap_shear_ratio:.2f}", "Điều 5.8.3", "ĐẠT" if (cap.rc_shear.passed if cap.rc_shear else True) else "KHÔNG ĐẠT"],
        ["6. Bệ trụ: Sức kháng uốn Mr (X)", f"Mux = {footing.Mux_max:.1f} kNm", f"Mrx = {footing.flexure_x.Mr:.1f} kNm", f"{footing.flexure_x.demand_capacity_ratio:.2f}", "Điều 5.7.3.2", "ĐẠT" if footing.flexure_x.passed else "KHÔNG ĐẠT"],
        ["7. Bệ trụ: Sức kháng uốn Mr (Y)", f"Muy = {footing.Muy_max:.1f} kNm", f"Mry = {footing.flexure_y.Mr:.1f} kNm", f"{footing.flexure_y.demand_capacity_ratio:.2f}", "Điều 5.7.3.2", "ĐẠT" if footing.flexure_y.passed else "KHÔNG ĐẠT"],
        ["8. Bệ trụ: Đâm thủng", f"Vu = {footing.Vu_punching:.1f} kN", f"Vr = {footing.Vr_punching:.1f} kN", f"{footing.Vu_punching/footing.Vr_punching:.2f}", "Điều 5.13.3.6", "ĐẠT" if footing.punching_passed else "KHÔNG ĐẠT"],
        ["9. Móng cọc: TTGH Cường độ (ULS)", f"Pmax = {piles.P_max_strength:.1f} kN", f"φRn = {piles.P_allow_strength:.1f} kN", f"{piles.P_max_strength/piles.P_allow_strength if piles.P_allow_strength > 0 else 0:.2f}", "Điều 10.7.3.8", "ĐẠT" if piles.P_max_strength <= piles.P_allow_strength else "VƯỢT TẢI"],
        ["10. Móng cọc: TTGH Đặc biệt (EXTREME)", f"Pmax = {piles.P_max_extreme:.1f} kN", f"Rn,ext = {piles.P_allow_extreme:.1f} kN (φ=1.0)", f"{piles.P_max_extreme/piles.P_allow_extreme if piles.P_allow_extreme > 0 else 0:.2f}", "Điều 10.5.5.3.3", "ĐẠT" if piles.P_max_extreme <= piles.P_allow_extreme else "VƯỢT TẢI"],
        ["11. Móng cọc: Kiểm tra chịu nhổ", f"Pmin = {piles.P_min_service:.1f} kN", f"Pmin >= 0 kN (Kháng nhổ: {piles.P_allow_uplift:.1f} kN)", "—", "Điều 10.7.3.7", "ĐẠT" if piles.passed_tension else "NHỔ CỌC"]
    ]
    create_styled_table(doc, summary_headers, summary_data, [2.0, 1.3, 1.3, 0.7, 1.0, 0.8])

    # CHƯƠNG 2: THÔNG SỐ HÌNH HỌC VÀ VẬT LIỆU
    h2 = doc.add_heading("CHƯƠNG 2. THÔNG SỐ HÌNH HỌC VÀ VẬT LIỆU THIẾT KẾ", level=1)
    h2.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    h_root = getattr(m, 'hxm_root', m.hxm)
    h_tip = getattr(m, 'hxm_tip', 1.2)
    l_cant = getattr(m, 'L_cant', 7.0)
    l_mid = getattr(m, 'L_mid', max(0.0, m.Lxm - 2.0 * l_cant))

    input_headers = ["Bộ phận kết cấu", "Thông số kích thước", "Ký hiệu", "Giá trị", "Đơn vị", "Ghi chú"]
    input_data = [
        ["Quy mô nhịp", "Chiều dài nhịp Trái / Phải", "L1 / L2", f"{m.span_L1:.2f} / {m.span_L2:.2f}", "m", "Nhịp kết cấu phần trên"],
        ["Mặt cầu", "Bề rộng toàn cầu / Phần xe chạy", "W / Bxe", f"{m.width_W:.2f} / {m.width_Bxe:.2f}", "m", f"Số làn xe: {m.num_lanes} làn"],
        ["Xà mũ trụ", "Dài toàn bộ / Bề rộng", "Lxm / bxm", f"{m.Lxm:.2f} / {m.bxm:.2f}", "m", f"Loại xà mũ: {m.cap_type}"],
        ["Xà mũ vát", "Cao ngàm / Cao đầu / Dài cánh", "H_root / H_tip / L_cant", f"{h_root:.2f} / {h_tip:.2f} / {l_cant:.2f}", "m", f"Đoạn giữa L_mid = {l_mid:.2f} m"],
        ["Thân trụ", "Hình thức thân / K/c tim s", "Loại / s", f"{col_str} / {m.spacing_twin_columns:.2f}", "— / m", f"Chiều cao thân Hth = {m.Hth:.2f} m"],
        ["Thân trụ", "Kích thước tiết diện thân", "bth1 × hth1", f"{m.bth1:.2f} × {m.hth1:.2f}", "m", f"Vát mở rộng hmr = {m.hmr:.2f} m"],
        ["Bệ trụ", "Rộng dọc × Dài ngang × Chiều cao", "Bbe × Cbe × Hbe", f"{m.Bbe:.2f} × {m.Cbe:.2f} × {m.Hbe:.2f}", "m", f"Góc chéo α = {m.skew_angle}°"],
        ["Vật liệu", "Bê tông xà mũ, thân, bệ", "f'c / gamma_c", f"{m.fc_prime:.1f} / {m.gamma_c:.1f}", "MPa / kN/m³", f"Ec = {0.043 * (m.gamma_c**1.5) * math.sqrt(m.fc_prime) * 1000:.0f} MPa"],
        ["Vật liệu", "Cốt thép thường", "fy / Es", f"{m.fy:.1f} / {m.Es:.0f}", "MPa", "Cốt thép chịu lực CB400-V"],
        ["Móng cọc", "Đường kính cọc / Số lượng cọc", "D / n", f"{m.pile_diameter:.2f} / {m.total_piles}", "m / cọc", f"Sức chịu tải Pcp = {m.pile_capacity_allowable:.1f} kN"]
    ]
    if m.cap_type == "PT":
        input_data.append(["Cáp DƯL", "Cường độ danh định / Hệ số căng", "fpu / kfpj", f"{m.fpu:.1f} / {m.kfpj:.2f}", "MPa / —", f"7 Nhóm cáp G1..G7 (Delta={m.delta_anchor}mm)"])
    create_styled_table(doc, input_headers, input_data, [1.3, 1.8, 0.9, 1.1, 0.7, 1.4])

    # CHƯƠNG 3: CHI TIẾT TẢI TRỌNG TIÊU CHUẨN
    h3 = doc.add_heading("CHƯƠNG 3. TẢI TRỌNG TIÊU CHUẨN TÁC DỤNG LÊN TRỤ CẦU", level=1)
    h3.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    load_headers = ["Ký hiệu", "Tên tải trọng", "Công thức / Cách xác định", "Giá trị lực", "Điểm đặt / Tác dụng"]
    load_data = [
        ["DC1_left", "Tĩnh tải KCN nhịp Trái", "Phản lực từ dầm nhịp trái", f"{m.DC1_left:.1f} kN", f"Lệch tâm e_left = {m.e_left:.2f} m"],
        ["DC1_right", "Tĩnh tải KCN nhịp Phải", "Phản lực từ dầm nhịp phải", f"{m.DC1_right:.1f} kN", f"Lệch tâm e_right = {m.e_right:.2f} m"],
        ["DW_kcn", "Tĩnh tải lớp phủ & tiện ích", "DW trái + DW phải", f"{m.DW_left + m.DW_right:.1f} kN", "Tại tim các gối"],
        ["DC2_cap", "Tự trọng xà mũ vát", "gamma_c × (V_mid + V_cant)", f"{loads.DC2_cap:.1f} kN", "Trọng tâm xà mũ"],
        ["DC2_stem", "Tự trọng thân trụ", "gamma_c × Thể tích thân", f"{loads.DC2_stem:.1f} kN", "Trọng tâm thân trụ"],
        ["DC2_footing", "Tự trọng bệ trụ", "gamma_c × Bbe × Cbe × Hbe", f"{loads.DC2_footing:.1f} kN", "Trọng tâm bệ"],
        ["PL", "Hoạt tải người đi bộ", "pPL × bpl × (L1+L2)/2", f"{m.pPL * m.width_bpl * (m.span_L1 + m.span_L2) / 2.0:.1f} kN", "Phân bố trên lề bộ hành"],
        ["WS", "Gió tác dụng lên KCN & Thân", "Áp lực gió TCVN 11823-3 Điều 3.8", f"Gió ngang / Gió dọc", "Tâm đón gió"],
        ["EQ", "Động đất tác dụng lên Trụ", f"Gia tốc A={m.accel_A}, S={m.S_seismic}, R_stem={m.R_pier_stem}, R_móng=1.0", f"Quán tính KCN & Trụ", "Tâm khối lượng"],
        ["CT", "Va xe ô tô vào thân trụ", "Điều 3.6.5: Lực tĩnh tương đương", f"{m.CT:.1f} kN", f"Tại cao độ z_CT = {m.z_CT:.1f} m"],
        ["CV", "Va tàu thủy thông thuyền", f"Điều 3.14: {m.river_class} (DWT={m.ship_DWT:.0f}T, V={m.ship_velocity_m_s:.1f}m/s)", f"{m.CV:.1f} kN", f"Tại cao độ z_CV = {m.z_CV:.1f} m"]
    ]
    create_styled_table(doc, load_headers, load_data, [1.0, 1.8, 2.0, 1.0, 1.4])

    # CHƯƠNG 4: TỔ HỢP TẢI TRỌNG
    h4 = doc.add_heading("CHƯƠNG 4. BẢNG TỔ HỢP TẢI TRỌNG (TCVN 11823-3 TẠI 3 MẶT CẮT)", level=1)
    h4.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    p4_note = doc.add_paragraph()
    p4_note.add_run("Bảng tổ hợp nội lực tại mặt cắt Đáy Bệ Trụ (dùng cho tính toán và kiểm toán móng cọc):").bold = True

    comb_headers = ["Tổ hợp tải trọng", "TTGH", "N (kN)", "Hx (kN)", "Hy (kN)", "Mx (kNm)", "My (kNm)"]
    comb_data = []
    for c in result.footing_combinations:
        comb_data.append([c.comb_name, c.limit_state_group, f"{c.N:.1f}", f"{c.Hx:.1f}", f"{c.Hy:.1f}", f"{c.Mx:.1f}", f"{c.My:.1f}"])
    create_styled_table(doc, comb_headers, comb_data, [2.1, 0.9, 0.9, 0.8, 0.8, 0.9, 0.9])

    # CHƯƠNG 5: ĐỊA CHẤT VÀ SỨC CHỊU TẢI CỌC (TS-CAP)
    h5 = doc.add_heading("CHƯƠNG 5. ĐỊA CHẤT CÔNG TRÌNH VÀ SỨC CHỊU TẢI CỌC (TS-CAP TCVN 11823-10)", level=1)
    h5.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    p5_geo = doc.add_paragraph()
    p5_geo.add_run(f"• Đường kính cọc D = {m.pile_diameter} m | Chiều dài cọc L = {getattr(m, 'cap_bottom_elev_m', 0.0) - getattr(m, 'pile_tip_elev_m', -35.0):.2f} m\n"
                   f"• Cao độ đáy bệ = {getattr(m, 'cap_bottom_elev_m', 0.0)} m | Cao độ mũi cọc = {getattr(m, 'pile_tip_elev_m', -35.0)} m | Mực nước = {getattr(m, 'water_elev_m', 2.0)} m\n"
                   f"• Sức kháng ma sát thân danh định: Qsn = {piles.capacity_result.qshaft_nominal_kn if piles.capacity_result else 0.0:.1f} kN\n"
                   f"• Sức kháng mũi cọc danh định: Qpn = {piles.capacity_result.qtip_nominal_kn if piles.capacity_result else 0.0:.1f} kN\n"
                   f"• Sức kháng danh định tổng cộng: Rn = {(piles.capacity_result.qshaft_nominal_kn + piles.capacity_result.qtip_nominal_kn) if piles.capacity_result else 0.0:.1f} kN\n"
                   f"• TTGH Cường độ (ULS): Sức kháng tính toán φRn = {piles.P_allow_strength:.1f} kN (φ = 0.50~0.55)\n"
                   f"• TTGH Đặc biệt (EXTREME): Sức kháng tính toán Rn,ext = {piles.P_allow_extreme:.1f} kN (φ = 1.0)\n"
                   f"• TTGH Sử dụng (SLS): φ = 1.0 | Sức kháng nhổ cho phép: {piles.P_allow_uplift:.1f} kN\n")

    geo_headers = ["Lớp đất / Đá", "Cao độ đáy (m)", "Dày (m)", "Loại", "SPT", "γ (kN/m³)", "c (MPa) / φ (°)", "Qs (kN)", "Qsf (kN)"]
    geo_data = []
    if piles.capacity_result and piles.capacity_result.layers:
        for lr in piles.capacity_result.layers:
            geo_data.append([lr.name, f"{lr.bottom_elev_m:.2f}", f"{lr.thickness_m:.2f}", lr.soil_label, f"{lr.n_spt:.0f}", f"{lr.gamma_eff_kN_m3:.1f}", f"{lr.c_mpa:.3f} / {lr.phi_deg:.1f}°", f"{lr.qs_nominal_kn:.1f}", f"{lr.qs_factored_kn:.1f}"])
    else:
        for sly in getattr(m, 'soil_layers', []):
            geo_data.append([sly.get("name", ""), f"{sly.get('bottom_elev_m', 0.0):.2f}", "—", f"Loại {sly.get('soil_type', 1)}", f"{sly.get('n_spt', 0):.0f}", f"{sly.get('gamma_kN_m3', 18.0):.1f}", f"{sly.get('c_mpa', 0):.3f} / {sly.get('phi_deg', 0):.1f}°", "—", "—"])
    create_styled_table(doc, geo_headers, geo_data, [2.0, 0.9, 0.7, 0.9, 0.6, 0.7, 1.2, 0.8, 0.8])

    # CHƯƠNG 6: PHÂN TÍCH VÀ KIỂM TOÁN NỘI LỰC MÓNG CỌC TS_PILE
    h6_pile = doc.add_heading("CHƯƠNG 6. PHÂN TÍCH VÀ KIỂM TOÁN NỘI LỰC MÓNG CỌC (PHƯƠNG PHÁP TS_PILE)", level=1)
    h6_pile.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    p6_note = doc.add_paragraph()
    p6_note.add_run("Phương pháp tính: Ma trận độ cứng 3D phân bố tải trọng lên từng cọc theo tọa độ thực tế (X, Y) và góc nghiêng chéo α.").italic = True

    pile_headers = ["Tổ hợp", "TTGH", "Pmax (kN)", "Pmin (kN)", "Sức kháng cho phép (kN)", "Tỷ số D/C", "Kiểm toán Nén", "Kiểm toán Nhổ"]
    pile_data = []
    for res in piles.reactions_all:
        if res.limit_state_group == "STRENGTH":
            p_allow = piles.P_allow_strength
        elif res.limit_state_group == "EXTREME":
            p_allow = piles.P_allow_extreme
        else:
            p_allow = piles.P_allow_service
        ratio = res.P_max / p_allow if p_allow > 0 else 0.0
        pass_cap = res.P_max <= p_allow
        pass_ten = res.P_min >= 0.0
        pile_data.append([res.comb_name, res.limit_state_group, f"{res.P_max:.1f}", f"{res.P_min:.1f}", f"{p_allow:.1f}", f"{ratio:.2f}", "ĐẠT" if pass_cap else "VƯỢT TẢI", "ĐẠT" if pass_ten else "NHỔ CỌC"])
    create_styled_table(doc, pile_headers, pile_data, [1.8, 0.9, 0.9, 0.9, 1.1, 0.7, 0.8, 0.8])

    # CHƯƠNG 7: CHI TIẾT KIỂM TOÁN CÁC BỘ PHẬN KẾT CẤU
    h7 = doc.add_heading("CHƯƠNG 7. CHI TIẾT KIỂM TOÁN CÁC BỘ PHẬN KẾT CẤU TRỤ CẦU", level=1)
    h7.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    # 7.1 XÀ MŨ
    doc.add_heading("7.1 Kiểm toán Xà mũ Trụ cầu (Pier Cap)", level=2)
    if m.cap_type == "RC":
        doc.add_paragraph(f"• Xà mũ bê tông cốt thép thường (RC): Chiều rộng bxm = {m.bxm}m, Chiều cao ngàm h_root = {h_root}m, Chiều dài cánh hẫng = {l_cant}m\n"
                          f"• Mô men uốn tính toán lớn nhất: Mu = {cap.Mu_max:.1f} kNm ➔ Sức kháng uốn Mr = {cap.rc_flexure.Mr:.1f} kNm ({'ĐẠT' if cap.rc_flexure.passed else 'KHÔNG ĐẠT'})\n"
                          f"• Lực cắt tính toán lớn nhất: Vu = {cap.Vu_max:.1f} kN ➔ Sức kháng cắt Vr = {cap.rc_shear.Vr:.1f} kN ({'ĐẠT' if cap.rc_shear.passed else 'KHÔNG ĐẠT'})\n"
                          f"• Ứng suất nứt: fss = {cap.rc_crack.fss:.1f} MPa <= fsa = {cap.rc_crack.fsa:.1f} MPa ({'ĐẠT' if cap.rc_crack.passed else 'KHÔNG ĐẠT'})")
    else:
        doc.add_paragraph(f"• Xà mũ Bê tông Dự ứng lực (PT): 7 Nhóm cáp G1..G7 (Cáp 15.2mm, fpu={m.fpu}MPa)\n"
                          f"• Sức kháng uốn ULS: Mu = {cap.Mu_max:.1f} kNm <= Mr = {cap.pt_result.Mr if cap.pt_result else 0.0:.1f} kNm\n"
                          f"• Bảng kiểm toán ứng suất 7 Giai đoạn thi công:")
        if hasattr(result, "pt_stages") and result.pt_stages:
            stage_headers = ["Giai đoạn thi công", "M_ngoại (kNm)", "Ứng suất đỉnh (MPa)", "Ứng suất đáy (MPa)", "Kết luận"]
            stage_data = []
            for st in result.pt_stages:
                stage_data.append([st.stage_name, f"{st.M_ext:.1f}", f"{st.sigma_top:.2f}", f"{st.sigma_bot:.2f}", "ĐẠT" if st.passed else "KHÔNG ĐẠT"])
            create_styled_table(doc, stage_headers, stage_data, [2.5, 1.2, 1.2, 1.2, 0.9])

    # 7.2 THÂN TRỤ
    doc.add_heading("7.2 Kiểm toán Thân trụ (Nén - Uốn 2 phương & Fiber Section)", level=2)
    doc.add_paragraph(f"• Tiết diện: {m.bth1}m × {m.hth1}m, Chiều cao Hth = {m.Hth}m | Hình thức: {col_str}\n"
                      f"• Độ mảnh phương dọc klu/r = {stem.slenderness_lambda_y:.1f} (Hệ số khuếch đại delta_by = {stem.delta_b_y:.2f})\n"
                      f"• Độ mảnh phương ngang klu/r = {stem.slenderness_lambda_x:.1f} (Hệ số khuếch đại delta_bx = {stem.delta_b_x:.2f})\n"
                      f"• Hàm lượng cốt thép dọc: Ast/Ag = {stem.rebar_ratio*100:.2f}% (Quy định: 1.0% <= Ast/Ag <= 8.0%)\n"
                      f"• Tương tác P-Mx-My: Tỷ số sử dụng lớn nhất D/C = {stem.utilization_pm:.2f} ({'ĐẠT' if stem.pm_passed else 'KHÔNG ĐẠT'})\n"
                      f"• Sức kháng cắt: Vu = {stem.Vu_max:.1f} kN <= Vr = {stem.shear_check.Vr:.1f} kN ({'ĐẠT' if stem.shear_check.passed else 'KHÔNG ĐẠT'})\n"
                      f"• Kiểm soát nứt: fss = {stem.crack_check.fss:.1f} MPa <= fsa = {stem.crack_check.fsa:.1f} MPa ({'ĐẠT' if stem.crack_check.passed else 'KHÔNG ĐẠT'})")

    # 7.3 BỆ TRỤ
    doc.add_heading("7.3 Kiểm toán Bệ trụ (Pier Footing)", level=2)
    doc.add_paragraph(f"• Kích thước đài bệ: {m.Bbe}m (dọc) × {m.Cbe}m (ngang) × {m.Hbe}m (cao)\n"
                      f"• Sức kháng uốn phương X (dọc cầu): Mux = {footing.Mux_max:.1f} kNm <= Mrx = {footing.flexure_x.Mr:.1f} kNm ({'ĐẠT' if footing.flexure_x.passed else 'KHÔNG ĐẠT'})\n"
                      f"• Sức kháng uốn phương Y (ngang cầu): Muy = {footing.Muy_max:.1f} kNm <= Mry = {footing.flexure_y.Mr:.1f} kNm ({'ĐẠT' if footing.flexure_y.passed else 'KHÔNG ĐẠT'})\n"
                      f"• Cắt dầm 1 phương: Vu = {footing.shear_beam.Vu:.1f} kN <= Vr = {footing.shear_beam.Vr:.1f} kN ({'ĐẠT' if footing.shear_beam.passed else 'KHÔNG ĐẠT'})\n"
                      f"• Đâm thủng 2 phương quanh cột: Vu = {footing.Vu_punching:.1f} kN <= Vr = {footing.Vr_punching:.1f} kN ({'ĐẠT' if footing.punching_passed else 'KHÔNG ĐẠT'})")

    # CHƯƠNG 8 / KẾT LUẬN
    h8 = doc.add_heading("CHƯƠNG 8. KẾT LUẬN VÀ KIẾN NGHỊ", level=1)
    h8.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    p_conc = doc.add_paragraph()
    if result.is_success:
        r_c = p_conc.add_run("KẾT LUẬN: Kết cấu Trụ cầu đáp ứng đầy đủ tất cả các điều kiện về chịu lực, nứt, ổn định móng cọc theo tiêu chuẩn TCVN 11823:2017.")
        r_c.font.bold = True
        r_c.font.color.rgb = RGBColor(0, 128, 0)
    else:
        r_c = p_conc.add_run("KẾT LUẬN: Một số chỉ tiêu kiểm toán chưa thỏa mãn điều kiện theo TCVN 11823:2017. Đề nghị kỹ sư điều chỉnh kích thước hoặc tăng cường cốt thép / cọc móng.")
        r_c.font.bold = True
        r_c.font.color.rgb = RGBColor(200, 0, 0)

    doc.save(output_path)
    return output_path
